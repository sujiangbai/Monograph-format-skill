#!/usr/bin/env python3
"""Create a structural and editable-object inventory for a monograph DOCX."""

from __future__ import annotations

import argparse
import collections
import json
import sys
from pathlib import Path

from docx.oxml.ns import qn

from _common import (
    FormatMonographError,
    content_fingerprint,
    equation_inventory,
    field_cache_inventory,
    field_inventory,
    load_document,
    protected_object_manifest,
    semantic_title_heading_role,
    style_effective_font,
    theme_font_inventory,
    word_xml_counts,
    write_json,
)
from structure_map import candidate_structure_map
from docx_pagination import pagination_inventory


def length_points(value) -> float | None:
    return None if value is None else round(value.pt, 3)


def length_mm(value) -> float | None:
    return None if value is None else round(value.mm, 3)


def style_definition(document, style) -> dict:
    r_pr = style.element.rPr
    r_fonts = None if r_pr is None else r_pr.rFonts
    attrs = {}
    if r_fonts is not None:
        for name in (
            "ascii",
            "hAnsi",
            "eastAsia",
            "cs",
            "asciiTheme",
            "hAnsiTheme",
            "eastAsiaTheme",
            "cstheme",
        ):
            attrs[name] = r_fonts.get(qn(f"w:{name}"))
    spacing = style.paragraph_format.line_spacing
    return {
        "font_name": style.font.name,
        "font_name_ascii": attrs.get("ascii") or attrs.get("hAnsi"),
        "font_name_east_asia": attrs.get("eastAsia"),
        "font_name_complex_script": attrs.get("cs"),
        "font_theme_references": {
            name: attrs.get(name)
            for name in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme")
            if attrs.get(name)
        },
        "effective_fonts": {
            name: {
                "name": style_effective_font(document, style, attribute)[0],
                "source": style_effective_font(document, style, attribute)[1],
            }
            for name, attribute in (
                ("ascii", "ascii"),
                ("east_asia", "eastAsia"),
                ("complex_script", "cs"),
            )
        },
        "font_size_pt": length_points(style.font.size),
        "bold": style.font.bold,
        "line_spacing": (
            length_points(spacing)
            if spacing is not None and hasattr(spacing, "pt")
            else (None if spacing is None else float(spacing))
        ),
        "line_spacing_rule": (
            None
            if style.paragraph_format.line_spacing_rule is None
            else str(style.paragraph_format.line_spacing_rule)
        ),
    }


def row_has_property(row, name: str) -> bool:
    tr_pr = row._tr.trPr
    return tr_pr is not None and tr_pr.find(qn(f"w:{name}")) is not None


def inventory(path: Path) -> dict:
    document = load_document(path)
    style_counts = collections.Counter(
        paragraph.style.name if paragraph.style else "<none>"
        for paragraph in document.paragraphs
    )
    headings = [
        {
            "index": index,
            "style": paragraph.style.name,
            "semantic_role": semantic_title_heading_role(paragraph.style),
            "text": paragraph.text[:160],
        }
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.style
        and (semantic_title_heading_role(paragraph.style) or "").startswith(
            "heading_"
        )
    ]
    outline_paragraphs = [
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph._p.pPr is not None
        and paragraph._p.pPr.find(qn("w:outlineLvl")) is not None
    ]
    sections = []
    for index, section in enumerate(document.sections):
        sections.append(
            {
                "index": index,
                "orientation": str(section.orientation),
                "page_width_mm": length_mm(section.page_width),
                "page_height_mm": length_mm(section.page_height),
                "margin_top_mm": length_mm(section.top_margin),
                "margin_bottom_mm": length_mm(section.bottom_margin),
                "margin_left_mm": length_mm(section.left_margin),
                "margin_right_mm": length_mm(section.right_margin),
                "gutter_mm": length_mm(section.gutter),
                "header_distance_mm": length_mm(section.header_distance),
                "footer_distance_mm": length_mm(section.footer_distance),
                "different_first_page_header_footer": bool(
                    section.different_first_page_header_footer
                ),
            }
        )

    tables = []
    for index, table in enumerate(document.tables):
        tables.append(
            {
                "index": index,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "style": table.style.name if table.style else None,
                "repeat_header_row": bool(
                    table.rows and row_has_property(table.rows[0], "tblHeader")
                ),
                "repeat_header_rows": [
                    row_index
                    for row_index, row in enumerate(table.rows)
                    if row_has_property(row, "tblHeader")
                ],
                "prevent_split_rows": [
                    row_index
                    for row_index, row in enumerate(table.rows)
                    if row_has_property(row, "cantSplit")
                ],
                "complex_merge": len(
                    {cell._tc for row in table.rows for cell in row.cells}
                )
                < sum(len(row.cells) for row in table.rows),
                "floating_object_count": len(
                    table._tbl.xpath(".//wp:anchor | .//w:object | .//w:pict")
                ),
                "visible_control_mark_candidates": sum(
                    1
                    for character in "\n".join(
                        cell.text for row in table.rows for cell in row.cells
                    )
                    if 0x2400 <= ord(character) <= 0x2426
                    or (ord(character) < 32 and character not in "\t\n\r")
                ),
            }
        )

    settings = document.settings.element
    object_manifest = protected_object_manifest(path)
    return {
        "file": str(path.resolve()),
        "content_fingerprint_sha256": content_fingerprint(path),
        "paragraph_count": len(document.paragraphs),
        "nonempty_paragraph_count": sum(bool(p.text.strip()) for p in document.paragraphs),
        "style_counts": dict(sorted(style_counts.items())),
        "style_definitions": {
            name: style_definition(document, document.styles[name])
            for name in (
                "Normal",
                "Heading 1",
                "Heading 2",
                "Heading 3",
                "Heading 4",
                "TOC 1",
                "TOC 2",
                "TOC 3",
                "Caption",
                "Quote",
                "Bibliography",
            )
            if name in document.styles
        },
        "theme_fonts": theme_font_inventory(document),
        "headings": headings,
        "direct_outline_level_paragraphs": outline_paragraphs,
        "section_count": len(document.sections),
        "sections": sections,
        "table_count": len(document.tables),
        "tables": tables,
        "inline_shape_count": len(document.inline_shapes),
        "settings": {
            "odd_and_even_pages_header_footer": bool(
                document.settings.odd_and_even_pages_header_footer
            ),
            "mirror_margins": settings.find(qn("w:mirrorMargins")) is not None,
            "update_fields_on_open": settings.find(qn("w:updateFields")) is not None,
        },
        "fields": field_inventory(path),
        "field_cache": field_cache_inventory(path),
        "pagination": pagination_inventory(path),
        "equations": equation_inventory(path),
        "protected_objects": {
            "embedding_count": len(object_manifest["embeddings"]),
            "media_count": len(object_manifest["media"]),
            "omml_hash_count": len(object_manifest["omml_sha256"]),
            "manifest": object_manifest,
        },
        "ooxml": word_xml_counts(path),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument(
        "--structure-map-output",
        type=Path,
        help="Write a text-free candidate structure map for caller QA.",
    )
    args = parser.parse_args()
    try:
        result = inventory(args.input)
        write_json(args.output, result)
        if args.structure_map_output:
            write_json(args.structure_map_output, candidate_structure_map(args.input))
    except FormatMonographError as exc:
        print(str(exc), file=sys.stderr)
        return 1
    print(
        json.dumps(
            {
                "output": str(args.output),
                "structure_map_output": (
                    str(args.structure_map_output) if args.structure_map_output else None
                ),
                "fingerprint": result["content_fingerprint_sha256"],
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
