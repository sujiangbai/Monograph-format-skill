"""Apply and inspect approved DOCX pagination sections."""

from __future__ import annotations

import copy
import posixpath
import re
import zipfile
from pathlib import Path
from typing import Any, Callable

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from _common import (
    NS,
    STRUCTURAL_INDENT_ATTRIBUTES,
    FormatMonographError,
    font_alias_keys,
    semantic_title_heading_role,
    style_effective_font,
)


FOOTER_PART = re.compile(r"word/footer\d+\.xml$")
HEADER_FOOTER_PART = re.compile(r"word/(?:header|footer)\d+\.xml$")


def _on_off_enabled(element: Any | None) -> bool:
    if element is None:
        return False
    value = element.get(qn("w:val"))
    if value is None:
        return True
    normalized = value.strip().lower()
    if normalized in {"0", "false", "off", "no"}:
        return False
    return normalized in {"1", "true", "on", "yes"}


def _section_properties(document: Any) -> list[Any]:
    return document.element.body.xpath("./w:p/w:pPr/w:sectPr | ./w:sectPr")


def _body_position(document: Any, paragraph: Any) -> int:
    children = list(document.element.body)
    try:
        return children.index(paragraph._p)
    except ValueError as exc:
        raise FormatMonographError(
            "Pagination locators must resolve to top-level body paragraphs."
        ) from exc


def section_index_for_paragraph(document: Any, paragraph: Any) -> int:
    position = _body_position(document, paragraph)
    section_index = 0
    for child in list(document.element.body)[:position]:
        p_pr = child.find(qn("w:pPr")) if child.tag == qn("w:p") else None
        if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
            section_index += 1
    return section_index


def _nonzero_indent_attributes(ind: Any | None) -> dict[str, str]:
    if ind is None:
        return {}
    result = {}
    for attribute in STRUCTURAL_INDENT_ATTRIBUTES:
        value = ind.get(qn(f"w:{attribute}"))
        if value is not None and value != "0":
            result[attribute] = value
    return result


def _effective_style_indent_attributes(style: Any | None) -> dict[str, str | None]:
    values: dict[str, str | None] = {
        attribute: None for attribute in STRUCTURAL_INDENT_ATTRIBUTES
    }
    unresolved = set(values)
    current = style
    visited: set[str] = set()
    while current is not None and current.style_id not in visited and unresolved:
        visited.add(current.style_id)
        p_pr = current.element.pPr
        ind = None if p_pr is None else p_pr.find(qn("w:ind"))
        if ind is not None:
            for attribute in tuple(unresolved):
                value = ind.get(qn(f"w:{attribute}"))
                if value is not None:
                    values[attribute] = value
                    unresolved.remove(attribute)
        current = current.base_style
    return values


def _set_page_numbering(sect_pr: Any, start: int | None, number_format: str) -> None:
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None and start is None:
        return
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), number_format)
    if start is None:
        pg_num.attrib.pop(qn("w:start"), None)
    else:
        pg_num.set(qn("w:start"), str(start))


def _field_instructions(root: Any) -> list[str]:
    def xpath(expression: str) -> list[Any]:
        try:
            return root.xpath(expression, namespaces=NS)
        except TypeError:
            return root.xpath(expression)

    result = [
        value.strip()
        for value in xpath(".//w:fldSimple/@w:instr")
        if value.strip()
    ]
    stack: list[dict[str, Any]] = []
    for element in root.iter():
        if element.tag == qn("w:fldChar"):
            kind = element.get(qn("w:fldCharType"))
            if kind == "begin":
                stack.append({"collecting": True, "parts": []})
            elif kind == "separate" and stack:
                field = stack[-1]
                instruction = "".join(field["parts"]).strip()
                if instruction:
                    result.append(instruction)
                field["collecting"] = False
            elif kind == "end" and stack:
                field = stack.pop()
                if field["collecting"]:
                    instruction = "".join(field["parts"]).strip()
                    if instruction:
                        result.append(instruction)
        elif (
            element.tag == qn("w:instrText")
            and stack
            and stack[-1]["collecting"]
        ):
            stack[-1]["parts"].append(element.text or "")
    return result


def _field_types(root: Any) -> set[str]:
    return {
        instruction.split(maxsplit=1)[0].upper()
        for instruction in _field_instructions(root)
        if instruction.strip()
    }


def _page_field_count(root: Any) -> int:
    def xpath(expression: str) -> list[Any]:
        try:
            return root.xpath(expression, namespaces=NS)
        except TypeError:
            return root.xpath(expression)

    simple = xpath(
        ".//w:fldSimple[starts-with(translate(normalize-space(@w:instr), 'page', 'PAGE'), 'PAGE')]"
    )
    complex_instructions = [
        value
        for value in xpath(".//w:instrText/text()")
        if value.strip().upper().startswith("PAGE")
    ]
    return len(simple) + len(complex_instructions)


def _page_only_footer(footer: Any) -> bool:
    root = footer._element
    if root.xpath(".//w:tbl | .//w:object | .//*[local-name()='blip']"):
        return False
    instructions = _field_instructions(root)
    field_types = _field_types(root)
    page_count = _page_field_count(root)
    if page_count == 0:
        return not field_types and not root.xpath(
            ".//w:drawing | .//w:pict | .//w:t[normalize-space(.) != '']"
        )
    if field_types - {"PAGE", "="}:
        return False
    formulas = [
        re.sub(r"\s+", "", instruction).upper()
        for instruction in instructions
        if instruction.lstrip().startswith("=")
    ]
    if formulas and (formulas != ["=-1"] or page_count != 1):
        return False
    for visual in root.xpath(".//w:drawing | .//w:pict"):
        if (
            not visual.xpath(".//*[local-name()='txbxContent']")
            or not _page_only_payload(visual)
        ):
            return False
        if visual.xpath(".//*[local-name()='blip']"):
            return False
        for image_data in visual.xpath(".//*[local-name()='imagedata']"):
            if any(str(value).strip() for value in image_data.attrib.values()):
                return False
    visible = [
        value.strip() for value in root.xpath(".//w:t/text()") if value.strip()
    ]
    return all(
        value.isdigit() or re.fullmatch(r"[IVXLCDMivxlcdm]+", value)
        for value in visible
    )


def _page_only_payload(root: Any) -> bool:
    field_types = _field_types(root)
    page_count = _page_field_count(root)
    if field_types - {"PAGE", "="} or page_count < 1:
        return False
    formulas = [
        re.sub(r"\s+", "", instruction).upper()
        for instruction in _field_instructions(root)
        if instruction.lstrip().startswith("=")
    ]
    return not formulas or (formulas == ["=-1"] and page_count == 1)


_PAGINATION_REWRITE_TAGS = {
    qn("w:ftr"), qn("w:p"), qn("w:pPr"), qn("w:pStyle"), qn("w:jc"),
    qn("w:r"), qn("w:rPr"), qn("w:rFonts"), qn("w:sz"), qn("w:szCs"),
    qn("w:fldSimple"), qn("w:fldChar"), qn("w:instrText"), qn("w:t"),
}


def _pagination_rewrite_shape(root: Any) -> str:
    """Classify the only footer shapes this batch is authorized to rewrite."""
    def xpath(expression: str) -> list[Any]:
        try:
            return root.xpath(expression, namespaces=NS)
        except TypeError:
            return root.xpath(expression)

    if any(element.tag not in _PAGINATION_REWRITE_TAGS for element in root.iter()):
        return "blocked"
    simple_fields = xpath(".//w:fldSimple")
    instructions = [
        re.sub(r"\s+", " ", field.get(qn("w:instr"), "")).strip().upper()
        for field in simple_fields
    ]
    stack: list[dict[str, Any]] = []
    complex_results: list[Any] = []
    for element in root.iter():
        if element.tag == qn("w:fldChar"):
            kind = (element.get(qn("w:fldCharType")) or "").lower()
            if kind == "begin":
                if stack:
                    return "blocked"
                stack.append({"parts": [], "separated": False})
            elif kind == "separate":
                if len(stack) != 1 or stack[-1]["separated"]:
                    return "blocked"
                stack[-1]["separated"] = True
            elif kind == "end":
                if len(stack) != 1 or not stack[-1]["separated"]:
                    return "blocked"
                instructions.append(
                    re.sub(r"\s+", " ", "".join(stack[-1]["parts"])).strip().upper()
                )
                stack.pop()
            else:
                return "blocked"
        elif element.tag == qn("w:instrText"):
            if len(stack) != 1 or stack[-1]["separated"]:
                return "blocked"
            stack[-1]["parts"].append(element.text or "")
        elif element.tag == qn("w:t") and (element.text or "").strip() and stack:
            if not stack[-1]["separated"]:
                return "blocked"
            complex_results.append(element)
    if stack or any(instruction != "PAGE" for instruction in instructions):
        return "blocked"
    visible = [element for element in xpath(".//w:t") if (element.text or "").strip()]
    if not instructions:
        return "blank" if not visible else "blocked"
    if len(instructions) != 1:
        return "blocked"
    if not visible or any(not (element.text or "").strip().isdigit() for element in visible):
        return "blocked"
    for element in visible:
        try:
            ancestors = element.xpath("ancestor::w:fldSimple", namespaces=NS)
        except TypeError:
            ancestors = element.xpath("ancestor::w:fldSimple")
        if (
            len(ancestors) != 1 or ancestors[0] not in simple_fields
        ) and element not in complex_results:
            return "blocked"
    return "page"


def _pagination_footer_shape(footer: Any) -> str:
    if footer.part.rels:
        return "blocked"
    return _pagination_rewrite_shape(footer._element)


def _replace_with_page_field(footer: Any, alignment: Any) -> None:
    paragraphs = list(footer.paragraphs)
    paragraph = paragraphs[0] if paragraphs else footer.add_paragraph()
    for extra in paragraphs[1:]:
        footer._element.remove(extra._p)
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.alignment = alignment
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "18")
    run_properties.extend((fonts, size, size_cs))
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _ensure_page_field(
    footer: Any, alignment: Any, *, replace_static_page_text: bool = False
) -> bool:
    del replace_static_page_text
    if _pagination_footer_shape(footer) == "blocked":
        raise FormatMonographError(
            "Footer is not blank or a relationship-free pure PAGE footer."
        )
    before = etree.tostring(footer._element)
    _replace_with_page_field(footer, alignment)
    return etree.tostring(footer._element) != before


def _existing_boundary_before(document: Any, paragraph: Any, label: str) -> int:
    position = _body_position(document, paragraph)
    if position == 0:
        return 0
    previous = list(document.element.body)[position - 1]
    p_pr = previous.find(qn("w:pPr")) if previous.tag == qn("w:p") else None
    if p_pr is None or p_pr.find(qn("w:sectPr")) is None:
        raise FormatMonographError(
            f"Approved {label} must begin at an existing, exact section boundary."
        )
    return section_index_for_paragraph(document, paragraph)


def _explicit_references(sect_pr: Any, kind: str) -> dict[str, str]:
    return {
        reference.get(qn("w:type"), "default"): reference.get(qn("r:id"))
        for reference in sect_pr.findall(qn(f"w:{kind}Reference"))
        if reference.get(qn("r:id"))
    }


def _effective_footer_references(document: Any) -> list[dict[str, str]]:
    inherited: dict[str, str] = {}
    result = []
    for sect_pr in _section_properties(document):
        inherited.update(_explicit_references(sect_pr, "footer"))
        result.append(dict(inherited))
    return result


def _footer_part(document: Any, rel_id: str, *, section: int, kind: str) -> Any:
    part = document.part.related_parts.get(rel_id)
    if part is None or not hasattr(part, "element"):
        raise FormatMonographError(
            f"Section {section} {kind} footer relationship cannot be resolved."
        )
    return part


def _footer_part_identity(part: Any) -> str:
    return str(part.partname)


def _validate_footer_part(document: Any, rel_id: str, *, section: int, kind: str) -> Any:
    part = _footer_part(document, rel_id, section=section, kind=kind)
    if part.rels or _pagination_rewrite_shape(part.element) == "blocked":
        raise FormatMonographError(
            f"Section {section} {kind} footer is not blank or a relationship-free pure PAGE footer."
        )
    return part


def preflight_pagination_sections(
    document: Any,
    settings: dict[str, Any],
    resolver: Callable[[Any, dict[str, Any]], Any],
    *,
    front_matter: dict[str, Any] | None = None,
    approved_headings: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """Validate all pagination inputs without mutating the document."""
    if not settings.get("approved"):
        return {}
    if settings.get("number_format") != "decimal":
        raise FormatMonographError("Approved pagination number_format must be decimal.")
    if settings.get("start_at") != {"toc": 1, "body": 1}:
        raise FormatMonographError(
            "Approved pagination starts must be exactly toc=1 and body=1."
        )
    if settings.get("continue_after_body_start") is not True:
        raise FormatMonographError(
            "Approved pagination must continue after the body start."
        )
    toc_locator = settings.get("toc_start")
    body_locator = settings.get("body_start")
    if not isinstance(toc_locator, dict) or not isinstance(body_locator, dict):
        raise FormatMonographError(
            "Approved pagination_sections requires toc_start and body_start locators."
        )
    toc_paragraph = resolver(document, toc_locator)
    body_paragraph = resolver(document, body_locator)
    front_matter = front_matter or {}
    toc_anchor = getattr(document, "_format_monograph_toc_heading", None)
    toc_heading_locator = front_matter.get("toc_heading")
    if toc_anchor is None and front_matter.get("approved") and isinstance(
        toc_heading_locator, dict
    ):
        toc_anchor = resolver(document, toc_heading_locator)
    if toc_anchor is None:
        toc_anchor = toc_paragraph
    toc_index = _existing_boundary_before(document, toc_anchor, "TOC pagination start")
    body_index = _existing_boundary_before(document, body_paragraph, "body pagination start")
    if _body_position(document, toc_anchor) >= _body_position(document, body_paragraph):
        raise FormatMonographError("TOC pagination start must precede body pagination start.")
    if body_index <= toc_index:
        raise FormatMonographError("TOC and body must begin in distinct existing sections.")

    sections = _section_properties(document)
    explicit_front_starts = []
    for index, sect_pr in enumerate(sections[: toc_index + 1]):
        pg_num = sect_pr.find(qn("w:pgNumType"))
        if pg_num is not None and pg_num.get(qn("w:start")) is not None:
            explicit_front_starts.append((index, pg_num.get(qn("w:start"))))
    if explicit_front_starts:
        if explicit_front_starts[0][1] != "1":
            raise FormatMonographError(
                "Existing front-matter page-number starts are ambiguous or inconsistent."
            )
        front_index = explicit_front_starts[0][0]
    else:
        front_index = toc_index

    if (
        front_matter.get("approved")
        and front_matter.get("separate_title_page")
        and isinstance(front_matter.get("book_title"), dict)
    ):
        title = resolver(document, front_matter["book_title"])
        title_index = section_index_for_paragraph(document, title)
        if title_index >= front_index:
            raise FormatMonographError(
                "The approved separate title page must precede the first numbered front section."
            )

    settings_element = document.settings.element
    even_odd_enabled = _on_off_enabled(
        settings_element.find(qn("w:evenAndOddHeaders"))
    )
    if not even_odd_enabled and any(
        sect_pr.findall(qn("w:headerReference")) for sect_pr in sections
    ):
        raise FormatMonographError(
            "Enabling odd/even footers would change existing header behavior."
        )

    effective = _effective_footer_references(document)
    part_areas: dict[str, set[str]] = {}
    for index, references in enumerate(effective):
        area = "numbered" if index >= front_index else "unnumbered"
        for kind, rel_id in references.items():
            part = _footer_part(document, rel_id, section=index, kind=kind)
            part_areas.setdefault(_footer_part_identity(part), set()).add(area)
    for partname, areas in part_areas.items():
        if areas == {"numbered", "unnumbered"}:
            raise FormatMonographError(
                f"Footer part {partname} is shared across unnumbered and numbered sections."
            )
    if not even_odd_enabled:
        for index, references in enumerate(effective[:front_index]):
            for kind, rel_id in references.items():
                part = document.part.related_parts.get(rel_id)
                if (
                    part is None
                    or not hasattr(part, "element")
                    or part.rels
                    or _pagination_rewrite_shape(part.element) != "blank"
                ):
                    raise FormatMonographError(
                        f"Enabling odd/even footers would change section {index} "
                        f"{kind} footer behavior."
                    )
    if not explicit_front_starts:
        for index, references in enumerate(effective[:toc_index]):
            for kind, rel_id in references.items():
                part = document.part.related_parts.get(rel_id)
                if (
                    part is not None
                    and hasattr(part, "element")
                    and not part.rels
                    and _pagination_rewrite_shape(part.element) == "page"
                ):
                    raise FormatMonographError(
                        "Existing PAGE footer before the TOC has no unambiguous "
                        "front-matter numbering start."
                    )
    first_page_sections = {toc_index, body_index}
    for entry in approved_headings or []:
        if not entry.get("approved") or int(entry.get("level", 0)) != 1:
            continue
        locator = entry.get("locator")
        if not isinstance(locator, dict):
            continue
        paragraph = resolver(document, locator)
        try:
            index = _existing_boundary_before(
                document, paragraph, "approved chapter heading"
            )
        except FormatMonographError:
            continue
        if index >= body_index:
            first_page_sections.add(index)

    relationship_roles: dict[str, set[str]] = {}
    for index in range(front_index, len(sections)):
        references = effective[index]
        kinds = ["default", "even"]
        if _on_off_enabled(sections[index].find(qn("w:titlePg"))):
            kinds.append("first")
        for kind in kinds:
            rel_id = references.get(kind)
            if rel_id:
                part = _validate_footer_part(
                    document, rel_id, section=index, kind=kind
                )
                relationship_roles.setdefault(_footer_part_identity(part), set()).add(
                    "left" if kind == "even" else "right"
                )
        if (
            _on_off_enabled(sections[index].find(qn("w:titlePg")))
            and index not in first_page_sections
        ):
            rel_id = references.get("first")
            part = None if rel_id is None else document.part.related_parts.get(rel_id)
            if (
                part is None
                or not hasattr(part, "element")
                or _pagination_rewrite_shape(part.element) != "page"
                or _page_footer_format(part.element)
                != {
                    "alignment": "right",
                    "font_ascii": "Times New Roman",
                    "font_east_asia": "宋体",
                    "font_size_half_points": "18",
                    "font_size_cs_half_points": "18",
                }
            ):
                raise FormatMonographError(
                    f"Section {index} has titlePg but is not an approved numbered first-page target."
                )
    for partname, roles in relationship_roles.items():
        if roles == {"left", "right"}:
            raise FormatMonographError(
                f"Footer part {partname} is shared by conflicting alignment roles."
            )

    create_footer_definitions = [
        (front_index, kind)
        for kind in ("default", "even")
        if not effective[front_index].get(kind)
    ]
    first_available = bool(effective[front_index].get("first"))
    for index in range(front_index, len(sections)):
        if _explicit_references(sections[index], "footer").get("first"):
            first_available = True
        if (
            index in first_page_sections
            and _on_off_enabled(sections[index].find(qn("w:titlePg")))
            and not first_available
        ):
            create_footer_definitions.append((index, "first"))
            first_available = True

    return {
        "front_section": front_index,
        "toc_section": toc_index,
        "body_section": body_index,
        "section_count": len(sections),
        "first_page_sections": sorted(first_page_sections),
        "create_footer_definitions": create_footer_definitions,
    }


def _ensure_visible_footers(document: Any, plan: dict[str, Any]) -> int:
    changed = 0
    document.settings.odd_and_even_pages_header_footer = True
    sections = list(document.sections)
    start_index = int(plan["front_section"])
    for index, kind in plan.get("create_footer_definitions", []):
        footer = {
            "default": sections[index].footer,
            "even": sections[index].even_page_footer,
            "first": sections[index].first_page_footer,
        }[kind]
        if not footer._has_definition:
            footer._add_definition()
    for index, section in enumerate(sections[start_index:], start=start_index):
        changed += int(_ensure_page_field(section.footer, WD_ALIGN_PARAGRAPH.RIGHT))
        changed += int(_ensure_page_field(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT))
        if (
            _on_off_enabled(section._sectPr.find(qn("w:titlePg")))
            and index in plan.get("first_page_sections", [])
        ):
            first = section.first_page_footer
            changed += int(_ensure_page_field(first, WD_ALIGN_PARAGRAPH.RIGHT))
    return changed


def apply_pagination_sections(
    document: Any,
    settings: dict[str, Any],
    resolver: Callable[[Any, dict[str, Any]], Any],
    *,
    replace_static_page_text: bool = False,
    front_matter: dict[str, Any] | None = None,
    approved_headings: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    del replace_static_page_text
    plan = preflight_pagination_sections(
        document,
        settings,
        resolver,
        front_matter=front_matter,
        approved_headings=approved_headings,
    )
    if not plan:
        return {}
    sections = _section_properties(document)
    starts = settings.get("start_at", {})
    number_format = str(settings.get("number_format", "decimal"))
    front_index = int(plan["front_section"])
    toc_index = int(plan["toc_section"])
    body_index = int(plan["body_section"])
    _set_page_numbering(sections[front_index], int(starts.get("toc", 1)), number_format)
    for sect_pr in sections[front_index + 1 : body_index]:
        _set_page_numbering(sect_pr, None, number_format)
    _set_page_numbering(sections[body_index], int(starts.get("body", 1)), number_format)
    if settings.get("continue_after_body_start", True):
        for sect_pr in sections[body_index + 1 :]:
            _set_page_numbering(sect_pr, None, number_format)
    fields_changed = _ensure_visible_footers(document, plan)
    return {**plan, "page_fields_changed": fields_changed}


def finalize_pagination_sections(
    document: Any,
    settings: dict[str, Any],
    resolver: Callable[[Any, dict[str, Any]], Any],
    *,
    front_matter: dict[str, Any] | None = None,
    approved_headings: list[dict[str, Any]] | None = None,
) -> bool:
    preflight_pagination_sections(
        document,
        settings,
        resolver,
        front_matter=front_matter,
        approved_headings=approved_headings,
    )
    return False


def _relationship_targets(package: zipfile.ZipFile) -> dict[str, str]:
    name = "word/_rels/document.xml.rels"
    if name not in package.namelist():
        return {}
    root = etree.fromstring(package.read(name))
    result = {}
    for relationship in root:
        identifier = relationship.get("Id")
        target = relationship.get("Target", "")
        if identifier:
            result[identifier] = posixpath.normpath(posixpath.join("word", target))
    return result


def _page_footer_format(root: Any) -> dict[str, Any]:
    try:
        fields = root.xpath(".//w:fldSimple", namespaces=NS)
    except TypeError:
        fields = root.xpath(".//w:fldSimple")
    if len(fields) != 1:
        return {}
    paragraph = fields[0]
    while paragraph is not None and paragraph.tag != qn("w:p"):
        paragraph = paragraph.getparent()
    justification = None
    if paragraph is not None:
        p_pr = paragraph.find(qn("w:pPr"))
        jc = None if p_pr is None else p_pr.find(qn("w:jc"))
        justification = None if jc is None else jc.get(qn("w:val"))
    run = fields[0].find(qn("w:r"))
    run_properties = None if run is None else run.find(qn("w:rPr"))
    fonts = None if run_properties is None else run_properties.find(qn("w:rFonts"))
    size = None if run_properties is None else run_properties.find(qn("w:sz"))
    size_cs = None if run_properties is None else run_properties.find(qn("w:szCs"))
    return {
        "alignment": justification,
        "font_ascii": None if fonts is None else fonts.get(qn("w:ascii")),
        "font_east_asia": None if fonts is None else fonts.get(qn("w:eastAsia")),
        "font_size_half_points": None if size is None else size.get(qn("w:val")),
        "font_size_cs_half_points": (
            None if size_cs is None else size_cs.get(qn("w:val"))
        ),
    }


def pagination_inventory(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as package:
        document = etree.fromstring(package.read("word/document.xml"))
        settings = (
            etree.fromstring(package.read("word/settings.xml"))
            if "word/settings.xml" in package.namelist()
            else None
        )
        relationships = _relationship_targets(package)
        referenced_parts: set[str] = set()
        inherited: dict[str, str] = {}
        sections = []
        for index, sect_pr in enumerate(
            document.xpath("./w:body/w:p/w:pPr/w:sectPr | ./w:body/w:sectPr", namespaces=NS)
        ):
            explicit = {}
            for reference in sect_pr.xpath(
                "./w:footerReference", namespaces=NS
            ):
                kind = reference.get(qn("w:type"), "default")
                rel_id = reference.get(qn("r:id"))
                if rel_id:
                    explicit[kind] = rel_id
                    inherited[kind] = rel_id
                    target = relationships.get(rel_id)
                    if target:
                        referenced_parts.add(target)
            effective = dict(inherited)
            footer_fields = {}
            footer_page_field_counts = {}
            footer_non_page_payload = {}
            footer_formats = {}
            for kind, rel_id in effective.items():
                target = relationships.get(rel_id)
                if target and target in package.namelist():
                    footer_root = etree.fromstring(package.read(target))
                    field_types = _field_types(footer_root)
                    footer_fields[kind] = sorted(field_types)
                    footer_page_field_counts[kind] = _page_field_count(footer_root)
                    visible = [
                        value.strip()
                        for value in footer_root.xpath(".//w:t/text()", namespaces=NS)
                        if value.strip()
                    ]
                    footer_non_page_payload[kind] = bool(
                        _pagination_rewrite_shape(footer_root) != "page"
                        or footer_root.xpath(
                            ".//w:tbl | .//w:drawing | .//w:object | .//w:pict",
                            namespaces=NS,
                        )
                        or any(not value.isdigit() for value in visible)
                    )
                    footer_formats[kind] = _page_footer_format(footer_root)
                else:
                    footer_fields[kind] = []
                    footer_page_field_counts[kind] = 0
                    footer_non_page_payload[kind] = False
                    footer_formats[kind] = {}
            pg_num = sect_pr.find(qn("w:pgNumType"))
            vertical_alignment = sect_pr.find(qn("w:vAlign"))
            sections.append(
                {
                    "index": index,
                    "page_number_start": (
                        None if pg_num is None else pg_num.get(qn("w:start"))
                    ),
                    "page_number_format": (
                        None if pg_num is None else pg_num.get(qn("w:fmt"))
                    ),
                    "different_first_page": _on_off_enabled(
                        sect_pr.find(qn("w:titlePg"))
                    ),
                    "vertical_alignment": (
                        None
                        if vertical_alignment is None
                        else vertical_alignment.get(qn("w:val"), "top")
                    ),
                    "footer_references": explicit,
                    "effective_footer_references": effective,
                    "footer_fields": footer_fields,
                    "footer_page_field_counts": footer_page_field_counts,
                    "footer_non_page_payload": footer_non_page_payload,
                    "footer_formats": footer_formats,
                    "missing_page_footer_types": [
                        kind
                        for kind in ("default", "even")
                        if "PAGE" not in footer_fields.get(kind, [])
                    ],
                }
            )
        all_parts = {
            name for name in package.namelist() if HEADER_FOOTER_PART.fullmatch(name)
        }
        all_referenced = {
            relationships.get(rel_id)
            for sect_pr in document.xpath(".//w:sectPr", namespaces=NS)
            for rel_id in sect_pr.xpath(
                "./w:headerReference/@r:id | ./w:footerReference/@r:id", namespaces=NS
            )
        }
        return {
            "odd_and_even_pages_header_footer": _on_off_enabled(
                None if settings is None else settings.find(qn("w:evenAndOddHeaders"))
            ),
            "sections": sections,
            "page_number_restarts": [
                item["index"] for item in sections if item["page_number_start"] is not None
            ],
            "orphan_header_footer_parts": sorted(all_parts - set(all_referenced)),
        }


def _resolve_audit_boundary(
    document: Any,
    locator: dict[str, Any],
    resolver: Callable[[Any, dict[str, Any]], Any],
    *,
    allow_unique_toc_field: bool = False,
    alternate_locators: list[dict[str, Any]] | None = None,
) -> Any:
    try:
        return resolver(document, locator)
    except FormatMonographError as original_error:
        for alternate in alternate_locators or []:
            try:
                return resolver(document, alternate)
            except FormatMonographError:
                continue
        if not allow_unique_toc_field:
            raise original_error
        matches = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph._p.xpath(
                ".//w:fldSimple[starts-with(translate(normalize-space(@w:instr), "
                "'toc', 'TOC'), 'TOC')] | "
                ".//w:instrText[starts-with(translate(normalize-space(.), "
                "'toc', 'TOC'), 'TOC')]"
            )
        ]
        if len(matches) == 1:
            return matches[0]
        raise original_error


def audit_pagination_sections(
    path: Path,
    document: Any,
    settings: dict[str, Any],
    resolver: Callable[[Any, dict[str, Any]], Any],
    structure_map: dict[str, Any] | None = None,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    inventory = pagination_inventory(path)
    if not settings.get("approved"):
        return [], inventory
    toc = section_index_for_paragraph(
        document,
        _resolve_audit_boundary(
            document,
            settings["toc_start"],
            resolver,
            allow_unique_toc_field=True,
        ),
    )
    body_locator = settings["body_start"]
    alternate_body_locators = []
    for heading in (structure_map or {}).get("headings", []):
        heading_locator = heading.get("locator", {})
        if (
            heading.get("approved")
            and heading_locator.get("text_sha256")
            == body_locator.get("text_sha256")
            and heading.get("normalized_text_sha256")
        ):
            alternate = copy.deepcopy(body_locator)
            alternate["text_sha256"] = heading["normalized_text_sha256"]
            alternate_body_locators.append(alternate)
    body_paragraph = _resolve_audit_boundary(
        document,
        body_locator,
        resolver,
        alternate_locators=alternate_body_locators,
    )
    body = section_index_for_paragraph(document, body_paragraph)
    failures = []
    front_matter = (structure_map or {}).get("front_matter", {})
    if front_matter.get("approved"):
        title_paragraph = _resolve_audit_boundary(
            document,
            front_matter["book_title"],
            resolver,
        )
        title_section = section_index_for_paragraph(document, title_paragraph)
        if front_matter.get("insert_toc_heading_if_missing"):
            toc_headings = [
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.strip() == front_matter.get("toc_heading_text")
                and paragraph.style is not None
                and paragraph.style.name == "Monograph TOC Heading"
            ]
            if len(toc_headings) != 1:
                failures.append(
                    {
                        "property": "toc_heading",
                        "expected": front_matter.get("toc_heading_text"),
                        "actual_count": len(toc_headings),
                    }
                )
            elif section_index_for_paragraph(document, toc_headings[0]) != toc:
                failures.append(
                    {"property": "toc_heading_section", "expected": toc}
                )
        title_style = title_paragraph.style
        if title_style is None or not (
            title_style.name == "Monograph Book Title"
            or (
                title_style.name.startswith("Monograph Approved ")
                and semantic_title_heading_role(title_style) == "title"
            )
        ):
            failures.append(
                {
                    "property": "book_title_style",
                    "expected": "Monograph Book Title or approved Title ancestry",
                }
            )
        title_alignment = title_paragraph.alignment
        if title_alignment is None and title_paragraph.style is not None:
            title_alignment = title_paragraph.style.paragraph_format.alignment
        if title_alignment != WD_ALIGN_PARAGRAPH.CENTER:
            failures.append(
                {"property": "book_title_alignment", "expected": "center"}
            )
        title_p_pr = title_paragraph._p.pPr
        title_direct_num_pr = (
            None if title_p_pr is None else title_p_pr.find(qn("w:numPr"))
        )
        if title_direct_num_pr is not None:
            failures.append(
                {"property": "book_title_direct_numbering", "expected": "absent"}
            )
        direct_indent = (
            None if title_p_pr is None else title_p_pr.find(qn("w:ind"))
        )
        direct_indent_conflicts = _nonzero_indent_attributes(direct_indent)
        if direct_indent_conflicts:
            failures.append(
                {
                    "property": "book_title_direct_indent",
                    "expected": 0,
                    "actual": direct_indent_conflicts,
                }
            )
        expected_title_format = {
            "font_name_east_asia": "黑体",
            "font_name_ascii": "Times New Roman",
            "font_size_pt": 22,
            "bold": True,
        }
        expected_title_format.update(front_matter.get("book_title_format", {}))
        if title_style is not None:
            effective_indent = _effective_style_indent_attributes(title_style)
            style_indent_conflicts = {
                attribute: value
                for attribute, value in effective_indent.items()
                if value not in {None, "0"}
            }
            if style_indent_conflicts:
                failures.append(
                    {
                        "property": "book_title_style_indent",
                        "expected": 0,
                        "actual": style_indent_conflicts,
                    }
                )
            actual_ascii = style_effective_font(document, title_style, "ascii")[0]
            actual_east_asia = style_effective_font(document, title_style, "eastAsia")[0]
            if not (
                font_alias_keys(actual_ascii)
                & font_alias_keys(expected_title_format["font_name_ascii"])
            ):
                failures.append(
                    {
                        "property": "book_title_font_name_ascii",
                        "expected": expected_title_format["font_name_ascii"],
                        "actual": actual_ascii,
                    }
                )
            if not (
                font_alias_keys(actual_east_asia)
                & font_alias_keys(expected_title_format["font_name_east_asia"])
            ):
                failures.append(
                    {
                        "property": "book_title_font_name_east_asia",
                        "expected": expected_title_format["font_name_east_asia"],
                        "actual": actual_east_asia,
                    }
                )
            actual_size = None if title_style.font.size is None else title_style.font.size.pt
            if actual_size != float(expected_title_format["font_size_pt"]):
                failures.append(
                    {
                        "property": "book_title_font_size_pt",
                        "expected": float(expected_title_format["font_size_pt"]),
                        "actual": actual_size,
                    }
                )
        style_bold = (
            title_paragraph.style is not None
            and title_paragraph.style.font.bold is True
        )
        if any(
            run.text and run.bold is not True and not (run.bold is None and style_bold)
            for run in title_paragraph.runs
        ):
            failures.append({"property": "book_title_bold", "expected": True})
        expected_spacing = expected_title_format.get("line_spacing_pt")
        if expected_spacing is not None and title_style is not None:
            actual_spacing = title_style.paragraph_format.line_spacing
            actual_spacing_pt = getattr(actual_spacing, "pt", None)
            if actual_spacing_pt != float(expected_spacing):
                failures.append(
                    {
                        "property": "book_title_line_spacing_pt",
                        "expected": float(expected_spacing),
                        "actual": actual_spacing_pt,
                    }
                )
            expected_rule = expected_title_format.get("line_spacing_rule")
            rule_names = {
                WD_LINE_SPACING.AT_LEAST: "at_least",
                WD_LINE_SPACING.EXACTLY: "exact",
                WD_LINE_SPACING.SINGLE: "single",
                WD_LINE_SPACING.ONE_POINT_FIVE: "one_point_five",
                WD_LINE_SPACING.DOUBLE: "double",
            }
            actual_rule = rule_names.get(title_style.paragraph_format.line_spacing_rule)
            if expected_rule and actual_rule != expected_rule:
                failures.append(
                    {
                        "property": "book_title_line_spacing_rule",
                        "expected": expected_rule,
                        "actual": actual_rule,
                    }
                )
    sections = inventory["sections"]
    explicit_front_starts = [
        item for item in sections[: toc + 1] if item["page_number_start"] is not None
    ]
    if len(explicit_front_starts) != 1 or explicit_front_starts[0]["page_number_start"] != "1":
        failures.append(
            {
                "property": "front_page_number_start",
                "expected": "one unambiguous decimal start at 1",
                "actual": [item["index"] for item in explicit_front_starts],
            }
        )
    front = toc if not explicit_front_starts else explicit_front_starts[0]["index"]
    starts = settings.get("start_at", {})
    expected = {front: str(starts.get("toc", 1)), body: str(starts.get("body", 1))}
    for index, start in expected.items():
        if index >= len(sections) or sections[index]["page_number_start"] != start:
            failures.append(
                {"section": index, "property": "page_number_start", "expected": start}
            )
    for item in sections[front + 1 : body]:
        if item["page_number_start"] is not None:
            failures.append(
                {
                    "section": item["index"],
                    "property": "unexpected_front_page_number_restart",
                    "actual": item["page_number_start"],
                }
            )
    if settings.get("continue_after_body_start", True):
        for item in sections[body + 1 :]:
            if item["page_number_start"] is not None:
                failures.append(
                    {
                        "section": item["index"],
                        "property": "unexpected_page_number_restart",
                        "actual": item["page_number_start"],
                    }
                )
    if not inventory["odd_and_even_pages_header_footer"]:
        failures.append({"property": "odd_and_even_pages_header_footer", "expected": True})
    for item in sections[front:]:
        if item["missing_page_footer_types"]:
            failures.append(
                {
                    "section": item["index"],
                    "property": "visible_page_footer_types",
                    "missing": item["missing_page_footer_types"],
                }
            )
        for kind in ("default", "even"):
            count = item["footer_page_field_counts"].get(kind, 0)
            if count != 1:
                failures.append(
                    {
                        "section": item["index"],
                        "property": "page_footer_field_count",
                        "footer_type": kind,
                        "expected": 1,
                        "actual": count,
                    }
                )
            if item["footer_non_page_payload"].get(kind, False):
                failures.append(
                    {
                        "section": item["index"],
                        "property": "page_footer_non_page_payload",
                        "footer_type": kind,
                        "expected": False,
                    }
                )
            expected_alignment = "right" if kind == "default" else "left"
            expected_format = {
                "alignment": expected_alignment,
                "font_ascii": "Times New Roman",
                "font_east_asia": "宋体",
                "font_size_half_points": "18",
                "font_size_cs_half_points": "18",
            }
            if item["footer_formats"].get(kind, {}) != expected_format:
                failures.append(
                    {
                        "section": item["index"],
                        "property": "page_footer_format",
                        "footer_type": kind,
                        "expected": expected_format,
                        "actual": item["footer_formats"].get(kind, {}),
                    }
                )
        if item["different_first_page"]:
            count = item["footer_page_field_counts"].get("first", 0)
            if count != 1 or item["footer_non_page_payload"].get("first", False):
                failures.append(
                    {
                        "section": item["index"],
                        "property": "first_page_footer",
                        "expected": "one bare PAGE field",
                        "actual": count,
                    }
                )
            expected_first = {
                "alignment": "right",
                "font_ascii": "Times New Roman",
                "font_east_asia": "宋体",
                "font_size_half_points": "18",
                "font_size_cs_half_points": "18",
            }
            if item["footer_formats"].get("first", {}) != expected_first:
                failures.append(
                    {
                        "section": item["index"],
                        "property": "first_page_footer_format",
                        "expected": expected_first,
                        "actual": item["footer_formats"].get("first", {}),
                    }
                )
    if inventory["orphan_header_footer_parts"]:
        failures.append(
            {
                "property": "orphan_header_footer_parts",
                "actual_count": len(inventory["orphan_header_footer_parts"]),
            }
        )
    return failures, inventory
