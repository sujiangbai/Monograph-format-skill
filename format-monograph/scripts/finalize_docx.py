#!/usr/bin/env python3
"""Finalize field caches and re-audit a formatted DOCX without overwriting inputs."""

from __future__ import annotations

import argparse
import copy
import ctypes
import errno
import hashlib
import json
import os
import posixpath
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
import stat
import uuid
from dataclasses import dataclass
from urllib.parse import unquote, urlsplit
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Callable

from lxml import etree
from docx.enum.section import WD_HEADER_FOOTER
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.hdrftr import FooterPart

from _common import (
    NS,
    FormatMonographError,
    ensure_docx,
    field_cache_inventory,
    font_alias_keys,
    isolated_approved_style_name,
    load_document,
    protected_payload_manifest,
    run_effective_font,
    style_effective_font,
    style_name_for_selector,
)
from backend_evidence import (
    atomic_write_bytes,
    backend_audit_binding,
    backend_audit_bytes,
    backend_audit_path,
    canonical_backend_projection,
)
from render_docx import locate_soffice
from libreoffice_runtime import macos_internal_macro_soffice
from structure_map import (
    approved_data_tables,
    approved_role_paragraphs,
    has_semantic_structure_map,
    toc_result_contract,
)
from structure_map import (
    load_structure_map,
    structure_content_fingerprint,
    validate_structure_map_source,
)
from validate_profile import validate
from field_writeback import (
    DEFAULT_ALLOWED_FIELD_TYPES,
    parse_fields,
    selective_field_result_writeback,
)
from field_completion import (
    FINALIZATION_EVIDENCE_VERSION,
    completion_evidence,
    final_ready_evidence_errors,
    finalization_evidence_shape_errors,
)
from external_command import ExternalCommandError, parse_external_command
from target_software import MICROSOFT_WORD, resolve_target_id
from docx_pagination import _page_only_footer, _replace_with_page_field
from toc_index_identity import (
    authorization_with_hash,
    canonical_index_descriptor,
    canonical_json_hash,
)


LIBREOFFICE_MACRO_RESULT_TIMEOUT_SECONDS = 300
LIBREOFFICE_MACRO_SHUTDOWN_TIMEOUT_SECONDS = 15
LIBREOFFICE_FIELD_SCRIPT_URI = (
    "vnd.sun.star.script:libreoffice_fields_macro.py$"
    "refresh_from_environment?language=Python&location=user"
)
LIBREOFFICE_LOG_TAIL_CHARS = 8192
LIBREOFFICE_TOC_BASELINE = re.compile(
    r'^TOC \\o "(?P<first>[1-9])-(?P<last>[1-9])" \\h \\z$'
)
FINAL_READY_FIELD_STATES = frozenset({"absent", "selective_verified"})
LIBREOFFICE_DELIVERY_STATUS = "libreoffice_refreshed"
LIBREOFFICE_WRITEBACK_STATUS = "libreoffice_selective"
TRUSTED_STATUS_IDENTITY_VERSION = 1
PUBLICATION_RECOVERY_VERSION = 1
PUBLICATION_RECOVERY_DIRECTORY = ".format-monograph-recovery"
PUBLICATION_RECORD_VERSION = 1
PUBLICATION_RECORD_NAME = "publication-record.json"
PUBLICATION_CLEANUP_POLICY = "manual_only"
# Darwin <sys/stdio.h> and Linux <linux/fs.h> public rename flags.
DARWIN_RENAME_EXCL = 0x00000004
LINUX_RENAME_NOREPLACE = 0x00000001
FIELD_CONTRACT_PART = re.compile(
    r"^word/(?:document|header\d+|footer\d+|footnotes|endnotes)\.xml$"
)
EXTERNAL_FIELD_TYPES = frozenset(
    {"DATABASE", "DDE", "INCLUDEPICTURE", "INCLUDETEXT", "LINK"}
)
EXTERNAL_COMMAND_CWD = Path(__file__).resolve().parent.parent


class LibreOfficeContractError(FormatMonographError):
    def __init__(self, message: str, evidence: dict[str, Any]) -> None:
        super().__init__(message)
        self.evidence = evidence


class FinalizationPublishError(FormatMonographError):
    def __init__(
        self,
        message: str,
        *,
        preserve_staging: bool = False,
    ) -> None:
        super().__init__(message)
        self.preserve_staging = preserve_staging


@dataclass
class _DirectoryAuthority:
    path: Path
    fd: int
    device: int
    inode: int
    owner: int
    mode: int
    backend: str = "posix"
    security_identity: str | None = None
    api: Any = None


@dataclass
class _ArtifactLocation:
    authority: _DirectoryAuthority
    name: str

    @property
    def path(self) -> Path:
        return self.authority.path / self.name


class _WindowsFileInformation(ctypes.Structure):
    _fields_ = (
        ("dwFileAttributes", ctypes.c_uint32),
        ("ftCreationTimeLow", ctypes.c_uint32),
        ("ftCreationTimeHigh", ctypes.c_uint32),
        ("ftLastAccessTimeLow", ctypes.c_uint32),
        ("ftLastAccessTimeHigh", ctypes.c_uint32),
        ("ftLastWriteTimeLow", ctypes.c_uint32),
        ("ftLastWriteTimeHigh", ctypes.c_uint32),
        ("dwVolumeSerialNumber", ctypes.c_uint32),
        ("nFileSizeHigh", ctypes.c_uint32),
        ("nFileSizeLow", ctypes.c_uint32),
        ("nNumberOfLinks", ctypes.c_uint32),
        ("nFileIndexHigh", ctypes.c_uint32),
        ("nFileIndexLow", ctypes.c_uint32),
    )


class _WindowsPublisherApi:
    """Win32 diagnostics only; production publishing remains fail-closed."""

    FILE_ATTRIBUTE_DIRECTORY = 0x00000010
    FILE_ATTRIBUTE_REPARSE_POINT = 0x00000400
    FILE_FLAG_BACKUP_SEMANTICS = 0x02000000
    FILE_FLAG_OPEN_REPARSE_POINT = 0x00200000
    FILE_READ_ATTRIBUTES = 0x00000080
    FILE_SHARE_READ_WRITE = 0x00000003
    OPEN_EXISTING = 3
    INVALID_FILE_ATTRIBUTES = 0xFFFFFFFF
    TOKEN_QUERY = 0x0008
    TOKEN_USER = 1
    OWNER_SECURITY_INFORMATION = 0x00000001
    DACL_SECURITY_INFORMATION = 0x00000004

    def __init__(self) -> None:
        if os.name != "nt" or not hasattr(ctypes, "WinDLL"):
            raise FinalizationPublishError(
                "Windows publisher authority is unavailable on this platform."
            )
        self.kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
        self.advapi32 = ctypes.WinDLL("advapi32", use_last_error=True)
        self._configure()

    def _configure(self) -> None:
        self.kernel32.CreateFileW.argtypes = (
            ctypes.c_wchar_p,
            ctypes.c_uint32,
            ctypes.c_uint32,
            ctypes.c_void_p,
            ctypes.c_uint32,
            ctypes.c_uint32,
            ctypes.c_void_p,
        )
        self.kernel32.CreateFileW.restype = ctypes.c_void_p
        self.kernel32.GetFileInformationByHandle.argtypes = (
            ctypes.c_void_p,
            ctypes.POINTER(_WindowsFileInformation),
        )
        self.kernel32.GetFileInformationByHandle.restype = ctypes.c_int
        self.kernel32.GetFileAttributesW.argtypes = (ctypes.c_wchar_p,)
        self.kernel32.GetFileAttributesW.restype = ctypes.c_uint32
        self.kernel32.MoveFileExW.argtypes = (
            ctypes.c_wchar_p,
            ctypes.c_wchar_p,
            ctypes.c_uint32,
        )
        self.kernel32.MoveFileExW.restype = ctypes.c_int
        self.kernel32.CloseHandle.argtypes = (ctypes.c_void_p,)
        self.kernel32.CloseHandle.restype = ctypes.c_int
        self.kernel32.LocalFree.argtypes = (ctypes.c_void_p,)
        self.kernel32.LocalFree.restype = ctypes.c_void_p
        self.kernel32.GetCurrentProcess.restype = ctypes.c_void_p
        self.advapi32.OpenProcessToken.argtypes = (
            ctypes.c_void_p,
            ctypes.c_uint32,
            ctypes.POINTER(ctypes.c_void_p),
        )
        self.advapi32.OpenProcessToken.restype = ctypes.c_int
        self.advapi32.GetTokenInformation.argtypes = (
            ctypes.c_void_p,
            ctypes.c_uint32,
            ctypes.c_void_p,
            ctypes.c_uint32,
            ctypes.POINTER(ctypes.c_uint32),
        )
        self.advapi32.GetTokenInformation.restype = ctypes.c_int
        self.advapi32.GetFileSecurityW.argtypes = (
            ctypes.c_wchar_p,
            ctypes.c_uint32,
            ctypes.c_void_p,
            ctypes.c_uint32,
            ctypes.POINTER(ctypes.c_uint32),
        )
        self.advapi32.GetFileSecurityW.restype = ctypes.c_int
        self.advapi32.GetSecurityDescriptorOwner.argtypes = (
            ctypes.c_void_p,
            ctypes.POINTER(ctypes.c_void_p),
            ctypes.POINTER(ctypes.c_int),
        )
        self.advapi32.GetSecurityDescriptorOwner.restype = ctypes.c_int
        self.advapi32.EqualSid.argtypes = (ctypes.c_void_p, ctypes.c_void_p)
        self.advapi32.EqualSid.restype = ctypes.c_int
        self.advapi32.ConvertSecurityDescriptorToStringSecurityDescriptorW.argtypes = (
            ctypes.c_void_p,
            ctypes.c_uint32,
            ctypes.c_uint32,
            ctypes.POINTER(ctypes.c_wchar_p),
            ctypes.POINTER(ctypes.c_uint32),
        )
        self.advapi32.ConvertSecurityDescriptorToStringSecurityDescriptorW.restype = ctypes.c_int

    @staticmethod
    def _winerror(label: str) -> OSError:
        code = ctypes.get_last_error()
        return OSError(code, f"{label} failed with WinError {code}")

    def open_directory(self, path: Path) -> int:
        handle = self.kernel32.CreateFileW(
            str(path),
            self.FILE_READ_ATTRIBUTES,
            self.FILE_SHARE_READ_WRITE,
            None,
            self.OPEN_EXISTING,
            self.FILE_FLAG_BACKUP_SEMANTICS | self.FILE_FLAG_OPEN_REPARSE_POINT,
            None,
        )
        invalid = ctypes.c_void_p(-1).value
        if handle in {None, invalid}:
            raise self._winerror(f"CreateFileW({path})")
        try:
            attributes, _, _ = self.query_directory(handle)
            if not attributes & self.FILE_ATTRIBUTE_DIRECTORY:
                raise FinalizationPublishError(
                    f"Windows authority is not a directory: {path}."
                )
            if attributes & self.FILE_ATTRIBUTE_REPARSE_POINT:
                raise FinalizationPublishError(
                    f"Windows directory authority is a reparse point: {path}."
                )
            return int(handle)
        except BaseException:
            self.close(handle)
            raise

    def query_directory(self, handle: int) -> tuple[int, int, int]:
        information = _WindowsFileInformation()
        if not self.kernel32.GetFileInformationByHandle(
            ctypes.c_void_p(handle), ctypes.byref(information)
        ):
            raise self._winerror("GetFileInformationByHandle")
        inode = (information.nFileIndexHigh << 32) | information.nFileIndexLow
        return (
            information.dwFileAttributes,
            information.dwVolumeSerialNumber,
            inode,
        )

    def close(self, handle: int) -> None:
        if handle and not self.kernel32.CloseHandle(ctypes.c_void_p(handle)):
            raise self._winerror("CloseHandle")

    def path_is_reparse(self, path: Path) -> bool:
        attributes = self.kernel32.GetFileAttributesW(str(path))
        if attributes == self.INVALID_FILE_ATTRIBUTES:
            code = ctypes.get_last_error()
            if code in {2, 3}:
                raise FileNotFoundError(code, "path is missing", str(path))
            raise OSError(code, "GetFileAttributesW failed", str(path))
        return bool(attributes & self.FILE_ATTRIBUTE_REPARSE_POINT)

    def atomic_noreplace_move(self, source: Path, target: Path) -> None:
        if self.kernel32.MoveFileExW(str(source), str(target), 0):
            return
        code = ctypes.get_last_error()
        if code in {80, 183}:
            raise FileExistsError(code, "destination exists", str(target))
        if code in {5, 32}:
            raise PermissionError(code, "access denied", str(target))
        if code in {17, 1142}:
            raise OSError(errno.EXDEV, "cross-volume move", str(source), str(target))
        if code in {1, 50, 120}:
            raise FinalizationPublishError(
                f"Windows MoveFileExW no-replace move is unavailable (WinError {code})."
            )
        raise OSError(code, "MoveFileExW failed", str(source), str(target))

    def _current_user_sid(self) -> tuple[ctypes.Array, ctypes.c_void_p]:
        token = ctypes.c_void_p()
        if not self.advapi32.OpenProcessToken(
            self.kernel32.GetCurrentProcess(), self.TOKEN_QUERY, ctypes.byref(token)
        ):
            raise self._winerror("OpenProcessToken")
        try:
            needed = ctypes.c_uint32()
            self.advapi32.GetTokenInformation(
                token, self.TOKEN_USER, None, 0, ctypes.byref(needed)
            )
            buffer = ctypes.create_string_buffer(needed.value)
            if not self.advapi32.GetTokenInformation(
                token,
                self.TOKEN_USER,
                buffer,
                needed.value,
                ctypes.byref(needed),
            ):
                raise self._winerror("GetTokenInformation")
            sid = ctypes.c_void_p.from_buffer(buffer).value
            return buffer, ctypes.c_void_p(sid)
        finally:
            self.close(int(token.value))

    def private_security_identity(self, path: Path) -> str:
        del path
        raise FinalizationPublishError(
            "Windows private-directory authorization is unavailable: an "
            "AccessCheck-based effective-permission proof has not been implemented."
        )


_WINDOWS_PUBLISHER_API: _WindowsPublisherApi | None = None
_PUBLISHER_AUTHORITY_BACKEND_OVERRIDE: str | None = None


def _windows_publisher_api() -> _WindowsPublisherApi:
    global _WINDOWS_PUBLISHER_API
    if _WINDOWS_PUBLISHER_API is None:
        _WINDOWS_PUBLISHER_API = _WindowsPublisherApi()
    return _WINDOWS_PUBLISHER_API


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def artifact_identity(
    path: Path,
    *,
    page_count: int | None = None,
    reported_path: Path | None = None,
) -> dict[str, Any]:
    identity: dict[str, Any] = {
        "path": str((reported_path or path).resolve(strict=False)),
        "sha256": file_sha256(path),
        "size_bytes": path.stat().st_size,
    }
    if page_count is not None:
        identity["page_count"] = page_count
    return identity


def standard_json_bytes(value: Any) -> bytes:
    """Serialize one standard JSON representation for status and console use."""
    try:
        return (
            json.dumps(
                value,
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
                allow_nan=False,
            )
            + "\n"
        ).encode("utf-8")
    except (TypeError, ValueError) as exc:
        raise FormatMonographError(
            "Finalization evidence is not standard JSON."
        ) from exc


def trusted_status_byte_identity(payload: bytes) -> dict[str, Any]:
    return {
        "version": TRUSTED_STATUS_IDENTITY_VERSION,
        "sha256": hashlib.sha256(payload).hexdigest(),
        "size_bytes": len(payload),
    }


def _path_has_control_characters(value: str) -> bool:
    return any(ord(character) < 32 or ord(character) == 127 for character in value)


def _resolve_declared_path(path: Path, label: str) -> tuple[Path, Path]:
    text = os.fspath(path)
    if not text or _path_has_control_characters(text):
        raise FormatMonographError(f"{label} path is empty or contains a control character.")
    try:
        lexical = Path(os.path.abspath(text))
        resolved = path.resolve(strict=False)
    except (OSError, RuntimeError, TypeError, ValueError) as exc:
        raise FormatMonographError(f"{label} path cannot be resolved.") from exc
    return lexical, resolved


def _existing_regular_target(path: Path, label: str) -> bool:
    try:
        if path.is_symlink():
            raise FormatMonographError(f"{label} target must not be a symlink.")
        if not path.exists():
            return False
        if not stat.S_ISREG(path.lstat().st_mode):
            raise FormatMonographError(f"{label} target is not a regular file.")
        return True
    except FormatMonographError:
        raise
    except (OSError, RuntimeError, TypeError, ValueError) as exc:
        raise FormatMonographError(f"{label} target cannot be inspected.") from exc


def resolve_finalization_path_contract(
    *,
    input_path: Path,
    source_path: Path | None,
    profile_path: Path,
    structure_map_path: Path,
    output_path: Path,
    pdf_output: Path | None,
    status_output: Path | None,
    force: bool,
) -> dict[str, Any]:
    """Resolve every path before filesystem mutation and reject all aliasing."""
    inputs = {
        "input": input_path,
        "profile": profile_path,
        "structure_map": structure_map_path,
    }
    if source_path is not None:
        inputs["source"] = source_path
    outputs = {"output": output_path}
    if pdf_output is not None:
        outputs["pdf"] = pdf_output
    if status_output is not None:
        outputs["status"] = status_output
        outputs["audit"] = backend_audit_path(status_output)

    input_paths: dict[str, Path] = {}
    input_lexical: dict[str, Path] = {}
    for label, path in inputs.items():
        lexical, resolved = _resolve_declared_path(path, label)
        input_lexical[label] = lexical
        input_paths[label] = resolved

    output_paths: dict[str, Path] = {}
    output_lexical: dict[str, Path] = {}
    for label, path in outputs.items():
        lexical, resolved = _resolve_declared_path(path, label)
        try:
            if path.is_symlink() or lexical.is_symlink():
                raise FormatMonographError(
                    f"{label} target must not be a symlink."
                )
        except FormatMonographError:
            raise
        except (OSError, RuntimeError, TypeError, ValueError) as exc:
            raise FormatMonographError(
                f"{label} target cannot be inspected."
            ) from exc
        output_lexical[label] = lexical
        output_paths[label] = resolved

    labels = list(output_paths)
    for index, left in enumerate(labels):
        for right in labels[index + 1 :]:
            if output_lexical[left] == output_lexical[right]:
                raise FormatMonographError(
                    f"Output paths {left} and {right} are lexically identical."
                )
            if output_paths[left] == output_paths[right]:
                raise FormatMonographError(
                    f"Output paths {left} and {right} resolve to the same target."
                )
    for output_label, resolved_output in output_paths.items():
        for input_label, resolved_input in input_paths.items():
            if output_lexical[output_label] == input_lexical[input_label]:
                raise FormatMonographError(
                    f"{output_label} path is lexically identical to {input_label}."
                )
            if resolved_output == resolved_input:
                raise FormatMonographError(
                    f"{output_label} path resolves to input {input_label}."
                )

    parents = {path.parent.resolve(strict=False) for path in output_paths.values()}
    if len(parents) != 1:
        raise FormatMonographError(
            "Finalization output, PDF, status, and backend audit must share one directory."
        )
    for label, path in output_paths.items():
        exists = _existing_regular_target(path, label)
        if exists and not force:
            raise FormatMonographError(
                f"{label} target exists; use --force to replace it."
            )
    return {
        "inputs": input_paths,
        "outputs": output_paths,
        "parent": next(iter(parents)),
    }


def _target_snapshot(
    path: Path, *, include_ctime: bool = False
) -> dict[str, Any] | None:
    flags = os.O_RDONLY
    lexical_before: os.stat_result | None = None
    if hasattr(os, "O_CLOEXEC"):
        flags |= os.O_CLOEXEC
    if hasattr(os, "O_NOFOLLOW"):
        flags |= os.O_NOFOLLOW
    else:
        try:
            lexical_before = path.lstat()
        except FileNotFoundError:
            return None
        if not stat.S_ISREG(lexical_before.st_mode):
            raise FormatMonographError(
                f"{path.name} target is not a regular file."
            )
    try:
        descriptor = os.open(path, flags)
    except FileNotFoundError:
        return None
    except (OSError, RuntimeError, TypeError, ValueError) as exc:
        raise FormatMonographError(
            f"{path.name} target cannot be inspected."
        ) from exc
    try:
        before = os.fstat(descriptor)
        if not stat.S_ISREG(before.st_mode):
            raise FormatMonographError(
                f"{path.name} target is not a regular file."
            )
        if lexical_before is not None and (
            lexical_before.st_dev != before.st_dev
            or lexical_before.st_ino != before.st_ino
        ):
            raise FormatMonographError(
                f"{path.name} target changed before it was opened."
            )
        digest = hashlib.sha256()
        while True:
            chunk = os.read(descriptor, 1024 * 1024)
            if not chunk:
                break
            digest.update(chunk)
        after = os.fstat(descriptor)
        stable_fields = (
            "st_dev",
            "st_ino",
            "st_mode",
            "st_size",
            "st_mtime_ns",
            "st_ctime_ns",
        )
        if any(getattr(before, name) != getattr(after, name) for name in stable_fields):
            raise FormatMonographError(
                f"{path.name} target changed while it was inspected."
            )
        snapshot = {
            "device": after.st_dev,
            "inode": after.st_ino,
            "mode": after.st_mode,
            "size": after.st_size,
            "mtime_ns": after.st_mtime_ns,
            "sha256": digest.hexdigest(),
        }
        if include_ctime:
            snapshot["ctime_ns"] = after.st_ctime_ns
        return snapshot
    finally:
        os.close(descriptor)


def _path_lexists(path: Path) -> bool:
    try:
        path.lstat()
        return True
    except FileNotFoundError:
        return False


def _read_regular_bytes_and_snapshot(
    path: Path,
) -> tuple[bytes, dict[str, Any]]:
    flags = os.O_RDONLY
    lexical_before: os.stat_result | None = None
    if hasattr(os, "O_CLOEXEC"):
        flags |= os.O_CLOEXEC
    if hasattr(os, "O_NOFOLLOW"):
        flags |= os.O_NOFOLLOW
    else:
        lexical_before = path.lstat()
        if not stat.S_ISREG(lexical_before.st_mode):
            raise FormatMonographError(
                f"{path.name} target is not a regular file."
            )
    descriptor = os.open(path, flags)
    try:
        before = os.fstat(descriptor)
        if not stat.S_ISREG(before.st_mode):
            raise FormatMonographError(
                f"{path.name} target is not a regular file."
            )
        if lexical_before is not None and (
            lexical_before.st_dev != before.st_dev
            or lexical_before.st_ino != before.st_ino
        ):
            raise FormatMonographError(
                f"{path.name} target changed before it was opened."
            )
        chunks: list[bytes] = []
        digest = hashlib.sha256()
        while True:
            chunk = os.read(descriptor, 1024 * 1024)
            if not chunk:
                break
            chunks.append(chunk)
            digest.update(chunk)
        after = os.fstat(descriptor)
        stable_fields = (
            "st_dev",
            "st_ino",
            "st_mode",
            "st_size",
            "st_mtime_ns",
            "st_ctime_ns",
        )
        if any(getattr(before, name) != getattr(after, name) for name in stable_fields):
            raise FormatMonographError(
                f"{path.name} target changed while it was inspected."
            )
        return b"".join(chunks), {
            "device": after.st_dev,
            "inode": after.st_ino,
            "mode": after.st_mode,
            "size": after.st_size,
            "mtime_ns": after.st_mtime_ns,
            "ctime_ns": after.st_ctime_ns,
            "sha256": digest.hexdigest(),
        }
    finally:
        os.close(descriptor)


def _read_artifact_location(
    location: _ArtifactLocation,
) -> tuple[bytes, dict[str, Any]]:
    """Read one regular entry through its already-open directory authority."""
    _assert_authority_fd(location.authority, "artifact parent")
    _validated_entry_name(location.name, "artifact")
    if location.authority.backend == "windows":
        if location.authority.api.path_is_reparse(location.path):
            raise FormatMonographError(
                f"{location.name} target is a Windows reparse point."
            )
        result = _read_regular_bytes_and_snapshot(location.path)
        _assert_authority_fd(
            location.authority, "artifact parent after inspection"
        )
        return result
    flags = os.O_RDONLY
    if hasattr(os, "O_CLOEXEC"):
        flags |= os.O_CLOEXEC
    if not hasattr(os, "O_NOFOLLOW"):
        raise FinalizationPublishError(
            "Regular-file no-follow inspection is unavailable on this platform."
        )
    flags |= os.O_NOFOLLOW
    descriptor = os.open(location.name, flags, dir_fd=location.authority.fd)
    try:
        before = os.fstat(descriptor)
        if not stat.S_ISREG(before.st_mode):
            raise FormatMonographError(
                f"{location.name} target is not a regular file."
            )
        chunks: list[bytes] = []
        digest = hashlib.sha256()
        while True:
            chunk = os.read(descriptor, 1024 * 1024)
            if not chunk:
                break
            chunks.append(chunk)
            digest.update(chunk)
        after = os.fstat(descriptor)
        stable_fields = (
            "st_dev",
            "st_ino",
            "st_mode",
            "st_size",
            "st_mtime_ns",
            "st_ctime_ns",
        )
        if any(getattr(before, key) != getattr(after, key) for key in stable_fields):
            raise FormatMonographError(
                f"{location.name} target changed while it was inspected."
            )
        return b"".join(chunks), {
            "device": after.st_dev,
            "inode": after.st_ino,
            "mode": after.st_mode,
            "size": after.st_size,
            "mtime_ns": after.st_mtime_ns,
            "ctime_ns": after.st_ctime_ns,
            "sha256": digest.hexdigest(),
        }
    finally:
        os.close(descriptor)


def _artifact_snapshot_at(
    location: _ArtifactLocation,
) -> dict[str, Any] | None:
    try:
        snapshot = _read_artifact_location(location)[1]
        snapshot.pop("ctime_ns", None)
        return snapshot
    except FileNotFoundError:
        return None


def _artifact_full_snapshot_at(
    location: _ArtifactLocation,
) -> dict[str, Any] | None:
    try:
        return _read_artifact_location(location)[1]
    except FileNotFoundError:
        return None


def _validated_entry_name(name: str, label: str) -> bytes:
    if (
        not isinstance(name, str)
        or not name
        or name in {".", ".."}
        or os.sep in name
        or (os.altsep and os.altsep in name)
        or _path_has_control_characters(name)
    ):
        raise FinalizationPublishError(f"{label} entry name is invalid.")
    return os.fsencode(name)


def _open_directory_authority(
    path: Path,
    *,
    parent_fd: int | None = None,
    name: str | None = None,
    require_private_owner: bool = False,
) -> _DirectoryAuthority:
    if os.name == "nt" or _PUBLISHER_AUTHORITY_BACKEND_OVERRIDE == "windows":
        if os.name == "nt" and _PUBLISHER_AUTHORITY_BACKEND_OVERRIDE != "windows":
            raise FinalizationPublishError(
                "Windows publication is unavailable: neither AccessCheck-based "
                "private-directory authorization nor an authority-bound no-replace "
                "rename has been implemented."
            )
        if parent_fd is not None and name is None:
            raise FinalizationPublishError("Directory authority entry is missing.")
        if name is not None:
            _validated_entry_name(name, "directory")
        api = _windows_publisher_api()
        handle = api.open_directory(path)
        try:
            attributes, device, inode = api.query_directory(handle)
            if attributes & api.FILE_ATTRIBUTE_REPARSE_POINT:
                raise FinalizationPublishError(
                    f"Windows directory authority is a reparse point: {path}."
                )
            security_identity = (
                api.private_security_identity(path)
                if require_private_owner
                else None
            )
            return _DirectoryAuthority(
                path=path,
                fd=handle,
                device=device,
                inode=inode,
                owner=0,
                mode=0,
                backend="windows",
                security_identity=security_identity,
                api=api,
            )
        except BaseException:
            api.close(handle)
            raise
    if not hasattr(os, "O_DIRECTORY") or not hasattr(os, "O_NOFOLLOW"):
        raise FinalizationPublishError(
            "Directory-FD authority is unavailable on this platform."
        )
    flags = os.O_RDONLY | os.O_DIRECTORY | os.O_NOFOLLOW
    if hasattr(os, "O_CLOEXEC"):
        flags |= os.O_CLOEXEC
    if parent_fd is None:
        descriptor = os.open(path, flags)
    else:
        if name is None:
            raise FinalizationPublishError("Directory authority entry is missing.")
        _validated_entry_name(name, "directory")
        descriptor = os.open(name, flags, dir_fd=parent_fd)
    try:
        details = os.fstat(descriptor)
        if not stat.S_ISDIR(details.st_mode):
            raise FinalizationPublishError(
                f"Directory authority is not a directory: {path}."
            )
        if require_private_owner:
            if not hasattr(os, "geteuid") or details.st_uid != os.geteuid():
                raise FinalizationPublishError(
                    f"Recovery directory owner is not the effective user: {path}."
                )
            if stat.S_IMODE(details.st_mode) & 0o022:
                raise FinalizationPublishError(
                    f"Recovery directory is group/world writable: {path}."
                )
        return _DirectoryAuthority(
            path=path,
            fd=descriptor,
            device=details.st_dev,
            inode=details.st_ino,
            owner=details.st_uid,
            mode=stat.S_IMODE(details.st_mode),
        )
    except BaseException:
            os.close(descriptor)
            raise


def _assert_publication_platform_available() -> None:
    if os.name == "nt" and _PUBLISHER_AUTHORITY_BACKEND_OVERRIDE != "windows":
        raise FinalizationPublishError(
            "Windows publication is unavailable before output-parent creation: "
            "neither AccessCheck-based private-directory authorization nor an "
            "authority-bound no-replace rename has been implemented."
        )


def _close_authority(authority: _DirectoryAuthority | None) -> None:
    if authority is not None:
        if authority.backend == "windows":
            authority.api.close(authority.fd)
        else:
            os.close(authority.fd)


def _close_authorities_collect(
    authorities: list[tuple[str, _DirectoryAuthority | None]],
) -> list[str]:
    """Attempt every close and return deterministic diagnostics."""
    errors: list[str] = []
    for label, authority in authorities:
        if authority is None or authority.fd < 0:
            continue
        try:
            _close_authority(authority)
        except (OSError, RuntimeError, TypeError, ValueError) as exc:
            errors.append(f"{label} authority close failed: {exc}")
        finally:
            authority.fd = -1
    return errors


def _assert_authority_fd(authority: _DirectoryAuthority, label: str) -> None:
    if authority.backend == "windows":
        attributes, device, inode = authority.api.query_directory(authority.fd)
        if (
            not attributes & authority.api.FILE_ATTRIBUTE_DIRECTORY
            or attributes & authority.api.FILE_ATTRIBUTE_REPARSE_POINT
            or device != authority.device
            or inode != authority.inode
        ):
            raise FinalizationPublishError(
                f"{label} Windows directory handle authority changed."
            )
        return
    details = os.fstat(authority.fd)
    if (
        not stat.S_ISDIR(details.st_mode)
        or details.st_dev != authority.device
        or details.st_ino != authority.inode
        or details.st_uid != authority.owner
        or stat.S_IMODE(details.st_mode) != authority.mode
    ):
        raise FinalizationPublishError(f"{label} directory FD authority changed.")


def _assert_authority_entry(
    parent: _DirectoryAuthority,
    name: str,
    authority: _DirectoryAuthority,
    label: str,
    *,
    require_private_owner: bool = False,
) -> None:
    _validated_entry_name(name, "directory")
    _assert_authority_fd(parent, f"{label} parent")
    _assert_authority_fd(authority, label)
    if parent.backend == "windows" or authority.backend == "windows":
        if parent.backend != "windows" or authority.backend != "windows":
            raise FinalizationPublishError(
                f"{label} mixes incompatible authority backends."
            )
        current_path = parent.path / name
        current_handle = parent.api.open_directory(current_path)
        try:
            attributes, device, inode = parent.api.query_directory(current_handle)
            if (
                not attributes & parent.api.FILE_ATTRIBUTE_DIRECTORY
                or attributes & parent.api.FILE_ATTRIBUTE_REPARSE_POINT
                or device != authority.device
                or inode != authority.inode
            ):
                raise FinalizationPublishError(
                    f"{label} directory entry no longer names its opened authority."
                )
            if require_private_owner:
                current_security = parent.api.private_security_identity(current_path)
                if current_security != authority.security_identity:
                    raise FinalizationPublishError(
                        f"{label} Windows private security descriptor changed."
                    )
        finally:
            parent.api.close(current_handle)
        return
    try:
        details = os.stat(name, dir_fd=parent.fd, follow_symlinks=False)
    except FileNotFoundError as exc:
        raise FinalizationPublishError(
            f"{label} directory entry disappeared."
        ) from exc
    if (
        not stat.S_ISDIR(details.st_mode)
        or details.st_dev != authority.device
        or details.st_ino != authority.inode
        or details.st_uid != authority.owner
        or stat.S_IMODE(details.st_mode) != authority.mode
    ):
        raise FinalizationPublishError(
            f"{label} directory entry no longer names its opened authority."
        )
    if require_private_owner:
        if not hasattr(os, "geteuid") or details.st_uid != os.geteuid():
            raise FinalizationPublishError(
                f"{label} directory owner is not the effective user."
            )
        if stat.S_IMODE(details.st_mode) & 0o022:
            raise FinalizationPublishError(
                f"{label} directory is group/world writable."
            )


def _authority_stat_entry(
    authority: _DirectoryAuthority, name: str
) -> Any:
    _validated_entry_name(name, "artifact")
    _assert_authority_fd(authority, "entry parent")
    if authority.backend == "windows":
        path = authority.path / name
        details = path.lstat()
        if authority.api.path_is_reparse(path):
            raise FinalizationPublishError(
                f"Windows entry is a reparse point: {path}."
            )
        if stat.S_ISDIR(details.st_mode):
            handle = authority.api.open_directory(path)
            try:
                _, device, inode = authority.api.query_directory(handle)
            finally:
                authority.api.close(handle)
            details = SimpleNamespace(
                st_mode=details.st_mode,
                st_dev=device,
                st_ino=inode,
                st_uid=0,
                st_size=details.st_size,
                st_mtime_ns=details.st_mtime_ns,
                st_ctime_ns=details.st_ctime_ns,
            )
        _assert_authority_fd(authority, "entry parent after stat")
        return details
    return os.stat(name, dir_fd=authority.fd, follow_symlinks=False)


def _authority_mkdir(
    authority: _DirectoryAuthority, name: str, mode: int
) -> None:
    _validated_entry_name(name, "directory")
    _assert_authority_fd(authority, "mkdir parent")
    if authority.backend == "windows":
        (authority.path / name).mkdir(mode=mode)
        _assert_authority_fd(authority, "mkdir parent after create")
        return
    os.mkdir(name, mode, dir_fd=authority.fd)


def _create_bound_directory(
    parent: _DirectoryAuthority,
    *,
    prefix: str,
    label: str,
    mode: int = 0o700,
    attempts: int = 64,
) -> _DirectoryAuthority:
    """Exclusively create and bind one new child without adopting an old entry."""
    if not re.fullmatch(r"[A-Za-z0-9._-]+", prefix):
        raise FinalizationPublishError(f"{label} prefix is invalid.")
    _assert_authority_fd(parent, f"{label} parent before create")
    for _ in range(attempts):
        name = f"{prefix}{uuid.uuid4().hex}"
        _validated_entry_name(name, label)
        try:
            _authority_mkdir(parent, name, mode)
        except FileExistsError:
            continue
        path = parent.path / name
        try:
            created = _authority_stat_entry(parent, name)
            if (
                not stat.S_ISDIR(created.st_mode)
                or created.st_dev != parent.device
                or created.st_uid != parent.owner
                or stat.S_IMODE(created.st_mode) != mode
            ):
                raise FinalizationPublishError(
                    f"New {label} entry does not have the required creation identity."
                )
            authority = _open_directory_authority(
                path,
                parent_fd=parent.fd,
                name=name,
                require_private_owner=True,
            )
            try:
                identity_changed = (
                    authority.device != created.st_dev
                    or authority.inode != created.st_ino
                )
                if authority.backend != "windows":
                    identity_changed = identity_changed or (
                        authority.owner != created.st_uid
                        or authority.mode != stat.S_IMODE(created.st_mode)
                    )
                if identity_changed:
                    raise FinalizationPublishError(
                        f"New {label} entry changed between create and bind."
                    )
                _assert_authority_entry(
                    parent,
                    name,
                    authority,
                    f"newly bound {label}",
                    require_private_owner=True,
                )
                return authority
            except BaseException:
                _close_authorities_collect([(label, authority)])
                raise
        except BaseException as exc:
            raise FinalizationPublishError(
                f"Failed to bind newly created {label} at {path}; the entry is "
                "retained and is never adopted or removed."
            ) from exc
    raise FinalizationPublishError(
        f"Could not allocate a unique {label} name after {attempts} attempts."
    )


def _authority_listdir(authority: _DirectoryAuthority) -> list[str]:
    _assert_authority_fd(authority, "list directory")
    if authority.backend == "windows":
        entries = os.listdir(authority.path)
        _assert_authority_fd(authority, "list directory after read")
        return entries
    return os.listdir(authority.fd)


def _raise_atomic_move_error(
    error_code: int, source: Path, target: Path
) -> None:
    if error_code in {errno.EEXIST, errno.ENOTEMPTY}:
        raise FileExistsError(error_code, os.strerror(error_code), str(target))
    if error_code in {errno.EACCES, errno.EPERM}:
        raise PermissionError(error_code, os.strerror(error_code), str(target))
    if error_code == errno.EXDEV:
        raise OSError(error_code, os.strerror(error_code), str(source), str(target))
    unsupported = {
        errno.EINVAL,
        getattr(errno, "ENOTSUP", errno.EINVAL),
        getattr(errno, "EOPNOTSUPP", errno.EINVAL),
        getattr(errno, "ENOSYS", errno.EINVAL),
    }
    if error_code in unsupported:
        raise FinalizationPublishError(
            f"Atomic no-replace move is unsupported for {source.name} -> "
            f"{target.name} (errno={error_code})."
        )
    raise OSError(error_code, os.strerror(error_code), str(source), str(target))


def _platform_atomic_noreplace_move(
    source: _ArtifactLocation,
    target: _ArtifactLocation,
) -> None:
    source_name = _validated_entry_name(source.name, "source")
    target_name = _validated_entry_name(target.name, "target")
    if source.authority.backend == "windows" or os.name == "nt":
        if os.name == "nt" and _PUBLISHER_AUTHORITY_BACKEND_OVERRIDE != "windows":
            raise FinalizationPublishError(
                "Windows pathname-based MoveFileExW publishing is disabled; an "
                "authority-bound no-replace rename is required."
            )
        api = source.authority.api or _windows_publisher_api()
        if target.authority.api is not None and target.authority.api is not api:
            raise FinalizationPublishError(
                "Windows atomic move mixes distinct authority backends."
            )
        api.atomic_noreplace_move(source.path, target.path)
        return
    if sys.platform == "darwin":
        library = ctypes.CDLL(None, use_errno=True)
        function = getattr(library, "renameatx_np", None)
        if function is None:
            raise FinalizationPublishError(
                "macOS renameatx_np(RENAME_EXCL) is unavailable."
            )
        function.argtypes = (
            ctypes.c_int,
            ctypes.c_char_p,
            ctypes.c_int,
            ctypes.c_char_p,
            ctypes.c_uint,
        )
        function.restype = ctypes.c_int
        result = function(
            source.authority.fd,
            source_name,
            target.authority.fd,
            target_name,
            DARWIN_RENAME_EXCL,
        )
        if result != 0:
            _raise_atomic_move_error(ctypes.get_errno(), source.path, target.path)
        return
    if sys.platform.startswith("linux"):
        library = ctypes.CDLL(None, use_errno=True)
        function = getattr(library, "renameat2", None)
        if function is None:
            raise FinalizationPublishError(
                "Linux renameat2(RENAME_NOREPLACE) is unavailable."
            )
        function.argtypes = (
            ctypes.c_int,
            ctypes.c_char_p,
            ctypes.c_int,
            ctypes.c_char_p,
            ctypes.c_uint,
        )
        function.restype = ctypes.c_int
        result = function(
            source.authority.fd,
            source_name,
            target.authority.fd,
            target_name,
            LINUX_RENAME_NOREPLACE,
        )
        if result != 0:
            _raise_atomic_move_error(ctypes.get_errno(), source.path, target.path)
        return
    raise FinalizationPublishError(
        "Atomic no-replace move is unavailable on this platform."
    )


def _atomic_noreplace_move(
    source: _ArtifactLocation,
    target: _ArtifactLocation,
) -> None:
    _assert_absolute_authority(source.authority, "source parent before atomic move")
    _assert_absolute_authority(target.authority, "target parent before atomic move")
    if source.authority.device != target.authority.device:
        raise OSError(
            errno.EXDEV,
            "atomic no-replace move requires one filesystem",
            str(source.path),
            str(target.path),
        )
    try:
        source_details = _authority_stat_entry(source.authority, source.name)
    except FileNotFoundError as exc:
        raise FinalizationPublishError(
            f"Atomic move source is missing: {source.path}."
        ) from exc
    if not stat.S_ISREG(source_details.st_mode):
        raise FinalizationPublishError(
            f"Atomic move source is not a regular non-symlink file: {source.path}."
        )
    if source_details.st_dev != source.authority.device:
        raise OSError(
            errno.EXDEV,
            "atomic move source is on another filesystem",
            str(source.path),
        )
    _platform_atomic_noreplace_move(source, target)
    _assert_absolute_authority(source.authority, "source parent after atomic move")
    _assert_absolute_authority(target.authority, "target parent after atomic move")
    try:
        if source.authority.backend == "windows":
            source.path.lstat()
            _assert_authority_fd(
                source.authority, "source parent after atomic move"
            )
        else:
            os.stat(
                source.name,
                dir_fd=source.authority.fd,
                follow_symlinks=False,
            )
    except FileNotFoundError:
        return
    raise FinalizationPublishError(
        f"Atomic move source pathname was concurrently recreated: {source.path}.",
        preserve_staging=True,
    )


def _status_binding_errors(
    value: dict[str, Any],
    targets: dict[str, Path],
    expected: dict[str, dict[str, Any] | None],
) -> list[str]:
    errors: list[str] = []
    binding = value.get("artifact_binding") or {}
    finalized = binding.get("finalized_docx") or {}
    output_expected = expected.get("output") or {}
    output_path = str(targets["output"].resolve(strict=False))
    if value.get("output") != output_path:
        errors.append("status output path differs from publish target")
    workflow = value.get("workflow_state") or {}
    if workflow.get("output_sha256") != output_expected.get("sha256"):
        errors.append("status workflow output hash differs from staged artifact")
    if finalized.get("path") != output_path:
        errors.append("status finalized DOCX path differs from publish target")
    if finalized.get("sha256") != output_expected.get("sha256"):
        errors.append("status finalized DOCX hash differs from staged artifact")
    if finalized.get("size_bytes") != output_expected.get("size"):
        errors.append("status finalized DOCX size differs from staged artifact")

    pdf_binding = binding.get("word_verification_pdf")
    pdf_expected = expected.get("pdf")
    if pdf_expected is None:
        if pdf_binding is not None or value.get("target_pdf") is not None:
            errors.append("status unexpectedly binds an absent target PDF")
    else:
        pdf_path = str(targets["pdf"].resolve(strict=False))
        if value.get("target_pdf") != pdf_path:
            errors.append("status target PDF path differs from publish target")
        if not isinstance(pdf_binding, dict):
            errors.append("status target PDF binding is missing")
        else:
            if pdf_binding.get("path") != pdf_path:
                errors.append("status target PDF path differs from publish target")
            if pdf_binding.get("sha256") != pdf_expected.get("sha256"):
                errors.append("status target PDF hash differs from staged artifact")
            if pdf_binding.get("size_bytes") != pdf_expected.get("size"):
                errors.append("status target PDF size differs from staged artifact")

    audit_binding = value.get("backend_audit") or {}
    audit_identity = audit_binding.get("artifact") or {}
    audit_expected = expected.get("audit")
    if audit_expected is not None:
        if audit_identity.get("path") != str(
            targets["audit"].resolve(strict=False)
        ):
            errors.append("status backend audit path differs from publish target")
        if audit_identity.get("sha256") != audit_expected.get("sha256"):
            errors.append("status backend audit hash differs from staged artifact")
        if audit_identity.get("size_bytes") != audit_expected.get("size"):
            errors.append("status backend audit size differs from staged artifact")
    return errors


def _trusted_status_errors(
    status_path: Path,
    trusted_identity: dict[str, Any] | None,
    targets: dict[str, Path],
    expected: dict[str, dict[str, Any] | None],
    *,
    status_location: _ArtifactLocation | None = None,
) -> list[str]:
    errors: list[str] = []
    if not isinstance(trusted_identity, dict) or set(trusted_identity) != {
        "version",
        "sha256",
        "size_bytes",
    }:
        return ["trusted status byte identity is missing or malformed"]
    if not (
        type(trusted_identity.get("version")) is int
        and trusted_identity.get("version") == TRUSTED_STATUS_IDENTITY_VERSION
    ):
        errors.append("trusted status byte identity version is unsupported")
    trusted_hash = trusted_identity.get("sha256")
    trusted_size = trusted_identity.get("size_bytes")
    if not (
        isinstance(trusted_hash, str)
        and len(trusted_hash) == 64
        and all(character in "0123456789abcdef" for character in trusted_hash)
    ):
        errors.append("trusted status byte identity hash is invalid")
    if type(trusted_size) is not int or trusted_size < 1:
        errors.append("trusted status byte identity size is invalid")
    try:
        if status_location is None:
            payload, snapshot = _read_regular_bytes_and_snapshot(status_path)
        else:
            payload, snapshot = _read_artifact_location(status_location)
    except (
        FormatMonographError,
        OSError,
        RuntimeError,
        TypeError,
        ValueError,
    ) as exc:
        return errors + [f"status entity cannot be read safely: {exc}"]
    if snapshot["sha256"] != trusted_hash:
        errors.append("status entity hash differs from trusted producer bytes")
    if snapshot["size"] != trusted_size:
        errors.append("status entity size differs from trusted producer bytes")
    try:
        value = json.loads(
            payload.decode("utf-8"),
            parse_constant=lambda token: (_ for _ in ()).throw(
                ValueError(f"non-standard JSON constant {token}")
            ),
        )
    except (UnicodeDecodeError, json.JSONDecodeError, ValueError) as exc:
        return errors + [f"status entity is not standard JSON: {exc}"]
    if not isinstance(value, dict):
        return errors + ["status entity root is not an object"]
    shape_errors = finalization_evidence_shape_errors(value)
    errors.extend(f"status closed-shape: {error}" for error in shape_errors)
    recalculated = final_ready_evidence_errors(completion_evidence(value))
    expected_validation = {
        "status": "pass" if not recalculated else "incomplete",
        "errors": recalculated,
    }
    stored_validation = (value.get("field_completion") or {}).get(
        "evidence_validation"
    )
    if stored_validation != expected_validation:
        errors.append(
            "status field completion validation differs from production recalculation"
        )
    errors.extend(_status_binding_errors(value, targets, expected))
    return list(dict.fromkeys(errors))


def _published_artifact_errors(
    names: list[str],
    target_locations: dict[str, _ArtifactLocation],
    expected: dict[str, dict[str, Any] | None],
) -> list[str]:
    errors: list[str] = []
    for name in names:
        expected_snapshot = expected[name]
        try:
            current = _artifact_snapshot_at(target_locations[name])
            if expected_snapshot is None:
                if current is not None:
                    errors.append(f"{name} target exists but should be absent")
            elif current != expected_snapshot:
                errors.append(f"{name} target differs from staged entity")
        except (FormatMonographError, OSError, RuntimeError, TypeError, ValueError):
            errors.append(f"{name} target cannot be verified")
    return errors


def _restore_preserved_regular(
    source: _ArtifactLocation, target: _ArtifactLocation
) -> str | None:
    try:
        _atomic_noreplace_move(source, target)
        return None
    except FileExistsError:
        return f"restore target is occupied: {target.path.name}"
    except (FinalizationPublishError, OSError, RuntimeError, TypeError, ValueError) as exc:
        return f"restore failed for {target.path.name}: {exc}"


def _recovery_manifest_payload(
    transaction_id: str,
    recovery_root: Path,
    backups: dict[str, _ArtifactLocation],
    targets: dict[str, Path],
    startup_snapshots: dict[str, dict[str, Any] | None],
) -> bytes:
    entries = []
    for name in ("output", "pdf", "audit", "status"):
        location = backups.get(name)
        snapshot = startup_snapshots.get(name)
        if location is None or snapshot is None:
            continue
        recovery_path = location.path
        entries.append(
            {
                "artifact": name,
                "original_target": str(targets[name].resolve(strict=False)),
                "recovery_path": str(recovery_path.absolute()),
                "startup_snapshot": dict(snapshot),
                "recovery_inode_may_continue_changing": True,
            }
        )
    return standard_json_bytes(
        {
            "version": PUBLICATION_RECOVERY_VERSION,
            "transaction_id": transaction_id,
            "recovery_directory": str(recovery_root.absolute()),
            "entries": entries,
            "business_gate": False,
            "cleanup_policy": (
                "manual_only_after_the_operator_proves_no_process_retains_an_open_"
                "descriptor_to_any_recovery_inode"
            ),
        }
    )


def _assert_absolute_authority(
    authority: _DirectoryAuthority, label: str
) -> None:
    _assert_authority_fd(authority, label)
    if authority.backend == "windows":
        current_handle = authority.api.open_directory(authority.path)
        try:
            attributes, device, inode = authority.api.query_directory(current_handle)
            if (
                not attributes & authority.api.FILE_ATTRIBUTE_DIRECTORY
                or attributes & authority.api.FILE_ATTRIBUTE_REPARSE_POINT
                or device != authority.device
                or inode != authority.inode
            ):
                raise FinalizationPublishError(
                    f"{label} path no longer names its opened authority."
                )
            if authority.security_identity is not None:
                current_security = authority.api.private_security_identity(
                    authority.path
                )
                if current_security != authority.security_identity:
                    raise FinalizationPublishError(
                        f"{label} Windows private security descriptor changed."
                    )
        finally:
            authority.api.close(current_handle)
        return
    try:
        details = authority.path.lstat()
    except FileNotFoundError as exc:
        raise FinalizationPublishError(f"{label} path disappeared.") from exc
    if (
        not stat.S_ISDIR(details.st_mode)
        or details.st_dev != authority.device
        or details.st_ino != authority.inode
    ):
        raise FinalizationPublishError(
            f"{label} path no longer names its opened directory authority."
        )


def _assert_recovery_chain(
    output_parent: _DirectoryAuthority,
    recovery_parent: _DirectoryAuthority,
    transaction: _DirectoryAuthority,
    phase: str,
) -> None:
    _assert_absolute_authority(output_parent, f"output parent at {phase}")
    _assert_authority_entry(
        output_parent,
        PUBLICATION_RECOVERY_DIRECTORY,
        recovery_parent,
        f"recovery parent at {phase}",
        require_private_owner=True,
    )
    _assert_authority_entry(
        recovery_parent,
        transaction.path.name,
        transaction,
        f"recovery transaction at {phase}",
        require_private_owner=True,
    )


def _open_recovery_authorities(
    output_parent: _DirectoryAuthority,
    transaction_id: str,
    event_hook: Callable[[str, str | None], None] | None,
) -> tuple[_DirectoryAuthority, _DirectoryAuthority]:
    recovery_path = output_parent.path / PUBLICATION_RECOVERY_DIRECTORY
    try:
        recovery_entry = _authority_stat_entry(
            output_parent, PUBLICATION_RECOVERY_DIRECTORY
        )
    except FileNotFoundError:
        _authority_mkdir(
            output_parent, PUBLICATION_RECOVERY_DIRECTORY, 0o700
        )
        recovery_entry = _authority_stat_entry(
            output_parent, PUBLICATION_RECOVERY_DIRECTORY
        )
    if event_hook:
        event_hook("before_recovery_parent_open", None)
    recovery = _open_directory_authority(
        recovery_path,
        parent_fd=output_parent.fd,
        name=PUBLICATION_RECOVERY_DIRECTORY,
        require_private_owner=True,
    )
    try:
        if (
            recovery_entry.st_dev != recovery.device
            or recovery_entry.st_ino != recovery.inode
        ):
            raise FinalizationPublishError(
                "Recovery parent changed between entry inspection and open."
            )
        if event_hook:
            event_hook("after_recovery_parent_open", None)
        _assert_authority_entry(
            output_parent,
            PUBLICATION_RECOVERY_DIRECTORY,
            recovery,
            "recovery parent after open",
            require_private_owner=True,
        )
        _authority_mkdir(recovery, transaction_id, 0o700)
        transaction_entry = _authority_stat_entry(recovery, transaction_id)
        if event_hook:
            event_hook("after_recovery_transaction_mkdir", None)
        transaction_path = recovery_path / transaction_id
        transaction = _open_directory_authority(
            transaction_path,
            parent_fd=recovery.fd,
            name=transaction_id,
            require_private_owner=True,
        )
        try:
            if (
                transaction_entry.st_dev != transaction.device
                or transaction_entry.st_ino != transaction.inode
            ):
                raise FinalizationPublishError(
                    "Recovery transaction changed between mkdir and open."
                )
            if event_hook:
                event_hook("after_recovery_transaction_create", None)
                event_hook("after_recovery_directory", None)
            _assert_recovery_chain(
                output_parent, recovery, transaction, "transaction create"
            )
            return recovery, transaction
        except BaseException:
            _close_authority(transaction)
            raise
    except BaseException:
        _close_authority(recovery)
        raise


def _write_authority_record_exclusive(
    authority: _DirectoryAuthority, name: str, payload: bytes
) -> dict[str, Any]:
    _validated_entry_name(name, "record")
    if authority.backend == "windows":
        _assert_authority_fd(authority, "record parent")
        path = authority.path / name
        descriptor = os.open(path, os.O_WRONLY | os.O_CREAT | os.O_EXCL, 0o600)
        try:
            view = memoryview(payload)
            while view:
                written = os.write(descriptor, view)
                if written <= 0:
                    raise OSError("record write made no progress")
                view = view[written:]
            details = os.fstat(descriptor)
            if not stat.S_ISREG(details.st_mode) or details.st_size != len(payload):
                raise FinalizationPublishError(
                    "Authority record is not the expected regular file."
                )
            identity = {
                "device": details.st_dev,
                "inode": details.st_ino,
                "mode": details.st_mode,
                "size": details.st_size,
                "mtime_ns": details.st_mtime_ns,
                "ctime_ns": details.st_ctime_ns,
                "sha256": hashlib.sha256(payload).hexdigest(),
            }
        finally:
            os.close(descriptor)
        _assert_authority_fd(authority, "record parent after write")
        if _read_artifact_location(_ArtifactLocation(authority, name))[0] != payload:
            raise FinalizationPublishError(
                "Authority record differs from its producer bytes."
            )
        return identity
    flags = os.O_WRONLY | os.O_CREAT | os.O_EXCL | os.O_NOFOLLOW
    if hasattr(os, "O_CLOEXEC"):
        flags |= os.O_CLOEXEC
    descriptor = os.open(name, flags, 0o600, dir_fd=authority.fd)
    try:
        view = memoryview(payload)
        while view:
            written = os.write(descriptor, view)
            if written <= 0:
                raise OSError("recovery manifest write made no progress")
            view = view[written:]
        details = os.fstat(descriptor)
        if not stat.S_ISREG(details.st_mode) or details.st_size != len(payload):
            raise FinalizationPublishError(
                "Recovery manifest is not the expected regular file."
            )
        identity = {
            "device": details.st_dev,
            "inode": details.st_ino,
            "mode": details.st_mode,
            "size": details.st_size,
            "mtime_ns": details.st_mtime_ns,
            "ctime_ns": details.st_ctime_ns,
            "sha256": hashlib.sha256(payload).hexdigest(),
        }
    finally:
        os.close(descriptor)
    return identity


def _copy_authority_regular_exclusive(
    source: _ArtifactLocation,
    target: _ArtifactLocation,
) -> dict[str, Any]:
    """Copy one producer-owned regular file without following or replacing entries."""
    _validated_entry_name(source.name, "producer artifact")
    _validated_entry_name(target.name, "staged artifact")
    _assert_absolute_authority(source.authority, "producer parent before import")
    _assert_absolute_authority(target.authority, "staging parent before import")
    if source.authority.backend == "windows" or target.authority.backend == "windows":
        if _PUBLISHER_AUTHORITY_BACKEND_OVERRIDE != "windows":
            raise FinalizationPublishError(
                "Windows staged producer import is unavailable without a "
                "handle-relative exclusive-create implementation."
            )
        source_path: str | bytes = str(source.path)
        target_path: str | bytes = str(target.path)
        source_dir_fd = None
        target_dir_fd = None
    else:
        source_path = source.name
        target_path = target.name
        source_dir_fd = source.authority.fd
        target_dir_fd = target.authority.fd
    read_flags = os.O_RDONLY | getattr(os, "O_NOFOLLOW", 0)
    write_flags = (
        os.O_WRONLY
        | os.O_CREAT
        | os.O_EXCL
        | getattr(os, "O_NOFOLLOW", 0)
    )
    if hasattr(os, "O_CLOEXEC"):
        read_flags |= os.O_CLOEXEC
        write_flags |= os.O_CLOEXEC
    source_fd = os.open(source_path, read_flags, dir_fd=source_dir_fd)
    target_fd = -1
    close_errors: list[str] = []
    digest = hashlib.sha256()
    size = 0
    write_identity: dict[str, Any] | None = None
    try:
        source_before = os.fstat(source_fd)
        if not stat.S_ISREG(source_before.st_mode):
            raise FinalizationPublishError(
                f"Producer artifact is not a regular non-symlink file: {source.path}."
            )
        target_fd = os.open(
            target_path,
            write_flags,
            0o600,
            dir_fd=target_dir_fd,
        )
        while True:
            chunk = os.read(source_fd, 1024 * 1024)
            if not chunk:
                break
            digest.update(chunk)
            size += len(chunk)
            view = memoryview(chunk)
            while view:
                written = os.write(target_fd, view)
                if written <= 0:
                    raise OSError("staged artifact write made no progress")
                view = view[written:]
        source_after = os.fstat(source_fd)
        target_details = os.fstat(target_fd)
        if (
            not stat.S_ISREG(target_details.st_mode)
            or target_details.st_size != size
            or source_after.st_dev != source_before.st_dev
            or source_after.st_ino != source_before.st_ino
            or source_after.st_size != source_before.st_size
            or source_after.st_mtime_ns != source_before.st_mtime_ns
        ):
            raise FinalizationPublishError(
                "Producer artifact identity changed during staged import."
            )
        write_identity = {
            "device": target_details.st_dev,
            "inode": target_details.st_ino,
            "mode": target_details.st_mode,
            "size": target_details.st_size,
            "mtime_ns": target_details.st_mtime_ns,
            "ctime_ns": target_details.st_ctime_ns,
            "sha256": digest.hexdigest(),
        }
    finally:
        for label, descriptor in (("staged", target_fd), ("producer", source_fd)):
            if descriptor < 0:
                continue
            try:
                os.close(descriptor)
            except OSError as exc:
                close_errors.append(f"{label} artifact close failed: {exc}")
    if close_errors:
        raise FinalizationPublishError("; ".join(close_errors), preserve_staging=True)
    _assert_absolute_authority(source.authority, "producer parent after import")
    _assert_absolute_authority(target.authority, "staging parent after import")
    if write_identity is None:
        raise FinalizationPublishError(
            "Staged artifact writer did not produce an identity.",
            preserve_staging=True,
        )
    reopened_identity = _artifact_full_snapshot_at(target)
    if reopened_identity != write_identity:
        raise FinalizationPublishError(
            "Staged artifact differs from its write-FD identity after reopen.",
            preserve_staging=True,
        )
    return write_identity


def _write_manifest_exclusive(
    transaction: _DirectoryAuthority, payload: bytes
) -> dict[str, Any]:
    return _write_authority_record_exclusive(
        transaction, "recovery-manifest.json", payload
    )


def _promote_backups_to_recovery(
    transaction_id: str,
    backups: dict[str, _ArtifactLocation],
    targets: dict[str, Path],
    startup_snapshots: dict[str, dict[str, Any] | None],
    published_expected: dict[str, dict[str, Any] | None],
    output_parent: _DirectoryAuthority,
    staging: _DirectoryAuthority,
    recovery_parent: _DirectoryAuthority,
    transaction: _DirectoryAuthority,
    event_hook: Callable[[str, str | None], None] | None,
) -> Path | None:
    if not backups:
        return None
    if not re.fullmatch(r"[0-9a-f]{32}", transaction_id):
        raise FinalizationPublishError("Publication transaction ID is invalid.")
    recovery_root = transaction.path
    _assert_recovery_chain(
        output_parent, recovery_parent, transaction, "before backup moves"
    )
    for name in ("output", "pdf", "audit", "status"):
        backup = backups.get(name)
        snapshot = startup_snapshots.get(name)
        if backup is None or snapshot is None:
            continue
        details = _authority_stat_entry(backup.authority, backup.name)
        if (
            not stat.S_ISREG(details.st_mode)
            or details.st_dev != snapshot["device"]
            or details.st_ino != snapshot["inode"]
        ):
            raise FinalizationPublishError(
                f"Captured {name} recovery inode identity changed."
            )
        if event_hook:
            event_hook("before_recovery_backup_move", name)
        _assert_recovery_chain(
            output_parent, recovery_parent, transaction, f"before {name} move"
        )
        recovery_location = _ArtifactLocation(
            transaction, f"{name}.previous"
        )
        _atomic_noreplace_move(backup, recovery_location)
        backups[name] = recovery_location
        moved = _authority_stat_entry(transaction, recovery_location.name)
        if moved.st_dev != snapshot["device"] or moved.st_ino != snapshot["inode"]:
            raise FinalizationPublishError(
                f"Promoted {name} recovery inode identity changed."
            )
        current_target = published_expected.get(name)
        if current_target is not None and (
            moved.st_dev == current_target["device"]
            and moved.st_ino == current_target["inode"]
        ):
            raise FinalizationPublishError(
                f"Recovery {name} aliases the newly published artifact."
            )
        if event_hook:
            event_hook("after_recovery_promotion", name)
        _assert_recovery_chain(
            output_parent, recovery_parent, transaction, f"after {name} move"
        )
        rechecked = _authority_stat_entry(transaction, recovery_location.name)
        if (
            not stat.S_ISREG(rechecked.st_mode)
            or rechecked.st_dev != snapshot["device"]
            or rechecked.st_ino != snapshot["inode"]
        ):
            raise FinalizationPublishError(
                f"Recovery {name} entry changed after its atomic move."
            )
    manifest_payload = _recovery_manifest_payload(
        transaction_id,
        recovery_root,
        backups,
        targets,
        startup_snapshots,
    )
    if event_hook:
        event_hook("before_recovery_manifest", None)
    _assert_recovery_chain(
        output_parent, recovery_parent, transaction, "before manifest"
    )
    manifest_identity = _write_manifest_exclusive(transaction, manifest_payload)
    if event_hook:
        event_hook("after_recovery_manifest", None)
    _assert_recovery_chain(
        output_parent, recovery_parent, transaction, "after manifest"
    )
    manifest_location = _ArtifactLocation(transaction, "recovery-manifest.json")
    expected_transaction_entries = sorted(
        [location.name for location in backups.values()]
        + ["recovery-manifest.json"]
    )
    if _artifact_full_snapshot_at(manifest_location) != manifest_identity:
        raise FinalizationPublishError(
            "Publication recovery manifest identity changed after publication."
        )
    transaction_entries = sorted(_authority_listdir(transaction))
    if transaction_entries != expected_transaction_entries:
        raise FinalizationPublishError(
            "Recovery transaction entries changed after manifest publication: "
            + ", ".join(transaction_entries)
        )
    if event_hook:
        event_hook("before_recovery_return", None)
    _assert_recovery_chain(
        output_parent, recovery_parent, transaction, "before return"
    )
    if _artifact_full_snapshot_at(manifest_location) != manifest_identity:
        raise FinalizationPublishError(
            "Publication recovery manifest identity changed before return."
        )
    transaction_entries = sorted(_authority_listdir(transaction))
    if transaction_entries != expected_transaction_entries:
        raise FinalizationPublishError(
            "Recovery transaction entries changed before return: "
            + ", ".join(transaction_entries)
        )
    return recovery_root


def _assert_publish_chain(
    output_parent: _DirectoryAuthority,
    staging: _DirectoryAuthority,
    phase: str,
) -> None:
    _assert_absolute_authority(output_parent, f"output parent at {phase}")
    _assert_authority_entry(
        output_parent,
        staging.path.name,
        staging,
        f"publication staging at {phase}",
        require_private_owner=True,
    )


def _publication_record_payload(
    transaction_id: str,
    staging: _DirectoryAuthority,
    recovery_root: Path | None,
) -> bytes:
    return standard_json_bytes(
        {
            "version": PUBLICATION_RECORD_VERSION,
            "transaction_id": transaction_id,
            "status": "committed",
            "retained_staging_directory": str(staging.path.absolute()),
            "recovery_directory": (
                str(recovery_root.absolute()) if recovery_root is not None else None
            ),
            "cleanup_policy": PUBLICATION_CLEANUP_POLICY,
            "business_gate": False,
        }
    )


def _rollback_published_artifacts(
    order: list[str],
    target_locations: dict[str, _ArtifactLocation],
    backups: dict[str, _ArtifactLocation],
    published_expected: dict[str, dict[str, Any] | None],
    initial_snapshots: dict[str, dict[str, Any] | None],
    output_parent: _DirectoryAuthority,
    staging: _DirectoryAuthority,
    event_hook: Callable[[str, str | None], None] | None,
) -> tuple[bool, list[str]]:
    preserve_staging = False
    errors: list[str] = []
    for name in reversed(order):
        target = target_locations[name]
        expected = published_expected.get(name)
        quarantine: _ArtifactLocation | None = None
        if event_hook:
            event_hook("before_rollback_capture", name)
        try:
            _assert_absolute_authority(
                output_parent, f"output parent before {name} rollback"
            )
            current = _artifact_snapshot_at(target)
            if current is not None:
                quarantine = _ArtifactLocation(
                    staging, f"quarantine-{name}-{uuid.uuid4().hex}"
                )
                _atomic_noreplace_move(target, quarantine)
                if event_hook:
                    event_hook("after_rollback_capture", name)
        except (
            FinalizationPublishError,
            FormatMonographError,
            OSError,
            RuntimeError,
            TypeError,
            ValueError,
        ) as exc:
            preserve_staging = True
            errors.append(f"{name} rollback capture failed: {exc}")
            continue

        captured_is_ours = False
        if quarantine is not None and expected is not None:
            try:
                captured_is_ours = _artifact_snapshot_at(quarantine) == expected
            except (FormatMonographError, OSError, RuntimeError, TypeError, ValueError):
                captured_is_ours = False
        if quarantine is not None and not captured_is_ours:
            preserve_staging = True
            errors.append(f"{name} rollback captured an unknown concurrent entity")
            restore_unknown = _restore_preserved_regular(quarantine, target)
            if restore_unknown:
                errors.append(restore_unknown)
        elif quarantine is not None:
            preserve_staging = True
            errors.append(
                f"{name} verified rollback quarantine retained; failure paths "
                "never unlink artifact entries"
            )

        backup = backups.get(name)
        if backup is not None:
            try:
                backup_is_original = (
                    _artifact_snapshot_at(backup) == initial_snapshots.get(name)
                )
            except (
                FormatMonographError,
                OSError,
                RuntimeError,
                TypeError,
                ValueError,
            ):
                backup_is_original = False
            if not backup_is_original:
                preserve_staging = True
                errors.append(
                    f"{name} recovery entry no longer matches the captured old inode"
                )
            else:
                restore_error = _restore_preserved_regular(backup, target)
                if restore_error:
                    preserve_staging = True
                    errors.append(f"{name} old backup retained: {restore_error}")
        elif initial_snapshots.get(name) is None and not captured_is_ours:
            # Unknown concurrent entities remain either at target or quarantine.
            pass

    if not preserve_staging:
        for name in order:
            try:
                if (
                    _artifact_snapshot_at(target_locations[name])
                    != initial_snapshots.get(name)
                ):
                    preserve_staging = True
                    errors.append(f"{name} rollback result differs from startup snapshot")
            except (FormatMonographError, OSError, RuntimeError, TypeError, ValueError):
                preserve_staging = True
                errors.append(f"{name} rollback result cannot be verified")
    return preserve_staging, errors


def publish_staged_artifacts(
    staged: dict[str, Path],
    targets: dict[str, Path],
    snapshots: dict[str, dict[str, Any] | None],
    staging_root: Path,
    *,
    trusted_status_identity: dict[str, Any] | None = None,
    transaction_id: str | None = None,
    event_hook: Callable[[str, str | None], None] | None = None,
    output_parent_authority: _DirectoryAuthority | None = None,
    staging_authority: _DirectoryAuthority | None = None,
    retained_staging_authorities: list[_DirectoryAuthority] | None = None,
    staged_expected: dict[str, dict[str, Any] | None] | None = None,
) -> dict[str, Any]:
    """Fail-close multi-file publish using descriptor-bound no-replace moves."""
    transaction_id = transaction_id or uuid.uuid4().hex
    if not re.fullmatch(r"[0-9a-f]{32}", transaction_id):
        raise FinalizationPublishError("Publication transaction ID is invalid.")
    order = [name for name in ("output", "pdf", "audit", "status") if name in targets]
    if not order or "output" not in targets:
        raise FinalizationPublishError("Publication target set is incomplete.")
    output_parent: _DirectoryAuthority | None = output_parent_authority
    staging: _DirectoryAuthority | None = staging_authority
    retained_authorities = list(retained_staging_authorities or [])
    recovery_parent: _DirectoryAuthority | None = None
    transaction: _DirectoryAuthority | None = None
    target_locations: dict[str, _ArtifactLocation] = {}
    staged_locations: dict[str, _ArtifactLocation] = {}
    expected: dict[str, dict[str, Any] | None] = {}
    backups: dict[str, _ArtifactLocation] = {}
    processed: list[str] = []
    recovery_root: Path | None = None
    planned_recovery_root = targets["output"].parent / PUBLICATION_RECOVERY_DIRECTORY / transaction_id
    phase = "startup"
    active_name: str | None = None
    try:
        if not isinstance(staged_expected, dict) or set(staged_expected) != set(order):
            raise FinalizationPublishError(
                "Publisher requires an exact producer-supplied staged identity set."
            )
        imported_expected = copy.deepcopy(staged_expected)
        parent_path = targets["output"].parent
        if output_parent is None:
            output_parent = _open_directory_authority(
                parent_path, require_private_owner=True
            )
        elif output_parent.path != parent_path:
            raise FinalizationPublishError(
                "Provided output-parent authority does not match the publication target."
            )
        _assert_absolute_authority(output_parent, "output parent at startup")
        if staging_root.parent.resolve(strict=False) != parent_path.resolve(strict=False):
            raise FinalizationPublishError(
                "Publication staging must be a direct child of the output parent."
            )
        if event_hook:
            event_hook("before_staging_authority_open", None)
        if staging is None:
            staging = _open_directory_authority(
                staging_root,
                parent_fd=output_parent.fd,
                name=staging_root.name,
                require_private_owner=True,
            )
        elif staging.path != staging_root:
            raise FinalizationPublishError(
                "Provided staging authority does not match the publication transaction."
            )
        if event_hook:
            event_hook("after_staging_authority_open", None)
        _assert_publish_chain(output_parent, staging, "startup")
        for retained in retained_authorities:
            _assert_authority_entry(
                staging,
                retained.path.name,
                retained,
                "retained producer workspace at startup",
                require_private_owner=True,
            )
        if output_parent.device != staging.device:
            raise OSError(
                errno.EXDEV,
                "publication staging and output parent use different filesystems",
                str(staging_root),
            )
        for name in order:
            if targets[name].parent.resolve(strict=False) != parent_path.resolve(strict=False):
                raise FinalizationPublishError(
                    f"{name} target is outside the opened output parent."
                )
            if staged[name].parent.resolve(strict=False) != staging_root.resolve(strict=False):
                raise FinalizationPublishError(
                    f"Staged {name} artifact is outside the opened staging directory."
                )
            if staged[name].name != targets[name].name:
                raise FinalizationPublishError(
                    f"Staged {name} basename differs from its publish target."
                )
            target_locations[name] = _ArtifactLocation(output_parent, targets[name].name)
            staged_locations[name] = _ArtifactLocation(staging, staged[name].name)
            try:
                imported_snapshot = _artifact_full_snapshot_at(
                    staged_locations[name]
                )
            except (
                FormatMonographError,
                OSError,
                RuntimeError,
                TypeError,
                ValueError,
            ) as exc:
                raise FinalizationPublishError(
                    f"Staged {name} artifact cannot be inspected."
                ) from exc
            if imported_snapshot != imported_expected[name]:
                raise FinalizationPublishError(
                    f"Staged {name} artifact differs from its imported identity."
                )
            if name != "pdf" and imported_snapshot is None:
                raise FinalizationPublishError(
                    f"Staged {name} artifact is missing."
                )
            expected[name] = copy.deepcopy(imported_snapshot)
            if expected[name] is not None:
                expected[name].pop("ctime_ns", None)
        for name in order:
            if _artifact_snapshot_at(target_locations[name]) != snapshots[name]:
                raise FinalizationPublishError(
                    f"{name} target differs from its startup snapshot."
                )
        if "status" in order:
            status_errors = _trusted_status_errors(
                staged["status"],
                trusted_status_identity,
                targets,
                expected,
                status_location=staged_locations["status"],
            )
            if status_errors:
                raise FinalizationPublishError("; ".join(status_errors))
        if event_hook:
            event_hook("after_start_snapshot", None)
        for name in [item for item in order if item != "status"]:
            active_name = name
            target = target_locations[name]
            snapshot = snapshots[name]
            phase = "before_capture"
            if event_hook:
                event_hook(phase, name)
            _assert_publish_chain(output_parent, staging, phase)
            if snapshot is not None:
                backup = _ArtifactLocation(
                    staging, f"backup-{name}-{uuid.uuid4().hex}"
                )
                _atomic_noreplace_move(target, backup)
                backups[name] = backup
                if event_hook:
                    event_hook("after_capture_rename", name)
                captured = _artifact_snapshot_at(backup)
                if captured != snapshot:
                    raise FinalizationPublishError(
                        f"{name} captured entity differs from startup snapshot"
                    )
            phase = "after_capture"
            if event_hook:
                event_hook(phase, name)
            phase = "before_publish"
            if event_hook:
                event_hook(phase, name)
            _assert_publish_chain(output_parent, staging, phase)
            if expected[name] is not None:
                if event_hook:
                    event_hook("before_atomic_publish", name)
                _atomic_noreplace_move(staged_locations[name], target)
                processed.append(name)
                if event_hook:
                    event_hook("after_atomic_publish", name)
            elif _artifact_snapshot_at(target) is not None:
                raise FinalizationPublishError(
                    f"{name} target was concurrently created before absent publish"
                )
            else:
                processed.append(name)
            phase = "after_publish"
            if event_hook:
                event_hook(phase, name)

        phase = "before_commit_validation"
        if event_hook:
            event_hook(phase, None)
        data_names = [name for name in processed if name != "status"]
        validation_errors = _published_artifact_errors(
            data_names, target_locations, expected
        )
        if "status" in staged:
            validation_errors.extend(
                _trusted_status_errors(
                    staged["status"],
                    trusted_status_identity,
                    targets,
                    expected,
                    status_location=staged_locations["status"],
                )
            )
        if validation_errors:
            raise FinalizationPublishError("; ".join(validation_errors))

        if "status" in order:
            name = "status"
            active_name = name
            target = target_locations[name]
            snapshot = snapshots[name]
            phase = "before_capture"
            if event_hook:
                event_hook(phase, name)
            _assert_publish_chain(output_parent, staging, phase)
            if snapshot is not None:
                backup = _ArtifactLocation(
                    staging, f"backup-{name}-{uuid.uuid4().hex}"
                )
                _atomic_noreplace_move(target, backup)
                backups[name] = backup
                if event_hook:
                    event_hook("after_capture_rename", name)
                if _artifact_snapshot_at(backup) != snapshot:
                    raise FinalizationPublishError(
                        "status captured entity differs from startup snapshot"
                    )
            phase = "after_capture"
            if event_hook:
                event_hook(phase, name)
            phase = "before_publish"
            if event_hook:
                event_hook(phase, name)
            _assert_publish_chain(output_parent, staging, phase)
            if event_hook:
                event_hook("before_atomic_publish", name)
            _atomic_noreplace_move(staged_locations[name], target)
            processed.append(name)
            if event_hook:
                event_hook("after_atomic_publish", name)
            phase = "after_publish"
            if event_hook:
                event_hook(phase, name)

        phase = "post_commit_validation"
        if event_hook:
            event_hook("before_post_commit_validation", None)
        validation_errors = _published_artifact_errors(
            processed, target_locations, expected
        )
        if "status" in staged:
            validation_errors.extend(
                _trusted_status_errors(
                    targets["status"],
                    trusted_status_identity,
                    targets,
                    expected,
                    status_location=target_locations["status"],
                )
            )
        if validation_errors:
            raise FinalizationPublishError("; ".join(validation_errors))
        if event_hook:
            event_hook("after_post_commit_validation", None)
        phase = "recovery_promotion"
        if backups:
            recovery_parent, transaction = _open_recovery_authorities(
                output_parent, transaction_id, event_hook
            )
            recovery_root = _promote_backups_to_recovery(
                transaction_id,
                backups,
                targets,
                snapshots,
                expected,
                output_parent,
                staging,
                recovery_parent,
                transaction,
                event_hook,
            )
        phase = "staging_retention"
        if event_hook:
            event_hook("before_staging_retention", None)
        _assert_publish_chain(output_parent, staging, "before staging retention")
        staging_entries = sorted(_authority_listdir(staging))
        expected_retained_entries = sorted(
            authority.path.name for authority in retained_authorities
        )
        if staging_entries != expected_retained_entries:
            raise FinalizationPublishError(
                "Publication staging contains unexpected entries before its "
                "commit record: " + ", ".join(staging_entries),
                preserve_staging=True,
            )
        for retained in retained_authorities:
            _assert_authority_entry(
                staging,
                retained.path.name,
                retained,
                "retained producer workspace before commit record",
                require_private_owner=True,
            )
        publication_payload = _publication_record_payload(
            transaction_id, staging, recovery_root
        )
        publication_identity = _write_authority_record_exclusive(
            staging, PUBLICATION_RECORD_NAME, publication_payload
        )
        if event_hook:
            event_hook("after_staging_retention", None)
        _assert_publish_chain(output_parent, staging, "after staging retention")
        publication_location = _ArtifactLocation(
            staging, PUBLICATION_RECORD_NAME
        )
        if _artifact_full_snapshot_at(publication_location) != publication_identity:
            raise FinalizationPublishError(
                "Retained publication record identity changed before return.",
                preserve_staging=True,
            )
        expected_committed_entries = sorted(
            expected_retained_entries + [PUBLICATION_RECORD_NAME]
        )
        committed_entries = sorted(_authority_listdir(staging))
        if committed_entries != expected_committed_entries:
            raise FinalizationPublishError(
                "Retained staging entries changed after publication record commit: "
                + ", ".join(committed_entries),
                preserve_staging=True,
            )
        for retained in retained_authorities:
            _assert_authority_entry(
                staging,
                retained.path.name,
                retained,
                "retained producer workspace after commit record",
                require_private_owner=True,
            )
        publication = {
            "version": PUBLICATION_RECORD_VERSION,
            "transaction_id": transaction_id,
            "recovery_directory": (
                str(recovery_root.absolute())
                if recovery_root is not None
                else None
            ),
            "retained_previous_targets": len(backups),
            "staging_removed": False,
            "staging_retained": True,
            "retained_staging_directory": str(staging.path.absolute()),
            "publication_record": str(publication_location.path.absolute()),
            "cleanup_policy": PUBLICATION_CLEANUP_POLICY,
            "business_gate": False,
            "commit_state": "committed",
            "cleanup_errors": [],
        }
        close_errors = _close_authorities_collect(
            [("transaction", transaction), ("recovery", recovery_parent)]
            + [
                (f"retained producer {authority.path.name}", authority)
                for authority in retained_authorities
            ]
            + [("staging", staging), ("output parent", output_parent)]
        )
        transaction = None
        recovery_parent = None
        retained_authorities = []
        staging = None
        output_parent = None
        if close_errors:
            publication["commit_state"] = "committed_with_cleanup_errors"
            publication["cleanup_errors"] = close_errors
        return publication
    except (
        FinalizationPublishError,
        FormatMonographError,
        OSError,
        RuntimeError,
        TypeError,
        ValueError,
    ) as exc:
        rollback_order = list(processed)
        for name in backups:
            if name not in rollback_order:
                rollback_order.append(name)
        if (
            active_name is not None
            and active_name in target_locations
            and active_name not in rollback_order
        ):
            rollback_order.append(active_name)
        if output_parent is not None and staging is not None:
            preserve, rollback_errors = _rollback_published_artifacts(
                rollback_order,
                target_locations,
                backups,
                expected,
                snapshots,
                output_parent,
                staging,
                event_hook,
            )
        else:
            preserve, rollback_errors = False, []
        for name, location in staged_locations.items():
            try:
                current_staged = _artifact_snapshot_at(location)
                if current_staged is not None and current_staged != expected.get(name):
                    preserve = True
                    rollback_errors.append(
                        f"staged {name} entry differs from the producer snapshot"
                    )
            except (
                FormatMonographError,
                OSError,
                RuntimeError,
                TypeError,
                ValueError,
            ):
                preserve = True
                rollback_errors.append(
                    f"staged {name} entry is no longer a known regular artifact"
                )
        preserve = (
            preserve
            or isinstance(exc, FileExistsError)
            or phase == "recovery_promotion"
            or (output_parent is not None and staging is not None)
            or (
                isinstance(exc, FinalizationPublishError)
                and exc.preserve_staging
            )
        )
        preserve = True
        location = f"phase={phase}, target={active_name or 'all'}"
        detail = f"Publish failed at {location}: {exc}"
        if rollback_errors:
            detail += "; " + "; ".join(rollback_errors)
        detail += (
            f"; publication transaction {transaction_id}; retained staging at "
            f"{staging_root.absolute()}; cleanup_policy={PUBLICATION_CLEANUP_POLICY}"
        )
        if phase == "recovery_promotion":
            detail += (
                f"; recovery transaction {transaction_id} may be retained at "
                f"{planned_recovery_root}; path authority must be re-established "
                "before any manual action"
            )
        close_errors = _close_authorities_collect(
            [("transaction", transaction), ("recovery", recovery_parent)]
            + [
                (f"retained producer {authority.path.name}", authority)
                for authority in retained_authorities
            ]
            + [("staging", staging), ("output parent", output_parent)]
        )
        transaction = None
        recovery_parent = None
        retained_authorities = []
        staging = None
        output_parent = None
        if close_errors:
            detail += "; close diagnostics: " + "; ".join(close_errors)
        raise FinalizationPublishError(
            detail,
            preserve_staging=preserve,
        ) from exc
    finally:
        _close_authorities_collect(
            [("transaction", transaction), ("recovery", recovery_parent)]
            + [
                (f"retained producer {authority.path.name}", authority)
                for authority in retained_authorities
            ]
            + [("staging", staging), ("output parent", output_parent)]
        )


def controlled_field_result_writeback(
    baseline_path: Path, refreshed_path: Path, output_path: Path
) -> list[str]:
    """Compatibility wrapper around V0.3.2 selective field-result writeback."""
    report = selective_field_result_writeback(
        baseline_path,
        refreshed_path,
        output_path,
        allowed_field_types=DEFAULT_ALLOWED_FIELD_TYPES,
    )
    return list(report["patched_parts"])


def rewrite_field_flags(path: Path, *, deferred: bool) -> None:
    temp_path = path.with_name(f".{path.name}.fields.tmp")
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(
        temp_path, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/document.xml":
                root = etree.fromstring(data)
                for field in root.xpath(
                    ".//w:fldSimple | .//w:fldChar[@w:fldCharType='begin']",
                    namespaces=NS,
                ):
                    dirty = f"{{{NS['w']}}}dirty"
                    if deferred:
                        field.set(dirty, "true")
                    else:
                        field.attrib.pop(dirty, None)
                data = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
            elif info.filename == "word/settings.xml":
                root = etree.fromstring(data)
                updates = root.xpath("./w:updateFields", namespaces=NS)
                if deferred and not updates:
                    update = etree.Element(f"{{{NS['w']}}}updateFields")
                    update.set(f"{{{NS['w']}}}val", "true")
                    root.append(update)
                elif not deferred:
                    for update in updates:
                        root.remove(update)
                data = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
            target.writestr(info, data)
    os.replace(temp_path, path)


def libreoffice_macro_command(soffice: str, profile: Path) -> list[str]:
    return [
        soffice,
        "--headless",
        "--invisible",
        "--nologo",
        "--nodefault",
        "--norestore",
        f"-env:UserInstallation={profile.resolve().as_uri()}",
        LIBREOFFICE_FIELD_SCRIPT_URI,
    ]


def toc_index_authorization(
    toc_contract: list[dict[str, Any]] | None, baseline_path: Path
) -> dict | None:
    if toc_contract is None:
        return None
    contract_hash = canonical_json_hash(toc_contract)
    try:
        descriptor = canonical_index_descriptor(baseline_path, contract_hash)
    except (KeyError, ValueError, zipfile.BadZipFile) as exc:
        raise FormatMonographError(str(exc)) from exc
    if len(descriptor["indexes"]) != 1:
        raise FormatMonographError(
            "Approved TOC authorization requires exactly one stable OOXML TOC identity."
        )
    return authorization_with_hash(descriptor)


def _process_log_tail(stream) -> str:
    stream.flush()
    stream.seek(0)
    value = stream.read()
    return value[-LIBREOFFICE_LOG_TAIL_CHARS:]


def _instruction_nodes(root: etree._Element, record) -> list[etree._Element]:
    if record.form != "complex" or record.begin is None or record.separate is None:
        return []
    elements = list(root.iter())
    try:
        start = elements.index(record.begin)
        end = elements.index(record.separate)
    except ValueError as exc:
        raise FormatMonographError(
            "LibreOffice field instruction boundaries are missing."
        ) from exc
    return [
        element
        for element in elements[start + 1 : end]
        if element.tag == qn("w:instrText")
    ]


def _field_contract_manifest(root: etree._Element) -> list[tuple]:
    return [
        (record.form, record.parent_order, record.field_type, record.instruction)
        for record in parse_fields(root)
    ]


def _canonical_xml_semantics(element: etree._Element) -> tuple:
    return (
        element.tag,
        tuple(sorted(element.attrib.items())),
        (element.text or "").strip(),
        tuple(_canonical_xml_semantics(child) for child in element),
    )


def section_pagination_manifest(root: etree._Element) -> list[tuple]:
    """Capture section boundaries and every sectPr semantic value."""
    tree = root.getroottree()
    return [
        (tree.getpath(section), _canonical_xml_semantics(section))
        for section in root.xpath(".//w:sectPr", namespaces=NS)
    ]


def _element_semantic_values(
    element: etree._Element, path: str | None = None
) -> dict:
    name = etree.QName(element).localname
    path = path or name
    values = {f"{path}/@{etree.QName(key).localname}": value for key, value in element.attrib.items()}
    if (element.text or "").strip():
        values[f"{path}/#text"] = (element.text or "").strip()
    counts: dict[str, int] = {}
    for child in element:
        child_name = etree.QName(child).localname
        counts[child_name] = counts.get(child_name, 0) + 1
        values.update(
            _element_semantic_values(
                child, f"{path}/{child_name}[{counts[child_name]}]"
            )
        )
    return values


def section_pagination_differences(
    baseline_root: etree._Element, refreshed_root: etree._Element
) -> list[dict[str, Any]]:
    baseline_sections = baseline_root.xpath(".//w:sectPr", namespaces=NS)
    refreshed_sections = refreshed_root.xpath(".//w:sectPr", namespaces=NS)
    differences: list[dict[str, Any]] = []
    if len(baseline_sections) != len(refreshed_sections):
        differences.append(
            {
                "property": "section_count",
                "baseline": len(baseline_sections),
                "libreoffice": len(refreshed_sections),
            }
        )
        return differences
    baseline_tree = baseline_root.getroottree()
    refreshed_tree = refreshed_root.getroottree()
    for ordinal, (baseline_section, refreshed_section) in enumerate(
        zip(baseline_sections, refreshed_sections)
    ):
        baseline_path = baseline_tree.getpath(baseline_section)
        refreshed_path = refreshed_tree.getpath(refreshed_section)
        if baseline_path != refreshed_path:
            differences.append(
                {
                    "section": ordinal,
                    "property": "boundary_path",
                    "baseline": baseline_path,
                    "libreoffice": refreshed_path,
                }
            )
        baseline_values = _element_semantic_values(baseline_section)
        refreshed_values = _element_semantic_values(refreshed_section)
        for key in sorted(set(baseline_values) | set(refreshed_values)):
            if baseline_values.get(key) != refreshed_values.get(key):
                differences.append(
                    {
                        "section": ordinal,
                        "property": key,
                        "baseline": baseline_values.get(key),
                        "libreoffice": refreshed_values.get(key),
                    }
                )
    return differences


def package_field_contract_manifest(path: Path) -> dict[str, list[tuple]]:
    with zipfile.ZipFile(path) as package:
        return {
            name: _field_contract_manifest(etree.fromstring(package.read(name)))
            for name in package.namelist()
            if FIELD_CONTRACT_PART.fullmatch(name)
        }


def package_external_connection_inventory(path: Path) -> list[dict[str, str]]:
    """Inventory active external package links before LibreOffice can load them."""
    connections = []
    with zipfile.ZipFile(path) as package:
        package_names = set(package.namelist())
        for name in package.namelist():
            data = package.read(name)
            if name.endswith(".rels"):
                root = etree.fromstring(data)
                for relationship in root:
                    target = relationship.get("Target", "")
                    reason = _relationship_target_rejection_reason(
                        name,
                        target,
                        relationship.get("TargetMode"),
                        package_names,
                    )
                    if reason is not None:
                        connections.append(
                            {
                                "part": name,
                                "kind": "relationship",
                                "target": target,
                                "reason": reason,
                            }
                        )
            elif name.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                except etree.XMLSyntaxError:
                    continue
                for element in root.iter():
                    for attribute, value in element.attrib.items():
                        if etree.QName(attribute).localname in {"href", "src"}:
                            reason = _xml_package_link_rejection_reason(
                                name, value, package_names
                            )
                        else:
                            reason = None
                        if reason is not None:
                            connections.append(
                                {
                                    "part": name,
                                    "kind": "xml_link",
                                    "target": value,
                                    "reason": reason,
                                }
                            )
                if FIELD_CONTRACT_PART.fullmatch(name):
                    for record in parse_fields(root):
                        if record.field_type in EXTERNAL_FIELD_TYPES:
                            connections.append(
                                {
                                    "part": name,
                                    "kind": "external_field",
                                    "target": record.instruction,
                                }
                            )
    unique = {
        (item["part"], item["kind"], item["target"]): item for item in connections
    }
    return [unique[key] for key in sorted(unique)]


def _relationship_source_part(relationship_part: str) -> str | None:
    normalized = relationship_part.lstrip("/")
    if normalized == "_rels/.rels":
        return ""
    directory, filename = posixpath.split(normalized)
    if posixpath.basename(directory) != "_rels" or not filename.endswith(".rels"):
        return None
    source_directory = posixpath.dirname(directory)
    source_name = filename[: -len(".rels")]
    return posixpath.join(source_directory, source_name)


def _relationship_target_rejection_reason(
    relationship_part: str,
    value: str,
    target_mode: str | None,
    package_names: set[str],
) -> str | None:
    candidate = unquote(value.strip())
    if isinstance(target_mode, str) and target_mode.casefold() == "external":
        return "target_mode_external"
    if not candidate:
        return "empty_relationship_target"
    parsed = urlsplit(candidate)
    if parsed.scheme:
        return f"uri_scheme:{parsed.scheme.casefold()}"
    if parsed.netloc or candidate.startswith(("//", "\\\\")):
        return "network_path"
    source_part = _relationship_source_part(relationship_part)
    if source_part is None:
        return "invalid_relationship_part"
    if source_part and source_part not in package_names:
        return "missing_relationship_source_part"
    path = parsed.path.replace("\\", "/")
    if not path:
        resolved = source_part
    elif path.startswith("/"):
        resolved = posixpath.normpath(path.lstrip("/"))
    else:
        resolved = posixpath.normpath(
            posixpath.join(posixpath.dirname(source_part), path)
        )
    if resolved.startswith("../") or resolved == "..":
        return "outside_package"
    if not resolved or resolved not in package_names:
        return "unresolved_package_reference"
    return None


def _xml_package_link_rejection_reason(
    part_name: str, value: str, package_names: set[str]
) -> str | None:
    candidate = unquote(value.strip())
    if not candidate or candidate.startswith("#"):
        return None
    parsed = urlsplit(candidate)
    if parsed.scheme:
        return f"uri_scheme:{parsed.scheme.casefold()}"
    if parsed.netloc or candidate.startswith(("//", "\\\\")):
        return "network_path"
    path = parsed.path.replace("\\", "/")
    if not path:
        return None
    resolved = (
        posixpath.normpath(path.lstrip("/"))
        if path.startswith("/")
        else posixpath.normpath(posixpath.join(posixpath.dirname(part_name), path))
    )
    if resolved.startswith("../") or resolved == "..":
        return "outside_package"
    if resolved not in package_names:
        return "unresolved_package_reference"
    return None


def restore_known_libreoffice_toc_instruction_order(
    baseline_path: Path, refreshed_path: Path, output_path: Path
) -> dict:
    """Restore exact baseline contracts around strictly matched LO field caches."""
    restored = 0
    restored_pairs = []
    with zipfile.ZipFile(baseline_path) as baseline, zipfile.ZipFile(
        refreshed_path
    ) as refreshed, zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as target:
        if "word/document.xml" not in baseline.namelist() or "word/document.xml" not in refreshed.namelist():
            raise FormatMonographError("LibreOffice refresh lost the main document part.")
        baseline_root = etree.fromstring(baseline.read("word/document.xml"))
        refreshed_root = etree.fromstring(refreshed.read("word/document.xml"))
        baseline_records = parse_fields(baseline_root)
        refreshed_records = parse_fields(refreshed_root)
        observed_instruction_pairs = [
            {
                "baseline": baseline_record.instruction,
                "libreoffice": refreshed_record.instruction,
            }
            for baseline_record, refreshed_record in zip(
                baseline_records, refreshed_records
            )
            if baseline_record.instruction != refreshed_record.instruction
        ]
        baseline_sections = section_pagination_manifest(baseline_root)
        refreshed_sections = section_pagination_manifest(refreshed_root)
        if baseline_sections != refreshed_sections:
            section_differences = section_pagination_differences(
                baseline_root, refreshed_root
            )
            raise LibreOfficeContractError(
                "LibreOffice changed pagination semantics or section boundaries.",
                {
                    "status": "rejected",
                    "pagination_semantics_identical": False,
                    "section_boundaries_identical": not any(
                        item.get("property")
                        in {"section_count", "boundary_path"}
                        for item in section_differences
                    ),
                    "section_differences": section_differences,
                    "observed_field_instruction_pairs": observed_instruction_pairs,
                },
            )
        section_serialization_differences = int(
            [
                etree.tostring(value)
                for value in baseline_root.xpath(".//w:sectPr", namespaces=NS)
            ]
            != [
                etree.tostring(value)
                for value in refreshed_root.xpath(".//w:sectPr", namespaces=NS)
            ]
        )
        if len(baseline_records) != len(refreshed_records):
            raise FormatMonographError(
                "LibreOffice changed the field occurrence count before instruction restoration."
            )
        for baseline_record, refreshed_record in zip(
            baseline_records, refreshed_records
        ):
            if (
                baseline_record.form != refreshed_record.form
                or baseline_record.parent_order != refreshed_record.parent_order
                or baseline_record.field_type != refreshed_record.field_type
            ):
                raise FormatMonographError(
                    "LibreOffice changed field ordering or boundaries before instruction restoration."
                )
            if baseline_record.instruction == refreshed_record.instruction:
                continue
            match = LIBREOFFICE_TOC_BASELINE.fullmatch(baseline_record.instruction)
            expected = (
                None
                if match is None
                else "TOC \\z \\o "
                f'"{match.group("first")}-{match.group("last")}" \\h'
            )
            if (
                baseline_record.form != "complex"
                or baseline_record.parent_order is not None
                or baseline_record.field_type != "TOC"
                or refreshed_record.instruction != expected
            ):
                raise FormatMonographError(
                    "LibreOffice changed a field instruction outside the exact approved TOC permutation."
                )
            nodes = _instruction_nodes(refreshed_root, refreshed_record)
            if not nodes:
                raise FormatMonographError(
                    "LibreOffice TOC instruction has no restorable text node."
                )
            nodes[0].text = f" {baseline_record.instruction} "
            for node in nodes[1:]:
                node.text = ""
            restored += 1
            restored_pairs.append(
                {
                    "baseline": baseline_record.instruction,
                    "libreoffice": refreshed_record.instruction,
                }
            )
        if _field_contract_manifest(baseline_root) != _field_contract_manifest(
            refreshed_root
        ):
            raise FormatMonographError(
                "LibreOffice field contract is not identical after TOC restoration."
            )
        for info in refreshed.infolist():
            data = refreshed.read(info.filename)
            if info.filename == "word/document.xml":
                data = etree.tostring(
                    refreshed_root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
            target.writestr(info, data)
    return {
        "status": "restored_exact_known_permutation" if restored else "not_needed",
        "restored_toc_instructions": restored,
        "restored_toc_instruction_pairs": restored_pairs,
        "field_contract_identical": True,
        "pagination_semantics_identical": True,
        "section_boundaries_identical": True,
        "section_serialization_differences_observed": section_serialization_differences,
    }


def libreoffice_macro_refresh(
    input_path: Path,
    output_path: Path,
    soffice: str,
    renderer_source: str,
    *,
    toc_authorization: dict | None = None,
    toc_contract: list[dict[str, Any]] | None = None,
) -> dict:
    external_connections = package_external_connection_inventory(input_path)
    if external_connections:
        raise FormatMonographError(
            "LibreOffice internal macro refuses documents with active external connections: "
            + json.dumps(external_connections, ensure_ascii=False)
        )
    helper = Path(__file__).with_name("libreoffice_fields_macro.py")
    with tempfile.TemporaryDirectory(prefix="format-monograph-fields-") as temp_name:
        temp = Path(temp_name)
        profile = temp / "lo-profile"
        script_dir = profile / "user" / "Scripts" / "python"
        script_dir.mkdir(parents=True)
        shutil.copy2(helper, script_dir / helper.name)
        result_path = temp / "field-result.json"
        env = os.environ.copy()
        env.update(
            {
                "FORMAT_MONOGRAPH_FIELD_INPUT": str(input_path.resolve()),
                "FORMAT_MONOGRAPH_FIELD_OUTPUT": str(output_path.resolve()),
                "FORMAT_MONOGRAPH_FIELD_RESULT": str(result_path.resolve()),
                "FORMAT_MONOGRAPH_TOC_AUTHORIZATION": (
                    json.dumps(toc_authorization, sort_keys=True)
                    if toc_authorization is not None
                    else ""
                ),
                "FORMAT_MONOGRAPH_TOC_CONTRACT": (
                    json.dumps(toc_contract, sort_keys=True)
                    if toc_contract is not None
                    else ""
                ),
            }
        )
        with tempfile.TemporaryFile(
            mode="w+t", encoding="utf-8"
        ) as stdout_log, tempfile.TemporaryFile(
            mode="w+t", encoding="utf-8"
        ) as stderr_log:
            process = subprocess.Popen(
                libreoffice_macro_command(soffice, profile),
                stdout=stdout_log,
                stderr=stderr_log,
                text=True,
                env=env,
            )
            deadline = time.monotonic() + LIBREOFFICE_MACRO_RESULT_TIMEOUT_SECONDS
            try:
                while time.monotonic() < deadline:
                    if result_path.is_file():
                        break
                    if process.poll() is not None:
                        break
                    time.sleep(0.1)
                if not result_path.is_file():
                    outcome = (
                        "timed out"
                        if process.poll() is None
                        else "exited without a result"
                    )
                    raise FormatMonographError(
                        f"LibreOffice internal field macro {outcome}. "
                        f"returncode={process.returncode} "
                        f"stdout={_process_log_tail(stdout_log).strip()} "
                        f"stderr={_process_log_tail(stderr_log).strip()}"
                    )
                try:
                    details = json.loads(result_path.read_text(encoding="utf-8"))
                except (OSError, json.JSONDecodeError) as exc:
                    raise FormatMonographError(
                        "LibreOffice internal field macro returned an invalid result file."
                    ) from exc
                if not isinstance(details, dict):
                    raise FormatMonographError(
                        "LibreOffice internal field macro returned a non-object result."
                    )
                if not details.get("ok") or not output_path.is_file():
                    raise FormatMonographError(
                        "LibreOffice internal field macro failed. "
                        f"details={json.dumps(details, ensure_ascii=False)}"
                    )
            finally:
                if process.poll() is None:
                    process.terminate()
                    try:
                        process.wait(
                            timeout=LIBREOFFICE_MACRO_SHUTDOWN_TIMEOUT_SECONDS
                        )
                    except subprocess.TimeoutExpired:
                        process.kill()
                        process.wait(
                            timeout=LIBREOFFICE_MACRO_SHUTDOWN_TIMEOUT_SECONDS
                        )
        rewrite_field_flags(output_path, deferred=False)
        details.pop("ok", None)
        details.update(
            {
                "backend": "libreoffice_uno",
                "renderer": soffice,
                "renderer_source": renderer_source,
                "uno_mode": "internal_python_macro",
            }
        )
        return details


def libreoffice_refresh(
    input_path: Path,
    output_path: Path,
    renderer: str | None,
    *,
    toc_authorization: dict | None = None,
    toc_contract: list[dict[str, Any]] | None = None,
) -> dict:
    soffice, renderer_source = locate_soffice(renderer)
    macro_soffice = macos_internal_macro_soffice(soffice)
    if macro_soffice:
        return libreoffice_macro_refresh(
            input_path,
            output_path,
            macro_soffice,
            renderer_source,
            toc_authorization=toc_authorization,
            toc_contract=toc_contract,
        )
    raise FormatMonographError(
        "LibreOffice field refresh requires the verified macOS internal-Python "
        "macro host; the legacy UNO server/helper backend is disabled."
    )


def field_contract_preserved(before: dict, after: dict) -> bool:
    if before["main_toc_fields"] > after["main_toc_fields"]:
        return False
    after_types = after.get("field_types", {})
    return all(
        int(after_types.get(name, 0)) >= int(count)
        for name, count in before.get("field_types", {}).items()
    )


def effective_font_failures(
    path: Path, profile: dict, structure_map: dict | None = None
) -> list[dict[str, Any]]:
    document = load_document(path)
    result = []
    property_attributes = {
        "font_name": "ascii",
        "font_name_ascii": "ascii",
        "font_name_east_asia": "eastAsia",
        "font_name_complex_script": "cs",
    }
    for rule in profile.get("rules", []):
        if rule.get("status") != "approved" or rule.get("application") != "automatic":
            continue
        selector = rule.get("selector", {})
        selector_kind = selector.get("kind")

        def audit_style_fonts(style: Any, style_label: str) -> None:
            for property_name, attribute in property_attributes.items():
                expected = rule.get("properties", {}).get(property_name)
                if not expected:
                    continue
                actual, source = style_effective_font(document, style, attribute)
                if not actual or not (
                    font_alias_keys(str(actual)) & font_alias_keys(str(expected))
                ):
                    result.append(
                        {
                            "rule": rule.get("id"),
                            "style": style_label,
                            "property": property_name,
                            "expected": str(expected),
                            "actual": actual,
                            "source": source,
                        }
                    )
        if selector_kind == "table_role":
            try:
                targets = (
                    approved_data_tables(document, structure_map)
                    if structure_map and has_semantic_structure_map(structure_map)
                    else [(table, {}) for table in document.tables]
                )
            except FormatMonographError:
                result.append(
                    {
                        "rule": rule.get("id"),
                        "selector": selector_kind,
                        "reason": "semantic_target_unresolvable",
                    }
                )
                continue
            for table_index, (table, entry) in enumerate(targets):
                for property_name, attribute in property_attributes.items():
                    expected = rule.get("properties", {}).get(property_name)
                    if not expected:
                        continue
                    for row_index, row in enumerate(table.rows):
                        if entry.get("caption_row") is not None and row_index == int(
                            entry["caption_row"]
                        ):
                            continue
                        for cell_index, cell in enumerate(row.cells):
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if not run.text:
                                        continue
                                    actual, source = run_effective_font(
                                        document, paragraph, run, attribute
                                    )
                                    if not actual or not (
                                        font_alias_keys(str(actual))
                                        & font_alias_keys(str(expected))
                                    ):
                                        result.append(
                                            {
                                                "rule": rule.get("id"),
                                                "table": table_index,
                                                "row": row_index,
                                                "cell": cell_index,
                                                "property": property_name,
                                                "expected": str(expected),
                                                "actual": actual,
                                                "source": source,
                                            }
                                        )
            continue
        if (
            structure_map
            and has_semantic_structure_map(structure_map)
            and selector_kind
            in {"paragraph_role", "caption_role", "bibliography_role"}
        ):
            derived_name = isolated_approved_style_name(selector)
            derived_style = None
            if derived_name is not None:
                try:
                    derived_style = document.styles[derived_name]
                except KeyError:
                    pass
            try:
                targets = approved_role_paragraphs(document, structure_map, selector)
            except FormatMonographError:
                result.append(
                    {
                        "rule": rule.get("id"),
                        "selector": selector_kind,
                        "reason": "semantic_target_unresolvable",
                    }
                )
                continue
            if not targets and derived_name is not None:
                if derived_style is not None:
                    result.append(
                        {
                            "rule": rule.get("id"),
                            "style": derived_name,
                            "reason": "derived_style_without_approved_target",
                        }
                    )
                continue
            if derived_style is not None:
                expected_base = style_name_for_selector(selector)
                if (
                    derived_style.type != WD_STYLE_TYPE.PARAGRAPH
                    or derived_style.base_style is None
                    or derived_style.base_style.name != expected_base
                ):
                    result.append(
                        {
                            "rule": rule.get("id"),
                            "style": derived_name,
                            "reason": "derived_style_binding_mismatch",
                            "expected_base": expected_base,
                            "actual_base": (
                                None
                                if derived_style.base_style is None
                                else derived_style.base_style.name
                            ),
                        }
                    )
                for paragraph_index, paragraph in enumerate(targets):
                    if (
                        paragraph.style is None
                        or paragraph.style.style_id != derived_style.style_id
                    ):
                        result.append(
                            {
                                "rule": rule.get("id"),
                                "paragraph": paragraph_index,
                                "reason": "derived_target_style_mismatch",
                                "expected_style": derived_name,
                                "actual_style": (
                                    None
                                    if paragraph.style is None
                                    else paragraph.style.name
                                ),
                            }
                        )
                audit_style_fonts(derived_style, derived_name)
            if targets:
                for property_name, attribute in property_attributes.items():
                    expected = rule.get("properties", {}).get(property_name)
                    if not expected:
                        continue
                    for paragraph_index, paragraph in enumerate(targets):
                        for run_index, run in enumerate(paragraph.runs):
                            if not run.text:
                                continue
                            actual, source = run_effective_font(
                                document, paragraph, run, attribute
                            )
                            if not actual or not (
                                font_alias_keys(str(actual))
                                & font_alias_keys(str(expected))
                            ):
                                result.append(
                                    {
                                        "rule": rule.get("id"),
                                        "paragraph": paragraph_index,
                                        "run": run_index,
                                        "property": property_name,
                                        "expected": str(expected),
                                        "actual": actual,
                                        "source": source,
                                    }
                                )
                continue
        style_name = style_name_for_selector(rule.get("selector", {}))
        if not style_name:
            continue
        try:
            style = document.styles[style_name]
        except KeyError:
            result.append(
                {"rule": rule.get("id"), "style": style_name, "reason": "missing_style"}
            )
            continue
        audit_style_fonts(style, style_name)
    return result


def use_deferred_output(input_path: Path, output_path: Path, reason: str) -> dict:
    output_path.unlink(missing_ok=True)
    shutil.copy2(input_path, output_path)
    rewrite_field_flags(output_path, deferred=True)
    return {"backend": "deferred_on_open", "fallback_from": reason}


def use_deferred_output_with_evidence(
    input_path: Path,
    output_path: Path,
    reason: str,
    attempted_backend: dict,
    *,
    stage: str,
    error: str,
    failed_checks: list[str],
) -> dict:
    deferred = use_deferred_output(input_path, output_path, reason)
    attempted = copy.deepcopy(attempted_backend)
    attempted["failure"] = {
        "status": "rejected",
        "stage": stage,
        "error": error,
        "failed_checks": list(failed_checks),
    }
    deferred["attempted_backend"] = attempted
    return deferred


def _external_command(value: str) -> list[str]:
    try:
        return parse_external_command(value)
    except ExternalCommandError as exc:
        raise FormatMonographError(str(exc)) from exc


def _require_external_target(response: dict[str, Any], target_id: str) -> None:
    reported = response.get("software")
    if not str(reported or "").strip():
        raise FormatMonographError(
            "External field updater must report its target software."
        )
    reported_id = resolve_target_id(reported)
    if target_id != MICROSOFT_WORD or reported_id != target_id:
        raise FormatMonographError(
            "External updater target software does not match the approved "
            f"target ID ({reported_id!r} != {target_id!r})."
        )
    response["target_id"] = reported_id


def _require_word_target_id(target_id: str) -> None:
    if target_id != MICROSOFT_WORD:
        raise FormatMonographError(
            "External Word backend requires target ID "
            f"{MICROSOFT_WORD!r}; got {target_id!r}."
        )


def _invoke_external_command(
    command: str, request: dict[str, Any], label: str
) -> subprocess.CompletedProcess[str]:
    try:
        return subprocess.run(
            _external_command(command),
            cwd=EXTERNAL_COMMAND_CWD,
            input=json.dumps(request, ensure_ascii=False),
            capture_output=True,
            text=True,
            timeout=600,
            check=False,
        )
    except OSError as exc:
        raise FormatMonographError(
            f"{label} command is unavailable: {exc}"
        ) from exc


def external_refresh(
    input_path: Path,
    output_path: Path,
    command: str,
    profile_path: Path,
    structure_map_path: Path,
    pdf_output: Path | None,
    target_software: str,
    *,
    allowed_field_types: set[str] | frozenset[str] = DEFAULT_ALLOWED_FIELD_TYPES,
) -> dict:
    target_software = resolve_target_id(target_software)
    _require_word_target_id(target_software)
    request = {
        "protocol_version": "1.1",
        "operation": "refresh_fields",
        "input_path": str(input_path.resolve()),
        "output_path": str(output_path.resolve()),
        "profile_path": str(profile_path.resolve()),
        "structure_map_path": str(structure_map_path.resolve()),
        "allowed_field_types": sorted(allowed_field_types),
        "target_software": target_software,
        "pdf_output_path": str(pdf_output.resolve()) if pdf_output else None,
    }
    input_hash = file_sha256(input_path)
    completed = _invoke_external_command(command, request, "External field updater")
    if file_sha256(input_path) != input_hash:
        raise FormatMonographError("External field updater changed its input DOCX.")
    if completed.returncode != 0:
        raise FormatMonographError(
            "External field updater failed. "
            f"stdout={completed.stdout.strip()} stderr={completed.stderr.strip()}"
        )
    try:
        response = json.loads(completed.stdout)
    except json.JSONDecodeError as exc:
        raise FormatMonographError(
            "External field updater did not return one JSON response."
        ) from exc
    if not isinstance(response, dict):
        raise FormatMonographError("External field updater response must be an object.")
    _require_external_target(response, target_software)
    required_true = ("repaginated", "saved", "field_cache_verified")
    if response.get("status") != "success" or any(
        response.get(name) is not True for name in required_true
    ):
        raise FormatMonographError(
            "External field updater did not confirm repagination, save, and cache verification."
        )
    if response.get("structural_changes_applied") != 0:
        raise FormatMonographError(
            "External field updater changed pagination or document structure."
        )
    if not output_path.is_file() or output_path.stat().st_size == 0:
        raise FormatMonographError("External field updater did not create the output DOCX.")
    updated_types = set(response.get("updated_field_types", []))
    allowed_types = set(request["allowed_field_types"])
    if not updated_types <= allowed_types:
        raise FormatMonographError(
            "External field updater reported a non-approved field type."
        )
    response.setdefault("backend", "external")
    response["command"] = _external_command(command)[0]
    return response


def external_measure(
    input_path: Path,
    command: str,
    profile_path: Path,
    structure_map_path: Path,
    target_software: str,
) -> dict:
    target_software = resolve_target_id(target_software)
    _require_word_target_id(target_software)
    request = {
        "protocol_version": "1.1",
        "operation": "measure_layout",
        "input_path": str(input_path.resolve()),
        "profile_path": str(profile_path.resolve()),
        "structure_map_path": str(structure_map_path.resolve()),
        "allowed_field_types": sorted(DEFAULT_ALLOWED_FIELD_TYPES),
        "target_software": target_software,
        "block_spacer_style_name": "Monograph Figure Table Spacer",
    }
    input_hash = file_sha256(input_path)
    completed = _invoke_external_command(command, request, "External layout measurer")
    if file_sha256(input_path) != input_hash:
        raise FormatMonographError("External layout measurer changed its input DOCX.")
    if completed.returncode != 0:
        raise FormatMonographError(
            "External layout measurement failed. "
            f"stdout={completed.stdout.strip()} stderr={completed.stderr.strip()}"
        )
    try:
        response = json.loads(completed.stdout)
    except json.JSONDecodeError as exc:
        raise FormatMonographError(
            "External layout measurer did not return one JSON response."
        ) from exc
    required = {
        "status": "success",
        "operation": "measure_layout",
        "repaginated": True,
        "saved": False,
        "read_only_verified": True,
        "structural_changes_applied": 0,
    }
    if not isinstance(response, dict) or any(
        response.get(name) != value for name, value in required.items()
    ):
        raise FormatMonographError(
            "External layout measurer did not satisfy the read-only contract."
        )
    _require_external_target(response, target_software)
    ordinals = response.get("page_boundary_spacer_ordinals", [])
    if not isinstance(ordinals, list) or any(
        not isinstance(value, int) or value < 0 for value in ordinals
    ):
        raise FormatMonographError(
            "External layout measurer returned invalid spacer ordinals."
        )
    sections = response.get("sections", [])
    if not isinstance(sections, list) or any(
        not isinstance(item, dict) for item in sections
    ):
        raise FormatMonographError(
            "External layout measurer returned invalid section metrics."
        )
    page_count = response.get("page_count")
    if not isinstance(page_count, int) or isinstance(page_count, bool) or page_count < 1:
        raise FormatMonographError(
            "External layout measurer returned an invalid page count."
        )
    response.setdefault("backend", "external")
    response["command"] = _external_command(command)[0]
    return response


def approved_front_matter_section_indexes(
    input_path: Path,
    structure_map: dict[str, Any],
) -> set[int]:
    front_matter = structure_map.get("front_matter", {})
    pagination = structure_map.get("pagination_sections", {})
    if not front_matter.get("approved") or not pagination.get("approved"):
        return set()
    with zipfile.ZipFile(input_path) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    sections = root.xpath(".//w:sectPr", namespaces=NS)
    restart_indexes = [
        index
        for index, section in enumerate(sections)
        if section.find(qn("w:pgNumType")) is not None
        and section.find(qn("w:pgNumType")).get(qn("w:start")) is not None
    ]
    if len(restart_indexes) != 2 or restart_indexes[1] != restart_indexes[0] + 1:
        raise FormatMonographError(
            "Approved title/TOC/body pagination requires exactly two adjacent restarts."
        )
    return set(restart_indexes)


def approved_front_matter_section_types(
    section_indexes: set[int],
    measurement: dict[str, Any],
    input_path: Path,
) -> dict[int, str]:
    if not section_indexes:
        return {}
    document = load_document(input_path)
    metrics = {
        int(item["section_index"]): item
        for item in measurement.get("sections", [])
        if isinstance(item, dict) and "section_index" in item
    }
    result = {}
    for index in sorted(section_indexes):
        previous = metrics.get(index - 1)
        if previous is None or previous.get("last_content_page") is None:
            raise FormatMonographError(
                "Target layout measurement omitted a front-matter section boundary."
            )
        desired_page = int(previous["last_content_page"]) + 1
        target = "evenPage" if desired_page % 2 == 0 else "oddPage"
        section_type = document.sections[index]._sectPr.find(qn("w:type"))
        current = (
            "nextPage"
            if section_type is None
            else section_type.get(qn("w:val"), "nextPage")
        )
        if current != target:
            result[index] = target
    return result


def apply_measured_layout_adjustments(
    input_path: Path,
    output_path: Path,
    ordinals: list[int],
    section_types: dict[int, str] | None = None,
) -> int:
    section_types = section_types or {}
    if any(value not in {"evenPage", "oddPage"} for value in section_types.values()):
        raise FormatMonographError("Measured section type is not an approved parity start.")
    selected = set(ordinals)
    if len(selected) != len(ordinals):
        raise FormatMonographError("Measured spacer ordinals must be unique.")
    with zipfile.ZipFile(input_path) as package:
        document_data = package.read("word/document.xml")
        root = etree.fromstring(document_data)
        sections = root.xpath(".//w:sectPr", namespaces=NS)
        for index, value in section_types.items():
            if not 0 <= index < len(sections):
                raise FormatMonographError(
                    "Measured section index is outside the DOCX section set."
                )
            section_type = sections[index].find(qn("w:type"))
            if section_type is None:
                section_type = etree.Element(qn("w:type"))
                sections[index].insert(0, section_type)
            section_type.set(qn("w:val"), value)
        spacers = root.xpath(
            ".//w:p[w:pPr/w:pStyle[@w:val='MonographFigureTableSpacer']]",
            namespaces=NS,
        )
        if selected and max(selected) >= len(spacers):
            raise FormatMonographError(
                "Measured spacer ordinal is outside the approved spacer set."
            )
        for ordinal in sorted(selected, reverse=True):
            spacer = spacers[ordinal]
            if spacer.xpath(
                ".//w:t[normalize-space(.) != ''] | .//w:drawing | .//w:object | "
                ".//w:pict | .//w:fldChar | .//w:instrText | .//w:sectPr",
                namespaces=NS,
            ):
                raise FormatMonographError(
                    "Measured page-boundary spacer contains authored or structural payload."
                )
            parent = spacer.getparent()
            if parent is None:
                raise FormatMonographError("Measured spacer has no document parent.")
            parent.remove(spacer)
        patched = etree.tostring(
            root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )
        temp = output_path.with_name(f".{output_path.name}.spacers.tmp")
        temp.unlink(missing_ok=True)
        try:
            with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
                for info in package.infolist():
                    target.writestr(
                        info,
                        patched if info.filename == "word/document.xml" else package.read(info.filename),
                    )
            temp.replace(output_path)
        finally:
            temp.unlink(missing_ok=True)
    if protected_payload_manifest(input_path) != protected_payload_manifest(output_path):
        output_path.unlink(missing_ok=True)
        raise FormatMonographError(
            "Core spacer normalization changed a protected payload."
        )
    return len(selected) + len(section_types)


def remove_measured_block_spacers(
    input_path: Path,
    output_path: Path,
    ordinals: list[int],
) -> int:
    return apply_measured_layout_adjustments(
        input_path,
        output_path,
        ordinals,
    )


def _append_page_offset_formula(paragraph: Any, offset: int) -> None:
    if offset != 1:
        raise FormatMonographError("Only the approved PAGE-minus-one offset is supported.")

    def marker(kind: str) -> Any:
        run = OxmlElement("w:r")
        field = OxmlElement("w:fldChar")
        field.set(qn("w:fldCharType"), kind)
        run.append(field)
        return run

    def instruction(value: str) -> Any:
        run = OxmlElement("w:r")
        node = OxmlElement("w:instrText")
        node.set(qn("xml:space"), "preserve")
        node.text = value
        run.append(node)
        return run

    paragraph._p.append(marker("begin"))
    paragraph._p.append(instruction(" = "))
    paragraph._p.append(marker("begin"))
    paragraph._p.append(instruction(" PAGE "))
    paragraph._p.append(marker("separate"))
    inner_result = OxmlElement("w:r")
    inner_text = OxmlElement("w:t")
    inner_text.text = "2"
    inner_result.append(inner_text)
    paragraph._p.append(inner_result)
    paragraph._p.append(marker("end"))
    paragraph._p.append(instruction(" - 1 "))
    paragraph._p.append(marker("separate"))
    outer_result = OxmlElement("w:r")
    outer_text = OxmlElement("w:t")
    outer_text.text = "1"
    outer_result.append(outer_text)
    paragraph._p.append(outer_result)
    paragraph._p.append(marker("end"))


def _isolate_page_footer(
    document: Any,
    section: Any,
    footer_type: WD_HEADER_FOOTER,
) -> Any:
    section._sectPr.remove_footerReference(footer_type)
    footer_part = FooterPart.new(document.part.package)
    relationship_id = document.part.relate_to(footer_part, RT.FOOTER)
    section._sectPr.add_footerReference(footer_type, relationship_id)
    return (
        section.footer
        if footer_type == WD_HEADER_FOOTER.PRIMARY
        else section.even_page_footer
    )


def _drop_unused_footer_relationships(document: Any) -> None:
    used = {
        reference.get(qn("r:id"))
        for section in document.sections
        for reference in section._sectPr.xpath("./w:footerReference")
    }
    for relationship_id, relationship in list(document.part.rels.items()):
        if relationship.reltype == RT.FOOTER and relationship_id not in used:
            document.part.drop_rel(relationship_id)


def apply_page_display_offsets(
    input_path: Path,
    output_path: Path,
    section_offsets: dict[int, int],
) -> int:
    if any(value != 1 for value in section_offsets.values()):
        raise FormatMonographError("Only a one-page parity offset is supported.")
    document = load_document(input_path)
    changed = 0
    for index, offset in sorted(section_offsets.items()):
        if not 0 <= index < len(document.sections):
            raise FormatMonographError("Page display offset section is out of range.")
        section = document.sections[index]
        for footer_type, alignment in (
            (WD_HEADER_FOOTER.PRIMARY, WD_ALIGN_PARAGRAPH.RIGHT),
            (WD_HEADER_FOOTER.EVEN_PAGE, WD_ALIGN_PARAGRAPH.LEFT),
        ):
            footer = (
                section.footer
                if footer_type == WD_HEADER_FOOTER.PRIMARY
                else section.even_page_footer
            )
            if not _page_only_footer(footer):
                raise FormatMonographError(
                    "Page display offset requires a page-only footer."
                )
            if index + 1 < len(document.sections):
                next_section = document.sections[index + 1]
                if next_section._sectPr.get_footerReference(footer_type) is None:
                    next_footer = (
                        next_section.footer
                        if footer_type == WD_HEADER_FOOTER.PRIMARY
                        else next_section.even_page_footer
                    )
                    if not _page_only_footer(next_footer):
                        raise FormatMonographError(
                            "The following section inherits a non-page footer."
                        )
                    next_footer = _isolate_page_footer(
                        document,
                        next_section,
                        footer_type,
                    )
                    _replace_with_page_field(next_footer, alignment)
                    changed += 1
            footer = _isolate_page_footer(document, section, footer_type)
            paragraphs = list(footer.paragraphs)
            paragraph = paragraphs[0] if paragraphs else footer.add_paragraph()
            for extra in paragraphs[1:]:
                footer._element.remove(extra._p)
            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)
            paragraph.alignment = alignment
            _append_page_offset_formula(paragraph, offset)
            changed += 1
    _drop_unused_footer_relationships(document)
    document.save(output_path)
    if protected_payload_manifest(input_path) != protected_payload_manifest(output_path):
        output_path.unlink(missing_ok=True)
        raise FormatMonographError(
            "Core page display offset changed a protected payload."
        )
    return changed


def external_verify(
    input_path: Path,
    command: str,
    profile_path: Path,
    structure_map_path: Path,
    pdf_output: Path,
    target_software: str,
    *,
    expected_page_count: int | None = None,
    allowed_field_types: set[str] | frozenset[str] = DEFAULT_ALLOWED_FIELD_TYPES,
) -> dict:
    target_software = resolve_target_id(target_software)
    _require_word_target_id(target_software)
    request = {
        "protocol_version": "1.1",
        "operation": "verify_only",
        "input_path": str(input_path.resolve()),
        "profile_path": str(profile_path.resolve()),
        "structure_map_path": str(structure_map_path.resolve()),
        "allowed_field_types": sorted(allowed_field_types),
        "target_software": target_software,
        "pdf_output_path": str(pdf_output.resolve()),
        "expected_page_count": expected_page_count,
    }
    input_hash = file_sha256(input_path)
    completed = _invoke_external_command(command, request, "External verifier")
    if file_sha256(input_path) != input_hash:
        raise FormatMonographError("External verifier changed its input DOCX.")
    if completed.returncode != 0:
        raise FormatMonographError(
            "External read-only verification failed. "
            f"stdout={completed.stdout.strip()} stderr={completed.stderr.strip()}"
        )
    try:
        response = json.loads(completed.stdout)
    except json.JSONDecodeError as exc:
        raise FormatMonographError(
            "External verifier did not return one JSON response."
        ) from exc
    required = {
        "status": "success",
        "operation": "verify_only",
        "repaginated": True,
        "saved": False,
        "read_only_verified": True,
        "pdf_exported": True,
    }
    if not isinstance(response, dict) or any(
        response.get(name) != value for name, value in required.items()
    ):
        raise FormatMonographError(
            "External verifier did not satisfy the read-only verification contract."
        )
    _require_external_target(response, target_software)
    if not pdf_output.is_file() or pdf_output.stat().st_size == 0:
        raise FormatMonographError("External verifier did not create its target PDF.")
    actual_page_count = response.get("page_count")
    if expected_page_count is not None:
        if (
            not isinstance(actual_page_count, int)
            or isinstance(actual_page_count, bool)
            or actual_page_count < 1
        ):
            raise FormatMonographError(
                "External verifier omitted a valid page count."
            )
        if actual_page_count != int(expected_page_count):
            raise FormatMonographError(
                "Selective output page count differs from the field calculation session."
            )
    response.setdefault("backend", "external")
    response["command"] = _external_command(command)[0]
    return response


class _CommittedStdoutSink:
    encoding = "utf-8"
    errors = "strict"

    def write(self, value: str) -> int:
        return len(value)

    def flush(self) -> None:
        pass

    def isatty(self) -> bool:
        return False


def _detach_stdout() -> None:
    """Prevent interpreter-exit flush from changing a committed transaction result."""
    sys.stdout = _CommittedStdoutSink()


def report_committed_result(
    payload: bytes, *, detach_after_success: bool = False
) -> None:
    """Best-effort reporting after durable commit; never roll back the commit."""
    try:
        sys.stdout.write(payload.decode("utf-8"))
        sys.stdout.flush()
    except (BrokenPipeError, OSError, UnicodeError, ValueError, TypeError):
        _detach_stdout()
        return
    if detach_after_success:
        _detach_stdout()


def report_committed_diagnostic(payload: bytes) -> None:
    """Best-effort stderr diagnostic after commit; never reverse success."""
    try:
        sys.stderr.write(payload.decode("utf-8"))
        sys.stderr.flush()
    except (BrokenPipeError, OSError, UnicodeError, ValueError, TypeError):
        return


def main(
    *,
    detach_stdout_after_report: bool = False,
    publication_event_hook: Callable[[str, str | None], None] | None = None,
) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--source", type=Path)
    parser.add_argument("--profile", required=True, type=Path)
    parser.add_argument("--structure-map", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--status-output", type=Path)
    parser.add_argument(
        "--field-updater",
        choices=("auto", "external", "libreoffice", "deferred"),
        default="auto",
    )
    parser.add_argument(
        "--field-updater-command",
        help="External updater command or a JSON array of command arguments.",
    )
    parser.add_argument(
        "--pdf-output",
        type=Path,
        help="Optional target-software PDF path requested from an external backend.",
    )
    parser.add_argument(
        "--target-software",
        help="Target application requested from an external field backend.",
    )
    parser.add_argument("--renderer")
    parser.add_argument("--approve-deferred", action="store_true")
    parser.add_argument("--force", action="store_true")
    args = parser.parse_args()
    staging_root: Path | None = None
    audit_output: Path | None = None
    target_outputs: dict[str, Path] = {}
    staged_outputs: dict[str, Path] = {}
    staged_expected: dict[str, dict[str, Any] | None] = {}
    producer_outputs: dict[str, Path] = {}
    target_snapshots: dict[str, dict[str, Any] | None] = {}
    transaction_id = uuid.uuid4().hex
    output_parent_authority: _DirectoryAuthority | None = None
    staging_authority: _DirectoryAuthority | None = None
    producer_authority: _DirectoryAuthority | None = None

    try:
        path_contract = resolve_finalization_path_contract(
            input_path=args.input,
            source_path=args.source,
            profile_path=args.profile,
            structure_map_path=args.structure_map,
            output_path=args.output,
            pdf_output=args.pdf_output,
            status_output=args.status_output,
            force=args.force,
        )
        resolved_inputs = path_contract["inputs"]
        target_outputs = path_contract["outputs"]
        args.input = resolved_inputs["input"]
        args.source = resolved_inputs.get("source")
        args.profile = resolved_inputs["profile"]
        args.structure_map = resolved_inputs["structure_map"]
        ensure_docx(args.input)
        if args.source:
            ensure_docx(args.source)
        errors, profile = validate(args.profile)
        if errors or profile["approval"]["status"] != "approved":
            raise FormatMonographError(
                "Finalization requires an approved valid profile: " + "; ".join(errors)
            )
        effective_target_software = resolve_target_id(
            args.target_software or profile["target_applications"][0]
        )
        structure_map = load_structure_map(args.structure_map)
        if args.source:
            validate_structure_map_source(args.source, structure_map)

        baseline = args.source or args.input
        baseline_fp = structure_content_fingerprint(baseline, structure_map)
        input_fp = structure_content_fingerprint(args.input, structure_map)
        if baseline_fp != input_fp:
            raise FormatMonographError(
                "Formatted input failed the stable pre-finalization content audit."
            )
        baseline_objects = protected_payload_manifest(baseline)
        approved_toc_contract = toc_result_contract(
            load_document(args.input), structure_map
        )
        input_font_failures = effective_font_failures(
            args.input, profile, structure_map
        )
        if input_font_failures:
            raise FormatMonographError(
                "Formatted input failed deterministic effective-font validation: "
                + json.dumps(input_font_failures, ensure_ascii=False)
            )
        input_fields = field_cache_inventory(args.input)
        backend: dict = {"backend": "not_needed"}
        delivery_status = input_fields["status"]

        publish_parent = path_contract["parent"]
        _assert_publication_platform_available()
        publish_parent.mkdir(parents=True, exist_ok=True)
        output_parent_authority = _open_directory_authority(
            publish_parent, require_private_owner=True
        )
        _assert_absolute_authority(
            output_parent_authority, "output parent before staging creation"
        )
        target_snapshots = {
            name: _target_snapshot(path) for name, path in target_outputs.items()
        }
        staging_authority = _create_bound_directory(
            output_parent_authority,
            prefix=f".format-monograph-finalize-{transaction_id}-",
            label="publication staging",
        )
        staging_root = staging_authority.path
        _assert_publish_chain(
            output_parent_authority,
            staging_authority,
            "immediately after staging creation",
        )
        if publication_event_hook:
            publication_event_hook("after_staging_create", None)
        _assert_publish_chain(
            output_parent_authority,
            staging_authority,
            "after staging creation hook",
        )
        producer_authority = _create_bound_directory(
            staging_authority,
            prefix="producer-",
            label="producer workspace",
        )
        producer_name = producer_authority.path.name
        producer_root = producer_authority.path
        if publication_event_hook:
            publication_event_hook("after_producer_workspace_create", None)
        _assert_authority_entry(
            staging_authority,
            producer_name,
            producer_authority,
            "producer workspace after creation",
            require_private_owner=True,
        )
        if _authority_listdir(producer_authority):
            raise FinalizationPublishError(
                "Producer workspace was not empty before producer execution.",
                preserve_staging=True,
            )
        staged_outputs = {
            name: staging_root / path.name for name, path in target_outputs.items()
        }
        producer_outputs = {
            name: producer_root / path.name for name, path in target_outputs.items()
        }
        args.output = producer_outputs["output"]
        args.pdf_output = producer_outputs.get("pdf")
        args.status_output = producer_outputs.get("status")
        audit_output = producer_outputs.get("audit")
        if publication_event_hook:
            publication_event_hook("before_producer_execution", None)
        _assert_publish_chain(
            output_parent_authority,
            staging_authority,
            "before producer execution",
        )
        _assert_authority_entry(
            staging_authority,
            producer_name,
            producer_authority,
            "producer workspace before execution",
            require_private_owner=True,
        )
        if _authority_listdir(producer_authority):
            raise FinalizationPublishError(
                "Producer workspace gained unknown entries before execution.",
                preserve_staging=True,
            )

        external_requested = args.field_updater == "external" or (
            args.field_updater == "auto" and bool(args.field_updater_command)
        )
        if external_requested:
            if not args.field_updater_command:
                raise FormatMonographError(
                    "External field update requires --field-updater-command."
                )
            try:
                if effective_target_software != MICROSOFT_WORD:
                    raise FormatMonographError(
                        "External Word field update requires target ID "
                        f"{MICROSOFT_WORD!r}; got {effective_target_software!r}."
                    )
                with tempfile.TemporaryDirectory(
                    prefix="format-monograph-external-fields-"
                ) as refresh_name:
                    refresh_root = Path(refresh_name)
                    refresh_input = args.input
                    measurements = []
                    section_adjustments = 0
                    spacers_removed = 0
                    pagination_section_indexes = (
                        approved_front_matter_section_indexes(
                            args.input,
                            structure_map,
                        )
                    )
                    remove_boundary_spacers = bool(
                        structure_map.get("block_spacing", {}).get("approved")
                        and structure_map.get("block_spacing", {}).get(
                            "same_page_only"
                        )
                    )
                    for measurement_pass in range(10):
                        measurement = external_measure(
                            refresh_input,
                            args.field_updater_command,
                            args.profile,
                            args.structure_map,
                            effective_target_software,
                        )
                        measurements.append(measurement)
                        ordinals = (
                            measurement.get("page_boundary_spacer_ordinals", [])
                            if remove_boundary_spacers
                            else []
                        )
                        section_types = approved_front_matter_section_types(
                            pagination_section_indexes,
                            measurement,
                            refresh_input,
                        )
                        if not ordinals and not section_types:
                            break
                        normalized = refresh_root / (
                            f"layout-normalized-{measurement_pass + 1}.docx"
                        )
                        apply_measured_layout_adjustments(
                            refresh_input,
                            normalized,
                            ordinals,
                            section_types,
                        )
                        section_adjustments += len(section_types)
                        spacers_removed += len(ordinals)
                        refresh_input = normalized
                    else:
                        raise FormatMonographError(
                            "Page-boundary spacer normalization did not converge."
                        )
                    display_offsets = {}
                    normalized_document = load_document(refresh_input)
                    for index in sorted(pagination_section_indexes):
                        section_type = normalized_document.sections[
                            index
                        ]._sectPr.find(qn("w:type"))
                        if (
                            section_type is not None
                            and section_type.get(qn("w:val")) == "evenPage"
                        ):
                            display_offsets[index] = 1
                    allowed_field_types = set(DEFAULT_ALLOWED_FIELD_TYPES)
                    if display_offsets:
                        offset_input = refresh_root / "page-display-offsets.docx"
                        apply_page_display_offsets(
                            refresh_input,
                            offset_input,
                            display_offsets,
                        )
                        refresh_input = offset_input
                        allowed_field_types.add("=")
                    writeback_field_types = set(allowed_field_types)
                    if approved_toc_contract is not None:
                        writeback_field_types.add("TC")
                    refreshed = refresh_root / "refreshed.docx"
                    backend = external_refresh(
                        refresh_input,
                        refreshed,
                        args.field_updater_command,
                        args.profile,
                        args.structure_map,
                        None,
                        effective_target_software,
                        allowed_field_types=allowed_field_types,
                    )
                    writeback = selective_field_result_writeback(
                        refresh_input,
                        refreshed,
                        args.output,
                        allowed_field_types=writeback_field_types,
                        toc_contract=approved_toc_contract,
                    )
                    backend["layout_measurements"] = measurements
                    backend["removed_page_boundary_spacers"] = spacers_removed
                    backend["core_section_start_adjustments"] = section_adjustments
                    backend["core_page_display_offsets"] = {
                        str(index): value for index, value in display_offsets.items()
                    }
                    backend["selective_writeback"] = writeback
                    if args.pdf_output is None:
                        raise FormatMonographError(
                            "Target Word verification requires a persistent --pdf-output artifact."
                        )
                    verification_pdf = args.pdf_output
                    backend["read_only_verification"] = external_verify(
                        args.output,
                        args.field_updater_command,
                        args.profile,
                        args.structure_map,
                        verification_pdf,
                        effective_target_software,
                        expected_page_count=backend.get("page_count"),
                        allowed_field_types=allowed_field_types,
                    )
                delivery_status = "selective_verified"
            except FormatMonographError as exc:
                if args.field_updater != "auto" or not args.approve_deferred:
                    raise
                backend = use_deferred_output_with_evidence(
                    args.input,
                    args.output,
                    "external_error",
                    backend,
                    stage="external_field_workflow",
                    error=str(exc),
                    failed_checks=["external_field_workflow"],
                )
                delivery_status = "deferred"
        elif input_fields["status"] in {"absent", "refreshed"}:
            shutil.copy2(args.input, args.output)
        elif args.field_updater == "deferred":
            if not args.approve_deferred:
                raise FormatMonographError(
                    "Deferred field update requires caller QA and --approve-deferred."
                )
            shutil.copy2(args.input, args.output)
            rewrite_field_flags(args.output, deferred=True)
            backend = {"backend": "deferred_on_open"}
            delivery_status = "deferred"
        else:
            with tempfile.TemporaryDirectory(
                prefix="format-monograph-libreoffice-fields-"
            ) as refresh_name:
                refreshed = Path(refresh_name) / "refreshed.docx"
                try:
                    backend = libreoffice_refresh(
                        args.input,
                        refreshed,
                        args.renderer,
                        toc_authorization=toc_index_authorization(
                            approved_toc_contract, args.input
                        ),
                        toc_contract=approved_toc_contract,
                    )
                except FormatMonographError as exc:
                    if args.field_updater != "auto" or not args.approve_deferred:
                        raise
                    backend = use_deferred_output_with_evidence(
                        args.input,
                        args.output,
                        "libreoffice_error",
                        {"backend": "libreoffice_uno"},
                        stage="libreoffice_refresh",
                        error=str(exc),
                        failed_checks=["libreoffice_refresh"],
                    )
                    delivery_status = "deferred"
                else:
                    try:
                        restored = Path(refresh_name) / "instruction-restored.docx"
                        backend["instruction_restoration"] = (
                            restore_known_libreoffice_toc_instruction_order(
                                args.input,
                                refreshed,
                                restored,
                            )
                        )
                        writeback = selective_field_result_writeback(
                            args.input,
                            restored,
                            args.output,
                            allowed_field_types=(
                                set(DEFAULT_ALLOWED_FIELD_TYPES) | {"TC"}
                                if approved_toc_contract is not None
                                else DEFAULT_ALLOWED_FIELD_TYPES
                            ),
                            toc_contract=approved_toc_contract,
                        )
                        if writeback.get("status") != "selective_verified":
                            raise FormatMonographError(
                                "LibreOffice core selective writeback was not verified."
                            )
                        writeback["core_writeback_status"] = writeback["status"]
                        writeback["status"] = LIBREOFFICE_WRITEBACK_STATUS
                        writeback["word_verification"] = "not_performed"
                        backend["selective_writeback"] = writeback
                        if package_field_contract_manifest(
                            args.input
                        ) != package_field_contract_manifest(args.output):
                            args.output.unlink(missing_ok=True)
                            raise FormatMonographError(
                                "LibreOffice delivery field instructions or boundaries differ from the baseline."
                            )
                        backend["delivery_field_contract_identical"] = True
                        backend["completion_scope"] = "libreoffice_non_final"
                        backend["word_verification_required"] = True
                        backend["word_verification_completed"] = False
                        delivery_status = LIBREOFFICE_DELIVERY_STATUS
                    except FormatMonographError as exc:
                        if args.field_updater != "auto" or not args.approve_deferred:
                            raise
                        attempted_backend = copy.deepcopy(backend)
                        if isinstance(exc, LibreOfficeContractError):
                            attempted_backend["contract_comparison"] = exc.evidence
                        attempted_backend["selective_writeback"] = {
                            "status": "rejected",
                            "error": str(exc),
                        }
                        backend = use_deferred_output_with_evidence(
                            args.input,
                            args.output,
                            "libreoffice_contract_or_integrity",
                            attempted_backend,
                            stage="selective_writeback",
                            error=str(exc),
                            failed_checks=["selective_writeback"],
                        )
                        delivery_status = "deferred"

        output_fields = field_cache_inventory(args.output)
        strict_backend = backend.get("backend") not in {
            "not_needed",
            "deferred_on_open",
        }
        if strict_backend:
            writeback_status = backend.get("selective_writeback", {}).get("status")
            if backend.get("backend") == "libreoffice_uno":
                selective_ok = (
                    writeback_status == LIBREOFFICE_WRITEBACK_STATUS
                    and backend.get("selective_writeback", {}).get(
                        "core_writeback_status"
                    )
                    == "selective_verified"
                )
                delivery_status = (
                    LIBREOFFICE_DELIVERY_STATUS
                    if selective_ok
                    else output_fields["status"]
                )
            else:
                selective_ok = writeback_status == "selective_verified"
                delivery_status = (
                    "selective_verified" if selective_ok else output_fields["status"]
                )
            field_contract_ok = field_contract_preserved(input_fields, output_fields)
            refreshed_ok = (
                not input_fields["main_toc_fields"] or delivery_status == "refreshed"
            )
            if backend.get("backend") not in {"libreoffice_uno"}:
                refreshed_ok = bool(
                    backend.get("field_cache_verified")
                    and backend.get("read_only_verification", {}).get(
                        "read_only_verified"
                    )
                    and selective_ok
                    and backend.get("selective_writeback", {}).get(
                        "unapproved_dirty_fields", 0
                    )
                    == 0
                )
            elif selective_ok:
                refreshed_ok = True
        else:
            field_contract_ok = True
            refreshed_ok = True

        output_fp = structure_content_fingerprint(args.output, structure_map)
        output_objects = protected_payload_manifest(args.output)
        output_font_failures = effective_font_failures(
            args.output, profile, structure_map
        )
        content_ok = baseline_fp == output_fp
        objects_ok = baseline_objects == output_objects
        fonts_ok = not output_font_failures
        if strict_backend and not (
            field_contract_ok and refreshed_ok and content_ok and objects_ok and fonts_ok
        ):
            if args.field_updater == "auto" and args.approve_deferred:
                failed_checks = [
                    name
                    for name, passed in (
                        ("field_contract", field_contract_ok),
                        ("field_refresh", refreshed_ok),
                        ("content_integrity", content_ok),
                        ("protected_object_integrity", objects_ok),
                        ("effective_font_integrity", fonts_ok),
                    )
                    if not passed
                ]
                backend = use_deferred_output_with_evidence(
                    args.input,
                    args.output,
                    "libreoffice_contract_or_integrity",
                    backend,
                    stage="post_writeback_integrity",
                    error=(
                        "Field refresh failed post-writeback checks: "
                        + ", ".join(failed_checks)
                    ),
                    failed_checks=failed_checks,
                )
                delivery_status = "deferred"
                output_fields = field_cache_inventory(args.output)
                output_fp = structure_content_fingerprint(args.output, structure_map)
                output_objects = protected_payload_manifest(args.output)
                output_font_failures = effective_font_failures(
                    args.output, profile, structure_map
                )
                content_ok = baseline_fp == output_fp
                objects_ok = baseline_objects == output_objects
                fonts_ok = not output_font_failures
            else:
                args.output.unlink(missing_ok=True)
                raise FormatMonographError(
                    "Field refresh did not preserve the editable-field "
                    "contract and document integrity."
                )
        if not content_ok or not objects_ok or not fonts_ok:
            args.output.unlink(missing_ok=True)
            raise FormatMonographError(
                "Finalization integrity failed "
                f"(content={'pass' if content_ok else 'fail'}, "
                f"protected_objects={'pass' if objects_ok else 'fail'}, "
                f"effective_fonts={'pass' if fonts_ok else 'fail'})."
            )

        audit_payload = backend_audit_bytes(backend)
        canonical_backend = canonical_backend_projection(backend)
        audit_binding = backend_audit_binding(
            target_outputs.get("audit"), audit_payload
        )
        field_gate_completed = delivery_status in FINAL_READY_FIELD_STATES
        word_verification_completed = delivery_status == "selective_verified"
        artifact_binding = {
            "version": 1,
            "finalized_docx": artifact_identity(
                args.output, reported_path=target_outputs["output"]
            ),
            "word_verification_pdf": (
                artifact_identity(
                    args.pdf_output,
                    page_count=(canonical_backend.get("read_only_verification") or {}).get(
                        "page_count"
                    ),
                    reported_path=target_outputs.get("pdf"),
                )
                if delivery_status == "selective_verified"
                and args.pdf_output is not None
                and args.pdf_output.is_file()
                else None
            ),
        }
        result = {
            "finalization_evidence_version": FINALIZATION_EVIDENCE_VERSION,
            "status": "pass",
            "delivery_field_status": delivery_status,
            "input_field_cache": input_fields,
            "output_field_cache": output_fields,
            "field_backend": canonical_backend,
            "backend_audit": audit_binding,
            "artifact_binding": artifact_binding,
            "field_writeback_status": (
                (canonical_backend.get("selective_writeback") or {}).get("status")
                or ("deferred" if delivery_status == "deferred" else "not_needed")
            ),
            "field_completion": {
                "field_gate_completed": field_gate_completed,
                "final_ready_eligible": field_gate_completed,
                "word_verification_required": delivery_status != "absent",
                "word_verification_completed": word_verification_completed,
                "completion_scope": canonical_backend.get("completion_scope")
                or (
                    "target_word_verified"
                    if word_verification_completed
                    else "no_fields"
                    if delivery_status == "absent"
                    else "incomplete"
                ),
            },
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
            "workflow_state": {
                "source_sha256": file_sha256(baseline),
                "input_sha256": file_sha256(args.input),
                "profile_sha256": file_sha256(args.profile),
                "structure_map_sha256": file_sha256(args.structure_map),
                "output_sha256": file_sha256(args.output),
                "stage": "finalized",
            },
            "target_pdf": (
                str(target_outputs["pdf"])
                if args.pdf_output and args.pdf_output.is_file()
                else None
            ),
            "target_layout_status": (
                "target_pdf_ready_for_visual_qa"
                if args.pdf_output and args.pdf_output.is_file()
                else "not_verified"
            ),
            "target_software": effective_target_software,
            "output": str(target_outputs["output"]),
            "publication": {
                "version": PUBLICATION_RECORD_VERSION,
                "transaction_id": transaction_id,
                "retained_staging_directory": str(staging_root.absolute()),
                "cleanup_policy": PUBLICATION_CLEANUP_POLICY,
                "business_gate": False,
            },
        }
        evidence_errors = final_ready_evidence_errors(completion_evidence(result))
        if evidence_errors and result["field_completion"]["final_ready_eligible"]:
            result["field_completion"]["field_gate_completed"] = False
            result["field_completion"]["final_ready_eligible"] = False
            evidence_errors = final_ready_evidence_errors(
                completion_evidence(result)
            )
        result["field_completion"]["evidence_validation"] = {
            "status": "pass" if not evidence_errors else "incomplete",
            "errors": evidence_errors,
        }
        shape_errors = finalization_evidence_shape_errors(result)
        if shape_errors:
            raise FormatMonographError(
                "Finalization produced invalid versioned evidence: "
                + "; ".join(shape_errors)
            )
        status_payload = standard_json_bytes(result)
        trusted_status_identity = trusted_status_byte_identity(status_payload)
        if args.status_output:
            if audit_output is None:
                raise FormatMonographError("Backend audit output path was not resolved.")
            atomic_write_bytes(audit_output, audit_payload)
            atomic_write_bytes(args.status_output, status_payload)
        for name in ("output", "pdf", "audit", "status"):
            if name not in producer_outputs:
                continue
            if publication_event_hook:
                publication_event_hook("before_staged_artifact_import", name)
            _assert_publish_chain(
                output_parent_authority,
                staging_authority,
                f"before staged {name} import",
            )
            _assert_authority_entry(
                staging_authority,
                producer_name,
                producer_authority,
                f"producer workspace before staged {name} import",
                require_private_owner=True,
            )
            staged_location = _ArtifactLocation(
                staging_authority, staged_outputs[name].name
            )
            if _artifact_snapshot_at(staged_location) is not None:
                raise FinalizationPublishError(
                    f"Staged {name} entry already exists; refusing to replace it.",
                    preserve_staging=True,
                )
            producer_location = _ArtifactLocation(
                producer_authority, producer_outputs[name].name
            )
            try:
                producer_snapshot = _artifact_snapshot_at(producer_location)
            except (FormatMonographError, OSError, RuntimeError, TypeError, ValueError) as exc:
                raise FinalizationPublishError(
                    f"Producer {name} artifact cannot be inspected.",
                    preserve_staging=True,
                ) from exc
            if producer_snapshot is None:
                if name == "pdf":
                    staged_expected[name] = None
                    continue
                raise FinalizationPublishError(
                    f"Producer {name} artifact is missing.",
                    preserve_staging=True,
                )
            imported = _copy_authority_regular_exclusive(
                producer_location,
                staged_location,
            )
            if (
                imported.get("sha256") != producer_snapshot.get("sha256")
                or imported.get("size") != producer_snapshot.get("size")
            ):
                raise FinalizationPublishError(
                    f"Staged {name} bytes differ from its producer identity.",
                    preserve_staging=True,
                )
            staged_expected[name] = imported
            if publication_event_hook:
                publication_event_hook("after_staged_artifact_import", name)
        if publication_event_hook:
            publication_event_hook("before_publisher", None)
        publication_result = publish_staged_artifacts(
            staged_outputs,
            target_outputs,
            target_snapshots,
            staging_root,
            trusted_status_identity=(
                trusted_status_identity if args.status_output else None
            ),
            transaction_id=transaction_id,
            event_hook=publication_event_hook,
            output_parent_authority=output_parent_authority,
            staging_authority=staging_authority,
            retained_staging_authorities=[producer_authority],
            staged_expected=staged_expected,
        )
        output_parent_authority = None
        staging_authority = None
        producer_authority = None
        staging_root = None
        if publication_result.get("cleanup_errors"):
            report_committed_diagnostic(
                standard_json_bytes(
                    {
                        "publication_commit_state": publication_result[
                            "commit_state"
                        ],
                        "cleanup_errors": publication_result["cleanup_errors"],
                    }
                )
            )
        report_committed_result(
            status_payload, detach_after_success=detach_stdout_after_report
        )
        return 0
    except (
        FormatMonographError,
        OSError,
        ValueError,
        subprocess.SubprocessError,
        zipfile.BadZipFile,
    ) as exc:
        close_errors = _close_authorities_collect(
            [
                ("producer", producer_authority),
                ("staging", staging_authority),
                ("output parent", output_parent_authority),
            ]
        )
        producer_authority = None
        staging_authority = None
        output_parent_authority = None
        detail = str(exc)
        if close_errors:
            detail += "; close diagnostics: " + "; ".join(close_errors)
        if staging_root is not None:
            detail += (
                f"; publication transaction {transaction_id}; retained staging "
                f"at {staging_root.absolute()}; cleanup_policy="
                f"{PUBLICATION_CLEANUP_POLICY}"
            )
        print(detail, file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main(detach_stdout_after_report=True))
