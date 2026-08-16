"""Candidate generation, validation, application, and audit normalization."""

from __future__ import annotations

import copy
import hashlib
import json
import posixpath
import re
import zipfile
from pathlib import Path
from typing import Any

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.image.image import Image as DocxImage
from docx.image.exceptions import UnrecognizedImageError
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree

from _common import (
    CONTENT_PART,
    FIELD_MARKER_PATTERN,
    NS,
    STRUCTURAL_INDENT_ATTRIBUTES,
    FormatMonographError,
    _numbering_level_for_style,
    _heading_prefix_pattern,
    _paragraph_text_without_field_results,
    _unique_row_cells,
    apply_table_properties,
    apply_style_properties,
    clear_controlled_direct_format,
    content_fingerprint,
    ensure_paragraph_style,
    normalize_structural_paragraph,
)
from docx_pagination import apply_pagination_sections, section_index_for_paragraph


HEADING_PATTERNS = (
    (4, re.compile(r"^\s*\d+\.\d+\.\d+\.\d+\.?\s*\S")),
    (3, re.compile(r"^\s*\d+\.\d+\.\d+\.?\s*\S")),
    (2, re.compile(r"^\s*\d+\.\d+\.?\s*\S")),
    (1, re.compile(r"^\s*第\s*[0-9一二三四五六七八九十百]+\s*章\s*\S")),
)
CAPTION_PATTERN = re.compile(
    r"^\s*(图|表)\s*(\d+(?:\.\d+){1,3})\s*[-－—–]\s*(\d+)\s*(\S.*)$"
)
LOOSE_CAPTION_PATTERN = re.compile(r"^\s*(图|表)\s*(\S.*)?$")
CAPTION_IDENTIFIER_PATTERN = re.compile(
    r"^\s*(图|表)\s*([0-9]+(?:[.．-][0-9]+)*)"
    r"(?:\s+|\s*[-－—–]\s*(?=\D))(\S.*)$"
)
DRAWING_MARK_PATTERN = re.compile(
    r"(?:^|\s)\d+(?:-\d+)+\s*(?:剖面|断面|截面|节点|详图|大样|立面|平面)"
)
BOOK_TITLE_STYLE = "Monograph Book Title"
TOC_HEADING_STYLE = "Monograph TOC Heading"
BLOCK_SPACER_STYLE = "Monograph Figure Table Spacer"
DEFAULT_BOOK_TITLE_FORMAT = {
    "font_name_east_asia": "黑体",
    "font_name_ascii": "Times New Roman",
    "font_name_complex_script": "Times New Roman",
    "font_size_pt": 22,
    "bold": True,
    "alignment": "center",
    "first_line_indent_chars": 0,
    "left_indent_pt": 0,
    "right_indent_pt": 0,
    "line_spacing_rule": "at_least",
    "line_spacing_pt": 33,
    "space_before_pt": 0,
    "space_after_pt": 0,
}
EMU_PER_INCH = 914400
EMU_PER_MM = 36000
ARCHITECTURE_TERMS = ("建筑", "平面图", "立面图", "剖面图", "详图", "大样")
CIVIL_ENGINEERING_TERMS = (
    "土木",
    "结构",
    "钢结构",
    "混凝土",
    "梁",
    "柱",
    "桁架",
    "基础",
    "节点",
    "截面",
    "断面",
)
STRUCTURE_MAP_VERSIONS = {"1.0", "1.1", "1.2", "1.3", "1.4"}
STRUCTURE_MAP_VERSIONS.add("1.5")
SEMANTIC_STRUCTURE_MAP_VERSIONS = {"1.1", "1.2", "1.3", "1.4", "1.5"}
CAPTION_ACTIONS = {
    "preserve",
    "style_only",
    "replace_identifier",
    "convert_to_seq",
    "move_caption",
}

ROLE_ALIASES = {
    "body_text": "body",
    "heading_1": "chapter_title",
    "heading1": "chapter_title",
    "heading_2": "level_2_section",
    "heading2": "level_2_section",
    "heading_3": "level_3_section",
    "heading3": "level_3_section",
    "heading_4": "level_4_section",
    "heading4": "level_4_section",
    "appendix_heading": "chapter_title",
    "caption": "all_captions",
    "all": "all_captions",
    "table_of_contents_level_1": "toc_level_1",
    "table_of_contents_level_2": "toc_level_2",
    "table_of_contents_level_3": "toc_level_3",
}

ROLE_STYLE_NAMES = {
    "body": "Normal",
    "title": "Title",
    "heading_1": "Heading 1",
    "heading_2": "Heading 2",
    "heading_3": "Heading 3",
    "heading_4": "Heading 4",
    "chapter_title": "Heading 1",
    "level_2_section": "Heading 2",
    "level_3_section": "Heading 3",
    "level_4_section": "Heading 4",
    "appendix_heading": "Heading 1",
    "figure_caption": "Caption",
    "figure_caption_unnumbered": "Caption",
    "figure_panel_label": "Caption",
    "table_caption": "Caption",
    "equation_caption": "Caption",
    "long_quote": "Quote",
    "reference_entry": "Bibliography",
    "answer": "Normal",
    "teaching_callout": "Normal",
    "toc_level_1": "TOC 1",
    "toc_level_2": "TOC 2",
    "toc_level_3": "TOC 3",
}

APPENDIX_PATTERN = re.compile(
    r"^\s*(?:附\s*录|APPENDIX)\s*[A-Z0-9一二三四五六七八九十百]*",
    re.IGNORECASE,
)


def text_sha256(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _heading_authored_text_hash(value: str, level: int) -> str:
    match = _heading_prefix_pattern(level).match(value)
    return text_sha256(value[match.end() :] if match else value)


def has_semantic_structure_map(structure_map: dict[str, Any]) -> bool:
    return structure_map.get("schema_version") in SEMANTIC_STRUCTURE_MAP_VERSIONS


def has_caption_actions_map(structure_map: dict[str, Any]) -> bool:
    return structure_map.get("schema_version") in {"1.2", "1.3", "1.4", "1.5"}


def _caption_domain(value: str, context: str = "") -> tuple[str, str]:
    evidence = f"{value}\n{context}"
    architecture = any(term in evidence for term in ARCHITECTURE_TERMS)
    civil = any(term in evidence for term in CIVIL_ENGINEERING_TERMS)
    if architecture and civil:
        return "mixed", "high"
    if architecture:
        return "architecture", "high"
    if civil:
        return "civil_engineering", "high"
    return "unknown", "low"


def _caption_identifier_semantics(
    identifier: str | None, title: str | None, domain_context: str
) -> str:
    if not identifier:
        return "unknown"
    drawing_evidence = bool(DRAWING_MARK_PATTERN.search(title or ""))
    if drawing_evidence:
        return "mixed" if "." in identifier or "．" in identifier else "drawing_mark"
    if domain_context != "unknown" and "-" in identifier and not any(
        marker in identifier for marker in (".", "．")
    ):
        return "drawing_mark"
    if any(marker in identifier for marker in (".", "．")):
        return "publication_number"
    return "unknown"


def _paragraph_has_payload(element: Any) -> bool:
    text = "".join(element.xpath(".//w:t/text()")).strip()
    if text:
        return True
    protected = (
        ".//w:tbl | .//w:drawing | .//w:object | .//w:pict | "
        ".//w:bookmarkStart | .//w:commentRangeStart | .//w:footnoteReference | "
        ".//w:endnoteReference | .//w:fldChar"
    )
    return bool(element.xpath(protected))


def _paragraph_style_signature(paragraph: Any) -> str:
    p_pr = paragraph._p.pPr
    r_pr = None if p_pr is None else p_pr.find(qn("w:rPr"))
    source = r_pr if r_pr is not None else (p_pr if p_pr is not None else paragraph._p)
    payload = etree.tostring(source)
    return hashlib.sha256(payload).hexdigest()


def _nearest_nonempty_hash(values: list[str], index: int, step: int) -> str | None:
    cursor = index + step
    while 0 <= cursor < len(values):
        if values[cursor].strip():
            return text_sha256(values[cursor])
        cursor += step
    return None


def _body_locator(
    index: int, values: list[str] | None = None
) -> dict[str, Any]:
    locator: dict[str, Any] = {"kind": "body_paragraph", "paragraph": index}
    if values is not None:
        locator.update(
            {
                "text_sha256": text_sha256(values[index]),
                "previous_nonempty_sha256": _nearest_nonempty_hash(values, index, -1),
                "next_nonempty_sha256": _nearest_nonempty_hash(values, index, 1),
            }
        )
    return locator


def _cell_locator(table: int, row: int, cell: int, paragraph: int) -> dict[str, Any]:
    return {
        "kind": "table_cell_paragraph",
        "table": table,
        "row": row,
        "cell": cell,
        "paragraph": paragraph,
    }


def _locator_key(locator: dict[str, Any]) -> tuple[Any, ...]:
    if locator.get("kind") == "body_paragraph":
        return ("body", int(locator["paragraph"]))
    if locator.get("kind") == "table_cell_paragraph":
        return (
            "cell",
            int(locator["table"]),
            int(locator["row"]),
            int(locator["cell"]),
            int(locator["paragraph"]),
        )
    raise FormatMonographError(f"Unsupported paragraph locator: {locator}")


def resolve_paragraph_locator(document: Any, locator: dict[str, Any]) -> Any:
    cached = getattr(document, "_format_monograph_locator_cache", {}).get(
        _locator_key(locator)
    )
    if cached is not None:
        return cached
    kind = locator.get("kind")
    if kind == "body_paragraph":
        index = int(locator["paragraph"])
        expected = locator.get("text_sha256")
        if 0 <= index < len(document.paragraphs):
            paragraph = document.paragraphs[index]
            if expected is None or text_sha256(paragraph.text) == expected:
                return paragraph
        if expected is None:
            raise FormatMonographError(f"Body paragraph locator is out of range: {index}")

        values = [paragraph.text for paragraph in document.paragraphs]
        candidates = [
            candidate
            for candidate, value in enumerate(values)
            if text_sha256(value) == expected
        ]
        before = locator.get("previous_nonempty_sha256")
        after = locator.get("next_nonempty_sha256")

        def context_score(candidate: int) -> int:
            return int(
                before is not None
                and _nearest_nonempty_hash(values, candidate, -1) == before
            ) + int(
                after is not None
                and _nearest_nonempty_hash(values, candidate, 1) == after
            )

        if len(candidates) > 1:
            scored = [(context_score(candidate), candidate) for candidate in candidates]
            best = max(score for score, _ in scored)
            candidates = [candidate for score, candidate in scored if score == best]
        if len(candidates) == 1:
            return document.paragraphs[candidates[0]]

        # A caller-approved identifier replacement changes the paragraph hash. Its
        # unchanged neighboring paragraphs still provide a stable, text-free anchor.
        contextual = [
            candidate
            for candidate in range(len(values))
            if context_score(candidate) == int(before is not None) + int(after is not None)
        ]
        if len(contextual) == 1:
            return document.paragraphs[contextual[0]]
        raise FormatMonographError(
            f"Stable body paragraph locator is ambiguous or missing: {index}"
        )
    if kind == "table_cell_paragraph":
        table_index = int(locator["table"])
        row_index = int(locator["row"])
        cell_index = int(locator["cell"])
        paragraph_index = int(locator["paragraph"])
        if not 0 <= table_index < len(document.tables):
            raise FormatMonographError(f"Table locator is out of range: {table_index}")
        table = document.tables[table_index]
        if not 0 <= row_index < len(table.rows):
            raise FormatMonographError(f"Table row locator is out of range: {row_index}")
        row = table.rows[row_index]
        if not 0 <= cell_index < len(row.cells):
            raise FormatMonographError(f"Table cell locator is out of range: {cell_index}")
        paragraphs = row.cells[cell_index].paragraphs
        if not 0 <= paragraph_index < len(paragraphs):
            raise FormatMonographError(
                f"Table-cell paragraph locator is out of range: {paragraph_index}"
            )
        return paragraphs[paragraph_index]
    raise FormatMonographError(f"Unsupported paragraph locator kind: {kind}")


def _verified_locator_paragraph(document: Any, entry: dict[str, Any]) -> Any:
    locator = copy.deepcopy(entry["locator"])
    if locator.get("kind") == "body_paragraph" and entry.get("text_sha256"):
        locator.setdefault("text_sha256", entry["text_sha256"])
    key = _locator_key(locator)
    verified = getattr(document, "_format_monograph_verified_locators", set())
    if key in verified:
        return resolve_paragraph_locator(document, locator)
    try:
        paragraph = resolve_paragraph_locator(document, locator)
    except FormatMonographError:
        normalized = entry.get("normalized_text_sha256")
        candidates = (
            [
                candidate
                for candidate in document.paragraphs
                if text_sha256(candidate.text) == normalized
            ]
            if normalized and locator.get("kind") == "body_paragraph"
            else []
        )
        if len(candidates) != 1:
            raise
        paragraph = candidates[0]
        cache = getattr(document, "_format_monograph_locator_cache", {})
        cache[key] = paragraph
        setattr(document, "_format_monograph_locator_cache", cache)
    expected = entry.get("text_sha256")
    actual = text_sha256(paragraph.text)
    if expected and actual != expected and actual != entry.get(
        "normalized_text_sha256"
    ):
        raise FormatMonographError(
            f"Structure-map paragraph hash mismatch at {key}."
        )
    return paragraph


def _replacement_target_matches(value: str, entry: dict[str, Any]) -> bool:
    span = entry["identifier_span"]
    start, end = int(span["start"]), int(span["end"])
    replacement = str(entry["replacement_identifier"])
    replacement_end = start + len(replacement)
    title_start = int(entry["title_span_start"]) + replacement_end - end
    return bool(
        value[start:replacement_end] == replacement
        and text_sha256(value[:start]) == entry["identifier_prefix_sha256"]
        and text_sha256(value[replacement_end:]) == entry["identifier_suffix_sha256"]
        and text_sha256(value[title_start:]) == entry["title_text_sha256"]
    )


def _resolve_replaced_caption_target(document: Any, entry: dict[str, Any]) -> Any:
    candidates = [
        paragraph
        for paragraph in document.paragraphs
        if _replacement_target_matches(paragraph.text, entry)
    ]
    if len(candidates) != 1:
        raise FormatMonographError(
            "Approved caption replacement target is ambiguous or missing."
        )
    return candidates[0]


def prime_structure_map_locators(document: Any, structure_map: dict[str, Any]) -> None:
    if not has_semantic_structure_map(structure_map):
        return
    cache = getattr(document, "_format_monograph_locator_cache", {})
    verified = getattr(document, "_format_monograph_verified_locators", set())
    for entry in structure_map.get("paragraph_roles", []):
        if not entry.get("approved"):
            continue
        paragraph = resolve_paragraph_locator(document, entry["locator"])
        if text_sha256(paragraph.text) != entry["text_sha256"]:
            raise FormatMonographError(
                f"Structure-map paragraph hash mismatch at {_locator_key(entry['locator'])}."
            )
        key = _locator_key(entry["locator"])
        cache[key] = paragraph
        verified.add(key)
    if structure_map.get("schema_version") in {"1.3", "1.4", "1.5"}:
        for group in structure_map.get("pagination_groups", []):
            if not group.get("approved"):
                continue
            for name in ("anchor", "caption"):
                locator = group.get(name)
                if not isinstance(locator, dict):
                    continue
                paragraph = resolve_paragraph_locator(document, locator)
                expected = locator.get("text_sha256")
                if expected and text_sha256(paragraph.text) != expected:
                    raise FormatMonographError(
                        f"Pagination locator hash mismatch at {_locator_key(locator)}."
                    )
                key = _locator_key(locator)
                cache[key] = paragraph
                verified.add(key)
    if structure_map.get("schema_version") in {"1.4", "1.5"}:
        front_matter = structure_map.get("front_matter", {})
        if front_matter.get("approved"):
            locator = front_matter.get("book_title")
            if isinstance(locator, dict):
                paragraph = resolve_paragraph_locator(document, locator)
                expected = locator.get("text_sha256")
                if expected and text_sha256(paragraph.text) != expected:
                    raise FormatMonographError(
                        f"Book-title locator hash mismatch at {_locator_key(locator)}."
                    )
                key = _locator_key(locator)
                cache[key] = paragraph
                verified.add(key)
        for image in structure_map.get("images", []):
            if not image.get("approved"):
                continue
            locator = image.get("locator")
            if not isinstance(locator, dict):
                continue
            paragraph = resolve_paragraph_locator(document, locator)
            key = _locator_key(locator)
            cache[key] = paragraph
            verified.add(key)
        pagination = structure_map.get("pagination_sections", {})
        if pagination.get("approved"):
            for name in ("toc_start", "body_start"):
                locator = pagination.get(name)
                if not isinstance(locator, dict):
                    continue
                paragraph = resolve_paragraph_locator(document, locator)
                expected = locator.get("text_sha256")
                if expected and text_sha256(paragraph.text) != expected:
                    raise FormatMonographError(
                        f"Pagination-section locator hash mismatch at {_locator_key(locator)}."
                    )
                key = _locator_key(locator)
                cache[key] = paragraph
                verified.add(key)
        for entry in structure_map.get("trailing_empty_sections", []):
            if not entry.get("approved_delete"):
                continue
            locator = entry.get("previous_boundary_locator")
            if not isinstance(locator, dict):
                continue
            paragraph = resolve_paragraph_locator(document, locator)
            key = _locator_key(locator)
            cache[key] = paragraph
            verified.add(key)
    setattr(document, "_format_monograph_locator_cache", cache)
    setattr(document, "_format_monograph_verified_locators", verified)


def normalized_role(value: str) -> str:
    return ROLE_ALIASES.get(value, value)


def approved_role_paragraphs(
    document: Any, structure_map: dict[str, Any], selector: dict[str, str]
) -> list[Any]:
    if not has_semantic_structure_map(structure_map):
        return []
    wanted = normalized_role(selector["value"])
    caption_roles = {
        "figure_caption",
        "figure_caption_unnumbered",
        "figure_panel_label",
        "table_caption",
        "equation_caption",
    }
    result = []
    seen: set[int] = set()
    for entry in structure_map.get("paragraph_roles", []):
        if not entry.get("approved"):
            continue
        locator = entry.get("locator", {})
        if locator.get("kind") == "body_paragraph" and any(
            toc.get("approved")
            and int(toc["start_paragraph"])
            <= int(locator.get("paragraph", -1))
            <= int(toc["end_paragraph"])
            for toc in structure_map.get("toc_ranges", [])
        ):
            continue
        role_value = (
            entry.get("canonical_role")
            if structure_map.get("schema_version") == "1.5"
            else entry.get("role")
        )
        role = normalized_role(str(role_value or entry.get("role", "unknown")))
        matches = role == wanted or (wanted == "all_captions" and role in caption_roles)
        if not matches:
            continue
        try:
            paragraph = _verified_locator_paragraph(document, entry)
        except FormatMonographError:
            locator = entry.get("locator", {})
            expected_style = ROLE_STYLE_NAMES.get(role)
            if role.startswith("heading_") and locator.get("kind") == "body_paragraph":
                normalized_hash = entry.get("normalized_text_sha256")
                candidates = [
                    paragraph
                    for paragraph in document.paragraphs
                    if paragraph.style
                    and paragraph.style.name == expected_style
                    and (
                        normalized_hash is None
                        or text_sha256(paragraph.text) == normalized_hash
                    )
                ]
                if len(candidates) != 1:
                    raise
                paragraph = candidates[0]
            elif role in caption_roles:
                caption_entry = next(
                    (
                        item
                        for item in structure_map.get("captions", [])
                        if item.get("approved")
                        and item.get("locator", {}).get("kind")
                        in {"body_paragraph", "table_cell_paragraph"}
                        and _locator_key(item["locator"])
                        == _locator_key(locator)
                    ),
                    None,
                )
                if caption_entry is None:
                    raise
                if (
                    has_caption_actions_map(structure_map)
                    and caption_entry.get("action")
                    not in {"replace_identifier", "convert_to_seq", "move_caption"}
                ):
                    raise
                if caption_entry.get("action") == "replace_identifier":
                    paragraph = _resolve_replaced_caption_target(
                        document, caption_entry
                    )
                elif (
                    locator.get("kind") == "table_cell_paragraph"
                    and caption_entry.get("migrate_outside_table")
                ):
                    table_index = int(locator["table"])
                    if not 0 <= table_index < len(document.tables):
                        raise
                    previous = document.tables[table_index]._tbl.getprevious()
                    if previous is None or previous.tag != qn("w:p"):
                        raise
                    paragraph = Paragraph(previous, document.tables[table_index]._parent)
                    if caption_entry.get("action") == "move_caption":
                        if text_sha256(paragraph.text) != caption_entry.get(
                            "text_sha256"
                        ):
                            raise FormatMonographError(
                                "Moved caption target content does not match approval."
                            )
                    elif caption_entry.get("action") == "convert_to_seq":
                        instructions = " ".join(
                            paragraph._p.xpath(
                                ".//w:fldSimple/@w:instr | .//w:instrText/text()"
                            )
                        ).upper()
                        if "STYLEREF " not in instructions or "SEQ " not in instructions:
                            raise FormatMonographError(
                                "Converted caption target fields do not match approval."
                            )
                else:
                    paragraph = resolve_paragraph_locator(document, locator)
            else:
                raise
        if id(paragraph._p) not in seen:
            result.append(paragraph)
            seen.add(id(paragraph._p))
    return result


def approved_data_tables(
    document: Any, structure_map: dict[str, Any]
) -> list[tuple[Any, dict[str, Any]]]:
    if not has_semantic_structure_map(structure_map):
        return []
    result = []
    for entry in structure_map.get("tables", []):
        if not entry.get("approved") or entry.get("kind") != "data":
            continue
        index = int(entry["table"])
        if not 0 <= index < len(document.tables):
            raise FormatMonographError(f"Structure-map table is out of range: {index}")
        normalized = copy.deepcopy(entry)
        deleted = getattr(document, "_format_monograph_deleted_table_rows", {}).get(index)
        if deleted is None:
            deleted = next(
                (
                    int(caption["locator"]["row"])
                    for caption in structure_map.get("captions", [])
                    if caption.get("approved")
                    and caption.get("migrate_outside_table")
                    and caption.get("locator", {}).get("kind")
                    == "table_cell_paragraph"
                    and int(caption["locator"]["table"]) == index
                ),
                None,
            )
        if deleted is not None:
            normalized["repeat_header_rows"] = [
                int(row) - 1 if int(row) > deleted else int(row)
                for row in normalized.get("repeat_header_rows", [])
                if int(row) != deleted
            ]
            normalized["header_rows"] = [
                int(row) - 1 if int(row) > deleted else int(row)
                for row in normalized.get("header_rows", [])
                if int(row) != deleted
            ]
            normalized["caption_row"] = None
        result.append((document.tables[index], normalized))
    return result


def _section_evidence(document: Any, sect_pr: Any) -> dict[str, Any]:
    header_ids = sect_pr.xpath("./w:headerReference/@r:id")
    footer_ids = sect_pr.xpath("./w:footerReference/@r:id")

    def related_has_payload(rel_id: str) -> bool:
        part = document.part.related_parts.get(rel_id)
        if part is None:
            return True
        try:
            root = etree.fromstring(part.blob)
        except (ValueError, etree.XMLSyntaxError):
            return True
        return bool(
            "".join(root.xpath(".//w:t/text()", namespaces=NS)).strip()
            or root.xpath(".//w:drawing | .//w:object | .//w:pict", namespaces=NS)
        )

    independent = {
        "header_reference_count": len(header_ids),
        "footer_reference_count": len(footer_ids),
        "header_footer_has_payload": any(
            related_has_payload(rel_id) for rel_id in header_ids + footer_ids
        ),
        "page_number_start": (
            sect_pr.xpath("string(./w:pgNumType/@w:start)") or None
        ),
        "different_first_page": bool(sect_pr.find(qn("w:titlePg")) is not None),
        "section_type": (
            sect_pr.xpath("string(./w:type/@w:val)") or None
        ),
        "has_page_geometry": bool(
            sect_pr.find(qn("w:pgSz")) is not None
            or sect_pr.find(qn("w:pgMar")) is not None
        ),
    }
    independent["empty_header_footer_references"] = bool(
        independent["header_reference_count"] or independent["footer_reference_count"]
    ) and not independent["header_footer_has_payload"]
    independent["safe_to_delete"] = not any(
        (
            independent["header_footer_has_payload"],
            independent["page_number_start"],
            independent["different_first_page"],
        )
    )
    return independent


def _section_properties_sha256(sect_pr: Any) -> str:
    canonical = etree.tostring(sect_pr, method="c14n", exclusive=True)
    return hashlib.sha256(canonical).hexdigest()


def _trailing_empty_sections(document: Any) -> list[dict[str, Any]]:
    body = document.element.body
    children = list(body)
    paragraph_index = {id(p._p): index for index, p in enumerate(document.paragraphs)}
    boundaries: list[tuple[int, int]] = []
    for child_index, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        p_pr = child.find(qn("w:pPr"))
        if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
            boundaries.append((child_index, paragraph_index[id(child)]))

    candidates = []
    starts = [0] + [child_index + 1 for child_index, _ in boundaries]
    ends = [child_index for child_index, _ in boundaries] + [len(children) - 2]
    for section_index in range(len(starts) - 1, 0, -1):
        start, end = starts[section_index], ends[section_index]
        span = children[start : end + 1] if end >= start else []
        if any(_paragraph_has_payload(element) for element in span):
            break
        previous_boundary = boundaries[section_index - 1][1]
        paragraph = document.paragraphs[previous_boundary]
        candidate_sect_pr = list(document.sections)[section_index]._sectPr
        evidence = _section_evidence(document, candidate_sect_pr)
        candidates.append(
            {
                "section": section_index,
                "previous_boundary_paragraph": previous_boundary,
                "previous_boundary_sha256": text_sha256(paragraph.text),
                "previous_boundary_locator": _body_locator(
                    previous_boundary, [item.text for item in document.paragraphs]
                ),
                "section_properties_sha256": _section_properties_sha256(candidate_sect_pr),
                "approved_delete": False,
                "confidence": "high" if evidence["safe_to_delete"] else "low",
                "evidence": evidence,
            }
        )
    return candidates


def _candidate_pagination_sections(
    document: Any, headings: list[dict[str, Any]]
) -> dict[str, Any]:
    body_start = next(
        (entry["locator"] for entry in headings if int(entry["level"]) == 1), None
    )
    toc_start = None
    for index, paragraph in enumerate(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        instructions = " ".join(
            paragraph._p.xpath(".//w:fldSimple/@w:instr | .//w:instrText/text()")
        ).upper()
        if "[[TOC]]" in paragraph.text or "TOC" in instructions or style_name.upper().startswith("TOC"):
            toc_start = _body_locator(
                index, [item.text for item in document.paragraphs]
            )
            break
    return {
        "approved": False,
        "toc_start": toc_start,
        "body_start": body_start,
        "number_format": "decimal",
        "start_at": {"toc": 1, "body": 1},
        "continue_after_body_start": True,
        "odd_position": "outer_right",
        "even_position": "outer_left",
        "show_on_first_page": True,
    }


def _candidate_front_matter(
    document: Any, pagination: dict[str, Any], body_values: list[str]
) -> dict[str, Any]:
    toc_locator = pagination.get("toc_start")
    toc_index = (
        int(toc_locator["paragraph"])
        if isinstance(toc_locator, dict) and "paragraph" in toc_locator
        else None
    )
    title_index = next(
        (
            index
            for index, value in enumerate(body_values)
            if value.strip() and (toc_index is None or index < toc_index)
        ),
        None,
    )
    return {
        "approved": False,
        "book_title": (
            _body_locator(title_index, body_values) if title_index is not None else None
        ),
        "separate_title_page": True,
        "title_page_numbering": "none",
        "title_page_vertical_alignment": "center",
        "book_title_format": dict(DEFAULT_BOOK_TITLE_FORMAT),
        "toc_heading_text": "目    录",
        "insert_toc_heading_if_missing": True,
    }


def _candidate_block_spacing() -> dict[str, Any]:
    return {
        "approved": False,
        "mode": "actual_blank_paragraph",
        "blank_lines": 1,
        "same_page_only": True,
        "after": ["approved_data_table", "approved_figure_caption"],
    }


def _chapter_number(value: str) -> int | None:
    match = re.match(r"^\s*第\s*([0-9一二三四五六七八九十百]+)\s*章", value)
    if not match:
        return None
    token = match.group(1)
    if token.isdigit():
        return int(token)
    digits = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
    if token == "十":
        return 10
    if "十" in token:
        left, right = token.split("十", 1)
        return (digits.get(left, 1) * 10) + digits.get(right, 0)
    return digits.get(token)


def _heading_number(value: str, level: int) -> tuple[int, ...] | None:
    if level == 1:
        chapter = _chapter_number(value)
        return None if chapter is None else (chapter,)
    match = re.match(r"^\s*(\d+(?:[.-]\d+){%d})" % (level - 1), value)
    if not match:
        return None
    return tuple(int(part) for part in re.split(r"[.-]", match.group(1)))


def _numbering_anomalies(headings: list[dict[str, Any]]) -> list[dict[str, Any]]:
    anomalies = []
    seen: set[tuple[int, ...]] = set()
    last_by_parent: dict[tuple[int, ...], int] = {}
    for entry in headings:
        number = tuple(entry.get("cached_number", []))
        if not number:
            continue
        locator = entry.get("locator", _body_locator(int(entry["paragraph"])))
        if number in seen:
            anomalies.append(
                {
                    "kind": "duplicate_number",
                    "locator": locator,
                    "observed": list(number),
                    "status": "open",
                }
            )
        parent = number[:-1]
        if len(number) > 1 and parent not in seen:
            anomalies.append(
                {
                    "kind": "missing_parent",
                    "locator": locator,
                    "observed": list(number),
                    "status": "open",
                }
            )
        previous = last_by_parent.get(parent)
        expected_last = previous + 1 if previous is not None else 1
        if len(number) > 1 and number[-1] != expected_last:
            anomalies.append(
                {
                    "kind": "sequence_gap",
                    "locator": locator,
                    "observed": list(number),
                    "expected_last": expected_last,
                    "status": "open",
                }
            )
        seen.add(number)
        last_by_parent[parent] = number[-1]
    return anomalies


def _role_for_paragraph(paragraph: Any, detected_level: int | None) -> str:
    if detected_level:
        return f"heading_{detected_level}"
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name == "Title":
        return "title"
    match = re.fullmatch(r"Heading (\d+)", style_name)
    if match and 1 <= int(match.group(1)) <= 4:
        return f"heading_{match.group(1)}"
    if style_name == "Quote":
        return "long_quote"
    if style_name == "Bibliography":
        return "reference_entry"
    return "unknown"


def _caption_entry(
    value: str, locator: dict[str, Any], context: str = ""
) -> dict[str, Any] | None:
    loose = LOOSE_CAPTION_PATTERN.match(value)
    if not loose:
        return None

    label = loose.group(1)
    domain_context, domain_confidence = _caption_domain(value, context)
    identifier = CAPTION_IDENTIFIER_PATTERN.match(value)
    entry: dict[str, Any] = {
        "locator": locator,
        "text_sha256": text_sha256(value),
        "label": label,
        "sequence_name": "Figure" if label == "图" else "Table",
        "numbering_mode": "manual_text",
        "identifier_semantics": "unknown",
        "domain_context": domain_context,
        "domain_confidence": domain_confidence,
        "action": "preserve",
        "completeness": "candidate",
        "hierarchy_status": "not_evaluated",
        "approved": False,
    }
    if identifier:
        raw_identifier = identifier.group(2)
        start, end = identifier.span(2)
        title = identifier.group(3)
        entry.update(
            {
                "identifier_span": {"start": start, "end": end},
                "identifier_sha256": text_sha256(raw_identifier),
                "identifier_prefix_sha256": text_sha256(value[:start]),
                "identifier_suffix_sha256": text_sha256(value[end:]),
                "title_text_sha256": text_sha256(title),
                "title_span_start": identifier.start(3),
                "identifier_semantics": _caption_identifier_semantics(
                    raw_identifier, title, domain_context
                ),
            }
        )

    legacy = CAPTION_PATTERN.match(value)
    if legacy:
        _, hierarchy, sequence, _ = legacy.groups()
        entry.update(
            {
                "heading_level": len(hierarchy.split(".")),
                "cached_hierarchy": hierarchy,
                "cached_sequence": sequence,
            }
        )
    return entry


def _table_text_hash(table: Any) -> str:
    value = "\u241e".join(
        "\u241f".join(cell.text for cell in row.cells) for row in table.rows
    )
    return text_sha256(value)


def _table_classification(table: Any, *, figure_panel: bool) -> str:
    if figure_panel:
        return "figure_panel"
    if len(table.rows) == 1 and len(table.columns) == 1:
        return "callout"
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    if borders is None:
        return "unknown"

    def visible(name: str) -> bool:
        element = borders.find(qn(f"w:{name}"))
        return bool(
            element is not None
            and element.get(qn("w:val"), "nil") not in {"nil", "none"}
        )

    if any(visible(name) for name in ("left", "right", "insideV")):
        return "grid"
    if visible("top") and visible("bottom"):
        return "three_line"
    return "unknown"


def _unnumbered_figure_caption_indexes(document: Any) -> set[int]:
    result: set[int] = set()
    for index, paragraph in enumerate(document.paragraphs):
        value = paragraph.text.strip()
        if not value or len(value) > 80 or "\n" in value:
            continue
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            continue
        if LOOSE_CAPTION_PATTERN.match(value) or any(
            pattern.match(value) for _, pattern in HEADING_PATTERNS
        ):
            continue
        if value.endswith(("\u3002", "\uff01", "\uff1f", "\uff1b", "\uff1a")):
            continue
        previous = paragraph._p.getprevious()
        if previous is None or previous.tag != qn("w:p"):
            continue
        if previous.xpath(".//w:drawing | .//w:pict"):
            result.add(index)
    return result


def _figure_panel_rows(table: Any) -> tuple[list[int], list[int]]:
    image_rows = [
        index
        for index, row in enumerate(table.rows)
        if row._tr.xpath(".//w:drawing | .//w:pict")
    ]
    label_rows: list[int] = []
    for image_row in image_rows:
        label_row = image_row + 1
        if label_row >= len(table.rows):
            continue
        row = table.rows[label_row]
        if row._tr.xpath(".//w:drawing | .//w:pict"):
            continue
        values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if not values or any(len(value) > 80 or "\n" in value for value in values):
            continue
        label_rows.append(label_row)
    return image_rows, sorted(set(label_rows))


def _drawing_state(paragraph: Any, drawing: Any) -> dict[str, Any]:
    inline = drawing.xpath("./wp:inline")
    anchored = drawing.xpath("./wp:anchor")
    container = inline[0] if inline else (anchored[0] if anchored else None)
    extent = None if container is None else container.find(qn("wp:extent"))
    rel_ids = drawing.xpath(".//a:blip/@r:embed")
    rel_id = rel_ids[0] if len(rel_ids) == 1 else None
    part = paragraph.part.related_parts.get(rel_id) if rel_id else None
    blob = None if part is None else part.blob
    content_type = None if part is None else part.content_type
    extension = ""
    if part is not None:
        extension = str(part.partname).rsplit(".", 1)[-1].lower()
    vector = extension in {"emf", "wmf", "svg"} or content_type in {
        "image/x-emf",
        "image/x-wmf",
        "image/svg+xml",
    }
    pixels = None
    dpi = None
    if blob is not None and not vector:
        try:
            image = DocxImage.from_blob(blob)
            pixels = {"width": int(image.px_width), "height": int(image.px_height)}
            dpi = {
                "horizontal": int(image.horz_dpi),
                "vertical": int(image.vert_dpi),
            }
        except (ValueError, TypeError, AttributeError, UnrecognizedImageError):
            pass
    cx = None if extent is None else int(extent.get("cx", "0"))
    cy = None if extent is None else int(extent.get("cy", "0"))
    effective_dpi = None
    if pixels and cx and cy:
        effective_dpi = min(
            pixels["width"] * EMU_PER_INCH / cx,
            pixels["height"] * EMU_PER_INCH / cy,
        )
    crop = bool(
        drawing.xpath(
            ".//*[local-name()='srcRect' and (@l or @t or @r or @b)]"
        )
    )
    return {
        "object_type": "inline" if inline else ("floating" if anchored else "unknown"),
        "relationship_id": rel_id,
        "media_sha256": None if blob is None else hashlib.sha256(blob).hexdigest(),
        "media_kind": "vector" if vector else "raster",
        "content_type": content_type,
        "source_extent_emu": {"cx": cx, "cy": cy},
        "source_pixels": pixels,
        "source_dpi": dpi,
        "source_effective_dpi": effective_dpi,
        "has_crop": crop,
        "supported": bool(inline and rel_id and blob is not None and cx and cy and not crop),
    }


def _effective_line_spacing(paragraph: Any) -> dict[str, Any]:
    sources: list[tuple[str, Any]] = [("paragraph", paragraph._p.pPr)]
    style = paragraph.style
    while style is not None:
        sources.append((f"style:{style.style_id}", style.element.pPr))
        style = style.base_style
    for source, properties in sources:
        if properties is None:
            continue
        spacing = properties.find(qn("w:spacing"))
        if spacing is None or spacing.get(qn("w:line")) is None:
            continue
        try:
            line = int(spacing.get(qn("w:line"), "0"))
        except ValueError:
            line = 0
        return {
            "source": source,
            "rule": spacing.get(qn("w:lineRule"), "auto"),
            "line_twips": line,
        }
    return {"source": "default", "rule": "auto", "line_twips": 240}


def _image_paragraph_payload(paragraph: Any) -> tuple[str, str | None]:
    if _paragraph_text_without_field_results(paragraph._p).strip():
        return "mixed_text", "image_shares_paragraph_with_authored_text"
    if paragraph._p.xpath(".//w:object | .//w:txbxContent"):
        return "mixed_object", "image_paragraph_contains_another_protected_object"
    if paragraph._p.xpath(".//w:pict"):
        return "legacy_picture", "legacy_picture_requires_qa"
    drawings = paragraph._p.xpath(".//w:drawing")
    if not drawings:
        return "unknown", "image_paragraph_has_no_supported_drawing"
    if any(drawing.xpath("./wp:anchor") for drawing in drawings):
        return "floating", "floating_image_line_spacing_must_remain_unchanged"
    if any(len(drawing.xpath("./wp:inline")) != 1 for drawing in drawings):
        return "unknown", "image_container_is_ambiguous"
    return "image_only", None


def _row_height_state(row: Any) -> dict[str, Any] | None:
    properties = row._tr.trPr
    if properties is None:
        return None
    heights = properties.findall(qn("w:trHeight"))
    if not heights:
        return None
    height = heights[-1]
    try:
        value = int(height.get(qn("w:val"), "0"))
    except ValueError:
        value = 0
    return {
        "rule": height.get(qn("w:hRule"), "atLeast"),
        "height_twips": value,
    }


def _image_visibility_candidate(
    paragraph: Any,
    state: dict[str, Any],
    table_context: dict[str, Any] | None,
) -> dict[str, Any]:
    payload, blocked_reason = _image_paragraph_payload(paragraph)
    spacing = _effective_line_spacing(paragraph)
    image_height_twips = round(int(state["source_extent_emu"].get("cy") or 0) / 635)
    line_risk = bool(
        payload == "image_only"
        and spacing["rule"] == "exact"
        and spacing["line_twips"] < image_height_twips
    )
    actions: list[str] = []
    if line_risk:
        actions.append("auto_single_line_spacing")

    row_state = None
    row_risk = False
    if table_context is not None:
        row_state = _row_height_state(table_context["row"])
        row_risk = bool(
            row_state
            and row_state["rule"] == "exact"
            and row_state["height_twips"] < image_height_twips
        )
        if row_risk:
            if table_context["table_entry"].get("complex_merge"):
                blocked_reason = blocked_reason or "complex_merged_table_row_requires_qa"
            elif table_context["table_entry"].get("has_floating_objects"):
                blocked_reason = blocked_reason or "table_row_contains_floating_objects"
            elif payload != "image_only":
                blocked_reason = blocked_reason or "mixed_table_image_paragraph_requires_qa"
            else:
                actions.append("relax_exact_table_row_height")

    if not actions and not blocked_reason:
        blocked_reason = None
    return {
        "approved": False,
        "action": actions,
        "paragraph_payload": payload,
        "source_line_spacing": spacing,
        "image_height_twips": image_height_twips,
        "fixed_line_clipping_candidate": line_risk,
        "table_row": (
            None
            if table_context is None
            else {
                "table": table_context["table_index"],
                "row": table_context["row_index"],
                "source_height": row_state,
                "fixed_height_clipping_candidate": row_risk,
            }
        ),
        "blocked_reason": blocked_reason,
    }


def _candidate_image_entry(
    paragraph: Any,
    locator: dict[str, Any],
    drawing: Any,
    drawing_index: int,
    image_index: int,
    placement: str,
    table_context: dict[str, Any] | None = None,
) -> dict[str, Any]:
    state = _drawing_state(paragraph, drawing)
    source_extent = state["source_extent_emu"]
    ratio = (
        None
        if not source_extent["cx"] or not source_extent["cy"]
        else source_extent["cx"] / source_extent["cy"]
    )
    figure_panel = placement == "table_figure_panel"
    standalone = placement == "standalone"
    max_width = 95 if figure_panel else (100 if ratio and ratio >= 1.6 else 90)
    reason = None
    if state["object_type"] != "inline":
        reason = "floating_or_unknown_drawing"
    elif not state["relationship_id"] or not state["media_sha256"]:
        reason = "missing_or_ambiguous_media_relationship"
    elif not source_extent["cx"] or not source_extent["cy"]:
        reason = "missing_display_extent"
    elif state["has_crop"]:
        reason = "existing_crop_requires_qa"
    elif not (figure_panel or standalone):
        reason = "table_embedded_image_requires_separate_qa"
    return {
        "image": f"IMG-{image_index:04d}",
        "locator": locator,
        "drawing_index": drawing_index,
        "placement": placement,
        "position_policy": "preserve_anchor",
        **state,
        "source_aspect_ratio": ratio,
        "resize": {
            "approved": False,
            "fit_mode": "fit_within_bounds",
            "aspect_ratio_locked": True,
            "alignment": "center",
            "max_width_percent": max_width,
            "max_height_percent": None if figure_panel else 65,
            "same_row_equal_height": figure_panel,
            "allow_upscale": False,
            "raster_upscale_max_percent": 100,
            "minimum_effective_dpi": 220,
        },
        "visibility": _image_visibility_candidate(
            paragraph, state, table_context
        ),
        "approval_blocked_reason": reason,
        "approved": False,
    }


def _candidate_images(
    document: Any,
    body_values: list[str],
    tables: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    image_index = 1
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        for drawing_index, drawing in enumerate(paragraph._p.xpath(".//w:drawing")):
            result.append(
                _candidate_image_entry(
                    paragraph,
                    _body_locator(paragraph_index, body_values),
                    drawing,
                    drawing_index,
                    image_index,
                    "standalone",
                    None,
                )
            )
            image_index += 1
    for table_index, table in enumerate(document.tables):
        table_entry = tables[table_index]
        image_rows = {int(value) for value in table_entry.get("image_rows", [])}
        seen_cells: set[Any] = set()
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    placement = (
                        "table_figure_panel"
                        if table_entry.get("layout_purpose") == "figure_panel"
                        and row_index in image_rows
                        else "table_embedded_unknown"
                    )
                    for drawing_index, drawing in enumerate(
                        paragraph._p.xpath(".//w:drawing")
                    ):
                        result.append(
                            _candidate_image_entry(
                                paragraph,
                                _cell_locator(
                                    table_index,
                                    row_index,
                                    cell_index,
                                    paragraph_index,
                                ),
                                drawing,
                                drawing_index,
                                image_index,
                                placement,
                                {
                                    "table_index": table_index,
                                    "row_index": row_index,
                                    "row": row,
                                    "table_entry": table_entry,
                                },
                            )
                        )
                        image_index += 1
    return result


def _approved_context_hashes(
    document: Any, structure_map: dict[str, Any], source_hash: str | None
) -> set[str | None]:
    accepted: set[str | None] = {source_hash}
    if source_hash is None:
        return accepted
    for caption in structure_map.get("captions", []):
        locator = caption.get("locator", {})
        if (
            caption.get("approved")
            and caption.get("action") == "replace_identifier"
            and locator.get("text_sha256") == source_hash
        ):
            try:
                accepted.add(text_sha256(resolve_paragraph_locator(document, locator).text))
            except FormatMonographError:
                continue
    return accepted


def _resolve_image_drawing(
    document: Any,
    entry: dict[str, Any],
    structure_map: dict[str, Any] | None = None,
) -> tuple[Any, Any]:
    locator = entry["locator"]
    drawing_index = int(entry["drawing_index"])
    try:
        paragraph = resolve_paragraph_locator(document, locator)
        drawings = paragraph._p.xpath(".//w:drawing")
        if 0 <= drawing_index < len(drawings):
            drawing = drawings[drawing_index]
            if _drawing_state(paragraph, drawing)["media_sha256"] == entry.get(
                "media_sha256"
            ):
                return paragraph, drawing
    except FormatMonographError:
        pass
    if locator.get("kind") == "body_paragraph":
        values = [candidate.text for candidate in document.paragraphs]
        expected_text = locator.get("text_sha256")
        before = locator.get("previous_nonempty_sha256")
        after = locator.get("next_nonempty_sha256")
        accepted_before = _approved_context_hashes(
            document, structure_map or {}, before
        )
        accepted_after = _approved_context_hashes(
            document, structure_map or {}, after
        )
        candidates: list[tuple[Any, Any]] = []
        for index, candidate in enumerate(document.paragraphs):
            if expected_text and text_sha256(candidate.text) != expected_text:
                continue
            if (
                before is not None
                and _nearest_nonempty_hash(values, index, -1) not in accepted_before
            ):
                continue
            if (
                after is not None
                and _nearest_nonempty_hash(values, index, 1) not in accepted_after
            ):
                continue
            drawings = candidate._p.xpath(".//w:drawing")
            if not 0 <= drawing_index < len(drawings):
                continue
            drawing = drawings[drawing_index]
            if _drawing_state(candidate, drawing)["media_sha256"] != entry.get(
                "media_sha256"
            ):
                continue
            candidates.append((candidate, drawing))
        if len(candidates) != 1:
            raise FormatMonographError(
                f"Approved image anchor is ambiguous or missing: {entry.get('image')}."
            )
        return candidates[0]
    raise FormatMonographError(
        f"Approved image drawing is missing or changed at {entry.get('image')}."
    )


def _section_index_for_table(document: Any, table: Any) -> int:
    children = list(document.element.body)
    try:
        position = children.index(table._tbl)
    except ValueError as exc:
        raise FormatMonographError("Approved image table is not a top-level body table.") from exc
    return sum(
        1
        for child in children[:position]
        if child.tag == qn("w:p")
        and child.find("./w:pPr/w:sectPr", namespaces=NS) is not None
    )


def _usable_section_extent(section: Any) -> tuple[int, int]:
    width = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    height = int(section.page_height) - int(section.top_margin) - int(section.bottom_margin)
    if width <= 0 or height <= 0:
        raise FormatMonographError("Section has no positive usable image area.")
    return width, height


def _image_available_extent(
    document: Any, structure_map: dict[str, Any], entry: dict[str, Any], paragraph: Any
) -> tuple[int, int | None]:
    resize = entry["resize"]
    locator = entry["locator"]
    if entry.get("placement") == "standalone":
        section = document.sections[section_index_for_paragraph(document, paragraph)]
        usable_width, usable_height = _usable_section_extent(section)
        return (
            round(usable_width * float(resize["max_width_percent"]) / 100),
            round(usable_height * float(resize["max_height_percent"]) / 100),
        )

    table_index = int(locator["table"])
    row_index = int(locator["row"])
    cell_index = int(locator["cell"])
    table = document.tables[table_index]
    cell = table.rows[row_index].cells[cell_index]
    section = document.sections[_section_index_for_table(document, table)]
    usable_width, _ = _usable_section_extent(section)
    cell_width = int(cell.width) if cell.width is not None else 0
    if cell_width <= 0:
        unique_cells = max(1, len({item._tc for item in table.rows[row_index].cells}))
        cell_width = usable_width // unique_cells
    table_entry = next(
        item
        for item in structure_map.get("tables", [])
        if int(item["table"]) == table_index
    )
    margins = table_entry.get("visual", {}).get("cell_margins_mm", {})
    horizontal_margin = float(margins.get("left", 0)) + float(margins.get("right", 0))
    cell_width = max(1, cell_width - round(horizontal_margin * EMU_PER_MM))
    return round(cell_width * float(resize["max_width_percent"]) / 100), None


def _image_upscale_limit(entry: dict[str, Any]) -> float:
    if entry.get("resize", {}).get("allow_upscale") is False:
        return 1.0
    if entry.get("media_kind") == "vector":
        return float("inf")
    resize = entry["resize"]
    policy_cap = float(resize["raster_upscale_max_percent"]) / 100
    effective_dpi = entry.get("source_effective_dpi")
    if effective_dpi is None:
        return 1.0
    dpi_cap = float(effective_dpi) / float(resize["minimum_effective_dpi"])
    return max(1.0, min(policy_cap, dpi_cap))


def _planned_image_extents(
    document: Any, structure_map: dict[str, Any]
) -> dict[str, tuple[int, int]]:
    approved = [
        entry
        for entry in structure_map.get("images", [])
        if entry.get("approved") and entry.get("resize", {}).get("approved")
    ]
    result: dict[str, tuple[int, int]] = {}
    groups: dict[tuple[int, int], list[dict[str, Any]]] = {}
    for entry in approved:
        paragraph, _ = _resolve_image_drawing(document, entry, structure_map)
        source = entry["source_extent_emu"]
        source_cx, source_cy = int(source["cx"]), int(source["cy"])
        max_width, max_height = _image_available_extent(
            document, structure_map, entry, paragraph
        )
        desired_scale = max_width / source_cx
        if max_height is not None:
            desired_scale = min(desired_scale, max_height / source_cy)
        scale = min(desired_scale, _image_upscale_limit(entry))
        target = (max(1, round(source_cx * scale)), max(1, round(source_cy * scale)))
        result[str(entry["image"])] = target
        locator = entry["locator"]
        if (
            entry.get("placement") == "table_figure_panel"
            and entry["resize"].get("same_row_equal_height")
        ):
            groups.setdefault(
                (int(locator["table"]), int(locator["row"])), []
            ).append(entry)
    for entries in groups.values():
        if len(entries) < 2:
            continue
        common_height = min(result[str(entry["image"])][1] for entry in entries)
        for entry in entries:
            ratio = float(entry["source_aspect_ratio"])
            result[str(entry["image"])] = (
                max(1, round(common_height * ratio)),
                common_height,
            )
    return result


def _set_drawing_extent(drawing: Any, cx: int, cy: int) -> None:
    inline = drawing.xpath("./wp:inline")
    if len(inline) != 1:
        raise FormatMonographError("Approved image is no longer an inline drawing.")
    extent = inline[0].find(qn("wp:extent"))
    if extent is None:
        raise FormatMonographError("Approved inline image has no display extent.")
    extent.set("cx", str(cx))
    extent.set("cy", str(cy))
    shape_extents = drawing.xpath(
        ".//*[local-name()='xfrm']/*[local-name()='ext']"
    )
    for shape_extent in shape_extents:
        shape_extent.set("cx", str(cx))
        shape_extent.set("cy", str(cy))


def _effective_paragraph_alignment(paragraph: Any) -> Any:
    alignment = paragraph.alignment
    style = paragraph.style
    while alignment is None and style is not None:
        alignment = style.paragraph_format.alignment
        style = style.base_style
    return alignment


def _apply_images(document: Any, structure_map: dict[str, Any]) -> int:
    planned = _planned_image_extents(document, structure_map)
    changed = 0
    for entry in structure_map.get("images", []):
        if not entry.get("approved") or not entry.get("resize", {}).get("approved"):
            continue
        paragraph, drawing = _resolve_image_drawing(document, entry, structure_map)
        state = _drawing_state(paragraph, drawing)
        if state["media_sha256"] != entry.get("media_sha256"):
            raise FormatMonographError(
                f"Approved image media changed before resize: {entry.get('image')}."
            )
        if state["object_type"] != "inline" or state["has_crop"]:
            raise FormatMonographError(
                f"Approved image anchor or crop state changed: {entry.get('image')}."
            )
        cx, cy = planned[str(entry["image"])]
        _set_drawing_extent(drawing, cx, cy)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        changed += 1
    return changed


def _set_image_paragraph_auto_spacing(paragraph: Any) -> bool:
    properties = paragraph._p.get_or_add_pPr()
    spacing = properties.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        properties.append(spacing)
    before = (spacing.get(qn("w:line")), spacing.get(qn("w:lineRule")))
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    return before != ("240", "auto")


def _relax_exact_row_height(row: Any) -> bool:
    properties = row._tr.get_or_add_trPr()
    heights = properties.findall(qn("w:trHeight"))
    if not heights:
        return False
    changed = False
    for height in heights:
        if height.get(qn("w:hRule"), "atLeast") == "exact":
            height.set(qn("w:hRule"), "atLeast")
            changed = True
    return changed


def _apply_image_visibility(document: Any, structure_map: dict[str, Any]) -> int:
    changed = 0
    seen_paragraphs: set[int] = set()
    seen_rows: set[tuple[int, int]] = set()
    for entry in structure_map.get("images", []):
        visibility = entry.get("visibility", {})
        if not visibility.get("approved"):
            continue
        paragraph, drawing = _resolve_image_drawing(document, entry, structure_map)
        state = _drawing_state(paragraph, drawing)
        if (
            state["object_type"] != "inline"
            or state["media_sha256"] != entry.get("media_sha256")
        ):
            raise FormatMonographError(
                f"Approved image visibility anchor changed: {entry.get('image')}."
            )
        actions = set(visibility.get("action", []))
        paragraph_key = id(paragraph._p)
        if (
            "auto_single_line_spacing" in actions
            and paragraph_key not in seen_paragraphs
        ):
            changed += int(_set_image_paragraph_auto_spacing(paragraph))
            seen_paragraphs.add(paragraph_key)
        if "relax_exact_table_row_height" in actions:
            row_info = visibility["table_row"]
            row_key = (int(row_info["table"]), int(row_info["row"]))
            if row_key in seen_rows:
                continue
            table_index, row_index = row_key
            if not 0 <= table_index < len(document.tables):
                raise FormatMonographError("Approved image table is out of range.")
            table = document.tables[table_index]
            if not 0 <= row_index < len(table.rows):
                raise FormatMonographError("Approved image table row is out of range.")
            changed += int(_relax_exact_row_height(table.rows[row_index]))
            seen_rows.add(row_key)
    return changed


def audit_structure_image_operations(
    document: Any, structure_map: dict[str, Any]
) -> list[dict[str, Any]]:
    failures: list[dict[str, Any]] = []
    try:
        planned = _planned_image_extents(document, structure_map)
    except FormatMonographError as exc:
        return [{"reason": "image_plan_unresolvable", "detail": str(exc)}]
    for entry in structure_map.get("images", []):
        if not entry.get("approved") or not entry.get("resize", {}).get("approved"):
            continue
        image_id = str(entry["image"])
        try:
            paragraph, drawing = _resolve_image_drawing(document, entry, structure_map)
            state = _drawing_state(paragraph, drawing)
        except FormatMonographError as exc:
            failures.append(
                {"image": image_id, "reason": "anchor_not_preserved", "detail": str(exc)}
            )
            continue
        expected_cx, expected_cy = planned[image_id]
        actual = state["source_extent_emu"]
        actual_cx, actual_cy = int(actual["cx"] or 0), int(actual["cy"] or 0)
        source_ratio = float(entry["source_aspect_ratio"])
        actual_ratio = actual_cx / actual_cy if actual_cy else 0
        checks = {
            "anchor_and_order_preserved": True,
            "position_policy_preserved": entry.get("position_policy")
            == "preserve_anchor",
            "inline_not_floating": state["object_type"] == "inline",
            "media_unchanged": state["media_sha256"] == entry.get("media_sha256"),
            "not_cropped": not state["has_crop"],
            "extent_matches_plan": (actual_cx, actual_cy)
            == (expected_cx, expected_cy),
            "aspect_ratio_locked": bool(actual_cy)
            and abs(actual_ratio - source_ratio) <= 0.00001,
            "centered_in_original_container": _effective_paragraph_alignment(paragraph)
            == WD_ALIGN_PARAGRAPH.CENTER,
        }
        if not all(checks.values()):
            failures.append(
                {
                    "image": image_id,
                    "reason": "approved_image_resize_mismatch",
                    "checks": checks,
                    "expected_extent_emu": {"cx": expected_cx, "cy": expected_cy},
                    "actual_extent_emu": {"cx": actual_cx, "cy": actual_cy},
                }
            )
    for entry in structure_map.get("images", []):
        visibility = entry.get("visibility", {})
        if not visibility.get("approved"):
            continue
        image_id = str(entry["image"])
        try:
            paragraph, drawing = _resolve_image_drawing(document, entry, structure_map)
            state = _drawing_state(paragraph, drawing)
        except FormatMonographError as exc:
            failures.append(
                {"image": image_id, "reason": "visibility_anchor_not_preserved", "detail": str(exc)}
            )
            continue
        actions = set(visibility.get("action", []))
        checks = {
            "position_policy_preserved": entry.get("position_policy")
            == "preserve_anchor",
            "inline_not_floating": state["object_type"] == "inline",
            "media_unchanged": state["media_sha256"] == entry.get("media_sha256"),
            "crop_state_unchanged": state["has_crop"] == entry.get("has_crop"),
        }
        if not entry.get("resize", {}).get("approved"):
            checks["display_extent_unchanged"] = state["source_extent_emu"] == entry.get(
                "source_extent_emu"
            )
        if "auto_single_line_spacing" in actions:
            spacing = _effective_line_spacing(paragraph)
            checks["line_spacing_expands_for_image"] = (
                spacing["source"] == "paragraph"
                and spacing["rule"] == "auto"
                and spacing["line_twips"] == 240
            )
        if "relax_exact_table_row_height" in actions:
            row_info = visibility["table_row"]
            table = document.tables[int(row_info["table"])]
            row_state = _row_height_state(table.rows[int(row_info["row"])])
            checks["table_row_not_exact"] = bool(
                row_state and row_state["rule"] != "exact"
            )
        if not all(checks.values()):
            failures.append(
                {
                    "image": image_id,
                    "reason": "approved_image_visibility_mismatch",
                    "checks": checks,
                }
            )
    return failures


def _candidate_pagination_groups(
    document: Any, captions: list[dict[str, Any]], body_values: list[str]
) -> list[dict[str, Any]]:
    groups: list[dict[str, Any]] = []
    table_indexes = {id(table._tbl): index for index, table in enumerate(document.tables)}
    for caption in captions:
        locator = caption.get("locator", {})
        if locator.get("kind") != "body_paragraph":
            continue
        index = int(locator["paragraph"])
        if caption.get("sequence_name") == "Figure" and index > 0:
            previous = document.paragraphs[index - 1]
            if previous._p.xpath(".//w:drawing | .//w:pict"):
                groups.append(
                    {
                        "kind": "figure_with_caption",
                        "anchor": _body_locator(index - 1, body_values),
                        "caption": copy.deepcopy(locator),
                        "approved": False,
                    }
                )
        elif caption.get("sequence_name") == "Table":
            sibling = document.paragraphs[index]._p.getnext()
            while sibling is not None and sibling.tag == qn("w:p") and not _paragraph_has_payload(sibling):
                sibling = sibling.getnext()
            if sibling is not None and sibling.tag == qn("w:tbl"):
                table_index = table_indexes.get(id(sibling))
                if table_index is not None:
                    groups.append(
                        {
                            "kind": "table_caption_with_table",
                            "anchor": copy.deepcopy(locator),
                            "table": table_index,
                            "table_text_sha256": _table_text_hash(
                                document.tables[table_index]
                            ),
                            "approved": False,
                        }
                    )
    return groups


def _candidate_qa_groups(
    headings: list[dict[str, Any]],
    appendices: list[dict[str, Any]],
    numbering_anomalies: list[dict[str, Any]],
    tables: list[dict[str, Any]],
    images: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Group repeated decisions without storing manuscript text."""
    groups: list[dict[str, Any]] = []
    frozen: list[dict[str, Any]] = []
    if headings:
        groups.append(
            {
                "id": "qa:heading-hierarchy",
                "kind": "heading_hierarchy",
                "item_count": len(headings),
                "decision_scope": "group_with_exceptions",
                "status": "open",
            }
        )
    if appendices:
        groups.append(
            {
                "id": "qa:appendix-recognition",
                "kind": "appendix_recognition",
                "item_count": len(appendices),
                "decision_scope": "group_with_exceptions",
                "status": "open",
            }
        )
        frozen.extend(
            {
                "id": f"freeze:appendix:{index}",
                "kind": "appendix",
                "locator": item["locator"],
                "status": "open",
            }
            for index, item in enumerate(appendices)
        )
    if numbering_anomalies:
        groups.append(
            {
                "id": "qa:heading-numbering",
                "kind": "heading_numbering",
                "item_count": len(numbering_anomalies),
                "decision_scope": "group_with_exceptions",
                "status": "open",
            }
        )
        frozen.extend(
            {
                "id": f"freeze:heading:{index}",
                "kind": "heading_numbering_and_toc",
                "locator": item.get("locator"),
                "status": "open",
            }
            for index, item in enumerate(numbering_anomalies)
        )
    for classification in (
        "three_line",
        "grid",
        "figure_panel",
        "callout",
        "unknown",
    ):
        matching_tables = [
            item for item in tables if item.get("classification") == classification
        ]
        if not matching_tables:
            continue
        groups.append(
            {
                "id": f"qa:table-kind:{classification}",
                "kind": "table_classification",
                "classification": classification,
                "item_count": len(matching_tables),
                "decision_scope": "group_with_exceptions",
                "status": "open",
            }
        )
        frozen.extend(
            {
                "id": f"freeze:table:{item['table']}",
                "kind": "table",
                "table": int(item["table"]),
                "status": "open",
            }
            for item in matching_tables
        )
    supported_images = [item for item in images if item.get("supported")]
    if supported_images:
        groups.append(
            {
                "id": "qa:image-resize-candidates",
                "kind": "image_handling",
                "item_count": len(supported_images),
                "decision_scope": "group_with_exceptions",
                "status": "open",
            }
        )
    ambiguous_images = [
        item
        for item in images
        if not item.get("supported") or item.get("approval_blocked_reason")
    ]
    if ambiguous_images:
        groups.append(
            {
                "id": "qa:ambiguous-images",
                "kind": "image_handling",
                "item_count": len(ambiguous_images),
                "decision_scope": "group_with_exceptions",
                "status": "open",
            }
        )
        frozen.extend(
            {
                "id": f"freeze:image:{item['image']}",
                "kind": "image",
                "image": item["image"],
                "status": "open",
            }
            for item in ambiguous_images
        )
    return groups, frozen


def _candidate_trial_selection(
    headings: list[dict[str, Any]],
    appendices: list[dict[str, Any]],
    tables: list[dict[str, Any]],
    images: list[dict[str, Any]],
) -> dict[str, Any]:
    selected_headings: list[dict[str, Any]] = []
    for level in (1, 2, 3, 4):
        selected_headings.extend(
            {"locator": item["locator"], "level": level}
            for item in [
                candidate
                for candidate in headings
                if int(candidate["level"]) == level
            ][:2]
        )
    table_ids: list[int] = []
    for kind in ("three_line", "grid", "figure_panel", "callout", "unknown"):
        table_ids.extend(
            int(item["table"])
            for item in [
                candidate
                for candidate in tables
                if candidate.get("classification") == kind
            ][:2]
            if int(item["table"]) not in table_ids
        )
    image_ids: list[str] = []
    for placement in ("standalone", "table_figure_panel", "table_embedded_unknown"):
        image_ids.extend(
            str(item["image"])
            for item in [
                candidate
                for candidate in images
                if candidate.get("placement") == placement
            ][:2]
            if str(item["image"]) not in image_ids
        )
    return {
        "mode": "representative_continuous_fragments",
        "whole_book_candidate": False,
        "max_rendered_pages_per_candidate": 30,
        "max_instances_per_object_kind": 2,
        "include_front_matter": True,
        "heading_samples": selected_headings,
        "appendix_samples": [item["locator"] for item in appendices[:2]],
        "table_samples": table_ids,
        "image_samples": image_ids,
        "status": "candidate",
    }


def _effective_outline_level(paragraph: Any) -> int | None:
    sources = [paragraph._p.pPr]
    style = paragraph.style
    while style is not None:
        sources.append(style.element.pPr)
        style = style.base_style
    for properties in sources:
        if properties is None:
            continue
        outline = properties.find(qn("w:outlineLvl"))
        if outline is None:
            continue
        try:
            return int(outline.get(qn("w:val"), "9"))
        except ValueError:
            return None
    if paragraph.style and paragraph.style.name in {
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
    }:
        return int(paragraph.style.name[-1]) - 1
    return None


def _toc_object_kinds(paragraph: Any) -> list[str]:
    checks = (
        ("drawing", ".//w:drawing"),
        ("legacy_picture", ".//w:pict"),
        ("ole_object", ".//w:object"),
        ("text_box", ".//w:txbxContent"),
    )
    return [name for name, expression in checks if paragraph._p.xpath(expression)]


def _candidate_toc_source(
    document: Any,
    headings: list[dict[str, Any]],
    appendices: list[dict[str, Any]],
    body_values: list[str],
) -> dict[str, Any]:
    candidate_indexes = {
        int(entry["locator"]["paragraph"])
        for entry in [*headings, *appendices]
        if entry.get("locator", {}).get("kind") == "body_paragraph"
    }
    contaminants: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs):
        outline_level = _effective_outline_level(paragraph)
        object_kinds = _toc_object_kinds(paragraph)
        in_toc_outline = outline_level is not None and 0 <= outline_level <= 3
        unapproved_outline = in_toc_outline and index not in candidate_indexes
        object_contamination = bool(object_kinds) and (
            index in candidate_indexes or in_toc_outline
        )
        if not (unapproved_outline or object_contamination):
            continue
        contaminants.append(
            {
                "locator": _body_locator(index, body_values),
                "object_kinds": object_kinds,
                "outline_level": outline_level,
                "candidate_heading": index in candidate_indexes,
                "reason": (
                    "toc_source_unapproved_outline_paragraph"
                    if unapproved_outline
                    else "toc_source_contains_non_text_object"
                ),
            }
        )
    return {
        "approved": False,
        "mode": "auto",
        "levels": 4,
        "tc_identifier": "M",
        "reject_non_text_results": True,
        "contaminants": contaminants,
    }


def candidate_structure_map(path: Path) -> dict[str, Any]:
    from _common import load_document

    document = load_document(path)
    headings = []
    captions = []
    paragraph_roles = []
    appendices = []
    chapter_starts = []
    body_values = [paragraph.text for paragraph in document.paragraphs]
    unnumbered_figure_captions = _unnumbered_figure_caption_indexes(document)
    for index, paragraph in enumerate(document.paragraphs):
        value = paragraph.text
        detected_level = None
        appendix_match = APPENDIX_PATTERN.match(value)
        if appendix_match and value.strip():
            appendix_number = value[appendix_match.start() : appendix_match.end()]
            appendices.append(
                {
                    "locator": _body_locator(index, body_values),
                    "text_sha256": text_sha256(value),
                    "existing_number_sha256": text_sha256(appendix_number),
                    "numbering_mode": "preserve_existing",
                    "numbering_anomalies": [],
                    "include_in_toc": None,
                    "toc_policy": "preserve_if_present",
                    "approved": False,
                }
            )
        for level, pattern in (() if appendix_match else HEADING_PATTERNS):
            if pattern.match(value):
                detected_level = level
                entry = {
                    "paragraph": index,
                    "locator": _body_locator(index, body_values),
                    "text_sha256": text_sha256(value),
                    "level": level,
                    "cached_number": list(_heading_number(value, level) or ()),
                    "source_style": paragraph.style.name if paragraph.style else None,
                    "direct_format_sha256": _paragraph_style_signature(paragraph),
                    "normalized_text_sha256": _heading_authored_text_hash(
                        value, level
                    ),
                    "approved": False,
                }
                headings.append(entry)
                if level == 1:
                    chapter = _chapter_number(value)
                    if chapter is not None:
                        chapter_starts.append(chapter)
                break
        nearby = "\n".join(body_values[max(0, index - 2) : index + 3])
        caption = _caption_entry(value, _body_locator(index, body_values), nearby)
        if caption:
            caption["paragraph"] = index
            captions.append(caption)
            role = "figure_caption" if caption["label"] == "图" else "table_caption"
        elif appendix_match and value.strip():
            role = "appendix_heading"
        elif index in unnumbered_figure_captions:
            role = "figure_caption_unnumbered"
        else:
            role = _role_for_paragraph(paragraph, detected_level)
        if value.strip():
            role_entry = {
                    "locator": _body_locator(index, body_values),
                    "text_sha256": text_sha256(value),
                    "role": role,
                    "canonical_role": normalized_role(role),
                    "source_style": paragraph.style.name if paragraph.style else None,
                    "direct_format_sha256": _paragraph_style_signature(paragraph),
                    "approved": False,
                }
            if detected_level is not None:
                role_entry["normalized_text_sha256"] = _heading_authored_text_hash(
                    value, detected_level
                )
            paragraph_roles.append(role_entry)

    tables = []
    for table_index, table in enumerate(document.tables):
        caption_row = None
        header_rows: list[int] = []
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        seen_cells: set[Any] = set()
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    locator = _cell_locator(
                        table_index, row_index, cell_index, paragraph_index
                    )
                    caption = _caption_entry(paragraph.text, locator, table_text)
                    if not caption:
                        continue
                    caption["migrate_outside_table"] = False
                    captions.append(caption)
                    paragraph_roles.append(
                        {
                            "locator": locator,
                            "text_sha256": text_sha256(paragraph.text),
                            "role": (
                                "figure_caption"
                                if caption["label"] == "图"
                                else "table_caption"
                            ),
                            "source_style": paragraph.style.name if paragraph.style else None,
                            "direct_format_sha256": _paragraph_style_signature(paragraph),
                            "approved": False,
                        }
                    )
                    if row_index == 0 and len({item._tc for item in row.cells}) == 1:
                        caption_row = 0
                        if len(table.rows) > 1:
                            header_rows = [1]

        image_rows, label_rows = _figure_panel_rows(table)
        if image_rows and label_rows:
            for row_index in label_rows:
                seen_label_cells: set[Any] = set()
                for cell_index, cell in enumerate(table.rows[row_index].cells):
                    if cell._tc in seen_label_cells:
                        continue
                    seen_label_cells.add(cell._tc)
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        if not paragraph.text.strip():
                            continue
                        paragraph_roles.append(
                            {
                                "locator": _cell_locator(
                                    table_index,
                                    row_index,
                                    cell_index,
                                    paragraph_index,
                                ),
                                "text_sha256": text_sha256(paragraph.text),
                                "role": "figure_panel_label",
                                "source_style": (
                                    paragraph.style.name if paragraph.style else None
                                ),
                                "direct_format_sha256": _paragraph_style_signature(
                                    paragraph
                                ),
                                "approved": False,
                            }
                        )

        first_row = "\u241f".join(cell.text for cell in table.rows[0].cells) if table.rows else ""
        unique_cells = {
            cell._tc for row in table.rows for cell in row.cells
        }
        visible_controls = sum(
            1
            for character in table_text
            if 0x2400 <= ord(character) <= 0x2426
            or (ord(character) < 32 and character not in "\t\n\r")
        )
        has_floating_objects = bool(
            table._tbl.xpath(".//wp:anchor | .//w:object | .//w:pict")
        )
        table_is_floating = table._tbl.tblPr.find(qn("w:tblpPr")) is not None
        classification = _table_classification(
            table, figure_panel=bool(image_rows and label_rows)
        )
        tables.append(
            {
                "table": table_index,
                "row_count": len(table.rows),
                "table_text_sha256": _table_text_hash(table),
                "first_row_sha256": text_sha256(first_row),
                "kind": "layout" if image_rows and label_rows else "unknown",
                "classification": classification,
                "layout_purpose": (
                    "figure_panel" if image_rows and label_rows else None
                ),
                "image_rows": image_rows,
                "label_rows": label_rows,
                "caption_row": caption_row,
                "header_rows": header_rows,
                "repeat_header_rows": [],
                "prevent_normal_row_split": False,
                "complex_merge": len(unique_cells)
                < sum(len(row.cells) for row in table.rows),
                "has_floating_objects": has_floating_objects,
                "table_is_floating": table_is_floating,
                "position_policy": "preserve_anchor",
                "visible_control_mark_candidates": visible_controls,
                "visual": {
                    "approved": False,
                    "alignment": "center" if not table_is_floating else None,
                    "available_width_percent": 100,
                    "allow_autofit": True,
                    "cell_margins_mm": {
                        "top": 1.0,
                        "right": 1.5,
                        "bottom": 1.0,
                        "left": 1.5,
                    },
                    "vertical_alignment": "center",
                    "border_preset": (
                        "borderless" if image_rows and label_rows else "preserve"
                    ),
                    "all_cell_alignment": (
                        "center" if image_rows and label_rows else None
                    ),
                    "text_wrapping": "none" if not table_is_floating else None,
                    "column_roles": (
                        []
                        if image_rows and label_rows
                        else ["unknown"] * len(table.columns)
                    ),
                    "orientation": "portrait",
                    "landscape_approved": False,
                },
                "approved": False,
            }
        )

    pagination_sections = _candidate_pagination_sections(document, headings)
    images = _candidate_images(document, body_values, tables)
    numbering_anomalies = _numbering_anomalies(headings)
    qa_groups, frozen_scopes = _candidate_qa_groups(
        headings, appendices, numbering_anomalies, tables, images
    )
    first_chapter = min(
        (int(item["paragraph"]) for item in headings if int(item["level"]) == 1),
        default=len(document.paragraphs),
    )
    original_toc_contains_appendix = any(
        int(item["locator"].get("paragraph", len(document.paragraphs))) < first_chapter
        for item in appendices
    )
    for appendix in appendices:
        appendix["original_toc_contains_appendix"] = original_toc_contains_appendix
    return {
        "schema_version": "1.5",
        "status": "candidate",
        "source_content_fingerprint_sha256": content_fingerprint(path),
        "paragraph_roles": paragraph_roles,
        "numbering": {
            "mode": "single_chapter" if len(set(chapter_starts)) == 1 else "whole_book",
            "chapter_start": chapter_starts[0] if chapter_starts else None,
            "heading_levels": 4,
            "expected_progression": "strict",
            "approved": False,
            "anomalies": numbering_anomalies,
        },
        "toc_source": _candidate_toc_source(
            document, headings, appendices, body_values
        ),
        "toc_ranges": [],
        "headings": headings,
        "captions": captions,
        "tables": tables,
        "images": images,
        "appendices": appendices,
        "qa_groups": qa_groups,
        "frozen_scopes": frozen_scopes,
        "trial_selection": _candidate_trial_selection(
            headings, appendices, tables, images
        ),
        "pagination_groups": _candidate_pagination_groups(document, captions, body_values),
        "table_cell_cleanups": [],
        "front_matter": _candidate_front_matter(
            document, pagination_sections, body_values
        ),
        "block_spacing": _candidate_block_spacing(),
        "pagination_sections": pagination_sections,
        "trailing_empty_sections": _trailing_empty_sections(document),
        "conflicts": [],
    }


def load_structure_map(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise FormatMonographError(f"Invalid structure map: {path}: {exc}") from exc
    if value.get("schema_version") not in STRUCTURE_MAP_VERSIONS:
        raise FormatMonographError(
            "Structure map schema_version must be 1.0, 1.1, 1.2, 1.3, 1.4, or 1.5."
        )
    if value.get("status") != "approved":
        raise FormatMonographError("Structure map status must be approved.")
    if value.get("conflicts"):
        raise FormatMonographError("Structure map contains unresolved conflicts.")
    if has_semantic_structure_map(value):
        for entry in value.get("paragraph_roles", []):
            if entry.get("approved") and entry.get("role") == "unknown":
                raise FormatMonographError("Approved paragraph role cannot be unknown.")
            if "locator" not in entry or "text_sha256" not in entry:
                raise FormatMonographError(
                    "Semantic structure-map paragraph roles require locator and text_sha256."
                )
        for entry in value.get("captions", []):
            if not entry.get("approved"):
                continue
            if value.get("schema_version") == "1.1":
                if entry.get("completeness") != "complete":
                    raise FormatMonographError("Incomplete captions cannot be approved.")
                if entry.get("hierarchy_status") not in {"match", "accepted"}:
                    raise FormatMonographError(
                        "Caption hierarchy must match or be explicitly accepted."
                    )
                continue

            action = entry.get("action")
            if action not in CAPTION_ACTIONS:
                raise FormatMonographError(
                    f"Structure map caption action is invalid: {action}"
                )
            if entry.get("numbering_mode") not in {"manual_text", "seq_field"}:
                raise FormatMonographError(
                    "Structure map 1.2+ captions require numbering_mode."
                )
            if entry.get("identifier_semantics") not in {
                "publication_number",
                "drawing_mark",
                "mixed",
                "unknown",
            }:
                raise FormatMonographError(
                    "Structure map caption identifier_semantics is invalid."
                )
            if entry.get("domain_context") not in {
                "general",
                "architecture",
                "civil_engineering",
                "mixed",
                "unknown",
            }:
                raise FormatMonographError(
                    "Structure map caption domain_context is invalid."
                )
            if entry.get("domain_confidence") not in {"high", "medium", "low"}:
                raise FormatMonographError(
                    "Structure map caption domain_confidence is invalid."
                )
            if action != "move_caption" and entry.get("migrate_outside_table"):
                raise FormatMonographError(
                    "Caption relocation requires the separate move_caption action."
                )
            if action == "replace_identifier":
                required = {
                    "identifier_span",
                    "identifier_sha256",
                    "identifier_prefix_sha256",
                    "identifier_suffix_sha256",
                    "title_text_sha256",
                    "title_span_start",
                    "replacement_identifier",
                }
                missing = sorted(required - set(entry))
                if missing or entry.get("replacement_confirmed") is not True:
                    raise FormatMonographError(
                        "Manual caption identifier replacement requires an exact boundary, "
                        "hashes, replacement_identifier, and replacement_confirmed=true."
                    )
                if entry.get("numbering_mode") != "manual_text":
                    raise FormatMonographError(
                        "Manual caption identifier replacement requires numbering_mode=manual_text."
                    )
            elif action == "convert_to_seq":
                required = {
                    "heading_level",
                    "cached_hierarchy",
                    "cached_sequence",
                    "sequence_name",
                }
                if required - set(entry):
                    raise FormatMonographError(
                        "SEQ conversion requires an unambiguous legacy caption boundary."
                    )
                if entry.get("numbering_mode") != "seq_field":
                    raise FormatMonographError(
                        "SEQ conversion requires numbering_mode=seq_field."
                    )
                if entry.get("hierarchy_status") not in {"match", "accepted"}:
                    raise FormatMonographError(
                        "SEQ conversion hierarchy must match or be explicitly accepted."
                    )
            elif action == "move_caption":
                if not entry.get("migrate_outside_table"):
                    raise FormatMonographError(
                        "move_caption requires migrate_outside_table=true."
                    )
                if entry.get("numbering_mode") != "manual_text":
                    raise FormatMonographError(
                        "move_caption preserves manual text; field conversion is separate."
                    )
        numbering = value.get("numbering", {})
        if numbering.get("approved"):
            chapter_start = numbering.get("chapter_start")
            if not isinstance(chapter_start, int) or chapter_start < 1:
                raise FormatMonographError(
                    "Approved numbering requires a positive integer chapter_start."
                )
            unresolved = [
                item
                for item in numbering.get("anomalies", [])
                if item.get("status", "open") != "accepted"
            ]
            if unresolved:
                raise FormatMonographError(
                    "Approved numbering contains unresolved progression anomalies."
                )
        if value.get("schema_version") in {"1.3", "1.4", "1.5"}:
            for entry in value.get("paragraph_roles", []):
                if not entry.get("approved"):
                    continue
                locator = entry.get("locator", {})
                if locator.get("kind") == "body_paragraph" and not locator.get(
                    "text_sha256"
                ):
                    raise FormatMonographError(
                        "Structure map 1.3 body locators require a stable text hash."
                    )
            for table in value.get("tables", []):
                if not table.get("approved"):
                    continue
                visual = table.get("visual", {})
                approved_figure_panel = (
                    visual.get("approved")
                    and table.get("layout_purpose") == "figure_panel"
                    and visual.get("border_preset") == "borderless"
                )
                if (
                    table.get("kind") == "layout"
                    and not table.get("pagination_only")
                    and not approved_figure_panel
                ):
                    raise FormatMonographError(
                        "Approved layout tables must be pagination-only or an explicitly "
                        "approved borderless figure panel."
                    )
                if table.get("repeat_caption_with_header") and not (
                    table.get("caption_row") == 0 and 1 in table.get("header_rows", [])
                ):
                    raise FormatMonographError(
                        "repeat_caption_with_header requires caption row 0 and header row 1."
                    )
            for group in value.get("pagination_groups", []):
                if not group.get("approved"):
                    continue
                if group.get("kind") not in {
                    "figure_with_caption",
                    "table_caption_with_table",
                    "keep_rows_together",
                }:
                    raise FormatMonographError("Invalid approved pagination group kind.")
                anchor = group.get("anchor")
                if group.get("kind") != "keep_rows_together" and not isinstance(
                    anchor, dict
                ):
                    raise FormatMonographError(
                        "Approved pagination group requires an anchor locator."
                    )
        if value.get("schema_version") in {"1.4", "1.5"}:
            front_matter = value.get("front_matter", {})
            if front_matter.get("approved"):
                if not isinstance(front_matter.get("book_title"), dict):
                    raise FormatMonographError(
                        "Approved front_matter requires a book_title locator."
                    )
                if front_matter.get("separate_title_page") is not True:
                    raise FormatMonographError(
                        "Approved book title must occupy a separate page."
                    )
                if front_matter.get("title_page_numbering") != "none":
                    raise FormatMonographError(
                        "Approved title page must be unnumbered and excluded from TOC pagination."
                    )
                if front_matter.get("toc_heading_text") not in {"目录", "目    录"}:
                    raise FormatMonographError(
                        "The technical-textbook TOC heading must be 目录 or 目    录."
                    )
                book_title_format = front_matter.get("book_title_format", {})
                if not isinstance(book_title_format, dict):
                    raise FormatMonographError(
                        "front_matter.book_title_format must be an object."
                    )
                unsupported_title_properties = set(book_title_format) - {
                    "font_name_ascii",
                    "font_name_east_asia",
                    "font_name_complex_script",
                    "font_size_pt",
                    "bold",
                    "alignment",
                    "first_line_indent_chars",
                    "line_spacing_rule",
                    "line_spacing_pt",
                    "space_before_pt",
                    "space_after_pt",
                }
                if unsupported_title_properties:
                    raise FormatMonographError(
                        "Unsupported book title format properties: "
                        + ", ".join(sorted(unsupported_title_properties))
                    )
                if front_matter.get("title_page_vertical_alignment") not in {
                    None,
                    "center",
                }:
                    raise FormatMonographError(
                        "Approved technical-textbook title pages must be vertically centered."
                    )
                if front_matter.get("insert_toc_heading_if_missing") is not True:
                    raise FormatMonographError(
                        "Approved front_matter must insert a missing TOC heading."
                    )
            block_spacing = value.get("block_spacing", {})
            if block_spacing.get("approved"):
                if block_spacing.get("mode") != "actual_blank_paragraph":
                    raise FormatMonographError(
                        "Approved figure/table spacing requires an actual blank paragraph."
                    )
                if block_spacing.get("blank_lines") != 1:
                    raise FormatMonographError(
                        "Approved figure/table spacing requires exactly one blank line."
                    )
                if block_spacing.get("same_page_only") is not True:
                    raise FormatMonographError(
                        "Figure/table blank spacing must be removed at a page boundary."
                    )
            pagination = value.get("pagination_sections", {})
            if pagination.get("approved"):
                required = {
                    "toc_start",
                    "body_start",
                    "number_format",
                    "start_at",
                    "continue_after_body_start",
                    "odd_position",
                    "even_position",
                    "show_on_first_page",
                }
                if required - set(pagination):
                    raise FormatMonographError(
                        "Approved pagination_sections is missing required settings."
                    )
                if pagination.get("number_format") != "decimal":
                    raise FormatMonographError(
                        "V0.2.5 technical-textbook pagination requires decimal numbering."
                    )
                starts = pagination.get("start_at", {})
                if starts != {"toc": 1, "body": 1}:
                    raise FormatMonographError(
                        "TOC and body pagination must each start at 1."
                    )
                if pagination.get("odd_position") != "outer_right" or pagination.get(
                    "even_position"
                ) != "outer_left":
                    raise FormatMonographError(
                        "Mirrored pagination requires odd outer-right and even outer-left."
                    )
                if pagination.get("show_on_first_page") is not True:
                    raise FormatMonographError(
                        "Approved TOC and body first pages must show page numbers."
                    )
                for name in ("toc_start", "body_start"):
                    if not isinstance(pagination.get(name), dict):
                        raise FormatMonographError(
                            f"Approved pagination_sections requires {name} locator."
                        )
            for table in value.get("tables", []):
                visual = table.get("visual", {})
                if table.get("position_policy") not in {None, "preserve_anchor"}:
                    raise FormatMonographError(
                        "Tables cannot be relocated; position_policy must be preserve_anchor."
                    )
                if not table.get("approved") or not visual.get("approved"):
                    continue
                if (
                    value.get("schema_version") == "1.5"
                    and table.get("table_is_floating")
                    and visual.get("text_wrapping") == "none"
                ):
                    raise FormatMonographError(
                        "A floating table cannot be changed to no-wrap without a separate position-preserving approval."
                    )
                if table.get("kind") == "layout":
                    if table.get("layout_purpose") != "figure_panel":
                        raise FormatMonographError(
                            "Visually formatted layout tables require layout_purpose=figure_panel."
                        )
                    if visual.get("border_preset") != "borderless":
                        raise FormatMonographError(
                            "Figure-panel layout tables must be borderless."
                        )
                    if visual.get("all_cell_alignment") != "center":
                        raise FormatMonographError(
                            "Figure-panel layout-table cells must be centered."
                        )
                    if visual.get("text_wrapping") != "none":
                        raise FormatMonographError(
                            "Figure-panel layout tables must use no text wrapping."
                        )
                    if not table.get("image_rows") or not table.get("label_rows"):
                        raise FormatMonographError(
                            "Figure-panel layout tables require approved image and label rows."
                        )
                    continue
                if table.get("kind") != "data":
                    raise FormatMonographError(
                        "Only approved data or figure-panel layout tables may receive visual formatting."
                    )
                if value.get("schema_version") == "1.5" and (
                    visual.get("alignment") != "center"
                    or visual.get("text_wrapping") != "none"
                ):
                    raise FormatMonographError(
                        "Approved schema 1.5 data tables must be centered and use no text wrapping."
                    )
                roles = visual.get("column_roles", [])
                if not roles or any(
                    role not in {"numeric", "unit", "short_code", "narrative"}
                    for role in roles
                ):
                    raise FormatMonographError(
                        "Approved table visuals require a known role for every column."
                    )
                if visual.get("border_preset") not in {
                    "preserve",
                    "three_line",
                    "full_grid",
                    "technical_textbook",
                    "borderless",
                }:
                    raise FormatMonographError("Invalid approved table border preset.")
                header_rows = sorted({int(row) for row in table.get("header_rows", [])})
                if visual.get("border_preset") == "technical_textbook":
                    if not header_rows:
                        raise FormatMonographError(
                            "Technical-textbook borders require approved header rows."
                        )
                    if header_rows != list(
                        range(header_rows[0], header_rows[-1] + 1)
                    ):
                        raise FormatMonographError(
                            "Multi-row table headers must be contiguous."
                        )
                    for key in ("major_border_pt", "minor_border_pt"):
                        if key in visual and not 0.25 <= float(visual[key]) <= 4:
                            raise FormatMonographError(
                                f"Approved table {key} must be between 0.25 and 4 pt."
                            )
                    for row_index in visual.get("horizontal_rule_rows", []):
                        if not 0 < int(row_index) < int(table.get("row_count", 10**9)):
                            raise FormatMonographError(
                                "Approved horizontal_rule_rows contains an invalid row."
                            )
                if visual.get("orientation") == "landscape" and visual.get(
                    "landscape_approved"
                ) is not True:
                    raise FormatMonographError(
                        "Landscape table layout requires landscape_approved=true."
                    )
            table_entries = {
                int(entry["table"]): entry for entry in value.get("tables", [])
            }
            for cleanup in value.get("table_cell_cleanups", []):
                if not cleanup.get("approved"):
                    continue
                table = table_entries.get(int(cleanup.get("table", -1)))
                if table is None or not table.get("approved"):
                    raise FormatMonographError(
                        "Approved table-cell cleanup requires an approved table."
                    )
                if table.get("complex_merge"):
                    raise FormatMonographError(
                        "Table-cell whitespace cleanup is blocked for complex merged tables."
                    )
                if cleanup.get("action") != "remove_leading_empty_paragraphs":
                    raise FormatMonographError(
                        "Unsupported table-cell cleanup action."
                    )
                if int(cleanup.get("count", 0)) < 1:
                    raise FormatMonographError(
                        "Table-cell cleanup requires a positive paragraph count."
                    )
                if not cleanup.get("cell_text_sha256") or not cleanup.get(
                    "result_cell_text_sha256"
                ):
                    raise FormatMonographError(
                        "Table-cell cleanup requires source and result cell hashes."
                    )
            table_entries = {
                int(entry["table"]): entry for entry in value.get("tables", [])
            }
            image_ids: set[str] = set()
            for image in value.get("images", []):
                image_id = str(image.get("image", ""))
                if not image_id or image_id in image_ids:
                    raise FormatMonographError(
                        "Structure-map image IDs must be unique and non-empty."
                    )
                image_ids.add(image_id)
                if image.get("position_policy") != "preserve_anchor":
                    raise FormatMonographError(
                        "Images cannot be relocated; position_policy must be preserve_anchor."
                    )
                visibility = image.get("visibility", {})
                if visibility.get("approved"):
                    actions = visibility.get("action", [])
                    if not actions or any(
                        action not in {
                            "auto_single_line_spacing",
                            "relax_exact_table_row_height",
                        }
                        for action in actions
                    ):
                        raise FormatMonographError(
                            "Approved image visibility requires a supported action."
                        )
                    if visibility.get("paragraph_payload") != "image_only":
                        raise FormatMonographError(
                            "Automatic image visibility repair requires an image-only paragraph."
                        )
                    if visibility.get("blocked_reason"):
                        raise FormatMonographError(
                            "Blocked image visibility candidates cannot be approved."
                        )
                    if image.get("object_type") != "inline":
                        raise FormatMonographError(
                            "Automatic image visibility repair supports inline drawings only."
                        )
                    if "auto_single_line_spacing" in actions and not visibility.get(
                        "fixed_line_clipping_candidate"
                    ):
                        raise FormatMonographError(
                            "Line-spacing repair requires fixed-line clipping evidence."
                        )
                    if "relax_exact_table_row_height" in actions:
                        row = visibility.get("table_row")
                        if not isinstance(row, dict) or not row.get(
                            "fixed_height_clipping_candidate"
                        ):
                            raise FormatMonographError(
                                "Table-row visibility repair requires fixed-height clipping evidence."
                            )
                        table = table_entries.get(int(row.get("table", -1)))
                        if table is None or table.get("complex_merge") or table.get(
                            "has_floating_objects"
                        ):
                            raise FormatMonographError(
                                "Complex or floating table rows require image visibility QA."
                            )
                if not image.get("approved"):
                    continue
                resize = image.get("resize", {})
                if resize.get("approved") is not True:
                    raise FormatMonographError(
                        "Approved image resizing requires resize.approved=true."
                    )
                if image.get("supported") is not True or image.get(
                    "approval_blocked_reason"
                ):
                    raise FormatMonographError(
                        "Unsupported, cropped, floating, or ambiguous images cannot be resized."
                    )
                if image.get("object_type") != "inline":
                    raise FormatMonographError(
                        "Approved image resizing supports inline drawings only."
                    )
                if resize.get("fit_mode") != "fit_within_bounds":
                    raise FormatMonographError("Unsupported image fit mode.")
                if resize.get("aspect_ratio_locked") is not True:
                    raise FormatMonographError(
                        "Approved images must keep their source aspect ratio."
                    )
                if resize.get("alignment") != "center":
                    raise FormatMonographError(
                        "Approved technical-textbook images must be centered in place."
                    )
                width = float(resize.get("max_width_percent", 0))
                height = resize.get("max_height_percent")
                if not 1 <= width <= 100 or (
                    height is not None and not 1 <= float(height) <= 100
                ):
                    raise FormatMonographError(
                        "Approved image width and height percentages must be between 1 and 100."
                    )
                if not 100 <= float(resize.get("raster_upscale_max_percent", 0)) <= 125:
                    raise FormatMonographError(
                        "Raster image enlargement must stay between 100% and 125%."
                    )
                if float(resize.get("minimum_effective_dpi", 0)) < 220:
                    raise FormatMonographError(
                        "Approved raster enlargement requires at least 220 effective DPI."
                    )
                extent = image.get("source_extent_emu", {})
                if int(extent.get("cx") or 0) <= 0 or int(extent.get("cy") or 0) <= 0:
                    raise FormatMonographError(
                        "Approved image resizing requires a positive source extent."
                    )
                if not image.get("media_sha256") or not isinstance(
                    image.get("locator"), dict
                ):
                    raise FormatMonographError(
                        "Approved image resizing requires a media hash and paragraph locator."
                    )
                placement = image.get("placement")
                if placement == "table_figure_panel":
                    locator = image["locator"]
                    table = table_entries.get(int(locator.get("table", -1)))
                    if not (
                        table
                        and table.get("approved")
                        and table.get("layout_purpose") == "figure_panel"
                        and table.get("position_policy", "preserve_anchor")
                        == "preserve_anchor"
                    ):
                        raise FormatMonographError(
                            "Figure-panel image resizing requires an approved stationary figure-panel table."
                        )
                elif placement != "standalone":
                    raise FormatMonographError(
                        "Data-table and unknown embedded images require separate QA and cannot be approved here."
                    )
        if value.get("schema_version") == "1.5":
            toc_source = value.get("toc_source", {})
            if toc_source.get("approved"):
                if toc_source.get("mode") not in {
                    "auto",
                    "heading_styles",
                    "tc_plain_text",
                }:
                    raise FormatMonographError(
                        "Approved TOC source mode must be auto, heading_styles, or tc_plain_text."
                    )
                if int(toc_source.get("levels", 0)) not in {1, 2, 3, 4}:
                    raise FormatMonographError(
                        "Approved TOC source levels must be between 1 and 4."
                    )
                if not re.fullmatch(
                    r"[A-Za-z]", str(toc_source.get("tc_identifier", ""))
                ):
                    raise FormatMonographError(
                        "Approved TOC TC identifier must be one ASCII letter."
                    )
                if toc_source.get("reject_non_text_results") is not True:
                    raise FormatMonographError(
                        "Approved TOC sources must reject every non-text result object."
                    )
            qa_ids: set[str] = set()
            for item in value.get("qa_groups", []):
                item_id = str(item.get("id", ""))
                if not item_id or item_id in qa_ids:
                    raise FormatMonographError(
                        "Structure-map 1.5 QA group IDs must be unique."
                    )
                qa_ids.add(item_id)
                if item.get("decision_scope") not in {
                    "group_with_exceptions",
                    "individual",
                }:
                    raise FormatMonographError(
                        "Structure-map 1.5 QA groups require a decision scope."
                    )
            freeze_ids: set[str] = set()
            for item in value.get("frozen_scopes", []):
                item_id = str(item.get("id", ""))
                if not item_id or item_id in freeze_ids:
                    raise FormatMonographError(
                        "Structure-map 1.5 frozen-scope IDs must be unique."
                    )
                freeze_ids.add(item_id)
            for appendix in value.get("appendices", []):
                if appendix.get("numbering_mode") != "preserve_existing":
                    raise FormatMonographError(
                        "Appendix numbering must preserve the manuscript value."
                    )
                if appendix.get("approved") and appendix.get("include_in_toc") not in {
                    True,
                    False,
                }:
                    raise FormatMonographError(
                        "Approved appendices require an explicit TOC inclusion decision."
                    )
                if appendix.get("approved") and not isinstance(
                    appendix.get("locator"), dict
                ):
                    raise FormatMonographError(
                        "Approved appendices require a stable locator."
                    )
            trial = value.get("trial_selection", {})
            if trial:
                if trial.get("whole_book_candidate") is not False:
                    raise FormatMonographError(
                        "Trial output cannot default to a whole-book candidate."
                    )
                pages = int(trial.get("max_rendered_pages_per_candidate", 0))
                if not 1 <= pages <= 30:
                    raise FormatMonographError(
                        "Trial candidates must be limited to at most 30 rendered pages."
                    )
            for image in value.get("images", []):
                if image.get("approved") and image.get("resize", {}).get(
                    "allow_upscale"
                ) is not False:
                    raise FormatMonographError(
                        "Structure-map 1.5 images cannot be enlarged automatically."
                    )
    return value


def validate_structure_map_source(path: Path, structure_map: dict[str, Any]) -> None:
    actual = content_fingerprint(path)
    expected = structure_map.get("source_content_fingerprint_sha256")
    if actual != expected:
        raise FormatMonographError(
            "Structure map source fingerprint does not match the input DOCX."
        )
    from _common import load_document

    document = load_document(path)
    if has_semantic_structure_map(structure_map):
        prime_structure_map_locators(document, structure_map)
    for key in ("headings", "captions", "appendices"):
        for entry in structure_map.get(key, []):
            if not entry.get("approved"):
                continue
            if entry.get("locator"):
                _verified_locator_paragraph(document, entry)
            else:
                _verified_paragraph(document, entry)
    for entry in structure_map.get("tables", []):
        if not entry.get("approved"):
            continue
        index = int(entry["table"])
        if not 0 <= index < len(document.tables):
            raise FormatMonographError(f"Structure-map table is out of range: {index}")
        if entry.get("table_text_sha256") and _table_text_hash(document.tables[index]) != entry["table_text_sha256"]:
            raise FormatMonographError(f"Structure-map table hash mismatch: {index}")
    for entry in structure_map.get("images", []):
        if not (
            entry.get("approved")
            or entry.get("visibility", {}).get("approved")
        ):
            continue
        paragraph, drawing = _resolve_image_drawing(
            document, entry, structure_map
        )
        state = _drawing_state(paragraph, drawing)
        if state["media_sha256"] != entry.get("media_sha256"):
            raise FormatMonographError(
                f"Structure-map image media hash mismatch: {entry.get('image')}."
            )
        if state["relationship_id"] != entry.get("relationship_id"):
            raise FormatMonographError(
                f"Structure-map image relationship mismatch: {entry.get('image')}."
            )
        if state["source_extent_emu"] != entry.get("source_extent_emu"):
            raise FormatMonographError(
                f"Structure-map image source extent mismatch: {entry.get('image')}."
            )
    pagination = structure_map.get("pagination_sections", {})
    front_matter = structure_map.get("front_matter", {})
    if front_matter.get("approved"):
        title = resolve_paragraph_locator(document, front_matter["book_title"])
        expected = front_matter["book_title"].get("text_sha256")
        if expected and text_sha256(title.text) != expected:
            raise FormatMonographError("Approved book-title locator hash mismatch.")
    if pagination.get("approved"):
        starts = [
            resolve_paragraph_locator(document, pagination[name])
            for name in ("toc_start", "body_start")
        ]
        if starts[0]._p.getparent().index(starts[0]._p) >= starts[1]._p.getparent().index(
            starts[1]._p
        ):
            raise FormatMonographError(
                "Approved TOC pagination start must precede the body start."
            )
    numbering = structure_map.get("numbering", {})
    if numbering.get("approved"):
        chapter_start = int(numbering["chapter_start"])
        approved_chapters = [
            _verified_locator_paragraph(document, entry).text
            for entry in structure_map.get("headings", [])
            if entry.get("approved") and int(entry.get("level", 0)) == 1
        ]
        detected = [_chapter_number(value) for value in approved_chapters]
        if detected and detected[0] != chapter_start:
            raise FormatMonographError(
                "Approved chapter_start does not match the first approved chapter heading."
            )


def _verified_paragraph(document: Any, entry: dict[str, Any]) -> Any:
    index = int(entry["paragraph"])
    if not 0 <= index < len(document.paragraphs):
        raise FormatMonographError(f"Structure-map paragraph is out of range: {index}")
    paragraph = document.paragraphs[index]
    if text_sha256(paragraph.text) != entry["text_sha256"]:
        raise FormatMonographError(
            f"Structure-map paragraph hash mismatch at paragraph {index}."
        )
    return paragraph


def _clear_paragraph(paragraph: Any) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _text_run(value: str) -> Any:
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    if value[:1].isspace() or value[-1:].isspace():
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = value
    run.append(text)
    return run


def _simple_field(instruction: str, placeholder: str) -> Any:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    field.set(qn("w:dirty"), "true")
    field.append(_text_run(placeholder))
    return field


def _complex_field_runs(instruction: str, placeholder: str) -> list[Any]:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)

    instruction_run = OxmlElement("w:r")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(
        "{http://www.w3.org/XML/1998/namespace}space", "preserve"
    )
    instruction_text.text = f" {instruction} "
    instruction_run.append(instruction_text)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    return [
        begin_run,
        instruction_run,
        separate_run,
        _text_run(placeholder),
        end_run,
    ]


def _apply_toc_ranges(document: Any, structure_map: dict[str, Any]) -> int:
    approved = [
        entry for entry in structure_map.get("toc_ranges", []) if entry.get("approved")
    ]
    for entry in approved:
        start, end = int(entry["start_paragraph"]), int(entry["end_paragraph"])
        hashes = entry.get("paragraph_sha256", [])
        if start < 0 or end < start or end >= len(document.paragraphs):
            raise FormatMonographError("Approved TOC range is out of bounds.")
        if len(hashes) != end - start + 1:
            raise FormatMonographError("Approved TOC range hash count is invalid.")
        for offset, index in enumerate(range(start, end + 1)):
            if text_sha256(document.paragraphs[index].text) != hashes[offset]:
                raise FormatMonographError(
                    f"Approved TOC range hash mismatch at paragraph {index}."
                )

    changed = 0
    for entry in sorted(
        approved, key=lambda item: int(item["start_paragraph"]), reverse=True
    ):
        start, end = int(entry["start_paragraph"]), int(entry["end_paragraph"])
        anchor = document.paragraphs[start]
        _clear_paragraph(anchor)
        p_pr = anchor._p.get_or_add_pPr()
        for name in ("pStyle", "outlineLvl", "keepNext", "pageBreakBefore"):
            element = p_pr.find(qn(f"w:{name}"))
            if element is not None:
                p_pr.remove(element)
        levels = int(entry.get("levels", 4))
        instruction = f'TOC \\o "1-{levels}" \\h \\z'
        if structure_map.get("schema_version") != "1.5":
            instruction += " \\u"
        anchor._p.extend(
            _complex_field_runs(instruction, "Update table of contents")
        )
        for index in range(end, start, -1):
            paragraph = document.paragraphs[index]
            parent = paragraph._p.getparent()
            if parent is None or parent.tag != qn("w:body"):
                raise FormatMonographError(
                    "Approved static TOC range must contain direct body paragraphs."
                )
            parent.remove(paragraph._p)
        changed += end - start + 1
    return changed


def _approved_toc_sources(
    document: Any, structure_map: dict[str, Any]
) -> list[tuple[Any, int, str, dict[str, Any]]]:
    maximum = int(structure_map.get("toc_source", {}).get("levels", 4))
    sources: list[tuple[Any, int, str, dict[str, Any]]] = []
    for entry in structure_map.get("headings", []):
        level = int(entry.get("level", 0))
        if not entry.get("approved") or not 1 <= level <= maximum:
            continue
        paragraph = (
            _verified_locator_paragraph(document, entry)
            if entry.get("locator")
            else _verified_paragraph(document, entry)
        )
        sources.append((paragraph, level, "heading", entry))
    for entry in structure_map.get("appendices", []):
        if not entry.get("approved") or entry.get("include_in_toc") is not True:
            continue
        sources.append(
            (_verified_locator_paragraph(document, entry), 1, "appendix", entry)
        )
    order = {id(paragraph._p): index for index, paragraph in enumerate(document.paragraphs)}
    sources.sort(key=lambda item: order.get(id(item[0]._p), 10**9))
    return sources


def toc_result_contract(
    document: Any, structure_map: dict[str, Any]
) -> list[dict[str, Any]] | None:
    toc_source = structure_map.get("toc_source", {})
    if not toc_source.get("approved"):
        return None
    result: list[dict[str, Any]] = []
    for paragraph, level, kind, entry in _approved_toc_sources(document, structure_map):
        value = paragraph.text
        if kind == "heading":
            match = _heading_prefix_pattern(level).match(value)
            if match:
                value = value[match.end() :]
        normalized = " ".join(value.split())
        if not normalized:
            raise FormatMonographError("Approved TOC source has no stable text hash.")
        result.append(
            {
                "level": level,
                "kind": kind,
                "text_sha256": text_sha256(normalized),
            }
        )
    if not result:
        raise FormatMonographError("Approved TOC source contains no approved headings.")
    return result


def _toc_style_mode_is_safe(
    document: Any,
    sources: list[tuple[Any, int, str, dict[str, Any]]],
    maximum: int,
) -> bool:
    approved = {id(paragraph._p): level for paragraph, level, _, _ in sources}
    for paragraph in document.paragraphs:
        outline = _effective_outline_level(paragraph)
        if outline is None or not 0 <= outline < maximum:
            continue
        expected = approved.get(id(paragraph._p))
        if expected != outline + 1 or _toc_object_kinds(paragraph):
            return False
    return all(
        _effective_outline_level(paragraph) == level - 1
        and not _toc_object_kinds(paragraph)
        for paragraph, level, _, _ in sources
    )


def _field_instruction_identifier(instruction: str) -> str | None:
    match = re.search(r'\\f\s+"?([A-Za-z])"?', instruction, re.IGNORECASE)
    return None if match is None else match.group(1).upper()


def _managed_tc_field(instruction: str, identifier: str) -> bool:
    return bool(
        re.match(r"\s*TC(?:\s|$)", instruction, re.IGNORECASE)
        and _field_instruction_identifier(instruction) == identifier.upper()
    )


def _field_paragraph(element: Any) -> Any | None:
    current = element
    while current is not None and current.tag != qn("w:p"):
        current = current.getparent()
    return current


def _clear_managed_tc_fields(
    document: Any,
    identifier: str,
    allowed_paragraphs: set[int],
) -> int:
    from field_writeback import parse_fields

    changed = 0
    for record in reversed(parse_fields(document.element)):
        if record.field_type != "TC" or not _managed_tc_field(
            record.instruction, identifier
        ):
            continue
        marker = record.simple if record.form == "simple" else record.begin
        assert marker is not None
        paragraph = _field_paragraph(marker)
        if paragraph is None or id(paragraph) not in allowed_paragraphs:
            raise FormatMonographError(
                "The reserved TOC TC identifier is already used outside approved headings."
            )
        if record.form == "simple":
            assert record.simple is not None
            parent = record.simple.getparent()
            if parent is not None:
                parent.remove(record.simple)
                changed += 1
            continue
        assert record.begin is not None and record.end is not None
        start_run, end_run = record.begin.getparent(), record.end.getparent()
        if (
            paragraph is None
            or start_run is None
            or end_run is None
            or start_run.getparent() is not paragraph
            or end_run.getparent() is not paragraph
        ):
            raise FormatMonographError(
                "A managed complex TC field cannot be removed without changing authored content."
            )
        start, end = paragraph.index(start_run), paragraph.index(end_run)
        if start > end:
            raise FormatMonographError("Managed TC field boundaries are out of order.")
        for index in range(end, start - 1, -1):
            paragraph.remove(paragraph[index])
        changed += 1
    return changed


def _ensure_managed_tc_fields(
    document: Any,
    sources: list[tuple[Any, int, str, dict[str, Any]]],
    identifier: str,
    numbering_approved: bool,
) -> int:
    from field_writeback import parse_fields

    expected = {}
    source_values = []
    for ordinal, (paragraph, level, kind, entry) in enumerate(sources, start=1):
        value = _toc_display_text(
            paragraph, level, kind, entry, numbering_approved
        )
        instruction = _tc_instruction(value, level, identifier)
        bookmark_name = _tc_bookmark_name(value, level, identifier, ordinal)
        expected[id(paragraph._p)] = (instruction, bookmark_name)
        source_values.append((paragraph, level, value, bookmark_name))
    actual: dict[int, list[tuple[str, str, str | None]]] = {
        key: [] for key in expected
    }
    for record in parse_fields(document.element):
        if record.field_type != "TC" or not _managed_tc_field(
            record.instruction, identifier
        ):
            continue
        marker = record.simple if record.form == "simple" else record.begin
        assert marker is not None
        paragraph = _field_paragraph(marker)
        key = None if paragraph is None else id(paragraph)
        if key not in expected:
            raise FormatMonographError(
                "The reserved TOC TC identifier is already used outside approved headings."
            )
        actual[key].append(
            (record.form, record.instruction, _tc_field_bookmark(record))
        )
    if all(
        values == [("complex", *expected[key])]
        for key, values in actual.items()
    ):
        return 0
    allowed = set(expected)
    changed = _clear_managed_tc_fields(document, identifier, allowed)
    used_ids = {
        int(value)
        for value in document.element.xpath(
            ".//w:bookmarkStart/@w:id | .//w:bookmarkEnd/@w:id"
        )
        if str(value).isdigit()
    }
    next_id = max(used_ids, default=-1) + 1
    for paragraph, level, value, bookmark_name in source_values:
        while next_id in used_ids:
            next_id += 1
        paragraph._p.extend(
            _tc_field_runs(
                value,
                level,
                identifier,
                next_id,
                bookmark_name,
            )
        )
        used_ids.add(next_id)
        next_id += 1
        changed += 1
    return changed


def _toc_display_text(
    paragraph: Any,
    level: int,
    kind: str,
    entry: dict[str, Any],
    numbering_approved: bool,
) -> str:
    value = paragraph.text.strip()
    if kind != "heading" or not numbering_approved:
        return value
    if _heading_prefix_pattern(level).match(value):
        return value
    number = tuple(int(item) for item in entry.get("cached_number", []))
    if len(number) != level:
        raise FormatMonographError("Approved TC heading has no complete cached number.")
    prefix = f"第{number[0]}章" if level == 1 else ".".join(map(str, number))
    return f"{prefix} {value}".strip()


def _tc_bookmark_name(
    text: str, level: int, identifier: str, ordinal: int
) -> str:
    digest = hashlib.sha256(
        f"{identifier.upper()}:{level}:{text}".encode("utf-8")
    ).hexdigest()[:8]
    return f"_Toc{ordinal:06d}{int(digest, 16):010d}"


def _tc_field_bookmark(record: Any) -> str | None:
    if record.form != "complex" or record.begin is None or record.end is None:
        return None
    paragraph = _field_paragraph(record.begin)
    start_run, end_run = record.begin.getparent(), record.end.getparent()
    if (
        paragraph is None
        or start_run is None
        or end_run is None
        or start_run.getparent() is not paragraph
        or end_run.getparent() is not paragraph
    ):
        return None
    start, end = paragraph.index(start_run), paragraph.index(end_run)
    starts: dict[str, str] = {}
    ends: set[str] = set()
    for child in paragraph[start : end + 1]:
        for element in child.iter():
            if element.tag == qn("w:bookmarkStart"):
                identifier = element.get(qn("w:id"))
                name = element.get(qn("w:name"))
                if identifier is not None and name:
                    starts[identifier] = name
            elif element.tag == qn("w:bookmarkEnd"):
                identifier = element.get(qn("w:id"))
                if identifier is not None:
                    ends.add(identifier)
    matched = [name for identifier, name in starts.items() if identifier in ends]
    return matched[0] if len(matched) == 1 else None


def _tc_field_runs(
    text: str,
    level: int,
    identifier: str,
    bookmark_id: int,
    bookmark_name: str,
) -> list[Any]:
    escaped = text.replace("\\", "\\\\").replace('"', '\\"')
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    instruction_run = OxmlElement("w:r")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(
        "{http://www.w3.org/XML/1998/namespace}space", "preserve"
    )
    instruction_text.text = ' TC "'
    instruction_run.append(instruction_text)
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), bookmark_name)
    title_run = OxmlElement("w:r")
    title_text = OxmlElement("w:instrText")
    title_text.text = escaped
    title_run.append(title_text)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))
    suffix_run = OxmlElement("w:r")
    suffix_text = OxmlElement("w:instrText")
    suffix_text.set(
        "{http://www.w3.org/XML/1998/namespace}space", "preserve"
    )
    suffix_text.text = f'" \\f {identifier.upper()} \\l "{level}" '
    suffix_run.append(suffix_text)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    runs = [
        begin_run,
        instruction_run,
        bookmark_start,
        title_run,
        bookmark_end,
        suffix_run,
        end_run,
    ]
    return runs


def _tc_instruction(text: str, level: int, identifier: str) -> str:
    escaped = text.replace("\\", "\\\\").replace('"', '\\"')
    return f'TC "{escaped}" \\f {identifier.upper()} \\l "{level}"'


def _set_single_toc_instruction(document: Any, instruction: str) -> int:
    from field_writeback import parse_fields

    records = [
        record for record in parse_fields(document.element) if record.field_type == "TOC"
    ]
    if len(records) != 1:
        raise FormatMonographError(
            "Approved TOC source strategy requires exactly one TOC field."
        )
    record = records[0]
    if record.instruction == " ".join(instruction.split()):
        return 0
    if record.form == "simple":
        assert record.simple is not None
        record.simple.set(qn("w:instr"), instruction)
        record.simple.set(qn("w:dirty"), "true")
        return 1
    assert record.begin is not None and record.separate is not None
    elements = list(document.element.iter())
    start, end = elements.index(record.begin), elements.index(record.separate)
    nodes = [
        element
        for element in elements[start + 1 : end]
        if element.tag == qn("w:instrText")
    ]
    if not nodes:
        raise FormatMonographError("TOC field has no instruction text node.")
    nodes[0].text = f" {instruction} "
    for node in nodes[1:]:
        node.text = ""
    record.begin.set(qn("w:dirty"), "true")
    return 1


def _apply_toc_source_strategy(
    document: Any, structure_map: dict[str, Any]
) -> tuple[int, str | None]:
    toc_source = structure_map.get("toc_source", {})
    if not toc_source.get("approved"):
        return 0, None
    sources = _approved_toc_sources(document, structure_map)
    if not sources:
        raise FormatMonographError("Approved TOC source contains no approved headings.")
    maximum = int(toc_source.get("levels", 4))
    requested = str(toc_source.get("mode", "auto"))
    style_safe = _toc_style_mode_is_safe(document, sources, maximum)
    if requested == "heading_styles" and not style_safe:
        raise FormatMonographError(
            "Heading-style TOC sources contain an unapproved or non-text object."
        )
    selected = (
        "heading_styles"
        if requested == "heading_styles" or requested == "auto" and style_safe
        else "tc_plain_text"
    )
    identifier = str(toc_source.get("tc_identifier", "M")).upper()
    allowed_paragraphs = {id(paragraph._p) for paragraph, _, _, _ in sources}
    if selected == "heading_styles":
        changed = _clear_managed_tc_fields(
            document, identifier, allowed_paragraphs
        )
        changed += _set_single_toc_instruction(
            document, f'TOC \\o "1-{maximum}" \\h \\z'
        )
        return changed, selected

    numbering_approved = bool(structure_map.get("numbering", {}).get("approved"))
    changed = _ensure_managed_tc_fields(
        document, sources, identifier, numbering_approved
    )
    changed += _set_single_toc_instruction(
        document, f"TOC \\f {identifier} \\h \\z"
    )
    return changed, selected


def audit_structure_toc_source_operations(
    document: Any, structure_map: dict[str, Any]
) -> list[dict[str, Any]]:
    toc_source = structure_map.get("toc_source", {})
    if not toc_source.get("approved"):
        return []
    from field_writeback import parse_fields

    try:
        sources = _approved_toc_sources(document, structure_map)
        contract = toc_result_contract(document, structure_map)
        records = parse_fields(document.element)
        toc_records = [record for record in records if record.field_type == "TOC"]
        if len(toc_records) != 1 or contract is None:
            raise FormatMonographError(
                "Approved TOC source requires exactly one TOC field and a source contract."
            )
        instruction = toc_records[0].instruction
        identifier = _field_instruction_identifier(instruction)
        maximum = int(toc_source.get("levels", 4))
        style_mode = bool(re.search(r"\\o\s+", instruction, re.IGNORECASE))
        if style_mode:
            if not _toc_style_mode_is_safe(document, sources, maximum):
                raise FormatMonographError(
                    "Heading-style TOC contains an unapproved or non-text source."
                )
            if any(record.field_type == "TC" for record in records):
                raise FormatMonographError(
                    "Heading-style TOC cannot retain managed TC sources."
                )
            return []
        expected_identifier = str(toc_source.get("tc_identifier", "M")).upper()
        if identifier != expected_identifier:
            raise FormatMonographError(
                "Plain-text TOC does not use the approved TC identifier."
            )
        expected = {}
        numbering_approved = bool(
            structure_map.get("numbering", {}).get("approved")
        )
        for ordinal, (paragraph, level, kind, entry) in enumerate(
            sources, start=1
        ):
            value = _toc_display_text(
                paragraph, level, kind, entry, numbering_approved
            )
            expected[id(paragraph._p)] = (
                _tc_instruction(value, level, expected_identifier),
                _tc_bookmark_name(
                    value, level, expected_identifier, ordinal
                ),
            )
        actual: dict[int, list[tuple[str, str, str | None]]] = {
            key: [] for key in expected
        }
        for record in records:
            if record.field_type != "TC" or not _managed_tc_field(
                record.instruction, expected_identifier
            ):
                continue
            marker = record.simple if record.form == "simple" else record.begin
            paragraph = None if marker is None else _field_paragraph(marker)
            key = None if paragraph is None else id(paragraph)
            if key not in expected:
                raise FormatMonographError(
                    "Managed TC source is outside an approved pure-text heading."
                )
            actual[key].append(
                (record.form, record.instruction, _tc_field_bookmark(record))
            )
        if any(
            values != [("complex", *expected[key])]
            for key, values in actual.items()
        ):
            raise FormatMonographError(
                "Managed TC sources do not match approved headings one-to-one."
            )
        return []
    except (FormatMonographError, KeyError, TypeError, ValueError) as exc:
        return [{"reason": "toc_source_contract_mismatch", "detail": str(exc)}]


def _apply_headings(document: Any, structure_map: dict[str, Any]) -> int:
    changed = 0
    clear_direct_numbering = bool(
        structure_map.get("numbering", {}).get("approved")
    )
    for entry in structure_map.get("headings", []):
        if not entry.get("approved"):
            continue
        level = int(entry["level"])
        if level not in {1, 2, 3, 4}:
            raise FormatMonographError(f"Invalid approved heading level: {level}")
        paragraph = (
            _verified_locator_paragraph(document, entry)
            if entry.get("locator")
            else _verified_paragraph(document, entry)
        )
        style = ensure_paragraph_style(document, f"Heading {level}")
        paragraph.style = style
        normalize_structural_paragraph(
            paragraph,
            style=style,
            clear_direct_numbering=clear_direct_numbering,
        )
        changed += 1
    return changed


def _apply_appendices(document: Any, structure_map: dict[str, Any]) -> int:
    changed = 0
    for entry in structure_map.get("appendices", []):
        if not entry.get("approved"):
            continue
        paragraph = _verified_locator_paragraph(document, entry)
        if entry.get("numbering_mode") != "preserve_existing":
            raise FormatMonographError("Appendix numbering cannot be rebuilt automatically.")
        if entry.get("include_in_toc") is True:
            paragraph.style = ensure_paragraph_style(document, "Heading 1")
            changed += 1
    return changed


def _remove_prefix_from_runs(paragraph: Any, length: int) -> None:
    remaining = length
    for run in paragraph.runs:
        if remaining <= 0:
            break
        take = min(len(run.text), remaining)
        run.text = run.text[take:]
        remaining -= take
    if remaining:
        raise FormatMonographError("Caption prefix spans an unsupported object.")


def _move_caption_before_table(
    document: Any, paragraph: Any, locator: dict[str, Any]
) -> Any:
    if locator.get("kind") != "table_cell_paragraph":
        return paragraph
    table_index = int(locator["table"])
    row_index = int(locator["row"])
    table = document.tables[table_index]
    row = table.rows[row_index]
    unique_cells = {cell._tc: cell for cell in row.cells}
    if len(unique_cells) != 1:
        raise FormatMonographError("Approved caption row must contain one merged cell.")
    cell = next(iter(unique_cells.values()))
    other_text = [item.text for item in cell.paragraphs if item._p is not paragraph._p]
    if any(value.strip() for value in other_text):
        raise FormatMonographError("Approved caption row contains additional authored text.")

    key = _locator_key(locator)
    table._tbl.addprevious(paragraph._p)
    moved = Paragraph(paragraph._p, table._parent)
    table._tbl.remove(row._tr)
    cache = getattr(document, "_format_monograph_locator_cache", {})
    cache[key] = moved
    setattr(document, "_format_monograph_locator_cache", cache)
    deleted_rows = getattr(document, "_format_monograph_deleted_table_rows", {})
    deleted_rows[table_index] = row_index
    setattr(document, "_format_monograph_deleted_table_rows", deleted_rows)
    return moved


def _replace_caption_identifier(paragraph: Any, entry: dict[str, Any]) -> None:
    span = entry["identifier_span"]
    start, end = int(span["start"]), int(span["end"])
    value = paragraph.text
    if not 0 <= start < end <= len(value):
        raise FormatMonographError("Approved caption identifier span is out of range.")
    if text_sha256(value[start:end]) != entry["identifier_sha256"]:
        raise FormatMonographError("Approved caption identifier hash does not match.")
    if text_sha256(value[:start]) != entry["identifier_prefix_sha256"]:
        raise FormatMonographError("Caption text before the identifier changed.")
    if text_sha256(value[end:]) != entry["identifier_suffix_sha256"]:
        raise FormatMonographError("Caption title or separator changed before replacement.")

    runs = list(paragraph.runs)
    if "".join(run.text for run in runs) != value:
        raise FormatMonographError(
            "Caption identifier crosses an unsupported inline object or hyperlink."
        )
    replacement = str(entry["replacement_identifier"])
    cursor = 0
    inserted = False
    for run in runs:
        run_start, run_end = cursor, cursor + len(run.text)
        cursor = run_end
        overlap_start, overlap_end = max(start, run_start), min(end, run_end)
        if overlap_start >= overlap_end:
            continue
        unsupported = [
            child
            for child in run._r
            if child.tag not in {qn("w:rPr"), qn("w:t")}
        ]
        if unsupported:
            raise FormatMonographError(
                "Caption identifier shares a run with an unsupported inline object."
            )
        local_start, local_end = overlap_start - run_start, overlap_end - run_start
        before, after = run.text[:local_start], run.text[local_end:]
        run.text = before + (replacement if not inserted else "") + after
        inserted = True
    if not inserted:
        raise FormatMonographError("Approved caption identifier could not be replaced.")


def _convert_caption_to_seq(document: Any, paragraph: Any, entry: dict[str, Any]) -> None:
    match = CAPTION_PATTERN.match(paragraph.text)
    if not match:
        raise FormatMonographError(
            f"Approved SEQ caption no longer matches at {entry.get('locator', entry.get('paragraph'))}."
        )
    label, hierarchy, sequence, _ = match.groups()
    if label != entry["label"]:
        raise FormatMonographError("Approved caption label does not match.")
    description_start = match.start(4)
    _remove_prefix_from_runs(paragraph, description_start)
    insertion = 1 if paragraph._p.pPr is not None else 0
    elements = [
        _text_run(f"{label} "),
        _simple_field(
            f"STYLEREF {int(entry['heading_level'])} \\s",
            entry.get("cached_hierarchy", hierarchy),
        ),
        _text_run("-"),
        _simple_field(
            f"SEQ {entry['sequence_name']} \\* ARABIC \\s {int(entry['heading_level'])}",
            entry.get("cached_sequence", sequence),
        ),
        _text_run(" "),
    ]
    for element in reversed(elements):
        paragraph._p.insert(insertion, element)
    paragraph.style = ensure_paragraph_style(document, "Caption")


def _apply_captions(document: Any, structure_map: dict[str, Any]) -> int:
    changed = 0
    version = structure_map.get("schema_version")
    for entry in structure_map.get("captions", []):
        if not entry.get("approved"):
            continue
        paragraph = (
            _verified_locator_paragraph(document, entry)
            if entry.get("locator")
            else _verified_paragraph(document, entry)
        )
        if version in {"1.2", "1.3", "1.4", "1.5"}:
            action = entry["action"]
            if action in {"preserve", "style_only"}:
                continue
            if action == "replace_identifier":
                _replace_caption_identifier(paragraph, entry)
                changed += 1
                continue
            if action == "move_caption":
                paragraph = _move_caption_before_table(
                    document, paragraph, entry["locator"]
                )
                paragraph.style = ensure_paragraph_style(document, "Caption")
                changed += 1
                continue
            if action != "convert_to_seq":
                raise FormatMonographError(f"Unsupported caption action: {action}")

        if entry.get("migrate_outside_table"):
            paragraph = _move_caption_before_table(
                document, paragraph, entry["locator"]
            )
        _convert_caption_to_seq(document, paragraph, entry)
        changed += 1
    return changed


def _set_row_property(row: Any, name: str, enabled: bool) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = tr_pr.find(qn(f"w:{name}"))
    if enabled and element is None:
        element = OxmlElement(f"w:{name}")
        element.set(qn("w:val"), "true")
        tr_pr.append(element)
    elif not enabled and element is not None:
        tr_pr.remove(element)


def _set_paragraph_property(paragraph: Any, name: str, enabled: bool) -> bool:
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn(f"w:{name}"))
    if enabled and element is None:
        element = OxmlElement(f"w:{name}")
        element.set(qn("w:val"), "true")
        p_pr.append(element)
        return True
    if not enabled and element is not None:
        p_pr.remove(element)
        return True
    return False


def _apply_outline_cleanup(document: Any, structure_map: dict[str, Any]) -> int:
    if structure_map.get("schema_version") not in {"1.3", "1.4", "1.5"}:
        return 0
    changed = 0
    for entry in structure_map.get("paragraph_roles", []):
        if not entry.get("approved"):
            continue
        role = normalized_role(str(entry.get("role", "unknown")))
        if role.startswith("heading_") or role in {"title", "subtitle"}:
            continue
        paragraph = _verified_locator_paragraph(document, entry)
        touched = False
        p_pr = paragraph._p.get_or_add_pPr()
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is not None:
            p_pr.remove(outline)
            touched = True
        if paragraph.style and paragraph.style.name in {
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Heading 4",
        }:
            style_name = ROLE_STYLE_NAMES.get(role, "Normal")
            paragraph.style = ensure_paragraph_style(document, style_name)
            touched = True
        changed += int(touched)
    return changed


def _apply_pagination_groups(document: Any, structure_map: dict[str, Any]) -> int:
    if structure_map.get("schema_version") not in {"1.3", "1.4", "1.5"}:
        return 0
    changed = 0
    for group in structure_map.get("pagination_groups", []):
        if not group.get("approved"):
            continue
        kind = group["kind"]
        if kind in {"figure_with_caption", "table_caption_with_table"}:
            paragraph = resolve_paragraph_locator(document, group["anchor"])
            changed += int(_set_paragraph_property(paragraph, "keepNext", True))
            continue
        if kind == "keep_rows_together":
            table_index = int(group["table"])
            if not 0 <= table_index < len(document.tables):
                raise FormatMonographError("Pagination table locator is out of range.")
            table = document.tables[table_index]
            if group.get("table_text_sha256") and _table_text_hash(table) != group[
                "table_text_sha256"
            ]:
                raise FormatMonographError("Pagination table hash mismatch.")
            for row in table.rows:
                before = row._tr.get_or_add_trPr().find(qn("w:cantSplit"))
                _set_row_property(row, "cantSplit", True)
                changed += int(before is None)
    return changed


def _special_paragraph_style(
    document: Any, name: str, *, size_pt: float, bold: bool
) -> Any:
    style = ensure_paragraph_style(document, name)
    style.base_style = document.styles["Normal"]
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    return style


def _apply_front_matter(document: Any, structure_map: dict[str, Any]) -> int:
    settings = structure_map.get("front_matter", {})
    if not settings.get("approved"):
        return 0
    title = resolve_paragraph_locator(document, settings["book_title"])
    title_format = dict(DEFAULT_BOOK_TITLE_FORMAT)
    title_format.update(settings.get("book_title_format", {}))
    title_style = _special_paragraph_style(
        document,
        BOOK_TITLE_STYLE,
        size_pt=float(title_format["font_size_pt"]),
        bold=bool(title_format["bold"]),
    )
    apply_style_properties(title_style, title_format)
    title.style = title_style
    clear_controlled_direct_format(title, title_format)
    normalize_structural_paragraph(
        title,
        style=title_style,
        clear_direct_numbering=True,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)

    toc = resolve_paragraph_locator(
        document, structure_map["pagination_sections"]["toc_start"]
    )
    previous = toc._p.getprevious()
    toc_heading = None
    if previous is not None and previous.tag == qn("w:p"):
        candidate = Paragraph(previous, document)
        if (
            candidate.text.strip() == settings["toc_heading_text"]
            or (
                candidate.style is not None
                and candidate.style.name == TOC_HEADING_STYLE
            )
        ):
            toc_heading = candidate
    inserted = False
    if toc_heading is None:
        element = OxmlElement("w:p")
        toc._p.addprevious(element)
        toc_heading = Paragraph(element, document)
        toc_heading.add_run(settings["toc_heading_text"])
        inserted = True
    elif toc_heading.text != settings["toc_heading_text"]:
        toc_heading.text = settings["toc_heading_text"]
    toc_heading.style = _special_paragraph_style(
        document, TOC_HEADING_STYLE, size_pt=18, bold=True
    )
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.paragraph_format.first_line_indent = Pt(0)
    for run in toc_heading.runs:
        run.bold = True
    setattr(document, "_format_monograph_book_title", title)
    setattr(document, "_format_monograph_toc_heading", toc_heading)
    return 1 + int(inserted)


def _new_block_spacer(document: Any) -> Any:
    style = _special_paragraph_style(
        document, BLOCK_SPACER_STYLE, size_pt=10.5, bold=False
    )
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(18)
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), style.style_id)
    p_pr.append(p_style)
    paragraph.append(p_pr)
    return paragraph


def _is_block_spacer(element: Any) -> bool:
    return bool(
        element is not None
        and element.tag == qn("w:p")
        and element.xpath(f"./w:pPr/w:pStyle[@w:val='{BLOCK_SPACER_STYLE.replace(' ', '')}']")
    )


def _starts_new_page(element: Any) -> bool:
    if element is None or element.tag != qn("w:p"):
        return False
    sect_pr = element.find("./w:pPr/w:sectPr", namespaces=NS)
    if sect_pr is not None:
        section_type = sect_pr.find(qn("w:type"))
        return section_type is None or section_type.get(qn("w:val")) != "continuous"
    return element.find("./w:pPr/w:pageBreakBefore", namespaces=NS) is not None


def _apply_block_spacing(document: Any, structure_map: dict[str, Any]) -> int:
    settings = structure_map.get("block_spacing", {})
    if not settings.get("approved"):
        return 0
    body = document.element.body
    targets: list[Any] = []
    requested = set(settings.get("after", []))
    if "approved_data_table" in requested:
        targets.extend(
            table._tbl
            for table, _ in approved_data_tables(document, structure_map)
            if table._tbl.getparent() is body
        )
    if "approved_figure_caption" in requested:
        for group in structure_map.get("pagination_groups", []):
            if not group.get("approved") or group.get("kind") != "figure_with_caption":
                continue
            caption = resolve_paragraph_locator(document, group["caption"])
            if caption._p.getparent() is body:
                targets.append(caption._p)

    changed = 0
    for target in sorted(set(targets), key=lambda item: list(body).index(item), reverse=True):
        following = target.getnext()
        if following is None or following.tag == qn("w:sectPr"):
            continue
        if _is_block_spacer(following):
            continue
        if _starts_new_page(following):
            continue
        target.addnext(_new_block_spacer(document))
        changed += 1
    return changed


def _section_properties_for_body_child(document: Any, child: Any) -> Any:
    children = list(document.element.body)
    position = children.index(child)
    for candidate in children[position:]:
        if candidate.tag == qn("w:p"):
            p_pr = candidate.find(qn("w:pPr"))
            sect_pr = None if p_pr is None else p_pr.find(qn("w:sectPr"))
            if sect_pr is not None:
                return sect_pr
        elif candidate.tag == qn("w:sectPr"):
            return candidate
    raise FormatMonographError("Table section properties could not be resolved.")


def _section_boundary(properties: Any) -> Any:
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    sect_pr = copy.deepcopy(properties)
    section_type = sect_pr.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        sect_pr.insert(0, section_type)
    section_type.set(qn("w:val"), "nextPage")
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is not None:
        pg_num.attrib.pop(qn("w:start"), None)
    p_pr.append(sect_pr)
    paragraph.append(p_pr)
    return paragraph


def _wrap_table_landscape(document: Any, table: Any) -> None:
    body = document.element.body
    if table._tbl.getparent() is not body:
        raise FormatMonographError(
            "A landscape table must be a top-level body table, not a nested layout table."
        )
    source = _section_properties_for_body_child(document, table._tbl)
    portrait_boundary = _section_boundary(source)
    landscape_boundary = _section_boundary(source)
    landscape = landscape_boundary.find("./w:pPr/w:sectPr", namespaces=NS)
    page_size = landscape.find(qn("w:pgSz"))
    if page_size is None:
        raise FormatMonographError("Landscape table requires explicit page size properties.")
    width, height = page_size.get(qn("w:w")), page_size.get(qn("w:h"))
    if width and height:
        page_size.set(qn("w:w"), height)
        page_size.set(qn("w:h"), width)
    page_size.set(qn("w:orient"), "landscape")
    position = list(body).index(table._tbl)
    body.insert(position, portrait_boundary)
    body.insert(position + 2, landscape_boundary)


def _table_matches_approved_cleanup_result(
    table: Any, table_index: int, structure_map: dict[str, Any]
) -> bool:
    cleanups = [
        entry
        for entry in structure_map.get("table_cell_cleanups", [])
        if entry.get("approved") and int(entry.get("table", -1)) == table_index
    ]
    if not cleanups:
        return False
    for entry in cleanups:
        row_index = int(entry["row"])
        cell_index = int(entry["cell"])
        if not 0 <= row_index < len(table.rows):
            return False
        if not 0 <= cell_index < len(table.rows[row_index].cells):
            return False
        cell = table.rows[row_index].cells[cell_index]
        if (
            text_sha256(cell.text) != entry.get("result_cell_text_sha256")
            or not cell.paragraphs
            or not _paragraph_has_payload(cell.paragraphs[0]._p)
        ):
            return False
    return True


def _apply_tables(document: Any, structure_map: dict[str, Any]) -> int:
    changed = 0
    for entry in structure_map.get("tables", []):
        if not entry.get("approved"):
            continue
        figure_panel = (
            entry.get("kind") == "layout"
            and entry.get("layout_purpose") == "figure_panel"
            and entry.get("visual", {}).get("approved")
        )
        if (
            has_semantic_structure_map(structure_map)
            and entry.get("kind") != "data"
            and not (
                structure_map.get("schema_version") in {"1.3", "1.4", "1.5"}
                and entry.get("kind") == "layout"
                and (entry.get("pagination_only") or figure_panel)
            )
        ):
            raise FormatMonographError(
                "Only tables approved with kind=data may receive data-table rules."
            )
        index = int(entry["table"])
        if not 0 <= index < len(document.tables):
            raise FormatMonographError(f"Structure-map table is out of range: {index}")
        table = document.tables[index]
        expected_table_hash = entry.get("table_text_sha256")
        if (
            expected_table_hash
            and _table_text_hash(table) != expected_table_hash
            and not _table_matches_approved_cleanup_result(
                table, index, structure_map
            )
        ):
            raise FormatMonographError(f"Structure-map table hash mismatch: {index}")
        first_row = "\u241f".join(cell.text for cell in table.rows[0].cells) if table.rows else ""
        if text_sha256(first_row) != entry["first_row_sha256"]:
            raise FormatMonographError(f"Structure-map table hash mismatch: {index}")
        repeat_rows = entry.get("repeat_header_rows", [])
        if not repeat_rows and entry.get("repeat_header"):
            repeat_rows = [0]
        if entry.get("repeat_caption_with_header"):
            repeat_rows = sorted({0, *[int(row) for row in repeat_rows], 1})
        for row_index in repeat_rows:
            row_index = int(row_index)
            if not 0 <= row_index < len(table.rows):
                raise FormatMonographError(
                    f"Approved repeat-header row is out of range: {index}:{row_index}"
                )
            _set_row_property(table.rows[row_index], "tblHeader", True)
        if entry.get("prevent_normal_row_split") or entry.get("keep_rows_together"):
            caption_row = entry.get("caption_row")
            for row_index, row in enumerate(table.rows):
                if caption_row is not None and row_index == int(caption_row):
                    continue
                _set_row_property(row, "cantSplit", True)
        visual = entry.get("visual", {})
        if visual.get("approved") and (figure_panel or entry.get("kind") == "data"):
            apply_table_properties(document, visual, [(table, entry)])
        if figure_panel:
            for row_index in entry.get("label_rows", []):
                row_index = int(row_index)
                if not 0 <= row_index < len(table.rows):
                    raise FormatMonographError(
                        "Figure-panel label row is out of range."
                    )
                for cell in _unique_row_cells(table.rows[row_index]):
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            paragraph.style = ensure_paragraph_style(
                                document, "Caption"
                            )
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if visual.get("approved") and visual.get("orientation") == "landscape":
            _wrap_table_landscape(document, table)
        changed += 1
    return changed


def _apply_table_cell_cleanups(document: Any, structure_map: dict[str, Any]) -> int:
    changed = 0
    for entry in structure_map.get("table_cell_cleanups", []):
        if not entry.get("approved"):
            continue
        table_index = int(entry["table"])
        row_index = int(entry["row"])
        cell_index = int(entry["cell"])
        if not 0 <= table_index < len(document.tables):
            raise FormatMonographError("Table-cell cleanup table is out of range.")
        table = document.tables[table_index]
        if not 0 <= row_index < len(table.rows):
            raise FormatMonographError("Table-cell cleanup row is out of range.")
        if not 0 <= cell_index < len(table.rows[row_index].cells):
            raise FormatMonographError("Table-cell cleanup cell is out of range.")
        cell = table.rows[row_index].cells[cell_index]
        count = int(entry["count"])
        current_cell_hash = text_sha256(cell.text)
        if (
            current_cell_hash == entry.get("result_cell_text_sha256")
            and cell.paragraphs
            and _paragraph_has_payload(cell.paragraphs[0]._p)
        ):
            continue
        if entry.get("table_text_sha256") != _table_text_hash(table):
            raise FormatMonographError("Table-cell cleanup source hash mismatch.")
        if current_cell_hash != entry.get("cell_text_sha256"):
            raise FormatMonographError("Table-cell cleanup cell hash mismatch.")
        if count >= len(cell.paragraphs):
            raise FormatMonographError(
                "Table-cell cleanup must retain at least one paragraph."
            )
        candidates = cell.paragraphs[:count]
        if any(_paragraph_has_payload(paragraph._p) for paragraph in candidates):
            raise FormatMonographError(
                "Table-cell cleanup can remove only leading empty paragraphs."
            )
        for paragraph in candidates:
            cell._tc.remove(paragraph._p)
            changed += 1
        if text_sha256(cell.text) != entry.get("result_cell_text_sha256"):
            raise FormatMonographError("Table-cell cleanup result hash mismatch.")
    return changed


def _final_section_is_empty(document: Any, boundary_index: int) -> bool:
    body = document.element.body
    boundary = document.paragraphs[boundary_index]._p
    children = list(body)
    start = children.index(boundary) + 1
    return not any(_paragraph_has_payload(element) for element in children[start:-1])


def _remove_final_empty_section(document: Any, entry: dict[str, Any]) -> None:
    locator = entry.get("previous_boundary_locator")
    if isinstance(locator, dict):
        boundary = resolve_paragraph_locator(document, locator)
        boundary_index = next(
            (
                index
                for index, paragraph in enumerate(document.paragraphs)
                if paragraph._p is boundary._p
            ),
            -1,
        )
        if boundary_index < 0:
            raise FormatMonographError(
                "Trailing-section boundary no longer belongs to the document."
            )
    else:
        boundary_index = int(entry["previous_boundary_paragraph"])
        if not 0 <= boundary_index < len(document.paragraphs):
            raise FormatMonographError("Trailing-section boundary is out of range.")
        boundary = document.paragraphs[boundary_index]
    if text_sha256(boundary.text) != entry["previous_boundary_sha256"]:
        raise FormatMonographError("Trailing-section boundary hash mismatch.")
    if not _final_section_is_empty(document, boundary_index):
        raise FormatMonographError("Approved trailing section is no longer empty.")

    p_pr = boundary._p.pPr
    previous = None if p_pr is None else p_pr.find(qn("w:sectPr"))
    final = document.element.body.sectPr
    if previous is None or final is None:
        raise FormatMonographError("Trailing-section boundary is missing section properties.")
    expected_section_hash = entry.get("section_properties_sha256")
    if expected_section_hash and _section_properties_sha256(final) != expected_section_hash:
        raise FormatMonographError("Trailing-section properties hash mismatch.")
    blocked = ("pgNumType", "titlePg")
    has_header_footer_refs = any(
        final.find(qn(f"w:{name}")) is not None
        for name in ("headerReference", "footerReference")
    )
    evidence = entry.get("evidence", {})
    if (
        any(final.find(qn(f"w:{name}")) is not None for name in blocked)
        or (has_header_footer_refs and evidence.get("header_footer_has_payload", True))
    ):
        raise FormatMonographError(
            "Trailing section has independent header, footer, or page-number settings."
        )

    body = document.element.body
    children = list(body)
    boundary_position = children.index(boundary._p)
    for child in list(children[boundary_position + 1 : -1]):
        body.remove(child)
    body.replace(final, copy.deepcopy(previous))
    p_pr.remove(previous)
    if not _paragraph_has_payload(boundary._p) and len(p_pr) == 0:
        body.remove(boundary._p)


def _apply_trailing_sections(document: Any, structure_map: dict[str, Any]) -> int:
    approved = [
        entry
        for entry in structure_map.get("trailing_empty_sections", [])
        if entry.get("approved_delete")
    ]
    approved.sort(key=lambda item: int(item["section"]), reverse=True)
    for entry in approved:
        if (
            has_semantic_structure_map(structure_map)
            and not entry.get("evidence", {}).get("safe_to_delete")
        ):
            raise FormatMonographError(
                "Approved trailing section lacks safe deletion evidence."
            )
        if int(entry["section"]) != len(document.sections) - 1:
            raise FormatMonographError(
                "Trailing empty sections must be approved and removed from the end inward."
            )
        _remove_final_empty_section(document, entry)
    return len(approved)


def _cleanup_orphan_header_footer_relationships(document: Any) -> int:
    referenced = {
        rel_id
        for sect_pr in document.element.xpath(".//w:sectPr")
        for rel_id in sect_pr.xpath(
            "./w:headerReference/@r:id | ./w:footerReference/@r:id"
        )
    }
    removable = [
        rel_id
        for rel_id, relationship in document.part.rels.items()
        if relationship.reltype.rsplit("/", 1)[-1] in {"header", "footer"}
        and rel_id not in referenced
    ]
    for rel_id in removable:
        document.part.drop_rel(rel_id)
    return len(removable)


def apply_structure_map(document: Any, structure_map: dict[str, Any]) -> list[dict[str, Any]]:
    trailing_targets = _apply_trailing_sections(document, structure_map)
    orphan_parts = _cleanup_orphan_header_footer_relationships(document)
    toc_targets = _apply_toc_ranges(document, structure_map)
    front_matter_targets = _apply_front_matter(document, structure_map)
    heading_targets = _apply_headings(document, structure_map)
    appendix_targets = _apply_appendices(document, structure_map)
    outline_targets = _apply_outline_cleanup(document, structure_map)
    toc_source_targets, toc_source_mode = _apply_toc_source_strategy(
        document, structure_map
    )
    table_targets = _apply_tables(document, structure_map)
    table_cell_cleanup_targets = _apply_table_cell_cleanups(document, structure_map)
    caption_targets = _apply_captions(document, structure_map)
    pagination_targets = _apply_pagination_groups(document, structure_map)
    spacing_targets = _apply_block_spacing(document, structure_map)
    pagination_sections = apply_pagination_sections(
        document,
        structure_map.get("pagination_sections", {}),
        resolve_paragraph_locator,
        front_matter=structure_map.get("front_matter", {}),
        replace_static_page_text=any(
            entry.get("approved_delete")
            and entry.get("evidence", {}).get("approved_derived_footer_only")
            for entry in structure_map.get("trailing_empty_sections", [])
        ),
    )
    image_targets = _apply_images(document, structure_map)
    image_visibility_targets = _apply_image_visibility(document, structure_map)
    caption_actions: dict[str, int] = {}
    if has_caption_actions_map(structure_map):
        for entry in structure_map.get("captions", []):
            if not entry.get("approved") or entry.get("action") not in {
                "replace_identifier",
                "convert_to_seq",
                "move_caption",
            }:
                continue
            action = str(entry["action"])
            caption_actions[action] = caption_actions.get(action, 0) + 1
    changes = [
        {"kind": "structure_toc", "targets": toc_targets},
        {"kind": "structure_front_matter", "targets": front_matter_targets},
        {"kind": "structure_headings", "targets": heading_targets},
        {"kind": "structure_appendices", "targets": appendix_targets},
        {"kind": "structure_outline_cleanup", "targets": outline_targets},
        {
            "kind": "structure_toc_source",
            "targets": toc_source_targets,
            "details": {"mode": toc_source_mode},
        },
        {"kind": "structure_tables", "targets": table_targets},
        {
            "kind": "structure_table_cell_cleanup",
            "targets": table_cell_cleanup_targets,
        },
        {
            "kind": "structure_captions",
            "targets": caption_targets,
            "actions": caption_actions,
        },
        {"kind": "structure_pagination", "targets": pagination_targets},
        {"kind": "structure_image_resize", "targets": image_targets},
        {
            "kind": "structure_image_visibility",
            "targets": image_visibility_targets,
        },
        {"kind": "structure_block_spacing", "targets": spacing_targets},
        {
            "kind": "structure_trailing_sections",
            "targets": trailing_targets,
        },
        {
            "kind": "structure_orphan_header_footer_cleanup",
            "targets": orphan_parts,
        },
        {
            "kind": "structure_pagination_sections",
            "targets": 1 if pagination_sections else 0,
            "details": pagination_sections,
        },
    ]
    return [
        change
        for change in changes
        if change["targets"]
        or (
            change["kind"] == "structure_toc_source"
            and change.get("details", {}).get("mode") is not None
        )
    ]


def _indent_conflicts(ind: Any | None) -> dict[str, str]:
    if ind is None:
        return {}
    return {
        attribute: value
        for attribute in STRUCTURAL_INDENT_ATTRIBUTES
        if (value := ind.get(qn(f"w:{attribute}"))) not in {None, "0"}
    }


def _direct_numbering_level(document: Any, paragraph: Any) -> Any | None:
    p_pr = paragraph._p.pPr
    num_pr = None if p_pr is None else p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        return None
    num_id = num_id_element.get(qn("w:val"))
    ilvl_element = num_pr.find(qn("w:ilvl"))
    ilvl = "0" if ilvl_element is None else ilvl_element.get(qn("w:val"), "0")
    root = document.part.numbering_part.element
    num = next(
        (item for item in root.findall(qn("w:num")) if item.get(qn("w:numId")) == num_id),
        None,
    )
    abstract_ref = None if num is None else num.find(qn("w:abstractNumId"))
    if abstract_ref is None:
        return None
    abstract_id = abstract_ref.get(qn("w:val"))
    abstract = next(
        (
            item
            for item in root.findall(qn("w:abstractNum"))
            if item.get(qn("w:abstractNumId")) == abstract_id
        ),
        None,
    )
    if abstract is None:
        return None
    return next(
        (
            item
            for item in abstract.findall(qn("w:lvl"))
            if item.get(qn("w:ilvl")) == ilvl
        ),
        None,
    )


def audit_structure_heading_operations(
    document: Any, structure_map: dict[str, Any]
) -> list[dict[str, Any]]:
    failures: list[dict[str, Any]] = []
    numbering_approved = bool(structure_map.get("numbering", {}).get("approved"))
    required_zero = (
        "left",
        "leftChars",
        "right",
        "rightChars",
        "firstLine",
        "firstLineChars",
    )
    for entry in structure_map.get("headings", []):
        if not entry.get("approved"):
            continue
        level = int(entry["level"])
        try:
            paragraph = (
                _verified_locator_paragraph(document, entry)
                if entry.get("locator")
                else _verified_paragraph(document, entry)
            )
        except FormatMonographError:
            normalized_hash = entry.get("normalized_text_sha256")
            candidates = [
                paragraph
                for paragraph in document.paragraphs
                if paragraph.style is not None
                and paragraph.style.name == f"Heading {level}"
                and normalized_hash
                and text_sha256(paragraph.text) == normalized_hash
            ]
            if len(candidates) != 1:
                failures.append(
                    {"level": level, "property": "heading_locator", "expected": "unique"}
                )
                continue
            paragraph = candidates[0]

        if paragraph.style is None or paragraph.style.name != f"Heading {level}":
            failures.append(
                {
                    "level": level,
                    "property": "heading_style",
                    "expected": f"Heading {level}",
                }
            )
            continue
        p_pr = paragraph._p.pPr
        direct_ind = None if p_pr is None else p_pr.find(qn("w:ind"))
        conflicts = _indent_conflicts(direct_ind)
        if conflicts:
            failures.append(
                {
                    "level": level,
                    "property": "heading_direct_indent",
                    "expected": 0,
                    "actual": conflicts,
                }
            )

        style_p_pr = paragraph.style.element.pPr
        style_ind = None if style_p_pr is None else style_p_pr.find(qn("w:ind"))
        style_values = {
            attribute: (
                None
                if style_ind is None
                else style_ind.get(qn(f"w:{attribute}"))
            )
            for attribute in required_zero
        }
        if _indent_conflicts(style_ind) or any(
            value != "0" for value in style_values.values()
        ):
            failures.append(
                {
                    "level": level,
                    "property": "heading_style_indent",
                    "expected": 0,
                    "actual": style_values,
                }
            )

        direct_num_pr = None if p_pr is None else p_pr.find(qn("w:numPr"))
        if numbering_approved and direct_num_pr is not None:
            failures.append(
                {
                    "level": level,
                    "property": "heading_direct_numbering",
                    "expected": "absent",
                }
            )
        if numbering_approved:
            numbering_level = _numbering_level_for_style(document, level - 1)
        elif direct_num_pr is not None:
            numbering_level = _direct_numbering_level(document, paragraph)
        else:
            numbering_level = None
        direct_num_id_element = (
            None if direct_num_pr is None else direct_num_pr.find(qn("w:numId"))
        )
        direct_num_id = (
            None
            if direct_num_id_element is None
            else direct_num_id_element.get(qn("w:val"))
        )
        if numbering_approved and numbering_level is None:
            failures.append(
                {
                    "level": level,
                    "property": "heading_numbering_level",
                    "expected": "resolved_style_linked_level",
                }
            )
        elif (
            not numbering_approved
            and direct_num_pr is not None
            and direct_num_id not in {None, "0"}
            and numbering_level is None
        ):
            failures.append(
                {
                    "level": level,
                    "property": "heading_direct_numbering_requires_qa",
                    "expected": "resolvable_zero_indent_numbering_or_approval",
                }
            )
        if numbering_level is not None:
            numbering_p_pr = numbering_level.find(qn("w:pPr"))
            numbering_ind = (
                None
                if numbering_p_pr is None
                else numbering_p_pr.find(qn("w:ind"))
            )
            numbering_conflicts = _indent_conflicts(numbering_ind)
            numbering_values = {
                attribute: (
                    None
                    if numbering_ind is None
                    else numbering_ind.get(qn(f"w:{attribute}"))
                )
                for attribute in required_zero
            }
            tabs = (
                []
                if numbering_p_pr is None
                else numbering_p_pr.findall(qn("w:tabs"))
            )
            if numbering_conflicts or tabs or (
                numbering_approved
                and any(value != "0" for value in numbering_values.values())
            ):
                failures.append(
                    {
                        "level": level,
                        "property": (
                            "heading_numbering_indent"
                            if numbering_approved
                            else "heading_direct_numbering_requires_qa"
                        ),
                        "expected": {"indent": 0, "tabs": 0},
                        "actual": {
                            "indent": numbering_values,
                            "conflicts": numbering_conflicts,
                            "tabs": len(tabs),
                        },
                    }
                )
    return failures


def audit_structure_table_operations(
    document: Any, structure_map: dict[str, Any]
) -> list[dict[str, Any]]:
    def effective_alignment(paragraph: Any) -> Any:
        alignment = paragraph.alignment
        style = paragraph.style
        while alignment is None and style is not None:
            alignment = style.paragraph_format.alignment
            style = style.base_style
        return alignment

    failures: list[dict[str, Any]] = []
    for entry in structure_map.get("tables", []):
        visual = entry.get("visual", {})
        if not (
            entry.get("approved")
            and entry.get("kind") == "layout"
            and entry.get("layout_purpose") == "figure_panel"
            and visual.get("approved")
        ):
            continue
        table_index = int(entry["table"])
        if not 0 <= table_index < len(document.tables):
            failures.append({"table": table_index, "reason": "table_out_of_range"})
            continue
        table = document.tables[table_index]
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        borderless = borders is not None and all(
            (borders.find(qn(f"w:{name}")) is not None)
            and borders.find(qn(f"w:{name}")).get(qn("w:val")) == "nil"
            for name in ("top", "left", "bottom", "right", "insideH", "insideV")
        )
        cell_overrides = table._tbl.xpath(
            ".//w:tcBorders/*[not(@w:val='nil') and not(@w:val='none')]"
        )
        label_paragraphs = [
            paragraph
            for row_index in entry.get("label_rows", [])
            if 0 <= int(row_index) < len(table.rows)
            for cell in table.rows[int(row_index)].cells
            for paragraph in cell.paragraphs
            if paragraph.text.strip()
        ]
        checks = {
            "borderless": borderless,
            "no_cell_border_overrides": not cell_overrides,
            "no_text_wrapping": table._tbl.tblPr.find(qn("w:tblpPr")) is None,
            "table_centered": table.alignment == WD_TABLE_ALIGNMENT.CENTER,
            "cells_vertically_centered": all(
                cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for row in table.rows
                for cell in row.cells
            ),
            "labels_present": bool(label_paragraphs),
            "labels_centered_and_styled": bool(label_paragraphs)
            and all(
                effective_alignment(paragraph) == WD_ALIGN_PARAGRAPH.CENTER
                and paragraph.style is not None
                for paragraph in label_paragraphs
            ),
        }
        if not all(checks.values()):
            failures.append(
                {
                    "table": table_index,
                    "reason": "figure_panel_visual_mismatch",
                    "checks": checks,
                }
            )
    for entry in structure_map.get("table_cell_cleanups", []):
        if not entry.get("approved"):
            continue
        table_index = int(entry["table"])
        row_index = int(entry["row"])
        cell_index = int(entry["cell"])
        try:
            cell = document.tables[table_index].rows[row_index].cells[cell_index]
        except IndexError:
            failures.append(
                {"table": table_index, "reason": "cleanup_target_out_of_range"}
            )
            continue
        if not cell.paragraphs or not _paragraph_has_payload(cell.paragraphs[0]._p):
            failures.append(
                {"table": table_index, "reason": "leading_empty_paragraph_remains"}
            )
    return failures


def _approved_indexes(structure_map: dict[str, Any], key: str) -> dict[int, dict[str, Any]]:
    return {
        int(entry["paragraph"]): entry
        for entry in structure_map.get(key, [])
        if entry.get("approved") and "paragraph" in entry
    }


def _normalize_manual_identifier(
    value: str, entries: list[dict[str, Any]]
) -> str:
    for entry in entries:
        span = entry["identifier_span"]
        start, end = int(span["start"]), int(span["end"])
        if text_sha256(value) == entry["text_sha256"]:
            if (
                text_sha256(value[:start]) == entry["identifier_prefix_sha256"]
                and text_sha256(value[start:end]) == entry["identifier_sha256"]
                and text_sha256(value[end:]) == entry["identifier_suffix_sha256"]
            ):
                return value[:start] + "[[CAPTION_IDENTIFIER]]" + value[end:]
        replacement = str(entry["replacement_identifier"])
        replacement_end = start + len(replacement)
        if (
            value[start:replacement_end] == replacement
            and text_sha256(value[:start]) == entry["identifier_prefix_sha256"]
            and text_sha256(value[replacement_end:])
            == entry["identifier_suffix_sha256"]
        ):
            return value[:start] + "[[CAPTION_IDENTIFIER]]" + value[replacement_end:]
    return value


def audit_caption_identifier_replacements(
    original_path: Path, formatted_path: Path, structure_map: dict[str, Any]
) -> list[dict[str, Any]]:
    if not has_caption_actions_map(structure_map):
        return []
    from _common import load_document

    original = load_document(original_path)
    formatted = load_document(formatted_path)
    results = []
    for entry in structure_map.get("captions", []):
        if not entry.get("approved") or entry.get("action") != "replace_identifier":
            continue
        source_locator = copy.deepcopy(entry["locator"])
        if source_locator.get("kind") == "body_paragraph" and entry.get("text_sha256"):
            source_locator.setdefault("text_sha256", entry["text_sha256"])
        source = resolve_paragraph_locator(original, source_locator)
        target = _resolve_replaced_caption_target(formatted, entry)
        span = entry["identifier_span"]
        start, end = int(span["start"]), int(span["end"])
        title_start = int(entry["title_span_start"])
        source_ok = (
            text_sha256(source.text) == entry["text_sha256"]
            and text_sha256(source.text[start:end]) == entry["identifier_sha256"]
            and text_sha256(source.text[title_start:]) == entry["title_text_sha256"]
        )
        expected = (
            source.text[:start]
            + str(entry["replacement_identifier"])
            + source.text[end:]
        )
        replacement_end = start + len(str(entry["replacement_identifier"]))
        title_shift = replacement_end - end
        title_preserved = text_sha256(target.text[title_start + title_shift :]) == entry[
            "title_text_sha256"
        ]
        passed = source_ok and target.text == expected and title_preserved
        results.append(
            {
                "locator": entry["locator"],
                "status": "pass" if passed else "fail",
                "identifier_changed_as_approved": target.text == expected,
                "title_preserved": title_preserved,
            }
        )
    return results


def _toc_field_paragraphs(root: etree._Element) -> set[Any]:
    bodies = root.xpath("/w:document/w:body", namespaces=NS)
    if not bodies:
        return set()
    result: set[Any] = set()
    stack: list[dict[str, Any]] = []
    for paragraph in bodies[0].xpath("./w:p", namespaces=NS):
        if paragraph.xpath(
            './/w:fldSimple[starts-with(translate(normalize-space(@w:instr), "toc", "TOC"), "TOC ")]',
            namespaces=NS,
        ):
            result.add(paragraph)
        if any(item.get("is_toc") and item.get("result") for item in stack):
            result.add(paragraph)
        for element in paragraph.iter():
            if element.tag == qn("w:fldChar"):
                kind = element.get(qn("w:fldCharType"))
                if kind == "begin":
                    stack.append({"parts": [], "is_toc": False, "result": False})
                elif kind == "separate" and stack:
                    instruction = "".join(stack[-1]["parts"]).strip().upper()
                    stack[-1]["is_toc"] = instruction.startswith("TOC ")
                    stack[-1]["result"] = True
                    if stack[-1]["is_toc"]:
                        result.add(paragraph)
                elif kind == "end" and stack:
                    if any(item.get("is_toc") and item.get("result") for item in stack):
                        result.add(paragraph)
                    stack.pop()
            elif element.tag == qn("w:instrText") and stack and not stack[-1]["result"]:
                stack[-1]["parts"].append(element.text or "")
            if any(item.get("is_toc") and item.get("result") for item in stack):
                result.add(paragraph)
    return result


def _legacy_toc_empty_anchors(
    root: etree._Element,
    toc_field_paragraphs: set[Any],
    structure_map: dict[str, Any],
) -> set[Any]:
    """Recognize empty anchors left by pre-1.3 static-TOC migration."""
    if structure_map.get("schema_version") in {"1.3", "1.4", "1.5"}:
        return set()
    if not toc_field_paragraphs:
        return set()
    approved_lengths = [
        int(entry["end_paragraph"]) - int(entry["start_paragraph"])
        for entry in structure_map.get("toc_ranges", [])
        if entry.get("approved")
    ]
    remaining = max(approved_lengths, default=0)
    if remaining <= 0:
        return set()
    body = root.find(qn("w:body"))
    if body is None:
        return set()
    children = list(body)
    field_positions = [
        index for index, child in enumerate(children) if child in toc_field_paragraphs
    ]
    if not field_positions:
        return set()
    result = set()
    for child in children[max(field_positions) + 1 :]:
        if remaining <= 0 or child.tag != qn("w:p"):
            break
        value = FIELD_MARKER_PATTERN.sub(
            "", _paragraph_text_without_field_results(child)
        )
        if value:
            break
        result.add(child)
        remaining -= 1
    return result


def _pagination_boundary_paragraphs(
    root: etree._Element, structure_map: dict[str, Any]
) -> set[Any]:
    pagination = structure_map.get("pagination_sections", {})
    if structure_map.get("schema_version") not in {"1.4", "1.5"}:
        return set()
    result = set()
    has_landscape_table = any(
        entry.get("approved")
        and entry.get("visual", {}).get("approved")
        and entry.get("visual", {}).get("orientation") == "landscape"
        for entry in structure_map.get("tables", [])
    )
    if has_landscape_table:
        result.update(
            paragraph
            for paragraph in root.xpath("/w:document/w:body/w:p", namespaces=NS)
            if paragraph.find("./w:pPr/w:sectPr", namespaces=NS) is not None
            and not _paragraph_text_without_field_results(paragraph)
        )
    if not pagination.get("approved"):
        return result
    if structure_map.get("front_matter", {}).get("approved"):
        toc_style = TOC_HEADING_STYLE.replace(" ", "")
        toc_headings = root.xpath(
            f"/w:document/w:body/w:p[w:pPr/w:pStyle[@w:val='{toc_style}']]",
            namespaces=NS,
        )
        if len(toc_headings) == 1:
            previous = toc_headings[0].getprevious()
            if (
                previous is not None
                and previous.tag == qn("w:p")
                and previous.find("./w:pPr/w:sectPr", namespaces=NS) is not None
                and not _paragraph_text_without_field_results(previous)
            ):
                result.add(previous)
    body_locator = pagination.get("body_start", {})
    expected = body_locator.get("text_sha256")
    if not expected:
        return result
    expected_hashes = {expected}
    heading = next(
        (
            entry
            for entry in structure_map.get("headings", [])
            if entry.get("approved")
            and entry.get("locator", {}).get("text_sha256") == expected
        ),
        None,
    )
    if heading and heading.get("normalized_text_sha256"):
        expected_hashes.add(heading["normalized_text_sha256"])
    paragraphs = root.xpath("/w:document/w:body/w:p", namespaces=NS)
    matches = [
        paragraph
        for paragraph in paragraphs
        if text_sha256(_paragraph_text_without_field_results(paragraph))
        in expected_hashes
    ]
    if len(matches) != 1:
        return result
    previous = matches[0].getprevious()
    if (
        previous is not None
        and previous.tag == qn("w:p")
        and previous.find("./w:pPr/w:sectPr", namespaces=NS) is not None
        and not _paragraph_text_without_field_results(previous)
    ):
        result.add(previous)
    return result


def _approved_tail_paragraphs(
    root: etree._Element, structure_map: dict[str, Any]
) -> set[Any]:
    approved = [
        entry
        for entry in structure_map.get("trailing_empty_sections", [])
        if entry.get("approved_delete")
    ]
    if not approved:
        return set()
    body = root.find(qn("w:body"))
    if body is None:
        return set()
    direct_paragraphs = [child for child in body if child.tag == qn("w:p")]
    first_removed_positions = []
    children = list(body)
    for entry in approved:
        expected_hash = entry.get("previous_boundary_sha256")
        boundary_is_empty = expected_hash == text_sha256("")
        matching = [
            paragraph
            for paragraph in direct_paragraphs
            if expected_hash and text_sha256(_paragraph_text_without_field_results(paragraph))
            == expected_hash
        ]
        if len(matching) == 1:
            boundary_paragraph = matching[0]
        else:
            boundary = int(entry["previous_boundary_paragraph"])
            if not 0 <= boundary < len(direct_paragraphs):
                continue
            candidate = direct_paragraphs[boundary]
            if expected_hash and text_sha256(
                _paragraph_text_without_field_results(candidate)
            ) != expected_hash:
                continue
            boundary_paragraph = candidate
        if boundary_is_empty and any(
            item.get("previous_boundary_sha256") != text_sha256("")
            for item in approved
        ):
            continue
        boundary_position = children.index(boundary_paragraph)
        first_removed_positions.append(
            boundary_position if boundary_is_empty else boundary_position + 1
        )
    if not first_removed_positions:
        return set()
    first_removed = min(first_removed_positions)
    return {
        paragraph
        for child in children[first_removed:]
        for paragraph in child.xpath(".//w:p | self::w:p", namespaces=NS)
    }


def _approved_deleted_header_footer_parts(
    package: zipfile.ZipFile,
    document_root: etree._Element,
    structure_map: dict[str, Any],
) -> set[str]:
    deleted_sections = {
        int(entry["section"])
        for entry in structure_map.get("trailing_empty_sections", [])
        if entry.get("approved_delete")
    }
    if not deleted_sections:
        return set()
    sections = document_root.xpath(".//w:sectPr", namespaces=NS)
    deleted_ids: set[str] = set()
    retained_ids: set[str] = set()
    relationship_attribute = qn("r:id")
    for index, section in enumerate(sections):
        target = deleted_ids if index in deleted_sections else retained_ids
        for reference in section.xpath(
            "./w:headerReference | ./w:footerReference", namespaces=NS
        ):
            relationship_id = reference.get(relationship_attribute)
            if relationship_id:
                target.add(relationship_id)
    removable_ids = deleted_ids - retained_ids
    if not removable_ids:
        return set()
    relationships_name = "word/_rels/document.xml.rels"
    if relationships_name not in package.namelist():
        return set()
    relationships = etree.fromstring(package.read(relationships_name))
    result = set()
    for relationship in relationships:
        if relationship.get("Id") not in removable_ids:
            continue
        target = relationship.get("Target", "")
        normalized = posixpath.normpath(posixpath.join("word", target))
        if re.fullmatch(r"word/(?:header|footer)\d+\.xml", normalized):
            result.add(normalized)
    return result


def _approved_table_cell_cleanup_paragraphs(
    root: etree._Element, structure_map: dict[str, Any]
) -> set[Any]:
    tables = root.xpath("/w:document/w:body/w:tbl", namespaces=NS)
    ignored: set[Any] = set()
    for entry in structure_map.get("table_cell_cleanups", []):
        if not entry.get("approved"):
            continue
        table_index = int(entry["table"])
        row_index = int(entry["row"])
        cell_index = int(entry["cell"])
        if not 0 <= table_index < len(tables):
            continue
        rows = tables[table_index].xpath("./w:tr", namespaces=NS)
        if not 0 <= row_index < len(rows):
            continue
        cells = rows[row_index].xpath("./w:tc", namespaces=NS)
        if not 0 <= cell_index < len(cells):
            continue
        paragraphs = cells[cell_index].xpath("./w:p", namespaces=NS)
        for paragraph in paragraphs[: int(entry["count"])]:
            text = "".join(
                paragraph.xpath(".//w:t/text()", namespaces=NS)
            ).strip()
            protected = paragraph.xpath(
                ".//w:tbl | .//w:drawing | .//w:object | .//w:pict | "
                ".//w:bookmarkStart | .//w:commentRangeStart | "
                ".//w:footnoteReference | .//w:endnoteReference | .//w:fldChar",
                namespaces=NS,
            )
            if not text and not protected:
                ignored.add(paragraph)
    return ignored


def structure_content_inventory(
    path: Path, structure_map: dict[str, Any]
) -> dict[str, list[str]]:
    toc_indexes = {
        index
        for entry in structure_map.get("toc_ranges", [])
        if entry.get("approved")
        for index in range(int(entry["start_paragraph"]), int(entry["end_paragraph"]) + 1)
    }
    heading_entries = [
        entry for entry in structure_map.get("headings", []) if entry.get("approved")
    ]
    caption_entries = [
        entry
        for entry in structure_map.get("captions", [])
        if entry.get("approved")
        and (
            not has_caption_actions_map(structure_map)
            or entry.get("action") == "convert_to_seq"
        )
    ]
    identifier_replacements = [
        entry
        for entry in structure_map.get("captions", [])
        if entry.get("approved") and entry.get("action") == "replace_identifier"
    ]
    migrated_caption_hashes = {
        entry["text_sha256"]
        for entry in structure_map.get("captions", [])
        if entry.get("approved")
        and entry.get("migrate_outside_table")
        and entry.get("locator", {}).get("kind") == "table_cell_paragraph"
    }
    manual_migrated_caption_hashes = {
        entry["text_sha256"]
        for entry in structure_map.get("captions", [])
        if entry.get("approved")
        and entry.get("action") == "move_caption"
        and entry.get("migrate_outside_table")
    }
    has_migrated_captions = bool(
        migrated_caption_hashes - manual_migrated_caption_hashes
    )
    approved_front_matter = structure_map.get("front_matter", {}).get("approved")
    approved_block_spacing = structure_map.get("block_spacing", {}).get("approved")
    toc_heading_style_id = TOC_HEADING_STYLE.replace(" ", "")
    block_spacer_style_id = BLOCK_SPACER_STYLE.replace(" ", "")
    result: dict[str, list[str]] = {}
    with zipfile.ZipFile(path) as package:
        document_root = etree.fromstring(package.read("word/document.xml"))
        ignored_parts = _approved_deleted_header_footer_parts(
            package, document_root, structure_map
        )
        for name in sorted(package.namelist()):
            if not CONTENT_PART.match(name) or name in ignored_parts:
                continue
            root = document_root if name == "word/document.xml" else etree.fromstring(
                package.read(name)
            )
            toc_field_paragraphs = (
                _toc_field_paragraphs(root) if name == "word/document.xml" else set()
            )
            legacy_toc_anchors = (
                _legacy_toc_empty_anchors(root, toc_field_paragraphs, structure_map)
                if name == "word/document.xml"
                else set()
            )
            approved_tail_paragraphs = (
                _approved_tail_paragraphs(root, structure_map)
                if name == "word/document.xml"
                else set()
            )
            pagination_boundaries = (
                _pagination_boundary_paragraphs(root, structure_map)
                if name == "word/document.xml"
                else set()
            )
            approved_cell_cleanup_paragraphs = (
                _approved_table_cell_cleanup_paragraphs(root, structure_map)
                if name == "word/document.xml"
                else set()
            )
            values = []
            body_index = -1
            for paragraph in root.xpath(".//w:p", namespaces=NS):
                if paragraph in approved_cell_cleanup_paragraphs:
                    continue
                style_ids = paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
                if approved_front_matter and style_ids == [toc_heading_style_id]:
                    continue
                if approved_block_spacing and style_ids == [block_spacer_style_id]:
                    continue
                direct_body_paragraph = (
                    name == "word/document.xml"
                    and paragraph.getparent() is not None
                    and paragraph.getparent().tag == qn("w:body")
                )
                if direct_body_paragraph:
                    body_index += 1
                current_index = body_index if direct_body_paragraph else None
                if paragraph in approved_tail_paragraphs or paragraph in pagination_boundaries:
                    continue
                if paragraph in toc_field_paragraphs or paragraph in legacy_toc_anchors:
                    continue
                if (
                    current_index is not None
                    and not toc_field_paragraphs
                    and current_index in toc_indexes
                ):
                    continue
                value = _paragraph_text_without_field_results(paragraph)
                value = FIELD_MARKER_PATTERN.sub("", value)
                value = _normalize_manual_identifier(value, identifier_replacements)
                original_caption = CAPTION_PATTERN.match(value)
                if text_sha256(value) in manual_migrated_caption_hashes:
                    value = "[[MOVED_MANUAL_CAPTION]]"
                elif text_sha256(value) in migrated_caption_hashes and original_caption:
                    value = f"{original_caption.group(1)} {original_caption.group(4)}"
                elif (has_migrated_captions or caption_entries) and direct_body_paragraph:
                    styles = paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
                    generated = re.match(r"^\s*(图|表)\s*[-－—–]\s*(.*)$", value)
                    if styles == ["Caption"] and generated:
                        value = f"{generated.group(1)} {generated.group(2)}"
                heading_entry = next(
                    (
                        entry
                        for entry in heading_entries
                        if text_sha256(value) == entry.get("text_sha256")
                    ),
                    None,
                )
                if heading_entry is not None:
                    match = _heading_prefix_pattern(int(heading_entry["level"])).match(value)
                    if match:
                        value = value[match.end() :]
                else:
                    styles = paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
                    style_match = (
                        re.fullmatch(r"Heading([1-4])", styles[0]) if styles else None
                    )
                    if style_match:
                        match = _heading_prefix_pattern(int(style_match.group(1))).match(
                            value
                        )
                        if match:
                            value = value[match.end() :]
                if any(
                    text_sha256(value) == entry.get("text_sha256")
                    for entry in caption_entries
                ):
                    match = CAPTION_PATTERN.match(value)
                    if match:
                        value = f"{match.group(1)} {match.group(4)}"
                    else:
                        value = re.sub(r"^(图|表)\s+[-－—–]?\s*", r"\1 ", value)
                values.append(value)
            if name == "word/document.xml" or any(values):
                result[name] = values
    approved_derived_footer_only = bool(
        structure_map.get("pagination_sections", {}).get("approved")
        and any(
            entry.get("approved_delete")
            and entry.get("evidence", {}).get("approved_derived_footer_only")
            for entry in structure_map.get("trailing_empty_sections", [])
        )
    )
    canonical: dict[str, list[str]] = {}
    headers: list[str] = []
    footers: list[str] = []
    for name, values in result.items():
        if re.fullmatch(r"word/header\d+\.xml", name):
            headers.extend(values)
        elif re.fullmatch(r"word/footer\d+\.xml", name):
            if not approved_derived_footer_only:
                footers.extend(values)
        else:
            canonical[name] = values
    canonical["word/_headers_by_content"] = sorted(headers)
    canonical["word/_footers_by_content"] = sorted(footers)
    return canonical


def structure_content_fingerprint(path: Path, structure_map: dict[str, Any]) -> str:
    result = structure_content_inventory(path, structure_map)
    encoded = json.dumps(result, ensure_ascii=False, separators=(",", ":"), sort_keys=True)
    return hashlib.sha256(encoded.encode("utf-8")).hexdigest()
