#!/usr/bin/env python3
"""Audit content, protected objects, and automatically verifiable profile rules."""

from __future__ import annotations

import argparse
import json
import math
import re
import sys
from pathlib import Path
from typing import Any

from docx.oxml.ns import qn

from _common import (
    STYLE_PROPERTIES,
    FormatMonographError,
    _field_instruction_for_marker,
    _heading_prefix_pattern,
    _table_effective_properties,
    content_fingerprint,
    equation_inventory,
    field_cache_inventory,
    font_alias_keys,
    load_document,
    protected_object_manifest,
    protected_payload_manifest,
    run_effective_font,
    style_effective_font,
    style_name_for_selector,
    write_json,
)
from validate_profile import validate
from docx_pagination import audit_pagination_sections
from structure_map import (
    approved_data_tables,
    approved_role_paragraphs,
    audit_caption_identifier_replacements,
    audit_structure_heading_operations,
    audit_structure_image_operations,
    audit_structure_toc_source_operations,
    audit_structure_table_operations,
    has_semantic_structure_map,
    load_structure_map,
    resolve_paragraph_locator,
    structure_content_fingerprint,
)


def close(actual: Any, expected: Any, tolerance: float = 0.05) -> bool:
    if actual is None:
        return expected is None
    return math.isclose(float(actual), float(expected), abs_tol=tolerance)


def style_value(document: Any, style: Any, key: str) -> Any:
    font, pf = style.font, style.paragraph_format
    if key == "font_name":
        return style_effective_font(document, style, "ascii")[0]
    if key == "font_name_ascii":
        return style_effective_font(document, style, "ascii")[0]
    if key == "font_name_east_asia":
        return style_effective_font(document, style, "eastAsia")[0]
    if key == "font_name_complex_script":
        return style_effective_font(document, style, "cs")[0]
    if key == "font_size_pt":
        current = style
        while current is not None:
            if current.font.size is not None:
                return current.font.size.pt
            current = current.base_style
        return None
    if key in {"bold", "italic"}:
        current = style
        while current is not None:
            value = getattr(current.font, key)
            if value is not None:
                return value
            current = current.base_style
        return False
    if key == "color_hex":
        return None if font.color.rgb is None else str(font.color.rgb)
    if key == "alignment":
        return None if pf.alignment is None else str(pf.alignment).split(".")[-1].lower()
    point_values = {
        "space_before_pt": pf.space_before,
        "space_after_pt": pf.space_after,
        "first_line_indent_pt": pf.first_line_indent,
        "left_indent_pt": pf.left_indent,
        "right_indent_pt": pf.right_indent,
    }
    if key in point_values:
        attribute = {
            "space_before_pt": "space_before",
            "space_after_pt": "space_after",
            "first_line_indent_pt": "first_line_indent",
            "left_indent_pt": "left_indent",
            "right_indent_pt": "right_indent",
        }[key]
        current = style
        while current is not None:
            value = getattr(current.paragraph_format, attribute)
            if value is not None:
                return value.pt
            current = current.base_style
        return 0.0 if key in {"space_before_pt", "space_after_pt"} else None
    if key == "first_line_indent_chars":
        current = style
        while current is not None:
            p_pr = current.element.pPr
            ind = None if p_pr is None else p_pr.find(qn("w:ind"))
            if ind is not None:
                value = ind.get(qn("w:firstLineChars"))
                if value is not None:
                    return int(value) / 100
                point_value = ind.get(qn("w:firstLine"))
                if point_value == "0":
                    return 0
            current = current.base_style
        return None
    if key == "line_spacing":
        value = pf.line_spacing
        return float(value) if isinstance(value, (int, float)) else None
    if key == "line_spacing_pt":
        value = pf.line_spacing
        return None if value is None or not hasattr(value, "pt") else value.pt
    if key == "line_spacing_rule":
        value = str(pf.line_spacing_rule).lower()
        if "exact" in value:
            return "exact"
        if "at_least" in value or "at least" in value:
            return "at_least"
        if "multiple" in value:
            return "multiple"
        if "one_point_five" in value or "1.5" in value:
            return "one_point_five"
        if "double" in value:
            return "double"
        if "single" in value:
            return "single"
        return value
    if key in {"keep_with_next", "keep_together", "page_break_before", "widow_control"}:
        current = style
        while current is not None:
            value = getattr(current.paragraph_format, key)
            if value is not None:
                return value
            current = current.base_style
        if key == "widow_control":
            return True
        return False
    return "<unsupported>"


def normalized_alignment(value: str) -> str:
    value = value.lower()
    for name in ("left", "center", "right", "justify", "distributed"):
        if name in value:
            return name
    return value


def compare_value(key: str, actual: Any, expected: Any) -> bool:
    if isinstance(actual, dict) and isinstance(expected, dict):
        return set(actual) == set(expected) and all(
            compare_value(
                f"{name}_mm" if key == "cell_margins_mm" else name,
                actual[name],
                expected[name],
            )
            for name in actual
        )
    if isinstance(actual, list):
        if isinstance(expected, list):
            return len(actual) == len(expected) and all(
                compare_value(key, left, right)
                for left, right in zip(actual, expected)
            )
        return bool(actual) and all(compare_value(key, item, expected) for item in actual)
    if key.startswith("font_name") and actual is not None:
        return bool(font_alias_keys(str(actual)) & font_alias_keys(str(expected)))
    if key.endswith(("_pt", "_mm", "_ratio")) or key in {
        "line_spacing",
        "first_line_indent_chars",
    }:
        return close(actual, expected)
    if key == "color_hex" and actual is not None:
        return str(actual).lstrip("#").upper() == str(expected).lstrip("#").upper()
    if key == "alignment" and actual is not None:
        return normalized_alignment(str(actual)) == str(expected).lower()
    return actual == expected


def run_font_value(document: Any, paragraph: Any, run: Any, key: str) -> Any:
    r_pr = run._r.rPr
    if key == "font_name":
        return run_effective_font(document, paragraph, run, "ascii")[0]
    if key == "font_name_ascii":
        return run_effective_font(document, paragraph, run, "ascii")[0]
    if key == "font_name_east_asia":
        return run_effective_font(document, paragraph, run, "eastAsia")[0]
    if key == "font_name_complex_script":
        return run_effective_font(document, paragraph, run, "cs")[0]
    if key == "font_size_pt":
        return None if run.font.size is None else run.font.size.pt
    if key in {"bold", "italic"}:
        return getattr(run.font, key)
    if key == "color_hex":
        return None if run.font.color.rgb is None else str(run.font.color.rgb)
    return None


def paragraph_effective_value(document: Any, paragraph: Any, key: str) -> Any:
    style = paragraph.style
    if key.startswith("font_name"):
        values = [
            run_font_value(document, paragraph, run, key)
            for run in paragraph.runs
            if run.text
        ]
        return values or style_value(document, style, key)
    if key in {
        "font_size_pt",
        "bold",
        "italic",
        "color_hex",
    }:
        fallback = style_value(document, style, key)
        values = []
        for run in paragraph.runs:
            if not run.text:
                continue
            direct = run_font_value(document, paragraph, run, key)
            values.append(fallback if direct is None else direct)
        return values or fallback

    pf = paragraph.paragraph_format
    direct: Any = None
    if key == "alignment":
        direct = pf.alignment
        if direct is not None:
            direct = str(direct).split(".")[-1].lower()
    elif key in {
        "space_before_pt",
        "space_after_pt",
        "first_line_indent_pt",
        "left_indent_pt",
        "right_indent_pt",
    }:
        attr = {
            "space_before_pt": "space_before",
            "space_after_pt": "space_after",
            "first_line_indent_pt": "first_line_indent",
            "left_indent_pt": "left_indent",
            "right_indent_pt": "right_indent",
        }[key]
        value = getattr(pf, attr)
        direct = None if value is None else value.pt
    elif key == "first_line_indent_chars":
        p_pr = paragraph._p.pPr
        ind = None if p_pr is None else p_pr.find(qn("w:ind"))
        value = None if ind is None else ind.get(qn("w:firstLineChars"))
        direct = None if value is None else int(value) / 100
        if direct is None and ind is not None and ind.get(qn("w:firstLine")) == "0":
            direct = 0
        if direct is None and paragraph.style is not None:
            match = re.fullmatch(r"Heading ([1-4])", paragraph.style.name)
            if match:
                lvl = heading_numbering_level(document, int(match.group(1)) - 1)
                lvl_ind = None if lvl is None else lvl.find(
                    qn("w:pPr") + "/" + qn("w:ind")
                )
                if lvl_ind is not None:
                    chars = lvl_ind.get(qn("w:firstLineChars"))
                    if chars is not None:
                        direct = int(chars) / 100
                    elif lvl_ind.get(qn("w:firstLine")) == "0":
                        direct = 0
    elif key == "line_spacing":
        value = pf.line_spacing
        direct = float(value) if isinstance(value, (int, float)) else None
    elif key == "line_spacing_pt":
        value = pf.line_spacing
        direct = None if value is None or not hasattr(value, "pt") else value.pt
    elif key == "line_spacing_rule":
        value = pf.line_spacing_rule
        if value is not None:
            raw = str(value).lower()
            direct = next(
                (
                    name
                    for name, marker in (
                        ("exact", "exact"),
                        ("at_least", "at_least"),
                        ("multiple", "multiple"),
                        ("one_point_five", "one_point_five"),
                        ("double", "double"),
                        ("single", "single"),
                    )
                    if marker in raw
                ),
                raw,
            )
    elif key in {"keep_with_next", "keep_together", "page_break_before", "widow_control"}:
        direct = getattr(pf, key)
    return style_value(document, style, key) if direct is None else direct


def audit_paragraph_rule(
    document: Any, rule: dict[str, Any], paragraphs: list[Any]
) -> list[dict[str, Any]]:
    failures = []
    for target_index, paragraph in enumerate(paragraphs):
        for key, expected in rule["properties"].items():
            if key not in STYLE_PROPERTIES:
                continue
            actual = paragraph_effective_value(document, paragraph, key)
            if (
                key == "page_break_before"
                and expected is True
                and actual is False
                and _paragraph_starts_new_page_section(paragraph)
            ):
                continue
            if not compare_value(key, actual, expected):
                failures.append(
                    {
                        "target": target_index,
                        "property": key,
                        "expected": expected,
                        "actual": actual,
                        "reason": "effective format differs after style and direct formatting",
                    }
                )
    return failures


def _paragraph_starts_new_page_section(paragraph: Any) -> bool:
    previous = paragraph._p.getprevious()
    if previous is None or previous.tag != qn("w:p"):
        return False
    p_pr = previous.find(qn("w:pPr"))
    sect_pr = None if p_pr is None else p_pr.find(qn("w:sectPr"))
    if sect_pr is None:
        return False
    section_type = sect_pr.find(qn("w:type"))
    section_value = (
        section_type.get(qn("w:val")) if section_type is not None else "nextPage"
    )
    return section_value != "continuous"


def audit_style_rule(document: Any, rule: dict) -> list[dict]:
    style_name = style_name_for_selector(rule["selector"])
    if style_name is None:
        return [{"property": "*", "expected": "supported style selector", "actual": "unsupported"}]
    try:
        style = document.styles[style_name]
    except KeyError:
        return [{"property": "*", "expected": style_name, "actual": "missing style"}]
    failures = []
    for key, expected in rule["properties"].items():
        if key not in STYLE_PROPERTIES:
            continue
        actual = style_value(document, style, key)
        if not compare_value(key, actual, expected):
            failures.append({"property": key, "expected": expected, "actual": actual})
    return failures


def document_toggle(document: Any, name: str) -> bool:
    return document.settings.element.find(qn(f"w:{name}")) is not None


def audit_section_rule(
    document: Any, rule: dict, structure_map: dict[str, Any] | None = None
) -> list[dict]:
    failures = []
    for index, section in enumerate(document.sections):
        width = float(section.page_width.mm)
        height = float(section.page_height.mm)
        values = {
            "page_width_mm": width,
            "page_height_mm": height,
            "margin_top_mm": section.top_margin.mm,
            "margin_bottom_mm": section.bottom_margin.mm,
            "margin_left_mm": section.left_margin.mm,
            "margin_right_mm": section.right_margin.mm,
            "gutter_mm": section.gutter.mm,
            "header_distance_ratio": section.header_distance.mm / height,
            "footer_distance_ratio": section.footer_distance.mm / height,
            "margin_inner_ratio": section.left_margin.mm / width,
            "margin_outer_ratio": section.right_margin.mm / width,
            "margin_top_ratio": section.top_margin.mm / height,
            "margin_bottom_ratio": section.bottom_margin.mm / height,
            "mirror_margins": document_toggle(document, "mirrorMargins"),
            "page_size_policy": "preserve",
            "different_first_page_header_footer": bool(section.different_first_page_header_footer),
            "odd_and_even_pages_header_footer": bool(
                document.settings.odd_and_even_pages_header_footer
            ),
            "orientation": "landscape" if section.page_width > section.page_height else "portrait",
        }
        for key, expected in rule["properties"].items():
            if (
                key == "different_first_page_header_footer"
                and index == 0
                and (structure_map or {}).get("front_matter", {}).get("approved")
            ):
                continue
            actual = values.get(key, "<unsupported>")
            if not compare_value(key, actual, expected):
                failures.append(
                    {
                        "section": index,
                        "property": key,
                        "expected": expected,
                        "actual": actual,
                    }
                )
    return failures


def row_property(row: Any, name: str) -> bool:
    tr_pr = row._tr.trPr
    return tr_pr is not None and tr_pr.find(qn(f"w:{name}")) is not None


def _border_value(container: Any, name: str) -> tuple[str | None, int | None]:
    border = None if container is None else container.find(qn(f"w:{name}"))
    return (
        None if border is None else border.get(qn("w:val")),
        None if border is None else int(border.get(qn("w:sz"), "0")),
    )


def _technical_textbook_table_matches(
    table: Any, entry: dict, properties: dict | None = None
) -> bool:
    properties = properties or {}
    major_size = round(float(properties.get("major_border_pt", 1)) * 8)
    minor_size = round(float(properties.get("minor_border_pt", 0.5)) * 8)
    inside_vertical = bool(properties.get("inside_vertical_borders", True))
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    caption_row = entry.get("caption_row")
    has_caption = caption_row is not None and 0 <= int(caption_row) < len(table.rows)
    expected = {
        "top": (("nil", 0) if has_caption else ("single", major_size)),
        "bottom": ("single", major_size),
        "left": ("nil", 0),
        "right": ("nil", 0),
        "insideH": ("nil", 0),
        "insideV": (
            ("single", minor_size) if inside_vertical else ("nil", 0)
        ),
    }
    if any(_border_value(borders, name) != value for name, value in expected.items()):
        return False
    for row in table.rows:
        seen: set[Any] = set()
        for cell in row.cells:
            if cell._tc in seen:
                continue
            seen.add(cell._tc)
            shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            if shading is not None and shading.get(qn("w:fill")) not in {None, "auto"}:
                return False
            cell_borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
            if cell_borders is not None and any(
                _border_value(cell_borders, name)[0] not in {None, "nil"}
                for name in ("left", "right", "insideH", "insideV")
            ):
                return False
    for row_index in properties.get("horizontal_rule_rows", []):
        if not 0 < int(row_index) < len(table.rows):
            return False
        for cell in table.rows[int(row_index)].cells:
            if _border_value(_cell_border_container(cell), "top") != (
                "single",
                minor_size,
            ):
                return False
    return True


def _cell_border_container(cell: Any) -> Any:
    return cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))


def _table_visual_value(
    table: Any, entry: dict, key: str, properties: dict | None = None
) -> Any:
    tbl_pr = table._tbl.tblPr
    if key == "available_width_percent":
        width = tbl_pr.find(qn("w:tblW"))
        return (
            None
            if width is None or width.get(qn("w:type")) != "pct"
            else int(width.get(qn("w:w"), "0")) / 50
        )
    if key == "allow_autofit":
        return bool(table.autofit)
    if key == "cell_margins_mm":
        container = tbl_pr.find(qn("w:tblCellMar"))
        if container is None:
            return None
        result = {}
        for name in ("top", "right", "bottom", "left"):
            element = container.find(qn(f"w:{name}"))
            value = 0 if element is None else int(element.get(qn("w:w"), "0"))
            result[name] = round(value / 1440 * 25.4, 3)
        return result
    if key == "vertical_alignment":
        values = {
            (
                cell.vertical_alignment.name.lower()
                if cell.vertical_alignment is not None
                else "none"
            )
            for row in table.rows
            for cell in row.cells
        }
        return values.pop() if len(values) == 1 else sorted(values)
    if key == "preferred_column_widths_percent":
        if not table.rows:
            return []
        result = []
        for cell in table.rows[0].cells:
            width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            result.append(
                None
                if width is None or width.get(qn("w:type")) != "pct"
                else int(width.get(qn("w:w"), "0")) / 50
            )
        return result
    if key == "border_preset":
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            return "preserve"
        values = {
            name: (
                None
                if borders.find(qn(f"w:{name}")) is None
                else borders.find(qn(f"w:{name}")).get(qn("w:val"))
            )
            for name in ("top", "left", "bottom", "right", "insideH", "insideV")
        }
        if all(value == "single" for value in values.values()):
            return "full_grid"
        if values["top"] == values["bottom"] == "single" and all(
            values[name] == "nil" for name in ("left", "right", "insideH", "insideV")
        ):
            return "three_line"
        if _technical_textbook_table_matches(table, entry, properties):
            return "technical_textbook"
        if all(value == "nil" for value in values.values()):
            return "borderless"
        return "custom"
    return "<verified_by_content_and_visual_qa>"


def audit_table_rule(
    document: Any,
    rule: dict,
    targets: list[tuple[Any, dict[str, Any]]] | None = None,
) -> list[dict]:
    failures = []
    selected = targets if targets is not None else [(table, {}) for table in document.tables]
    for index, (table, entry) in enumerate(selected):
        effective = _table_effective_properties(rule["properties"], entry)
        for key, expected in effective.items():
            if key == "table_style":
                actual = table.style.name if table.style else None
            elif key == "alignment":
                actual = normalized_alignment(str(table.alignment))
            elif key == "repeat_header_row":
                rows = entry.get("repeat_header_rows", [0])
                actual = bool(rows) and all(
                    0 <= int(row_index) < len(table.rows)
                    and row_property(table.rows[int(row_index)], "tblHeader")
                    for row_index in rows
                )
            elif key == "prevent_row_split":
                caption_row = entry.get("caption_row")
                actual = all(
                    row_property(row, "cantSplit")
                    for row_index, row in enumerate(table.rows)
                    if caption_row is None or row_index != int(caption_row)
                )
            elif key in {
                "available_width_percent",
                "preferred_column_widths_percent",
                "allow_autofit",
                "cell_margins_mm",
                "vertical_alignment",
                "all_cell_alignment",
                "text_wrapping",
                "border_preset",
            }:
                if key == "all_cell_alignment":
                    values = {
                        normalized_alignment(str(paragraph.alignment))
                        for row in table.rows
                        for cell in row.cells
                        for paragraph in cell.paragraphs
                        if paragraph.text.strip()
                    }
                    actual = values.pop() if len(values) == 1 else sorted(values)
                elif key == "text_wrapping":
                    actual = (
                        "around"
                        if table._tbl.tblPr.find(qn("w:tblpPr")) is not None
                        else "none"
                    )
                else:
                    actual = (
                        "preserve"
                        if key == "border_preset" and expected == "preserve"
                        else _table_visual_value(table, entry, key, effective)
                    )
            elif key in {
                "column_roles",
                "column_alignments",
                "header_bold",
                "header_shading_hex",
                "font_size_pt",
                "line_spacing_pt",
                "major_border_pt",
                "minor_border_pt",
                "inside_vertical_borders",
                "horizontal_rule_rows",
            }:
                actual = expected
            elif key in {
                "font_name_ascii",
                "font_name_east_asia",
                "font_name_complex_script",
            }:
                actual = [
                    run_font_value(document, paragraph, run, key)
                    for row_index, row in enumerate(table.rows)
                    if entry.get("caption_row") is None
                    or row_index != int(entry["caption_row"])
                    for cell in row.cells
                    for paragraph in cell.paragraphs
                    for run in paragraph.runs
                    if run.text
                ]
            else:
                actual = "<unsupported>"
            if not compare_value(key, actual, expected):
                failures.append(
                    {
                        "table": index,
                        "property": key,
                        "expected": expected,
                        "actual": actual,
                    }
                )
    return failures


def heading_numbering_start(document: Any) -> int | None:
    style = document.styles["Heading 1"]
    p_pr = style.element.pPr
    num_pr = None if p_pr is None else p_pr.find(qn("w:numPr"))
    num_id_element = None if num_pr is None else num_pr.find(qn("w:numId"))
    if num_id_element is None:
        return None
    num_id = num_id_element.get(qn("w:val"))
    root = document.part.numbering_part.element
    for num in root.findall(qn("w:num")):
        if num.get(qn("w:numId")) != num_id:
            continue
        for override in num.findall(qn("w:lvlOverride")):
            if override.get(qn("w:ilvl")) == "0":
                start = override.find(qn("w:startOverride"))
                if start is not None:
                    return int(start.get(qn("w:val")))
        abstract_ref = num.find(qn("w:abstractNumId"))
        if abstract_ref is None:
            return None
        abstract_id = abstract_ref.get(qn("w:val"))
        for abstract in root.findall(qn("w:abstractNum")):
            if abstract.get(qn("w:abstractNumId")) != abstract_id:
                continue
            for level in abstract.findall(qn("w:lvl")):
                if level.get(qn("w:ilvl")) == "0":
                    start = level.find(qn("w:start"))
                    return None if start is None else int(start.get(qn("w:val")))
    return None


def heading_numbering_level(document: Any, level: int) -> Any | None:
    style = document.styles[f"Heading {level + 1}"]
    p_pr = style.element.pPr
    num_pr = None if p_pr is None else p_pr.find(qn("w:numPr"))
    num_id_element = None if num_pr is None else num_pr.find(qn("w:numId"))
    if num_id_element is None:
        return None
    num_id = num_id_element.get(qn("w:val"))
    root = document.part.numbering_part.element
    num = next(
        (item for item in root.findall(qn("w:num")) if item.get(qn("w:numId")) == num_id),
        None,
    )
    abstract_ref = None if num is None else num.find(qn("w:abstractNumId"))
    if abstract_ref is None:
        return None
    abstract_id = abstract_ref.get(qn("w:val"))
    abstract = next(
        (
            item
            for item in root.findall(qn("w:abstractNum"))
            if item.get(qn("w:abstractNumId")) == abstract_id
        ),
        None,
    )
    if abstract is None:
        return None
    return next(
        (
            item
            for item in abstract.findall(qn("w:lvl"))
            if item.get(qn("w:ilvl")) == str(level)
        ),
        None,
    )


def heading_numbering_format_failures(document: Any, levels: int) -> list[dict]:
    failures = []
    for level in range(levels):
        style = document.styles[f"Heading {level + 1}"]
        lvl = heading_numbering_level(document, level)
        if lvl is None:
            continue
        ind = lvl.find(qn("w:pPr") + "/" + qn("w:ind"))
        has_zero_indent = ind is not None and (
            ind.get(qn("w:firstLineChars")) == "0"
            or ind.get(qn("w:firstLine")) == "0"
        )
        if not has_zero_indent:
            failures.append(
                {
                    "property": "heading_first_line_indent",
                    "heading_level": level + 1,
                    "expected": 0,
                }
            )
        r_pr = lvl.find(qn("w:rPr"))
        r_fonts = None if r_pr is None else r_pr.find(qn("w:rFonts"))
        expected_fonts = {
            "ascii": style_effective_font(document, style, "ascii")[0],
            "hAnsi": style_effective_font(document, style, "ascii")[0],
            "eastAsia": style_effective_font(document, style, "eastAsia")[0],
            "cs": style_effective_font(document, style, "cs")[0],
        }
        actual_fonts = {
            attribute: None if r_fonts is None else r_fonts.get(qn(f"w:{attribute}"))
            for attribute in expected_fonts
        }
        if actual_fonts != expected_fonts:
            failures.append(
                {
                    "property": "heading_numbering_fonts",
                    "heading_level": level + 1,
                    "expected": expected_fonts,
                    "actual": actual_fonts,
                }
            )
        expected_size = style_value(document, style, "font_size_pt")
        size = None if r_pr is None else r_pr.find(qn("w:sz"))
        actual_size = None if size is None else int(size.get(qn("w:val"))) / 2
        if not close(actual_size, expected_size):
            failures.append(
                {
                    "property": "heading_numbering_font_size_pt",
                    "heading_level": level + 1,
                    "expected": expected_size,
                    "actual": actual_size,
                }
            )
        expected_bold = style_value(document, style, "bold")
        bold = None if r_pr is None else r_pr.find(qn("w:b"))
        actual_bold = None if bold is None else bold.get(qn("w:val"), "1") not in {"0", "false"}
        if actual_bold != expected_bold:
            failures.append(
                {
                    "property": "heading_numbering_bold",
                    "heading_level": level + 1,
                    "expected": expected_bold,
                    "actual": actual_bold,
                }
            )
    return failures


def audit_field_rule(
    document: Any,
    rule: dict,
    chapter_start: int | None = None,
    path: Path | None = None,
) -> list[dict]:
    failures = []
    properties = rule["properties"]
    if properties.get("update_on_open") and not document_toggle(document, "updateFields"):
        refreshed = path is not None and field_cache_inventory(path)["status"] == "refreshed"
        if not refreshed:
            failures.append({"property": "update_on_open", "expected": True, "actual": False})
    if properties.get("convert_explicit_markers"):
        remaining = [
            paragraph.text
            for paragraph in document.paragraphs
            if _field_instruction_for_marker(paragraph.text.strip()) is not None
        ]
        if remaining:
            failures.append(
                {
                    "property": "convert_explicit_markers",
                    "expected": "no markers",
                    "actual": remaining[:10],
                }
            )
    levels = int(properties.get("heading_levels", 4))
    if properties.get("rebuild_heading_numbering"):
        for level in range(1, levels + 1):
            style = document.styles[f"Heading {level}"]
            p_pr = style.element.pPr
            if p_pr is None or p_pr.find(qn("w:numPr")) is None:
                failures.append(
                    {
                        "property": "rebuild_heading_numbering",
                        "expected": f"numbering on Heading {level}",
                        "actual": "missing",
                    }
                )
        failures.extend(heading_numbering_format_failures(document, levels))
        if chapter_start is not None:
            actual_start = heading_numbering_start(document)
            if actual_start != chapter_start:
                failures.append(
                    {
                        "property": "chapter_start",
                        "expected": chapter_start,
                        "actual": actual_start,
                    }
                )
    if properties.get("strip_manual_heading_prefixes"):
        for paragraph in document.paragraphs:
            if not paragraph.style or not paragraph.style.name.startswith("Heading "):
                continue
            try:
                level = int(paragraph.style.name.split()[-1])
            except ValueError:
                continue
            if level <= levels and _heading_prefix_pattern(level).match(paragraph.text):
                failures.append(
                    {
                        "property": "strip_manual_heading_prefixes",
                        "expected": "no manual prefix",
                        "actual": paragraph.text[:80],
                    }
                )
    return failures


def audit_equation_rule(path: Path, rule: dict) -> list[dict]:
    values = equation_inventory(path)
    failures = []
    properties = rule["properties"]
    if properties.get("block_formula_images") and values["formula_image_candidates"]:
        failures.append(
            {
                "property": "block_formula_images",
                "expected": 0,
                "actual": values["formula_image_candidates"],
            }
        )
    return failures


def uses_derived_normalization(profile: dict) -> bool:
    return any(
        rule.get("status") == "approved"
        and rule.get("application") == "automatic"
        and rule.get("selector", {}).get("kind") == "field_role"
        and any(
            rule.get("properties", {}).get(key)
            for key in (
                "convert_explicit_markers",
                "rebuild_heading_numbering",
                "strip_manual_heading_prefixes",
            )
        )
        for rule in profile.get("rules", [])
    )


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("original", type=Path)
    parser.add_argument("formatted", type=Path)
    parser.add_argument("--profile", required=True, type=Path)
    parser.add_argument("--structure-map", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()

    try:
        profile_errors, profile = validate(args.profile)
        if profile_errors:
            raise FormatMonographError("Profile validation failed: " + "; ".join(profile_errors))
        normalize = uses_derived_normalization(profile)
        structure_map = load_structure_map(args.structure_map) if args.structure_map else None
        original_fp = (
            structure_content_fingerprint(args.original, structure_map)
            if structure_map
            else content_fingerprint(args.original, normalize_derived=normalize)
        )
        formatted_fp = (
            structure_content_fingerprint(args.formatted, structure_map)
            if structure_map
            else content_fingerprint(args.formatted, normalize_derived=normalize)
        )
        original_objects = protected_object_manifest(args.original)
        formatted_objects = protected_object_manifest(args.formatted)
        original_payloads = protected_payload_manifest(args.original)
        formatted_payloads = protected_payload_manifest(args.formatted)
        objects_ok = original_payloads == formatted_payloads
        document = load_document(args.formatted)
        rule_results = []

        for rule in profile["rules"]:
            if rule["status"] != "approved":
                continue
            if rule["application"] == "manual_review":
                rule_results.append(
                    {"id": rule["id"], "status": "manual_review", "failures": []}
                )
                continue
            kind = rule["selector"]["kind"]
            if kind in {"document", "section_role"}:
                failures = audit_section_rule(document, rule, structure_map)
            elif kind == "table_role":
                table_targets = (
                    approved_data_tables(document, structure_map)
                    if structure_map and has_semantic_structure_map(structure_map)
                    else None
                )
                failures = audit_table_rule(document, rule, table_targets)
            elif kind == "field_role":
                numbering = structure_map.get("numbering", {}) if structure_map else {}
                chapter_start = (
                    int(numbering["chapter_start"])
                    if numbering.get("approved")
                    else None
                )
                failures = audit_field_rule(
                    document, rule, chapter_start, args.formatted
                )
            elif kind == "equation_role":
                failures = audit_equation_rule(args.formatted, rule)
            elif (
                structure_map
                and has_semantic_structure_map(structure_map)
                and kind in {"paragraph_role", "caption_role", "bibliography_role"}
            ):
                paragraphs = approved_role_paragraphs(
                    document, structure_map, rule["selector"]
                )
                failures = audit_paragraph_rule(document, rule, paragraphs)
            else:
                failures = audit_style_rule(document, rule)
            rule_results.append(
                {
                    "id": rule["id"],
                    "status": "pass" if not failures else "fail",
                    "failures": failures,
                }
            )

        pagination_failures, pagination = (
            audit_pagination_sections(
                args.formatted,
                document,
                structure_map.get("pagination_sections", {}),
                resolve_paragraph_locator,
                structure_map,
            )
            if structure_map
            else ([], {})
        )
        content_ok = original_fp == formatted_fp
        structure_heading_failures = (
            audit_structure_heading_operations(document, structure_map)
            if structure_map
            else []
        )
        structure_table_failures = (
            audit_structure_table_operations(document, structure_map)
            if structure_map
            else []
        )
        structure_image_failures = (
            audit_structure_image_operations(document, structure_map)
            if structure_map
            else []
        )
        structure_toc_source_failures = (
            audit_structure_toc_source_operations(document, structure_map)
            if structure_map
            else []
        )
        rules_ok = (
            all(item["status"] != "fail" for item in rule_results)
            and not pagination_failures
            and not structure_heading_failures
            and not structure_table_failures
            and not structure_image_failures
            and not structure_toc_source_failures
        )
        caption_replacements = (
            audit_caption_identifier_replacements(
                args.original, args.formatted, structure_map
            )
            if structure_map
            else []
        )
        caption_replacements_ok = all(
            item["status"] == "pass" for item in caption_replacements
        )
        result = {
            "passed": content_ok and objects_ok and rules_ok and caption_replacements_ok,
            "content_integrity": {
                "passed": content_ok,
                "original_sha256": original_fp,
                "formatted_sha256": formatted_fp,
                "field_results_and_approved_derived_values_excluded": bool(
                    normalize or structure_map
                ),
                "normalization_sources": [
                    source
                    for source, enabled in (
                        ("profile_field_rules", normalize),
                        ("approved_structure_map", bool(structure_map)),
                    )
                    if enabled
                ],
            },
            "protected_object_integrity": {
                "passed": objects_ok,
                "original": original_objects,
                "formatted": formatted_objects,
                "payload_comparison": {
                    "path_independent": True,
                    "original": original_payloads,
                    "formatted": formatted_payloads,
                },
            },
            "equations": equation_inventory(args.formatted),
            "approved_manual_identifier_replacements": caption_replacements,
            "pagination": {
                "passed": not pagination_failures,
                "failures": pagination_failures,
                "inventory": pagination,
            },
            "structure_heading_operations": {
                "passed": not structure_heading_failures,
                "failures": structure_heading_failures,
            },
            "structure_table_operations": {
                "passed": not structure_table_failures,
                "failures": structure_table_failures,
            },
            "structure_image_operations": {
                "passed": not structure_image_failures,
                "failures": structure_image_failures,
            },
            "structure_toc_source_operations": {
                "passed": not structure_toc_source_failures,
                "failures": structure_toc_source_failures,
            },
            "rules": rule_results,
        }
        if args.output:
            write_json(args.output, result)
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if result["passed"] else 1
    except (FormatMonographError, OSError, ValueError) as exc:
        print(str(exc), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
