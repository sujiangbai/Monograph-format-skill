"""Shared helpers for the format-monograph command-line tools."""

from __future__ import annotations

import copy
import hashlib
import json
import os
import platform
import re
import shutil
import subprocess
import zipfile
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
O_NS = "urn:schemas-microsoft-com:office:office"
V_NS = "urn:schemas-microsoft-com:vml"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": W_NS, "m": M_NS, "o": O_NS, "v": V_NS, "r": R_NS}
CONTENT_PART = re.compile(
    r"^word/(?:document|header\d+|footer\d+|footnotes|endnotes)\.xml$"
)
FIELD_MARKER_PATTERN = re.compile(
    r"\[\[(?:TOC|PAGE|(?:REF|PAGEREF|SEQ):[A-Za-z0-9_.-]+)\]\]"
)

ROLE_STYLE_MAP = {
    "body": "Normal",
    "title": "Title",
    "subtitle": "Subtitle",
    "caption": "Caption",
    "bibliography": "Bibliography",
    "body_text": "Normal",
    "chapter_title": "Heading 1",
    "level_2_section": "Heading 2",
    "level_3_section": "Heading 3",
    "level_4_section": "Heading 4",
    "long_quote": "Quote",
    "heading_1": "Heading 1",
    "heading_2": "Heading 2",
    "heading_3": "Heading 3",
    "heading_4": "Heading 4",
    "figure_caption": "Caption",
    "table_caption": "Caption",
    "equation_caption": "Caption",
    "reference_entry": "Bibliography",
    "answer": "Normal",
    "teaching_callout": "Normal",
    "toc_level_1": "TOC 1",
    "toc_level_2": "TOC 2",
    "toc_level_3": "TOC 3",
    "table_of_contents_level_1": "TOC 1",
    "table_of_contents_level_2": "TOC 2",
    "table_of_contents_level_3": "TOC 3",
    **{f"heading{i}": f"Heading {i}" for i in range(1, 10)},
}

STYLE_PROPERTIES = {
    "font_name",
    "font_name_ascii",
    "font_name_east_asia",
    "font_name_complex_script",
    "font_size_pt",
    "bold",
    "italic",
    "color_hex",
    "alignment",
    "space_before_pt",
    "space_after_pt",
    "line_spacing",
    "line_spacing_rule",
    "line_spacing_pt",
    "first_line_indent_pt",
    "first_line_indent_chars",
    "left_indent_pt",
    "right_indent_pt",
    "keep_with_next",
    "keep_together",
    "page_break_before",
    "widow_control",
}

CAPTION_POLICY_PROPERTIES = {
    "numbering_mode",
    "preserve_identifier",
    "domain_context",
    "allow_automatic_renumbering",
    "preserve_table_cell_caption_position",
}

SECTION_PROPERTIES = {
    "page_width_mm",
    "page_height_mm",
    "orientation",
    "margin_top_mm",
    "margin_bottom_mm",
    "margin_left_mm",
    "margin_right_mm",
    "gutter_mm",
    "different_first_page_header_footer",
    "odd_and_even_pages_header_footer",
    "page_size_policy",
    "mirror_margins",
    "margin_inner_ratio",
    "margin_outer_ratio",
    "margin_top_ratio",
    "margin_bottom_ratio",
    "header_distance_ratio",
    "footer_distance_ratio",
}

TABLE_PROPERTIES = {
    "table_style",
    "alignment",
    "repeat_header_row",
    "prevent_row_split",
    "available_width_percent",
    "preferred_column_widths_percent",
    "allow_autofit",
    "cell_margins_mm",
    "vertical_alignment",
    "border_preset",
    "column_roles",
    "column_alignments",
    "header_bold",
    "header_shading_hex",
    "font_name_ascii",
    "font_name_east_asia",
    "font_name_complex_script",
    "font_size_pt",
    "line_spacing_pt",
}

FIELD_PROPERTIES = {
    "update_on_open",
    "mark_fields_dirty",
    "convert_explicit_markers",
    "rebuild_heading_numbering",
    "heading_levels",
    "strip_manual_heading_prefixes",
    "chapter_start",
}

FONT_ALIAS_GROUPS = (
    {"宋体", "simsun", "nsimsun", "新宋体"},
    {"黑体", "simhei"},
    {"楷体", "kaiti", "simkai", "楷体_gb2312"},
    {"仿宋", "fangsong", "simfang", "仿宋_gb2312"},
)

EQUATION_PROPERTIES = {
    "require_editable_equations",
    "preserve_editable_objects",
    "block_formula_images",
}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distributed": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}

TABLE_ALIGNMENTS = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}

FONT_THEME_ATTRIBUTES = {
    "ascii": "asciiTheme",
    "hAnsi": "hAnsiTheme",
    "eastAsia": "eastAsiaTheme",
    "cs": "cstheme",
}

A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


class FormatMonographError(RuntimeError):
    """Expected user-facing validation or processing error."""


def ensure_docx(path: Path) -> None:
    if path.suffix.lower() != ".docx":
        raise FormatMonographError(f"Expected a .docx file: {path}")
    if not path.is_file():
        raise FormatMonographError(f"DOCX file does not exist: {path}")
    if not zipfile.is_zipfile(path):
        raise FormatMonographError(f"File is not a valid DOCX package: {path}")
    with zipfile.ZipFile(path) as package:
        if "word/document.xml" not in package.namelist():
            raise FormatMonographError(f"DOCX package has no word/document.xml: {path}")


def load_json(path: Path) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise FormatMonographError(f"JSON file does not exist: {path}") from exc
    except json.JSONDecodeError as exc:
        raise FormatMonographError(
            f"Invalid JSON at line {exc.lineno}, column {exc.colno}: {exc.msg}"
        ) from exc


def write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )


def profile_schema_path() -> Path:
    return Path(__file__).resolve().parent.parent / "references" / "format-profile.schema.json"


def _font_directories() -> list[Path]:
    candidates: list[Path] = []
    system = platform.system()
    if system == "Windows":
        if os.environ.get("WINDIR"):
            candidates.append(Path(os.environ["WINDIR"]) / "Fonts")
        if os.environ.get("LOCALAPPDATA"):
            candidates.append(
                Path(os.environ["LOCALAPPDATA"]) / "Microsoft" / "Windows" / "Fonts"
            )
    elif system == "Darwin":
        candidates.extend(
            [
                Path("/System/Library/Fonts"),
                Path("/Library/Fonts"),
                Path.home() / "Library" / "Fonts",
            ]
        )
    else:
        candidates.extend(
            [
                Path("/usr/share/fonts"),
                Path("/usr/local/share/fonts"),
                Path.home() / ".fonts",
                Path.home() / ".local" / "share" / "fonts",
            ]
        )
    return [path for path in candidates if path.is_dir()]


def available_font_names() -> set[str]:
    names: set[str] = set()
    fc_list = shutil.which("fc-list")
    if fc_list:
        try:
            result = subprocess.run(
                [fc_list, "--format=%{family}\\n"],
                capture_output=True,
                text=True,
                check=False,
                timeout=20,
            )
            for line in result.stdout.splitlines():
                names.update(part.strip() for part in line.split(",") if part.strip())
        except (OSError, subprocess.TimeoutExpired):
            pass

    if platform.system() == "Windows":
        try:
            import winreg

            key_path = r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"
            with winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, key_path) as key:
                index = 0
                while True:
                    try:
                        label, _, _ = winreg.EnumValue(key, index)
                    except OSError:
                        break
                    label = re.sub(r"\s*\([^)]*\)\s*$", "", label)
                    names.update(
                        part.strip()
                        for part in re.split(r"\s*&\s*", label)
                        if part.strip()
                    )
                    index += 1
        except OSError:
            pass

    for directory in _font_directories():
        for pattern in ("*.ttf", "*.ttc", "*.otf", "*.dfont"):
            for path in directory.rglob(pattern):
                names.add(path.stem)
    return names


def _font_key(value: str) -> str:
    return re.sub(r"[\s_-]+", "", value).casefold()


def font_alias_keys(value: str) -> set[str]:
    key = _font_key(value)
    for group in FONT_ALIAS_GROUPS:
        normalized = {_font_key(item) for item in group}
        if key in normalized:
            return normalized
    return {key}


def resolve_font_name(value: str, available_names: set[str] | None = None) -> dict[str, Any]:
    available_names = available_font_names() if available_names is None else available_names
    available_by_key = {_font_key(name): name for name in available_names}
    for key in font_alias_keys(value):
        if key in available_by_key:
            return {
                "requested": value,
                "available": True,
                "matched_name": available_by_key[key],
                "match": "exact" if key == _font_key(value) else "verified_alias",
            }
    return {
        "requested": value,
        "available": False,
        "matched_name": None,
        "match": "missing",
    }


def required_profile_fonts(profile: dict[str, Any]) -> list[str]:
    keys = {
        "font_name",
        "font_name_ascii",
        "font_name_east_asia",
        "font_name_complex_script",
    }
    return sorted(
        {
            str(value)
            for rule in profile.get("rules", [])
            if rule.get("status") == "approved"
            and rule.get("application") == "automatic"
            for key, value in rule.get("properties", {}).items()
            if key in keys and str(value).strip()
        }
    )


def missing_profile_fonts(profile: dict[str, Any]) -> list[str]:
    return [
        item["requested"]
        for item in profile_font_resolutions(profile)
        if not item["available"]
    ]


def profile_font_resolutions(profile: dict[str, Any]) -> list[dict[str, Any]]:
    available = available_font_names()
    return [
        resolve_font_name(name, available)
        for name in required_profile_fonts(profile)
    ]


def _paragraph_text_without_field_results(paragraph: etree._Element) -> str:
    pieces: list[str] = []
    field_stack: list[str] = []

    for element in paragraph.iter():
        if element.tag == f"{{{W_NS}}}fldChar":
            kind = element.get(f"{{{W_NS}}}fldCharType")
            if kind == "begin":
                field_stack.append("instruction")
            elif kind == "separate" and field_stack:
                field_stack[-1] = "result"
            elif kind == "end" and field_stack:
                field_stack.pop()
            continue

        if element.tag in {f"{{{W_NS}}}t", f"{{{W_NS}}}delText"}:
            if any(state == "result" for state in field_stack):
                continue
            if any(parent.tag == f"{{{W_NS}}}fldSimple" for parent in element.iterancestors()):
                continue
            pieces.append(element.text or "")
        elif element.tag == f"{{{W_NS}}}tab":
            pieces.append("\t")
        elif element.tag in {f"{{{W_NS}}}br", f"{{{W_NS}}}cr"}:
            pieces.append("\n")

    return "".join(pieces)


def _normalize_derived_paragraph_text(paragraph: etree._Element, value: str) -> str:
    value = FIELD_MARKER_PATTERN.sub("", value)
    styles = paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
    if styles:
        match = re.fullmatch(r"Heading([1-4])", styles[0])
        if match:
            prefix = _heading_prefix_pattern(int(match.group(1))).match(value)
            if prefix:
                return value[prefix.end():]
    return value


def content_inventory(path: Path, normalize_derived: bool = False) -> dict[str, list[str]]:
    """Return authored text by OOXML part, excluding generated field results."""
    ensure_docx(path)
    result: dict[str, list[str]] = {}
    with zipfile.ZipFile(path) as package:
        for name in sorted(package.namelist()):
            if not CONTENT_PART.match(name):
                continue
            root = etree.fromstring(package.read(name))
            values = []
            for paragraph in root.xpath(".//w:p", namespaces=NS):
                value = _paragraph_text_without_field_results(paragraph)
                if normalize_derived:
                    value = _normalize_derived_paragraph_text(paragraph, value)
                values.append(value)
            result[name] = values
    return result


def content_fingerprint(path: Path, normalize_derived: bool = False) -> str:
    encoded = json.dumps(
        content_inventory(path, normalize_derived=normalize_derived),
        ensure_ascii=False,
        separators=(",", ":"),
        sort_keys=True,
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def word_xml_counts(path: Path) -> dict[str, int]:
    ensure_docx(path)
    counts = {
        "fields": 0,
        "equations": 0,
        "drawings": 0,
        "text_boxes": 0,
        "footnotes": 0,
        "endnotes": 0,
        "comments": 0,
    }
    with zipfile.ZipFile(path) as package:
        names = set(package.namelist())
        counts["footnotes"] = int("word/footnotes.xml" in names)
        counts["endnotes"] = int("word/endnotes.xml" in names)
        counts["comments"] = int("word/comments.xml" in names)
        for name in names:
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = etree.fromstring(package.read(name))
            counts["fields"] += len(root.xpath(".//w:fldChar | .//w:fldSimple", namespaces=NS))
            counts["equations"] += len(root.xpath(".//*[local-name()='oMath']"))
            counts["drawings"] += len(root.xpath(".//w:drawing", namespaces=NS))
            counts["text_boxes"] += len(root.xpath(".//w:txbxContent", namespaces=NS))
    return counts


def _iter_container_paragraphs(container: Any) -> Iterable[Any]:
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_paragraphs(cell)


def iter_document_paragraphs(document: Any) -> Iterable[Any]:
    # Accessing an undefined header/footer through python-docx materializes a new
    # part. Restrict automatic paragraph traversal to existing body containers
    # so a read-only rule count cannot alter the package.
    yield from _iter_container_paragraphs(document)


def style_name_for_selector(selector: dict[str, str]) -> str | None:
    kind, value = selector["kind"], selector["value"]
    if kind == "style_name":
        return value
    if kind == "paragraph_role":
        return ROLE_STYLE_MAP.get(value)
    if kind == "caption_role":
        return "Caption"
    if kind == "bibliography_role":
        return "Bibliography"
    return None


def supported_properties(rule: dict[str, Any]) -> set[str]:
    kind = rule["selector"]["kind"]
    if kind in {"document", "section_role"}:
        return SECTION_PROPERTIES
    if kind == "table_role":
        return TABLE_PROPERTIES
    if kind == "field_role":
        return FIELD_PROPERTIES
    if kind == "equation_role":
        return EQUATION_PROPERTIES
    if kind == "caption_role":
        return STYLE_PROPERTIES | CAPTION_POLICY_PROPERTIES
    if style_name_for_selector(rule["selector"]):
        return STYLE_PROPERTIES
    return set()


def unsupported_properties(rule: dict[str, Any]) -> list[str]:
    return sorted(set(rule["properties"]) - supported_properties(rule))


def theme_font_inventory(document: Any) -> dict[str, str]:
    cached = getattr(document, "_format_monograph_theme_fonts", None)
    if cached is not None:
        return cached
    result: dict[str, str] = {}
    theme_part = next(
        (
            part
            for part in document.part.package.parts
            if str(part.partname) == "/word/theme/theme1.xml"
        ),
        None,
    )
    if theme_part is not None:
        root = etree.fromstring(theme_part.blob)
        namespaces = {"a": A_NS}
        for family, element_name in (("major", "majorFont"), ("minor", "minorFont")):
            container = root.find(
                f".//{{{A_NS}}}fontScheme/{{{A_NS}}}{element_name}"
            )
            if container is None:
                continue
            latin = container.find(f"{{{A_NS}}}latin")
            east_asia = container.find(f"{{{A_NS}}}ea")
            complex_script = container.find(f"{{{A_NS}}}cs")
            hans = container.xpath("./a:font[@script='Hans']", namespaces=namespaces)
            latin_name = "" if latin is None else latin.get("typeface", "")
            east_asia_name = (
                "" if east_asia is None else east_asia.get("typeface", "")
            ) or (hans[0].get("typeface", "") if hans else "")
            complex_name = (
                "" if complex_script is None else complex_script.get("typeface", "")
            )
            result[f"{family}HAnsi"] = latin_name
            result[f"{family}Ascii"] = latin_name
            result[f"{family}EastAsia"] = east_asia_name
            result[f"{family}Bidi"] = complex_name
    setattr(document, "_format_monograph_theme_fonts", result)
    return result


def rpr_effective_font(
    document: Any, r_pr: Any, attribute: str
) -> tuple[str | None, str | None]:
    if r_pr is None:
        return None, None
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        return None, None
    explicit = r_fonts.get(qn(f"w:{attribute}"))
    if explicit:
        return explicit, "explicit"
    theme_attribute = FONT_THEME_ATTRIBUTES[attribute]
    theme_key = r_fonts.get(qn(f"w:{theme_attribute}"))
    if not theme_key:
        return None, None
    resolved = theme_font_inventory(document).get(theme_key) or None
    return resolved, f"theme:{theme_key}"


def style_effective_font(
    document: Any, style: Any, attribute: str
) -> tuple[str | None, str | None]:
    current = style
    visited: set[str] = set()
    while current is not None and current.style_id not in visited:
        visited.add(current.style_id)
        value, source = rpr_effective_font(document, current.element.rPr, attribute)
        if value:
            return value, f"style:{current.name}:{source}"
        current = current.base_style
    defaults = document.styles.element.find(
        qn("w:docDefaults") + "/" + qn("w:rPrDefault") + "/" + qn("w:rPr")
    )
    value, source = rpr_effective_font(document, defaults, attribute)
    return value, None if source is None else f"docDefaults:{source}"


def run_effective_font(
    document: Any, paragraph: Any, run: Any, attribute: str
) -> tuple[str | None, str | None]:
    value, source = rpr_effective_font(document, run._r.rPr, attribute)
    if value:
        return value, f"run:{source}"
    run_style = run.style
    if run_style is not None and run_style.name != "Default Paragraph Font":
        value, source = style_effective_font(document, run_style, attribute)
        if value:
            return value, f"character-{source}"
    return style_effective_font(document, paragraph.style, attribute)


def _font_attributes(font_element: Any) -> Any:
    r_pr = font_element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    return r_fonts


def _set_font_attributes(font_element: Any, name: str, attributes: tuple[str, ...]) -> None:
    r_fonts = _font_attributes(font_element)
    for attr in attributes:
        r_fonts.set(qn(f"w:{attr}"), name)
        theme_attribute = FONT_THEME_ATTRIBUTES.get(attr)
        if theme_attribute:
            r_fonts.attrib.pop(qn(f"w:{theme_attribute}"), None)


def _set_east_asian_font(font_element: Any, name: str) -> None:
    _set_font_attributes(font_element, name, ("ascii", "hAnsi", "eastAsia", "cs"))


def apply_style_properties(style: Any, properties: dict[str, Any]) -> None:
    font = style.font
    paragraph_format = style.paragraph_format

    if "font_name" in properties:
        font.name = properties["font_name"]
        _set_east_asian_font(style.element, properties["font_name"])
    if "font_name_ascii" in properties:
        font.name = properties["font_name_ascii"]
        _set_font_attributes(
            style.element, properties["font_name_ascii"], ("ascii", "hAnsi")
        )
        if "font_name_complex_script" not in properties:
            _set_font_attributes(style.element, properties["font_name_ascii"], ("cs",))
    if "font_name_east_asia" in properties:
        _set_font_attributes(
            style.element, properties["font_name_east_asia"], ("eastAsia",)
        )
    if "font_name_complex_script" in properties:
        _set_font_attributes(
            style.element, properties["font_name_complex_script"], ("cs",)
        )
    if "font_size_pt" in properties:
        font.size = Pt(float(properties["font_size_pt"]))
    if "bold" in properties:
        font.bold = bool(properties["bold"])
    if "italic" in properties:
        font.italic = bool(properties["italic"])
    if "color_hex" in properties:
        color = str(properties["color_hex"]).lstrip("#")
        if not re.fullmatch(r"[0-9A-Fa-f]{6}", color):
            raise FormatMonographError(f"Invalid color_hex: {properties['color_hex']}")
        font.color.rgb = RGBColor.from_string(color.upper())
    if "alignment" in properties:
        value = properties["alignment"]
        if value not in ALIGNMENTS:
            raise FormatMonographError(f"Unsupported paragraph alignment: {value}")
        paragraph_format.alignment = ALIGNMENTS[value]
    for key, attr in (
        ("space_before_pt", "space_before"),
        ("space_after_pt", "space_after"),
        ("first_line_indent_pt", "first_line_indent"),
        ("left_indent_pt", "left_indent"),
        ("right_indent_pt", "right_indent"),
    ):
        if key in properties:
            setattr(paragraph_format, attr, Pt(float(properties[key])))
    if "line_spacing" in properties:
        paragraph_format.line_spacing = float(properties["line_spacing"])
    if "line_spacing_rule" in properties:
        rules = {
            "single": WD_LINE_SPACING.SINGLE,
            "one_point_five": WD_LINE_SPACING.ONE_POINT_FIVE,
            "double": WD_LINE_SPACING.DOUBLE,
            "at_least": WD_LINE_SPACING.AT_LEAST,
            "exact": WD_LINE_SPACING.EXACTLY,
            "multiple": WD_LINE_SPACING.MULTIPLE,
        }
        value = properties["line_spacing_rule"]
        if value not in rules:
            raise FormatMonographError(f"Unsupported line_spacing_rule: {value}")
        paragraph_format.line_spacing_rule = rules[value]
    if "line_spacing_pt" in properties:
        paragraph_format.line_spacing = Pt(float(properties["line_spacing_pt"]))
        if "line_spacing_rule" not in properties:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if "first_line_indent_chars" in properties:
        p_pr = style.element.get_or_add_pPr()
        ind = p_pr.get_or_add_ind()
        ind.set(
            qn("w:firstLineChars"),
            str(int(round(float(properties["first_line_indent_chars"]) * 100))),
        )
        ind.attrib.pop(qn("w:firstLine"), None)
    for key, attr in (
        ("keep_with_next", "keep_with_next"),
        ("keep_together", "keep_together"),
        ("page_break_before", "page_break_before"),
        ("widow_control", "widow_control"),
    ):
        if key in properties:
            setattr(paragraph_format, attr, bool(properties[key]))


def ensure_paragraph_style(document: Any, style_name: str) -> Any:
    try:
        return document.styles[style_name]
    except KeyError:
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        if style_name != "Normal":
            style.base_style = document.styles["Normal"]
        return style


def _drop_attributes(element: Any, names: tuple[str, ...]) -> None:
    for name in names:
        element.attrib.pop(qn(f"w:{name}"), None)
    if not element.attrib and len(element) == 0 and element.getparent() is not None:
        element.getparent().remove(element)


def clear_controlled_direct_format(paragraph: Any, properties: dict[str, Any]) -> None:
    p_pr = paragraph._p.pPr
    if p_pr is not None:
        if "alignment" in properties:
            element = p_pr.find(qn("w:jc"))
            if element is not None:
                p_pr.remove(element)

        spacing_keys = {
            "space_before_pt": ("before", "beforeLines", "beforeAutospacing"),
            "space_after_pt": ("after", "afterLines", "afterAutospacing"),
            "line_spacing": ("line", "lineRule"),
            "line_spacing_pt": ("line", "lineRule"),
            "line_spacing_rule": ("line", "lineRule"),
        }
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is not None:
            attrs = tuple(
                attr
                for key, values in spacing_keys.items()
                if key in properties
                for attr in values
            )
            _drop_attributes(spacing, attrs)

        indent_keys = {
            "first_line_indent_pt": ("firstLine", "firstLineChars", "hanging", "hangingChars"),
            "first_line_indent_chars": ("firstLine", "firstLineChars", "hanging", "hangingChars"),
            "left_indent_pt": ("left", "leftChars", "start", "startChars"),
            "right_indent_pt": ("right", "rightChars", "end", "endChars"),
        }
        indent = p_pr.find(qn("w:ind"))
        if indent is not None:
            attrs = tuple(
                attr
                for key, values in indent_keys.items()
                if key in properties
                for attr in values
            )
            _drop_attributes(indent, attrs)

        for key, tag in (
            ("keep_with_next", "keepNext"),
            ("keep_together", "keepLines"),
            ("page_break_before", "pageBreakBefore"),
            ("widow_control", "widowControl"),
        ):
            element = p_pr.find(qn(f"w:{tag}"))
            if key in properties and element is not None:
                p_pr.remove(element)

    for run in paragraph.runs:
        r_pr = run._r.rPr
        if r_pr is None:
            continue
        if any(
            key in properties
            for key in (
                "font_name",
                "font_name_ascii",
                "font_name_east_asia",
                "font_name_complex_script",
            )
        ):
            r_fonts = r_pr.find(qn("w:rFonts"))
            if r_fonts is not None:
                attrs = []
                if "font_name" in properties:
                    attrs.extend(
                        (
                            "ascii",
                            "hAnsi",
                            "eastAsia",
                            "cs",
                            "asciiTheme",
                            "hAnsiTheme",
                            "eastAsiaTheme",
                            "cstheme",
                        )
                    )
                if "font_name_ascii" in properties:
                    attrs.extend(("ascii", "hAnsi", "asciiTheme", "hAnsiTheme"))
                    if "font_name_complex_script" not in properties:
                        attrs.extend(("cs", "cstheme"))
                if "font_name_east_asia" in properties:
                    attrs.extend(("eastAsia", "eastAsiaTheme"))
                if "font_name_complex_script" in properties:
                    attrs.extend(("cs", "cstheme"))
                _drop_attributes(r_fonts, tuple(attrs))
        for key, tags in (
            ("font_size_pt", ("sz", "szCs")),
            ("bold", ("b", "bCs")),
            ("italic", ("i", "iCs")),
            ("color_hex", ("color",)),
        ):
            if key not in properties:
                continue
            for tag in tags:
                element = r_pr.find(qn(f"w:{tag}"))
                if element is not None:
                    r_pr.remove(element)


def apply_style_rule_to_paragraphs(
    document: Any, rule: dict[str, Any], paragraphs: Iterable[Any]
) -> int:
    style_name = style_name_for_selector(rule["selector"])
    if style_name is None:
        raise FormatMonographError(
            f"Rule {rule['id']} has no paragraph style mapping."
        )
    style = ensure_paragraph_style(document, style_name)
    apply_style_properties(style, rule["properties"])
    targets = list(paragraphs)
    for paragraph in targets:
        paragraph.style = style
        clear_controlled_direct_format(paragraph, rule["properties"])
    return len(targets)


def _set_document_toggle(document: Any, name: str, enabled: bool) -> None:
    settings = document.settings.element
    element = settings.find(qn(f"w:{name}"))
    if enabled:
        if element is None:
            element = OxmlElement(f"w:{name}")
            settings.append(element)
        element.set(qn("w:val"), "true")
    elif element is not None:
        settings.remove(element)


def apply_section_properties(document: Any, properties: dict[str, Any]) -> int:
    sections = list(document.sections)
    page_size_policy = properties.get("page_size_policy")
    if page_size_policy not in {None, "preserve"}:
        raise FormatMonographError(f"Unsupported page_size_policy: {page_size_policy}")

    explicit_size = "page_width_mm" in properties or "page_height_mm" in properties
    for section in sections:
        if "page_width_mm" in properties:
            section.page_width = Mm(float(properties["page_width_mm"]))
        if "page_height_mm" in properties:
            section.page_height = Mm(float(properties["page_height_mm"]))
        if "orientation" in properties:
            orientation = str(properties["orientation"]).lower()
            if orientation not in {"portrait", "landscape"}:
                raise FormatMonographError(f"Unsupported orientation: {orientation}")
            target = WD_ORIENT.PORTRAIT if orientation == "portrait" else WD_ORIENT.LANDSCAPE
            if section.orientation != target and not explicit_size:
                section.page_width, section.page_height = section.page_height, section.page_width
            section.orientation = target

        for key, attr in (
            ("margin_top_mm", "top_margin"),
            ("margin_bottom_mm", "bottom_margin"),
            ("margin_left_mm", "left_margin"),
            ("margin_right_mm", "right_margin"),
            ("gutter_mm", "gutter"),
        ):
            if key in properties:
                setattr(section, attr, Mm(float(properties[key])))

        width_mm = float(section.page_width.mm)
        height_mm = float(section.page_height.mm)
        ratio_values = (
            ("margin_inner_ratio", "left_margin", width_mm),
            ("margin_outer_ratio", "right_margin", width_mm),
            ("margin_top_ratio", "top_margin", height_mm),
            ("margin_bottom_ratio", "bottom_margin", height_mm),
            ("header_distance_ratio", "header_distance", height_mm),
            ("footer_distance_ratio", "footer_distance", height_mm),
        )
        for key, attr, basis in ratio_values:
            if key in properties:
                ratio = float(properties[key])
                if not 0 <= ratio < 0.5:
                    raise FormatMonographError(f"{key} must be between 0 and 0.5.")
                setattr(section, attr, Mm(basis * ratio))

        if "different_first_page_header_footer" in properties:
            section.different_first_page_header_footer = bool(
                properties["different_first_page_header_footer"]
            )

    if "odd_and_even_pages_header_footer" in properties:
        document.settings.odd_and_even_pages_header_footer = bool(
            properties["odd_and_even_pages_header_footer"]
        )
    if "mirror_margins" in properties:
        _set_document_toggle(document, "mirrorMargins", bool(properties["mirror_margins"]))

    return len(sections)

def _set_repeat_table_header(row: Any, enabled: bool) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if enabled:
        if existing is None:
            existing = OxmlElement("w:tblHeader")
            tr_pr.append(existing)
        existing.set(qn("w:val"), "true")
    elif existing is not None:
        tr_pr.remove(existing)


def _set_prevent_row_split(row: Any, enabled: bool) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:cantSplit"))
    if enabled:
        if existing is None:
            existing = OxmlElement("w:cantSplit")
            tr_pr.append(existing)
        existing.set(qn("w:val"), "true")
    elif existing is not None:
        tr_pr.remove(existing)


def _table_effective_properties(
    properties: dict[str, Any], entry: dict[str, Any]
) -> dict[str, Any]:
    result = dict(properties)
    visual = entry.get("visual", {})
    if visual.get("approved"):
        result.update(
            {
                key: value
                for key, value in visual.items()
                if key not in {"approved", "orientation", "landscape_approved"}
            }
        )
    return result


def _set_table_width_percent(table: Any, value: float) -> None:
    if not 1 <= value <= 100:
        raise FormatMonographError("available_width_percent must be between 1 and 100.")
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.insert(0, width)
    width.set(qn("w:type"), "pct")
    width.set(qn("w:w"), str(round(value * 50)))


def _set_column_widths_percent(table: Any, values: list[Any]) -> None:
    if len(values) != len(table.columns):
        raise FormatMonographError(
            "preferred_column_widths_percent must match the table column count."
        )
    widths = [float(value) for value in values]
    if any(value <= 0 for value in widths) or abs(sum(widths) - 100.0) > 0.25:
        raise FormatMonographError(
            "preferred_column_widths_percent must be positive and total 100."
        )
    for column_index, percent in enumerate(widths):
        for row in table.rows:
            cell = row.cells[column_index]
            tc_pr = cell._tc.get_or_add_tcPr()
            width = tc_pr.find(qn("w:tcW"))
            if width is None:
                width = OxmlElement("w:tcW")
                tc_pr.append(width)
            width.set(qn("w:type"), "pct")
            width.set(qn("w:w"), str(round(percent * 50)))


def _set_table_cell_margins(table: Any, value: Any) -> None:
    margins = (
        {name: float(value) for name in ("top", "right", "bottom", "left")}
        if isinstance(value, (int, float))
        else {name: float(value.get(name, 0)) for name in ("top", "right", "bottom", "left")}
    )
    if any(amount < 0 or amount > 20 for amount in margins.values()):
        raise FormatMonographError("cell_margins_mm values must be between 0 and 20.")
    tbl_pr = table._tbl.tblPr
    container = tbl_pr.find(qn("w:tblCellMar"))
    if container is None:
        container = OxmlElement("w:tblCellMar")
        tbl_pr.append(container)
    for name, amount in margins.items():
        element = container.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            container.append(element)
        element.set(qn("w:w"), str(round(amount / 25.4 * 1440)))
        element.set(qn("w:type"), "dxa")


def _set_border(element: Any, name: str, *, style: str, size: int = 4) -> None:
    border = element.find(qn(f"w:{name}"))
    if border is None:
        border = OxmlElement(f"w:{name}")
        element.append(border)
    border.set(qn("w:val"), style)
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "000000")


def _set_table_borders(table: Any, preset: str, header_rows: list[int]) -> None:
    if preset == "preserve":
        return
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    if preset == "full_grid":
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            _set_border(borders, name, style="single", size=4)
        return
    if preset != "three_line":
        raise FormatMonographError(f"Unsupported table border preset: {preset}")
    for name in ("top", "bottom"):
        _set_border(borders, name, style="single", size=8)
    for name in ("left", "right", "insideH", "insideV"):
        _set_border(borders, name, style="nil", size=0)
    for row_index in header_rows:
        if not 0 <= row_index < len(table.rows):
            continue
        for cell in table.rows[row_index].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_borders = tc_pr.find(qn("w:tcBorders"))
            if cell_borders is None:
                cell_borders = OxmlElement("w:tcBorders")
                tc_pr.append(cell_borders)
            _set_border(cell_borders, "bottom", style="single", size=4)


def _set_cell_shading(cell: Any, color: str | None) -> None:
    if not color:
        return
    value = str(color).lstrip("#").upper()
    if not re.fullmatch(r"[0-9A-F]{6}", value):
        raise FormatMonographError("header_shading_hex must be a six-digit hex color.")
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), value)


def _format_table_text(
    table: Any,
    properties: dict[str, Any],
    header_rows: list[int],
    caption_row: int | None = None,
) -> None:
    vertical = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }
    vertical_name = str(properties.get("vertical_alignment", "center"))
    if vertical_name not in vertical:
        raise FormatMonographError(
            f"Unsupported table vertical alignment: {vertical_name}"
        )
    roles = list(properties.get("column_roles", []))
    if roles and len(roles) != len(table.columns):
        raise FormatMonographError("column_roles must match the table column count.")
    role_alignments = {
        "numeric": "right",
        "unit": "center",
        "short_code": "center",
        "narrative": "left",
    }
    role_alignments.update(properties.get("column_alignments", {}))
    for row_index, row in enumerate(table.rows):
        if caption_row is not None and row_index == caption_row:
            continue
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = vertical[vertical_name]
            if row_index in header_rows:
                _set_cell_shading(cell, properties.get("header_shading_hex"))
            alignment = "center" if row_index in header_rows else role_alignments.get(
                roles[column_index] if roles else "narrative", "left"
            )
            if alignment not in ALIGNMENTS:
                raise FormatMonographError(f"Unsupported table cell alignment: {alignment}")
            for paragraph in cell.paragraphs:
                paragraph.alignment = ALIGNMENTS[alignment]
                if "line_spacing_pt" in properties:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    paragraph.paragraph_format.line_spacing = Pt(
                        float(properties["line_spacing_pt"])
                    )
                for run in paragraph.runs:
                    if "font_name_ascii" in properties:
                        _set_font_attributes(
                            run._element,
                            str(properties["font_name_ascii"]),
                            ("ascii", "hAnsi"),
                        )
                        if "font_name_complex_script" not in properties:
                            _set_font_attributes(
                                run._element,
                                str(properties["font_name_ascii"]),
                                ("cs",),
                            )
                    if "font_name_east_asia" in properties:
                        _set_font_attributes(
                            run._element,
                            str(properties["font_name_east_asia"]),
                            ("eastAsia",),
                        )
                    if "font_name_complex_script" in properties:
                        _set_font_attributes(
                            run._element,
                            str(properties["font_name_complex_script"]),
                            ("cs",),
                        )
                    if "font_size_pt" in properties:
                        run.font.size = Pt(float(properties["font_size_pt"]))
                    if row_index in header_rows and "header_bold" in properties:
                        run.bold = bool(properties["header_bold"])


def apply_table_properties(
    document: Any,
    properties: dict[str, Any],
    targets: list[tuple[Any, dict[str, Any]]] | None = None,
) -> int:
    selected = targets if targets is not None else [(table, {}) for table in document.tables]
    for table, entry in selected:
        effective = _table_effective_properties(properties, entry)
        if "table_style" in effective:
            table.style = effective["table_style"]
        if "alignment" in effective:
            value = effective["alignment"]
            if value not in TABLE_ALIGNMENTS:
                raise FormatMonographError(f"Unsupported table alignment: {value}")
            table.alignment = TABLE_ALIGNMENTS[value]
        if "available_width_percent" in effective:
            _set_table_width_percent(table, float(effective["available_width_percent"]))
        if "preferred_column_widths_percent" in effective:
            _set_column_widths_percent(
                table, list(effective["preferred_column_widths_percent"])
            )
        if "allow_autofit" in effective:
            table.autofit = bool(effective["allow_autofit"])
        if "cell_margins_mm" in effective:
            _set_table_cell_margins(table, effective["cell_margins_mm"])
        if "repeat_header_row" in effective and table.rows:
            header_rows = entry.get("repeat_header_rows", [0])
            for row_index in header_rows:
                if 0 <= int(row_index) < len(table.rows):
                    _set_repeat_table_header(
                        table.rows[int(row_index)], bool(effective["repeat_header_row"])
                    )
        if "prevent_row_split" in effective:
            caption_row = entry.get("caption_row")
            for row_index, row in enumerate(table.rows):
                if caption_row is not None and row_index == int(caption_row):
                    continue
                _set_prevent_row_split(row, bool(effective["prevent_row_split"]))
        header_rows = [int(value) for value in entry.get("header_rows", [0])]
        _set_table_borders(
            table, str(effective.get("border_preset", "preserve")), header_rows
        )
        caption_row = entry.get("caption_row")
        _format_table_text(
            table,
            effective,
            header_rows,
            None if caption_row is None else int(caption_row),
        )
    return len(selected)



def _record_derived_change(document: Any, change: dict[str, Any]) -> None:
    changes = getattr(document, "_format_monograph_derived_changes", None)
    if changes is None:
        changes = []
        setattr(document, "_format_monograph_derived_changes", changes)
    changes.append(change)


def _set_update_fields_on_open(document: Any, enabled: bool) -> None:
    settings = document.settings.element
    element = settings.find(qn("w:updateFields"))
    if enabled:
        if element is None:
            element = OxmlElement("w:updateFields")
            settings.append(element)
        element.set(qn("w:val"), "true")
    elif element is not None:
        settings.remove(element)


def _mark_fields_dirty(document: Any) -> int:
    root = document.element
    fields = root.xpath(".//w:fldSimple | .//w:fldChar[@w:fldCharType='begin']")
    for field in fields:
        field.set(qn("w:dirty"), "true")
    return len(fields)


def _clear_paragraph_content(paragraph: Any) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def _append_simple_field(paragraph: Any, instruction: str, placeholder: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    field.set(qn("w:dirty"), "true")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = placeholder
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _field_instruction_for_marker(marker: str) -> tuple[str, str] | None:
    fixed = {
        "[[TOC]]": ('TOC \\o "1-3" \\h \\z \\u', "Update table of contents"),
        "[[PAGE]]": ("PAGE", "1"),
    }
    if marker in fixed:
        return fixed[marker]
    match = re.fullmatch(r"\[\[(REF|PAGEREF|SEQ):([A-Za-z0-9_.-]+)\]\]", marker)
    if not match:
        return None
    kind, value = match.groups()
    if kind == "REF":
        return f"REF {value} \\h", "0"
    if kind == "PAGEREF":
        return f"PAGEREF {value} \\h", "0"
    return f"SEQ {value} \\* ARABIC \\s 1", "1"


def _text_run_like(source_run: Any, value: str) -> Any:
    run = OxmlElement("w:r")
    r_pr = source_run._r.find(qn("w:rPr"))
    if r_pr is not None:
        run.append(copy.deepcopy(r_pr))
    text = OxmlElement("w:t")
    if value[:1].isspace() or value[-1:].isspace():
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = value
    run.append(text)
    return run


def _simple_field_like(source_run: Any, instruction: str, placeholder: str) -> Any:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    field.set(qn("w:dirty"), "true")
    field.append(_text_run_like(source_run, placeholder))
    return field


def _field_reference_target(marker: str) -> str | None:
    match = re.fullmatch(r"\[\[(?:REF|PAGEREF):([A-Za-z0-9_.-]+)\]\]", marker)
    return None if match is None else match.group(1)


def _convert_explicit_field_markers(document: Any) -> int:
    converted = 0
    bookmarks = set(document.element.xpath(".//w:bookmarkStart/@w:name"))
    for index, paragraph in enumerate(iter_document_paragraphs(document)):
        paragraph_text = paragraph.text
        matches = list(FIELD_MARKER_PATTERN.finditer(paragraph_text))
        if not matches:
            continue

        runs = list(paragraph.runs)
        ranges = []
        cursor = 0
        for run in runs:
            start = cursor
            cursor += len(run.text)
            ranges.append((start, cursor, run))

        assignments: dict[int, list[tuple[re.Match[str], Any]]] = {}
        for match in matches:
            marker = match.group(0)
            if marker == "[[TOC]]" and paragraph_text.strip() != marker:
                raise FormatMonographError(
                    f"TOC marker must occupy its own paragraph at paragraph {index}."
                )
            target = _field_reference_target(marker)
            if target is not None and target not in bookmarks:
                raise FormatMonographError(
                    f"Field marker {marker} references missing bookmark {target} "
                    f"at paragraph {index}."
                )
            containing = [
                (start, run)
                for start, end, run in ranges
                if start <= match.start() and match.end() <= end
            ]
            if len(containing) != 1:
                raise FormatMonographError(
                    f"Field marker spans multiple runs at paragraph {index}: {marker}"
                )
            start, run = containing[0]
            assignments.setdefault(id(run._r), []).append((match, run))

        for _, assigned in assignments.items():
            source_run = assigned[0][1]
            run_start = next(start for start, _, run in ranges if run is source_run)
            parent = source_run._r.getparent()
            insertion = parent.index(source_run._r)
            local_cursor = 0
            for match, _ in assigned:
                local_start = match.start() - run_start
                local_end = match.end() - run_start
                prefix = source_run.text[local_cursor:local_start]
                if prefix:
                    parent.insert(insertion, _text_run_like(source_run, prefix))
                    insertion += 1
                marker = match.group(0)
                instruction, placeholder = _field_instruction_for_marker(marker)
                parent.insert(
                    insertion,
                    _simple_field_like(source_run, instruction, placeholder),
                )
                insertion += 1
                _record_derived_change(
                    document,
                    {
                        "kind": "field_marker",
                        "paragraph": index,
                        "source": marker,
                        "field": instruction,
                    },
                )
                converted += 1
                local_cursor = local_end
            suffix = source_run.text[local_cursor:]
            if suffix:
                parent.insert(insertion, _text_run_like(source_run, suffix))
            parent.remove(source_run._r)
    return converted


def _remove_prefix_from_runs(paragraph: Any, length: int) -> None:
    remaining = length
    for run in paragraph.runs:
        if remaining <= 0:
            break
        text = run.text
        take = min(len(text), remaining)
        run.text = text[take:]
        remaining -= take
    if remaining:
        raise FormatMonographError("Heading prefix spans an unsupported non-text object.")


def _heading_prefix_pattern(level: int) -> re.Pattern[str]:
    if level == 1:
        return re.compile(r"^\s*第\s*[0-9一二三四五六七八九十百]+\s*章\s*")
    separators = level - 1
    return re.compile(rf"^\s*\d+(?:[.-]\d+){{{separators}}}\s*")


def _strip_manual_heading_prefixes(document: Any, levels: int) -> int:
    changed = 0
    for index, paragraph in enumerate(document.paragraphs):
        if not paragraph.style or not paragraph.style.name.startswith("Heading "):
            continue
        try:
            level = int(paragraph.style.name.split()[-1])
        except ValueError:
            continue
        if level > levels:
            continue
        text = paragraph.text
        match = _heading_prefix_pattern(level).match(text)
        if not match:
            if re.match(r"^\s*(?:第\s*\d+\s*章|\d+[.-]\d+)", text):
                raise FormatMonographError(
                    f"Ambiguous manual heading number at paragraph {index}: {text[:80]}"
                )
            continue
        prefix = match.group(0)
        _remove_prefix_from_runs(paragraph, len(prefix))
        _record_derived_change(
            document,
            {"kind": "heading_prefix", "paragraph": index, "removed": prefix},
        )
        changed += 1
    return changed


def _next_numbering_id(root: Any, tag: str, attribute: str) -> int:
    values = []
    for element in root.findall(qn(f"w:{tag}")):
        value = element.get(qn(f"w:{attribute}"))
        if value is not None and value.isdigit():
            values.append(int(value))
    return max(values, default=0) + 1


def _ensure_heading_numbering(document: Any, levels: int, chapter_start: int = 1) -> int:
    if not 1 <= levels <= 4:
        raise FormatMonographError("heading_levels must be between 1 and 4.")
    if chapter_start < 1:
        raise FormatMonographError("chapter_start must be a positive integer.")
    root = document.part.numbering_part.element
    abstract_id = _next_numbering_id(root, "abstractNum", "abstractNumId")
    num_id = _next_numbering_id(root, "num", "numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(levels):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), str(chapter_start if level == 0 else 1))
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), f"Heading{level + 1}")
        lvl_text = OxmlElement("w:lvlText")
        if level == 0:
            lvl_text.set(qn("w:val"), "第%1章")
        else:
            lvl_text.set(
                qn("w:val"),
                ".".join(f"%{number}" for number in range(1, level + 2)),
            )
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        lvl.extend([start, num_fmt, p_style, lvl_text, suff])
        abstract.append(lvl)
    root.insert(0, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if chapter_start != 1:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(chapter_start))
        override.append(start_override)
        num.append(override)
    root.append(num)

    for level in range(levels):
        style = ensure_paragraph_style(document, f"Heading {level + 1}")
        p_pr = style.element.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        for child_name in ("ilvl", "numId"):
            existing = num_pr.find(qn(f"w:{child_name}"))
            if existing is not None:
                num_pr.remove(existing)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        number = OxmlElement("w:numId")
        number.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, number])
    return levels


def apply_field_properties(document: Any, properties: dict[str, Any]) -> int:
    changed = 0
    if "update_on_open" in properties:
        _set_update_fields_on_open(document, bool(properties["update_on_open"]))
        changed += 1
    if properties.get("mark_fields_dirty"):
        changed += _mark_fields_dirty(document)
    if properties.get("convert_explicit_markers"):
        changed += _convert_explicit_field_markers(document)

    levels = int(properties.get("heading_levels", 4))
    if properties.get("strip_manual_heading_prefixes"):
        changed += _strip_manual_heading_prefixes(document, levels)
    if properties.get("rebuild_heading_numbering"):
        changed += _ensure_heading_numbering(
            document, levels, int(properties.get("chapter_start", 1))
        )
    return changed


def apply_equation_properties(document: Any, properties: dict[str, Any]) -> int:
    unsupported_values = {
        key: value
        for key, value in properties.items()
        if key in EQUATION_PROPERTIES and value not in {True, False}
    }
    if unsupported_values:
        raise FormatMonographError(
            "Equation policy properties must be boolean: "
            + ", ".join(sorted(unsupported_values))
        )
    return len(document.element.xpath(".//*[local-name()='oMath']"))


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _complex_field_instructions(root: etree._Element) -> list[str]:
    instructions: list[str] = []
    for paragraph in root.xpath(".//w:p", namespaces=NS):
        stack: list[dict[str, Any]] = []
        for element in paragraph.iter():
            if element.tag == qn("w:fldChar"):
                kind = element.get(qn("w:fldCharType"))
                if kind == "begin":
                    stack.append({"parts": [], "captured": False})
                elif kind == "separate" and stack and not stack[-1]["captured"]:
                    value = "".join(stack[-1]["parts"]).strip()
                    if value:
                        instructions.append(value)
                    stack[-1]["captured"] = True
                elif kind == "end" and stack:
                    field = stack.pop()
                    if not field["captured"]:
                        value = "".join(field["parts"]).strip()
                        if value:
                            instructions.append(value)
            elif (
                element.tag == qn("w:instrText")
                and stack
                and not stack[-1]["captured"]
            ):
                stack[-1]["parts"].append(element.text or "")
    return instructions


def field_inventory(path: Path) -> dict[str, Any]:
    ensure_docx(path)
    counts: dict[str, int] = {}
    instructions: list[str] = []
    bookmarks: set[str] = set()
    with zipfile.ZipFile(path) as package:
        for name in package.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = etree.fromstring(package.read(name))
            bookmarks.update(
                root.xpath(".//w:bookmarkStart/@w:name", namespaces=NS)
            )
            instructions.extend(
                value.strip()
                for value in root.xpath(".//w:fldSimple/@w:instr", namespaces=NS)
                if value.strip()
            )
            instructions.extend(_complex_field_instructions(root))

    references = []
    sequences = []
    for instruction in instructions:
        match = re.match(r"\s*([A-Za-z]+)(?:\s+([^\\\s]+))?", instruction)
        kind = match.group(1).upper() if match else "UNKNOWN"
        argument = match.group(2) if match else None
        counts[kind] = counts.get(kind, 0) + 1
        if kind in {"REF", "PAGEREF"} and argument:
            references.append({"type": kind, "target": argument})
        if kind == "SEQ" and argument:
            sequences.append(argument)

    unresolved = sorted(
        {
            item["target"]
            for item in references
            if item["target"] not in bookmarks
        }
    )
    return {
        "total": len(instructions),
        "types": dict(sorted(counts.items())),
        "bookmarks": sorted(bookmarks),
        "references": references,
        "unresolved_references": unresolved,
        "sequences": sorted(set(sequences)),
    }


def field_cache_inventory(path: Path) -> dict[str, Any]:
    inventory = field_inventory(path)
    toc_fields = int(inventory["types"].get("TOC", 0))
    toc_entries = 0
    dirty_fields = 0
    update_on_open = False
    with zipfile.ZipFile(path) as package:
        document = etree.fromstring(package.read("word/document.xml"))
        toc_entries = len(
            document.xpath(
                ".//w:p[starts-with(translate(./w:pPr/w:pStyle/@w:val, 'toc', 'TOC'), 'TOC')]",
                namespaces=NS,
            )
        )
        dirty_fields = len(
            document.xpath(
                ".//w:fldSimple[@w:dirty='true' or @w:dirty='1'] | "
                ".//w:fldChar[@w:fldCharType='begin'][@w:dirty='true' or @w:dirty='1']",
                namespaces=NS,
            )
        )
        if "word/settings.xml" in package.namelist():
            settings = etree.fromstring(package.read("word/settings.xml"))
            update_on_open = bool(settings.xpath(".//w:updateFields", namespaces=NS))

    if toc_fields == 0:
        status = "absent"
    elif toc_entries == 0:
        status = "code_only"
    elif dirty_fields:
        status = "stale"
    else:
        status = "refreshed"
    return {
        "status": status,
        "main_toc_fields": toc_fields,
        "toc_entries": toc_entries,
        "dirty_fields": dirty_fields,
        "update_on_open": update_on_open,
        "field_types": inventory["types"],
    }


def equation_inventory(path: Path) -> dict[str, Any]:
    ensure_docx(path)
    result = {
        "omml": 0,
        "mathtype_ole": 0,
        "legacy_equation_ole": 0,
        "other_ole": 0,
        "formula_image_candidates": 0,
    }
    with zipfile.ZipFile(path) as package:
        for name in package.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = etree.fromstring(package.read(name))
            result["omml"] += len(root.xpath(".//m:oMath", namespaces=NS))
            for ole in root.xpath(".//*[local-name()='OLEObject']"):
                prog_id = (ole.get("ProgID") or ole.get(f"{{{O_NS}}}ProgID") or "").lower()
                if "dsmt" in prog_id or "mathtype" in prog_id:
                    result["mathtype_ole"] += 1
                elif prog_id in {"equation.3", "equation.2", "equation"}:
                    result["legacy_equation_ole"] += 1
                else:
                    result["other_ole"] += 1

            for paragraph in root.xpath(".//w:p", namespaces=NS):
                if not paragraph.xpath(".//w:drawing | .//v:imagedata", namespaces=NS):
                    continue
                styles = paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
                metadata = paragraph.xpath(
                    ".//*[local-name()='docPr']/@name | "
                    ".//*[local-name()='docPr']/@title | "
                    ".//*[local-name()='docPr']/@descr"
                )
                hint = " ".join(styles + metadata).lower()
                if any(token in hint for token in ("equation", "formula", "公式")):
                    result["formula_image_candidates"] += 1
    result["editable_total"] = (
        result["omml"] + result["mathtype_ole"] + result["legacy_equation_ole"]
    )
    return result


def protected_object_manifest(path: Path) -> dict[str, Any]:
    ensure_docx(path)
    embeddings: dict[str, str] = {}
    media: dict[str, str] = {}
    omml: list[str] = []
    with zipfile.ZipFile(path) as package:
        for name in sorted(package.namelist()):
            data = package.read(name)
            if name.startswith("word/embeddings/"):
                embeddings[name] = _sha256_bytes(data)
            elif name.startswith("word/media/"):
                media[name] = _sha256_bytes(data)
            if name.startswith("word/") and name.endswith(".xml"):
                root = etree.fromstring(data)
                for equation in root.xpath(".//m:oMath", namespaces=NS):
                    canonical = etree.tostring(equation, method="c14n", exclusive=True)
                    omml.append(_sha256_bytes(canonical))
    return {
        "embeddings": embeddings,
        "media": media,
        "omml_sha256": sorted(omml),
    }


def protected_payload_manifest(path: Path) -> dict[str, list[str]]:
    """Compare protected payloads independently of package part renaming."""
    manifest = protected_object_manifest(path)
    return {
        "embeddings": sorted(manifest["embeddings"].values()),
        "media": sorted(manifest["media"].values()),
        "omml_sha256": sorted(manifest["omml_sha256"]),
    }

def apply_rule(
    document: Any,
    rule: dict[str, Any],
    *,
    paragraph_targets: list[Any] | None = None,
    table_targets: list[tuple[Any, dict[str, Any]]] | None = None,
    chapter_start: int | None = None,
) -> int:
    unsupported = unsupported_properties(rule)
    if unsupported:
        raise FormatMonographError(
            f"Rule {rule['id']} has unsupported automatic properties: {', '.join(unsupported)}"
        )

    selector = rule["selector"]
    if selector["kind"] in {"document", "section_role"}:
        return apply_section_properties(document, rule["properties"])
    if selector["kind"] == "table_role":
        return apply_table_properties(document, rule["properties"], table_targets)
    if selector["kind"] == "field_role":
        properties = dict(rule["properties"])
        if chapter_start is not None:
            properties["chapter_start"] = chapter_start
        return apply_field_properties(document, properties)
    if selector["kind"] == "equation_role":
        return apply_equation_properties(document, rule["properties"])

    style_name = style_name_for_selector(selector)
    if style_name is None:
        raise FormatMonographError(
            f"Rule {rule['id']} uses an unsupported automatic selector: {selector}"
        )
    if paragraph_targets is not None:
        return apply_style_rule_to_paragraphs(document, rule, paragraph_targets)
    style = ensure_paragraph_style(document, style_name)
    apply_style_properties(style, rule["properties"])
    return sum(
        1
        for paragraph in iter_document_paragraphs(document)
        if paragraph.style and paragraph.style.name == style_name
    )


def first_anchor_paragraph(document: Any, rule: dict[str, Any]) -> Any | None:
    style_name = style_name_for_selector(rule["selector"])
    paragraphs = list(document.paragraphs)
    if style_name:
        for paragraph in paragraphs:
            if (
                paragraph.text.strip()
                and paragraph.style
                and paragraph.style.name == style_name
                and paragraph.runs
            ):
                return paragraph
    for paragraph in paragraphs:
        if paragraph.text.strip() and paragraph.runs:
            return paragraph
    return None


def summarize_rule(rule: dict[str, Any]) -> str:
    pairs = ", ".join(f"{key}={value}" for key, value in sorted(rule["properties"].items()))
    sources = ", ".join(rule["source_ids"])
    return f"{rule['id']}: {pairs}. Sources: {sources}."


def load_document(path: Path) -> Any:
    ensure_docx(path)
    try:
        return Document(str(path))
    except Exception as exc:
        raise FormatMonographError(f"Unable to open DOCX {path}: {exc}") from exc
