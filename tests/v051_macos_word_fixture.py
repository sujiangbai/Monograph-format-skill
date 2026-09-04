"""Generate this batch's synthetic input; never read a user document or image."""
from __future__ import annotations

import argparse
import hashlib
import io
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from lxml import etree
from PIL import Image, ImageDraw

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
HEADINGS = ("Chapter Alpha", "Chapter Beta")
APPROVED = {"TOC", "PAGE", "NUMPAGES", "SECTIONPAGES", "PAGEREF", "REF"}


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def field(paragraph, code: str, cache: str) -> None:
    for kind in ("begin", "instruction", "separate", "result", "end"):
        run = paragraph.add_run()
        if kind == "result":
            run.text = cache
        elif kind == "instruction":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = f" {code} "
            run._r.append(node)
        else:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
            if kind == "begin":
                node.set(qn("w:dirty"), "false")
            run._r.append(node)


def bookmark(paragraph, name: str, identifier: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(identifier))
    start.set(qn("w:name"), name)
    paragraph._p.insert(1 if paragraph._p.pPr is not None else 0, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(identifier))
    paragraph._p.append(end)


def inspect_synthetic_package(path: Path) -> dict:
    """A small generated-fixture assertion, not a general DOCX security scanner."""
    fields = []
    with zipfile.ZipFile(path) as package:
        names = package.namelist()
        if len(names) != len(set(names)):
            raise ValueError("duplicate package members")
        for name in names:
            if re.search(r"vba|activex|embeddings/|\.bin$", name, re.I):
                raise ValueError("macro/active payload in synthetic fixture")
            if not name.endswith((".xml", ".rels")):
                continue
            data = package.read(name)
            if b"<!DOCTYPE" in data or b"<!ENTITY" in data:
                raise ValueError("XML entity in synthetic fixture")
            root = etree.fromstring(data)
            for node in root.iter():
                if node.get("TargetMode", "").lower() == "external":
                    raise ValueError("external relationship in synthetic fixture")
                for key, value in node.attrib.items():
                    local = etree.QName(key).localname
                    if local in {"Target", "href", "src", "link"} and (
                        re.match(r"^[A-Za-z][A-Za-z0-9+.-]*:", value)
                        or value.startswith(("//", "\\\\"))
                    ):
                        raise ValueError("external target in synthetic fixture")
                    if local == "ContentType" and re.search(r"macroEnabled|vba", value, re.I):
                        raise ValueError("macro content type in synthetic fixture")
            if name.startswith("word/"):
                for ordinal, instruction in enumerate(root.xpath("//w:instrText", namespaces=NS)):
                    exact = instruction.text or ""
                    kind = exact.strip().split()[0]
                    if kind not in APPROVED | {"QUOTE"}:
                        raise ValueError("unexplained field in synthetic fixture")
                    fields.append({"part": name, "ordinal": ordinal, "instruction": exact,
                                   "approved_for_later_refresh": kind in APPROVED})
        document = etree.fromstring(package.read("word/document.xml"))
        return {"status": "pass", "macro_payloads": 0, "external_connections": 0,
                "section_count": len(document.xpath("//w:sectPr", namespaces=NS)),
                "table_count": len(document.xpath("//w:tbl", namespaces=NS)),
                "image_count": sum(name.startswith("word/media/") for name in names),
                "fields": fields}


def create_fixture(directory: Path) -> dict:
    directory.mkdir(parents=True, exist_ok=False)
    output = directory / "v051-synthetic.docx"
    document = Document()  # python-docx's packaged blank template, not a local file.
    document.core_properties.author = "Synthetic capability probe"
    document.core_properties.last_modified_by = "Synthetic capability probe"
    # standard_business_brief tokens; no lists, metadata grid or decorative cover.
    styles = [("Normal", 11, "000000", 0, 6), ("Heading 1", 16, "2E74B5", 16, 8),
              ("Heading 2", 13, "2E74B5", 12, 6), ("Heading 3", 12, "1F4D78", 8, 4),
              ("Title", 23, "000000", 0, 4), ("Footer", 9, "555555", 0, 0)]
    for name, size, color, before, after in styles:
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
    document.add_paragraph("Synthetic Word probe", style="Title")
    document.add_paragraph("All text, data and graphics are generated by this fixture.")
    document.add_paragraph("Contents")
    field(document.add_paragraph(), 'TOC \\o "1-1"', "Chapter Alpha\t1\nChapter Beta\t2")
    front = document.sections[0]
    front.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    field(front.footer.paragraphs[0], "PAGE", "i")
    document.add_section(WD_SECTION.NEW_PAGE)
    body = document.sections[1]
    body.footer.is_linked_to_previous = False
    body.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    field(body.footer.paragraphs[0], "PAGE", "1")
    for section, numbering in zip(document.sections, ("lowerRoman", "decimal")):
        section.page_width, section.page_height = Inches(8.5), Inches(11)
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
        section.header_distance = section.footer_distance = Inches(0.492)
        page_number = OxmlElement("w:pgNumType")
        page_number.set(qn("w:start"), "1")
        page_number.set(qn("w:fmt"), numbering)
        section._sectPr.append(page_number)
    alpha = document.add_paragraph(HEADINGS[0], style="Heading 1")
    bookmark(alpha, "probe_alpha", 1)
    document.add_paragraph("Alpha preserves two  spaces, a tab\tand a line\nbreak.")
    picture = Image.new("RGB", (160, 96), "white")
    drawing = ImageDraw.Draw(picture)
    drawing.rectangle((8, 8, 70, 88), fill="#2E74B5")
    drawing.rectangle((90, 8, 152, 88), fill="#E8EEF5", outline="#2E74B5", width=2)
    buffer = io.BytesIO()
    picture.save(buffer, format="PNG")
    document.add_picture(io.BytesIO(buffer.getvalue()), width=Inches(2.0))
    document.add_paragraph("Synthetic image: two rectangles generated in memory.")
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = (2700, 6660)
    table._tbl.tblPr.find(qn("w:tblW")).set(qn("w:w"), "9360")
    table._tbl.tblPr.find(qn("w:tblW")).set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    table._tbl.tblPr.append(indent)
    margins = OxmlElement("w:tblCellMar")
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        margin = OxmlElement("w:" + side)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    table._tbl.tblPr.append(margins)
    for col, width in zip(table.columns, widths):
        col.width = Twips(width)
    for row, values in zip(table.rows, (("Item", "Synthetic value"), ("A", "11"), ("B", "22"))):
        for cell, width, value in zip(row.cells, widths, values):
            cell.width = Twips(width)
            cell.text = value
    for cell in table.rows[0].cells:
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "F2F4F7")
        cell._tc.get_or_add_tcPr().append(shade)
    beta = document.add_paragraph(HEADINGS[1], style="Heading 1")
    beta.paragraph_format.page_break_before = True  # named fixture override: two body pages.
    bookmark(beta, "probe_beta", 2)
    document.add_paragraph("Beta refers only to bookmarks in this generated document.")
    for label, code, cache in (("All pages: ", "NUMPAGES", "3"),
                               ("Body pages: ", "SECTIONPAGES", "2"),
                               ("Alpha page: ", "PAGEREF probe_alpha", "1"),
                               ("Alpha heading: ", "REF probe_alpha", HEADINGS[0]),
                               ("Unapproved field: ", 'QUOTE "unapproved_constant"', "DO NOT REFRESH")):
        paragraph = document.add_paragraph(label)
        field(paragraph, code, cache)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "false")
    document.settings.element.append(update)
    document.save(output)
    manifest = {"fixture": "v051-synthetic-v1", "docx": output.name,
                "sha256": digest(output), "size_bytes": output.stat().st_size,
                "source": "new literal text/data and in-memory procedural PNG only",
                "current_step_approved_updates": [],
                "later_toc_sources": [{"level": 1, "kind": "heading", "text_sha256": hashlib.sha256(text.encode()).hexdigest()}
                                      for text in HEADINGS],
                "package_preflight": inspect_synthetic_package(output)}
    (directory / "fixture.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    return manifest


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--out-dir", type=Path, required=True)
    args = parser.parse_args()
    print(json.dumps(create_fixture(args.out_dir), indent=2))
