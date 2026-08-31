from __future__ import annotations

import base64
import copy
import contextlib
import hashlib
import io
import json
import os
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "format-monograph"
SCRIPTS = SKILL / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import NS  # noqa: E402
import backend_evidence  # noqa: E402
import finalize_docx  # noqa: E402
from backend_evidence import (  # noqa: E402
    BackendEvidenceError,
    backend_audit_binding,
    backend_audit_bytes,
    backend_audit_path,
    canonical_backend_projection,
    canonical_backend_shape_errors,
    read_bound_backend_audit,
)
from field_completion import (  # noqa: E402
    FINALIZATION_EVIDENCE_VERSION,
    completion_evidence,
    final_ready_evidence_errors,
    finalization_evidence_shape_errors,
)
from external_command import (  # noqa: E402
    external_command_cache_reusable,
    external_command_identity,
    external_command_identity_errors,
)
from target_software import (  # noqa: E402
    LIBREOFFICE,
    MICROSOFT_WORD,
    UNSUPPORTED,
    resolve_target_id,
)
from finalize_docx import controlled_field_result_writeback  # noqa: E402
from run_monograph import (  # noqa: E402
    RunError,
    VERIFICATION_OUTPUT_VERSION,
    artifact_binding_errors,
    canonical_finalization_gate_summary,
    file_sha256,
    finalize as orchestrate_finalize,
    finalize_request_identity,
    finalization_consistency_errors,
    final_ready_field_evidence,
    has_target_layout_evidence,
    json_sha256,
    local_artifact_identity,
    render_page_count_errors,
    status as orchestrate_status,
    validate_visual_manifest,
    verify,
    verify_request_identity,
)
from structure_map import (  # noqa: E402
    candidate_structure_map,
    load_structure_map,
)


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def add_page_field(paragraph, value: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    instruction_run._r.append(instruction)
    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    paragraph.add_run(value)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_toc_field(paragraph) -> None:
    add_page_field(paragraph, "Update table of contents")
    instruction = paragraph._p.xpath(".//w:instrText")[0]
    instruction.text = ' TOC \\o "1-4" \\h \\z '


def rewrite_package(source: Path, output: Path, transform) -> None:
    with zipfile.ZipFile(source) as package, zipfile.ZipFile(
        output, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for info in package.infolist():
            target.writestr(info, transform(info.filename, package.read(info.filename)))


class V030WholeBookRuntimeTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def _field_cache(self, status: str) -> dict:
        has_toc = status != "absent"
        return {
            "status": status,
            "main_toc_fields": 1 if has_toc else 0,
            "toc_entries": 1 if has_toc and status != "code_only" else 0,
            "dirty_fields": 1 if status == "stale" else 0,
            "update_on_open": status == "stale",
            "field_types": {"TOC": 1} if has_toc else {},
        }

    def _persist_backend_audit(
        self, status_path: Path, raw_backend: dict
    ) -> tuple[dict, dict, Path]:
        payload = backend_audit_bytes(raw_backend)
        path = backend_audit_path(status_path)
        path.write_bytes(payload)
        return (
            canonical_backend_projection(raw_backend),
            backend_audit_binding(path, payload),
            path,
        )

    def _word_completion_evidence(self) -> dict:
        return {
            "delivery_status": "selective_verified",
            "input_cache_status": "stale",
            "output_cache_status": "refreshed",
            "backend": "external",
            "field_cache_verified": True,
            "calculation_page_count": 7,
            "writeback_status": "selective_verified",
            "selective_writeback_status": "selective_verified",
            "read_only_operation": "verify_only",
            "read_only_verified": True,
            "read_only_repaginated": True,
            "read_only_saved": False,
            "verification_pdf_exported": True,
            "verification_page_count": 7,
            "artifact_binding": {
                "version": 1,
                "finalized_docx": {
                    "path": "/tmp/finalized.docx",
                    "sha256": "a" * 64,
                    "size_bytes": 100,
                },
                "word_verification_pdf": {
                    "path": "/tmp/verification.pdf",
                    "sha256": "b" * 64,
                    "size_bytes": 100,
                    "page_count": 7,
                },
            },
            "word_verification_required": True,
            "word_verification_completed": True,
            "completion_scope": "target_word_verified",
            "field_gate_completed": True,
            "final_ready_eligible": True,
        }

    def test_complete_no_fields_and_word_evidence_are_the_only_valid_shapes(self) -> None:
        no_fields = {
            "delivery_status": "absent",
            "input_cache_status": "absent",
            "output_cache_status": "absent",
            "backend": "not_needed",
            "field_cache_verified": None,
            "calculation_page_count": None,
            "writeback_status": "not_needed",
            "selective_writeback_status": None,
            "read_only_operation": None,
            "read_only_verified": None,
            "read_only_repaginated": None,
            "read_only_saved": None,
            "verification_pdf_exported": None,
            "verification_page_count": None,
            "artifact_binding": {
                "version": 1,
                "finalized_docx": {
                    "path": "/tmp/finalized.docx",
                    "sha256": "a" * 64,
                    "size_bytes": 100,
                },
                "word_verification_pdf": None,
            },
            "word_verification_required": False,
            "word_verification_completed": False,
            "completion_scope": "no_fields",
            "field_gate_completed": True,
            "final_ready_eligible": True,
        }
        self.assertTrue(final_ready_field_evidence(no_fields))
        self.assertTrue(final_ready_field_evidence(self._word_completion_evidence()))

    def test_contradictory_final_ready_evidence_matrix_is_rejected(self) -> None:
        cases = {}
        libreoffice = self._word_completion_evidence()
        libreoffice.update(
            {
                "backend": "libreoffice_uno",
                "writeback_status": "libreoffice_selective",
                "selective_writeback_status": "libreoffice_selective",
            }
        )
        cases["selective_verified_plus_libreoffice"] = libreoffice
        deferred = self._word_completion_evidence()
        deferred.update(
            {
                "backend": "deferred_on_open",
                "writeback_status": "deferred",
                "selective_writeback_status": None,
            }
        )
        cases["selective_verified_plus_deferred"] = deferred
        missing_verification = self._word_completion_evidence()
        missing_verification["read_only_verified"] = False
        missing_verification["word_verification_completed"] = False
        cases["missing_word_verification"] = missing_verification
        wrong_scope = self._word_completion_evidence()
        wrong_scope["completion_scope"] = "libreoffice_non_final"
        cases["wrong_completion_scope"] = wrong_scope
        wrong_page_count = self._word_completion_evidence()
        wrong_page_count["verification_page_count"] = 8
        cases["mismatched_word_page_count"] = wrong_page_count
        for input_status, output_status in (
            ("absent", "refreshed"),
            ("stale", "absent"),
            ("stale", "deferred"),
            ("stale", "stale"),
            ("unknown", "refreshed"),
        ):
            changed_cache = self._word_completion_evidence()
            changed_cache["input_cache_status"] = input_status
            changed_cache["output_cache_status"] = output_status
            cases[f"cache_{input_status}_{output_status}"] = changed_cache
        for name, evidence in cases.items():
            with self.subTest(name=name):
                self.assertFalse(final_ready_field_evidence(evidence))
                self.assertTrue(final_ready_evidence_errors(evidence))

    def test_verify_reloads_full_finalization_and_rejects_contradictory_state(self) -> None:
        status_path = self.root / "field-finalization.json"
        finalization = {
            "finalization_evidence_version": FINALIZATION_EVIDENCE_VERSION,
            "delivery_field_status": "selective_verified",
            "input_field_cache": self._field_cache("stale"),
            "output_field_cache": self._field_cache("refreshed"),
            "field_writeback_status": "libreoffice_selective",
            "field_backend": {
                "backend": "libreoffice_uno",
                "field_cache_verified": True,
                "selective_writeback": {"status": "libreoffice_selective"},
                "read_only_verification": {
                    "operation": "verify_only",
                    "read_only_verified": True,
                    "repaginated": True,
                    "saved": False,
                    "pdf_exported": True,
                },
            },
            "field_completion": {
                "word_verification_required": True,
                "word_verification_completed": True,
                "completion_scope": "target_word_verified",
                "field_gate_completed": True,
                "final_ready_eligible": True,
                "evidence_validation": {"status": "pass", "errors": []},
            },
        }
        status_path.write_text(json.dumps(finalization), encoding="utf-8")
        canonical = completion_evidence(finalization)
        state = {
            "status": "candidate_ready",
            "artifacts": {"finalization_status": status_path.name},
            "field_writeback": {"completion_evidence": canonical},
            "blockers": [],
        }
        args = SimpleNamespace(work_dir=self.root)
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.current_inputs",
            return_value=(self.root / "source.docx", self.root / "profile.json"),
        ), patch("run_monograph.save_state"):
            self.assertEqual(2, verify(args))
        self.assertEqual("candidate_ready", state["status"])
        self.assertTrue(
            state["field_writeback"]["completion_evidence_errors"]
        )

    def _bound_word_finalization(self) -> tuple[dict, dict]:
        source = self.root / "source.docx"
        profile = self.root / "profile.json"
        structure = self.root / "structure.json"
        formatted = self.root / "formatted.docx"
        finalized = self.root / "finalized.docx"
        pdf = self.root / "verification.pdf"
        source.write_bytes(b"source")
        profile.write_text("{}", encoding="utf-8")
        structure.write_text("{}", encoding="utf-8")
        formatted.write_bytes(b"formatted")
        finalized.write_bytes(b"bound finalized docx")
        pdf.write_bytes(b"%PDF-bound verification")
        updater = self.root / "external-updater"
        updater.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        updater.chmod(0o755)
        identity = lambda path: {
            "path": str(path.resolve()),
            "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
            "size_bytes": path.stat().st_size,
        }
        binding = {
            "version": 1,
            "finalized_docx": identity(finalized),
            "word_verification_pdf": {**identity(pdf), "page_count": 7},
        }
        status_path = self.root / "finalization.json"
        raw_backend = {
            "backend": "external",
            "target_id": "microsoft_word",
            "field_cache_verified": True,
            "page_count": 7,
            "selective_writeback": {"status": "selective_verified"},
            "read_only_verification": {
                "target_id": "microsoft_word",
                "operation": "verify_only",
                "read_only_verified": True,
                "repaginated": True,
                "saved": False,
                "pdf_exported": True,
                "page_count": 7,
            },
        }
        canonical_backend, audit_binding, audit_path = self._persist_backend_audit(
            status_path, raw_backend
        )
        finalization = {
            "finalization_evidence_version": FINALIZATION_EVIDENCE_VERSION,
            "delivery_field_status": "selective_verified",
            "input_field_cache": self._field_cache("stale"),
            "output_field_cache": self._field_cache("refreshed"),
            "field_writeback_status": "selective_verified",
            "field_backend": canonical_backend,
            "backend_audit": audit_binding,
            "field_completion": {
                "word_verification_required": True,
                "word_verification_completed": True,
                "completion_scope": "target_word_verified",
                "field_gate_completed": True,
                "final_ready_eligible": True,
            },
            "artifact_binding": binding,
            "status": "pass",
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
            "workflow_state": {
                "stage": "finalized",
                "source_sha256": file_sha256(source),
                "input_sha256": file_sha256(formatted),
                "profile_sha256": file_sha256(profile),
                "structure_map_sha256": file_sha256(structure),
                "output_sha256": binding["finalized_docx"]["sha256"],
            },
            "output": str(finalized.resolve()),
            "target_pdf": str(pdf.resolve()),
            "target_layout_status": "target_pdf_ready_for_visual_qa",
            "target_software": "microsoft_word",
        }
        finalization["field_completion"]["evidence_validation"] = {
            "status": "pass",
            "errors": [],
        }
        canonical = completion_evidence(finalization)
        state = {
            "source": {"path": str(source), "sha256": file_sha256(source)},
            "profile": {"path": str(profile), "sha256": file_sha256(profile)},
            "structure_map": {
                "path": str(structure),
                "sha256": file_sha256(structure),
            },
            "artifacts": {
                "formatted": formatted.name,
                "finalized": finalized.name,
                "target_pdf": pdf.name,
                "backend_audit": audit_path.name,
            },
            "field_writeback": {
                "completion_evidence": canonical,
                "artifact_binding": binding,
            },
        }
        state["finalization_gate"] = canonical_finalization_gate_summary(
            self.root, finalization
        )
        state["finalization_request"] = finalize_request_identity(
            SimpleNamespace(
                field_updater="external",
                field_updater_command=str(updater),
                target_software="Microsoft Word",
                renderer=None,
                approve_deferred=False,
            )
        )
        state["stages"] = {
            "finalize": {
                "status": "complete",
                "input_key_sha256": json_sha256(
                    {
                        "formatted": file_sha256(formatted),
                        "map": file_sha256(structure),
                        "behavior": state["finalization_request"],
                    }
                ),
            }
        }
        return finalization, state

    def test_artifact_and_json_tamper_matrix_is_rejected(self) -> None:
        finalization, state = self._bound_word_finalization()
        self.assertEqual([], artifact_binding_errors(self.root, finalization, state))
        self.assertEqual([], finalization_consistency_errors(self.root, finalization, state)[1])
        self.assertEqual([], render_page_count_errors(completion_evidence(finalization), 7))

        finalized = self.root / "finalized.docx"
        original_docx = finalized.read_bytes()
        finalized.write_bytes(b"tampered docx")
        self.assertTrue(artifact_binding_errors(self.root, finalization, state))
        finalized.write_bytes(original_docx)

        pdf = self.root / "verification.pdf"
        original_pdf = pdf.read_bytes()
        pdf.write_bytes(b"tampered pdf")
        self.assertTrue(artifact_binding_errors(self.root, finalization, state))
        pdf.write_bytes(original_pdf)

        changed_json = json.loads(json.dumps(finalization))
        changed_json["field_completion"]["completion_scope"] = "incomplete"
        self.assertTrue(finalization_consistency_errors(self.root, changed_json, state)[1])

        changed_state = json.loads(json.dumps(state))
        changed_state["field_writeback"]["artifact_binding"]["finalized_docx"]["sha256"] = "0" * 64
        self.assertTrue(artifact_binding_errors(self.root, finalization, changed_state))

        self.assertTrue(render_page_count_errors(completion_evidence(finalization), 8))

    def test_finalize_orchestration_rejects_invalid_word_cache_shape(self) -> None:
        work = self.root / "orchestration"
        work.mkdir()
        source = work / "source.docx"
        profile = work / "profile.json"
        structure = work / "structure.json"
        formatted = work / "formatted.docx"
        source.write_bytes(b"source")
        profile.write_text("{}", encoding="utf-8")
        structure.write_text("{}", encoding="utf-8")
        formatted.write_bytes(b"formatted")
        final_dir = work / "final"
        final_dir.mkdir()
        finalized = final_dir / "source-finalized.docx"
        pdf = final_dir / "source-target.pdf"
        finalized.write_bytes(b"finalized")
        pdf.write_bytes(b"%PDF verification")
        identity = lambda path: {
            "path": str(path.resolve()),
            "sha256": file_sha256(path),
            "size_bytes": path.stat().st_size,
        }
        finalization, _ = self._bound_word_finalization()
        finalization["input_field_cache"]["status"] = "absent"
        finalization["output_field_cache"]["status"] = "refreshed"
        finalization["artifact_binding"] = {
            "version": 1,
            "finalized_docx": identity(finalized),
            "word_verification_pdf": {**identity(pdf), "page_count": 7},
        }
        finalization["workflow_state"] = {
            "stage": "finalized",
            "source_sha256": file_sha256(source),
            "input_sha256": file_sha256(formatted),
            "profile_sha256": file_sha256(profile),
            "structure_map_sha256": file_sha256(structure),
            "output_sha256": file_sha256(finalized),
        }
        finalization["output"] = str(finalized.resolve())
        finalization["target_pdf"] = str(pdf.resolve())
        raw_backend = read_bound_backend_audit(finalization)
        _, finalization["backend_audit"], _ = self._persist_backend_audit(
            final_dir / "finalization.json", raw_backend
        )
        (final_dir / "finalization.json").write_text(
            json.dumps(finalization), encoding="utf-8"
        )
        state = {
            "status": "candidate_ready",
            "source": {"path": str(source), "sha256": file_sha256(source)},
            "profile": {"path": str(profile), "sha256": file_sha256(profile)},
            "structure_map": {
                "path": str(structure),
                "sha256": file_sha256(structure),
            },
            "artifacts": {"formatted": str(formatted.relative_to(work))},
            "blockers": [],
            "qa_groups": [],
            "frozen_scopes": [],
            "stages": {},
            "metrics": {},
        }
        args = SimpleNamespace(
            work_dir=work,
            resume=False,
            field_updater="external",
            field_updater_command=sys.executable,
            target_software="Microsoft Word",
            renderer=None,
            approve_deferred=False,
        )
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.run_script",
            return_value=SimpleNamespace(returncode=0),
        ), patch("run_monograph.save_state"):
            self.assertEqual(0, orchestrate_finalize(args))
        self.assertEqual("candidate_ready", state["status"])
        self.assertTrue(state["blockers"])
        self.assertTrue(state["field_writeback"]["completion_evidence_errors"])

    def _cached_final_ready_run(self, name: str) -> tuple[Path, dict, SimpleNamespace, SimpleNamespace]:
        work = self.root / name
        work.mkdir()
        updater = work / "external-updater"
        updater.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        updater.chmod(0o755)
        source = work / "source.docx"
        profile = work / "profile.json"
        structure = work / "structure.json"
        formatted = work / "formatted.docx"
        source.write_bytes(b"source")
        profile.write_text("{}", encoding="utf-8")
        structure.write_text("{}", encoding="utf-8")
        formatted.write_bytes(b"formatted")
        final_dir = work / "final"
        render_dir = work / "rendered"
        final_dir.mkdir()
        render_dir.mkdir()
        finalized = final_dir / "source-finalized.docx"
        pdf = final_dir / "source-target.pdf"
        finalization_status = final_dir / "finalization.json"
        audit = final_dir / "audit.json"
        render_manifest = render_dir / "render-manifest.json"
        visual_manifest = work / "visual.json"
        finalized.write_bytes(b"finalized")
        pdf.write_bytes(b"%PDF persistent verification")
        audit.write_text('{"status":"pass"}', encoding="utf-8")
        render_manifest.write_text(
            json.dumps(
                {
                    "page_count": 7,
                    "target_pdf_source": str(pdf.resolve()),
                    "target_software": "microsoft_word",
                    "renderer": None,
                    "renderer_source": "target_pdf",
                    "target_layout_unverified": False,
                }
            ),
            encoding="utf-8",
        )
        visual_manifest.write_text(
            json.dumps(
                {
                    "all_pages_inspected": True,
                    "target_layout_verified": True,
                    "page_count": 7,
                    "issues": [],
                }
            ),
            encoding="utf-8",
        )
        identity = lambda path: {
            "path": str(path.resolve()),
            "sha256": file_sha256(path),
            "size_bytes": path.stat().st_size,
        }
        binding = {
            "version": 1,
            "finalized_docx": identity(finalized),
            "word_verification_pdf": {**identity(pdf), "page_count": 7},
        }
        raw_backend = {
            "backend": "external",
            "target_id": "microsoft_word",
            "field_cache_verified": True,
            "page_count": 7,
            "selective_writeback": {"status": "selective_verified"},
            "read_only_verification": {
                "target_id": "microsoft_word",
                "operation": "verify_only",
                "read_only_verified": True,
                "repaginated": True,
                "saved": False,
                "pdf_exported": True,
                "page_count": 7,
            },
        }
        canonical_backend, audit_binding, audit_path = self._persist_backend_audit(
            finalization_status, raw_backend
        )
        finalization = {
            "finalization_evidence_version": FINALIZATION_EVIDENCE_VERSION,
            "delivery_field_status": "selective_verified",
            "input_field_cache": self._field_cache("stale"),
            "output_field_cache": self._field_cache("refreshed"),
            "field_writeback_status": "selective_verified",
            "field_backend": canonical_backend,
            "backend_audit": audit_binding,
            "field_completion": {
                "word_verification_required": True,
                "word_verification_completed": True,
                "completion_scope": "target_word_verified",
                "field_gate_completed": True,
                "final_ready_eligible": True,
            },
            "artifact_binding": binding,
            "status": "pass",
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
            "workflow_state": {
                "stage": "finalized",
                "source_sha256": file_sha256(source),
                "input_sha256": file_sha256(formatted),
                "profile_sha256": file_sha256(profile),
                "structure_map_sha256": file_sha256(structure),
                "output_sha256": identity(finalized)["sha256"],
            },
            "output": str(finalized.resolve()),
            "target_pdf": str(pdf.resolve()),
            "target_layout_status": "target_pdf_ready_for_visual_qa",
            "target_software": "microsoft_word",
        }
        finalization["field_completion"]["evidence_validation"] = {
            "status": "pass",
            "errors": [],
        }
        finalization_status.write_text(json.dumps(finalization), encoding="utf-8")
        completion = completion_evidence(finalization)
        finalize_args = SimpleNamespace(
            work_dir=work,
            resume=True,
            field_updater="external",
            field_updater_command=str(updater),
            target_software="Microsoft Word",
            renderer=None,
            approve_deferred=False,
        )
        verify_args = SimpleNamespace(
            work_dir=work,
            resume=True,
            renderer=None,
            target_software="Microsoft Word",
            visual_qa_manifest=visual_manifest,
        )
        finalize_behavior = finalize_request_identity(finalize_args)
        verify_behavior = verify_request_identity(verify_args)
        finalize_key = json_sha256(
            {
                "formatted": file_sha256(formatted),
                "map": file_sha256(structure),
                "behavior": finalize_behavior,
            }
        )
        verify_key = json_sha256(
            {
                "finalized": file_sha256(finalized),
                "map": file_sha256(structure),
                "target_pdf": file_sha256(pdf),
                "visual_manifest": {
                    "path": str(visual_manifest.resolve()),
                    "sha256": file_sha256(visual_manifest),
                },
                "behavior": verify_behavior,
            }
        )
        state = {
            "status": "final_ready",
            "run_id": name,
            "source": {"path": str(source), "sha256": file_sha256(source)},
            "profile": {"path": str(profile), "sha256": file_sha256(profile)},
            "structure_map": {
                "path": str(structure),
                "sha256": file_sha256(structure),
            },
            "artifacts": {
                "formatted": "formatted.docx",
                "finalized": "final/source-finalized.docx",
                "finalization_status": "final/finalization.json",
                "target_pdf": "final/source-target.pdf",
                "backend_audit": str(audit_path.relative_to(work)),
                "final_audit": "final/audit.json",
                "render_manifest": "rendered/render-manifest.json",
                "visual_qa_manifest": str(visual_manifest.resolve()),
            },
            "field_writeback": {
                "completion_evidence": completion,
                "artifact_binding": binding,
            },
            "finalization_request": finalize_behavior,
            "finalization_gate": canonical_finalization_gate_summary(
                work, finalization
            ),
            "verification_outputs": {
                "version": VERIFICATION_OUTPUT_VERSION,
                "request": verify_behavior,
                "final_audit": local_artifact_identity(audit),
                "render_manifest": local_artifact_identity(render_manifest),
                "visual_manifest": local_artifact_identity(visual_manifest),
            },
            "visual_qa": {
                "all_pages_inspected": True,
                "target_layout_verified": True,
                "page_count": 7,
            },
            "blockers": [],
            "qa_groups": [],
            "frozen_scopes": [],
            "metrics": {"cache_hits": 0},
            "stages": {
                "finalize": {
                    "status": "complete",
                    "input_key_sha256": finalize_key,
                    "cache_hit": False,
                },
                "verify": {
                    "status": "complete",
                    "input_key_sha256": verify_key,
                    "cache_hit": False,
                },
            },
        }
        return work, state, finalize_args, verify_args

    def _cached_no_fields_final_ready_run(
        self, name: str
    ) -> tuple[Path, dict, SimpleNamespace, SimpleNamespace]:
        work, state, finalize_args, verify_args = self._cached_final_ready_run(name)
        finalization_path = work / "final/finalization.json"
        finalization = json.loads(finalization_path.read_text(encoding="utf-8"))
        canonical_backend, audit_binding, audit_path = self._persist_backend_audit(
            finalization_path, {"backend": "not_needed"}
        )
        finalization.update(
            {
                "delivery_field_status": "absent",
                "input_field_cache": self._field_cache("absent"),
                "output_field_cache": self._field_cache("absent"),
                "field_writeback_status": "not_needed",
                "field_backend": canonical_backend,
                "backend_audit": audit_binding,
                "field_completion": {
                    "word_verification_required": False,
                    "word_verification_completed": False,
                    "completion_scope": "no_fields",
                    "field_gate_completed": True,
                    "final_ready_eligible": True,
                    "evidence_validation": {"status": "pass", "errors": []},
                },
                "artifact_binding": {
                    "version": 1,
                    "finalized_docx": finalization["artifact_binding"][
                        "finalized_docx"
                    ],
                    "word_verification_pdf": None,
                },
                "target_pdf": None,
                "target_layout_status": "not_verified",
                "target_software": "libreoffice",
            }
        )
        finalization_path.write_text(json.dumps(finalization), encoding="utf-8")
        (work / "final/source-target.pdf").unlink()
        state["artifacts"].pop("target_pdf")
        state["artifacts"]["backend_audit"] = str(audit_path.relative_to(work))
        state["field_writeback"].update(
            {
                "completion_evidence": completion_evidence(finalization),
                "artifact_binding": finalization["artifact_binding"],
            }
        )
        state["finalization_gate"] = canonical_finalization_gate_summary(
            work, finalization
        )
        finalize_args.target_software = "LibreOffice Writer"
        finalize_behavior = finalize_request_identity(finalize_args)
        state["finalization_request"] = finalize_behavior
        state["stages"]["finalize"]["input_key_sha256"] = json_sha256(
            {
                "formatted": file_sha256(work / "formatted.docx"),
                "map": file_sha256(work / "structure.json"),
                "behavior": finalize_behavior,
            }
        )
        verify_args.target_software = "LibreOffice Writer"
        verify_behavior = verify_request_identity(verify_args)
        renderer = verify_behavior["renderer"]
        render_manifest = work / "rendered/render-manifest.json"
        render_manifest.write_text(
            json.dumps(
                {
                    "page_count": 7,
                    "target_pdf_source": None,
                    "target_software": "libreoffice",
                    "renderer": renderer["resolved_path"],
                    "renderer_source": renderer["source"],
                    "target_layout_unverified": False,
                }
            ),
            encoding="utf-8",
        )
        state["verification_outputs"].update(
            {
                "request": verify_behavior,
                "render_manifest": local_artifact_identity(render_manifest),
            }
        )
        state["stages"]["verify"]["input_key_sha256"] = json_sha256(
            {
                "finalized": file_sha256(work / "final/source-finalized.docx"),
                "map": file_sha256(work / "structure.json"),
                "target_pdf": None,
                "visual_manifest": {
                    "path": str((work / "visual.json").resolve()),
                    "sha256": file_sha256(work / "visual.json"),
                },
                "behavior": verify_behavior,
            }
        )
        return work, state, finalize_args, verify_args

    def _bind_external_finalize_command(
        self,
        work: Path,
        state: dict,
        args: SimpleNamespace,
        argv: list[str],
    ) -> dict:
        args.field_updater_command = json.dumps(argv)
        request = finalize_request_identity(
            args, target_id=MICROSOFT_WORD, renderer_used=False
        )
        state["finalization_request"] = request
        state["stages"]["finalize"]["input_key_sha256"] = json_sha256(
            {
                "formatted": file_sha256(work / "formatted.docx"),
                "map": file_sha256(work / "structure.json"),
                "behavior": request,
            }
        )
        return request

    def _verification_script_result(
        self,
        work: Path,
        *,
        target_software: str | None,
        renderer_identity: dict,
        use_target_pdf: bool,
        manifest_target: str | None = None,
        page_count: int = 7,
    ):
        def run(script, *_arguments):
            if script == "audit_docx.py":
                (work / "final/audit.json").write_text(
                    '{"status":"pass"}', encoding="utf-8"
                )
            elif script == "render_docx.py":
                target_pdf = (
                    str((work / "final/source-target.pdf").resolve())
                    if use_target_pdf
                    else None
                )
                (work / "rendered/render-manifest.json").write_text(
                    json.dumps(
                        {
                            "page_count": page_count,
                            "target_pdf_source": target_pdf,
                            "target_software": (
                                target_software
                                if manifest_target is None
                                else manifest_target
                            ),
                            "renderer": (
                                None
                                if use_target_pdf
                                else renderer_identity["resolved_path"]
                            ),
                            "renderer_source": (
                                "target_pdf"
                                if use_target_pdf
                                else renderer_identity["source"]
                            ),
                            "target_layout_unverified": bool(
                                target_software
                                and "word" in target_software.casefold()
                                and not use_target_pdf
                            ),
                        }
                    ),
                    encoding="utf-8",
                )
            return SimpleNamespace(returncode=0)

        return run

    def test_finalization_evidence_shape_accepts_all_production_business_shapes(self) -> None:
        word_work, _, _, _ = self._cached_final_ready_run("shape-word")
        word = json.loads(
            (word_work / "final/finalization.json").read_text(encoding="utf-8")
        )
        no_fields_work, _, _, _ = self._cached_no_fields_final_ready_run(
            "shape-no-fields"
        )
        no_fields = json.loads(
            (no_fields_work / "final/finalization.json").read_text(
                encoding="utf-8"
            )
        )
        deferred = copy.deepcopy(word)
        deferred.update(
            {
                "delivery_field_status": "deferred",
                "output_field_cache": self._field_cache("stale"),
                "field_backend": canonical_backend_projection(
                    {
                        "backend": "deferred_on_open",
                        "fallback_from": "external_error",
                        "attempted_backend": {"backend": "external"},
                    }
                ),
                "field_writeback_status": "deferred",
                "target_pdf": None,
                "target_layout_status": "not_verified",
            }
        )
        deferred["artifact_binding"]["word_verification_pdf"] = None
        deferred["field_completion"] = {
            "field_gate_completed": False,
            "final_ready_eligible": False,
            "word_verification_required": True,
            "word_verification_completed": False,
            "completion_scope": "incomplete",
            "evidence_validation": {
                "status": "incomplete",
                "errors": ["delivery_status is neither absent nor selective_verified"],
            },
        }
        libreoffice = copy.deepcopy(deferred)
        libreoffice.update(
            {
                "delivery_field_status": "libreoffice_refreshed",
                "field_backend": canonical_backend_projection(
                    {
                        "backend": "libreoffice_uno",
                        "delivery_field_contract_identical": True,
                        "selective_writeback": {
                            "status": "libreoffice_selective"
                        },
                    }
                ),
                "field_writeback_status": "libreoffice_selective",
                "target_software": LIBREOFFICE,
            }
        )
        libreoffice["field_completion"]["completion_scope"] = (
            "libreoffice_non_final"
        )
        refreshed = copy.deepcopy(no_fields)
        refreshed.update(
            {
                "delivery_field_status": "refreshed",
                "input_field_cache": self._field_cache("refreshed"),
                "output_field_cache": self._field_cache("refreshed"),
            }
        )
        refreshed["field_completion"] = {
            "field_gate_completed": False,
            "final_ready_eligible": False,
            "word_verification_required": True,
            "word_verification_completed": False,
            "completion_scope": "incomplete",
            "evidence_validation": {
                "status": "incomplete",
                "errors": ["delivery_status is neither absent nor selective_verified"],
            },
        }
        for name, evidence in (
            ("word", word),
            ("no_fields", no_fields),
            ("deferred", deferred),
            ("libreoffice_nonfinal", libreoffice),
            ("refreshed_candidate", refreshed),
        ):
            with self.subTest(name=name):
                self.assertEqual([], finalization_evidence_shape_errors(evidence))
        for evidence in (deferred, libreoffice, refreshed):
            self.assertTrue(final_ready_evidence_errors(completion_evidence(evidence)))

    def test_backend_projection_is_closed_and_audit_retains_bounded_diagnostics(self) -> None:
        raw = {
            "backend": "external",
            "target_id": MICROSOFT_WORD,
            "field_cache_verified": True,
            "page_count": 7,
            "producer_extension": {"trace": "retained only in audit"},
            "selective_writeback": {
                "status": "selective_verified",
                "matched_fields": 3,
                "updated_fields": 3,
                "diagnostic": "not canonical",
            },
            "read_only_verification": {
                "operation": "verify_only",
                "target_id": MICROSOFT_WORD,
                "read_only_verified": True,
                "repaginated": True,
                "saved": False,
                "pdf_exported": True,
                "page_count": 7,
                "application_build": "synthetic",
            },
        }
        canonical = canonical_backend_projection(raw)
        self.assertEqual([], canonical_backend_shape_errors(canonical))
        self.assertNotIn("producer_extension", canonical)
        self.assertNotIn("diagnostic", canonical["selective_writeback"])
        self.assertNotIn(
            "application_build", canonical["read_only_verification"]
        )

        status_path = self.root / "closed-finalization.json"
        payload = backend_audit_bytes(raw)
        audit_path = backend_audit_path(status_path)
        audit_path.write_bytes(payload)
        finalization = {
            "backend_audit": backend_audit_binding(audit_path, payload)
        }
        loaded = read_bound_backend_audit(finalization)
        self.assertEqual(raw, loaded)

    def test_backend_projection_rejects_noncanonical_business_values(self) -> None:
        base = {
            "backend": "external",
            "target_id": MICROSOFT_WORD,
            "page_count": 7,
            "selective_writeback": {
                "status": "selective_verified",
                "matched_fields": 1,
                "updated_fields": 1,
            },
            "read_only_verification": {
                "operation": "verify_only",
                "target_id": MICROSOFT_WORD,
                "read_only_verified": True,
                "repaginated": True,
                "saved": False,
                "pdf_exported": True,
                "page_count": 7,
            },
        }
        mutations = (
            ("unknown_backend", lambda value: value.update(backend="unknown")),
            ("unknown_target", lambda value: value.update(target_id="wordperfect")),
            ("boolean_page_count", lambda value: value.update(page_count=True)),
            (
                "negative_match_count",
                lambda value: value["selective_writeback"].update(
                    matched_fields=-1
                ),
            ),
            (
                "unknown_operation",
                lambda value: value["read_only_verification"].update(
                    operation="save_and_verify"
                ),
            ),
            ("nan_page_count", lambda value: value.update(page_count=float("nan"))),
            (
                "infinite_page_count",
                lambda value: value["read_only_verification"].update(
                    page_count=float("inf")
                ),
            ),
        )
        for name, mutate in mutations:
            value = copy.deepcopy(base)
            mutate(value)
            with self.subTest(name=name), self.assertRaises(BackendEvidenceError):
                canonical_backend_projection(value)

    def test_backend_audit_rejects_nonstandard_and_bounded_json_matrix(self) -> None:
        deep = {}
        cursor = deep
        for _ in range(34):
            cursor["child"] = {}
            cursor = cursor["child"]
        cases = (
            ("nan", {"backend": "external", "value": float("nan")}),
            ("infinity", {"backend": "external", "value": float("inf")}),
            ("non_string_key", {"backend": "external", 7: "bad"}),
            ("object", {"backend": "external", "value": object()}),
            ("depth", {"backend": "external", "value": deep}),
            (
                "string_length",
                {
                    "backend": "external",
                    "value": "x"
                    * (backend_evidence.BACKEND_AUDIT_MAX_STRING_LENGTH + 1),
                },
            ),
        )
        for name, raw in cases:
            with self.subTest(name=name), self.assertRaises(BackendEvidenceError):
                backend_audit_bytes(raw)
        with patch.object(backend_evidence, "BACKEND_AUDIT_MAX_NODES", 5):
            with self.assertRaises(BackendEvidenceError):
                backend_audit_bytes({"backend": "external", "items": [1, 2, 3]})
        with patch.object(backend_evidence, "BACKEND_AUDIT_MAX_BYTES", 32):
            with self.assertRaises(BackendEvidenceError):
                backend_audit_bytes({"backend": "external", "message": "bounded"})

    def test_backend_audit_tamper_is_rejected_by_finalize_verify_and_status(self) -> None:
        def mutate_bytes(work: Path, _evidence: dict) -> None:
            path = work / "final/finalization-backend-audit.json"
            path.write_bytes(path.read_bytes() + b" ")

        def mutate_path(work: Path, evidence: dict) -> None:
            evidence["backend_audit"]["artifact"]["path"] = str(
                (work / "missing-audit.json").resolve()
            )

        def mutate_hash(_work: Path, evidence: dict) -> None:
            evidence["backend_audit"]["artifact"]["sha256"] = "0" * 64

        def mutate_size(_work: Path, evidence: dict) -> None:
            evidence["backend_audit"]["artifact"]["size_bytes"] += 1

        for entrypoint in ("finalize", "verify", "status"):
            for name, mutate in (
                ("bytes", mutate_bytes),
                ("path", mutate_path),
                ("hash", mutate_hash),
                ("size", mutate_size),
            ):
                work, state, finalize_args, verify_args = self._cached_final_ready_run(
                    f"backend-audit-{entrypoint}-{name}"
                )
                status_path = work / "final/finalization.json"
                evidence = json.loads(status_path.read_text(encoding="utf-8"))
                mutate(work, evidence)
                status_path.write_text(json.dumps(evidence), encoding="utf-8")
                with self.subTest(entrypoint=entrypoint, mutation=name), patch(
                    "run_monograph.load_state", return_value=state
                ), patch("run_monograph.save_state"), patch(
                    "run_monograph.run_script",
                    return_value=SimpleNamespace(returncode=0),
                ) as run_script_mock, contextlib.redirect_stdout(
                    io.StringIO()
                ), contextlib.redirect_stderr(
                    io.StringIO()
                ):
                    if entrypoint == "finalize":
                        self.assertEqual(2, orchestrate_finalize(finalize_args))
                        self.assertEqual(1, run_script_mock.call_count)
                    elif entrypoint == "verify":
                        self.assertEqual(2, verify(verify_args))
                        run_script_mock.assert_not_called()
                    else:
                        self.assertEqual(
                            2,
                            orchestrate_status(
                                SimpleNamespace(work_dir=work, json=True)
                            ),
                        )
                        run_script_mock.assert_not_called()
                self.assertEqual(
                    "final_ready" if entrypoint == "finalize" else "candidate_ready",
                    state["status"],
                )

    def test_fresh_backend_audit_is_atomic_before_verification_invalidation_matrix(self) -> None:
        def bind_payload(evidence: dict, path: Path, payload: bytes) -> None:
            evidence["backend_audit"] = backend_audit_binding(path, payload)

        for name in (
            "missing",
            "truncated",
            "hash",
            "size",
            "nan",
            "oversize",
            "overdepth",
            "overnodes",
            "overstring",
            "root",
            "payload_nonobject",
            "arbitrary_path",
            "nul_path",
            "symlink",
        ):
            work, state, args, _ = self._cached_final_ready_run(
                f"fresh-backend-audit-{name}"
            )
            status_path = work / "final/finalization.json"
            expected = backend_audit_path(status_path)
            evidence = json.loads(status_path.read_text(encoding="utf-8"))
            if name == "missing":
                expected.unlink()
            elif name == "truncated":
                expected.write_bytes(b"{")
            elif name == "hash":
                evidence["backend_audit"]["artifact"]["sha256"] = "0" * 64
            elif name == "size":
                evidence["backend_audit"]["artifact"]["size_bytes"] += 1
            elif name == "nan":
                payload = (
                    b'{"backend_audit_version":1,"backend":'
                    b'{"backend":"external","diagnostic":NaN}}\n'
                )
                expected.write_bytes(payload)
                bind_payload(evidence, expected, payload)
            elif name == "oversize":
                payload = (
                    b'{"backend_audit_version":1,"backend":{"backend":"external",'
                    b'"diagnostic":"'
                    + b"x" * (backend_evidence.BACKEND_AUDIT_MAX_BYTES + 1)
                    + b'"}}\n'
                )
                expected.write_bytes(payload)
                bind_payload(evidence, expected, payload)
            elif name == "overdepth":
                nested: dict = {}
                cursor = nested
                for _ in range(34):
                    cursor["child"] = {}
                    cursor = cursor["child"]
                payload = json.dumps(
                    {"backend_audit_version": 1, "backend": nested},
                    separators=(",", ":"),
                ).encode("utf-8")
                expected.write_bytes(payload)
                bind_payload(evidence, expected, payload)
            elif name == "overnodes":
                payload = json.dumps(
                    {
                        "backend_audit_version": 1,
                        "backend": {
                            "items": [
                                0
                                for _ in range(
                                    backend_evidence.BACKEND_AUDIT_MAX_NODES
                                )
                            ]
                        },
                    },
                    separators=(",", ":"),
                ).encode("utf-8")
                expected.write_bytes(payload)
                bind_payload(evidence, expected, payload)
            elif name == "overstring":
                payload = json.dumps(
                    {
                        "backend_audit_version": 1,
                        "backend": {
                            "diagnostic": "x"
                            * (
                                backend_evidence.BACKEND_AUDIT_MAX_STRING_LENGTH
                                + 1
                            )
                        },
                    },
                    separators=(",", ":"),
                ).encode("utf-8")
                expected.write_bytes(payload)
                bind_payload(evidence, expected, payload)
            elif name == "root":
                payload = b'{"backend_audit_version":1,"backend":{},"extra":true}\n'
                expected.write_bytes(payload)
                bind_payload(evidence, expected, payload)
            elif name == "payload_nonobject":
                payload = b'{"backend_audit_version":1,"backend":[]}\n'
                expected.write_bytes(payload)
                bind_payload(evidence, expected, payload)
            elif name == "arbitrary_path":
                arbitrary = work / "arbitrary-audit.json"
                payload = expected.read_bytes()
                arbitrary.write_bytes(payload)
                bind_payload(evidence, arbitrary, payload)
            elif name == "nul_path":
                evidence["backend_audit"]["artifact"]["path"] = "bad\0audit.json"
            elif name == "symlink":
                real = work / "real-audit.json"
                payload = expected.read_bytes()
                real.write_bytes(payload)
                expected.unlink()
                expected.symlink_to(real)
                bind_payload(evidence, expected, payload)
                evidence["backend_audit"]["artifact"]["path"] = str(expected)
            status_path.write_text(json.dumps(evidence), encoding="utf-8")
            before = copy.deepcopy(state)
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state") as save, patch(
                "run_monograph.run_script",
                return_value=SimpleNamespace(returncode=0),
            ) as run_script_mock, contextlib.redirect_stdout(
                io.StringIO()
            ), contextlib.redirect_stderr(
                io.StringIO()
            ):
                self.assertEqual(2, orchestrate_finalize(args))
            self.assertEqual(1, run_script_mock.call_count)
            save.assert_not_called()
            self.assertEqual(before, state)
            self.assertIn("verify", state["stages"])

    def test_all_finalization_version_fields_require_exact_integer_type(self) -> None:
        work, _, _, _ = self._cached_final_ready_run("strict-version-shape")
        base = json.loads(
            (work / "final/finalization.json").read_text(encoding="utf-8")
        )
        deferred_backend = canonical_backend_projection(
            {
                "backend": "deferred_on_open",
                "fallback_from": "libreoffice_contract_or_integrity",
                "attempted_backend": {
                    "backend": "libreoffice_uno",
                    "failure": {
                        "status": "rejected",
                        "stage": "selective_writeback",
                        "failed_checks": ["selective_writeback"],
                    },
                },
            }
        )
        locations = {
            "finalization": lambda value: value.update(
                finalization_evidence_version=None
            ),
            "backend": lambda value: value["field_backend"].update(version=None),
            "selective": lambda value: value["field_backend"][
                "selective_writeback"
            ].update(version=None),
            "verification": lambda value: value["field_backend"][
                "read_only_verification"
            ].update(version=None),
            "artifact_binding": lambda value: value["artifact_binding"].update(
                version=None
            ),
            "audit_binding": lambda value: value["backend_audit"].update(
                version=None
            ),
        }
        for label, prepare in locations.items():
            for bad_version in (True, 1.0, "1", 2):
                evidence = copy.deepcopy(base)
                prepare(evidence)
                if label == "finalization":
                    evidence["finalization_evidence_version"] = bad_version
                elif label == "backend":
                    evidence["field_backend"]["version"] = bad_version
                elif label == "selective":
                    evidence["field_backend"]["selective_writeback"][
                        "version"
                    ] = bad_version
                elif label == "verification":
                    evidence["field_backend"]["read_only_verification"][
                        "version"
                    ] = bad_version
                elif label == "artifact_binding":
                    evidence["artifact_binding"]["version"] = bad_version
                else:
                    evidence["backend_audit"]["version"] = bad_version
                with self.subTest(label=label, version=bad_version):
                    self.assertTrue(finalization_evidence_shape_errors(evidence))
        for label in ("attempt", "failure"):
            for bad_version in (True, 1.0, "1", 2):
                evidence = copy.deepcopy(base)
                evidence["field_backend"] = copy.deepcopy(deferred_backend)
                if label == "attempt":
                    evidence["field_backend"]["attempt"]["version"] = bad_version
                else:
                    evidence["field_backend"]["attempt"]["failure"][
                        "version"
                    ] = bad_version
                with self.subTest(label=label, version=bad_version):
                    self.assertTrue(finalization_evidence_shape_errors(evidence))

    def test_invalid_version_and_control_paths_never_crash_three_entrypoints(self) -> None:
        mutations = (
            (
                "version_true",
                lambda evidence: evidence.update(finalization_evidence_version=True),
            ),
            (
                "version_float",
                lambda evidence: evidence.update(finalization_evidence_version=1.0),
            ),
            (
                "version_string",
                lambda evidence: evidence.update(finalization_evidence_version="1"),
            ),
            (
                "version_unknown",
                lambda evidence: evidence.update(finalization_evidence_version=2),
            ),
            ("nul", lambda evidence: evidence.update(output="bad\0output.docx")),
            (
                "control",
                lambda evidence: evidence["artifact_binding"][
                    "finalized_docx"
                ].update(path="bad\noutput.docx"),
            ),
        )
        for entrypoint in ("finalize", "verify", "status"):
            for name, mutate in mutations:
                work, state, finalize_args, verify_args = self._cached_final_ready_run(
                    f"invalid-path-version-{entrypoint}-{name}"
                )
                status_path = work / "final/finalization.json"
                evidence = json.loads(status_path.read_text(encoding="utf-8"))
                mutate(evidence)
                status_path.write_text(json.dumps(evidence), encoding="utf-8")
                before = copy.deepcopy(state)
                with self.subTest(entrypoint=entrypoint, name=name), patch(
                    "run_monograph.load_state", return_value=state
                ), patch("run_monograph.save_state") as save, patch(
                    "run_monograph.run_script",
                    return_value=SimpleNamespace(returncode=0),
                ) as run_script_mock, contextlib.redirect_stdout(
                    io.StringIO()
                ), contextlib.redirect_stderr(
                    io.StringIO()
                ):
                    if entrypoint == "finalize":
                        self.assertEqual(2, orchestrate_finalize(finalize_args))
                        run_script_mock.assert_called_once()
                        save.assert_not_called()
                        self.assertEqual(before, state)
                    elif entrypoint == "verify":
                        self.assertEqual(2, verify(verify_args))
                        run_script_mock.assert_not_called()
                        save.assert_called_once()
                        self.assertEqual("candidate_ready", state["status"])
                    else:
                        self.assertEqual(
                            2,
                            orchestrate_status(
                                SimpleNamespace(work_dir=work, json=True)
                            ),
                        )
                        run_script_mock.assert_not_called()
                        save.assert_called_once()
                        self.assertEqual("candidate_ready", state["status"])

    def test_invalid_fresh_finalization_shape_preserves_old_verify_state_matrix(self) -> None:
        top_level_required = (
            "finalization_evidence_version",
            "status",
            "delivery_field_status",
            "input_field_cache",
            "output_field_cache",
            "field_backend",
            "backend_audit",
            "artifact_binding",
            "field_writeback_status",
            "field_completion",
            "content_integrity",
            "protected_object_integrity",
            "effective_font_integrity",
            "workflow_state",
            "target_pdf",
            "target_layout_status",
            "target_software",
            "output",
        )
        mutations = [(f"missing_{name}", lambda value, key=name: value.pop(key)) for name in top_level_required]
        mutations.extend(
            (
                ("empty_object", lambda value: value.clear()),
                ("unknown_top_field", lambda value: value.update(unknown=True)),
                (
                    "unknown_evidence_version",
                    lambda value: value.update(finalization_evidence_version=2),
                ),
                (
                    "boolean_evidence_version",
                    lambda value: value.update(finalization_evidence_version=True),
                ),
                ("input_cache_type", lambda value: value.update(input_field_cache=[])),
                (
                    "unknown_cache_field",
                    lambda value: value["input_field_cache"].update(unknown=0),
                ),
                ("output_cache_type", lambda value: value.update(output_field_cache=[])),
                ("backend_type", lambda value: value.update(field_backend=[])),
                ("binding_type", lambda value: value.update(artifact_binding=[])),
                ("completion_type", lambda value: value.update(field_completion=[])),
                ("workflow_type", lambda value: value.update(workflow_state=[])),
                (
                    "unknown_input_cache",
                    lambda value: value["input_field_cache"].update(status="unknown"),
                ),
                (
                    "boolean_cache_count",
                    lambda value: value["input_field_cache"].update(dirty_fields=True),
                ),
                (
                    "unknown_delivery",
                    lambda value: value.update(delivery_field_status="unknown"),
                ),
                (
                    "unknown_backend",
                    lambda value: value["field_backend"].update(backend="unknown"),
                ),
                (
                    "unknown_backend_field",
                    lambda value: value["field_backend"].update(unknown="x"),
                ),
                (
                    "unknown_selective_field",
                    lambda value: value["field_backend"][
                        "selective_writeback"
                    ].update(unknown="x"),
                ),
                (
                    "unknown_verification_field",
                    lambda value: value["field_backend"][
                        "read_only_verification"
                    ].update(unknown="x"),
                ),
                (
                    "bad_selective_count",
                    lambda value: value["field_backend"][
                        "selective_writeback"
                    ].update(matched_fields=-1),
                ),
                (
                    "bad_verification_operation",
                    lambda value: value["field_backend"][
                        "read_only_verification"
                    ].update(operation="save_and_verify"),
                ),
                (
                    "nan_backend_page_count",
                    lambda value: value["field_backend"].update(
                        page_count=float("nan")
                    ),
                ),
                (
                    "infinite_verification_page_count",
                    lambda value: value["field_backend"][
                        "read_only_verification"
                    ].update(page_count=float("inf")),
                ),
                (
                    "unknown_backend_audit_field",
                    lambda value: value["backend_audit"].update(unknown="x"),
                ),
                (
                    "unknown_writeback",
                    lambda value: value.update(field_writeback_status="unknown"),
                ),
                (
                    "unknown_target",
                    lambda value: value.update(target_software="wordperfect"),
                ),
                (
                    "unknown_artifact_version",
                    lambda value: value["artifact_binding"].update(version=2),
                ),
                (
                    "unknown_artifact_field",
                    lambda value: value["artifact_binding"]["finalized_docx"].update(
                        unknown="x"
                    ),
                ),
                (
                    "empty_artifact_path",
                    lambda value: value["artifact_binding"]["finalized_docx"].update(
                        path=" "
                    ),
                ),
                (
                    "bad_artifact_hash",
                    lambda value: value["artifact_binding"]["finalized_docx"].update(
                        sha256="xyz"
                    ),
                ),
                (
                    "boolean_artifact_size",
                    lambda value: value["artifact_binding"]["finalized_docx"].update(
                        size_bytes=True
                    ),
                ),
                (
                    "boolean_pdf_page_count",
                    lambda value: value["artifact_binding"][
                        "word_verification_pdf"
                    ].update(page_count=True),
                ),
                (
                    "bad_workflow_hash",
                    lambda value: value["workflow_state"].update(
                        output_sha256="A" * 64
                    ),
                ),
                (
                    "unknown_workflow_stage",
                    lambda value: value["workflow_state"].update(stage="verified"),
                ),
                ("empty_output_path", lambda value: value.update(output="")),
                ("bad_target_pdf_type", lambda value: value.update(target_pdf=7)),
                (
                    "unknown_layout_status",
                    lambda value: value.update(target_layout_status="ready"),
                ),
                (
                    "unknown_completion_scope",
                    lambda value: value["field_completion"].update(
                        completion_scope="unknown"
                    ),
                ),
                (
                    "boolean_completion_flag",
                    lambda value: value["field_completion"].update(
                        final_ready_eligible=1
                    ),
                ),
                (
                    "validation_type",
                    lambda value: value["field_completion"].update(
                        evidence_validation=[]
                    ),
                ),
                (
                    "validation_missing_errors",
                    lambda value: value["field_completion"].update(
                        evidence_validation={"status": "pass"}
                    ),
                ),
                (
                    "validation_bad_error_type",
                    lambda value: value["field_completion"].update(
                        evidence_validation={"status": "incomplete", "errors": [7]}
                    ),
                ),
                (
                    "validation_status_conflict",
                    lambda value: value["field_completion"].update(
                        evidence_validation={"status": "pass", "errors": ["bad"]}
                    ),
                ),
            )
        )
        for name, mutate in mutations:
            work, state, args, _ = self._cached_final_ready_run(
                f"invalid-fresh-shape-{name}"
            )
            status_path = work / "final/finalization.json"
            evidence = json.loads(status_path.read_text(encoding="utf-8"))
            mutate(evidence)
            before = copy.deepcopy(state)

            def run(_script, *_arguments):
                status_path.write_text(json.dumps(evidence), encoding="utf-8")
                return SimpleNamespace(returncode=0)

            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state") as save, patch(
                "run_monograph.run_script", side_effect=run
            ) as run_script_mock, contextlib.redirect_stdout(
                io.StringIO()
            ), contextlib.redirect_stderr(
                io.StringIO()
            ):
                self.assertEqual(2, orchestrate_finalize(args))
            self.assertEqual(1, run_script_mock.call_count)
            save.assert_not_called()
            self.assertEqual(before, state)

    def test_non_object_missing_invalid_json_and_process_failure_are_atomic(self) -> None:
        cases = ("non_object", "missing", "invalid_json", "process_failure")
        for name in cases:
            work, state, args, _ = self._cached_final_ready_run(
                f"fresh-evidence-atomic-{name}"
            )
            status_path = work / "final/finalization.json"
            before = copy.deepcopy(state)

            def run(_script, *_arguments):
                if name == "non_object":
                    status_path.write_text("[]", encoding="utf-8")
                elif name == "missing":
                    status_path.unlink()
                elif name == "invalid_json":
                    status_path.write_text("{", encoding="utf-8")
                return SimpleNamespace(returncode=1 if name == "process_failure" else 0)

            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state") as save, patch(
                "run_monograph.run_script", side_effect=run
            ), contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(
                io.StringIO()
            ):
                self.assertEqual(2, orchestrate_finalize(args))
            save.assert_not_called()
            self.assertEqual(before, state)

    def test_real_producer_publisher_failure_preserves_old_disk_and_run_state_seam(
        self,
    ) -> None:
        from test_v024_finalization import V024FinalizationTests

        producer = V024FinalizationTests(
            methodName="test_deferred_finalization_requires_explicit_qa"
        )
        producer.setUp()
        try:
            formatted = producer.apply()
            work = self.root / "real-producer-publisher-seam"
            work.mkdir()
            final_dir = work / "final"
            final_dir.mkdir()
            final_docx = final_dir / f"{producer.source.stem}-finalized.docx"
            final_status = final_dir / "finalization.json"
            target_pdf = final_dir / f"{producer.source.stem}-target.pdf"
            audit = final_dir / "finalization-backend-audit.json"
            sentinels = {
                final_docx: b"old-finalized",
                final_status: b"old-finalization-status",
                target_pdf: b"old-target-pdf",
                audit: b"old-backend-audit",
            }
            for path, payload in sentinels.items():
                path.write_bytes(payload)
            state = {
                "status": "candidate_ready",
                "source": {
                    "path": str(producer.source),
                    "sha256": file_sha256(producer.source),
                },
                "profile": {
                    "path": str(producer.profile),
                    "sha256": file_sha256(producer.profile),
                },
                "structure_map": {
                    "path": str(producer.structure),
                    "sha256": file_sha256(producer.structure),
                },
                "artifacts": {"formatted": str(formatted)},
                "blockers": [],
                "qa_groups": [],
                "frozen_scopes": [],
                "stages": {},
                "metrics": {},
            }
            before = copy.deepcopy(state)
            args = SimpleNamespace(
                work_dir=work,
                resume=False,
                field_updater="deferred",
                field_updater_command=None,
                target_software=None,
                renderer=None,
                approve_deferred=True,
            )
            real_atomic_write_bytes = finalize_docx.atomic_write_bytes
            tampered = False

            def run(script, *arguments):
                nonlocal tampered

                def atomic_write(path, payload):
                    nonlocal tampered
                    real_atomic_write_bytes(path, payload)
                    if Path(path).name == "finalization.json":
                        value = json.loads(Path(path).read_text(encoding="utf-8"))
                        value["unknown_gate"] = True
                        Path(path).write_bytes(
                            finalize_docx.standard_json_bytes(value)
                        )
                        tampered = True

                with patch.object(
                    sys,
                    "argv",
                    [script, *(str(argument) for argument in arguments)],
                ), patch(
                    "finalize_docx.atomic_write_bytes", side_effect=atomic_write
                ), contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(
                    io.StringIO()
                ):
                    return SimpleNamespace(returncode=finalize_docx.main())

            with patch("run_monograph.load_state", return_value=state), patch(
                "run_monograph.run_script", side_effect=run
            ), patch("run_monograph.save_state") as save, contextlib.redirect_stdout(
                io.StringIO()
            ), contextlib.redirect_stderr(
                io.StringIO()
            ):
                self.assertEqual(2, orchestrate_finalize(args))
            self.assertTrue(tampered)
            save.assert_not_called()
            self.assertEqual(before, state)
            for path, payload in sentinels.items():
                self.assertEqual(payload, path.read_bytes())
            retained = list(final_dir.glob(".format-monograph-finalize-*"))
            self.assertEqual(1, len(retained))
            self.assertTrue((retained[0] / "finalization.json").exists())
            self.assertFalse(
                (final_dir / finalize_docx.PUBLICATION_RECOVERY_DIRECTORY).exists()
            )
        finally:
            producer.tearDown()

    def test_finalize_resume_revalidates_cache_tamper_matrix(self) -> None:
        for name, mutate in (
            ("docx", lambda work: (work / "final/source-finalized.docx").write_bytes(b"changed")),
            ("pdf", lambda work: (work / "final/source-target.pdf").write_bytes(b"changed")),
            (
                "json",
                lambda work: (work / "final/finalization.json").write_text(
                    '{"changed":true}', encoding="utf-8"
                ),
            ),
        ):
            work, state, args, _ = self._cached_final_ready_run(
                f"finalize-resume-{name}"
            )
            mutate(work)
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state") as save, patch(
                "run_monograph.run_script",
                return_value=SimpleNamespace(returncode=0),
            ) as run_script_mock, contextlib.redirect_stdout(
                io.StringIO()
            ), contextlib.redirect_stderr(
                io.StringIO()
            ):
                self.assertEqual(2 if name == "json" else 0, orchestrate_finalize(args))
                self.assertEqual(1, run_script_mock.call_count)
            self.assertEqual(
                "final_ready" if name == "json" else "candidate_ready",
                state["status"],
            )
            if name == "json":
                save.assert_not_called()
        _, state, args, _ = self._cached_final_ready_run("finalize-resume-valid")
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            return_value=SimpleNamespace(returncode=0),
        ) as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(0, orchestrate_finalize(args))
            self.assertEqual(1, run_script_mock.call_count)
        self.assertFalse(state["stages"]["finalize"]["cache_hit"])

    def test_verify_resume_revalidates_output_tamper_matrix(self) -> None:
        mutations = (
            (
                "render_page",
                lambda work: (work / "rendered/render-manifest.json").write_text(
                    '{"page_count":8,"target_pdf_source":"changed"}',
                    encoding="utf-8",
                ),
            ),
            (
                "render_content",
                lambda work: (work / "rendered/render-manifest.json").write_text(
                    '{"page_count":7,"target_pdf_source":"changed"}',
                    encoding="utf-8",
                ),
            ),
            ("pdf", lambda work: (work / "final/source-target.pdf").write_bytes(b"changed")),
            ("docx", lambda work: (work / "final/source-finalized.docx").write_bytes(b"changed")),
        )
        for name, mutate in mutations:
            work, state, _, args = self._cached_final_ready_run(
                f"verify-resume-{name}"
            )
            mutate(work)
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script"
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(2, verify(args))
                run_script_mock.assert_not_called()
            self.assertEqual("candidate_ready", state["status"])
        _, state, _, args = self._cached_final_ready_run("verify-resume-valid")
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch("run_monograph.run_script") as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(0, verify(args))
            run_script_mock.assert_not_called()
        self.assertTrue(state["stages"]["verify"]["cache_hit"])

    def test_status_downgrades_invalid_final_ready_without_regeneration(self) -> None:
        work, state, _, _ = self._cached_final_ready_run("status-invalid")
        (work / "final/source-target.pdf").write_bytes(b"changed")
        args = SimpleNamespace(work_dir=work, json=True)
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ) as save, patch("run_monograph.run_script") as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(2, orchestrate_status(args))
            save.assert_called_once()
            run_script_mock.assert_not_called()
        self.assertEqual("candidate_ready", state["status"])
        self.assertEqual("invalid", state["local_evidence_validation"]["status"])

    def test_external_status_is_read_only_and_revalidates_recorded_entities(self) -> None:
        work, state, _, _ = self._cached_final_ready_run(
            "external-status-valid"
        )
        args = SimpleNamespace(work_dir=work, json=True)
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch("run_monograph.run_script") as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(0, orchestrate_status(args))
            run_script_mock.assert_not_called()
        self.assertEqual("final_ready", state["status"])

        changed_work, changed_state, _, _ = self._cached_final_ready_run(
            "external-status-changed"
        )
        (changed_work / "external-updater").write_text(
            "#!/bin/sh\nexit 1\n", encoding="utf-8"
        )
        changed_args = SimpleNamespace(work_dir=changed_work, json=True)
        with patch("run_monograph.load_state", return_value=changed_state), patch(
            "run_monograph.save_state"
        ), patch("run_monograph.run_script") as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(2, orchestrate_status(changed_args))
            run_script_mock.assert_not_called()
        self.assertEqual("candidate_ready", changed_state["status"])
        self.assertTrue(
            any(
                "command resolution or dependency content changed" in error
                for error in changed_state["local_evidence_validation"]["errors"]
            )
        )

    def test_finalization_gate_single_field_conflict_matrix_is_rejected(self) -> None:
        mutations = (
            ("status", ("status",), "fail"),
            ("content", ("content_integrity",), "fail"),
            ("objects", ("protected_object_integrity",), "fail"),
            ("fonts", ("effective_font_integrity",), "fail"),
            ("source_hash", ("workflow_state", "source_sha256"), "0" * 64),
            ("input_hash", ("workflow_state", "input_sha256"), "0" * 64),
            ("profile_hash", ("workflow_state", "profile_sha256"), "0" * 64),
            (
                "structure_hash",
                ("workflow_state", "structure_map_sha256"),
                "0" * 64,
            ),
            ("output_hash", ("workflow_state", "output_sha256"), "0" * 64),
            ("stage", ("workflow_state", "stage"), "verified"),
            ("output_path", ("output",), "/tmp/different-finalized.docx"),
            ("target_pdf", ("target_pdf",), "/tmp/different-target.pdf"),
            ("target_layout", ("target_layout_status",), "not_verified"),
        )
        for name, path, value in mutations:
            work, state, _, _ = self._cached_final_ready_run(
                f"gate-conflict-{name}"
            )
            status_path = work / "final/finalization.json"
            finalization = json.loads(status_path.read_text(encoding="utf-8"))
            target = finalization
            for key in path[:-1]:
                target = target[key]
            target[path[-1]] = value
            self.assertTrue(final_ready_field_evidence(completion_evidence(finalization)))
            status_path.write_text(json.dumps(finalization), encoding="utf-8")
            args = SimpleNamespace(work_dir=work, json=True)
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script"
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(2, orchestrate_status(args))
                run_script_mock.assert_not_called()
            self.assertEqual("candidate_ready", state["status"])

    def test_evidence_validation_conflict_matrix_rejects_read_only_reuse(self) -> None:
        mutations = (
            (
                "status",
                lambda finalization: finalization["field_completion"][
                    "evidence_validation"
                ].update({"status": "incomplete"}),
            ),
            (
                "errors",
                lambda finalization: finalization["field_completion"][
                    "evidence_validation"
                ].update({"errors": ["forged"]}),
            ),
            (
                "missing",
                lambda finalization: finalization["field_completion"].pop(
                    "evidence_validation"
                ),
            ),
            (
                "recalculation_conflict",
                lambda finalization: finalization["field_completion"][
                    "evidence_validation"
                ].update({"status": "pass", "errors": ["not-recalculated"]}),
            ),
        )
        for mutation_name, mutate in mutations:
            for entrypoint in ("finalize", "verify", "status"):
                work, state, finalize_args, verify_args = self._cached_final_ready_run(
                    f"evidence-validation-{mutation_name}-{entrypoint}"
                )
                finalization_path = work / "final/finalization.json"
                finalization = json.loads(
                    finalization_path.read_text(encoding="utf-8")
                )
                mutate(finalization)
                self.assertTrue(
                    final_ready_field_evidence(completion_evidence(finalization))
                )
                finalization_path.write_text(
                    json.dumps(finalization), encoding="utf-8"
                )
                args = (
                    finalize_args
                    if entrypoint == "finalize"
                    else verify_args
                    if entrypoint == "verify"
                    else SimpleNamespace(work_dir=work, json=True)
                )
                handler = {
                    "finalize": orchestrate_finalize,
                    "verify": verify,
                    "status": orchestrate_status,
                }[entrypoint]
                with self.subTest(
                    mutation=mutation_name, entrypoint=entrypoint
                ), patch("run_monograph.load_state", return_value=state), patch(
                    "run_monograph.save_state"
                ) as save, patch(
                    "run_monograph.run_script",
                    return_value=SimpleNamespace(returncode=0),
                ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                    if entrypoint == "finalize":
                        self.assertEqual(2, handler(args))
                        self.assertEqual(1, run_script_mock.call_count)
                        save.assert_not_called()
                    else:
                        self.assertEqual(2, handler(args))
                        run_script_mock.assert_not_called()
                self.assertEqual(
                    "final_ready" if entrypoint == "finalize" else "candidate_ready",
                    state["status"],
                )

    def test_allowlisted_target_resolver_rejects_word_substrings(self) -> None:
        for alias in (
            "Microsoft Word",
            "Microsoft Word 2016",
            "Microsoft Word 2019",
            "Microsoft Word 2021",
            "Microsoft Word for Mac",
            "Microsoft 365",
            "microsoft_word",
        ):
            with self.subTest(alias=alias):
                self.assertEqual(MICROSOFT_WORD, resolve_target_id(alias))
        for alias in ("LibreOffice", "LibreOffice Writer", "libreoffice"):
            with self.subTest(alias=alias):
                self.assertEqual(LIBREOFFICE, resolve_target_id(alias))
        for value in (
            "WordPerfect",
            "word processor",
            "password",
            "random-word-target",
            "WPS Writer",
            "",
            None,
        ):
            with self.subTest(value=value):
                self.assertEqual(UNSUPPORTED, resolve_target_id(value))

    def test_unparseable_and_missing_external_commands_are_unavailable(self) -> None:
        invalid = external_command_identity("[not-json", self.root)
        missing = external_command_identity(
            "format-monograph-definitely-missing-adapter --fixed", self.root
        )
        self.assertEqual("invalid", invalid["status"])
        self.assertEqual("unavailable", missing["status"])
        self.assertTrue(external_command_identity_errors("invalid", invalid))
        self.assertTrue(external_command_identity_errors("missing", missing))

    def test_implicit_profile_word_target_hits_finalize_cache(self) -> None:
        work, state, args, _ = self._cached_final_ready_run(
            "implicit-profile-word"
        )
        profile = work / "profile.json"
        profile.write_text(
            json.dumps({"target_applications": ["Microsoft 365"]}),
            encoding="utf-8",
        )
        state["profile"]["sha256"] = file_sha256(profile)
        finalization_path = work / "final/finalization.json"
        finalization = json.loads(finalization_path.read_text(encoding="utf-8"))
        finalization["workflow_state"]["profile_sha256"] = file_sha256(profile)
        finalization_path.write_text(json.dumps(finalization), encoding="utf-8")
        args.target_software = None
        request = finalize_request_identity(
            args, target_id=MICROSOFT_WORD, renderer_used=False
        )
        state["finalization_request"] = request
        state["finalization_gate"] = canonical_finalization_gate_summary(
            work, finalization
        )
        state["stages"]["finalize"]["input_key_sha256"] = json_sha256(
            {
                "formatted": file_sha256(work / "formatted.docx"),
                "map": file_sha256(work / "structure.json"),
                "behavior": request,
            }
        )
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            return_value=SimpleNamespace(returncode=0),
        ) as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(0, orchestrate_finalize(args))
            self.assertEqual(1, run_script_mock.call_count)
        self.assertEqual(MICROSOFT_WORD, state["finalization_request"]["target_software"])
        self.assertFalse(state["stages"]["finalize"]["cache_hit"])

    def test_external_updater_entity_and_arguments_control_finalize_cache(self) -> None:
        for mutation in ("executable", "script", "fixed_argument"):
            work, state, args, _ = self._cached_final_ready_run(
                f"updater-identity-{mutation}"
            )
            executable = work / "adapter"
            script = work / "adapter.py"
            executable.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
            executable.chmod(0o755)
            script.write_text("print('v1')\n", encoding="utf-8")
            command = json.dumps(
                [str(executable), str(script), "--mode", "one"]
            )
            args.field_updater_command = command
            initial = finalize_request_identity(
                args, target_id=MICROSOFT_WORD, renderer_used=False
            )
            self.assertEqual("available", initial["external_updater"]["status"])
            self.assertEqual(1, len(initial["external_updater"]["file_arguments"]))
            state["finalization_request"] = initial
            state["stages"]["finalize"]["input_key_sha256"] = json_sha256(
                {
                    "formatted": file_sha256(work / "formatted.docx"),
                    "map": file_sha256(work / "structure.json"),
                    "behavior": initial,
                }
            )
            if mutation == "executable":
                executable.write_text("#!/bin/sh\nexit 1\n", encoding="utf-8")
                self.assertTrue(
                    external_command_identity_errors(
                        "cached", initial["external_updater"]
                    )
                )
            elif mutation == "script":
                script.write_text("print('v2')\n", encoding="utf-8")
                self.assertTrue(
                    external_command_identity_errors(
                        "cached", initial["external_updater"]
                    )
                )
            else:
                args.field_updater_command = json.dumps(
                    [str(executable), str(script), "--mode", "two"]
                )
            changed = finalize_request_identity(
                args, target_id=MICROSOFT_WORD, renderer_used=False
            )
            self.assertNotEqual(initial, changed)
            with self.subTest(mutation=mutation), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script",
                return_value=SimpleNamespace(returncode=0),
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(0, orchestrate_finalize(args))
                self.assertEqual(1, run_script_mock.call_count)
            self.assertFalse(state["stages"]["finalize"]["cache_hit"])

        work, state, args, _ = self._cached_final_ready_run(
            "updater-identity-unchanged"
        )
        unchanged = finalize_request_identity(
            args, target_id=MICROSOFT_WORD, renderer_used=False
        )
        self.assertEqual([], external_command_identity_errors("current", unchanged["external_updater"]))
        self.assertFalse(external_command_cache_reusable(unchanged["external_updater"]))
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            return_value=SimpleNamespace(returncode=0),
        ) as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(0, orchestrate_finalize(args))
            self.assertEqual(1, run_script_mock.call_count)
        self.assertFalse(state["stages"]["finalize"]["cache_hit"])

    def test_external_dependency_closure_argument_and_python_matrix(self) -> None:
        executable = self.root / "adapter"
        executable.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        executable.chmod(0o755)
        config = self.root / "config.json"
        config.write_text('{"mode":"one"}', encoding="utf-8")
        response = self.root / "arguments.rsp"
        response.write_text("--mode one", encoding="utf-8")
        directory = self.root / "inputs"
        directory.mkdir()
        (directory / "nested.txt").write_text("one", encoding="utf-8")
        direct_script = self.root / "adapter.py"
        direct_script.write_text("print('ok')\n", encoding="utf-8")
        module = self.root / "closed_module.py"
        module.write_text("VALUE = 1\n", encoding="utf-8")

        audited_cases = {
            "config_equals": [str(executable), f"--config={config}"],
            "config_separate": [str(executable), "--config", str(config)],
            "relative_file": [str(executable), "config.json"],
            "directory": [str(executable), str(directory)],
            "fixed_scalar": [str(executable), "--mode", "one"],
            "python_script": [sys.executable, str(direct_script)],
            "python_module": [sys.executable, "-m", "closed_module"],
        }
        for name, argv in audited_cases.items():
            with self.subTest(name=name):
                identity = external_command_identity(json.dumps(argv), self.root)
                self.assertEqual("available", identity["status"])
                self.assertEqual(3, identity["version"])
                self.assertEqual(2, identity["dependency_closure"]["version"])
                self.assertEqual(
                    "unproven", identity["dependency_closure"]["status"]
                )
                self.assertFalse(external_command_cache_reusable(identity))
                self.assertIn(
                    "external_program_not_hermetic",
                    identity["dependency_closure"]["incomplete_reasons"],
                )
                self.assertIn(
                    "runtime_dependencies_unproven",
                    identity["dependency_closure"]["incomplete_reasons"],
                )
                forged = json.loads(json.dumps(identity))
                forged["cache_reusable"] = True
                forged["dependency_closure"].update(
                    {"status": "complete", "cache_reusable": True}
                )
                self.assertFalse(external_command_cache_reusable(forged))
                self.assertTrue(external_command_identity_errors("forged", forged))

        incomplete_cases = {
            "response_file": [str(executable), f"@{response}"],
            "python_unresolved_module": [sys.executable, "-m", "json.tool"],
            "python_dynamic_code": [sys.executable, "-c", "print('ok')"],
        }
        for name, argv in incomplete_cases.items():
            with self.subTest(name=name):
                identity = external_command_identity(json.dumps(argv), self.root)
                self.assertEqual("available", identity["status"])
                self.assertEqual(
                    "unproven", identity["dependency_closure"]["status"]
                )
                self.assertFalse(external_command_cache_reusable(identity))
                self.assertEqual([], external_command_identity_errors(name, identity))

        imported_script = self.root / "imported.py"
        imported_script.write_text("import json\nprint(json.dumps({}))\n", encoding="utf-8")
        imported = external_command_identity(
            json.dumps([sys.executable, str(imported_script)]), self.root
        )
        self.assertFalse(external_command_cache_reusable(imported))
        self.assertIn(
            "python_external_import:json",
            imported["dependency_closure"]["incomplete_reasons"],
        )
        file_opening_script = self.root / "file-opening.py"
        file_opening_script.write_text(
            "print(open('implicit.cfg').read())\n", encoding="utf-8"
        )
        file_opening = external_command_identity(
            json.dumps([sys.executable, str(file_opening_script)]), self.root
        )
        self.assertFalse(external_command_cache_reusable(file_opening))
        self.assertIn(
            "python_unbounded_call:open",
            file_opening["dependency_closure"]["incomplete_reasons"],
        )

    def test_external_dependency_entity_change_matrix_forces_finalize_rerun(self) -> None:
        cases = ("config_equals", "config_separate", "response", "directory")
        for name in cases:
            work, state, args, _ = self._cached_final_ready_run(
                f"dependency-entity-{name}"
            )
            executable = work / "adapter"
            executable.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
            executable.chmod(0o755)
            dependency = work / ("tree" if name == "directory" else "config.json")
            if name == "directory":
                dependency.mkdir()
                (dependency / "entry.txt").write_text("one", encoding="utf-8")
                argv = [str(executable), str(dependency)]
            else:
                dependency.write_text("one", encoding="utf-8")
                if name == "config_equals":
                    argv = [str(executable), f"--config={dependency}"]
                elif name == "config_separate":
                    relative = os.path.relpath(dependency, SKILL)
                    argv = [str(executable), "--config", relative]
                else:
                    argv = [str(executable), f"@{dependency}"]
            args.field_updater_command = json.dumps(argv)
            initial = finalize_request_identity(
                args, target_id=MICROSOFT_WORD, renderer_used=False
            )
            state["finalization_request"] = initial
            state["stages"]["finalize"]["input_key_sha256"] = json_sha256(
                {
                    "formatted": file_sha256(work / "formatted.docx"),
                    "map": file_sha256(work / "structure.json"),
                    "behavior": initial,
                }
            )
            if name == "directory":
                (dependency / "entry.txt").write_text("two", encoding="utf-8")
            else:
                dependency.write_text("two", encoding="utf-8")
            changed = finalize_request_identity(
                args, target_id=MICROSOFT_WORD, renderer_used=False
            )
            self.assertNotEqual(initial, changed)
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script",
                return_value=SimpleNamespace(returncode=0),
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(0, orchestrate_finalize(args))
                self.assertGreater(run_script_mock.call_count, 0)
            self.assertFalse(state["stages"]["finalize"]["cache_hit"])

    def test_directory_and_symlink_dependency_mutation_matrix(self) -> None:
        executable = self.root / "adapter"
        executable.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        executable.chmod(0o755)
        for mutation in ("add", "delete", "rename", "content"):
            directory = self.root / f"directory-{mutation}"
            directory.mkdir()
            entry = directory / "entry.txt"
            entry.write_text("one", encoding="utf-8")
            command = json.dumps([str(executable), str(directory)])
            initial = external_command_identity(command, self.root)
            self.assertFalse(external_command_cache_reusable(initial))
            if mutation == "add":
                (directory / "added.txt").write_text("added", encoding="utf-8")
            elif mutation == "delete":
                entry.unlink()
            elif mutation == "rename":
                entry.rename(directory / "renamed.txt")
            else:
                entry.write_text("two", encoding="utf-8")
            with self.subTest(mutation=mutation):
                changed = external_command_identity(command, self.root)
                self.assertNotEqual(initial, changed)
                self.assertTrue(external_command_identity_errors("directory", initial))

        target_one = self.root / "target-one.json"
        target_two = self.root / "target-two.json"
        target_one.write_text("one", encoding="utf-8")
        target_two.write_text("two", encoding="utf-8")
        linked = self.root / "linked.json"
        linked.symlink_to(target_one)
        command = json.dumps([str(executable), str(linked)])
        initial = external_command_identity(command, self.root)
        linked.unlink()
        linked.symlink_to(target_two)
        changed = external_command_identity(command, self.root)
        self.assertNotEqual(initial, changed)
        self.assertTrue(external_command_identity_errors("symlink", initial))

        directory_target = self.root / "directory-symlink-target"
        directory_target.mkdir()
        directory_link = self.root / "directory-link"
        directory_link.symlink_to(directory_target, target_is_directory=True)
        unresolved = external_command_identity(
            json.dumps([str(executable), str(directory_link)]), self.root
        )
        self.assertFalse(external_command_cache_reusable(unresolved))

        tree = self.root / "tree-with-link"
        outside = self.root / "outside-tree"
        tree.mkdir()
        outside.mkdir()
        (outside / "must-not-be-scanned.txt").write_text("outside", encoding="utf-8")
        (tree / "escape").symlink_to(outside, target_is_directory=True)
        linked_tree = external_command_identity(
            json.dumps([str(executable), str(tree)]), self.root
        )
        entries = linked_tree["dependency_closure"]["entities"][0]["entries"]
        self.assertEqual(
            [{"path": "escape", "type": "symlink", "link_target": str(outside)}],
            entries[1:],
        )
        self.assertFalse(external_command_cache_reusable(linked_tree))

    def test_directory_and_file_symlink_mutations_force_finalize_rerun(self) -> None:
        for mutation in ("add", "delete", "rename", "content", "symlink_target"):
            work, state, args, _ = self._cached_final_ready_run(
                f"dependency-rerun-{mutation}"
            )
            executable = work / "adapter"
            executable.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
            executable.chmod(0o755)
            if mutation == "symlink_target":
                target_one = work / "target-one.json"
                target_two = work / "target-two.json"
                target_one.write_text("one", encoding="utf-8")
                target_two.write_text("two", encoding="utf-8")
                dependency = work / "linked.json"
                dependency.symlink_to(target_one)
            else:
                dependency = work / "tree"
                dependency.mkdir()
                entry = dependency / "entry.txt"
                entry.write_text("one", encoding="utf-8")
            args.field_updater_command = json.dumps(
                [str(executable), str(dependency)]
            )
            initial = finalize_request_identity(
                args, target_id=MICROSOFT_WORD, renderer_used=False
            )
            self.assertFalse(
                external_command_cache_reusable(initial["external_updater"])
            )
            state["finalization_request"] = initial
            state["stages"]["finalize"]["input_key_sha256"] = json_sha256(
                {
                    "formatted": file_sha256(work / "formatted.docx"),
                    "map": file_sha256(work / "structure.json"),
                    "behavior": initial,
                }
            )
            if mutation == "add":
                (dependency / "added.txt").write_text("added", encoding="utf-8")
            elif mutation == "delete":
                entry.unlink()
            elif mutation == "rename":
                entry.rename(dependency / "renamed.txt")
            elif mutation == "content":
                entry.write_text("two", encoding="utf-8")
            else:
                dependency.unlink()
                dependency.symlink_to(target_two)
            with self.subTest(mutation=mutation), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script",
                return_value=SimpleNamespace(returncode=0),
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(0, orchestrate_finalize(args))
                self.assertGreater(run_script_mock.call_count, 0)
            self.assertFalse(state["stages"]["finalize"]["cache_hit"])

    def test_path_executable_and_unprovable_python_module_cache_behavior(self) -> None:
        bin_dir = self.root / "bin"
        bin_dir.mkdir()
        target_one = bin_dir / "adapter-one"
        target_two = bin_dir / "adapter-two"
        for path, status in ((target_one, 0), (target_two, 1)):
            path.write_text(f"#!/bin/sh\nexit {status}\n", encoding="utf-8")
            path.chmod(0o755)
        linked = bin_dir / "adapter"
        linked.symlink_to(target_one)
        with patch.dict(os.environ, {"PATH": "bin"}):
            initial = external_command_identity("adapter --mode one", self.root)
            self.assertFalse(external_command_cache_reusable(initial))
            self.assertEqual(
                str(self.root.resolve() / "bin/adapter"),
                initial["executable"]["located_path"],
            )
            linked.unlink()
            linked.symlink_to(target_two)
            changed = external_command_identity("adapter --mode one", self.root)
        self.assertNotEqual(initial, changed)
        self.assertNotEqual(
            initial["executable"]["resolved_path"],
            changed["executable"]["resolved_path"],
        )

        work, state, args, _ = self._cached_final_ready_run(
            "unprovable-python-module"
        )
        args.field_updater_command = json.dumps([sys.executable, "-m", "json.tool"])
        request = finalize_request_identity(
            args, target_id=MICROSOFT_WORD, renderer_used=False
        )
        self.assertFalse(external_command_cache_reusable(request["external_updater"]))
        state["finalization_request"] = request
        state["stages"]["finalize"]["input_key_sha256"] = json_sha256(
            {
                "formatted": file_sha256(work / "formatted.docx"),
                "map": file_sha256(work / "structure.json"),
                "behavior": request,
            }
        )
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            return_value=SimpleNamespace(returncode=0),
        ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
            self.assertEqual(0, orchestrate_finalize(args))
            self.assertEqual(0, orchestrate_finalize(args))
            self.assertEqual(2, run_script_mock.call_count)
        self.assertFalse(state["stages"]["finalize"]["cache_hit"])
        self.assertEqual("candidate_ready", state["status"])
        self.assertEqual([], state["blockers"])
        self.assertEqual(
            [], state["field_writeback"]["completion_evidence_errors"]
        )
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch("run_monograph.run_script") as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(
                0, orchestrate_status(SimpleNamespace(work_dir=work, json=True))
            )
            run_script_mock.assert_not_called()

    def test_native_wrapper_hidden_dependency_matrix_always_reruns(self) -> None:
        cases = (
            "native",
            "implicit_cwd_config",
            "environment",
            "path_subcommand",
            "usr_bin_env",
        )
        for name in cases:
            work, state, args, _ = self._cached_final_ready_run(
                f"nonhermetic-native-{name}"
            )
            wrapper = work / "wrapper"
            wrapper.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
            wrapper.chmod(0o755)
            hidden_config = work / "implicit.cfg"
            hidden_config.write_text("one", encoding="utf-8")
            bin_dir = work / "bin"
            bin_dir.mkdir()
            subcommand = bin_dir / "subcommand"
            subcommand.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
            subcommand.chmod(0o755)
            if name == "native":
                argv = ["/usr/bin/true", "--fixed", "one"]
            elif name == "usr_bin_env":
                argv = ["/usr/bin/env", str(wrapper), "--mode", "one"]
            else:
                argv = [str(wrapper), "--mode", "one"]
            request = self._bind_external_finalize_command(
                work, state, args, argv
            )
            self.assertFalse(external_command_cache_reusable(request["external_updater"]))
            self.assertEqual(
                "unproven",
                request["external_updater"]["dependency_closure"]["status"],
            )
            with patch.dict(
                os.environ,
                {
                    "UPDATER_CONFIG": str(hidden_config),
                    "PATH": str(bin_dir),
                },
            ), patch("run_monograph.load_state", return_value=state), patch(
                "run_monograph.save_state"
            ), patch(
                "run_monograph.run_script",
                return_value=SimpleNamespace(returncode=0),
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(0, orchestrate_finalize(args))
                if name == "implicit_cwd_config":
                    hidden_config.write_text("two", encoding="utf-8")
                elif name == "environment":
                    os.environ["UPDATER_CONFIG"] = str(work / "different.cfg")
                elif name == "path_subcommand":
                    subcommand.write_text("#!/bin/sh\nexit 1\n", encoding="utf-8")
                self.assertEqual(0, orchestrate_finalize(args))
                self.assertEqual(2, run_script_mock.call_count)
            self.assertFalse(state["stages"]["finalize"]["cache_hit"])

    def test_python_alias_reflection_and_pythonpath_matrix_always_reruns(self) -> None:
        module = self.root / "local_module.py"
        module.write_text("VALUE = 1\n", encoding="utf-8")
        package = self.root / "local_package"
        package.mkdir()
        (package / "__init__.py").write_text("", encoding="utf-8")
        (package / "__main__.py").write_text("print('ok')\n", encoding="utf-8")
        sitecustomize = self.root / "sitecustomize.py"
        sitecustomize.write_text("VALUE = 1\n", encoding="utf-8")
        scripts = {
            "direct_script": "print('ok')\n",
            "reader_alias": "reader = open\nprint(reader('implicit.cfg').read())\n",
            "builtins_reflection": (
                "reader = getattr(__builtins__, 'open')\n"
                "print(reader('implicit.cfg').read())\n"
            ),
            "import_reflection": (
                "loader = getattr(__builtins__, '__import__')\n"
                "print(loader('json').dumps({}))\n"
            ),
        }
        commands: dict[str, list[str]] = {}
        for name, source in scripts.items():
            script = self.root / f"{name}.py"
            script.write_text(source, encoding="utf-8")
            commands[name] = [sys.executable, str(script)]
        commands.update(
            {
                "module": [sys.executable, "-m", "local_module"],
                "package": [sys.executable, "-m", "local_package"],
                "sitecustomize": [sys.executable, str(self.root / "direct_script.py")],
            }
        )
        for name, argv in commands.items():
            work, state, args, _ = self._cached_final_ready_run(
                f"nonhermetic-python-{name}"
            )
            request = self._bind_external_finalize_command(
                work, state, args, argv
            )
            updater = request["external_updater"]
            self.assertFalse(external_command_cache_reusable(updater))
            self.assertIn(
                "runtime_dependencies_unproven",
                updater["dependency_closure"]["incomplete_reasons"],
            )
            with self.subTest(name=name), patch.dict(
                os.environ, {"PYTHONPATH": str(self.root)}
            ), patch("run_monograph.load_state", return_value=state), patch(
                "run_monograph.save_state"
            ), patch(
                "run_monograph.run_script",
                return_value=SimpleNamespace(returncode=0),
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(0, orchestrate_finalize(args))
                if name == "sitecustomize":
                    sitecustomize.write_text("VALUE = 2\n", encoding="utf-8")
                self.assertEqual(0, orchestrate_finalize(args))
                self.assertEqual(2, run_script_mock.call_count)
            self.assertFalse(state["stages"]["finalize"]["cache_hit"])

    def test_non_external_no_fields_finalize_cache_remains_reusable(self) -> None:
        work, state, args, _ = self._cached_no_fields_final_ready_run(
            "non-external-no-fields-cache"
        )
        formatted = work / "formatted.docx"
        document = Document()
        document.add_paragraph("No fields")
        document.save(formatted)
        finalization_path = work / "final/finalization.json"
        finalization = json.loads(finalization_path.read_text(encoding="utf-8"))
        finalization["workflow_state"]["input_sha256"] = file_sha256(formatted)
        finalization_path.write_text(json.dumps(finalization), encoding="utf-8")
        args.field_updater = "auto"
        args.field_updater_command = None
        request = finalize_request_identity(
            args, target_id=LIBREOFFICE, renderer_used=False
        )
        self.assertIsNone(request["external_updater"])
        state["finalization_request"] = request
        state["finalization_gate"] = canonical_finalization_gate_summary(
            work, finalization
        )
        state["stages"]["finalize"]["input_key_sha256"] = json_sha256(
            {
                "formatted": file_sha256(formatted),
                "map": file_sha256(work / "structure.json"),
                "behavior": request,
            }
        )
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch("run_monograph.run_script") as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(0, orchestrate_finalize(args))
            run_script_mock.assert_not_called()
        self.assertTrue(state["stages"]["finalize"]["cache_hit"])
        self.assertEqual("final_ready", state["status"])
        self.assertIn("verify", state["stages"])
        self.assertIn("verification_outputs", state)
        self.assertIn("visual_qa", state)
        self.assertIn("render_manifest", state["artifacts"])

    def test_real_external_finalize_invalidates_verify_then_verify_rebuilds(self) -> None:
        work, state, finalize_args, verify_args = self._cached_final_ready_run(
            "external-finalize-invalidates-verify"
        )
        state["metrics"]["rendered_pages"] = 7
        state["local_evidence_validation"] = {
            "status": "valid",
            "source": "verify_resume",
            "errors": [],
        }
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            return_value=SimpleNamespace(returncode=0),
        ) as finalize_run, contextlib.redirect_stdout(io.StringIO()):
            self.assertEqual(0, orchestrate_finalize(finalize_args))
        self.assertEqual(1, finalize_run.call_count)
        self.assertEqual("finalize_docx.py", finalize_run.call_args.args[0])
        self.assertEqual("candidate_ready", state["status"])
        self.assertEqual([], state["blockers"])
        self.assertNotIn("verify", state["stages"])
        self.assertNotIn("verification_outputs", state)
        self.assertNotIn("visual_qa", state)
        self.assertNotIn("local_evidence_validation", state)
        self.assertNotIn("rendered_pages", state["metrics"])
        for name in ("final_audit", "render_manifest", "visual_qa_manifest"):
            self.assertNotIn(name, state["artifacts"])
        for name in ("formatted", "finalized", "finalization_status", "target_pdf"):
            self.assertIn(name, state["artifacts"])

        request = verify_request_identity(
            verify_args, target_id=MICROSOFT_WORD, renderer_used=False
        )
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            side_effect=self._verification_script_result(
                work,
                target_software=MICROSOFT_WORD,
                renderer_identity=request["renderer"],
                use_target_pdf=True,
            ),
        ) as verify_run, contextlib.redirect_stdout(io.StringIO()):
            self.assertEqual(0, verify(verify_args))
        self.assertEqual(
            ["audit_docx.py", "render_docx.py"],
            [call.args[0] for call in verify_run.call_args_list],
        )
        self.assertEqual("final_ready", state["status"])
        self.assertEqual([], state["blockers"])
        self.assertFalse(state["stages"]["verify"]["cache_hit"])
        self.assertEqual(
            VERIFICATION_OUTPUT_VERSION, state["verification_outputs"]["version"]
        )
        self.assertEqual(7, state["metrics"]["rendered_pages"])

    def test_verify_after_real_finalize_rejects_invalid_downstream_evidence(self) -> None:
        cases = (
            ("missing_visual", None, 7, MICROSOFT_WORD, 0),
            ("page_conflict", "valid", 8, MICROSOFT_WORD, 2),
            ("target_conflict", "valid", 7, LIBREOFFICE, 2),
        )
        for name, visual_kind, page_count, manifest_target, expected in cases:
            work, state, finalize_args, verify_args = self._cached_final_ready_run(
                f"finalize-then-verify-{name}"
            )
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script",
                return_value=SimpleNamespace(returncode=0),
            ), contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(0, orchestrate_finalize(finalize_args))
            if visual_kind is None:
                verify_args.visual_qa_manifest = None
            request = verify_request_identity(
                verify_args, target_id=MICROSOFT_WORD, renderer_used=False
            )
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script",
                side_effect=self._verification_script_result(
                    work,
                    target_software=MICROSOFT_WORD,
                    renderer_identity=request["renderer"],
                    use_target_pdf=True,
                    manifest_target=manifest_target,
                    page_count=page_count,
                ),
            ) as verify_run, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(expected, verify(verify_args))
            self.assertEqual(2, verify_run.call_count)
            self.assertEqual("candidate_ready", state["status"])
            self.assertTrue(state["blockers"])

        work, state, finalize_args, verify_args = self._cached_final_ready_run(
            "finalize-then-verify-invalid-visual"
        )
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script", return_value=SimpleNamespace(returncode=0)
        ), contextlib.redirect_stdout(io.StringIO()):
            self.assertEqual(0, orchestrate_finalize(finalize_args))
        verify_args.visual_qa_manifest.write_text(
            '{"all_pages_inspected":false,"target_layout_verified":true,'
            '"page_count":7,"issues":[]}',
            encoding="utf-8",
        )
        request = verify_request_identity(
            verify_args, target_id=MICROSOFT_WORD, renderer_used=False
        )
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            side_effect=self._verification_script_result(
                work,
                target_software=MICROSOFT_WORD,
                renderer_identity=request["renderer"],
                use_target_pdf=True,
            ),
        ) as verify_run, contextlib.redirect_stdout(io.StringIO()), self.assertRaises(
            RunError
        ):
            verify(verify_args)
        self.assertEqual(2, verify_run.call_count)
        self.assertNotEqual("final_ready", state["status"])

    def test_real_nonexternal_finalize_invalidates_verify_state(self) -> None:
        work, state, args, _ = self._cached_no_fields_final_ready_run(
            "nonexternal-finalize-invalidates-verify"
        )
        formatted = work / "formatted.docx"
        document = Document()
        document.add_paragraph("No fields")
        document.save(formatted)
        finalization_path = work / "final/finalization.json"
        finalization = json.loads(finalization_path.read_text(encoding="utf-8"))
        finalization["workflow_state"]["input_sha256"] = file_sha256(formatted)
        finalization_path.write_text(json.dumps(finalization), encoding="utf-8")
        args.resume = False
        args.field_updater = "auto"
        args.field_updater_command = None
        state["metrics"]["rendered_pages"] = 7
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script", return_value=SimpleNamespace(returncode=0)
        ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
            self.assertEqual(0, orchestrate_finalize(args))
        self.assertEqual(1, run_script_mock.call_count)
        self.assertNotIn("verify", state["stages"])
        self.assertNotIn("verification_outputs", state)
        self.assertNotIn("visual_qa", state)
        self.assertNotIn("rendered_pages", state["metrics"])

    def test_verify_resume_without_new_finalize_still_cache_hits(self) -> None:
        _, state, _, args = self._cached_final_ready_run(
            "ordinary-verify-resume-cache"
        )
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch("run_monograph.run_script") as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(0, verify(args))
        run_script_mock.assert_not_called()
        self.assertTrue(state["stages"]["verify"]["cache_hit"])
        self.assertEqual("final_ready", state["status"])

    def test_finalize_resume_behavior_identity_cache_miss_matrix(self) -> None:
        renderer = self.root / "alternate-soffice"
        renderer.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        renderer.chmod(0o755)
        cases = (
            ("approve_deferred", lambda args: setattr(args, "approve_deferred", True)),
            ("target", lambda args: setattr(args, "target_software", "LibreOffice Writer")),
        )
        for name, mutate in cases:
            _, state, args, _ = self._cached_final_ready_run(
                f"finalize-behavior-{name}"
            )
            mutate(args)
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script",
                return_value=SimpleNamespace(returncode=0),
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(0, orchestrate_finalize(args))
            self.assertEqual("finalize_docx.py", run_script_mock.call_args.args[0])
            self.assertFalse(state["stages"]["finalize"]["cache_hit"])

        _, state, args, _ = self._cached_final_ready_run(
            "finalize-unused-renderer"
        )
        args.renderer = str(renderer)
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            return_value=SimpleNamespace(returncode=0),
        ) as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(0, orchestrate_finalize(args))
            self.assertEqual(1, run_script_mock.call_count)
        self.assertFalse(state["stages"]["finalize"]["cache_hit"])

    def test_verify_resume_legal_visual_manifest_change_reruns(self) -> None:
        work, state, _, args = self._cached_final_ready_run("verify-visual-b")
        visual_b = work / "visual-b.json"
        visual_b.write_text(
            json.dumps(
                {
                    "all_pages_inspected": True,
                    "target_layout_verified": True,
                    "page_count": 7,
                    "issues": [],
                }
            ),
            encoding="utf-8",
        )
        args.visual_qa_manifest = visual_b
        request = verify_request_identity(args)
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            side_effect=self._verification_script_result(
                work,
                target_software="Microsoft Word",
                renderer_identity=request["renderer"],
                use_target_pdf=True,
            ),
        ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
            self.assertEqual(0, verify(args))
        self.assertEqual(2, run_script_mock.call_count)
        self.assertEqual("final_ready", state["status"])
        self.assertEqual(
            str(visual_b.resolve()), state["artifacts"]["visual_qa_manifest"]
        )
        self.assertFalse(state["stages"]["verify"]["cache_hit"])

    def test_verify_resume_request_switch_matrix_reruns_and_rechecks_target(self) -> None:
        renderer = self.root / "verify-alternate-soffice"
        renderer.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        renderer.chmod(0o755)
        cases = (
            ("word_to_libreoffice", None, "LibreOffice Writer", "libreoffice", 2),
            ("target_inconsistent", None, "LibreOffice Writer", "microsoft_word", 2),
        )
        for name, requested_renderer, requested_target, manifest_target, expected in cases:
            work, state, _, args = self._cached_final_ready_run(
                f"verify-switch-{name}"
            )
            args.renderer = requested_renderer
            args.target_software = requested_target
            request = verify_request_identity(args, renderer_used=True)
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script",
                side_effect=self._verification_script_result(
                    work,
                    target_software=request["target_software"],
                    renderer_identity=request["renderer"],
                    use_target_pdf=False,
                    manifest_target=manifest_target,
                ),
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(expected, verify(args))
            self.assertEqual(2, run_script_mock.call_count)
            self.assertFalse(state["stages"]["verify"]["cache_hit"])
            self.assertEqual(
                "final_ready" if expected == 0 else "candidate_ready",
                state["status"],
            )

        for name, requested_renderer, requested_target in (
            ("unused_renderer", str(renderer), "Microsoft Word"),
            ("implicit_target", None, None),
        ):
            _, state, _, args = self._cached_final_ready_run(
                f"verify-switch-{name}"
            )
            args.renderer = requested_renderer
            args.target_software = requested_target
            with self.subTest(name=name), patch(
                "run_monograph.load_state", return_value=state
            ), patch("run_monograph.save_state"), patch(
                "run_monograph.run_script"
            ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
                self.assertEqual(0, verify(args))
                run_script_mock.assert_not_called()
            self.assertTrue(state["stages"]["verify"]["cache_hit"])

    def test_verify_resume_libreoffice_to_word_reruns_without_false_final_ready(self) -> None:
        work, state, _, args = self._cached_no_fields_final_ready_run(
            "verify-libreoffice-to-word"
        )
        args.target_software = "Microsoft Word"
        request = verify_request_identity(args, renderer_used=True)
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch(
            "run_monograph.run_script",
            side_effect=self._verification_script_result(
                work,
                target_software="Microsoft Word",
                renderer_identity=request["renderer"],
                use_target_pdf=False,
            ),
        ) as run_script_mock, contextlib.redirect_stdout(io.StringIO()):
            self.assertEqual(0, verify(args))
        self.assertEqual(2, run_script_mock.call_count)
        self.assertEqual("candidate_ready", state["status"])
        self.assertFalse(state["stages"]["verify"]["cache_hit"])

    def test_word_target_pdf_ignores_deleted_unrelated_renderer(self) -> None:
        work, state, _, args = self._cached_final_ready_run(
            "word-pdf-unrelated-renderer"
        )
        renderer = work / "unused-renderer"
        renderer.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        renderer.chmod(0o755)
        args.renderer = str(renderer)
        renderer.unlink()
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch("run_monograph.run_script") as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(0, verify(args))
            run_script_mock.assert_not_called()
        self.assertIsNone(state["verification_outputs"]["request"]["renderer"])
        self.assertTrue(state["stages"]["verify"]["cache_hit"])

    def test_used_renderer_entity_change_invalidates_status(self) -> None:
        work, state, _, verify_args = self._cached_no_fields_final_ready_run(
            "used-renderer-change"
        )
        renderer = work / "used-renderer"
        renderer.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        renderer.chmod(0o755)
        verify_args.renderer = str(renderer)
        request = verify_request_identity(
            verify_args, target_id=LIBREOFFICE, renderer_used=True
        )
        render_manifest = work / "rendered/render-manifest.json"
        manifest = json.loads(render_manifest.read_text(encoding="utf-8"))
        manifest.update(
            {
                "target_software": LIBREOFFICE,
                "renderer": request["renderer"]["resolved_path"],
                "renderer_source": request["renderer"]["source"],
            }
        )
        render_manifest.write_text(json.dumps(manifest), encoding="utf-8")
        state["verification_outputs"]["request"] = request
        state["verification_outputs"]["render_manifest"] = local_artifact_identity(
            render_manifest
        )
        state["stages"]["verify"]["input_key_sha256"] = json_sha256(
            {
                "finalized": file_sha256(work / "final/source-finalized.docx"),
                "map": file_sha256(work / "structure.json"),
                "target_pdf": None,
                "visual_manifest": {
                    "path": str((work / "visual.json").resolve()),
                    "sha256": file_sha256(work / "visual.json"),
                },
                "behavior": request,
            }
        )
        renderer.write_text("#!/bin/sh\nexit 1\n", encoding="utf-8")
        args = SimpleNamespace(work_dir=work, json=True)
        with patch("run_monograph.load_state", return_value=state), patch(
            "run_monograph.save_state"
        ), patch("run_monograph.run_script") as run_script_mock, contextlib.redirect_stdout(
            io.StringIO()
        ):
            self.assertEqual(2, orchestrate_status(args))
            run_script_mock.assert_not_called()
        self.assertEqual("candidate_ready", state["status"])

    def synthetic_book(self) -> Path:
        image = self.root / "pixel.png"
        image.write_bytes(PNG_1X1)
        path = self.root / "synthetic-book.docx"
        document = Document()
        document.add_paragraph("Synthetic title")
        document.add_paragraph("Chapter body")
        document.add_paragraph("APPENDIX A Load cases")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Column A"
        table.cell(0, 1).text = "Column B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        document.add_paragraph().add_run().add_picture(str(image), width=Inches(1))
        document.save(path)
        return path

    def test_candidate_map_is_private_bounded_and_preserves_appendices(self) -> None:
        source = self.synthetic_book()
        structure = candidate_structure_map(source)
        serialized = json.dumps(structure, ensure_ascii=False)
        self.assertEqual("1.5", structure["schema_version"])
        self.assertNotIn("APPENDIX A Load cases", serialized)
        self.assertEqual("preserve_existing", structure["appendices"][0]["numbering_mode"])
        self.assertIsNone(structure["appendices"][0]["include_in_toc"])
        self.assertFalse(structure["trial_selection"]["whole_book_candidate"])
        self.assertLessEqual(
            structure["trial_selection"]["max_rendered_pages_per_candidate"], 30
        )
        self.assertEqual([], structure["tables"][0]["header_rows"])
        self.assertEqual("center", structure["tables"][0]["visual"]["alignment"])
        self.assertEqual("none", structure["tables"][0]["visual"]["text_wrapping"])
        self.assertTrue(
            all(not image["resize"]["allow_upscale"] for image in structure["images"])
        )

    def test_large_synthetic_book_keeps_trial_samples_bounded(self) -> None:
        source = self.root / "large-synthetic.docx"
        document = Document()
        document.add_paragraph("Synthetic title")
        for chapter in range(1, 21):
            document.add_paragraph(f"第{chapter}章 Synthetic chapter {chapter}")
            for section in range(1, 11):
                document.add_paragraph(
                    f"{chapter}.{section} Synthetic section {chapter}-{section}"
                )
                document.add_paragraph("Synthetic body paragraph")
            table = document.add_table(rows=3, cols=3)
            for row in table.rows:
                for cell in row.cells:
                    cell.text = "Synthetic cell"
        for appendix in ("A", "B", "C"):
            document.add_paragraph(f"APPENDIX {appendix} Synthetic appendix")
        document.save(source)

        structure = candidate_structure_map(source)
        trial = structure["trial_selection"]
        self.assertGreaterEqual(len(structure["headings"]), 200)
        self.assertEqual(3, len(structure["appendices"]))
        self.assertLessEqual(len(trial["heading_samples"]), 8)
        self.assertLessEqual(len(trial["appendix_samples"]), 2)
        self.assertLessEqual(len(trial["table_samples"]), 2)
        self.assertFalse(trial["whole_book_candidate"])

    def test_schema_14_remains_readable(self) -> None:
        legacy = candidate_structure_map(self.synthetic_book())
        legacy["schema_version"] = "1.4"
        legacy["status"] = "approved"
        path = self.root / "legacy-map.json"
        path.write_text(json.dumps(legacy), encoding="utf-8")
        self.assertEqual("1.4", load_structure_map(path)["schema_version"])

    def test_prepare_resume_uses_cache_and_source_change_invalidates_it(self) -> None:
        source = self.synthetic_book()
        profile = SKILL / "examples" / "profiles" / "technical-textbook-layout.v0.2.5.draft.json"
        work = self.root / "run"
        command = [
            sys.executable,
            str(SCRIPTS / "run_monograph.py"),
            "prepare",
            str(source),
            "--profile",
            str(profile),
            "--work-dir",
            str(work),
        ]
        first = subprocess.run(command, capture_output=True, text=True, check=False)
        self.assertEqual(0, first.returncode, first.stderr)
        first_state = json.loads((work / "run-state.json").read_text(encoding="utf-8"))
        second = subprocess.run(
            [*command, "--resume"], capture_output=True, text=True, check=False
        )
        self.assertEqual(0, second.returncode, second.stderr)
        second_state = json.loads((work / "run-state.json").read_text(encoding="utf-8"))
        self.assertEqual(1, second_state["metrics"]["cache_hits"])
        self.assertEqual(first_state["run_id"], second_state["run_id"])

        document = Document(source)
        document.add_paragraph("Changed synthetic source")
        document.save(source)
        third = subprocess.run(
            [*command, "--resume"], capture_output=True, text=True, check=False
        )
        self.assertEqual(0, third.returncode, third.stderr)
        third_state = json.loads((work / "run-state.json").read_text(encoding="utf-8"))
        self.assertNotEqual(first_state["run_id"], third_state["run_id"])
        self.assertEqual(0, third_state["metrics"]["cache_hits"])

    def test_environment_exposes_portable_capability_contract(self) -> None:
        completed = subprocess.run(
            [sys.executable, str(SCRIPTS / "check_environment.py"), "--json"],
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(0, completed.returncode, completed.stderr)
        result = json.loads(completed.stdout)
        expected = {
            "file_read",
            "python_execution",
            "docx_inspection",
            "profile_validation",
            "docx_editing",
            "font_discovery",
            "rendering",
            "target_word",
            "field_update",
            "multimodal_source_reading",
        }
        self.assertEqual(expected, set(result["portable_capabilities"]))
        self.assertIsNone(
            result["portable_capabilities"]["multimodal_source_reading"]["available"]
        )

    def test_controlled_writeback_accepts_only_field_result_changes(self) -> None:
        image = self.root / "pixel.png"
        image.write_bytes(PNG_1X1)
        baseline = self.root / "baseline.docx"
        document = Document()
        document.add_paragraph("Authored content")
        document.add_picture(str(image))
        add_page_field(document.add_paragraph(), "1")
        document.save(baseline)
        refreshed = self.root / "refreshed.docx"

        def approved_transform(name: str, data: bytes) -> bytes:
            if name == "word/document.xml":
                root = etree.fromstring(data)
                result = root.xpath(
                    ".//w:fldChar[@w:fldCharType='separate']/following::w:t[1]",
                    namespaces=NS,
                )[0]
                result.text = "2"
                return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            return data

        rewrite_package(baseline, refreshed, approved_transform)
        output = self.root / "output.docx"
        replaced = controlled_field_result_writeback(baseline, refreshed, output)
        self.assertIn("word/document.xml", replaced)
        with zipfile.ZipFile(baseline) as before, zipfile.ZipFile(output) as after:
            media = next(name for name in before.namelist() if name.startswith("word/media/"))
            self.assertEqual(before.read(media), after.read(media))
            self.assertIn(b">2<", after.read("word/document.xml"))

        rejected = self.root / "rejected.docx"

        def rejected_transform(name: str, data: bytes) -> bytes:
            if name != "word/document.xml":
                return data
            root = etree.fromstring(data)
            root.xpath(".//w:t", namespaces=NS)[0].text = "Changed authored content"
            return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

        rewrite_package(baseline, refreshed, rejected_transform)
        with self.assertRaises(Exception):
            controlled_field_result_writeback(baseline, refreshed, rejected)
        self.assertFalse(rejected.exists())

        def protected_transform(name: str, data: bytes) -> bytes:
            if name.startswith("word/media/"):
                return b"target-application-mutated-media"
            return approved_transform(name, data)

        rewrite_package(baseline, refreshed, protected_transform)
        with self.assertRaises(Exception):
            controlled_field_result_writeback(
                baseline, refreshed, self.root / "protected-rejected.docx"
            )

    def test_controlled_writeback_accepts_multi_paragraph_toc_cache(self) -> None:
        baseline = self.root / "toc-baseline.docx"
        document = Document()
        add_toc_field(document.add_paragraph())
        document.add_paragraph("Authored body")
        document.save(baseline)
        refreshed = self.root / "toc-refreshed.docx"

        def expand_toc(name: str, data: bytes) -> bytes:
            if name != "word/document.xml":
                return data
            root = etree.fromstring(data)
            body = root.xpath("/w:document/w:body", namespaces=NS)[0]
            paragraph = body.xpath("./w:p[.//w:instrText]", namespaces=NS)[0]
            end_run = paragraph.xpath(
                ".//w:r[w:fldChar[@w:fldCharType='end']]", namespaces=NS
            )[0]
            paragraph.remove(end_run)
            result = paragraph.xpath(
                ".//w:fldChar[@w:fldCharType='separate']/following::w:t[1]",
                namespaces=NS,
            )[0]
            result.text = "Chapter 1"
            second = OxmlElement("w:p")
            second_run = OxmlElement("w:r")
            second_text = OxmlElement("w:t")
            second_text.text = "Section 1.1"
            second_run.append(second_text)
            second.append(second_run)
            second.append(end_run)
            body.insert(body.index(paragraph) + 1, second)
            return etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        rewrite_package(baseline, refreshed, expand_toc)
        output = self.root / "toc-output.docx"
        replaced = controlled_field_result_writeback(baseline, refreshed, output)
        self.assertIn("word/document.xml", replaced)
        reloaded = Document(output)
        self.assertEqual("Authored body", reloaded.paragraphs[-1].text)
        self.assertIn("Section 1.1", [paragraph.text for paragraph in reloaded.paragraphs])

    def test_final_ready_manifest_requires_every_page_and_no_issues(self) -> None:
        manifest = self.root / "visual.json"
        manifest.write_text(
            json.dumps(
                {
                    "all_pages_inspected": True,
                    "target_layout_verified": True,
                    "page_count": 12,
                    "issues": [],
                }
            ),
            encoding="utf-8",
        )
        self.assertEqual(12, validate_visual_manifest(manifest, 12)["page_count"])
        with self.assertRaises(Exception):
            validate_visual_manifest(manifest, 11)
        self.assertFalse(
            has_target_layout_evidence(
                {
                    "target_software": None,
                    "target_pdf_source": None,
                    "target_layout_unverified": False,
                }
            )
        )
        self.assertTrue(
            has_target_layout_evidence(
                {
                    "target_software": "LibreOffice Writer",
                    "target_pdf_source": None,
                    "target_layout_unverified": False,
                }
            )
        )
        self.assertTrue(
            has_target_layout_evidence(
                {
                    "target_software": "microsoft_word",
                    "target_pdf_source": "target.pdf",
                    "target_layout_unverified": False,
                }
            )
        )


if __name__ == "__main__":
    unittest.main()
