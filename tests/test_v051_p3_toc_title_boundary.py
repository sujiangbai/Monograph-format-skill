from __future__ import annotations

import base64
import json
import subprocess
import sys
import tempfile
import unittest
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "format-monograph"
SCRIPTS = SKILL / "scripts"
PROFILE = SKILL / "examples" / "profiles" / "v051-foundation-format-slice.json"
sys.path.insert(0, str(SCRIPTS))

from _common import (  # noqa: E402
    FormatMonographError,
    available_font_names,
    isolated_approved_style_name,
    protected_object_manifest,
    style_effective_font,
)
from structure_map import (  # noqa: E402
    DEFAULT_TOC_HEADING_FORMAT,
    STANDARD_TOC_HEADING_TEXT,
    _apply_front_matter,
    _body_locator,
    apply_structure_map,
    candidate_structure_map,
    load_structure_map,
    prime_structure_map_locators,
    structure_content_fingerprint,
)
from validate_profile import validate  # noqa: E402


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


class V051P3TocTitleBoundaryTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    @staticmethod
    def _add_toc_field(paragraph: object, result: str = "Synthetic entry 1") -> None:
        begin_run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin_run._r.append(begin)
        instruction_run = paragraph.add_run()
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = ' TOC \\o "1-3" \\h \\z '
        instruction_run._r.append(instruction)
        separate_run = paragraph.add_run()
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        separate_run._r.append(separate)
        paragraph.add_run(result)
        end_run = paragraph.add_run()
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        end_run._r.append(end)

    @staticmethod
    def _add_omml(paragraph: object) -> None:
        math = OxmlElement("m:oMath")
        run = OxmlElement("m:r")
        text = OxmlElement("m:t")
        text.text = "x+y=7"
        run.append(text)
        math.append(run)
        paragraph._p.append(math)

    def _source(
        self,
        toc_title: str | None,
        *,
        split_standard_runs: bool = False,
        extra_title: str | None = None,
        second_toc: bool = False,
        protected_payload: bool = False,
    ) -> tuple[Path, dict[str, object]]:
        source = self.root / (
            "toc-title-" + str(len(list(self.root.glob("toc-title-*.docx")))) + ".docx"
        )
        document = Document()
        document.add_paragraph("Synthetic whole-book title")
        if extra_title is not None:
            document.add_paragraph(extra_title)
        if toc_title is not None:
            heading = document.add_paragraph()
            if split_standard_runs and toc_title == "目录":
                first = heading.add_run("目")
                second = heading.add_run("录")
                first.font.italic = True
                first.font.size = Pt(9)
                second.font.bold = False
            else:
                run = heading.add_run(toc_title)
                run.font.italic = True
                run.font.size = Pt(9)
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            heading.paragraph_format.first_line_indent = Pt(12)
            heading.paragraph_format.left_indent = Pt(9)
            heading.paragraph_format.right_indent = Pt(6)
            heading.paragraph_format.space_before = Pt(11)
            heading.paragraph_format.space_after = Pt(3)
            heading.paragraph_format.page_break_before = True
        toc = document.add_paragraph()
        self._add_toc_field(toc)
        body = document.add_paragraph("第1章 Synthetic body", style="Heading 1")
        if second_toc:
            duplicate = document.add_paragraph()
            self._add_toc_field(duplicate, "Synthetic duplicate")
        if protected_payload:
            body.add_run().add_break(WD_BREAK.PAGE)
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Cell A"
            table.cell(0, 1).text = "Cell B"
            image_path = self.root / "synthetic-pixel.png"
            image_path.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
                )
            )
            document.add_paragraph().add_run().add_picture(str(image_path))
            formula = document.add_paragraph("Synthetic formula: ")
            self._add_omml(formula)
            ole = document.add_paragraph("Synthetic editable object anchor")
            ole._p.append(OxmlElement("w:object"))
            document.add_section(WD_SECTION.NEW_PAGE)
            document.add_paragraph("  Authored whitespace stays  ")
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        structure["front_matter"]["approved"] = True
        structure["front_matter"]["insert_toc_heading_if_missing"] = False
        structure["front_matter"]["toc_heading_text"] = STANDARD_TOC_HEADING_TEXT
        return source, structure

    def _approve_existing_heading(
        self, source: Path, structure: dict[str, object], index: int
    ) -> None:
        values = [paragraph.text for paragraph in Document(source).paragraphs]
        structure["front_matter"]["toc_heading"] = _body_locator(index, values)
        structure["front_matter"][
            "toc_heading_action"
        ] = "normalize_and_format_standard"

    @staticmethod
    def _xml_contract(path: Path) -> dict[str, object]:
        with zipfile.ZipFile(path) as package:
            root = etree.fromstring(package.read("word/document.xml"))
            return {
                "paragraph_text": [
                    "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))
                    for paragraph in root.xpath("/w:document/w:body/w:p", namespaces=NS)
                ],
                "tables": [
                    etree.tostring(table, method="c14n", exclusive=True)
                    for table in root.xpath("/w:document/w:body/w:tbl", namespaces=NS)
                ],
                "toc_fields": [
                    etree.tostring(paragraph, method="c14n", exclusive=True)
                    for paragraph in root.xpath(
                        "/w:document/w:body/w:p[.//w:instrText[contains(translate(., 'toc', 'TOC'), 'TOC')]]",
                        namespaces=NS,
                    )
                ],
                "sections": [
                    etree.tostring(section, method="c14n", exclusive=True)
                    for section in root.xpath(".//w:sectPr", namespaces=NS)
                ],
                "omml": [
                    etree.tostring(math, method="c14n", exclusive=True)
                    for math in root.xpath(".//m:oMath", namespaces={**NS, "m": "http://schemas.openxmlformats.org/officeDocument/2006/math"})
                ],
                "objects": [
                    etree.tostring(item, method="c14n", exclusive=True)
                    for item in root.xpath(".//w:object", namespaces=NS)
                ],
                "relationships": package.read("word/_rels/document.xml.rels"),
                "media": {
                    name: package.read(name)
                    for name in package.namelist()
                    if name.startswith("word/media/")
                },
            }

    def test_candidate_defaults_and_profile_define_preservation_only_contract(self) -> None:
        source, structure = self._source("目录")
        front = structure["front_matter"]
        self.assertFalse(front["insert_toc_heading_if_missing"])
        self.assertEqual("preserve", front["toc_heading_action"])
        self.assertEqual(STANDARD_TOC_HEADING_TEXT, front["toc_heading_text"])
        self.assertEqual(1, front["toc_heading"]["paragraph"])

        front["book_title_format"].pop("left_indent_pt", None)
        front["book_title_format"].pop("right_indent_pt", None)
        path = self.root / "structure.json"
        path.write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
        self.assertEqual("preserve", load_structure_map(path)["front_matter"]["toc_heading_action"])
        structure["front_matter"]["insert_toc_heading_if_missing"] = True
        path.write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
        with self.assertRaisesRegex(FormatMonographError, "does not authorize creation"):
            load_structure_map(path)

        legacy = deepcopy(structure)
        legacy["schema_version"] = "1.4"
        path.write_text(json.dumps(legacy, ensure_ascii=False), encoding="utf-8")
        loaded_legacy = load_structure_map(path)
        document = Document(source)
        prime_structure_map_locators(document, loaded_legacy)
        before_document = etree.tostring(
            deepcopy(document.element), method="c14n", exclusive=True
        )
        before_styles = etree.tostring(
            deepcopy(document.styles.element), method="c14n", exclusive=True
        )
        with self.assertRaisesRegex(
            FormatMonographError, "creation_not_authorized"
        ):
            apply_structure_map(document, loaded_legacy)
        self.assertEqual(
            before_document,
            etree.tostring(document.element, method="c14n", exclusive=True),
        )
        self.assertEqual(
            before_styles,
            etree.tostring(document.styles.element, method="c14n", exclusive=True),
        )

        errors, profile = validate(PROFILE)
        self.assertEqual([], errors)
        rule = next(item for item in profile["rules"] if item["id"] == "FMT-TOCTITLE-501")
        self.assertEqual(DEFAULT_TOC_HEADING_FORMAT, rule["properties"])
        self.assertEqual(
            {"kind": "paragraph_role", "value": "toc_heading"},
            rule["selector"],
        )

    def test_standard_titles_format_and_repeat_without_structural_change(self) -> None:
        for original in ("目录", STANDARD_TOC_HEADING_TEXT):
            with self.subTest(original=original):
                source, structure = self._source(
                    original, split_standard_runs=original == "目录"
                )
                self._approve_existing_heading(source, structure, 1)
                document = Document(source)
                prime_structure_map_locators(document, structure)
                changes = apply_structure_map(document, structure)
                heading = document.paragraphs[1]
                self.assertEqual(STANDARD_TOC_HEADING_TEXT, heading.text)
                self.assertEqual(4, heading.text.count(" "))
                self.assertNotIn("\u00a0", heading.text)
                self.assertTrue(heading.style.name.startswith("Monograph Approved "))
                style = heading.style
                self.assertEqual(18, style.font.size.pt)
                self.assertIs(True, style.font.bold)
                self.assertIs(False, style.font.italic)
                self.assertEqual("000000", str(style.font.color.rgb))
                self.assertEqual(WD_ALIGN_PARAGRAPH.CENTER, style.paragraph_format.alignment)
                self.assertEqual(18, style.paragraph_format.space_after.pt)
                self.assertEqual(0, style.paragraph_format.space_before.pt)
                self.assertIs(True, style.paragraph_format.keep_with_next)
                self.assertIs(True, heading.paragraph_format.page_break_before)
                self.assertEqual("Times New Roman", style_effective_font(document, style, "ascii")[0])
                self.assertEqual("STHeiti Medium", style_effective_font(document, style, "eastAsia")[0])
                self.assertTrue(all(run.italic is None for run in heading.runs))
                evidence = next(
                    change["details"]["toc_heading"]
                    for change in changes
                    if change["kind"] == "structure_front_matter"
                )
                self.assertEqual("formatted_standard", evidence["status"])
                self.assertEqual(original == "目录", evidence["text_normalized"])

                first = self.root / f"first-{len(list(self.root.glob('first-*.docx')))}.docx"
                document.save(first)
                repeated = Document(first)
                prime_structure_map_locators(repeated, structure)
                second_changes = apply_structure_map(repeated, structure)
                second = self.root / f"second-{len(list(self.root.glob('second-*.docx')))}.docx"
                repeated.save(second)
                self.assertEqual(
                    self._xml_contract(first),
                    self._xml_contract(second),
                )
                second_evidence = next(
                    change["details"]["toc_heading"]
                    for change in second_changes
                    if change["kind"] == "structure_front_matter"
                )
                self.assertFalse(second_evidence["text_normalized"])

    def test_apply_profile_end_to_end_accepts_only_authorized_text_delta(self) -> None:
        portable_font = next(iter(sorted(available_font_names())), None)
        if portable_font is None:
            self.skipTest("No host font is available for strict integration testing.")
        source, structure = self._source("目录", split_standard_runs=True)
        self._approve_existing_heading(source, structure, 1)
        structure["front_matter"]["book_title_format"].pop("left_indent_pt", None)
        structure["front_matter"]["book_title_format"].pop("right_indent_pt", None)
        structure_path = self.root / "approved-structure.json"
        structure_path.write_text(
            json.dumps(structure, ensure_ascii=False), encoding="utf-8"
        )
        profile = json.loads(PROFILE.read_text(encoding="utf-8"))
        profile["profile_id"] = "v051-toc-title-portable-test"
        for rule in profile["rules"]:
            for key in (
                "font_name",
                "font_name_ascii",
                "font_name_east_asia",
                "font_name_complex_script",
            ):
                if key in rule.get("properties", {}):
                    rule["properties"][key] = portable_font
        profile_path = self.root / "portable-profile.json"
        profile_path.write_text(
            json.dumps(profile, ensure_ascii=False), encoding="utf-8"
        )
        output_dir = self.root / "output"
        result = subprocess.run(
            [
                sys.executable,
                str(SCRIPTS / "apply_profile.py"),
                str(source),
                "--profile",
                str(profile_path),
                "--structure-map",
                str(structure_path),
                "--output-dir",
                str(output_dir),
            ],
            cwd=SKILL,
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(0, result.returncode, result.stdout + result.stderr)
        output = output_dir / f"{source.stem}-formatted.docx"
        formatted = Document(output)
        self.assertEqual(STANDARD_TOC_HEADING_TEXT, formatted.paragraphs[1].text)
        self.assertTrue(
            formatted.paragraphs[1].style.name.startswith("Monograph Approved ")
        )
        self.assertEqual(
            portable_font,
            style_effective_font(
                formatted, formatted.paragraphs[1].style, "eastAsia"
            )[0],
        )
        self.assertEqual(
            structure_content_fingerprint(source, structure),
            structure_content_fingerprint(output, structure),
        )
        audit = subprocess.run(
            [
                sys.executable,
                str(SCRIPTS / "audit_docx.py"),
                str(source),
                str(output),
                "--profile",
                str(profile_path),
                "--structure-map",
                str(structure_path),
            ],
            cwd=SKILL,
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(0, audit.returncode, audit.stdout + audit.stderr)
        self.assertTrue(json.loads(audit.stdout)["passed"])

    def test_custom_missing_and_unapproved_titles_are_preserved(self) -> None:
        custom = "Synthetic Contents — Keep  exactly"
        source, structure = self._source(custom)
        self._approve_existing_heading(source, structure, 1)
        document = Document(source)
        before = etree.tostring(deepcopy(document.paragraphs[1]._p), method="c14n", exclusive=True)
        prime_structure_map_locators(document, structure)
        changes = apply_structure_map(document, structure)
        self.assertEqual(before, etree.tostring(document.paragraphs[1]._p, method="c14n", exclusive=True))
        evidence = next(
            change["details"]["toc_heading"]
            for change in changes
            if change["kind"] == "structure_front_matter"
        )
        self.assertEqual("preserved_custom", evidence["status"])

        missing_source, missing_structure = self._source(None)
        missing_document = Document(missing_source)
        before_count = len(missing_document.paragraphs)
        prime_structure_map_locators(missing_document, missing_structure)
        apply_structure_map(missing_document, missing_structure)
        self.assertEqual(before_count, len(missing_document.paragraphs))
        self.assertNotIn(STANDARD_TOC_HEADING_TEXT, [p.text for p in missing_document.paragraphs])

        unapproved_source, unapproved_structure = self._source("目录")
        unapproved_document = Document(unapproved_source)
        before = etree.tostring(deepcopy(unapproved_document.paragraphs[1]._p), method="c14n", exclusive=True)
        prime_structure_map_locators(unapproved_document, unapproved_structure)
        apply_structure_map(unapproved_document, unapproved_structure)
        self.assertEqual(before, etree.tostring(unapproved_document.paragraphs[1]._p, method="c14n", exclusive=True))

    def test_duplicate_title_or_toc_field_blocks_before_any_change(self) -> None:
        cases = (
            {"extra_title": "目录", "toc_title": STANDARD_TOC_HEADING_TEXT, "second_toc": False},
            {"extra_title": None, "toc_title": "目录", "second_toc": True},
        )
        for case in cases:
            with self.subTest(case=case):
                source, structure = self._source(**case)
                heading_index = 2 if case["extra_title"] is not None else 1
                self._approve_existing_heading(source, structure, heading_index)
                document = Document(source)
                prime_structure_map_locators(document, structure)
                before_document = etree.tostring(deepcopy(document.element), method="c14n", exclusive=True)
                before_styles = etree.tostring(deepcopy(document.styles.element), method="c14n", exclusive=True)
                with self.assertRaisesRegex(
                    FormatMonographError,
                    "TOC_TITLE_BOUNDARY_BLOCKED",
                ) as caught:
                    apply_structure_map(document, structure)
                self.assertNotIn("目录", str(caught.exception))
                self.assertEqual(before_document, etree.tostring(document.element, method="c14n", exclusive=True))
                self.assertEqual(before_styles, etree.tostring(document.styles.element, method="c14n", exclusive=True))

    def test_isolated_style_collisions_block_all_writes_from_both_entries(self) -> None:
        selector = {"kind": "paragraph_role", "value": "toc_heading"}
        derived_name = isolated_approved_style_name(selector)
        self.assertIsNotNone(derived_name)
        for collision in ("wrong_type", "wrong_base", "occupied"):
            for entry in ("apply_structure_map", "_apply_front_matter"):
                with self.subTest(collision=collision, entry=entry):
                    source, structure = self._source("目录")
                    self._approve_existing_heading(source, structure, 1)
                    document = Document(source)
                    if collision == "wrong_type":
                        document.styles.add_style(
                            derived_name, WD_STYLE_TYPE.CHARACTER
                        )
                    else:
                        expected_base = document.styles.add_style(
                            "Monograph TOC Heading", WD_STYLE_TYPE.PARAGRAPH
                        )
                        derived = document.styles.add_style(
                            derived_name, WD_STYLE_TYPE.PARAGRAPH
                        )
                        derived.base_style = (
                            document.styles["Normal"]
                            if collision == "wrong_base"
                            else expected_base
                        )
                        document.paragraphs[1].style = derived
                        if collision == "occupied":
                            document.add_paragraph(
                                "Synthetic unapproved style occupant",
                                style=derived,
                            )
                    prime_structure_map_locators(document, structure)
                    before_document = etree.tostring(
                        deepcopy(document.element), method="c14n", exclusive=True
                    )
                    before_styles = etree.tostring(
                        deepcopy(document.styles.element),
                        method="c14n",
                        exclusive=True,
                    )
                    operation = (
                        apply_structure_map
                        if entry == "apply_structure_map"
                        else _apply_front_matter
                    )
                    with self.assertRaisesRegex(
                        FormatMonographError, "isolated approved-role style"
                    ):
                        operation(document, structure)
                    self.assertEqual(
                        before_document,
                        etree.tostring(
                            document.element, method="c14n", exclusive=True
                        ),
                    )
                    self.assertEqual(
                        before_styles,
                        etree.tostring(
                            document.styles.element,
                            method="c14n",
                            exclusive=True,
                        ),
                    )

    def test_dangling_header_footer_style_references_block_all_writes(self) -> None:
        selector = {"kind": "paragraph_role", "value": "toc_heading"}
        derived_name = isolated_approved_style_name(selector)
        self.assertIsNotNone(derived_name)
        derived_style_id = derived_name.replace(" ", "")
        for story in ("header", "footer"):
            for entry in ("apply_structure_map", "_apply_front_matter"):
                with self.subTest(story=story, entry=entry):
                    source, structure = self._source("目录")
                    self._approve_existing_heading(source, structure, 1)
                    fixture = Document(source)
                    header = fixture.sections[0].header.paragraphs[0]
                    footer = fixture.sections[0].footer.paragraphs[0]
                    header.text = "Synthetic existing header"
                    footer.text = "Synthetic existing footer"
                    target = header if story == "header" else footer
                    p_style = OxmlElement("w:pStyle")
                    p_style.set(qn("w:val"), derived_style_id)
                    target._p.get_or_add_pPr().insert(0, p_style)
                    fixture.save(source)

                    document = Document(source)
                    with self.assertRaises(KeyError):
                        document.styles[derived_name]
                    prime_structure_map_locators(document, structure)
                    before_parts = tuple(
                        sorted(str(part.partname) for part in document.part.package.parts)
                    )
                    before_document = etree.tostring(
                        deepcopy(document.element), method="c14n", exclusive=True
                    )
                    before_styles = etree.tostring(
                        deepcopy(document.styles.element),
                        method="c14n",
                        exclusive=True,
                    )
                    before_stories = {
                        str(part.partname): part.blob
                        for part in document.part.package.parts
                        if str(part.partname).startswith(
                            ("/word/header", "/word/footer")
                        )
                    }
                    operation = (
                        apply_structure_map
                        if entry == "apply_structure_map"
                        else _apply_front_matter
                    )
                    with self.assertRaisesRegex(
                        FormatMonographError,
                        "referenced by an unapproved paragraph",
                    ):
                        operation(document, structure)
                    self.assertEqual(
                        before_parts,
                        tuple(
                            sorted(
                                str(part.partname)
                                for part in document.part.package.parts
                            )
                        ),
                    )
                    self.assertEqual(
                        before_document,
                        etree.tostring(
                            document.element, method="c14n", exclusive=True
                        ),
                    )
                    self.assertEqual(
                        before_styles,
                        etree.tostring(
                            document.styles.element,
                            method="c14n",
                            exclusive=True,
                        ),
                    )
                    self.assertEqual(
                        before_stories,
                        {
                            str(part.partname): part.blob
                            for part in document.part.package.parts
                            if str(part.partname).startswith(
                                ("/word/header", "/word/footer")
                            )
                        },
                    )

    def test_four_space_exception_preserves_all_other_payload_and_fingerprint(self) -> None:
        source, structure = self._source(
            "目录", split_standard_runs=True, protected_payload=True
        )
        self._approve_existing_heading(source, structure, 1)
        before = self._xml_contract(source)
        before_objects = protected_object_manifest(source)
        before_fingerprint = structure_content_fingerprint(source, structure)
        document = Document(source)
        prime_structure_map_locators(document, structure)
        apply_structure_map(document, structure)
        output = self.root / "integrity-output.docx"
        document.save(output)
        after = self._xml_contract(output)
        self.assertEqual(before["paragraph_text"][:1], after["paragraph_text"][:1])
        self.assertEqual(STANDARD_TOC_HEADING_TEXT, after["paragraph_text"][1])
        self.assertEqual(before["paragraph_text"][2:], after["paragraph_text"][2:])
        for key in ("tables", "toc_fields", "sections", "omml", "objects", "relationships", "media"):
            self.assertEqual(before[key], after[key], key)
        self.assertEqual(before_objects, protected_object_manifest(output))
        self.assertEqual(before_fingerprint, structure_content_fingerprint(output, structure))

        relocated_identity = deepcopy(structure)
        relocated_identity["front_matter"]["toc_heading"]["paragraph"] = 0
        self.assertEqual(
            structure_content_fingerprint(source, relocated_identity),
            structure_content_fingerprint(output, relocated_identity),
        )

        wrong_identity = deepcopy(structure)
        values = [paragraph.text for paragraph in Document(source).paragraphs]
        wrong_identity["front_matter"]["toc_heading"] = _body_locator(0, values)
        with self.assertRaisesRegex(
            FormatMonographError, "heading_not_immediately_before_toc"
        ):
            structure_content_fingerprint(source, wrong_identity)
        with self.assertRaisesRegex(
            FormatMonographError, "heading_not_immediately_before_toc"
        ):
            structure_content_fingerprint(output, wrong_identity)


if __name__ == "__main__":
    unittest.main()
