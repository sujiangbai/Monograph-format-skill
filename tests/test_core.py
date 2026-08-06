from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
SKILL = REPO / "format-monograph"
SCRIPTS = SKILL / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import content_fingerprint  # noqa: E402
from validate_profile import validate  # noqa: E402


def profile() -> dict:
    return {
        "schema_version": "1.0",
        "profile_id": "synthetic-monograph-v1",
        "name": "Synthetic monograph profile",
        "locale": "zh-CN",
        "scope": {"document_type": "monograph", "input_format": "DOCX"},
        "target_applications": ["Microsoft 365"],
        "source_precedence": ["written_requirement", "docx_template"],
        "sources": [
            {
                "id": "SRC-001",
                "type": "written_requirement",
                "label": "Synthetic requirements",
                "summary": "Generated requirements for automated tests.",
                "public": True,
            }
        ],
        "rules": [
            {
                "id": "FMT-PAGE-001",
                "category": "page",
                "selector": {"kind": "document", "value": "all"},
                "properties": {
                    "margin_top_mm": 25.0,
                    "margin_bottom_mm": 25.0,
                    "margin_left_mm": 30.0,
                    "margin_right_mm": 25.0,
                },
                "source_ids": ["SRC-001"],
                "evidence_summary": "Synthetic page geometry.",
                "confidence": "high",
                "status": "approved",
                "application": "automatic",
            },
            {
                "id": "FMT-HEAD-001",
                "category": "heading",
                "selector": {"kind": "paragraph_role", "value": "heading1"},
                "properties": {
                    "font_name": "Arial",
                    "font_size_pt": 16,
                    "bold": True,
                    "space_before_pt": 12,
                    "space_after_pt": 6,
                    "keep_with_next": True,
                },
                "source_ids": ["SRC-001"],
                "evidence_summary": "Synthetic heading rule.",
                "confidence": "high",
                "status": "approved",
                "application": "automatic",
            },
            {
                "id": "FMT-BODY-001",
                "category": "body",
                "selector": {"kind": "paragraph_role", "value": "body"},
                "properties": {
                    "font_name": "Arial",
                    "font_size_pt": 11,
                    "alignment": "justify",
                    "line_spacing": 1.5,
                    "first_line_indent_pt": 22,
                },
                "source_ids": ["SRC-001"],
                "evidence_summary": "Synthetic body rule.",
                "confidence": "high",
                "status": "approved",
                "application": "automatic",
            },
            {
                "id": "FMT-TABLE-001",
                "category": "table",
                "selector": {"kind": "table_role", "value": "all"},
                "properties": {
                    "table_style": "Table Grid",
                    "alignment": "center",
                    "repeat_header_row": True,
                },
                "source_ids": ["SRC-001"],
                "evidence_summary": "Synthetic table rule.",
                "confidence": "high",
                "status": "approved",
                "application": "automatic",
            },
            {
                "id": "FMT-NOTE-001",
                "category": "note",
                "selector": {"kind": "note_role", "value": "all"},
                "properties": {"note": "Review footnote separator manually."},
                "source_ids": ["SRC-001"],
                "evidence_summary": "Synthetic manual review rule.",
                "confidence": "medium",
                "status": "approved",
                "application": "manual_review",
            },
        ],
        "conflicts": [],
        "open_questions": [],
        "approval": {
            "status": "approved",
            "approved_by": "test-suite",
            "approved_at": "2026-08-06T00:00:00Z",
        },
    }


class CoreWorkflowTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.input = self.root / "sample.docx"
        document = Document()
        document.add_heading("Synthetic Monograph", level=0)
        document.add_heading("Chapter One", level=1)
        document.add_paragraph("Authored body text must remain unchanged.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Value A"
        table.cell(1, 1).text = "Value B"
        section = document.sections[0]
        section.header.paragraphs[0].text = "Synthetic header"
        section.footer.paragraphs[0].text = "Synthetic footer"
        document.save(self.input)
        self.original_bytes = self.input.read_bytes()
        self.profile = self.root / "profile.json"
        self.profile.write_text(
            json.dumps(profile(), ensure_ascii=False, indent=2), encoding="utf-8"
        )
        self.output = self.root / "out"

    def tearDown(self) -> None:
        self.temp.cleanup()

    def run_script(self, name: str, *args: object) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(SCRIPTS / name), *(str(arg) for arg in args)],
            cwd=SKILL,
            capture_output=True,
            text=True,
            check=False,
        )

    def test_profile_validates(self) -> None:
        errors, parsed = validate(self.profile)
        self.assertEqual([], errors)
        self.assertEqual("approved", parsed["approval"]["status"])

    def test_invalid_automatic_property_is_rejected(self) -> None:
        invalid = profile()
        invalid["rules"][0]["properties"]["unknown_property"] = 1
        path = self.root / "invalid.json"
        path.write_text(json.dumps(invalid), encoding="utf-8")
        errors, _ = validate(path)
        self.assertTrue(any("unsupported properties" in error for error in errors))

    def test_end_to_end_structural_workflow(self) -> None:
        inventory = self.root / "inventory.json"
        inspected = self.run_script(
            "inspect_docx.py", self.input, "--output", inventory
        )
        self.assertEqual(0, inspected.returncode, inspected.stderr)
        self.assertTrue(inventory.exists())

        applied = self.run_script(
            "apply_profile.py",
            self.input,
            "--profile",
            self.profile,
            "--output-dir",
            self.output,
        )
        self.assertEqual(0, applied.returncode, applied.stderr)
        formatted = self.output / "sample-formatted.docx"
        review = self.output / "sample-review.docx"
        report = self.output / "sample-format-report.md"
        self.assertTrue(formatted.exists())
        self.assertTrue(review.exists())
        self.assertTrue(report.exists())
        self.assertEqual(self.original_bytes, self.input.read_bytes())
        self.assertEqual(content_fingerprint(self.input), content_fingerprint(formatted))
        self.assertEqual(content_fingerprint(self.input), content_fingerprint(review))

        formatted_doc = Document(formatted)
        self.assertAlmostEqual(16.0, formatted_doc.styles["Heading 1"].font.size.pt)
        self.assertEqual("Arial", formatted_doc.styles["Heading 1"].font.name)
        self.assertTrue(formatted_doc.styles["Heading 1"].font.bold)
        self.assertEqual("Table Grid", formatted_doc.tables[0].style.name)

        review_doc = Document(review)
        self.assertGreaterEqual(len(review_doc.comments), 4)

        audit = self.root / "audit.json"
        audited = self.run_script(
            "audit_docx.py",
            self.input,
            formatted,
            "--profile",
            self.profile,
            "--output",
            audit,
        )
        self.assertEqual(0, audited.returncode, audited.stderr)
        self.assertTrue(json.loads(audit.read_text(encoding="utf-8"))["passed"])

    def test_environment_command_is_machine_readable(self) -> None:
        result = self.run_script("check_environment.py", "--json")
        self.assertEqual(0, result.returncode, result.stderr)
        parsed = json.loads(result.stdout)
        self.assertIn(parsed["mode"], {"full", "structural", "analysis"})


if __name__ == "__main__":
    unittest.main()
