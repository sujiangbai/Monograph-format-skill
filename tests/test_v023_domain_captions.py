from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "format-monograph"
SCRIPTS = SKILL / "scripts"
sys.path.insert(0, str(SCRIPTS))

from structure_map import candidate_structure_map, load_structure_map  # noqa: E402
from test_v11_execution import approved_v11_profile  # noqa: E402


def caption_profile(*, seq: bool = False) -> dict:
    profile = approved_v11_profile()
    profile["rules"] = [
        {
            "id": "FMT-CAP-923",
            "category": "caption",
            "selector": {"kind": "caption_role", "value": "all"},
            "properties": {
                "font_size_pt": 9,
                "alignment": "center",
                "numbering_mode": "seq_field" if seq else "manual_text",
                "preserve_identifier": not seq,
                "domain_context": "unknown",
                "allow_automatic_renumbering": seq,
                "preserve_table_cell_caption_position": True,
            },
            "source_ids": ["SRC-901"],
            "evidence_summary": "Synthetic caller-approved caption behavior.",
            "confidence": "high",
            "status": "approved",
            "application": "automatic",
        }
    ]
    return profile


class V023DomainCaptionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def run_script(self, name: str, *args: object) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(SCRIPTS / name), *(str(value) for value in args)],
            cwd=SKILL,
            capture_output=True,
            text=True,
            check=False,
        )

    def write_profile(self, profile: dict, name: str = "profile.json") -> Path:
        path = self.root / name
        path.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
        return path

    def approve_caption_roles(self, structure: dict) -> None:
        for entry in structure["paragraph_roles"]:
            if entry["role"] in {"figure_caption", "table_caption"}:
                entry["approved"] = True

    def test_candidate_map_is_domain_aware_and_does_not_infer_missing_number(self) -> None:
        source = self.root / "domain-captions.docx"
        document = Document()
        document.add_paragraph("第4章 Synthetic steel structure")
        document.add_paragraph("结构节点与剖面图 synthetic context")
        document.add_paragraph("图 4.2.1-1 1-1剖面图")
        document.add_paragraph("表 4.1-Synthetic title")
        document.save(source)

        structure = candidate_structure_map(source)
        self.assertEqual("1.3", structure["schema_version"])
        self.assertEqual(2, len(structure["captions"]))
        figure, table = structure["captions"]
        self.assertEqual("preserve", figure["action"])
        self.assertEqual("manual_text", figure["numbering_mode"])
        self.assertEqual("mixed", figure["identifier_semantics"])
        self.assertIn(figure["domain_context"], {"mixed", "civil_engineering"})
        self.assertEqual("publication_number", table["identifier_semantics"])
        self.assertEqual("candidate", table["completeness"])
        serialized = json.dumps(structure, ensure_ascii=False)
        self.assertNotIn("Synthetic title", serialized)
        self.assertNotIn("剖面图", serialized)

    def test_style_only_preserves_manual_text_and_table_caption_row(self) -> None:
        source = self.root / "style-only.docx"
        document = Document()
        document.add_paragraph("图 4.2.1-1 1-1剖面图")
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).merge(table.cell(0, 1)).text = "表 4.1-Synthetic table"
        table.cell(1, 0).text = "Header A"
        table.cell(1, 1).text = "Header B"
        table.cell(2, 0).text = "Value A"
        table.cell(2, 1).text = "Value B"
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        self.approve_caption_roles(structure)
        for entry in structure["captions"]:
            entry.update({"approved": True, "action": "style_only"})
        structure_path = self.root / "style-only-map.json"
        structure_path.write_text(
            json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        profile_path = self.write_profile(caption_profile())
        output = self.root / "style-only-out"
        result = self.run_script(
            "apply_profile.py",
            source,
            "--profile",
            profile_path,
            "--structure-map",
            structure_path,
            "--output-dir",
            output,
            "--allow-missing-fonts",
        )
        self.assertEqual(0, result.returncode, result.stdout + result.stderr)

        formatted = output / "style-only-formatted.docx"
        after = Document(formatted)
        self.assertEqual("图 4.2.1-1 1-1剖面图", after.paragraphs[0].text)
        self.assertEqual(3, len(after.tables[0].rows))
        self.assertEqual("表 4.1-Synthetic table", after.tables[0].cell(0, 0).text)
        self.assertEqual("Caption", after.tables[0].cell(0, 0).paragraphs[0].style.name)
        with zipfile.ZipFile(formatted) as package:
            self.assertNotIn(b"SEQ Figure", package.read("word/document.xml"))
            self.assertNotIn(b"SEQ Table", package.read("word/document.xml"))

    def test_confirmed_identifier_replacement_preserves_title_and_audits(self) -> None:
        source = self.root / "replace.docx"
        document = Document()
        document.add_paragraph("图 3.2-1 Synthetic title")
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        self.approve_caption_roles(structure)
        structure["captions"][0].update(
            {
                "approved": True,
                "action": "replace_identifier",
                "replacement_identifier": "4.2-1",
                "replacement_confirmed": True,
            }
        )
        structure_path = self.root / "replace-map.json"
        structure_path.write_text(
            json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        profile_path = self.write_profile(caption_profile())
        output = self.root / "replace-out"
        applied = self.run_script(
            "apply_profile.py",
            source,
            "--profile",
            profile_path,
            "--structure-map",
            structure_path,
            "--output-dir",
            output,
            "--allow-missing-fonts",
        )
        self.assertEqual(0, applied.returncode, applied.stdout + applied.stderr)
        formatted = output / "replace-formatted.docx"
        self.assertEqual("图 4.2-1 Synthetic title", Document(formatted).paragraphs[0].text)

        audit = self.run_script(
            "audit_docx.py",
            source,
            formatted,
            "--profile",
            profile_path,
            "--structure-map",
            structure_path,
        )
        self.assertEqual(0, audit.returncode, audit.stdout + audit.stderr)
        payload = json.loads(audit.stdout)
        self.assertEqual(
            "pass",
            payload["approved_manual_identifier_replacements"][0]["status"],
        )
        self.assertTrue(
            payload["approved_manual_identifier_replacements"][0]["title_preserved"]
        )

        damaged = self.root / "damaged.docx"
        changed = Document(formatted)
        changed.paragraphs[0].runs[-1].text += " changed"
        changed.save(damaged)
        failed = self.run_script(
            "audit_docx.py",
            source,
            damaged,
            "--profile",
            profile_path,
            "--structure-map",
            structure_path,
        )
        self.assertEqual(1, failed.returncode)

    def test_caption_move_is_separate_and_preserves_manual_identifier(self) -> None:
        source = self.root / "move.docx"
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).merge(table.cell(0, 1)).text = "表 2.1-Synthetic table"
        table.cell(1, 0).text = "Header A"
        table.cell(1, 1).text = "Header B"
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        self.approve_caption_roles(structure)
        structure["captions"][0].update(
            {
                "approved": True,
                "action": "move_caption",
                "migrate_outside_table": True,
            }
        )
        structure_path = self.root / "move-map.json"
        structure_path.write_text(
            json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        output = self.root / "move-out"
        applied = self.run_script(
            "apply_profile.py",
            source,
            "--profile",
            self.write_profile(caption_profile(), "move-profile.json"),
            "--structure-map",
            structure_path,
            "--output-dir",
            output,
            "--allow-missing-fonts",
        )
        self.assertEqual(0, applied.returncode, applied.stdout + applied.stderr)
        formatted = Document(output / "move-formatted.docx")
        self.assertEqual("表 2.1-Synthetic table", formatted.paragraphs[0].text)
        self.assertEqual(1, len(formatted.tables[0].rows))

    def test_seq_conversion_requires_profile_and_map_approval(self) -> None:
        source = self.root / "seq.docx"
        document = Document()
        document.add_paragraph("图 1.1-1 Synthetic caption")
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        structure["captions"][0].update(
            {
                "approved": True,
                "action": "convert_to_seq",
                "numbering_mode": "seq_field",
                "hierarchy_status": "accepted",
            }
        )
        structure_path = self.root / "seq-map.json"
        structure_path.write_text(
            json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        denied = self.run_script(
            "apply_profile.py",
            source,
            "--profile",
            self.write_profile(caption_profile(), "manual-profile.json"),
            "--structure-map",
            structure_path,
            "--output-dir",
            self.root / "denied",
            "--allow-missing-fonts",
        )
        self.assertEqual(1, denied.returncode)
        self.assertIn("SEQ caption conversion requires", denied.stderr)

        allowed_output = self.root / "allowed"
        allowed = self.run_script(
            "apply_profile.py",
            source,
            "--profile",
            self.write_profile(caption_profile(seq=True), "seq-profile.json"),
            "--structure-map",
            structure_path,
            "--output-dir",
            allowed_output,
            "--allow-missing-fonts",
        )
        self.assertEqual(0, allowed.returncode, allowed.stdout + allowed.stderr)
        with zipfile.ZipFile(allowed_output / "seq-formatted.docx") as package:
            self.assertIn(b"SEQ Figure", package.read("word/document.xml"))

    def test_legacy_structure_maps_remain_readable(self) -> None:
        for version in ("1.0", "1.1", "1.2"):
            value = {
                "schema_version": version,
                "status": "approved",
                "source_content_fingerprint_sha256": "0" * 64,
                "paragraph_roles": [],
                "numbering": {},
                "toc_ranges": [],
                "headings": [],
                "captions": [],
                "tables": [],
                "trailing_empty_sections": [],
                "conflicts": [],
            }
            path = self.root / f"legacy-{version}.json"
            path.write_text(json.dumps(value), encoding="utf-8")
            self.assertEqual(version, load_structure_map(path)["schema_version"])


if __name__ == "__main__":
    unittest.main()
