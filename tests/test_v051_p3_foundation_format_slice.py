from __future__ import annotations

import json
import base64
import hashlib
import subprocess
import sys
import tempfile
import unittest
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "format-monograph"
SCRIPTS = SKILL / "scripts"
PROFILE = SKILL / "examples" / "profiles" / "v051-foundation-format-slice.json"
sys.path.insert(0, str(SCRIPTS))

from _common import (  # noqa: E402
    FormatMonographError,
    apply_style_rule_to_paragraphs,
    available_font_names,
    content_fingerprint,
    missing_profile_fonts,
    protected_object_manifest,
    semantic_title_heading_role,
    style_effective_font,
)
from finalize_docx import effective_font_failures  # noqa: E402
from structure_map import (  # noqa: E402
    _apply_front_matter,
    approved_role_paragraphs,
    candidate_structure_map,
    load_structure_map,
    prime_structure_map_locators,
    text_sha256,
)
from validate_profile import validate  # noqa: E402


PORTABLE_TEST_FONT = next(
    iter(
        sorted(
            available_font_names()
            - {"STHeiti Medium", "Songti", "Times New Roman"}
        )
    ),
    None,
)


class V051P3FoundationFormatSliceTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.portable_font = PORTABLE_TEST_FONT
        self.portable_profile = self.root / "portable-foundation-profile.json"
        if self.portable_font is not None:
            profile = json.loads(PROFILE.read_text(encoding="utf-8"))
            profile["profile_id"] = "v051-foundation-format-slice-portable-test"
            for rule in profile["rules"]:
                for key in (
                    "font_name",
                    "font_name_ascii",
                    "font_name_east_asia",
                    "font_name_complex_script",
                ):
                    if key in rule.get("properties", {}):
                        rule["properties"][key] = self.portable_font
            self.portable_profile.write_text(
                json.dumps(profile, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )

    def tearDown(self) -> None:
        self.temp.cleanup()

    def _strict_portable_profile(self) -> Path:
        if self.portable_font is None:
            self.skipTest("No host font is available for strict integration testing.")
        return self.portable_profile

    @staticmethod
    def _rule(role: str, properties: dict[str, object]) -> dict[str, object]:
        return {
            "id": "FMT-V051-TEST",
            "selector": {"kind": "paragraph_role", "value": role},
            "properties": properties,
        }

    def test_semantic_target_formatting_does_not_mutate_shared_builtin_style(self) -> None:
        document = Document()
        approved = document.add_paragraph("Approved body", style="Normal")
        excluded = document.add_paragraph("Image anchor", style="Normal")
        document.styles["Normal"].font.size = Pt(11)
        document.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        baseline_style_xml = etree.tostring(
            deepcopy(document.styles["Normal"].element), method="c14n", exclusive=True
        )

        rule = self._rule(
            "body_text",
            {
                "font_name_ascii": "Times New Roman",
                "font_name_east_asia": "Songti SC",
                "font_size_pt": 10.5,
                "bold": False,
                "alignment": "justify",
                "line_spacing_rule": "at_least",
                "line_spacing_pt": 18,
                "first_line_indent_chars": 2,
                "left_indent_pt": 0,
                "right_indent_pt": 0,
                "space_before_pt": 0,
                "space_after_pt": 0,
                "keep_with_next": False,
                "keep_together": False,
                "widow_control": True,
            },
        )
        apply_style_rule_to_paragraphs(document, rule, [approved], isolate_targets=True)

        self.assertNotEqual("Normal", approved.style.name)
        self.assertEqual("Normal", excluded.style.name)
        self.assertEqual(
            baseline_style_xml,
            etree.tostring(
                document.styles["Normal"].element, method="c14n", exclusive=True
            ),
            "approved-only formatting must not mutate the shared Normal style",
        )
        self.assertEqual(11, excluded.style.font.size.pt)
        self.assertEqual(WD_ALIGN_PARAGRAPH.LEFT, excluded.style.paragraph_format.alignment)
        self.assertEqual(10.5, approved.style.font.size.pt)
        self.assertEqual(WD_ALIGN_PARAGRAPH.JUSTIFY, approved.style.paragraph_format.alignment)
        self.assertEqual(WD_LINE_SPACING.AT_LEAST, approved.style.paragraph_format.line_spacing_rule)
        self.assertEqual(18, approved.style.paragraph_format.line_spacing.pt)
        self.assertEqual("200", approved.style.element.pPr.ind.get(qn("w:firstLineChars")))
        self.assertEqual("Times New Roman", style_effective_font(document, approved.style, "ascii")[0])
        self.assertEqual("Songti SC", style_effective_font(document, approved.style, "eastAsia")[0])

    def test_heading_rule_preserves_existing_page_break_before(self) -> None:
        document = Document()
        heading = document.add_paragraph("Approved chapter", style="Heading 1")
        heading.paragraph_format.page_break_before = True
        apply_style_rule_to_paragraphs(
            document,
            self._rule(
                "chapter_title",
                {
                    "font_size_pt": 18,
                    "bold": True,
                    "alignment": "center",
                    "line_spacing_rule": "at_least",
                    "line_spacing_pt": 27,
                    "space_before_pt": 24,
                    "space_after_pt": 18,
                    "first_line_indent_chars": 0,
                    "left_indent_pt": 0,
                    "right_indent_pt": 0,
                    "keep_with_next": True,
                    "keep_together": True,
                },
            ),
            [heading],
            isolate_targets=True,
        )
        self.assertIs(True, heading.paragraph_format.page_break_before)

    @staticmethod
    def _derived_style_name(role: str, base_style: str) -> str:
        identity = f"paragraph_role:{role}:{base_style}"
        return "Monograph Approved " + hashlib.sha256(
            identity.encode("utf-8")
        ).hexdigest()[:12]

    def test_existing_derived_style_collisions_fail_closed_and_repeat_is_legal(self) -> None:
        rule = self._rule("body_text", {"font_size_pt": 10.5})
        derived_name = self._derived_style_name("body_text", "Normal")

        for case in ("wrong_type", "wrong_base", "unused", "header_only"):
            with self.subTest(case=case):
                document = Document()
                target = document.add_paragraph("Approved body", style="Normal")
                if case == "wrong_type":
                    document.styles.add_style(derived_name, WD_STYLE_TYPE.CHARACTER)
                else:
                    derived = document.styles.add_style(
                        derived_name, WD_STYLE_TYPE.PARAGRAPH
                    )
                    derived.base_style = document.styles[
                        "Quote" if case == "wrong_base" else "Normal"
                    ]
                    if case == "wrong_base":
                        target.style = derived
                    elif case == "header_only":
                        document.sections[0].header.paragraphs[0].style = derived
                with self.assertRaises(FormatMonographError):
                    apply_style_rule_to_paragraphs(
                        document, rule, [target], isolate_targets=True
                    )

        document = Document()
        target = document.add_paragraph("Approved body", style="Normal")
        apply_style_rule_to_paragraphs(
            document, rule, [target], isolate_targets=True
        )
        apply_style_rule_to_paragraphs(
            document, rule, [target], isolate_targets=True
        )
        self.assertEqual(derived_name, target.style.name)

    def _front_matter_case(self, toc_title: str | None) -> tuple[Document, dict]:
        source = self.root / f"front-{toc_title or 'missing'}.docx"
        document = Document()
        document.add_paragraph("Fictional Systems Handbook")
        if toc_title is not None:
            document.add_paragraph(toc_title)
        document.add_paragraph("[[TOC]]")
        document.add_paragraph("Chapter One", style="Heading 1")
        document.save(source)
        structure = candidate_structure_map(source)
        structure["schema_version"] = "1.5"
        structure["status"] = "approved"
        structure["front_matter"]["approved"] = True
        structure["front_matter"]["insert_toc_heading_if_missing"] = False
        structure["front_matter"]["toc_heading_text"] = toc_title or "Custom contents"
        structure["front_matter"]["book_title_format"].pop("left_indent_pt", None)
        structure["front_matter"]["book_title_format"].pop("right_indent_pt", None)
        path = self.root / "front-structure.json"
        path.write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
        loaded = load_structure_map(path)
        formatted = Document(source)
        prime_structure_map_locators(formatted, loaded)
        return formatted, loaded

    def test_front_matter_preservation_mode_never_inserts_or_rewrites_toc_heading(self) -> None:
        for toc_title in (None, "目录", "目    录", "Custom Contents"):
            with self.subTest(toc_title=toc_title):
                document, structure = self._front_matter_case(toc_title)
                before = [paragraph.text for paragraph in document.paragraphs]
                changed = _apply_front_matter(document, structure)
                after = [paragraph.text for paragraph in document.paragraphs]
                self.assertEqual(before, after)
                self.assertEqual(1, changed)
                self.assertFalse(
                    hasattr(document, "_format_monograph_toc_heading"),
                    "preservation mode must not style or designate a TOC heading",
                )

    def test_candidate_profile_has_exact_foundation_values_and_no_active_content_rules(self) -> None:
        errors, profile = validate(PROFILE)
        self.assertEqual([], errors)
        self.assertEqual("approved", profile["approval"]["status"])
        rules = {
            (rule["selector"]["kind"], rule["selector"]["value"]): rule["properties"]
            for rule in profile["rules"]
            if rule["status"] == "approved" and rule["application"] == "automatic"
        }
        self.assertEqual({"page_size_policy": "preserve"}, rules[("document", "all")])
        self.assertNotIn(("field_role", "all"), rules)
        self.assertNotIn(("index_role", "table_of_contents"), rules)

        expected = {
            "chapter_title": (18, "center", 27, 24, 18),
            "level_2_section": (14, "left", 21, 18, 9),
            "level_3_section": (12, "left", 18, 12, 6),
            "level_4_section": (10.5, "left", 18, 9, 3),
        }
        for role, values in expected.items():
            props = rules[("paragraph_role", role)]
            self.assertEqual(values, tuple(props[key] for key in (
                "font_size_pt", "alignment", "line_spacing_pt", "space_before_pt", "space_after_pt"
            )))
            self.assertEqual("at_least", props["line_spacing_rule"])
            self.assertTrue(props["bold"])
            self.assertFalse(props["italic"])
            self.assertEqual("000000", props["color_hex"])
            self.assertTrue(props["keep_with_next"])
            self.assertTrue(props["keep_together"])
            self.assertNotIn("page_break_before", props)

        body = rules[("paragraph_role", "body_text")]
        self.assertEqual(10.5, body["font_size_pt"])
        self.assertFalse(body["bold"])
        self.assertEqual("justify", body["alignment"])
        self.assertEqual("at_least", body["line_spacing_rule"])
        self.assertEqual(18, body["line_spacing_pt"])
        self.assertEqual(2, body["first_line_indent_chars"])
        self.assertTrue(body["widow_control"])
        self.assertFalse(body["keep_with_next"])
        self.assertFalse(body["keep_together"])
        title = rules[("paragraph_role", "title")]
        self.assertFalse(title["italic"])
        self.assertEqual("000000", title["color_hex"])

    @unittest.skipUnless(sys.platform == "darwin", "macOS live-font gate")
    def test_shipping_profile_strict_fonts_are_available_on_macos(self) -> None:
        profile = json.loads(PROFILE.read_text(encoding="utf-8"))
        self.assertEqual([], missing_profile_fonts(profile))

    def test_final_effective_font_audit_targets_derived_styles_and_keeps_legacy(self) -> None:
        profile_path = self._strict_portable_profile()
        profile = json.loads(profile_path.read_text(encoding="utf-8"))
        source, structure_path, _ = self._synthetic_source_and_map()
        output = self.root / "font-audit"
        applied = self._run_script(
            "apply_profile.py",
            source,
            "--profile",
            profile_path,
            "--structure-map",
            structure_path,
            "--output-dir",
            output,
        )
        self.assertEqual(0, applied.returncode, applied.stdout + applied.stderr)
        formatted = output / "fictional-foundation-slice-formatted.docx"
        structure = load_structure_map(structure_path)
        self.assertEqual([], effective_font_failures(formatted, profile, structure))

        normal_changed = self.root / "normal-cambria.docx"
        normal_document = Document(formatted)
        normal_document.styles["Normal"].font.name = "Cambria"
        normal_document.save(normal_changed)
        self.assertEqual(
            [],
            effective_font_failures(normal_changed, profile, structure),
            "an unapproved global Normal font must not replace approved-target auditing",
        )

        tampered_path = self.root / "tampered-derived-font.docx"
        tampered = Document(formatted)
        target = next(
            paragraph
            for paragraph in tampered.paragraphs
            if paragraph.text == "1.1 Approved Scope"
        )
        target.style.font.name = "Definitely Wrong Font"
        tampered.save(tampered_path)
        failures = effective_font_failures(tampered_path, profile, structure)
        self.assertTrue(
            any(
                item.get("rule") == "FMT-HEAD-502"
                and item.get("property") == "font_name_ascii"
                for item in failures
            ),
            failures,
        )

        rebound_path = self.root / "rebound-approved-target.docx"
        rebound = Document(formatted)
        rebound_target = next(
            paragraph
            for paragraph in rebound.paragraphs
            if paragraph.text == "1.1 Approved Scope"
        )
        rebound_target.style = rebound.styles["Heading 2"]
        rebound.save(rebound_path)
        rebound_failures = effective_font_failures(rebound_path, profile, structure)
        self.assertTrue(
            any(
                item.get("rule") == "FMT-HEAD-502"
                and item.get("reason") == "derived_target_style_mismatch"
                for item in rebound_failures
            ),
            rebound_failures,
        )

        body_rule = next(
            rule
            for rule in profile["rules"]
            if rule["selector"]
            == {"kind": "paragraph_role", "value": "body_text"}
        )
        no_target_path = self.root / "no-approved-body-target.docx"
        no_target = Document()
        no_target.styles["Normal"].font.name = "Cambria"
        self.assertEqual(
            0,
            apply_style_rule_to_paragraphs(
                no_target, body_rule, [], isolate_targets=True
            ),
        )
        self.assertEqual(
            0,
            apply_style_rule_to_paragraphs(
                no_target, body_rule, [], isolate_targets=True
            ),
        )
        derived_name = self._derived_style_name("body_text", "Normal")
        self.assertNotIn(derived_name, {style.name for style in no_target.styles})
        no_target.save(no_target_path)
        empty_semantic_map = {
            "schema_version": "1.5",
            "paragraph_roles": [],
            "toc_ranges": [],
        }
        self.assertEqual(
            [],
            effective_font_failures(
                no_target_path,
                {"rules": [body_rule]},
                empty_semantic_map,
            ),
        )

        unused_path = self.root / "unused-derived-without-approved-target.docx"
        unused = Document()
        unused_derived = unused.styles.add_style(
            derived_name, WD_STYLE_TYPE.PARAGRAPH
        )
        unused_derived.base_style = unused.styles["Normal"]
        with self.assertRaises(FormatMonographError):
            apply_style_rule_to_paragraphs(
                unused, body_rule, [], isolate_targets=True
            )
        unused.save(unused_path)
        unused_failures = effective_font_failures(
            unused_path,
            {"rules": [body_rule]},
            empty_semantic_map,
        )
        self.assertTrue(
            any(
                item.get("rule") == body_rule["id"]
                and item.get("reason")
                == "derived_style_without_approved_target"
                for item in unused_failures
            ),
            unused_failures,
        )

        legacy_path = self.root / "legacy-global-style.docx"
        legacy = Document()
        legacy.add_paragraph("Legacy body")
        legacy.styles["Normal"].font.name = self.portable_font
        legacy.save(legacy_path)
        legacy_rule = deepcopy(body_rule)
        legacy_rule["properties"] = {"font_name_ascii": self.portable_font}
        self.assertEqual(
            [], effective_font_failures(legacy_path, {"rules": [legacy_rule]})
        )

    @staticmethod
    def _add_omml(paragraph: object) -> None:
        math = OxmlElement("m:oMath")
        run = OxmlElement("m:r")
        text = OxmlElement("m:t")
        text.text = "x+y=7"
        run.append(text)
        math.append(run)
        paragraph._p.append(math)

    @staticmethod
    def _add_field_and_bookmark(paragraph: object) -> None:
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), "51")
        bookmark_start.set(qn("w:name"), "fictional_anchor")
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), "51")
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), " REF fictional_anchor \\h ")
        result_run = OxmlElement("w:r")
        result_text = OxmlElement("w:t")
        result_text.text = "anchor"
        result_run.append(result_text)
        field.append(result_run)
        paragraph._p.extend((bookmark_start, bookmark_end, field))

    def _synthetic_source_and_map(self) -> tuple[Path, Path, list[str]]:
        source = self.root / "fictional-foundation-slice.docx"
        image_path = self.root / "fictional-pixel.png"
        image_path.write_bytes(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
            )
        )
        document = Document()
        first = document.sections[0]
        first.page_width = Mm(182)
        first.page_height = Mm(257)
        first.top_margin = Mm(23)
        first.bottom_margin = Mm(27)
        first.left_margin = Mm(31)
        first.right_margin = Mm(19)
        first.gutter = Mm(4)
        first.header_distance = Mm(11)
        first.footer_distance = Mm(13)

        title = document.add_paragraph("Fictional Systems Handbook", style="Title")
        title_borders = OxmlElement("w:pBdr")
        title_bottom = OxmlElement("w:bottom")
        title_bottom.set(qn("w:val"), "single")
        title_bottom.set(qn("w:color"), "4472C4")
        title_borders.append(title_bottom)
        title._p.get_or_add_pPr().append(title_borders)
        break_one = document.add_paragraph("Title-page tail")
        break_one.add_run().add_break(WD_BREAK.PAGE)
        document.add_paragraph("Custom Contents")
        document.add_paragraph("[[TOC]]")
        break_two = document.add_paragraph("Contents-page tail")
        break_two.add_run().add_break(WD_BREAK.PAGE)

        h1 = document.add_paragraph("第1章 Fictional Foundations", style="Heading 1")
        h1.paragraph_format.page_break_before = False
        document.add_paragraph("1.1 Approved Scope", style="Heading 2")
        document.add_paragraph("1.1.1 Stable Identity", style="Heading 3")
        document.add_paragraph("1.1.1.1 Deterministic Case", style="Heading 4")
        document.add_paragraph("Approved Unnumbered Heading", style="Heading 3")
        approved_body = [
            "Approved body alpha keeps  double  spaces and\ta tab.",
            "Approved body beta is fictional and contains mixed 中文 text.",
        ]
        for value in approved_body:
            document.add_paragraph(value)
        field_paragraph = document.add_paragraph("Approved body field: ")
        self._add_field_and_bookmark(field_paragraph)
        approved_body.append(field_paragraph.text)
        break_three = document.add_paragraph("Body-page tail")
        break_three.add_run().add_break(WD_BREAK.PAGE)

        document.add_paragraph("图 1-1 Excluded caption", style="Caption")
        document.add_paragraph("Excluded bullet", style="List Bullet")
        document.add_paragraph("Excluded quotation", style="Quote")
        bibliography = document.styles.add_style(
            "Bibliography", WD_STYLE_TYPE.PARAGRAPH
        )
        bibliography.base_style = document.styles["Normal"]
        document.add_paragraph("Excluded reference", style="Bibliography")
        code_style = document.styles.add_style("Fictional Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = document.styles["Normal"]
        document.add_paragraph("print('excluded code')", style="Fictional Code")
        document.add_paragraph("Excluded unknown normal paragraph")
        image_anchor = document.add_paragraph()
        image_anchor.add_run().add_picture(str(image_path), width=Inches(0.3))
        formula = document.add_paragraph()
        self._add_omml(formula)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Excluded table header A"
        table.cell(0, 1).text = "Excluded table header B"
        table.cell(1, 0).text = "Excluded table value A"
        table.cell(1, 1).text = "Excluded table value B"
        second = document.add_section(WD_SECTION.CONTINUOUS)
        second.page_width = Mm(182)
        second.page_height = Mm(257)
        second.top_margin = Mm(17)
        second.bottom_margin = Mm(29)
        second.left_margin = Mm(22)
        second.right_margin = Mm(24)
        document.add_paragraph("Excluded note after continuous section")
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        approved_text = {
            "Fictional Systems Handbook": "title",
            "第1章 Fictional Foundations": "chapter_title",
            "1.1 Approved Scope": "level_2_section",
            "1.1.1 Stable Identity": "level_3_section",
            "1.1.1.1 Deterministic Case": "level_4_section",
            "Approved Unnumbered Heading": "level_3_section",
            **{value: "body" for value in approved_body},
        }
        for entry in structure["paragraph_roles"]:
            paragraph_index = int(entry["locator"].get("paragraph", -1))
            if paragraph_index < 0:
                continue
            text = document.paragraphs[paragraph_index].text
            if text not in approved_text:
                continue
            canonical = approved_text[text]
            entry["role"] = "body_text" if canonical == "body" else canonical
            entry["canonical_role"] = canonical
            entry["approved"] = True
        structure["front_matter"]["approved"] = False
        structure["pagination_sections"]["approved"] = False
        structure_path = self.root / "fictional-foundation-structure.json"
        structure_path.write_text(
            json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        return source, structure_path, approved_body

    @staticmethod
    def _package_contract(path: Path) -> dict[str, object]:
        with zipfile.ZipFile(path) as package:
            document = etree.fromstring(package.read("word/document.xml"))
            relationships = []
            for name in sorted(
                item for item in package.namelist() if item.endswith(".rels")
            ):
                root = etree.fromstring(package.read(name))
                relationships.append(
                    (
                        name,
                        sorted(
                            (
                                child.get("Id"),
                                child.get("Type"),
                                child.get("Target"),
                                child.get("TargetMode"),
                            )
                            for child in root
                        ),
                    )
                )
            return {
                "content": content_fingerprint(path),
                "protected": protected_object_manifest(path),
                "sections": [
                    etree.tostring(item, method="c14n", exclusive=True)
                    for item in document.xpath(".//w:sectPr", namespaces={"w": document.nsmap["w"]})
                ],
                "page_break_paragraphs": [
                    index
                    for index, paragraph in enumerate(
                        document.xpath("/w:document/w:body/w:p", namespaces={"w": document.nsmap["w"]})
                    )
                    if paragraph.xpath(".//w:br[@w:type='page']", namespaces={"w": document.nsmap["w"]})
                ],
                "tables": [
                    etree.tostring(item, method="c14n", exclusive=True)
                    for item in document.xpath(".//w:tbl", namespaces={"w": document.nsmap["w"]})
                ],
                "field_instructions": document.xpath(
                    ".//w:instrText/text() | .//w:fldSimple/@w:instr",
                    namespaces={"w": document.nsmap["w"]},
                ),
                "bookmarks": sorted(
                    zip(
                        document.xpath(".//w:bookmarkStart/@w:id", namespaces={"w": document.nsmap["w"]}),
                        document.xpath(".//w:bookmarkStart/@w:name", namespaces={"w": document.nsmap["w"]}),
                    )
                ),
                "relationships": relationships,
            }

    def _run_script(self, name: str, *args: object) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(SCRIPTS / name), *(str(value) for value in args)],
            cwd=SKILL,
            capture_output=True,
            text=True,
            check=False,
        )

    def test_prepare_approved_map_apply_audit_preserves_exclusions_and_is_idempotent(self) -> None:
        source, structure, approved_body = self._synthetic_source_and_map()
        profile_path = self._strict_portable_profile()
        before = self._package_contract(source)
        source_document = Document(source)
        excluded_texts = {
            "图 1-1 Excluded caption",
            "Excluded bullet",
            "Excluded quotation",
            "Excluded reference",
            "print('excluded code')",
            "Excluded unknown normal paragraph",
        }
        excluded_paragraph_xml = {
            paragraph.text: etree.tostring(
                paragraph._p, method="c14n", exclusive=True
            )
            for paragraph in source_document.paragraphs
            if paragraph.text in excluded_texts
        }
        excluded_style_xml = {
            name: etree.tostring(
                deepcopy(source_document.styles[name].element),
                method="c14n",
                exclusive=True,
            )
            for name in (
                "Caption",
                "List Bullet",
                "Quote",
                "Bibliography",
                "Fictional Code",
            )
        }
        output_one = self.root / "first"
        result = self._run_script(
            "apply_profile.py",
            source,
            "--profile",
            profile_path,
            "--structure-map",
            structure,
            "--output-dir",
            output_one,
        )
        self.assertEqual(0, result.returncode, result.stdout + result.stderr)
        formatted_one = output_one / "fictional-foundation-slice-formatted.docx"
        self.assertEqual(before, self._package_contract(formatted_one))

        audit = self._run_script(
            "audit_docx.py",
            source,
            formatted_one,
            "--profile",
            profile_path,
            "--structure-map",
            structure,
        )
        self.assertEqual(0, audit.returncode, audit.stdout + audit.stderr)
        self.assertTrue(json.loads(audit.stdout)["passed"])

        formatted = Document(formatted_one)
        approved_values = {
            "Fictional Systems Handbook",
            "第1章 Fictional Foundations",
            "1.1 Approved Scope",
            "1.1.1 Stable Identity",
            "1.1.1.1 Deterministic Case",
            "Approved Unnumbered Heading",
            *approved_body,
        }
        for paragraph in formatted.paragraphs:
            if paragraph.text in approved_values:
                self.assertTrue(paragraph.style.name.startswith("Monograph Approved "))
            elif paragraph.style.name in {
                "Normal",
                "Caption",
                "List Bullet",
                "Quote",
                "Bibliography",
                "Fictional Code",
            }:
                self.assertFalse(paragraph.style.name.startswith("Monograph Approved "))
        neutral_targets = {
            "Fictional Systems Handbook",
            "第1章 Fictional Foundations",
            "1.1 Approved Scope",
            "1.1.1 Stable Identity",
            "1.1.1.1 Deterministic Case",
            "Approved Unnumbered Heading",
        }
        for paragraph in formatted.paragraphs:
            if paragraph.text not in neutral_targets:
                continue
            self.assertEqual("000000", str(paragraph.style.font.color.rgb))
            self.assertIs(False, paragraph.style.font.italic)
            borders = paragraph.style.element.pPr.find(qn("w:pBdr"))
            self.assertIsNotNone(borders)
            self.assertEqual(
                {"top", "left", "bottom", "right", "between", "bar"},
                {child.tag.rsplit("}", 1)[-1] for child in borders},
            )
            self.assertTrue(
                all(child.get(qn("w:val")) == "nil" for child in borders)
            )
            direct_p_pr = paragraph._p.pPr
            self.assertIsNone(
                None if direct_p_pr is None else direct_p_pr.find(qn("w:pBdr"))
            )
        self.assertEqual(
            excluded_paragraph_xml,
            {
                paragraph.text: etree.tostring(
                    paragraph._p, method="c14n", exclusive=True
                )
                for paragraph in formatted.paragraphs
                if paragraph.text in excluded_texts
            },
        )
        self.assertEqual(
            excluded_style_xml,
            {
                name: etree.tostring(
                    formatted.styles[name].element,
                    method="c14n",
                    exclusive=True,
                )
                for name in excluded_style_xml
            },
        )

        output_two = self.root / "second"
        second = self._run_script(
            "apply_profile.py",
            formatted_one,
            "--profile",
            profile_path,
            "--structure-map",
            structure,
            "--output-dir",
            output_two,
        )
        self.assertEqual(0, second.returncode, second.stdout + second.stderr)
        formatted_two = output_two / "fictional-foundation-slice-formatted-formatted.docx"
        self.assertEqual(self._package_contract(formatted_one), self._package_contract(formatted_two))
        first_doc = Document(formatted_one)
        second_doc = Document(formatted_two)
        self.assertEqual(
            [(p.text, p.style.style_id) for p in first_doc.paragraphs],
            [(p.text, p.style.style_id) for p in second_doc.paragraphs],
        )
        self.assertEqual(
            etree.tostring(first_doc.styles.element, method="c14n", exclusive=True),
            etree.tostring(second_doc.styles.element, method="c14n", exclusive=True),
        )

    def test_formatted_inspect_and_prepare_preserve_title_and_heading_semantics(self) -> None:
        source, structure, _ = self._synthetic_source_and_map()
        profile_path = self._strict_portable_profile()
        output = self.root / "formatted-inspection"
        applied = self._run_script(
            "apply_profile.py",
            source,
            "--profile",
            profile_path,
            "--structure-map",
            structure,
            "--output-dir",
            output,
        )
        self.assertEqual(0, applied.returncode, applied.stdout + applied.stderr)
        formatted = output / "fictional-foundation-slice-formatted.docx"
        inventory_path = self.root / "formatted-inventory.json"
        candidate_path = self.root / "formatted-candidate.json"
        inspected = self._run_script(
            "inspect_docx.py",
            formatted,
            "--output",
            inventory_path,
            "--structure-map-output",
            candidate_path,
        )
        self.assertEqual(0, inspected.returncode, inspected.stdout + inspected.stderr)
        inventory = json.loads(inventory_path.read_text(encoding="utf-8"))
        candidate = json.loads(candidate_path.read_text(encoding="utf-8"))
        self.assertEqual(5, len(inventory["headings"]))
        self.assertEqual(
            {"heading_1", "heading_2", "heading_3", "heading_4"},
            {item["semantic_role"] for item in inventory["headings"]},
        )
        self.assertEqual(5, len(candidate["headings"]))
        roles_by_hash = {
            item["text_sha256"]: item["canonical_role"]
            for item in candidate["paragraph_roles"]
        }
        aliases = {
            "heading_1": "chapter_title",
            "heading_2": "level_2_section",
            "heading_3": "level_3_section",
            "heading_4": "level_4_section",
        }
        for paragraph in Document(formatted).paragraphs:
            semantic = semantic_title_heading_role(paragraph.style)
            if semantic is not None:
                self.assertEqual(
                    aliases.get(semantic, semantic),
                    roles_by_hash[text_sha256(paragraph.text)],
                )

        prepare_work = self.root / "prepare-work"
        prepared = self._run_script(
            "run_monograph.py",
            "prepare",
            formatted,
            "--profile",
            profile_path,
            "--work-dir",
            prepare_work,
        )
        self.assertEqual(0, prepared.returncode, prepared.stdout + prepared.stderr)
        prepared_map = json.loads(
            (prepare_work / "candidate-structure-map.json").read_text(
                encoding="utf-8"
            )
        )
        self.assertEqual(5, len(prepared_map["headings"]))

    def test_combined_front_matter_title_rule_and_pagination_audit_repeat(self) -> None:
        profile_path = self._strict_portable_profile()
        for insert_heading in (False, True):
            with self.subTest(insert_heading=insert_heading):
                source = self.root / f"front-combined-{insert_heading}.docx"
                document = Document()
                document.add_paragraph("Fictional Systems Handbook", style="Title")
                document.add_paragraph("[[TOC]]")
                document.add_paragraph("Chapter One", style="Heading 1")
                document.save(source)
                structure = candidate_structure_map(source)
                structure["status"] = "approved"
                for entry in structure["paragraph_roles"]:
                    paragraph = document.paragraphs[
                        int(entry["locator"]["paragraph"])
                    ]
                    if paragraph.text == "Fictional Systems Handbook":
                        entry["role"] = "title"
                        entry["canonical_role"] = "title"
                        entry["approved"] = True
                    elif paragraph.text == "Chapter One":
                        entry["role"] = "chapter_title"
                        entry["canonical_role"] = "chapter_title"
                        entry["approved"] = True
                structure["front_matter"]["approved"] = True
                structure["front_matter"][
                    "insert_toc_heading_if_missing"
                ] = insert_heading
                structure["front_matter"]["book_title_format"].pop(
                    "left_indent_pt", None
                )
                structure["front_matter"]["book_title_format"].pop(
                    "right_indent_pt", None
                )
                for key in (
                    "font_name_ascii",
                    "font_name_east_asia",
                    "font_name_complex_script",
                ):
                    structure["front_matter"]["book_title_format"][key] = (
                        self.portable_font
                    )
                structure["pagination_sections"]["approved"] = True
                structure_path = self.root / f"front-combined-{insert_heading}.json"
                structure_path.write_text(
                    json.dumps(structure, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )

                first_dir = self.root / f"front-{insert_heading}-first"
                first = self._run_script(
                    "apply_profile.py",
                    source,
                    "--profile",
                    profile_path,
                    "--structure-map",
                    structure_path,
                    "--output-dir",
                    first_dir,
                )
                self.assertEqual(0, first.returncode, first.stdout + first.stderr)
                formatted = first_dir / f"front-combined-{insert_heading}-formatted.docx"
                audited = self._run_script(
                    "audit_docx.py",
                    source,
                    formatted,
                    "--profile",
                    profile_path,
                    "--structure-map",
                    structure_path,
                )
                self.assertEqual(0, audited.returncode, audited.stdout + audited.stderr)
                audit_result = json.loads(audited.stdout)
                self.assertTrue(audit_result["passed"], audit_result)

                loaded = load_structure_map(structure_path)
                repeated_document = Document(formatted)
                prime_structure_map_locators(repeated_document, loaded)
                self.assertEqual(1, _apply_front_matter(repeated_document, loaded))
                profile = json.loads(profile_path.read_text(encoding="utf-8"))
                title_rule = next(
                    rule
                    for rule in profile["rules"]
                    if rule["selector"]
                    == {"kind": "paragraph_role", "value": "title"}
                )
                title_targets = approved_role_paragraphs(
                    repeated_document, loaded, title_rule["selector"]
                )
                self.assertEqual(
                    1,
                    apply_style_rule_to_paragraphs(
                        repeated_document,
                        title_rule,
                        title_targets,
                        isolate_targets=True,
                    ),
                )
                repeated = self.root / f"front-combined-{insert_heading}-repeat.docx"
                repeated_document.save(repeated)
                repeated_audit = self._run_script(
                    "audit_docx.py",
                    source,
                    repeated,
                    "--profile",
                    profile_path,
                    "--structure-map",
                    structure_path,
                )
                self.assertEqual(
                    0,
                    repeated_audit.returncode,
                    repeated_audit.stdout + repeated_audit.stderr,
                )


if __name__ == "__main__":
    unittest.main()
