from __future__ import annotations

import tempfile
import unittest
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

import sys

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "format-monograph" / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import FormatMonographError  # noqa: E402
from docx_pagination import (  # noqa: E402
    _ensure_page_field,
    apply_pagination_sections,
    audit_pagination_sections,
    pagination_inventory,
)
from structure_map import (  # noqa: E402
    apply_structure_map,
    candidate_structure_map,
    prime_structure_map_locators,
    text_sha256,
)


def _resolver(document, locator):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == locator["text"])


def _settings() -> dict:
    return {
        "approved": True,
        "toc_start": {"text": "Synthetic TOC"},
        "body_start": {"text": "Chapter One"},
        "number_format": "decimal",
        "start_at": {"toc": 1, "body": 1},
        "continue_after_body_start": True,
    }


def _approved_chapters() -> list[dict]:
    return [
        {
            "approved": True,
            "level": 1,
            "locator": {"text": "Chapter Two"},
        }
    ]


def _set_start(section, value: int) -> None:
    element = OxmlElement("w:pgNumType")
    element.set(qn("w:fmt"), "decimal")
    element.set(qn("w:start"), str(value))
    section._sectPr.append(element)


def _append_complex_page(paragraph, *, include_end: bool = True) -> None:
    for kind, text in (("begin", None), (None, " PAGE "), ("separate", None)):
        run = OxmlElement("w:r")
        if kind is None:
            instruction = OxmlElement("w:instrText")
            instruction.text = text
            run.append(instruction)
        else:
            marker = OxmlElement("w:fldChar")
            marker.set(qn("w:fldCharType"), kind)
            run.append(marker)
        paragraph._p.append(run)
    result_run = OxmlElement("w:r")
    result = OxmlElement("w:t")
    result.text = "1"
    result_run.append(result)
    paragraph._p.append(result_run)
    if include_end:
        end_run = OxmlElement("w:r")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        end_run.append(end)
        paragraph._p.append(end_run)


def _synthetic_book(*, numbered_preface: bool = False) -> tuple[object, dict[str, int]]:
    document = Document()
    document.add_paragraph("Synthetic title")
    document.add_section(WD_SECTION.NEW_PAGE)
    indexes = {"title": 0}
    if numbered_preface:
        indexes["front"] = 1
        document.add_paragraph("Synthetic preface")
        document.add_section(WD_SECTION.NEW_PAGE)
    indexes["toc"] = len(document.sections) - 1
    document.add_paragraph("Synthetic TOC")
    document.add_section(WD_SECTION.NEW_PAGE)
    indexes["body"] = len(document.sections) - 1
    document.add_paragraph("Chapter One", style="Heading 1")
    document.sections[indexes["body"]].different_first_page_header_footer = True
    document.add_paragraph("Body text")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Synthetic protected table"
    protected_object = document.add_paragraph()
    protected_object._p.append(OxmlElement("w:object"))
    document.add_section(WD_SECTION.NEW_PAGE)
    indexes["chapter_two"] = len(document.sections) - 1
    document.add_paragraph("Chapter Two", style="Heading 1")
    document.sections[indexes["chapter_two"]].different_first_page_header_footer = True
    document.add_section(WD_SECTION.NEW_PAGE)
    indexes["landscape"] = len(document.sections) - 1
    document.sections[indexes["landscape"]].orientation = WD_ORIENT.LANDSCAPE
    document.sections[indexes["landscape"]].different_first_page_header_footer = False
    document.add_paragraph("Landscape table placeholder")
    if numbered_preface:
        _set_start(document.sections[indexes["front"]], 1)
    return document, indexes


def _package_subset(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as package:
        return {
            name: package.read(name)
            for name in sorted(package.namelist())
            if name in {"word/document.xml", "word/settings.xml", "word/_rels/document.xml.rels"}
            or name.startswith("word/footer")
        }


def _mutation_snapshot(document) -> tuple[bytes, bytes, tuple, tuple]:
    rels = tuple(
        sorted(
            (rel_id, rel.reltype, rel.target_ref, rel.is_external)
            for rel_id, rel in document.part.rels.items()
        )
    )
    running_parts = tuple(
        sorted(
            (rel_id, etree.tostring(part.element))
            for rel_id, part in document.part.related_parts.items()
            if hasattr(part, "element")
            and str(getattr(part, "partname", "")).startswith(("/word/footer", "/word/header"))
        )
    )
    return etree.tostring(document.element), etree.tostring(document.settings.element), rels, running_parts


def _nonpagination_document_xml(document) -> bytes:
    root = deepcopy(document.element)
    for element in root.xpath(".//w:pgNumType | .//w:footerReference"):
        element.getparent().remove(element)
    return etree.tostring(root, method="c14n")


def _nonfooter_relationships(document) -> tuple:
    return tuple(
        sorted(
            (rel_id, rel.reltype, rel.target_ref, rel.is_external)
            for rel_id, rel in document.part.rels.items()
            if not rel.reltype.endswith("/footer")
        )
    )


class V051P3PageNumberFoundationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_existing_boundaries_front_continuation_body_restart_and_idempotence(self) -> None:
        for numbered_preface in (False, True):
            with self.subTest(numbered_preface=numbered_preface):
                document, indexes = _synthetic_book(numbered_preface=numbered_preface)
                section_fingerprint = [
                    (
                        section.orientation,
                        section.top_margin,
                        section.bottom_margin,
                        section.left_margin,
                        section.right_margin,
                        section.different_first_page_header_footer,
                    )
                    for section in document.sections
                ]
                protected_xml = _nonpagination_document_xml(document)
                protected_relationships = _nonfooter_relationships(document)
                result = apply_pagination_sections(
                    document,
                    _settings(),
                    _resolver,
                    approved_headings=_approved_chapters(),
                )
                self.assertEqual(indexes["toc"], result["toc_section"])
                self.assertEqual(indexes["body"], result["body_section"])
                first = self.root / f"first-{numbered_preface}.docx"
                document.save(first)
                inventory = pagination_inventory(first)
                expected_front = indexes.get("front", indexes["toc"])
                self.assertEqual("1", inventory["sections"][expected_front]["page_number_start"])
                if numbered_preface:
                    self.assertIsNone(inventory["sections"][indexes["toc"]]["page_number_start"])
                self.assertEqual("1", inventory["sections"][indexes["body"]]["page_number_start"])
                self.assertTrue(
                    all(
                        item["page_number_start"] is None
                        for item in inventory["sections"][indexes["body"] + 1 :]
                    )
                )
                for item in inventory["sections"][expected_front:]:
                    self.assertEqual(1, item["footer_page_field_counts"]["default"])
                    self.assertEqual(1, item["footer_page_field_counts"]["even"])
                    self.assertEqual("right", item["footer_formats"]["default"]["alignment"])
                    self.assertEqual("left", item["footer_formats"]["even"]["alignment"])
                    self.assertEqual("宋体", item["footer_formats"]["default"]["font_east_asia"])
                    self.assertEqual("Times New Roman", item["footer_formats"]["default"]["font_ascii"])
                    self.assertEqual("18", item["footer_formats"]["default"]["font_size_half_points"])
                for index in (indexes["body"], indexes["chapter_two"]):
                    self.assertTrue(inventory["sections"][index]["different_first_page"])
                    self.assertEqual(1, inventory["sections"][index]["footer_page_field_counts"]["first"])
                self.assertEqual(section_fingerprint, [
                    (
                        section.orientation,
                        section.top_margin,
                        section.bottom_margin,
                        section.left_margin,
                        section.right_margin,
                        section.different_first_page_header_footer,
                    )
                    for section in document.sections
                ])
                self.assertEqual(protected_xml, _nonpagination_document_xml(document))
                self.assertEqual(protected_relationships, _nonfooter_relationships(document))
                failures, _ = audit_pagination_sections(first, Document(first), _settings(), _resolver)
                self.assertEqual([], failures)

                repeated = Document(first)
                apply_pagination_sections(
                    repeated,
                    _settings(),
                    _resolver,
                    approved_headings=_approved_chapters(),
                )
                second = self.root / f"second-{numbered_preface}.docx"
                repeated.save(second)
                self.assertEqual(_package_subset(first), _package_subset(second))

    def test_complex_footer_matrix_fails_before_any_write(self) -> None:
        def text(footer):
            footer.paragraphs[0].text = "Chapter footer"

        def mixed(footer):
            footer.paragraphs[0].text = "Page "
            field = OxmlElement("w:fldSimple")
            field.set(qn("w:instr"), "PAGE")
            footer.paragraphs[0]._p.append(field)

        def raw(name):
            return lambda footer: footer.paragraphs[0]._p.append(OxmlElement(name))

        def other_field(footer):
            field = OxmlElement("w:fldSimple")
            field.set(qn("w:instr"), "DATE")
            footer.paragraphs[0]._p.append(field)

        def relationship(footer):
            footer.part.relate_to(
                "https://example.invalid/synthetic",
                "urn:synthetic:unknown-footer-relationship",
                is_external=True,
            )

        cases = {
            "text": text,
            "mixed": mixed,
            "table": raw("w:tbl"),
            "drawing": raw("w:drawing"),
            "vml": raw("w:pict"),
            "ole": raw("w:object"),
            "other_field": other_field,
            "unknown_relationship": relationship,
        }
        for name, mutate in cases.items():
            with self.subTest(name=name):
                document, indexes = _synthetic_book()
                footer = document.sections[indexes["toc"]].footer
                footer.is_linked_to_previous = False
                mutate(footer)
                before = _mutation_snapshot(document)
                with self.assertRaisesRegex(FormatMonographError, "footer"):
                    apply_pagination_sections(document, _settings(), _resolver)
                self.assertEqual(before, _mutation_snapshot(document))

    def test_missing_and_ambiguous_boundaries_fail_closed(self) -> None:
        missing = Document()
        missing.add_paragraph("Synthetic TOC")
        missing.add_paragraph("Chapter One")
        before = _mutation_snapshot(missing)
        with self.assertRaisesRegex(FormatMonographError, "existing, exact section boundary"):
            apply_pagination_sections(missing, _settings(), _resolver)
        self.assertEqual(before, _mutation_snapshot(missing))

        ambiguous, indexes = _synthetic_book(numbered_preface=True)
        ambiguous.sections[indexes["front"]]._sectPr.find(qn("w:pgNumType")).set(
            qn("w:start"), "2"
        )
        before = _mutation_snapshot(ambiguous)
        with self.assertRaisesRegex(FormatMonographError, "ambiguous or inconsistent"):
            apply_pagination_sections(ambiguous, _settings(), _resolver)
        self.assertEqual(before, _mutation_snapshot(ambiguous))

        implicit_front, _ = _synthetic_book()
        _ensure_page_field(implicit_front.sections[0].footer, 2)
        implicit_front.settings.odd_and_even_pages_header_footer = True
        before = _mutation_snapshot(implicit_front)
        with self.assertRaisesRegex(FormatMonographError, "shared across"):
            apply_pagination_sections(implicit_front, _settings(), _resolver)
        self.assertEqual(before, _mutation_snapshot(implicit_front))

    def test_frozen_pagination_settings_fail_before_write(self) -> None:
        cases = {
            "format": {**_settings(), "number_format": "roman"},
            "toc_start": {**_settings(), "start_at": {"toc": 2, "body": 1}},
            "body_start": {**_settings(), "start_at": {"toc": 1, "body": 2}},
            "continue": {**_settings(), "continue_after_body_start": False},
        }
        for name, settings in cases.items():
            with self.subTest(name=name):
                document, _ = _synthetic_book()
                before = _mutation_snapshot(document)
                with self.assertRaises(FormatMonographError):
                    apply_pagination_sections(document, settings, _resolver)
                self.assertEqual(before, _mutation_snapshot(document))

    def test_conflicting_footer_alignment_roles_fail_before_write(self) -> None:
        for conflict in ("first_even", "cross_section"):
            with self.subTest(conflict=conflict):
                document, indexes = _synthetic_book()
                if conflict == "first_even":
                    section = document.sections[indexes["body"]]
                    section.even_page_footer.is_linked_to_previous = False
                    section.first_page_footer.is_linked_to_previous = False
                    references = {
                        item.get(qn("w:type"), "default"): item
                        for item in section._sectPr.findall(qn("w:footerReference"))
                    }
                    references["first"].set(
                        qn("r:id"), references["even"].get(qn("r:id"))
                    )
                else:
                    toc = document.sections[indexes["toc"]]
                    chapter = document.sections[indexes["chapter_two"]]
                    toc.footer.is_linked_to_previous = False
                    chapter.even_page_footer.is_linked_to_previous = False
                    toc_ref = next(
                        item
                        for item in toc._sectPr.findall(qn("w:footerReference"))
                        if item.get(qn("w:type"), "default") == "default"
                    )
                    even_ref = next(
                        item
                        for item in chapter._sectPr.findall(qn("w:footerReference"))
                        if item.get(qn("w:type")) == "even"
                    )
                    even_ref.set(qn("r:id"), toc_ref.get(qn("r:id")))
                before = _mutation_snapshot(document)
                with self.assertRaisesRegex(FormatMonographError, "conflicting alignment"):
                    apply_pagination_sections(
                        document,
                        _settings(),
                        _resolver,
                        approved_headings=_approved_chapters(),
                    )
                self.assertEqual(before, _mutation_snapshot(document))

    def test_footer_part_identity_blocks_cross_area_and_aliased_relationships(self) -> None:
        document, _ = _synthetic_book()
        document.sections[0].footer.paragraphs[0].text = ""
        before = _mutation_snapshot(document)
        with self.assertRaisesRegex(FormatMonographError, "shared across"):
            apply_pagination_sections(
                document,
                _settings(),
                _resolver,
                approved_headings=_approved_chapters(),
            )
        self.assertEqual(before, _mutation_snapshot(document))

        document, indexes = _synthetic_book()
        toc = document.sections[indexes["toc"]]
        chapter = document.sections[indexes["chapter_two"]]
        toc.footer.is_linked_to_previous = False
        target_part = toc.footer.part
        alias_id = "rIdSyntheticFooterAlias"
        document.part.rels.add_relationship(RT.FOOTER, target_part, alias_id)
        chapter.even_page_footer.is_linked_to_previous = False
        even_ref = next(
            item
            for item in chapter._sectPr.findall(qn("w:footerReference"))
            if item.get(qn("w:type")) == "even"
        )
        even_ref.set(qn("r:id"), alias_id)
        before = _mutation_snapshot(document)
        with self.assertRaisesRegex(FormatMonographError, "conflicting alignment"):
            apply_pagination_sections(
                document,
                _settings(),
                _resolver,
                approved_headings=_approved_chapters(),
            )
        self.assertEqual(before, _mutation_snapshot(document))

    def test_multisection_toc_restart_is_removed_and_audited(self) -> None:
        document = Document()
        document.add_paragraph("Synthetic title")
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Synthetic TOC")
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Synthetic TOC continuation")
        _set_start(document.sections[2], 7)
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Chapter One")
        apply_pagination_sections(document, _settings(), _resolver)
        output = self.root / "multi-toc.docx"
        document.save(output)
        inventory = pagination_inventory(output)
        self.assertEqual("1", inventory["sections"][1]["page_number_start"])
        self.assertIsNone(inventory["sections"][2]["page_number_start"])
        failures, _ = audit_pagination_sections(
            output, Document(output), _settings(), _resolver
        )
        self.assertEqual([], failures)

        broken = Document(output)
        broken.sections[2]._sectPr.find(qn("w:pgNumType")).set(qn("w:start"), "7")
        broken_path = self.root / "multi-toc-broken.docx"
        broken.save(broken_path)
        failures, _ = audit_pagination_sections(
            broken_path, Document(broken_path), _settings(), _resolver
        )
        self.assertIn(
            "unexpected_front_page_number_restart",
            {failure["property"] for failure in failures},
        )

    def test_on_off_values_are_interpreted_consistently(self) -> None:
        false_values = ("0", "false", "off", "no", "FALSE")
        true_values = (None, "1", "true", "on", "yes", "TRUE")
        for value in false_values:
            with self.subTest(setting_false=value):
                document, _ = _synthetic_book()
                document.sections[0].header.paragraphs[0].text = "Running header"
                flag = OxmlElement("w:evenAndOddHeaders")
                flag.set(qn("w:val"), value)
                document.settings.element.append(flag)
                before = _mutation_snapshot(document)
                with self.assertRaisesRegex(FormatMonographError, "header behavior"):
                    apply_pagination_sections(document, _settings(), _resolver)
                self.assertEqual(before, _mutation_snapshot(document))
        for value in true_values:
            with self.subTest(setting_true=value):
                document, _ = _synthetic_book()
                document.sections[0].header.paragraphs[0].text = "Running header"
                flag = OxmlElement("w:evenAndOddHeaders")
                if value is not None:
                    flag.set(qn("w:val"), value)
                document.settings.element.append(flag)
                apply_pagination_sections(
                    document,
                    _settings(),
                    _resolver,
                    approved_headings=_approved_chapters(),
                )

        for value in false_values:
            with self.subTest(title_false=value):
                document, indexes = _synthetic_book()
                title_page = document.sections[indexes["body"]]._sectPr.find(qn("w:titlePg"))
                title_page.set(qn("w:val"), value)
                apply_pagination_sections(
                    document,
                    _settings(),
                    _resolver,
                    approved_headings=_approved_chapters(),
                )
                output = self.root / f"title-false-{value}.docx"
                document.save(output)
                inventory = pagination_inventory(output)
                self.assertFalse(inventory["sections"][indexes["body"]]["different_first_page"])
                self.assertEqual(
                    0,
                    inventory["sections"][indexes["body"]]["footer_page_field_counts"].get(
                        "first", 0
                    ),
                )
        for ordinal, value in enumerate(true_values):
            with self.subTest(title_true=value):
                document, indexes = _synthetic_book()
                title_page = document.sections[indexes["body"]]._sectPr.find(qn("w:titlePg"))
                if value is None:
                    title_page.attrib.pop(qn("w:val"), None)
                else:
                    title_page.set(qn("w:val"), value)
                apply_pagination_sections(
                    document,
                    _settings(),
                    _resolver,
                    approved_headings=_approved_chapters(),
                )
                output = self.root / f"title-true-{ordinal}.docx"
                document.save(output)
                inventory = pagination_inventory(output)
                self.assertTrue(inventory["sections"][indexes["body"]]["different_first_page"])
                self.assertEqual(
                    1,
                    inventory["sections"][indexes["body"]]["footer_page_field_counts"]["first"],
                )

    def test_complex_page_rewrite_accepts_only_balanced_unmixed_field(self) -> None:
        def append_simple_page(item) -> None:
            field = OxmlElement("w:fldSimple")
            field.set(qn("w:instr"), "PAGE")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = "1"
            run.append(text)
            field.append(run)
            item.paragraphs[0]._p.append(field)

        document, indexes = _synthetic_book()
        footer = document.sections[indexes["toc"]].footer
        footer.is_linked_to_previous = False
        _append_complex_page(footer.paragraphs[0])
        apply_pagination_sections(
            document,
            _settings(),
            _resolver,
            approved_headings=_approved_chapters(),
        )
        self.assertEqual(1, len(footer._element.xpath(".//w:fldSimple")))
        self.assertFalse(footer._element.xpath(".//w:fldChar | .//w:instrText"))

        for name, mutate in {
            "unbalanced": lambda item: _append_complex_page(
                item.paragraphs[0], include_end=False
            ),
            "mixed": lambda item: (
                _append_complex_page(item.paragraphs[0]),
                item.add_paragraph("Chapter text"),
            ),
            "two_complex": lambda item: (
                _append_complex_page(item.paragraphs[0]),
                _append_complex_page(item.paragraphs[0]),
            ),
            "two_simple": lambda item: (
                append_simple_page(item),
                append_simple_page(item),
            ),
            "simple_and_complex": lambda item: (
                append_simple_page(item),
                _append_complex_page(item.paragraphs[0]),
            ),
        }.items():
            with self.subTest(name=name):
                document, indexes = _synthetic_book()
                footer = document.sections[indexes["toc"]].footer
                footer.is_linked_to_previous = False
                mutate(footer)
                before = _mutation_snapshot(document)
                with self.assertRaisesRegex(FormatMonographError, "footer"):
                    apply_pagination_sections(
                        document,
                        _settings(),
                        _resolver,
                        approved_headings=_approved_chapters(),
                    )
                self.assertEqual(before, _mutation_snapshot(document))

    def test_unapproved_title_page_section_is_not_treated_as_chapter(self) -> None:
        document, indexes = _synthetic_book()
        document.sections[indexes["landscape"]].different_first_page_header_footer = True
        before = _mutation_snapshot(document)
        with self.assertRaisesRegex(FormatMonographError, "not an approved"):
            apply_pagination_sections(
                document,
                _settings(),
                _resolver,
                approved_headings=_approved_chapters(),
            )
        self.assertEqual(before, _mutation_snapshot(document))

    def test_separate_title_page_cannot_be_numbered_front_start(self) -> None:
        cases = (("same", "Synthetic title", True), ("later", "Chapter One", False))
        for name, title_text, stray_start in cases:
            with self.subTest(name=name):
                document, indexes = _synthetic_book()
                if stray_start:
                    _set_start(document.sections[indexes["title"]], 1)
                before = _mutation_snapshot(document)
                with self.assertRaisesRegex(FormatMonographError, "must precede"):
                    apply_pagination_sections(
                        document,
                        _settings(),
                        _resolver,
                        front_matter={
                            "approved": True,
                            "separate_title_page": True,
                            "book_title": {"text": title_text},
                        },
                        approved_headings=_approved_chapters(),
                    )
                self.assertEqual(before, _mutation_snapshot(document))

    def test_approved_toc_heading_is_the_existing_section_anchor(self) -> None:
        document = Document()
        document.add_paragraph("Synthetic title")
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("TOC Heading")
        document.add_paragraph("Synthetic TOC")
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Chapter One")
        result = apply_pagination_sections(
            document,
            _settings(),
            _resolver,
            front_matter={
                "approved": True,
                "toc_heading": {"text": "TOC Heading"},
            },
        )
        self.assertEqual(1, result["toc_section"])
        self.assertEqual(2, result["body_section"])

    def test_enabling_odd_even_does_not_change_existing_running_content(self) -> None:
        for kind in ("header", "unnumbered_footer"):
            with self.subTest(kind=kind):
                document, indexes = _synthetic_book()
                if kind == "header":
                    document.sections[0].header.paragraphs[0].text = "Running header"
                else:
                    document.sections[0].footer.paragraphs[0].text = "Title footer"
                    for footer in (
                        document.sections[indexes["toc"]].footer,
                        document.sections[indexes["toc"]].even_page_footer,
                    ):
                        footer.is_linked_to_previous = False
                before = _mutation_snapshot(document)
                with self.assertRaisesRegex(FormatMonographError, "odd/even"):
                    apply_pagination_sections(document, _settings(), _resolver)
                self.assertEqual(before, _mutation_snapshot(document))

    def test_structure_map_preflights_footer_before_other_mutations(self) -> None:
        source = self.root / "structure-preflight.docx"
        document, indexes = _synthetic_book()
        document.save(source)
        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        structure["pagination_sections"]["approved"] = True
        toc_paragraph = next(
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text == "Synthetic TOC"
        )
        structure["pagination_sections"]["toc_start"] = {
            "kind": "body_paragraph",
            "paragraph": toc_paragraph,
            "text_sha256": text_sha256("Synthetic TOC"),
        }
        loaded = Document(source)
        loaded.sections[indexes["toc"]].footer.is_linked_to_previous = False
        loaded.sections[indexes["toc"]].footer.paragraphs[0].text = "Protected"
        prime_structure_map_locators(loaded, structure)
        before = _mutation_snapshot(loaded)
        with self.assertRaisesRegex(FormatMonographError, "footer"):
            apply_structure_map(loaded, structure)
        self.assertEqual(before, _mutation_snapshot(loaded))


if __name__ == "__main__":
    unittest.main()
