from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from jsonschema import Draft202012Validator

REPO = Path(__file__).resolve().parents[1]
SCRIPTS = REPO / "format-monograph" / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import FormatMonographError, apply_table_properties  # noqa: E402
from audit_docx import audit_table_rule  # noqa: E402
from docx_pagination import audit_pagination_sections, pagination_inventory  # noqa: E402
from finalize_docx import external_refresh  # noqa: E402
from structure_map import (  # noqa: E402
    approved_data_tables,
    apply_structure_map,
    candidate_structure_map,
    load_structure_map,
    prime_structure_map_locators,
    resolve_paragraph_locator,
    structure_content_fingerprint,
)


class V025WordPaginationTableTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.source = self.root / "synthetic-v025.docx"
        document = Document()
        document.add_paragraph("[[TOC]]")
        document.add_paragraph("Synthetic TOC entry")
        document.add_paragraph("第1章 Synthetic body")
        table = document.add_table(rows=2, cols=4)
        for index, value in enumerate(("Value", "Unit", "Code", "Description")):
            table.cell(0, index).text = value
        for index, value in enumerate(("12.5", "kN", "A1", "Narrative value")):
            table.cell(1, index).text = value
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("第2章 Continued body")
        document.save(self.source)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def approved_map(self) -> dict:
        structure = candidate_structure_map(self.source)
        self.assertEqual("1.5", structure["schema_version"])
        structure["status"] = "approved"
        pagination = structure["pagination_sections"]
        self.assertIsNotNone(pagination["toc_start"])
        self.assertIsNotNone(pagination["body_start"])
        pagination["approved"] = True
        table = structure["tables"][0]
        table.update(
            {
                "approved": True,
                "kind": "data",
                "header_rows": [0],
                "repeat_header_rows": [0],
                "prevent_normal_row_split": True,
            }
        )
        table["visual"].update(
            {
                "approved": True,
                "preferred_column_widths_percent": [20, 15, 15, 50],
                "allow_autofit": False,
                "border_preset": "three_line",
                "column_roles": ["numeric", "unit", "short_code", "narrative"],
                "orientation": "landscape",
                "landscape_approved": True,
            }
        )
        return structure

    def build_formatted(self) -> tuple[dict, Path, list[dict]]:
        structure = self.approved_map()
        map_path = self.root / "approved-map.json"
        map_path.write_text(json.dumps(structure), encoding="utf-8")
        structure = load_structure_map(map_path)
        document = Document(self.source)
        prime_structure_map_locators(document, structure)
        changes = apply_structure_map(document, structure)
        apply_table_properties(
            document,
            {
                "alignment": "center",
                "repeat_header_row": True,
                "prevent_row_split": True,
                "available_width_percent": 100,
                "cell_margins_mm": {"top": 1, "right": 1.5, "bottom": 1, "left": 1.5},
                "vertical_alignment": "center",
                "font_name_ascii": "Times New Roman",
                "font_name_east_asia": "宋体",
                "font_size_pt": 9,
                "line_spacing_pt": 15,
                "header_bold": True,
            },
            approved_data_tables(document, structure),
        )
        output = self.root / "formatted.docx"
        document.save(output)
        return structure, output, changes

    def test_dual_pagination_and_table_visuals_are_structural_and_editable(self) -> None:
        structure, output, changes = self.build_formatted()

        inventory = pagination_inventory(output)
        self.assertTrue(inventory["odd_and_even_pages_header_footer"])
        self.assertEqual([0, 1], inventory["page_number_restarts"])
        self.assertFalse(inventory["orphan_header_footer_parts"])
        for section in inventory["sections"]:
            self.assertFalse(section["different_first_page"])
            self.assertFalse(section["missing_page_footer_types"])

        formatted = Document(output)
        failures, _ = audit_pagination_sections(
            output,
            formatted,
            structure["pagination_sections"],
            resolve_paragraph_locator,
        )
        self.assertFalse(failures)
        table = formatted.tables[0]
        self.assertFalse(table.autofit)
        self.assertIsNotNone(table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")))
        self.assertTrue(
            any(
                change["kind"] == "structure_pagination_sections"
                for change in changes
            )
        )
        table_failures = audit_table_rule(
            formatted,
            {
                "properties": {
                    "alignment": "center",
                    "available_width_percent": 100,
                    "preferred_column_widths_percent": [20, 15, 15, 50],
                    "allow_autofit": False,
                    "cell_margins_mm": {
                        "top": 1,
                        "right": 1.5,
                        "bottom": 1,
                        "left": 1.5,
                    },
                    "vertical_alignment": "center",
                    "border_preset": "three_line",
                }
            },
            approved_data_tables(formatted, structure),
        )
        self.assertFalse(table_failures)
        self.assertEqual(
            structure_content_fingerprint(self.source, structure),
            structure_content_fingerprint(output, structure),
        )

    def test_missing_even_footer_and_body_restart_fail_audit(self) -> None:
        structure, output, _ = self.build_formatted()
        document = Document(output)
        for section in document.sections:
            section.even_page_footer.is_linked_to_previous = True
        later = list(document.sections)[-1]._sectPr
        page_number = later.find(qn("w:pgNumType"))
        if page_number is None:
            page_number = OxmlElement("w:pgNumType")
            later.append(page_number)
        page_number.set(qn("w:start"), "1")
        broken = self.root / "broken-pagination.docx"
        document.save(broken)
        failures, _ = audit_pagination_sections(
            broken,
            Document(broken),
            structure["pagination_sections"],
            resolve_paragraph_locator,
        )
        properties = {item["property"] for item in failures}
        self.assertIn("visible_page_footer_types", properties)
        self.assertIn("unexpected_page_number_restart", properties)

    def test_external_protocol_is_strict_and_does_not_overwrite_input(self) -> None:
        helper = self.root / "backend.py"
        helper.write_text(
            "import json, shutil, sys\n"
            "request = json.load(sys.stdin)\n"
            "shutil.copy2(request['input_path'], request['output_path'])\n"
            "print(json.dumps({'status':'success','backend':'test_word','software':'Microsoft Word',"
            "'repaginated':True,'saved':True,'field_cache_verified':True,"
            "'structural_changes_applied':0,'updated_field_types':[]}))\n",
            encoding="utf-8",
        )
        output = self.root / "external.docx"
        response = external_refresh(
            self.source,
            output,
            json.dumps([sys.executable, str(helper)]),
            self.root / "profile.json",
            self.root / "map.json",
            None,
            "Microsoft Word",
        )
        self.assertEqual("test_word", response["backend"])
        self.assertTrue(output.is_file())
        self.assertNotEqual(self.source.resolve(), output.resolve())

        helper.write_text(
            helper.read_text(encoding="utf-8").replace(
                "'updated_field_types':[]", "'updated_field_types':['LINK']"
            ),
            encoding="utf-8",
        )
        with self.assertRaises(FormatMonographError):
            external_refresh(
                self.source,
                self.root / "rejected.docx",
                json.dumps([sys.executable, str(helper)]),
                self.root / "profile.json",
                self.root / "map.json",
                None,
                "Microsoft Word",
            )

    def test_target_pdf_can_be_rendered_without_libreoffice_conversion(self) -> None:
        import fitz

        pdf_path = self.root / "word-export.pdf"
        pdf = fitz.open()
        page = pdf.new_page()
        page.insert_text((72, 72), "Synthetic target PDF")
        pdf.save(pdf_path)
        pdf.close()
        output_dir = self.root / "rendered"
        completed = subprocess.run(
            [
                sys.executable,
                str(SCRIPTS / "render_docx.py"),
                str(self.source),
                "--output-dir",
                str(output_dir),
                "--target-pdf",
                str(pdf_path),
                "--target-software",
                "Microsoft Word",
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(0, completed.returncode, completed.stdout + completed.stderr)
        manifest = json.loads((output_dir / "render-manifest.json").read_text())
        self.assertEqual(1, manifest["page_count"])
        self.assertFalse(manifest["target_layout_unverified"])
        self.assertEqual("pending", manifest["visual_review"])

    def test_v025_profile_is_valid_and_old_profile_is_preserved(self) -> None:
        schema = json.loads(
            (REPO / "format-monograph" / "references" / "format-profile.schema.json").read_text(
                encoding="utf-8"
            )
        )
        profile = json.loads(
            (
                REPO
                / "format-monograph"
                / "examples"
                / "profiles"
                / "technical-textbook-layout.v0.2.5.draft.json"
            ).read_text(encoding="utf-8")
        )
        self.assertFalse(list(Draft202012Validator(schema).iter_errors(profile)))
        self.assertEqual("draft", profile["approval"]["status"])
        self.assertFalse(
            profile["rules"][0]["properties"]["different_first_page_header_footer"]
        )
        self.assertTrue(
            (
                REPO
                / "format-monograph"
                / "examples"
                / "profiles"
                / "technical-textbook-layout.v0.2.draft.json"
            ).is_file()
        )

    def test_word_adapter_keeps_com_out_of_core(self) -> None:
        adapter = (
            REPO
            / "adapters"
            / "microsoft-word"
            / "windows"
            / "word_field_updater.ps1"
        )
        text = adapter.read_text(encoding="utf-8")
        self.assertIn("AutomationSecurity = 3", text)
        self.assertIn("UpdateLinksAtOpen = $false", text)
        self.assertIn(
            "recent-file",
            (adapter.parent / "INSTALL.md").read_text(encoding="utf-8"),
        )
        for script in (REPO / "format-monograph" / "scripts").glob("*.py"):
            core = script.read_text(encoding="utf-8").casefold()
            self.assertNotIn("new-object -comobject word.application", core)


if __name__ == "__main__":
    unittest.main()
