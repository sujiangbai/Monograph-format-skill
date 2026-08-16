from __future__ import annotations

import base64
import copy
import hashlib
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "format-monograph" / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import NS, FormatMonographError  # noqa: E402
from field_writeback import parse_fields, selective_field_result_writeback  # noqa: E402
from structure_map import (  # noqa: E402
    apply_structure_map,
    audit_structure_image_operations,
    audit_structure_toc_source_operations,
    candidate_structure_map,
    toc_result_contract,
    validate_structure_map_source,
)


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def add_complex_field(paragraph, instruction: str, value: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instruction_run = paragraph.add_run()
    instruction_node = OxmlElement("w:instrText")
    instruction_node.text = f" {instruction} "
    instruction_run._r.append(instruction_node)
    separator_run = paragraph.add_run()
    separator = OxmlElement("w:fldChar")
    separator.set(qn("w:fldCharType"), "separate")
    separator_run._r.append(separator)
    paragraph.add_run(value)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_toc_placeholder(paragraph) -> None:
    add_complex_field(paragraph, 'TOC \\o "1-4" \\h \\z', "Update TOC")


def add_toc_result(
    document: Document,
    entries: list[tuple[int, str, str]],
    *,
    add_non_text: bool = False,
) -> None:
    paragraphs = []
    for level, title, page in entries:
        style_name = f"TOC {level}"
        if style_name not in document.styles:
            document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        paragraph = document.add_paragraph(style=style_name)
        paragraphs.append(paragraph)
        if len(paragraphs) == 1:
            begin_run = paragraph.add_run()
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            begin_run._r.append(begin)
            instruction_run = paragraph.add_run()
            instruction = OxmlElement("w:instrText")
            instruction.text = ' TOC \\f M \\h \\z '
            instruction_run._r.append(instruction)
            separator_run = paragraph.add_run()
            separator = OxmlElement("w:fldChar")
            separator.set(qn("w:fldCharType"), "separate")
            separator_run._r.append(separator)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), f"_TocSynthetic{len(paragraphs)}")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = title
        run.append(text)
        hyperlink.append(run)
        leading_tab = paragraph.add_run()
        leading_tab._r.append(OxmlElement("w:tab"))
        paragraph._p.append(hyperlink)
        tab_run = paragraph.add_run()
        tab_run._r.append(OxmlElement("w:tab"))
        add_complex_field(paragraph, f"PAGEREF _TocSynthetic{len(paragraphs)}", page)
        if add_non_text and len(paragraphs) == 1:
            paragraph._p.append(OxmlElement("w:drawing"))
    end_run = paragraphs[-1].add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_tc_source(paragraph, text: str, level: int) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f'TC "{text}" \\f M \\l "{level}"')
    field.set(qn("w:dirty"), "true")
    run = OxmlElement("w:r")
    run.append(OxmlElement("w:t"))
    field.append(run)
    paragraph._p.append(field)


def rewrite_package(source: Path, output: Path, transform) -> None:
    with zipfile.ZipFile(source) as package, zipfile.ZipFile(
        output, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for info in package.infolist():
            target.writestr(info, transform(info.filename, package.read(info.filename)))


def change_targets(changes: list[dict[str, object]], kind: str) -> int:
    return int(next((item["targets"] for item in changes if item["kind"] == kind), 0))


class V033TocImageIntegrityTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.image_path = self.root / "synthetic.png"
        self.image_path.write_bytes(PNG_1X1)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_inline_image_fixed_line_spacing_is_repaired_without_resizing(self) -> None:
        source = self.root / "inline-source.docx"
        document = Document()
        document.add_paragraph("Synthetic preceding text")
        image_paragraph = document.add_paragraph()
        image_paragraph.add_run().add_picture(str(self.image_path), width=Inches(1))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "120")
        spacing.set(qn("w:lineRule"), "exact")
        image_paragraph._p.get_or_add_pPr().append(spacing)
        document.add_paragraph("Synthetic following text")
        document.save(source)

        structure = candidate_structure_map(source)
        image = structure["images"][0]
        self.assertEqual(["auto_single_line_spacing"], image["visibility"]["action"])
        image["visibility"]["approved"] = True
        structure["status"] = "approved"
        validate_structure_map_source(source, structure)

        formatted = Document(source)
        body_before = list(formatted.element.body)
        extent_before = copy.deepcopy(
            formatted.element.body.xpath(".//wp:inline/wp:extent")[0].attrib
        )
        changes = apply_structure_map(formatted, structure)
        self.assertEqual(body_before, list(formatted.element.body))
        self.assertEqual(
            extent_before,
            formatted.element.body.xpath(".//wp:inline/wp:extent")[0].attrib,
        )
        direct_spacing = formatted.paragraphs[1]._p.pPr.find(qn("w:spacing"))
        self.assertEqual("240", direct_spacing.get(qn("w:line")))
        self.assertEqual("auto", direct_spacing.get(qn("w:lineRule")))
        self.assertEqual(
            1,
            change_targets(changes, "structure_image_visibility"),
        )
        second = apply_structure_map(formatted, structure)
        self.assertEqual(
            0,
            change_targets(second, "structure_image_visibility"),
        )
        self.assertFalse(audit_structure_image_operations(formatted, structure))

    def test_simple_fixed_table_row_is_relaxed_but_mixed_image_is_blocked(self) -> None:
        source = self.root / "table-source.docx"
        document = Document()
        table = document.add_table(rows=1, cols=1)
        paragraph = table.cell(0, 0).paragraphs[0]
        paragraph.add_run().add_picture(str(self.image_path), width=Inches(1))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "100")
        spacing.set(qn("w:lineRule"), "exact")
        paragraph._p.get_or_add_pPr().append(spacing)
        height = OxmlElement("w:trHeight")
        height.set(qn("w:val"), "100")
        height.set(qn("w:hRule"), "exact")
        table.rows[0]._tr.get_or_add_trPr().append(height)
        document.save(source)

        structure = candidate_structure_map(source)
        image = structure["images"][0]
        self.assertEqual(
            {"auto_single_line_spacing", "relax_exact_table_row_height"},
            set(image["visibility"]["action"]),
        )
        image["visibility"]["approved"] = True
        structure["status"] = "approved"
        validate_structure_map_source(source, structure)
        formatted = Document(source)
        apply_structure_map(formatted, structure)
        row_height = formatted.tables[0].rows[0]._tr.trPr.find(qn("w:trHeight"))
        self.assertEqual("atLeast", row_height.get(qn("w:hRule")))
        self.assertFalse(audit_structure_image_operations(formatted, structure))

        mixed_source = self.root / "mixed-source.docx"
        mixed = Document()
        paragraph = mixed.add_paragraph("Authored text ")
        paragraph.add_run().add_picture(str(self.image_path), width=Inches(1))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "100")
        spacing.set(qn("w:lineRule"), "exact")
        paragraph._p.get_or_add_pPr().append(spacing)
        mixed.save(mixed_source)
        candidate = candidate_structure_map(mixed_source)["images"][0]["visibility"]
        self.assertEqual([], candidate["action"])
        self.assertEqual("mixed_text", candidate["paragraph_payload"])
        self.assertIsNotNone(candidate["blocked_reason"])

    def test_image_outline_pollution_switches_whole_toc_to_plain_text_tc(self) -> None:
        source = self.root / "toc-source.docx"
        document = Document()
        toc = document.add_paragraph()
        add_toc_placeholder(toc)
        document.add_paragraph("第1章 Synthetic chapter", style="Heading 1")
        document.add_paragraph("1.1 Synthetic section", style="Heading 2")
        polluted = document.add_paragraph(style="Heading 2")
        polluted.add_run().add_picture(str(self.image_path), width=Inches(1))
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        for heading in structure["headings"]:
            heading["approved"] = True
        structure["toc_source"]["approved"] = True
        self.assertTrue(structure["toc_source"]["contaminants"])

        formatted = Document(source)
        first = apply_structure_map(formatted, structure)
        toc_change = next(
            item for item in first if item["kind"] == "structure_toc_source"
        )
        self.assertEqual("tc_plain_text", toc_change["details"]["mode"])
        records = parse_fields(formatted.element)
        tc_records = [item for item in records if item.field_type == "TC"]
        self.assertEqual(2, len(tc_records))
        self.assertTrue(
            all(
                item.form == "complex"
                and item.separate is None
                and item.begin is not None
                and item.begin.getparent().find(qn("w:rPr")) is None
                for item in tc_records
            )
        )
        self.assertTrue(
            all(
                item.begin.getparent().getparent().xpath(
                    ".//w:bookmarkStart[starts-with(@w:name, '_Toc')]"
                )
                for item in tc_records
            )
        )
        self.assertIn("TOC \\f M", next(item.instruction for item in records if item.field_type == "TOC"))
        polluted_element = formatted.paragraphs[3]._p
        self.assertFalse(
            any(
                record.field_type == "TC"
                and (
                    record.simple is not None
                    and record.simple.getparent() is polluted_element
                    or record.begin is not None
                    and record.begin.getparent().getparent() is polluted_element
                )
                for record in parse_fields(formatted.element)
            )
        )
        contract = toc_result_contract(formatted, structure)
        self.assertEqual([1, 2], [item["level"] for item in contract])
        self.assertTrue(all(set(item) == {"level", "kind", "text_sha256"} for item in contract))
        self.assertFalse(audit_structure_toc_source_operations(formatted, structure))

        snapshot = etree.tostring(formatted.element)
        second = apply_structure_map(formatted, structure)
        self.assertEqual(snapshot, etree.tostring(formatted.element))
        self.assertEqual(
            0,
            change_targets(second, "structure_toc_source"),
        )

    def test_clean_heading_sources_keep_heading_style_toc(self) -> None:
        source = self.root / "clean-toc-source.docx"
        document = Document()
        add_toc_placeholder(document.add_paragraph())
        document.add_paragraph("第1章 Synthetic chapter", style="Heading 1")
        document.add_paragraph("1.1 Synthetic section", style="Heading 2")
        document.save(source)
        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        for heading in structure["headings"]:
            heading["approved"] = True
        structure["toc_source"]["approved"] = True
        formatted = Document(source)
        changes = apply_structure_map(formatted, structure)
        toc_change = next(
            item for item in changes if item["kind"] == "structure_toc_source"
        )
        self.assertEqual("heading_styles", toc_change["details"]["mode"])
        self.assertFalse(any(item.field_type == "TC" for item in parse_fields(formatted.element)))
        self.assertFalse(audit_structure_toc_source_operations(formatted, structure))

    def _toc_package(self, path: Path, *, non_text: bool = False) -> list[dict[str, object]]:
        entries = [(1, "第1章 Synthetic chapter", "1"), (2, "1.1 Synthetic section", "2")]
        document = Document()
        add_toc_result(document, entries, add_non_text=non_text)
        first = document.add_paragraph("第1章 Synthetic chapter")
        second = document.add_paragraph("1.1 Synthetic section")
        add_tc_source(first, entries[0][1], 1)
        add_tc_source(second, entries[1][1], 2)
        document.save(path)
        return [
            {
                "level": level,
                "kind": "heading",
                "text_sha256": hashlib.sha256(
                    " ".join(title.split(maxsplit=1)[1:]).encode("utf-8")
                ).hexdigest(),
            }
            for level, title, _page in entries
        ]

    def test_selective_writeback_accepts_only_matching_text_toc_and_tc_sources(self) -> None:
        baseline = self.root / "baseline.docx"
        contract = self._toc_package(baseline)
        refreshed = self.root / "refreshed.docx"
        rewrite_package(baseline, refreshed, lambda _name, data: data)
        output = self.root / "output.docx"
        report = selective_field_result_writeback(
            baseline,
            refreshed,
            output,
            allowed_field_types={"TOC", "PAGEREF", "TC"},
            toc_contract=contract,
        )
        self.assertEqual("selective_verified", report["status"])
        self.assertEqual(2, report["approved_source_fields"])
        self.assertEqual("verified_text_only", report["toc_result_status"])
        self.assertEqual(2, report["toc_source_count"])

        mismatch = self.root / "mismatch.docx"

        def change_title(name: str, data: bytes) -> bytes:
            if name != "word/document.xml":
                return data
            root = etree.fromstring(data)
            title = root.xpath(
                ".//w:p[w:pPr/w:pStyle[@w:val='TOC1']]//w:hyperlink//w:t",
                namespaces=NS,
            )[0]
            title.text = "第1章 Wrong title"
            return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

        rewrite_package(baseline, mismatch, change_title)
        with self.assertRaisesRegex(FormatMonographError, "text or order"):
            selective_field_result_writeback(
                baseline,
                mismatch,
                self.root / "mismatch-output.docx",
                allowed_field_types={"TOC", "PAGEREF", "TC"},
                toc_contract=contract,
            )

    def test_selective_writeback_rejects_non_text_toc_payload(self) -> None:
        baseline = self.root / "clean.docx"
        contract = self._toc_package(baseline)
        for local_name in ("drawing", "pict", "object", "txbxContent", "tbl"):
            with self.subTest(local_name=local_name):
                refreshed = self.root / f"non-text-{local_name}.docx"

                def add_object(name: str, data: bytes, selected=local_name) -> bytes:
                    if name != "word/document.xml":
                        return data
                    root = etree.fromstring(data)
                    paragraph = root.xpath(
                        ".//w:p[w:pPr/w:pStyle[@w:val='TOC1']]", namespaces=NS
                    )[0]
                    paragraph.append(etree.Element(qn(f"w:{selected}")))
                    return etree.tostring(
                        root,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    )

                rewrite_package(baseline, refreshed, add_object)
                with self.assertRaisesRegex(FormatMonographError, "non-text"):
                    selective_field_result_writeback(
                        baseline,
                        refreshed,
                        self.root / f"non-text-{local_name}-output.docx",
                        allowed_field_types={"TOC", "PAGEREF", "TC"},
                        toc_contract=contract,
                    )

    def test_selective_writeback_rejects_empty_and_extra_toc_entries(self) -> None:
        baseline = self.root / "toc-cardinality.docx"
        contract = self._toc_package(baseline)

        def empty_title(name: str, data: bytes) -> bytes:
            if name != "word/document.xml":
                return data
            root = etree.fromstring(data)
            root.xpath(
                ".//w:p[w:pPr/w:pStyle[@w:val='TOC1']]//w:hyperlink//w:t",
                namespaces=NS,
            )[0].text = ""
            return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

        empty = self.root / "toc-empty.docx"
        rewrite_package(baseline, empty, empty_title)
        with self.assertRaisesRegex(FormatMonographError, "empty entry"):
            selective_field_result_writeback(
                baseline,
                empty,
                self.root / "toc-empty-output.docx",
                allowed_field_types={"TOC", "PAGEREF", "TC"},
                toc_contract=contract,
            )

        def duplicate_entry(name: str, data: bytes) -> bytes:
            if name != "word/document.xml":
                return data
            root = etree.fromstring(data)
            entries = root.xpath(
                ".//w:p[w:pPr/w:pStyle[starts-with(@w:val, 'TOC')]]",
                namespaces=NS,
            )
            duplicate = copy.deepcopy(entries[-1])
            outer_end = duplicate.xpath(
                ".//w:fldChar[@w:fldCharType='end']", namespaces=NS
            )[-1]
            outer_end.getparent().remove(outer_end)
            entries[-1].addprevious(duplicate)
            return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

        extra = self.root / "toc-extra.docx"
        rewrite_package(baseline, extra, duplicate_entry)
        with self.assertRaisesRegex(FormatMonographError, "entry count"):
            selective_field_result_writeback(
                baseline,
                extra,
                self.root / "toc-extra-output.docx",
                allowed_field_types={"TOC", "PAGEREF", "TC"},
                toc_contract=contract,
            )


if __name__ == "__main__":
    unittest.main()
