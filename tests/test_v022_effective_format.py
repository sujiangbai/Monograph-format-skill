from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "format-monograph"
SCRIPTS = SKILL / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import (  # noqa: E402
    apply_style_rule_to_paragraphs,
    missing_profile_fonts,
    resolve_font_name,
)
from audit_docx import audit_paragraph_rule, heading_numbering_start  # noqa: E402
from structure_map import candidate_structure_map, load_structure_map  # noqa: E402
from test_v11_execution import approved_v11_profile  # noqa: E402


class V022EffectiveFormatTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.source = self.root / "synthetic-chapter.docx"

        document = Document()
        title = document.add_paragraph("Synthetic title")
        title.runs[0].font.size = Pt(24)
        document.add_paragraph("第4章 Synthetic chapter")
        section = document.add_paragraph("4.1 Synthetic section")
        section.runs[0].font.size = Pt(30)
        body = document.add_paragraph("Synthetic body text.")
        body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        body.runs[0].font.name = "Courier New"
        body.runs[0].font.size = Pt(20)
        body.runs[0].font.bold = True
        body.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        table = document.add_table(rows=3, cols=2)
        caption = table.cell(0, 0).merge(table.cell(0, 1))
        caption.text = "表 4.1-1 Synthetic table"
        table.cell(1, 0).text = "Header A"
        table.cell(1, 1).text = "Header B"
        table.cell(2, 0).text = "Value A"
        table.cell(2, 1).text = "Value B"
        layout = document.add_table(rows=1, cols=1)
        layout.cell(0, 0).text = "Synthetic layout container"

        caption_style = document.styles["Caption"]
        caption_style.element.getparent().remove(caption_style.element)
        document.save(self.source)

        structure = candidate_structure_map(self.source)
        structure["schema_version"] = "1.1"
        structure["status"] = "approved"
        structure["numbering"].update(
            {"approved": True, "chapter_start": 4, "mode": "single_chapter"}
        )
        for entry in structure["headings"]:
            entry["approved"] = entry["paragraph"] in {1, 2}
        for entry in structure["paragraph_roles"]:
            locator = entry["locator"]
            if locator["kind"] == "body_paragraph" and locator["paragraph"] == 3:
                entry.update({"role": "body", "approved": True})
            elif entry["role"] in {"heading_1", "heading_2", "table_caption"}:
                entry["approved"] = True
        for entry in structure["captions"]:
            entry.update(
                {
                    "approved": True,
                    "migrate_outside_table": True,
                    "completeness": "complete",
                    "hierarchy_status": "match",
                }
            )
        structure["tables"][0].update(
            {
                "kind": "data",
                "approved": True,
                "caption_row": 0,
                "header_rows": [1],
                "repeat_header_rows": [1],
                "prevent_normal_row_split": True,
            }
        )
        for entry in structure["trailing_empty_sections"]:
            entry["approved_delete"] = False

        self.structure = self.root / "structure.json"
        self.structure.write_text(
            json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        profile = approved_v11_profile()
        profile["rules"] = [
            rule
            for rule in profile["rules"]
            if rule["selector"]["kind"]
            in {"paragraph_role", "table_role", "field_role"}
        ]
        body_rule = next(
            rule for rule in profile["rules"] if rule["category"] == "body"
        )
        body_rule["properties"]["bold"] = False
        body_rule["properties"].pop("font_name_east_asia", None)
        body_rule["properties"]["font_name_ascii"] = "Arial"
        heading_rule = next(
            rule for rule in profile["rules"] if rule["category"] == "heading"
        )
        heading_two = json.loads(json.dumps(heading_rule))
        heading_two["id"] = "FMT-HEAD-902"
        heading_two["selector"]["value"] = "level_2_section"
        heading_two["properties"]["font_size_pt"] = 14
        profile["rules"].append(heading_two)
        profile["rules"].append(
            {
                "id": "FMT-CAP-902",
                "category": "caption",
                "selector": {"kind": "caption_role", "value": "all"},
                "properties": {
                    "font_name_ascii": "Arial",
                    "font_size_pt": 9,
                    "alignment": "center",
                },
                "source_ids": ["SRC-901"],
                "evidence_summary": "Synthetic approved caption rule.",
                "confidence": "high",
                "status": "approved",
                "application": "automatic",
            }
        )
        self.profile = self.root / "profile.json"
        self.profile.write_text(
            json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        self.output = self.root / "out"

    def tearDown(self) -> None:
        self.temp.cleanup()

    def run_script(self, name: str, *args: object) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(SCRIPTS / name), *(str(value) for value in args)],
            cwd=SKILL,
            capture_output=True,
            text=True,
            check=False,
        )

    def test_structure_map_11_contains_roles_without_authored_text(self) -> None:
        value = self.structure.read_text(encoding="utf-8")
        self.assertIn('"schema_version": "1.1"', value)
        self.assertIn('"paragraph_roles"', value)
        self.assertIn('"chapter_start": 4', value)
        self.assertNotIn("Synthetic body text", value)
        self.assertNotIn("Synthetic table", value)

    def test_font_alias_is_not_reported_as_substitution(self) -> None:
        resolution = resolve_font_name("宋体", {"SimSun"})
        self.assertTrue(resolution["available"])
        self.assertEqual("verified_alias", resolution["match"])
        profile = approved_v11_profile()
        for rule in profile["rules"]:
            if rule["category"] == "body":
                rule["properties"] = {"font_name_east_asia": "宋体"}
            elif "font_name_east_asia" in rule.get("properties", {}):
                rule["properties"].pop("font_name_east_asia")
        with patch("_common.available_font_names", return_value={"SimSun"}):
            self.assertNotIn("宋体", missing_profile_fonts(profile))

    def test_effective_format_numbering_and_complex_table(self) -> None:
        source_bytes = self.source.read_bytes()
        result = self.run_script(
            "apply_profile.py",
            self.source,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure,
            "--output-dir",
            self.output,
            "--allow-missing-fonts",
        )
        self.assertEqual(0, result.returncode, result.stdout + result.stderr)
        self.assertEqual(source_bytes, self.source.read_bytes())

        formatted = self.output / "synthetic-chapter-formatted.docx"
        document = Document(formatted)
        self.assertEqual(24, document.paragraphs[0].runs[0].font.size.pt)
        self.assertEqual("Synthetic chapter", document.paragraphs[1].text)
        self.assertEqual("Synthetic section", document.paragraphs[2].text)
        self.assertEqual(4, heading_numbering_start(document))

        body = document.paragraphs[3]
        self.assertIsNone(body.runs[0].font.size)
        self.assertIsNone(body.runs[0].font.bold)
        self.assertEqual("C00000", str(body.runs[0].font.color.rgb))
        self.assertIsNone(body.paragraph_format.alignment)

        self.assertEqual(2, len(document.tables[0].rows))
        first_row_props = document.tables[0].rows[0]._tr.get_or_add_trPr()
        self.assertIsNotNone(first_row_props.find(qn("w:tblHeader")))
        self.assertIsNotNone(first_row_props.find(qn("w:cantSplit")))
        table_xml = document.tables[0]._tbl
        caption = table_xml.getprevious()
        self.assertEqual(qn("w:p"), caption.tag)
        self.assertEqual("Caption", document.paragraphs[4].style.name)
        layout_props = document.tables[1].rows[0]._tr.get_or_add_trPr()
        self.assertIsNone(layout_props.find(qn("w:tblHeader")))
        self.assertIsNone(layout_props.find(qn("w:cantSplit")))

        audit = self.run_script(
            "audit_docx.py",
            self.source,
            formatted,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure,
        )
        self.assertEqual(0, audit.returncode, audit.stdout + audit.stderr)
        payload = json.loads(audit.stdout)
        self.assertEqual(
            ["profile_field_rules", "approved_structure_map"],
            payload["content_integrity"]["normalization_sources"],
        )

    def test_effective_audit_detects_conflicting_direct_format(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("Synthetic")
        paragraph.runs[0].font.size = Pt(20)
        rule = {
            "id": "FMT-BODY-999",
            "selector": {"kind": "paragraph_role", "value": "body_text"},
            "properties": {"font_size_pt": 10.5, "alignment": "justify"},
        }
        document.styles["Normal"].font.size = Pt(10.5)
        document.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.assertTrue(audit_paragraph_rule(document, rule, [paragraph]))
        apply_style_rule_to_paragraphs(document, rule, [paragraph])
        self.assertEqual([], audit_paragraph_rule(document, rule, [paragraph]))

    def test_environment_accepts_explicit_renderer_path(self) -> None:
        result = self.run_script(
            "check_environment.py", "--json", "--renderer", sys.executable
        )
        self.assertEqual(0, result.returncode, result.stderr)
        rendering = json.loads(result.stdout)["rendering"]
        self.assertEqual("argument", rendering["source"])
        self.assertEqual(str(Path(sys.executable).resolve()), rendering["renderer"])

    def test_incomplete_caption_cannot_be_approved(self) -> None:
        source = self.root / "incomplete-caption.docx"
        document = Document()
        document.add_paragraph("第1章 Synthetic")
        document.add_paragraph("表 Synthetic caption without a number")
        document.save(source)
        structure = candidate_structure_map(source)
        structure["schema_version"] = "1.1"
        structure["status"] = "approved"
        self.assertEqual("candidate", structure["captions"][0]["completeness"])
        structure["captions"][0].update(
            {"approved": True, "completeness": "incomplete"}
        )
        path = self.root / "incomplete-structure.json"
        path.write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
        result = self.run_script(
            "validate_structure_map.py", path, "--source", source
        )
        self.assertEqual(1, result.returncode)
        self.assertIn("incomplete", result.stderr.lower())

    def test_trailing_section_reports_independent_header_payload(self) -> None:
        source = self.root / "header-section.docx"
        document = Document()
        document.add_paragraph("Synthetic body")
        document.add_section(WD_SECTION.NEW_PAGE)
        final = document.sections[-1]
        final.header.is_linked_to_previous = False
        final.header.paragraphs[0].text = "Synthetic running header"
        document.save(source)
        structure = candidate_structure_map(source)
        evidence = structure["trailing_empty_sections"][0]["evidence"]
        self.assertTrue(evidence["header_footer_has_payload"])
        self.assertFalse(evidence["safe_to_delete"])

    def test_empty_header_reference_can_be_approved_for_tail_cleanup(self) -> None:
        source = self.root / "empty-header-section.docx"
        document = Document()
        document.add_paragraph("Synthetic body")
        document.add_section(WD_SECTION.NEW_PAGE)
        final = document.sections[-1]
        final.header.is_linked_to_previous = False
        document.save(source)
        structure = candidate_structure_map(source)
        evidence = structure["trailing_empty_sections"][0]["evidence"]
        self.assertTrue(evidence["empty_header_footer_references"])
        self.assertFalse(evidence["header_footer_has_payload"])
        self.assertTrue(evidence["safe_to_delete"])

    def test_whole_book_numbering_progression_is_detected(self) -> None:
        source = self.root / "whole-book.docx"
        document = Document()
        document.add_paragraph("第1章 Synthetic one")
        document.add_paragraph("1.1 Synthetic section")
        document.add_paragraph("第2章 Synthetic two")
        document.add_paragraph("2.1 Synthetic section")
        document.save(source)
        numbering = candidate_structure_map(source)["numbering"]
        self.assertEqual("whole_book", numbering["mode"])
        self.assertEqual(1, numbering["chapter_start"])
        self.assertEqual([], numbering["anomalies"])

    def test_structure_map_10_remains_readable(self) -> None:
        value = {
            "schema_version": "1.0",
            "status": "approved",
            "source_content_fingerprint_sha256": "0" * 64,
            "toc_ranges": [],
            "headings": [],
            "captions": [],
            "tables": [],
            "trailing_empty_sections": [],
            "conflicts": [],
        }
        path = self.root / "legacy-map.json"
        path.write_text(json.dumps(value), encoding="utf-8")
        self.assertEqual("1.0", load_structure_map(path)["schema_version"])


if __name__ == "__main__":
    unittest.main()
