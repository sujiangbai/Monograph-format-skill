from __future__ import annotations

import importlib.util
import json
import os
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from types import ModuleType, SimpleNamespace
from unittest.mock import Mock, patch
from urllib.parse import unquote, urlparse

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "format-monograph" / "scripts"
TESTS = ROOT / "tests"
sys.path.insert(0, str(SCRIPTS))
sys.path.insert(0, str(TESTS))

from _common import FormatMonographError  # noqa: E402
from backend_evidence import (  # noqa: E402
    BackendEvidenceError,
    backend_audit_binding,
    backend_audit_bytes,
    canonical_backend_projection,
    read_bound_backend_audit,
)
from check_environment import libreoffice_field_status, resolve_renderer  # noqa: E402
from finalize_docx import (  # noqa: E402
    LIBREOFFICE_FIELD_SCRIPT_URI,
    libreoffice_refresh,
    libreoffice_macro_refresh,
    package_external_connection_inventory,
    package_field_contract_manifest,
    restore_known_libreoffice_toc_instruction_order,
    toc_index_authorization,
)
from libreoffice_runtime import macos_internal_macro_soffice  # noqa: E402
from render_smoke import (  # noqa: E402
    field_finalization_errors,
    field_finalization_summary,
    load_field_backend_audit,
)
from libreoffice_security_smoke import positive_control_calibration  # noqa: E402
from toc_index_identity import authorization_with_hash  # noqa: E402


class FakePropertyValue:
    def __init__(self) -> None:
        self.Name = ""
        self.Value = None


def load_macro_module():
    uno = ModuleType("uno")
    uno.systemPathToFileUrl = lambda value: value  # type: ignore[attr-defined]
    modules = {
        "uno": uno,
        "com": ModuleType("com"),
        "com.sun": ModuleType("com.sun"),
        "com.sun.star": ModuleType("com.sun.star"),
        "com.sun.star.beans": ModuleType("com.sun.star.beans"),
        "com.sun.star.document": ModuleType("com.sun.star.document"),
        "com.sun.star.document.MacroExecMode": ModuleType(
            "com.sun.star.document.MacroExecMode"
        ),
        "com.sun.star.document.UpdateDocMode": ModuleType(
            "com.sun.star.document.UpdateDocMode"
        ),
    }
    modules["com.sun.star.beans"].PropertyValue = FakePropertyValue  # type: ignore[attr-defined]
    modules["com.sun.star.document.MacroExecMode"].NEVER_EXECUTE = 0  # type: ignore[attr-defined]
    modules["com.sun.star.document.UpdateDocMode"].NO_UPDATE = 0  # type: ignore[attr-defined]
    name = f"libreoffice_fields_macro_test_{id(modules)}"
    spec = importlib.util.spec_from_file_location(
        name, SCRIPTS / "libreoffice_fields_macro.py"
    )
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    with patch.dict(sys.modules, modules):
        spec.loader.exec_module(module)
    return module


class LibreOfficeRuntimeTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.soffice = (
            self.root / "LibreOffice.app" / "Contents" / "MacOS" / "soffice"
        )
        self.soffice.parent.mkdir(parents=True)
        self.soffice.write_text("#!/bin/sh\n", encoding="utf-8")
        self.soffice.chmod(0o755)
        contents = self.soffice.parents[1]
        (contents / "Resources" / "Scripts" / "python").mkdir(parents=True)
        (contents / "Frameworks" / "LibreOfficePython.framework").mkdir(
            parents=True
        )

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_direct_macos_app_runtime_is_detected(self) -> None:
        self.assertEqual(
            str(self.soffice.resolve()),
            macos_internal_macro_soffice(str(self.soffice), system="Darwin"),
        )

    def test_homebrew_style_wrapper_resolves_app_runtime(self) -> None:
        wrapper = self.root / "bin" / "soffice"
        wrapper.parent.mkdir()
        wrapper.write_text(
            f'#!/bin/bash\nexec "{self.soffice}" "$@"\n', encoding="utf-8"
        )
        wrapper.chmod(0o755)
        self.assertEqual(
            str(self.soffice.resolve()),
            macos_internal_macro_soffice(str(wrapper), system="Darwin"),
        )

    def test_internal_macro_runtime_is_macos_only(self) -> None:
        self.assertIsNone(
            macos_internal_macro_soffice(str(self.soffice), system="Linux")
        )

    def test_legacy_uno_server_fallback_is_fail_closed_matrix(self) -> None:
        input_path = self.root / "input.docx"
        DocxDocument().save(input_path)
        output_path = self.root / "output.docx"
        unsupported = self.root / "soffice"
        unsupported.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        unsupported.chmod(0o755)
        wrapper = self.root / "ordinary-wrapper"
        wrapper.write_text("#!/bin/sh\nexec soffice \"$@\"\n", encoding="utf-8")
        wrapper.chmod(0o755)
        cases = (
            ("Linux", str(unsupported)),
            ("Windows", str(unsupported)),
            ("Darwin", str(unsupported)),
            ("Darwin", str(wrapper)),
        )
        for system, renderer in cases:
            output_path.unlink(missing_ok=True)
            with self.subTest(system=system, renderer=Path(renderer).name), patch(
                "libreoffice_runtime.platform.system", return_value=system
            ), patch("finalize_docx.subprocess.Popen") as popen, patch(
                "finalize_docx.libreoffice_macro_refresh"
            ) as macro_refresh:
                with self.assertRaisesRegex(
                    FormatMonographError, "legacy UNO server/helper backend is disabled"
                ):
                    libreoffice_refresh(input_path, output_path, renderer)
                popen.assert_not_called()
                macro_refresh.assert_not_called()
                self.assertFalse(output_path.exists())

    def test_legacy_helper_entrypoint_is_an_inert_fail_closed_stub(self) -> None:
        helper = SCRIPTS / "libreoffice_fields.py"
        completed = subprocess.run(
            [sys.executable, str(helper)],
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(2, completed.returncode)
        self.assertIn("legacy LibreOffice UNO server/helper backend is disabled", completed.stderr)
        source = helper.read_text(encoding="utf-8")
        self.assertNotIn("UpdateDocMode", source)
        self.assertNotIn("getTextFields", source)
        self.assertNotIn("getDocumentIndexes", source)

    def test_unsupported_host_rejects_active_fixture_matrix_before_launch(self) -> None:
        renderer = self.root / "fake-renderer"
        renderer.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
        renderer.chmod(0o755)
        fixtures = {
            "external_relationship": (
                "word/_rels/document.xml.rels",
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" Type="test" Target="http://127.0.0.1/probe" TargetMode="External"/>'
                "</Relationships>",
            ),
            "on_load_macro": (
                "content.xml",
                '<root xmlns:xlink="http://www.w3.org/1999/xlink" '
                'xlink:href="vnd.sun.star.script:Standard.Main?language=Basic&amp;location=document"/>',
            ),
            "unapproved_index": (
                "word/document.xml",
                '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:body><w:p><w:r><w:instrText>TOC \\o "1-9"</w:instrText>'
                "</w:r></w:p></w:body></w:document>",
            ),
        }
        for name, (part, payload) in fixtures.items():
            fixture = self.root / f"{name}.docx"
            with zipfile.ZipFile(fixture, "w") as package:
                package.writestr(part, payload)
            with self.subTest(name=name), patch(
                "libreoffice_runtime.platform.system", return_value="Linux"
            ), patch("finalize_docx.subprocess.Popen") as popen, patch(
                "finalize_docx.subprocess.run"
            ) as run:
                with self.assertRaises(FormatMonographError):
                    libreoffice_refresh(fixture, self.root / f"{name}-out.docx", str(renderer))
                popen.assert_not_called()
                run.assert_not_called()

    def test_environment_never_advertises_legacy_uno_field_refresh(self) -> None:
        with patch("libreoffice_runtime.platform.system", return_value="Linux"):
            status = libreoffice_field_status(str(self.soffice))
        self.assertFalse(status["available"])
        self.assertIsNone(status["backend"])
        self.assertEqual(
            "verified_internal_macro_host_unavailable", status["reason"]
        )

    def test_external_uri_scheme_matrix_and_internal_relative_link(self) -> None:
        namespace = "urn:test-links"
        for scheme in ("file", "ftp", "smb", "http", "https", "mailto"):
            package_path = self.root / f"{scheme}.odt"
            with zipfile.ZipFile(package_path, "w") as package:
                package.writestr(
                    "content.xml",
                    f'<root xmlns:x="{namespace}" x:href="{scheme}://host/path"/>',
                )
            inventory = package_external_connection_inventory(package_path)
            with self.subTest(scheme=scheme):
                self.assertEqual(1, len(inventory))
                self.assertEqual(f"uri_scheme:{scheme}", inventory[0]["reason"])

        file_path = self.root / "file-uri.odt"
        with zipfile.ZipFile(file_path, "w") as package:
            package.writestr(
                "content.xml",
                f'<root xmlns:x="{namespace}" x:href="file:///tmp/probe.png"/>',
            )
        self.assertEqual(
            "uri_scheme:file",
            package_external_connection_inventory(file_path)[0]["reason"],
        )

        internal = self.root / "internal.odt"
        with zipfile.ZipFile(internal, "w") as package:
            package.writestr(
                "content.xml",
                f'<root xmlns:x="{namespace}" x:href="Pictures/image.png"/>',
            )
            package.writestr("Pictures/image.png", b"image")
        self.assertEqual([], package_external_connection_inventory(internal))

    def test_relationship_target_mode_external_rejects_relative_target(self) -> None:
        package_path = self.root / "external-relationship.docx"
        with zipfile.ZipFile(package_path, "w") as package:
            package.writestr(
                "word/_rels/document.xml.rels",
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" Type="test" Target="relative.bin" TargetMode="External"/>'
                "</Relationships>",
            )
        inventory = package_external_connection_inventory(package_path)
        self.assertEqual("target_mode_external", inventory[0]["reason"])

    def _relationship_package(
        self,
        name: str,
        relationship_part: str,
        target_value: str,
        members: tuple[str, ...] = (),
    ) -> Path:
        package_path = self.root / name
        with zipfile.ZipFile(package_path, "w") as package:
            package.writestr(
                relationship_part,
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                f'<Relationship Id="rId1" Type="test" Target="{target_value}"/>'
                "</Relationships>",
            )
            for member in members:
                package.writestr(member, b"member")
        return package_path

    def test_internal_relationship_target_resolution_matrix(self) -> None:
        rejected = {
            "slash-network": ("//server/share.bin", "network_path"),
            "backslash-network": ("\\\\server\\share.bin", "network_path"),
            "escape": ("../../escape.bin", "outside_package"),
            "encoded-escape": ("%2e%2e/%2e%2e/escape.bin", "outside_package"),
            "unresolved": ("media/missing.png", "unresolved_package_reference"),
            "scheme": ("smb://server/share.bin", "uri_scheme:smb"),
        }
        for name, (target, reason) in rejected.items():
            package_path = self._relationship_package(
                f"{name}.docx",
                "word/_rels/document.xml.rels",
                target,
                ("word/document.xml",),
            )
            inventory = package_external_connection_inventory(package_path)
            with self.subTest(name=name):
                self.assertEqual(1, len(inventory))
                self.assertEqual(reason, inventory[0]["reason"])

        for name, relationship_part, target, members in (
            (
                "word-relative",
                "word/_rels/document.xml.rels",
                "media/image.png",
                ("word/document.xml", "word/media/image.png"),
            ),
            (
                "word-backslash",
                "word/_rels/document.xml.rels",
                "media\\image.png",
                ("word/document.xml", "word/media/image.png"),
            ),
            (
                "word-fragment",
                "word/_rels/document.xml.rels",
                "#bookmark",
                ("word/document.xml",),
            ),
            (
                "package-root",
                "_rels/.rels",
                "word/document.xml",
                ("word/document.xml",),
            ),
        ):
            package_path = self._relationship_package(
                f"{name}.docx", relationship_part, target, members
            )
            with self.subTest(name=name):
                self.assertEqual(
                    [], package_external_connection_inventory(package_path)
                )

        missing_source = self._relationship_package(
            "missing-source.docx",
            "word/_rels/document.xml.rels",
            "media/image.png",
            ("word/media/image.png",),
        )
        self.assertEqual(
            "missing_relationship_source_part",
            package_external_connection_inventory(missing_source)[0]["reason"],
        )

    @unittest.skipIf(os.name == "nt", "Windows does not expose a POSIX executable bit.")
    def test_non_executable_soffice_is_unavailable(self) -> None:
        self.soffice.chmod(0o644)
        self.assertIsNone(
            macos_internal_macro_soffice(str(self.soffice), system="Darwin")
        )
        self.assertEqual((None, "argument"), resolve_renderer(str(self.soffice)))

    def _macro_process(self) -> Mock:
        process = Mock()
        process.poll.return_value = None
        process.returncode = None
        process.communicate.return_value = ("", "")
        process.wait.return_value = 0
        return process

    def _toc_docx(self, path: Path, instruction: str = 'TOC \\o "1-3" \\h \\z') -> None:
        document = DocxDocument()
        paragraph = document.add_paragraph()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        paragraph.add_run()._r.append(begin)
        instruction_node = OxmlElement("w:instrText")
        instruction_node.set(qn("xml:space"), "preserve")
        instruction_node.text = f" {instruction} "
        paragraph.add_run()._r.append(instruction_node)
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        paragraph.add_run()._r.append(separate)
        paragraph.add_run("Synthetic heading 1")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        paragraph.add_run()._r.append(end)
        document.save(path)

    def _toc_authorization(self, path: Path) -> dict:
        return toc_index_authorization(self._toc_contract(), path)

    def _toc_contract(self) -> list[dict]:
        return [{"level": 1, "kind": "heading", "text_sha256": "a" * 64}]

    def test_macro_command_uses_isolated_profile_and_success_result(self) -> None:
        input_path = self.root / "input.docx"
        output_path = self.root / "output.docx"
        self._toc_docx(input_path)
        process = self._macro_process()
        observed = {}

        def launch(command, **kwargs):
            observed["command"] = command
            profile_url = next(
                value.split("=", 1)[1]
                for value in command
                if value.startswith("-env:UserInstallation=")
            )
            profile = Path(unquote(urlparse(profile_url).path))
            observed["profile"] = profile
            observed["helper_exists"] = (
                profile / "user" / "Scripts" / "python" / "libreoffice_fields_macro.py"
            ).is_file()
            env = kwargs["env"]
            observed["toc_authorization"] = json.loads(
                env["FORMAT_MONOGRAPH_TOC_AUTHORIZATION"]
            )
            observed["toc_contract"] = json.loads(
                env["FORMAT_MONOGRAPH_TOC_CONTRACT"]
            )
            Path(env["FORMAT_MONOGRAPH_FIELD_RESULT"]).write_text(
                json.dumps(
                    {
                        "ok": True,
                        "approved_indexes_updated": 1,
                        "text_fields_collection_refreshed": False,
                    }
                ),
                encoding="utf-8",
            )
            Path(env["FORMAT_MONOGRAPH_FIELD_OUTPUT"]).write_bytes(b"refreshed")
            return process

        with patch("finalize_docx.subprocess.Popen", side_effect=launch), patch(
            "finalize_docx.rewrite_field_flags"
        ) as rewrite:
            result = libreoffice_macro_refresh(
                input_path,
                output_path,
                str(self.soffice),
                "argument",
                toc_authorization=self._toc_authorization(input_path),
                toc_contract=self._toc_contract(),
            )

        self.assertTrue(observed["helper_exists"])
        self.assertEqual(str(self.soffice), observed["command"][0])
        self.assertIn("--headless", observed["command"])
        self.assertEqual(LIBREOFFICE_FIELD_SCRIPT_URI, observed["command"][-1])
        self.assertIn("format-monograph-fields-", str(observed["profile"]))
        self.assertEqual("libreoffice_uno", result["backend"])
        self.assertEqual("internal_python_macro", result["uno_mode"])
        self.assertEqual(2, observed["toc_authorization"]["version"])
        self.assertEqual(self._toc_contract(), observed["toc_contract"])
        self.assertEqual(3, observed["toc_authorization"]["indexes"][0]["uno"]["level"])
        self.assertEqual(
            'TOC \\o "1-3" \\h \\z',
            observed["toc_authorization"]["indexes"][0]["ooxml"]["instruction"],
        )
        self.assertFalse(result["text_fields_collection_refreshed"])
        rewrite.assert_called_once_with(output_path, deferred=False)
        process.terminate.assert_called_once_with()
        process.wait.assert_called_once()

    def test_macro_abnormal_exit_reports_only_log_tails_without_pipe_deadlock(self) -> None:
        input_path = self.root / "input.docx"
        DocxDocument().save(input_path)
        process = self._macro_process()
        process.poll.return_value = 7
        process.returncode = 7
        observed = {}

        def launch(_command, **kwargs):
            observed["stdout"] = kwargs["stdout"]
            observed["stderr"] = kwargs["stderr"]
            kwargs["stdout"].write("A" * 20000 + "STDOUT-TAIL")
            kwargs["stderr"].write("B" * 20000 + "STDERR-TAIL")
            return process

        with patch("finalize_docx.subprocess.Popen", side_effect=launch):
            with self.assertRaisesRegex(
                FormatMonographError, "exited without a result"
            ) as raised:
                libreoffice_macro_refresh(
                    input_path,
                    self.root / "output.docx",
                    str(self.soffice),
                    "argument",
                )
        self.assertIsNot(observed["stdout"], subprocess.PIPE)
        self.assertIsNot(observed["stderr"], subprocess.PIPE)
        self.assertIn("STDOUT-TAIL", str(raised.exception))
        self.assertIn("STDERR-TAIL", str(raised.exception))
        process.terminate.assert_not_called()

    def test_external_field_is_rejected_before_libreoffice_starts(self) -> None:
        input_path = self.root / "external-field.docx"
        self._toc_docx(
            input_path,
            'INCLUDETEXT "http://127.0.0.1:9/never-contact.docx"',
        )
        with patch("finalize_docx.subprocess.Popen") as launch:
            with self.assertRaisesRegex(
                FormatMonographError, "refuses documents with active external connections"
            ):
                libreoffice_macro_refresh(
                    input_path,
                    self.root / "external-output.docx",
                    str(self.soffice),
                    "argument",
                )
        launch.assert_not_called()

    def test_macro_error_result_is_rejected_and_process_is_terminated(self) -> None:
        input_path = self.root / "input.docx"
        output_path = self.root / "output.docx"
        DocxDocument().save(input_path)
        process = self._macro_process()

        def launch(_command, **kwargs):
            Path(kwargs["env"]["FORMAT_MONOGRAPH_FIELD_RESULT"]).write_text(
                json.dumps({"ok": False, "error": "synthetic failure"}),
                encoding="utf-8",
            )
            return process

        with patch("finalize_docx.subprocess.Popen", side_effect=launch):
            with self.assertRaisesRegex(FormatMonographError, "synthetic failure"):
                libreoffice_macro_refresh(
                    input_path, output_path, str(self.soffice), "argument"
                )
        process.terminate.assert_called_once_with()
        process.wait.assert_called_once()

    def test_macro_timeout_terminates_process(self) -> None:
        input_path = self.root / "input.docx"
        DocxDocument().save(input_path)
        process = self._macro_process()
        with patch("finalize_docx.subprocess.Popen", return_value=process), patch(
            "finalize_docx.time.monotonic", side_effect=[0, 301]
        ):
            with self.assertRaisesRegex(FormatMonographError, "timed out"):
                libreoffice_macro_refresh(
                    input_path,
                    self.root / "output.docx",
                    str(self.soffice),
                    "argument",
                )
        process.terminate.assert_called_once_with()
        process.wait.assert_called_once()

    def test_macro_shutdown_kills_process_that_does_not_terminate(self) -> None:
        input_path = self.root / "input.docx"
        output_path = self.root / "output.docx"
        DocxDocument().save(input_path)
        process = self._macro_process()
        process.wait.side_effect = [
            subprocess.TimeoutExpired("soffice", 15),
            0,
        ]

        def launch(_command, **kwargs):
            Path(kwargs["env"]["FORMAT_MONOGRAPH_FIELD_RESULT"]).write_text(
                json.dumps({"ok": True}), encoding="utf-8"
            )
            output_path.write_bytes(b"refreshed")
            return process

        with patch("finalize_docx.subprocess.Popen", side_effect=launch), patch(
            "finalize_docx.rewrite_field_flags"
        ):
            libreoffice_macro_refresh(
                input_path, output_path, str(self.soffice), "argument"
            )
        process.kill.assert_called_once_with()
        self.assertEqual(2, process.wait.call_count)

    def test_safe_media_descriptor_blocks_links_and_document_macros(self) -> None:
        module = load_macro_module()
        result_path = self.root / "macro-result.json"
        output_path = self.root / "macro-output.docx"

        class Indexes:
            def __init__(self) -> None:
                self.updated = False

            def getCount(self):
                return 1

            def getByIndex(self, _index):
                return SimpleNamespace(
                    supportsService=lambda service: service
                    == "com.sun.star.text.ContentIndex",
                    CreateFromOutline=True,
                    CreateFromMarks=False,
                    Level=3,
                    update=lambda: setattr(self, "updated", True),
                )

        class Fields:
            def __init__(self) -> None:
                self.refreshed = False

            def refresh(self):
                self.refreshed = True

            def createEnumeration(self):
                return SimpleNamespace(
                    hasMoreElements=lambda: False,
                    nextElement=lambda: None,
                )

        indexes = Indexes()
        fields = Fields()

        class Document:
            def __init__(self) -> None:
                self.calculated = False
                self.closed = False

            def getDocumentIndexes(self):
                return indexes

            def getTextFields(self):
                return fields

            def getGraphicObjects(self):
                return SimpleNamespace(getElementNames=lambda: ())

            def calculateAll(self):
                self.calculated = True

            def storeAsURL(self, target, _properties):
                Path(target).write_bytes(b"synthetic output")

            def close(self, _deliver_ownership):
                self.closed = True

        document = Document()

        class Desktop:
            def __init__(self) -> None:
                self.properties = {}
                self.external_link_updated = False
                self.document_macro_executed = False

            def loadComponentFromURL(self, _source, _target, _flags, properties):
                self.properties = {item.Name: item.Value for item in properties}
                self.external_link_updated = self.properties.get("UpdateDocMode") != 0
                self.document_macro_executed = (
                    self.properties.get("MacroExecutionMode") != 0
                )
                return document

        desktop = Desktop()
        module.XSCRIPTCONTEXT = SimpleNamespace(getDesktop=lambda: desktop)
        input_path = self.root / "input.docx"
        self._toc_docx(input_path)
        environment = {
            "FORMAT_MONOGRAPH_FIELD_INPUT": str(input_path),
            "FORMAT_MONOGRAPH_FIELD_OUTPUT": str(output_path),
            "FORMAT_MONOGRAPH_FIELD_RESULT": str(result_path),
            "FORMAT_MONOGRAPH_TOC_AUTHORIZATION": json.dumps(
                self._toc_authorization(input_path)
            ),
            "FORMAT_MONOGRAPH_TOC_CONTRACT": json.dumps(self._toc_contract()),
        }
        with patch.dict(os.environ, environment):
            self.assertTrue(module.refresh_from_environment())

        self.assertEqual(0, desktop.properties["UpdateDocMode"])
        self.assertEqual(0, desktop.properties["MacroExecutionMode"])
        self.assertFalse(desktop.external_link_updated)
        self.assertFalse(desktop.document_macro_executed)
        self.assertTrue(indexes.updated)
        self.assertFalse(fields.refreshed)
        self.assertTrue(document.calculated)
        self.assertTrue(document.closed)
        result = json.loads(result_path.read_text(encoding="utf-8"))
        self.assertTrue(result["ok"])
        self.assertFalse(result["text_fields_collection_refreshed"])

    def _run_macro_index_fixture(self, candidates, authorization, contract=None):
        module = load_macro_module()
        result_path = self.root / f"index-result-{len(candidates)}.json"
        output_path = self.root / f"index-output-{len(candidates)}.docx"

        class Indexes:
            def getCount(self):
                return len(candidates)

            def getByIndex(self, ordinal):
                return candidates[ordinal]

        class Document:
            def getDocumentIndexes(self):
                return Indexes()

            def getTextFields(self):
                return SimpleNamespace(
                    createEnumeration=lambda: SimpleNamespace(
                        hasMoreElements=lambda: False,
                        nextElement=lambda: None,
                    )
                )

            def getGraphicObjects(self):
                return SimpleNamespace(getElementNames=lambda: ())

            def calculateAll(self):
                return None

            def storeAsURL(self, target, _properties):
                Path(target).write_bytes(b"synthetic output")

            def close(self, _deliver_ownership):
                return None

        module.XSCRIPTCONTEXT = SimpleNamespace(
            getDesktop=lambda: SimpleNamespace(
                loadComponentFromURL=lambda *_args: Document()
            )
        )
        input_path = self.root / "index-input.docx"
        self._toc_docx(input_path)
        environment = {
            "FORMAT_MONOGRAPH_FIELD_INPUT": str(input_path),
            "FORMAT_MONOGRAPH_FIELD_OUTPUT": str(output_path),
            "FORMAT_MONOGRAPH_FIELD_RESULT": str(result_path),
            "FORMAT_MONOGRAPH_TOC_AUTHORIZATION": (
                json.dumps(authorization) if authorization is not None else ""
            ),
            "FORMAT_MONOGRAPH_TOC_CONTRACT": (
                json.dumps(contract) if contract is not None else ""
            ),
        }
        with patch.dict(os.environ, environment):
            returned = module.refresh_from_environment()
        return returned, json.loads(result_path.read_text(encoding="utf-8"))

    def test_unapproved_toc_index_is_skipped_before_update(self) -> None:
        candidate = SimpleNamespace(
            supportsService=Mock(return_value=True),
            CreateFromOutline=True,
            CreateFromMarks=False,
            Level=3,
            update=Mock(),
        )
        returned, result = self._run_macro_index_fixture([candidate], None)
        self.assertTrue(returned)
        candidate.update.assert_not_called()
        self.assertEqual(0, result["approved_indexes_updated"])
        self.assertEqual(1, result["skipped_indexes"])

    def test_multiple_toc_indexes_reject_before_any_update(self) -> None:
        candidates = [
            SimpleNamespace(
                supportsService=Mock(return_value=True),
                CreateFromOutline=True,
                CreateFromMarks=False,
                Level=3,
                update=Mock(),
            )
            for _ in range(2)
        ]
        input_path = self.root / "index-input.docx"
        self._toc_docx(input_path)
        authorization = self._toc_authorization(input_path)
        returned, result = self._run_macro_index_fixture(
            candidates, authorization, self._toc_contract()
        )
        self.assertFalse(returned)
        self.assertIn("index count", result["error"])
        for candidate in candidates:
            candidate.update.assert_not_called()

    def test_toc_index_identity_conflict_rejects_before_update(self) -> None:
        candidate = SimpleNamespace(
            supportsService=Mock(return_value=False),
            CreateFromOutline=True,
            CreateFromMarks=False,
            Level=3,
            update=Mock(),
        )
        input_path = self.root / "index-input.docx"
        self._toc_docx(input_path)
        authorization = self._toc_authorization(input_path)
        returned, result = self._run_macro_index_fixture(
            [candidate], authorization, self._toc_contract()
        )
        self.assertFalse(returned)
        self.assertIn("identity", result["error"])
        candidate.update.assert_not_called()

    def test_same_service_wrong_toc_properties_reject_before_update(self) -> None:
        input_path = self.root / "index-input.docx"
        self._toc_docx(input_path)
        authorization = self._toc_authorization(input_path)
        for property_name, wrong_value in (
            ("CreateFromOutline", False),
            ("CreateFromMarks", True),
            ("Level", 4),
        ):
            candidate = SimpleNamespace(
                supportsService=Mock(return_value=True),
                CreateFromOutline=True,
                CreateFromMarks=False,
                Level=3,
                update=Mock(),
            )
            setattr(candidate, property_name, wrong_value)
            returned, result = self._run_macro_index_fixture(
                [candidate], authorization, self._toc_contract()
            )
            with self.subTest(property_name=property_name):
                self.assertFalse(returned)
                self.assertIn("identity", result["error"])
                candidate.update.assert_not_called()

    def test_independent_toc_contract_anchor_rejects_tampering_before_update(self) -> None:
        input_path = self.root / "index-input.docx"
        self._toc_docx(input_path)
        authorization = self._toc_authorization(input_path)
        candidate = SimpleNamespace(
            supportsService=Mock(return_value=True),
            CreateFromOutline=True,
            CreateFromMarks=False,
            Level=3,
            update=Mock(),
        )
        descriptor = dict(authorization)
        descriptor.pop("authorization_id")
        descriptor["structure_contract_sha256"] = "0" * 64
        tampered = authorization_with_hash(descriptor)
        returned, result = self._run_macro_index_fixture(
            [candidate], tampered, self._toc_contract()
        )
        self.assertFalse(returned)
        self.assertIn("anchor hash", result["error"])
        candidate.update.assert_not_called()

    def test_toc_contract_anchor_shape_failures_are_zero_update(self) -> None:
        input_path = self.root / "index-input.docx"
        self._toc_docx(input_path)
        authorization = self._toc_authorization(input_path)
        for name, contract in (
            ("missing", None),
            ("non_list", {"level": 1}),
            ("empty", []),
            (
                "invalid_item",
                [{"level": 1, "kind": "heading", "text_sha256": "bad"}],
            ),
        ):
            candidate = SimpleNamespace(
                supportsService=Mock(return_value=True),
                CreateFromOutline=True,
                CreateFromMarks=False,
                Level=3,
                update=Mock(),
            )
            returned, result = self._run_macro_index_fixture(
                [candidate], authorization, contract
            )
            with self.subTest(name=name):
                self.assertFalse(returned)
                self.assertIn("contract anchor", result["error"])
                candidate.update.assert_not_called()

    def test_close_failure_removes_output_and_publishes_negative_result(self) -> None:
        module = load_macro_module()
        result_path = self.root / "close-result.json"
        output_path = self.root / "close-output.docx"

        class EmptyIndexes:
            def getCount(self):
                return 0

        class Document:
            def getDocumentIndexes(self):
                return EmptyIndexes()

            def getTextFields(self):
                return SimpleNamespace(
                    createEnumeration=lambda: SimpleNamespace(
                        hasMoreElements=lambda: False,
                        nextElement=lambda: None,
                    )
                )

            def getGraphicObjects(self):
                return SimpleNamespace(getElementNames=lambda: ())

            def calculateAll(self):
                return None

            def storeAsURL(self, target, _properties):
                Path(target).write_bytes(b"must not survive")

            def close(self, _deliver_ownership):
                raise RuntimeError("close rejected")

        module.XSCRIPTCONTEXT = SimpleNamespace(
            getDesktop=lambda: SimpleNamespace(
                loadComponentFromURL=lambda *_args: Document()
            )
        )
        environment = {
            "FORMAT_MONOGRAPH_FIELD_INPUT": str(self.root / "input.docx"),
            "FORMAT_MONOGRAPH_FIELD_OUTPUT": str(output_path),
            "FORMAT_MONOGRAPH_FIELD_RESULT": str(result_path),
        }
        with patch.dict(os.environ, environment):
            self.assertFalse(module.refresh_from_environment())
        result = json.loads(result_path.read_text(encoding="utf-8"))
        self.assertFalse(result["ok"])
        self.assertIn("document close failed", result["error"])
        self.assertFalse(output_path.exists())

    def test_macro_load_error_writes_negative_result(self) -> None:
        module = load_macro_module()
        result_path = self.root / "macro-error.json"
        module.XSCRIPTCONTEXT = SimpleNamespace(
            getDesktop=lambda: SimpleNamespace(
                loadComponentFromURL=Mock(side_effect=RuntimeError("load rejected"))
            )
        )
        environment = {
            "FORMAT_MONOGRAPH_FIELD_INPUT": str(self.root / "input.docx"),
            "FORMAT_MONOGRAPH_FIELD_OUTPUT": str(self.root / "output.docx"),
            "FORMAT_MONOGRAPH_FIELD_RESULT": str(result_path),
        }
        with patch.dict(os.environ, environment):
            self.assertFalse(module.refresh_from_environment())
        result = json.loads(result_path.read_text(encoding="utf-8"))
        self.assertFalse(result["ok"])
        self.assertIn("load rejected", result["error"])

    def test_deferred_output_never_satisfies_smoke_gate(self) -> None:
        payload = {
            "status": "pass",
            "delivery_field_status": "deferred",
            "field_writeback_status": "deferred",
            "field_backend": {
                "backend": "deferred_on_open",
                "fallback_from": "libreoffice_contract_or_integrity",
                "attempted_backend": {
                    "backend": "libreoffice_uno",
                    "selective_writeback": {
                        "status": "rejected",
                        "error": "Target application changed field instructions.",
                    },
                },
            },
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
            "field_completion": {
                "field_gate_completed": False,
                "final_ready_eligible": False,
                "word_verification_required": True,
                "word_verification_completed": False,
            },
        }
        raw_backend = payload["field_backend"]
        payload["field_backend"] = canonical_backend_projection(raw_backend)
        self.assertEqual([], field_finalization_errors(payload, raw_backend))
        summary = field_finalization_summary(payload, raw_backend)
        self.assertEqual("strictly_deferred", summary["gate_outcome"])
        self.assertFalse(summary["backend_result_accepted"])

    def test_safe_linux_libreoffice_unavailability_is_strictly_deferred(self) -> None:
        error = (
            "LibreOffice field refresh requires the verified macOS internal-Python "
            "macro host; the legacy UNO server/helper backend is disabled."
        )
        raw_backend = {
            "backend": "deferred_on_open",
            "fallback_from": "libreoffice_error",
            "attempted_backend": {
                "backend": "libreoffice_uno",
                "failure": {
                    "status": "rejected",
                    "stage": "libreoffice_refresh",
                    "error": error,
                    "failed_checks": ["libreoffice_refresh"],
                },
            },
        }
        payload = {
            "status": "pass",
            "delivery_field_status": "deferred",
            "field_writeback_status": "deferred",
            "field_backend": canonical_backend_projection(raw_backend),
            "field_completion": {
                "field_gate_completed": False,
                "final_ready_eligible": False,
                "word_verification_required": True,
                "word_verification_completed": False,
            },
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
        }
        self.assertEqual([], field_finalization_errors(payload, raw_backend))
        self.assertEqual(
            "strictly_deferred",
            field_finalization_summary(payload, raw_backend)["gate_outcome"],
        )

    def test_unknown_libreoffice_failure_is_not_strictly_deferred(self) -> None:
        raw_backend = {
            "backend": "deferred_on_open",
            "fallback_from": "libreoffice_error",
            "attempted_backend": {
                "backend": "libreoffice_uno",
                "failure": {
                    "status": "rejected",
                    "stage": "libreoffice_refresh",
                    "error": "unknown renderer failure",
                    "failed_checks": ["libreoffice_refresh"],
                },
            },
        }
        payload = {
            "status": "pass",
            "delivery_field_status": "deferred",
            "field_writeback_status": "deferred",
            "field_backend": canonical_backend_projection(raw_backend),
            "field_completion": {
                "field_gate_completed": False,
                "final_ready_eligible": False,
                "word_verification_required": True,
                "word_verification_completed": False,
            },
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
        }
        errors = field_finalization_errors(payload, raw_backend)
        self.assertIn(
            "LibreOffice backend neither completed non-finally nor deferred strictly",
            errors,
        )
        self.assertEqual(
            "invalid", field_finalization_summary(payload, raw_backend)["gate_outcome"]
        )

    def test_external_positive_control_calibration_matrix(self) -> None:
        completed = {
            "unsafe_full_update_load_completed": True,
            "unsafe_graphics_loaded": 1,
        }
        self.assertEqual(
            {"status": "calibrated", "reason": None},
            positive_control_calibration(completed, ["/linked-graphic.png"]),
        )
        self.assertEqual(
            {
                "status": "unavailable",
                "reason": "loopback_request_not_observed",
            },
            positive_control_calibration(completed, []),
        )
        self.assertEqual(
            {
                "status": "invalid",
                "reason": "unexpected_loopback_request_path",
            },
            positive_control_calibration(completed, ["/unexpected.png"]),
        )
        self.assertEqual(
            {
                "status": "invalid",
                "reason": "unsafe_positive_control_did_not_complete",
            },
            positive_control_calibration(
                {
                    "unsafe_full_update_load_completed": False,
                    "unsafe_graphics_loaded": 1,
                },
                [],
            ),
        )

    def test_nonfinal_libreoffice_output_satisfies_backend_smoke_gate(self) -> None:
        payload = {
            "status": "pass",
            "delivery_field_status": "libreoffice_refreshed",
            "field_writeback_status": "libreoffice_selective",
            "field_backend": {
                "backend": "libreoffice_uno",
                "selective_writeback": {"status": "libreoffice_selective"},
                "delivery_field_contract_identical": True,
            },
            "field_completion": {
                "field_gate_completed": False,
                "final_ready_eligible": False,
                "word_verification_required": True,
                "word_verification_completed": False,
            },
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
        }
        raw_backend = payload["field_backend"]
        payload["field_backend"] = canonical_backend_projection(raw_backend)
        self.assertEqual([], field_finalization_errors(payload, raw_backend))

    def test_post_integrity_deferred_preserves_attempt_and_satisfies_safety_smoke(self) -> None:
        payload = {
            "status": "pass",
            "delivery_field_status": "deferred",
            "field_writeback_status": "deferred",
            "field_backend": {
                "backend": "deferred_on_open",
                "fallback_from": "libreoffice_contract_or_integrity",
                "attempted_backend": {
                    "backend": "libreoffice_uno",
                    "approved_indexes_updated": 1,
                    "selective_writeback": {
                        "status": "libreoffice_selective",
                        "core_writeback_status": "selective_verified",
                    },
                    "failure": {
                        "status": "rejected",
                        "stage": "post_writeback_integrity",
                        "error": "synthetic content failure",
                        "failed_checks": ["content_integrity"],
                    },
                },
            },
            "field_completion": {
                "field_gate_completed": False,
                "final_ready_eligible": False,
                "word_verification_required": True,
                "word_verification_completed": False,
            },
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
        }
        raw_backend = payload["field_backend"]
        payload["field_backend"] = canonical_backend_projection(raw_backend)
        self.assertEqual([], field_finalization_errors(payload, raw_backend))
        summary = field_finalization_summary(payload, raw_backend)
        self.assertEqual("strictly_deferred", summary["gate_outcome"])
        self.assertEqual(1, summary["approved_indexes_updated"])
        self.assertNotIn("approved_indexes_updated", payload["field_backend"])
        self.assertEqual(
            "post_writeback_integrity", summary["failure"]["stage"]
        )

    def test_libreoffice_cannot_claim_word_selective_verified_status(self) -> None:
        payload = {
            "status": "pass",
            "delivery_field_status": "selective_verified",
            "field_writeback_status": "selective_verified",
            "field_backend": {
                "backend": "libreoffice_uno",
                "selective_writeback": {"status": "selective_verified"},
                "delivery_field_contract_identical": True,
            },
            "field_completion": {
                "field_gate_completed": True,
                "final_ready_eligible": True,
                "word_verification_required": False,
                "word_verification_completed": True,
            },
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
        }
        raw_backend = payload["field_backend"]
        payload["field_backend"] = canonical_backend_projection(raw_backend)
        errors = field_finalization_errors(payload, raw_backend)
        self.assertIn(
            "LibreOffice backend neither completed non-finally nor deferred strictly",
            errors,
        )
        self.assertIn("LibreOffice incorrectly became final_ready eligible", errors)

    def test_backend_audit_reader_and_render_smoke_reject_version_and_path_matrix(self) -> None:
        path = self.root / "backend-audit.json"
        payload = backend_audit_bytes({"backend": "libreoffice_uno"})
        path.write_bytes(payload)
        base = {"backend_audit": backend_audit_binding(path, payload)}
        for bad_version in (True, 1.0, "1", 2):
            value = json.loads(json.dumps(base))
            value["backend_audit"]["version"] = bad_version
            for reader in (read_bound_backend_audit, load_field_backend_audit):
                with self.subTest(
                    location="binding", version=bad_version, reader=reader.__name__
                ), self.assertRaises(BackendEvidenceError):
                    reader(value)

        for literal in ("true", "1.0", '"1"', "2"):
            root_payload = (
                '{"backend_audit_version":'
                + literal
                + ',"backend":{"backend":"libreoffice_uno"}}\n'
            ).encode("utf-8")
            path.write_bytes(root_payload)
            value = {"backend_audit": backend_audit_binding(path, root_payload)}
            for reader in (read_bound_backend_audit, load_field_backend_audit):
                with self.subTest(
                    location="root", version=literal, reader=reader.__name__
                ), self.assertRaises(BackendEvidenceError):
                    reader(value)

        duplicate_payload = (
            b'{"backend_audit_version":1,"backend_audit_version":1,'
            b'"backend":{"backend":"libreoffice_uno"}}\n'
        )
        path.write_bytes(duplicate_payload)
        duplicate = {
            "backend_audit": backend_audit_binding(path, duplicate_payload)
        }
        for reader in (read_bound_backend_audit, load_field_backend_audit):
            with self.subTest(reader=reader.__name__), self.assertRaises(
                BackendEvidenceError
            ):
                reader(duplicate)

        path.write_bytes(payload)
        for name, mutate in (
            (
                "unknown_binding_key",
                lambda value: value["backend_audit"].update(unknown=True),
            ),
            (
                "unknown_artifact_key",
                lambda value: value["backend_audit"]["artifact"].update(
                    unknown=True
                ),
            ),
            (
                "nul_path",
                lambda value: value["backend_audit"]["artifact"].update(
                    path="bad\0audit.json"
                ),
            ),
            (
                "control_path",
                lambda value: value["backend_audit"]["artifact"].update(
                    path="bad\naudit.json"
                ),
            ),
        ):
            value = json.loads(json.dumps(base))
            mutate(value)
            for reader in (read_bound_backend_audit, load_field_backend_audit):
                with self.subTest(name=name, reader=reader.__name__), self.assertRaises(
                    BackendEvidenceError
                ):
                    reader(value)

    def _toc_docx(
        self, path: Path, instruction: str = 'TOC \\o "1-3" \\h \\z'
    ) -> None:
        document = DocxDocument()
        paragraph = document.add_paragraph()
        for kind in ("begin",):
            run = paragraph.add_run()
            marker = OxmlElement("w:fldChar")
            marker.set(qn("w:fldCharType"), kind)
            run._r.append(marker)
        run = paragraph.add_run()
        node = OxmlElement("w:instrText")
        node.set(qn("xml:space"), "preserve")
        node.text = f" {instruction} "
        run._r.append(node)
        run = paragraph.add_run()
        marker = OxmlElement("w:fldChar")
        marker.set(qn("w:fldCharType"), "separate")
        run._r.append(marker)
        paragraph.add_run("Contents 1")
        run = paragraph.add_run()
        marker = OxmlElement("w:fldChar")
        marker.set(qn("w:fldCharType"), "end")
        run._r.append(marker)
        document.save(path)

    def test_exact_known_toc_permutation_is_restored_to_baseline_contract(self) -> None:
        baseline = self.root / "baseline.docx"
        refreshed = self.root / "refreshed.docx"
        restored = self.root / "restored.docx"
        self._toc_docx(baseline, 'TOC \\o "1-3" \\h \\z')
        self._toc_docx(refreshed, 'TOC \\z \\o "1-3" \\h')
        report = restore_known_libreoffice_toc_instruction_order(
            baseline, refreshed, restored
        )
        self.assertEqual("restored_exact_known_permutation", report["status"])
        self.assertEqual(1, report["restored_toc_instructions"])
        self.assertEqual(
            package_field_contract_manifest(baseline),
            package_field_contract_manifest(restored),
        )

    def test_unrecognized_toc_instruction_change_remains_rejected(self) -> None:
        baseline = self.root / "baseline.docx"
        refreshed = self.root / "refreshed.docx"
        self._toc_docx(baseline, 'TOC \\o "1-3" \\h \\z')
        self._toc_docx(refreshed, 'TOC \\o "1-2" \\h \\z')
        with self.assertRaisesRegex(
            FormatMonographError, "outside the exact approved TOC permutation"
        ):
            restore_known_libreoffice_toc_instruction_order(
                baseline, refreshed, self.root / "rejected.docx"
            )

    def test_libreoffice_page_margin_change_is_rejected_not_restored(self) -> None:
        baseline = self.root / "baseline.docx"
        refreshed = self.root / "refreshed.docx"
        restored = self.root / "restored.docx"
        self._toc_docx(baseline, 'TOC \\o "1-3" \\h \\z')
        self._toc_docx(refreshed, 'TOC \\z \\o "1-3" \\h')
        document = DocxDocument(refreshed)
        document.sections[0].top_margin = Inches(2)
        document.save(refreshed)
        with self.assertRaisesRegex(
            FormatMonographError,
            "changed pagination semantics or section boundaries",
        ):
            restore_known_libreoffice_toc_instruction_order(
                baseline, refreshed, restored
            )


if __name__ == "__main__":
    unittest.main()
