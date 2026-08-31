from __future__ import annotations

import argparse
import json
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
SCRIPTS = REPO / "format-monograph" / "scripts"
sys.path.insert(0, str(SCRIPTS))

from backend_evidence import (  # noqa: E402
    BackendEvidenceError,
    read_bound_backend_audit,
)
from test_v024_finalization import V024FinalizationTests  # noqa: E402

SCRIPT = REPO / "format-monograph" / "scripts" / "render_docx.py"
SAFE_LIBREOFFICE_UNAVAILABLE_ERROR = (
    "LibreOffice field refresh requires the verified macOS internal-Python "
    "macro host; the legacy UNO server/helper backend is disabled."
)


def load_field_backend_audit(finalization: dict) -> dict:
    """Load only a strictly bound, bounded diagnostic sidecar."""
    return read_bound_backend_audit(finalization)


def field_finalization_summary(
    finalization: dict, backend_audit: dict | None = None
) -> dict:
    backend = finalization.get("field_backend", {})
    attempt = backend.get("attempt") or {}
    canonical_selective = (
        attempt.get("selective_writeback_status")
        or (backend.get("selective_writeback") or {}).get("status")
    )
    canonical_failure = attempt.get("failure") or {}
    raw_backend = backend_audit or {}
    attempted = raw_backend.get("attempted_backend", raw_backend)
    selective = attempted.get("selective_writeback") or {}
    failure = attempted.get("failure") or {}
    completion = finalization.get("field_completion", {})
    contract_or_integrity_rejection = (
        backend.get("backend") == "deferred_on_open"
        and attempt.get("backend") == "libreoffice_uno"
        and attempt.get("fallback_from") == "libreoffice_contract_or_integrity"
        and (
            canonical_selective in {"rejected", "error"}
            or canonical_failure.get("status") == "rejected"
        )
    )
    safe_backend_unavailable = (
        backend.get("backend") == "deferred_on_open"
        and attempt.get("backend") == "libreoffice_uno"
        and attempt.get("fallback_from") == "libreoffice_error"
        and canonical_selective is None
        and canonical_failure.get("status") == "rejected"
        and canonical_failure.get("stage") == "libreoffice_refresh"
        and canonical_failure.get("failed_checks") == ["libreoffice_refresh"]
        and failure.get("status") == "rejected"
        and failure.get("stage") == "libreoffice_refresh"
        and failure.get("failed_checks") == ["libreoffice_refresh"]
        and failure.get("error") == SAFE_LIBREOFFICE_UNAVAILABLE_ERROR
    )
    deferred_rejection = contract_or_integrity_rejection or safe_backend_unavailable
    nonfinal_backend = (
        backend.get("backend") == "libreoffice_uno"
        and finalization.get("delivery_field_status") == "libreoffice_refreshed"
        and finalization.get("field_writeback_status") == "libreoffice_selective"
        and canonical_selective == "libreoffice_selective"
    )
    return {
        "gate_outcome": (
            "libreoffice_nonfinal"
            if nonfinal_backend
            else "strictly_deferred"
            if deferred_rejection
            else "invalid"
        ),
        "backend_result_accepted": nonfinal_backend,
        "backend": backend.get("backend"),
        "attempted_backend": attempt.get("backend"),
        "fallback_from": attempt.get("fallback_from"),
        "delivery_field_status": finalization.get("delivery_field_status"),
        "field_writeback_status": finalization.get("field_writeback_status"),
        "selective_writeback": selective,
        "failure": failure,
        "approved_indexes_updated": attempted.get("approved_indexes_updated"),
        "document_index_count": attempted.get("document_index_count"),
        "toc_authorization_id": attempted.get("toc_authorization_id"),
        "instruction_restoration": attempted.get("instruction_restoration"),
        "contract_comparison": attempted.get("contract_comparison"),
        "delivery_field_contract_identical": backend.get(
            "delivery_field_contract_identical"
        ),
        "field_gate_completed": completion.get("field_gate_completed"),
        "final_ready_eligible": completion.get("final_ready_eligible"),
        "word_verification_required": completion.get(
            "word_verification_required"
        ),
        "word_verification_completed": completion.get(
            "word_verification_completed"
        ),
    }


def field_finalization_errors(
    finalization: dict, backend_audit: dict | None = None
) -> list[str]:
    summary = field_finalization_summary(finalization, backend_audit)
    errors = []
    if finalization.get("status") != "pass":
        errors.append("finalization status did not pass")
    if summary["gate_outcome"] == "libreoffice_nonfinal":
        if summary["delivery_field_contract_identical"] is not True:
            errors.append("delivery field contract was not identical to the baseline")
    elif summary["gate_outcome"] == "strictly_deferred":
        if summary["delivery_field_status"] != "deferred":
            errors.append("strict rejection did not retain deferred delivery status")
        if summary["field_writeback_status"] != "deferred":
            errors.append("strict rejection did not retain deferred writeback status")
    else:
        errors.append("LibreOffice backend neither completed non-finally nor deferred strictly")
    if summary["field_gate_completed"] is not False:
        errors.append("LibreOffice incorrectly completed the Word field gate")
    if summary["final_ready_eligible"] is not False:
        errors.append("LibreOffice incorrectly became final_ready eligible")
    if summary["word_verification_required"] is not True:
        errors.append("Word verification was not marked as required")
    if summary["word_verification_completed"] is not False:
        errors.append("Word verification was incorrectly marked complete")
    for name in (
        "content_integrity",
        "protected_object_integrity",
        "effective_font_integrity",
    ):
        if finalization.get(name) != "pass":
            errors.append(f"{name} did not pass")
    return errors


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--renderer")
    args = parser.parse_args()
    with tempfile.TemporaryDirectory() as temp_name:
        root = Path(temp_name)
        source = root / "render-smoke.docx"
        output = root / "pages"
        document = Document()
        document.add_heading("Render Smoke Test", level=1)
        document.add_paragraph("This synthetic page verifies LibreOffice and PyMuPDF.")
        document.save(source)

        render_command = [
            sys.executable,
            str(SCRIPT),
            str(source),
            "--output-dir",
            str(output),
        ]
        if args.renderer:
            render_command.extend(["--renderer", args.renderer])
        completed = subprocess.run(
            render_command,
            cwd=REPO / "format-monograph",
            capture_output=True,
            text=True,
            check=False,
        )
        if completed.returncode != 0:
            print(completed.stdout)
            print(completed.stderr, file=sys.stderr)
            return completed.returncode

        manifest = json.loads((output / "render-manifest.json").read_text(encoding="utf-8"))
        pages = [Path(path) for path in manifest["pages"]]
        if manifest["page_count"] < 1 or not pages:
            print("Renderer produced no pages.", file=sys.stderr)
            return 1
        if any(not page.is_file() or page.stat().st_size == 0 for page in pages):
            print("Renderer produced an empty or missing PNG.", file=sys.stderr)
            return 1

        case = V024FinalizationTests(
            methodName="test_deferred_finalization_requires_explicit_qa"
        )
        case.setUp()
        try:
            formatted = case.apply()
            finalized = case.root / "field-finalized.docx"
            status_path = case.root / "field-finalization.json"
            finalization_command = [
                sys.executable,
                str(REPO / "format-monograph" / "scripts" / "finalize_docx.py"),
                str(formatted),
                "--source",
                str(case.source),
                "--profile",
                str(case.profile),
                "--structure-map",
                str(case.structure),
                "--output",
                str(finalized),
                "--status-output",
                str(status_path),
                "--field-updater",
                "auto",
                "--approve-deferred",
            ]
            if args.renderer:
                finalization_command.extend(["--renderer", args.renderer])
            finalized_result = subprocess.run(
                finalization_command,
                cwd=REPO / "format-monograph",
                capture_output=True,
                text=True,
                check=False,
            )
            if finalized_result.returncode != 0:
                print(finalized_result.stdout)
                print(finalized_result.stderr, file=sys.stderr)
                return finalized_result.returncode
            finalization = json.loads(status_path.read_text(encoding="utf-8"))
            try:
                backend_audit = load_field_backend_audit(finalization)
            except BackendEvidenceError as exc:
                print(f"Backend audit validation failed: {exc}", file=sys.stderr)
                return 1
            summary = field_finalization_summary(finalization, backend_audit)
            print(
                json.dumps(
                    {"field_finalization_smoke": summary},
                    ensure_ascii=False,
                    sort_keys=True,
                )
            )
            errors = field_finalization_errors(finalization, backend_audit)
            if errors:
                print(
                    "LibreOffice field finalization smoke failed: "
                    + "; ".join(errors),
                    file=sys.stderr,
                )
                return 1
        finally:
            case.tearDown()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
