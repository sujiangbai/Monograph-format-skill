from __future__ import annotations

import copy
import hashlib
import struct
import sys
import tempfile
import unittest
import zipfile
import zlib
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

REPO = Path(__file__).resolve().parents[1]
SCRIPTS = REPO / "format-monograph" / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import (  # noqa: E402
    FormatMonographError,
    apply_field_properties,
    apply_style_properties,
    apply_style_rule_to_paragraphs,
    apply_table_properties,
    ensure_paragraph_style,
    style_effective_font,
    style_name_for_selector,
)
from audit_docx import audit_field_rule, audit_paragraph_rule, audit_table_rule  # noqa: E402
from docx_pagination import (  # noqa: E402
    apply_pagination_sections,
    audit_pagination_sections,
    _ensure_page_field,
    _page_field_count,
    _resolve_audit_boundary,
    pagination_inventory,
)
from finalize_docx import effective_font_failures  # noqa: E402
from structure_map import (  # noqa: E402
    apply_structure_map,
    approved_data_tables,
    approved_role_paragraphs,
    audit_caption_identifier_replacements,
    audit_structure_heading_operations,
    audit_structure_image_operations,
    audit_structure_table_operations,
    candidate_structure_map,
    prime_structure_map_locators,
    resolve_paragraph_locator,
    structure_content_inventory,
    structure_content_fingerprint,
    text_sha256,
)


def set_theme_fonts(element, ascii_theme: str, east_asia_theme: str) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for name in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.attrib.pop(qn(f"w:{name}"), None)
    r_fonts.set(qn("w:asciiTheme"), ascii_theme)
    r_fonts.set(qn("w:hAnsiTheme"), ascii_theme)
    r_fonts.set(qn("w:eastAsiaTheme"), east_asia_theme)


def synthetic_png(width: int, height: int) -> bytes:
    def chunk(kind: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload))
            + kind
            + payload
            + struct.pack(">I", zlib.crc32(kind + payload) & 0xFFFFFFFF)
        )

    scanline = b"\x00" + b"\x66\x99\xcc" * width
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(scanline * height, 9))
        + chunk(b"IEND", b"")
    )


def replace_theme_fonts(path: Path) -> None:
    temp = path.with_suffix(".theme.docx")
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(
        temp, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/theme/theme1.xml":
                root = etree.fromstring(data)
                namespace = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
                for family, typeface in (
                    ("majorFont", "Synthetic Heading Theme"),
                    ("minorFont", "Synthetic Body Theme"),
                ):
                    fonts = root.xpath(
                        f"//a:fontScheme/a:{family}/a:font[@script='Hans']",
                        namespaces=namespace,
                    )
                    if fonts:
                        fonts[0].set("typeface", typeface)
                data = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
            target.writestr(info, data)
    temp.replace(path)


def add_page_field(paragraph) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def page_field_element(cache: str = "CXLI") -> OxmlElement:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = cache
    run.append(text)
    field.append(run)
    return field


class V026DeterministicFontTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_theme_fonts_are_resolved_then_removed_for_controlled_roles(self) -> None:
        source = self.root / "theme-source.docx"
        document = Document()
        body = document.add_paragraph("Synthetic body")
        heading = document.add_paragraph("Synthetic heading", style="Heading 2")
        set_theme_fonts(document.styles["Normal"].element, "minorHAnsi", "minorEastAsia")
        set_theme_fonts(document.styles["Heading 2"].element, "majorHAnsi", "majorEastAsia")
        set_theme_fonts(body.runs[0]._r, "minorHAnsi", "minorEastAsia")
        set_theme_fonts(heading.runs[0]._r, "majorHAnsi", "majorEastAsia")
        document.save(source)
        replace_theme_fonts(source)

        document = Document(source)
        self.assertEqual(
            "Synthetic Body Theme",
            style_effective_font(document, document.styles["Normal"], "eastAsia")[0],
        )
        self.assertEqual(
            "Synthetic Heading Theme",
            style_effective_font(document, document.styles["Heading 2"], "eastAsia")[0],
        )

        body_rule = {
            "id": "FMT-BODY-TEST",
            "selector": {"kind": "paragraph_role", "value": "body_text"},
            "properties": {
                "font_name_ascii": "Times New Roman",
                "font_name_east_asia": "宋体",
                "font_name_complex_script": "Times New Roman",
            },
        }
        heading_rule = {
            "id": "FMT-HEAD-TEST",
            "selector": {"kind": "paragraph_role", "value": "level_2_section"},
            "properties": {
                "font_name_ascii": "Arial",
                "font_name_east_asia": "黑体",
                "font_name_complex_script": "Arial",
            },
        }
        apply_style_rule_to_paragraphs(document, body_rule, [document.paragraphs[0]])
        apply_style_rule_to_paragraphs(document, heading_rule, [document.paragraphs[1]])
        output = self.root / "deterministic.docx"
        document.save(output)

        formatted = Document(output)
        for style_name, ascii_name, east_asia_name in (
            ("Normal", "Times New Roman", "宋体"),
            ("Heading 2", "Arial", "黑体"),
        ):
            r_fonts = formatted.styles[style_name].element.rPr.rFonts
            self.assertEqual(ascii_name, r_fonts.get(qn("w:ascii")))
            self.assertEqual(east_asia_name, r_fonts.get(qn("w:eastAsia")))
            for attribute in (
                "asciiTheme",
                "hAnsiTheme",
                "eastAsiaTheme",
                "cstheme",
            ):
                self.assertIsNone(r_fonts.get(qn(f"w:{attribute}")))
        self.assertFalse(audit_paragraph_rule(formatted, body_rule, [formatted.paragraphs[0]]))
        self.assertFalse(
            audit_paragraph_rule(formatted, heading_rule, [formatted.paragraphs[1]])
        )
        profile = {
            "rules": [
                {**body_rule, "status": "approved", "application": "automatic"},
                {**heading_rule, "status": "approved", "application": "automatic"},
            ]
        }
        self.assertFalse(effective_font_failures(output, profile))

    def test_all_controlled_text_styles_receive_mixed_script_fonts(self) -> None:
        source = self.root / "role-styles.docx"
        document = Document()
        cases = (
            ("Normal", {"kind": "paragraph_role", "value": "body_text"}, "宋体", "Times New Roman"),
            ("Heading 1", {"kind": "paragraph_role", "value": "heading_1"}, "黑体", "Arial"),
            ("Heading 2", {"kind": "paragraph_role", "value": "heading_2"}, "黑体", "Arial"),
            ("Heading 3", {"kind": "paragraph_role", "value": "heading_3"}, "黑体", "Arial"),
            ("Heading 4", {"kind": "paragraph_role", "value": "heading_4"}, "宋体", "Times New Roman"),
            ("TOC 1", {"kind": "paragraph_role", "value": "toc_level_1"}, "宋体", "Times New Roman"),
            ("TOC 2", {"kind": "paragraph_role", "value": "toc_level_2"}, "宋体", "Times New Roman"),
            ("TOC 3", {"kind": "paragraph_role", "value": "toc_level_3"}, "宋体", "Times New Roman"),
            ("Caption", {"kind": "caption_role", "value": "all"}, "宋体", "Times New Roman"),
            ("Quote", {"kind": "paragraph_role", "value": "long_quote"}, "宋体", "Times New Roman"),
            ("Bibliography", {"kind": "bibliography_role", "value": "entries"}, "宋体", "Times New Roman"),
            ("Footnote Text", {"kind": "style_name", "value": "Footnote Text"}, "宋体", "Times New Roman"),
        )
        for style_name, _selector, _east_asia, _ascii in cases:
            style = ensure_paragraph_style(document, style_name)
            set_theme_fonts(style.element, "minorHAnsi", "minorEastAsia")
        document.save(source)
        replace_theme_fonts(source)

        document = Document(source)
        profile_rules = []
        for index, (style_name, selector, east_asia, ascii_name) in enumerate(cases):
            paragraph = document.add_paragraph(f"Synthetic role {index}", style=style_name)
            rule = {
                "id": f"FMT-ROLE-{index}",
                "selector": selector,
                "properties": {
                    "font_name_ascii": ascii_name,
                    "font_name_east_asia": east_asia,
                    "font_name_complex_script": ascii_name,
                },
            }
            apply_style_rule_to_paragraphs(document, rule, [paragraph])
            profile_rules.append(
                {**rule, "status": "approved", "application": "automatic"}
            )
        output = self.root / "role-styles-formatted.docx"
        document.save(output)

        formatted = Document(output)
        for style_name, _selector, east_asia, ascii_name in cases:
            self.assertEqual(
                ascii_name,
                style_effective_font(formatted, formatted.styles[style_name], "ascii")[0],
            )
            self.assertEqual(
                east_asia,
                style_effective_font(formatted, formatted.styles[style_name], "eastAsia")[0],
            )
        self.assertFalse(
            effective_font_failures(output, {"rules": profile_rules})
        )

    def test_page_footer_rebuild_is_idempotent_and_protects_other_content(self) -> None:
        document = Document()
        footer = document.sections[0].footer
        paragraph = footer.paragraphs[0]
        add_page_field(paragraph)
        add_page_field(paragraph)
        before = (
            etree.tostring(document.element),
            etree.tostring(document.settings.element),
            tuple(
                sorted(
                    (rel_id, rel.reltype, rel.target_ref, rel.is_external)
                    for rel_id, rel in document.part.rels.items()
                )
            ),
            tuple(
                sorted(
                    (rel_id, etree.tostring(part.element))
                    for rel_id, part in document.part.related_parts.items()
                    if hasattr(part, "element")
                    and str(getattr(part, "partname", "")).startswith(
                        ("/word/footer", "/word/header")
                    )
                )
            ),
        )
        with self.assertRaises(FormatMonographError):
            _ensure_page_field(footer, 2)
        after = (
            etree.tostring(document.element),
            etree.tostring(document.settings.element),
            tuple(
                sorted(
                    (rel_id, rel.reltype, rel.target_ref, rel.is_external)
                    for rel_id, rel in document.part.rels.items()
                )
            ),
            tuple(
                sorted(
                    (rel_id, etree.tostring(part.element))
                    for rel_id, part in document.part.related_parts.items()
                    if hasattr(part, "element")
                    and str(getattr(part, "partname", "")).startswith(
                        ("/word/footer", "/word/header")
                    )
                )
            ),
        )
        self.assertEqual(before, after)

        single = Document()
        single_footer = single.sections[0].footer
        add_page_field(single_footer.paragraphs[0])
        self.assertTrue(_ensure_page_field(single_footer, 2))
        self.assertEqual(1, _page_field_count(single_footer._element))
        self.assertEqual(1, len(single_footer.paragraphs))
        self.assertFalse(_ensure_page_field(single_footer, 2))
        self.assertEqual(1, _page_field_count(single_footer._element))

        protected = Document().sections[0].footer
        protected.paragraphs[0].text = "Synthetic chapter footer"
        with self.assertRaises(FormatMonographError):
            _ensure_page_field(protected, 2)

        static_page = Document().sections[0].footer
        static_page.paragraphs[0].text = "315"
        with self.assertRaises(FormatMonographError):
            _ensure_page_field(static_page, 2, replace_static_page_text=True)

        duplicate_path = self.root / "duplicate-footer.docx"
        duplicated = Document()
        add_page_field(duplicated.sections[0].footer.paragraphs[0])
        add_page_field(duplicated.sections[0].footer.paragraphs[0])
        duplicated.save(duplicate_path)
        inventory = pagination_inventory(duplicate_path)
        self.assertEqual(
            2, inventory["sections"][0]["footer_page_field_counts"]["default"]
        )

        legacy = Document().sections[0].footer
        legacy_paragraph = legacy.paragraphs[0]
        for container_name in ("drawing", "pict"):
            container = OxmlElement(f"w:{container_name}")
            text_box = OxmlElement("w:txbxContent")
            inner_paragraph = OxmlElement("w:p")
            inner_paragraph.append(page_field_element())
            text_box.append(inner_paragraph)
            container.append(text_box)
            if container_name == "pict":
                container.append(
                    etree.Element("{urn:schemas-microsoft-com:vml}imagedata")
                )
            legacy_paragraph._p.append(container)
        with self.assertRaises(FormatMonographError):
            _ensure_page_field(legacy, 2)
        self.assertTrue(legacy._element.xpath(".//w:drawing | .//w:pict"))

        linked_vml_footer = Document().sections[0].footer
        pict = OxmlElement("w:pict")
        text_box = OxmlElement("w:txbxContent")
        inner_paragraph = OxmlElement("w:p")
        inner_paragraph.append(page_field_element())
        text_box.append(inner_paragraph)
        pict.append(text_box)
        image_data = etree.Element("{urn:schemas-microsoft-com:vml}imagedata")
        image_data.set(qn("r:id"), "rIdSyntheticImage")
        pict.append(image_data)
        linked_vml_footer.paragraphs[0]._p.append(pict)
        with self.assertRaises(FormatMonographError):
            _ensure_page_field(linked_vml_footer, 2)
        self.assertTrue(linked_vml_footer._element.xpath(".//w:pict"))

        image_footer = Document().sections[0].footer
        drawing = OxmlElement("w:drawing")
        drawing.append(OxmlElement("a:blip"))
        image_footer.paragraphs[0]._p.append(drawing)
        with self.assertRaises(FormatMonographError):
            _ensure_page_field(image_footer, 2)

    def test_table_fonts_are_explicit_and_effectively_audited(self) -> None:
        output = self.root / "table-fonts.docx"
        document = Document()
        table = document.add_table(rows=2, cols=2)
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell.text = f"Synthetic {row_index}-{cell_index}"
        properties = {
            "font_name_ascii": "Times New Roman",
            "font_name_east_asia": "宋体",
            "font_name_complex_script": "Times New Roman",
        }
        apply_table_properties(
            document,
            properties,
            [(table, {"header_rows": [0], "repeat_header_rows": [0]})],
        )
        document.save(output)

        formatted = Document(output)
        rule = {
            "id": "FMT-TABLE-FONT-TEST",
            "selector": {"kind": "table_role", "value": "all"},
            "properties": {**properties, "border_preset": "preserve"},
        }
        self.assertFalse(
            audit_table_rule(formatted, rule, [(formatted.tables[0], {})])
        )
        profile = {
            "rules": [{**rule, "status": "approved", "application": "automatic"}]
        }
        self.assertFalse(effective_font_failures(output, profile))
        r_fonts = formatted.tables[0].cell(0, 0).paragraphs[0].runs[0]._r.rPr.rFonts
        self.assertEqual("Times New Roman", r_fonts.get(qn("w:cs")))
        for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            self.assertIsNone(r_fonts.get(qn(f"w:{attribute}")))

        caption_document = Document()
        caption_table = caption_document.add_table(rows=2, cols=1)
        caption = caption_table.cell(0, 0).paragraphs[0]
        caption.text = "Synthetic table caption"
        caption.style = "Caption"
        caption.alignment = 1
        caption_entry = {"caption_row": 0, "header_rows": [1]}
        apply_table_properties(
            caption_document,
            {**properties, "column_roles": ["narrative"]},
            [(caption_table, caption_entry)],
        )
        self.assertEqual(1, caption.alignment)

        inherited_document = Document()
        normal_ind = (
            inherited_document.styles["Normal"]
            .element.get_or_add_pPr()
            .get_or_add_ind()
        )
        normal_ind.set(qn("w:firstLineChars"), "200")
        inherited_table = inherited_document.add_table(rows=1, cols=2)
        preserved = inherited_table.cell(0, 0).paragraphs[0]
        preserved.text = "Synthetic short label 3"
        preserved.alignment = 1
        intentional = inherited_table.cell(0, 1).paragraphs[0]
        intentional.text = "Synthetic intentional indent"
        intentional.paragraph_format.first_line_indent = Pt(6)
        apply_table_properties(
            inherited_document,
            {"font_size_pt": 9, "line_spacing_pt": 15},
            [(inherited_table, {"header_rows": []})],
        )
        self.assertEqual(1, preserved.alignment)
        self.assertEqual("Monograph Table Text", preserved.style.name)
        preserved_style_ind = preserved.style.element.pPr.find(qn("w:ind"))
        self.assertEqual("0", preserved_style_ind.get(qn("w:firstLine")))
        self.assertEqual("0", preserved_style_ind.get(qn("w:firstLineChars")))
        self.assertEqual(6, intentional.paragraph_format.first_line_indent.pt)

        inherited_output = self.root / "table-no-indent-style.docx"
        inherited_document.save(inherited_output)
        reloaded = Document(inherited_output)
        reloaded_preserved = reloaded.tables[0].cell(0, 0).paragraphs[0]
        self.assertEqual("Monograph Table Text", reloaded_preserved.style.name)
        reloaded_ind = reloaded_preserved.style.element.pPr.find(qn("w:ind"))
        self.assertEqual("0", reloaded_ind.get(qn("w:firstLine")))
        self.assertEqual("0", reloaded_ind.get(qn("w:firstLineChars")))

    def test_approved_derived_footers_are_path_and_cache_independent(self) -> None:
        source = self.root / "source.docx"
        document = Document()
        document.add_paragraph("Synthetic content")
        document.sections[0].footer.paragraphs[0].text = "315"
        document.save(source)

        formatted = self.root / "formatted.docx"
        document = Document(source)
        footer = document.sections[0].footer
        footer.paragraphs[0].clear()
        add_page_field(footer.paragraphs[0])
        document.save(formatted)

        structure_map = {
            "schema_version": "1.4",
            "pagination_sections": {"approved": True},
            "trailing_empty_sections": [
                {
                    "section": 99,
                    "approved_delete": True,
                    "previous_boundary_paragraph": 0,
                    "previous_boundary_sha256": hashlib.sha256(
                        "Synthetic content".encode("utf-8")
                    ).hexdigest(),
                    "evidence": {"approved_derived_footer_only": True},
                }
            ],
            "headings": [],
            "captions": [],
            "tables": [],
            "toc_ranges": [],
        }
        self.assertEqual(
            structure_content_fingerprint(source, structure_map),
            structure_content_fingerprint(formatted, structure_map),
        )

    def test_static_toc_role_locators_are_not_audited_after_toc_rebuild(self) -> None:
        document = Document()
        document.add_paragraph("TOC marker")
        document.add_paragraph("Old static entry")
        document.add_paragraph("Synthetic body")
        structure_map = {
            "schema_version": "1.4",
            "toc_ranges": [
                {"approved": True, "start_paragraph": 0, "end_paragraph": 1}
            ],
            "paragraph_roles": [
                {
                    "approved": True,
                    "role": "body",
                    "text_sha256": "unused-after-toc-rebuild",
                    "locator": {"kind": "body_paragraph", "paragraph": 1},
                }
            ],
        }
        self.assertEqual(
            [],
            approved_role_paragraphs(
                document,
                structure_map,
                {"kind": "paragraph_role", "value": "body_text"},
            ),
        )
        self.assertEqual(
            "TOC 1",
            style_name_for_selector(
                {"kind": "paragraph_role", "value": "toc_level_1"}
            ),
        )

        toc_document = Document()
        toc_paragraph = toc_document.add_paragraph()
        toc_paragraph._p.append(page_field_element("Update directory"))
        toc_paragraph._p[-1].set(qn("w:instr"), 'TOC \\o "1-3"')

        def missing_locator(_document, _locator):
            raise FormatMonographError("Synthetic removed static TOC locator")

        resolved_toc = _resolve_audit_boundary(
                toc_document,
                {"kind": "body_paragraph", "paragraph": 1},
                missing_locator,
                allow_unique_toc_field=True,
            )
        self.assertIs(toc_paragraph._p, resolved_toc._p)

        normalized_paragraph = toc_document.add_paragraph("Synthetic normalized title")

        def normalized_locator(document, locator):
            if locator.get("text_sha256") == "normalized":
                return document.paragraphs[1]
            raise FormatMonographError("Synthetic original heading locator")

        resolved_body = _resolve_audit_boundary(
            toc_document,
            {"kind": "body_paragraph", "text_sha256": "original"},
            normalized_locator,
            alternate_locators=[
                {"kind": "body_paragraph", "text_sha256": "normalized"}
            ],
        )
        self.assertIs(normalized_paragraph._p, resolved_body._p)

    def test_word_adapter_uses_compatible_open_calls(self) -> None:
        adapter = (
            REPO
            / "adapters"
            / "microsoft-word"
            / "windows"
            / "word_field_updater.ps1"
        ).read_text(encoding="utf-8")
        self.assertIn("Documents.Open($outputPath)", adapter)
        self.assertIn("Documents.Open($inputPath, $false, $true)", adapter)
        self.assertIn("Documents.Open($outputPath, $false, $true)", adapter)
        self.assertIn("try { $word.Options.UpdateLinksAtOpen", adapter)
        self.assertIn('"verify_only"', adapter)
        self.assertIn("read_only_verified = $openedReadOnly", adapter)
        self.assertIn("structural_changes_applied = 0", adapter)
        self.assertNotIn("Normalize-FrontMatterPagination", adapter)
        self.assertNotIn("Normalize-DisplayedPageReferences", adapter)
        self.assertNotIn("Set-DisplayedPageField", adapter)

    def test_localized_caption_style_keeps_semantic_audit_contract(self) -> None:
        source = self.root / "caption-source.docx"
        original_text = "Fig 1.1 Synthetic caption"
        replacement_text = "Fig 1.2 Synthetic caption"
        document = Document()
        document.add_paragraph(original_text)
        document.save(source)

        output = self.root / "caption-localized.docx"
        document = Document()
        paragraph = document.add_paragraph(replacement_text, style="Caption")
        rule = {
            "id": "FMT-CAP-LOCALIZED",
            "selector": {"kind": "caption_role", "value": "all"},
            "properties": {
                "font_name_ascii": "Times New Roman",
                "font_name_east_asia": "Times New Roman",
                "font_name_complex_script": "Times New Roman",
            },
        }
        apply_style_rule_to_paragraphs(document, rule, [paragraph])
        style = document.styles["Caption"]
        style.element.set(qn("w:styleId"), "LocalizedCaption")
        style.element.find(qn("w:name")).set(
            qn("w:val"), "Localized Caption"
        )
        paragraph._p.get_or_add_pPr().get_or_add_pStyle().val = "LocalizedCaption"
        document.save(output)

        start, end = 0, len("Fig 1.1")
        title_start = end + 1
        entry = {
            "approved": True,
            "role": "figure_caption",
            "text_sha256": hashlib.sha256(original_text.encode("utf-8")).hexdigest(),
            "locator": {
                "kind": "body_paragraph",
                "paragraph": 0,
                "text_sha256": hashlib.sha256(
                    original_text.encode("utf-8")
                ).hexdigest(),
            },
        }
        caption = {
            "approved": True,
            "action": "replace_identifier",
            "locator": entry["locator"],
            "text_sha256": entry["text_sha256"],
            "identifier_span": {"start": start, "end": end},
            "identifier_sha256": hashlib.sha256(
                original_text[start:end].encode("utf-8")
            ).hexdigest(),
            "identifier_prefix_sha256": hashlib.sha256(b"").hexdigest(),
            "identifier_suffix_sha256": hashlib.sha256(
                original_text[end:].encode("utf-8")
            ).hexdigest(),
            "title_span_start": title_start,
            "title_text_sha256": hashlib.sha256(
                original_text[title_start:].encode("utf-8")
            ).hexdigest(),
            "replacement_identifier": "Fig 1.2",
        }
        structure_map = {
            "schema_version": "1.4",
            "paragraph_roles": [entry],
            "captions": [caption],
            "toc_ranges": [],
        }
        formatted = Document(output)
        targets = approved_role_paragraphs(
            formatted, structure_map, rule["selector"]
        )
        self.assertEqual([replacement_text], [target.text for target in targets])
        profile = {
            "rules": [{**rule, "status": "approved", "application": "automatic"}]
        }
        self.assertFalse(effective_font_failures(output, profile, structure_map))
        self.assertEqual(
            "pass",
            audit_caption_identifier_replacements(
                source, output, structure_map
            )[0]["status"],
        )

    def test_unresolvable_semantic_font_target_is_an_integrity_failure(self) -> None:
        output = self.root / "missing-semantic-target.docx"
        document = Document()
        document.add_paragraph("Synthetic body")
        document.save(output)

        locator = {
            "kind": "table_cell_paragraph",
            "table": 1,
            "row": 0,
            "cell": 0,
            "paragraph": 0,
        }
        structure_map = {
            "schema_version": "1.4",
            "paragraph_roles": [
                {
                    "approved": True,
                    "role": "table_caption",
                    "text_sha256": hashlib.sha256(
                        b"Synthetic missing caption"
                    ).hexdigest(),
                    "locator": locator,
                }
            ],
            "captions": [
                {
                    "approved": True,
                    "action": "style_only",
                    "locator": locator,
                }
            ],
            "toc_ranges": [],
        }
        profile = {
            "rules": [
                {
                    "id": "FMT-CAP-MISSING-TARGET",
                    "selector": {"kind": "caption_role", "value": "all"},
                    "properties": {"font_name_ascii": "Times New Roman"},
                    "status": "approved",
                    "application": "automatic",
                }
            ]
        }
        self.assertEqual(
            [
                {
                    "rule": "FMT-CAP-MISSING-TARGET",
                    "selector": "caption_role",
                    "reason": "semantic_target_unresolvable",
                }
            ],
            effective_font_failures(output, profile, structure_map),
        )

    def test_pagination_requires_existing_exact_boundaries_before_mutation(self) -> None:
        document = Document()
        document.add_paragraph("Synthetic TOC")
        body = document.add_paragraph("Synthetic chapter", style="Heading 1")
        document.styles["Heading 1"].paragraph_format.page_break_before = True
        settings = {
            "approved": True,
            "toc_start": {"text": "Synthetic TOC"},
            "body_start": {"text": "Synthetic chapter"},
            "number_format": "decimal",
            "start_at": {"toc": 1, "body": 1},
            "continue_after_body_start": True,
        }

        def resolver(doc, locator):
            return next(p for p in doc.paragraphs if p.text == locator["text"])

        before = etree.tostring(document.element)
        with self.assertRaisesRegex(FormatMonographError, "existing, exact section boundary"):
            apply_pagination_sections(document, settings, resolver)
        self.assertEqual(before, etree.tostring(document.element))
        self.assertIsNone(body.paragraph_format.page_break_before)

        plain = Document()
        plain.add_paragraph("Synthetic TOC")
        plain_body = plain.add_paragraph("Synthetic chapter")
        with self.assertRaisesRegex(FormatMonographError, "existing, exact section boundary"):
            apply_pagination_sections(plain, settings, resolver)
        self.assertIsNone(plain_body.paragraph_format.page_break_before)

    def test_front_matter_blank_block_spacing_and_table_rules_are_structural(self) -> None:
        source = self.root / "front-matter-tables.docx"
        document = Document()
        title = document.add_paragraph("Synthetic whole-book title")
        title_p_pr = title._p.get_or_add_pPr()
        title_ind = OxmlElement("w:ind")
        for attribute, value in (
            ("left", "720"),
            ("leftChars", "200"),
            ("right", "360"),
            ("rightChars", "100"),
            ("firstLine", "240"),
            ("hangingChars", "100"),
        ):
            title_ind.set(qn(f"w:{attribute}"), value)
        title_p_pr.append(title_ind)
        title_num_pr = OxmlElement("w:numPr")
        title_p_pr.append(title_num_pr)
        title_section = document.sections[0]._sectPr
        vertical = OxmlElement("w:vAlign")
        vertical.set(qn("w:val"), "center")
        title_section.append(vertical)
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("[[TOC]]")
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("第1章 Synthetic body")
        table = document.add_table(rows=5, cols=4)
        table.cell(0, 0).merge(table.cell(0, 3)).text = "表 1.1 Synthetic table"
        table.cell(1, 0).merge(table.cell(1, 1)).text = "Grouped header"
        table.cell(1, 2).merge(table.cell(2, 2)).text = "Merged Y"
        table.cell(1, 3).merge(table.cell(2, 3)).text = "Merged Z"
        table.cell(2, 0).text = "Header A"
        table.cell(2, 1).text = "Header B"
        for row_index in (3, 4):
            for column_index in range(4):
                table.cell(row_index, column_index).text = f"R{row_index}C{column_index}"
        document.add_paragraph("Synthetic following body")
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        structure["pagination_sections"]["approved"] = True
        structure["front_matter"]["approved"] = True
        structure["block_spacing"]["approved"] = True
        table_entry = structure["tables"][0]
        table_entry.update(
            {
                "approved": True,
                "kind": "data",
                "caption_row": 0,
                "header_rows": [1, 2],
                "repeat_header_rows": [1, 2],
            }
        )
        table_entry["visual"].update(
            {
                "approved": True,
                "border_preset": "technical_textbook",
                "column_roles": ["narrative"] * 4,
            }
        )

        formatted = Document(source)
        prime_structure_map_locators(formatted, structure)
        apply_structure_map(formatted, structure)
        apply_table_properties(
            formatted,
            {"border_preset": "technical_textbook"},
            approved_data_tables(formatted, structure),
        )
        apply_structure_map(formatted, structure)
        output = self.root / "front-matter-tables-formatted.docx"
        formatted.save(output)

        source_inventory = structure_content_inventory(source, structure)
        output_inventory = structure_content_inventory(output, structure)
        self.assertEqual(source_inventory, output_inventory)
        reloaded = Document(output)
        self.assertEqual("Monograph Book Title", reloaded.paragraphs[0].style.name)
        title_p_pr = reloaded.paragraphs[0]._p.pPr
        self.assertIsNone(title_p_pr.find(qn("w:numPr")))
        self.assertIsNone(title_p_pr.find(qn("w:sectPr")))
        toc_paragraph = next(
            paragraph for paragraph in reloaded.paragraphs if paragraph.text == "[[TOC]]"
        )
        self.assertNotEqual("Monograph TOC Heading", toc_paragraph.style.name)
        direct_title_ind = title_p_pr.find(qn("w:ind"))
        if direct_title_ind is not None:
            self.assertTrue(
                all(value == "0" for value in direct_title_ind.attrib.values())
            )
        title_style = reloaded.paragraphs[0].style
        self.assertTrue(title_style.font.bold)
        self.assertEqual(22, title_style.font.size.pt)
        self.assertEqual(
            WD_LINE_SPACING.AT_LEAST,
            title_style.paragraph_format.line_spacing_rule,
        )
        self.assertEqual(33, title_style.paragraph_format.line_spacing.pt)
        title_style_ind = title_style.element.pPr.find(qn("w:ind"))
        for attribute in (
            "left",
            "leftChars",
            "right",
            "rightChars",
            "firstLine",
            "firstLineChars",
        ):
            self.assertEqual("0", title_style_ind.get(qn(f"w:{attribute}")))
        self.assertEqual(
            "center",
            reloaded.sections[0]
            ._sectPr.find(qn("w:vAlign"))
            .get(qn("w:val")),
        )
        self.assertEqual(
            "Times New Roman",
            style_effective_font(reloaded, title_style, "ascii")[0],
        )
        self.assertEqual("黑体", style_effective_font(reloaded, title_style, "eastAsia")[0])
        self.assertEqual(
            0,
            sum(
                paragraph.style is not None
                and paragraph.style.name == "Monograph TOC Heading"
                for paragraph in reloaded.paragraphs
            ),
        )
        self.assertEqual(
            ["nextPage", "nextPage"],
            [
                (
                    "nextPage"
                    if section._sectPr.find(qn("w:type")) is None
                    else section._sectPr.find(qn("w:type")).get(qn("w:val"))
                )
                for section in reloaded.sections[:2]
            ],
        )
        pagination_failures, _ = audit_pagination_sections(
            output,
            reloaded,
            structure["pagination_sections"],
            resolve_paragraph_locator,
            structure,
        )
        self.assertFalse(pagination_failures)
        self.assertEqual(
            1,
            sum(
                paragraph.style is not None
                and paragraph.style.name == "Monograph Figure Table Spacer"
                for paragraph in reloaded.paragraphs
            ),
        )
        audit_entry = dict(table_entry)
        audit_entry["visual"] = {"approved": False}
        self.assertFalse(
            audit_table_rule(
                reloaded,
                {"properties": {"border_preset": "technical_textbook"}},
                [(reloaded.tables[0], audit_entry)],
            )
        )
        borders = reloaded.tables[0]._tbl.tblPr.find(qn("w:tblBorders"))
        expected = {
            "top": ("nil", "0"),
            "bottom": ("single", "8"),
            "left": ("nil", "0"),
            "right": ("nil", "0"),
            "insideH": ("nil", "0"),
            "insideV": ("single", "4"),
        }
        for name, (style, size) in expected.items():
            border = borders.find(qn(f"w:{name}"))
            self.assertEqual(style, border.get(qn("w:val")))
            self.assertEqual(size, border.get(qn("w:sz")))
        for row in reloaded.tables[0].rows:
            for cell in row.cells:
                shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
                self.assertIsNotNone(shading)
                self.assertEqual("auto", shading.get(qn("w:fill")))

    def test_front_matter_blocks_authored_content_between_title_and_toc(self) -> None:
        source = self.root / "front-matter-authored-content.docx"
        document = Document()
        document.add_paragraph("Synthetic whole-book title")
        document.add_paragraph("Authored front-matter content")
        document.add_paragraph("[[TOC]]")
        document.add_paragraph("第1章 Synthetic body")
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        structure["pagination_sections"]["approved"] = True
        structure["front_matter"]["approved"] = True
        formatted = Document(source)
        prime_structure_map_locators(formatted, structure)
        with self.assertRaisesRegex(
            FormatMonographError,
            "existing, exact section boundary",
        ):
            apply_structure_map(formatted, structure)

    def test_technical_table_rebuilds_cell_borders_and_adds_semantic_separator(self) -> None:
        document = Document()
        table = document.add_table(rows=4, cols=2)
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell.text = f"R{row_index}C{column_index}"
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                for name in ("top", "left", "bottom", "right"):
                    border = OxmlElement(f"w:{name}")
                    border.set(qn("w:val"), "single")
                    border.set(qn("w:sz"), "12")
                    borders.append(border)
                tc_pr.append(borders)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "00FF00")
                tc_pr.append(shading)

        properties = {
            "border_preset": "technical_textbook",
            "major_border_pt": 1,
            "minor_border_pt": 0.5,
            "inside_vertical_borders": True,
            "horizontal_rule_rows": [3],
        }
        entry = {"header_rows": [0], "visual": {"approved": False}}
        apply_table_properties(document, properties, [(table, entry)])

        table_borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        expected = {
            "top": ("single", "8"),
            "bottom": ("single", "8"),
            "left": ("nil", "0"),
            "right": ("nil", "0"),
            "insideH": ("nil", "0"),
            "insideV": ("single", "4"),
        }
        for name, (style, size) in expected.items():
            border = table_borders.find(qn(f"w:{name}"))
            self.assertEqual(style, border.get(qn("w:val")))
            self.assertEqual(size, border.get(qn("w:sz")))
        for cell in table.rows[3].cells:
            border = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders")).find(
                qn("w:top")
            )
            self.assertEqual("single", border.get(qn("w:val")))
            self.assertEqual("4", border.get(qn("w:sz")))
        self.assertFalse(audit_table_rule(document, {"properties": properties}, [(table, entry)]))

    def test_figure_panel_unnumbered_caption_and_cell_cleanup_are_structural(self) -> None:
        source = self.root / "figure-panel-cleanup.docx"
        document = Document()
        image = document.add_paragraph()
        image._p.append(OxmlElement("w:drawing"))
        unnumbered = document.add_paragraph("Synthetic waveform A")
        unnumbered.alignment = WD_ALIGN_PARAGRAPH.CENTER

        panel = document.add_table(rows=2, cols=2)
        for cell in panel.rows[0].cells:
            cell.paragraphs[0]._p.append(OxmlElement("w:drawing"))
        for index, cell in enumerate(panel.rows[1].cells):
            cell.text = f"Panel {index + 1}"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        data = document.add_table(rows=2, cols=2)
        data.cell(0, 0).text = "Header A"
        data.cell(0, 1).text = "Header B"
        cleanup_cell = data.cell(1, 0)
        cleanup_cell.add_paragraph("Conclusion")
        data.cell(1, 1).text = "Result"
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        panel_entry = structure["tables"][0]
        self.assertEqual("layout", panel_entry["kind"])
        self.assertEqual("figure_panel", panel_entry["layout_purpose"])
        panel_entry["approved"] = True
        panel_entry["pagination_only"] = False
        panel_entry["visual"]["approved"] = True
        for role in structure["paragraph_roles"]:
            if role["role"] in {"figure_caption_unnumbered", "figure_panel_label"}:
                role["approved"] = True

        data_entry = structure["tables"][1]
        data_entry.update({"approved": True, "kind": "data", "header_rows": [0]})
        structure["table_cell_cleanups"] = [
            {
                "table": 1,
                "row": 1,
                "cell": 0,
                "action": "remove_leading_empty_paragraphs",
                "count": 1,
                "table_text_sha256": data_entry["table_text_sha256"],
                "cell_text_sha256": text_sha256(cleanup_cell.text),
                "result_cell_text_sha256": text_sha256("Conclusion"),
                "approved": True,
            }
        ]

        formatted = Document(source)
        prime_structure_map_locators(formatted, structure)
        apply_structure_map(formatted, structure)
        apply_structure_map(formatted, structure)
        output = self.root / "figure-panel-cleanup-formatted.docx"
        formatted.save(output)

        reloaded = Document(output)
        self.assertEqual(WD_TABLE_ALIGNMENT.CENTER, reloaded.tables[0].alignment)
        self.assertTrue(
            all(
                cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for row in reloaded.tables[0].rows
                for cell in row.cells
            )
        )
        self.assertTrue(
            all(
                paragraph.style.name == "Caption"
                for cell in reloaded.tables[0].rows[1].cells
                for paragraph in cell.paragraphs
                if paragraph.text
            )
        )
        self.assertEqual("Conclusion", reloaded.tables[1].cell(1, 0).text)
        self.assertEqual(
            structure_content_fingerprint(source, structure),
            structure_content_fingerprint(output, structure),
        )
        reloaded.styles["Caption"].paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )
        reloaded.styles["Caption"].element.find(qn("w:name")).set(
            qn("w:val"), "Localized Figure Label"
        )
        for cell in reloaded.tables[0].rows[1].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = None
        self.assertFalse(audit_structure_table_operations(reloaded, structure))
        roles = {
            role["role"]
            for role in structure["paragraph_roles"]
            if role["approved"]
        }
        self.assertIn("figure_caption_unnumbered", roles)
        self.assertIn("figure_panel_label", roles)

    def test_image_resizing_preserves_media_anchors_order_and_table_position(self) -> None:
        source = self.root / "image-anchor-source.docx"
        image_path = self.root / "synthetic-image.png"
        image_path.write_bytes(synthetic_png(600, 400))

        document = Document()
        document.add_paragraph("Synthetic chapter", style="Heading 1")
        standalone = document.add_paragraph()
        standalone.add_run().add_picture(str(image_path), width=Inches(1))
        document.add_paragraph("Synthetic following paragraph")
        panel = document.add_table(rows=2, cols=2)
        panel.cell(0, 0).paragraphs[0].add_run().add_picture(
            str(image_path), width=Inches(1)
        )
        panel.cell(0, 1).paragraphs[0].add_run().add_picture(
            str(image_path), width=Inches(1.2)
        )
        panel.cell(1, 0).text = "Panel A"
        panel.cell(1, 1).text = "Panel B"
        document.add_paragraph("Synthetic tail")
        document.save(source)

        structure = candidate_structure_map(source)
        structure["status"] = "approved"
        panel_entry = structure["tables"][0]
        self.assertEqual("figure_panel", panel_entry["layout_purpose"])
        panel_entry["approved"] = True
        panel_entry["position_policy"] = "preserve_anchor"
        panel_entry["visual"]["approved"] = True
        for image in structure["images"]:
            self.assertEqual("preserve_anchor", image["position_policy"])
            if image["placement"] in {"standalone", "table_figure_panel"}:
                self.assertTrue(image["supported"])
                image["approved"] = True
                image["resize"]["approved"] = True

        formatted = Document(source)
        # A derived empty paragraph may share the image paragraph's text hash and
        # shift its numeric index. Media plus unchanged authored neighbors must
        # still resolve the original image without moving it.
        formatted.paragraphs[1]._p.addprevious(OxmlElement("w:p"))
        body_children_before = list(formatted.element.body)
        standalone_paragraph = formatted.paragraphs[2]
        panel_table = formatted.tables[0]
        panel_paragraphs = [
            panel_table.cell(0, 0).paragraphs[0],
            panel_table.cell(0, 1).paragraphs[0],
        ]
        standalone_parent = standalone_paragraph._p.getparent()
        standalone_index = standalone_parent.index(standalone_paragraph._p)
        panel_parent = panel_table._tbl.getparent()
        panel_index = panel_parent.index(panel_table._tbl)
        panel_locations = [
            (paragraph._p.getparent(), paragraph._p.getparent().index(paragraph._p))
            for paragraph in panel_paragraphs
        ]
        original_extents = {
            image["image"]: dict(image["source_extent_emu"])
            for image in structure["images"]
            if image["approved"]
        }
        shifted_source = self.root / "image-anchor-shifted-source.docx"
        formatted.save(shifted_source)

        prime_structure_map_locators(formatted, structure)
        result = apply_structure_map(formatted, structure)
        self.assertEqual(3, next(change["targets"] for change in result if change["kind"] == "structure_image_resize"))
        self.assertEqual(body_children_before, list(formatted.element.body))
        self.assertIs(standalone_parent, standalone_paragraph._p.getparent())
        self.assertEqual(standalone_index, standalone_parent.index(standalone_paragraph._p))
        self.assertIs(panel_parent, panel_table._tbl.getparent())
        self.assertEqual(panel_index, panel_parent.index(panel_table._tbl))
        for paragraph, (parent, index) in zip(panel_paragraphs, panel_locations):
            self.assertIs(parent, paragraph._p.getparent())
            self.assertEqual(index, parent.index(paragraph._p))

        output = self.root / "image-anchor-formatted.docx"
        formatted.save(output)
        reloaded = Document(output)
        self.assertFalse(audit_structure_image_operations(reloaded, structure))
        resized_extents = {
            image["image"]: reloaded.part.document.element.xpath(
                ".//w:drawing"
            )[index]
            .xpath("./wp:inline/wp:extent")[0]
            for index, image in enumerate(
                [item for item in structure["images"] if item["approved"]]
            )
        }
        self.assertTrue(
            all(
                int(extent.get("cx"))
                <= int(original_extents[image_id]["cx"])
                and int(extent.get("cy"))
                <= int(original_extents[image_id]["cy"])
                for image_id, extent in resized_extents.items()
            )
        )
        panel_heights = [
            int(extent.get("cy"))
            for image_id, extent in resized_extents.items()
            if next(
                image
                for image in structure["images"]
                if image["image"] == image_id
            )["placement"]
            == "table_figure_panel"
        ]
        self.assertEqual(1, len(set(panel_heights)))
        self.assertEqual(
            structure_content_fingerprint(shifted_source, structure),
            structure_content_fingerprint(output, structure),
        )

    def test_heading_numbering_inherits_mixed_fonts_size_weight_and_zero_indent(self) -> None:
        document = Document()
        specifications = (
            ("Heading 1", "黑体", "Times New Roman", 18),
            ("Heading 2", "黑体", "Times New Roman", 14),
            ("Heading 3", "黑体", "Times New Roman", 12),
            ("Heading 4", "宋体", "Times New Roman", 10.5),
        )
        for style_name, east_asia, ascii_name, size in specifications:
            style = ensure_paragraph_style(document, style_name)
            apply_style_properties(
                style,
                {
                    "font_name_east_asia": east_asia,
                    "font_name_ascii": ascii_name,
                    "font_name_complex_script": ascii_name,
                    "font_size_pt": size,
                    "bold": True,
                    "first_line_indent_chars": 2,
                },
            )
            paragraph = document.add_paragraph("Synthetic heading", style=style_name)
            direct_ind = paragraph._p.get_or_add_pPr().get_or_add_ind()
            for attribute, value in (
                ("left", "720"),
                ("leftChars", "200"),
                ("right", "360"),
                ("rightChars", "100"),
                ("firstLine", "240"),
                ("hangingChars", "100"),
            ):
                direct_ind.set(qn(f"w:{attribute}"), value)
            direct_num_pr = OxmlElement("w:numPr")
            paragraph._p.get_or_add_pPr().append(direct_num_pr)

        properties = {
            "rebuild_heading_numbering": True,
            "heading_levels": 4,
            "chapter_start": 4,
        }
        apply_field_properties(document, properties)
        output = self.root / "heading-numbering-format.docx"
        document.save(output)
        reloaded = Document(output)

        self.assertFalse(
            audit_field_rule(
                reloaded,
                {"properties": properties},
                chapter_start=4,
                path=output,
            )
        )
        for paragraph in reloaded.paragraphs:
            self.assertIsNone(paragraph.paragraph_format.first_line_indent)
            self.assertIsNone(paragraph._p.pPr.find(qn("w:numPr")))
            direct_ind = paragraph._p.pPr.find(qn("w:ind"))
            if direct_ind is not None:
                self.assertFalse(direct_ind.attrib)
        for style_name, _east_asia, _ascii_name, _size in specifications:
            ind = reloaded.styles[style_name].element.pPr.find(qn("w:ind"))
            for attribute in (
                "left",
                "leftChars",
                "right",
                "rightChars",
                "firstLine",
                "firstLineChars",
            ):
                self.assertEqual("0", ind.get(qn(f"w:{attribute}")))

        for level in range(4):
            numbering_level = reloaded.styles[f"Heading {level + 1}"].element.pPr.find(
                qn("w:numPr")
            )
            self.assertIsNotNone(numbering_level)
        structure = {
            "schema_version": "1.5",
            "numbering": {"approved": True},
            "headings": [
                {
                    "approved": True,
                    "level": level,
                    "paragraph": level - 1,
                    "text_sha256": text_sha256(reloaded.paragraphs[level - 1].text),
                }
                for level in range(1, 5)
            ],
        }
        self.assertFalse(audit_structure_heading_operations(reloaded, structure))

    def test_unapproved_direct_heading_numbering_is_preserved_and_reported(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("1.1 Synthetic heading", style="Heading 2")
        apply_field_properties(
            document,
            {
                "rebuild_heading_numbering": True,
                "heading_levels": 4,
                "chapter_start": 1,
            },
        )
        style_num_pr = document.styles["Heading 2"].element.pPr.find(qn("w:numPr"))
        paragraph._p.get_or_add_pPr().append(copy.deepcopy(style_num_pr))
        num_id = style_num_pr.find(qn("w:numId")).get(qn("w:val"))
        abstract_id = next(
            item
            for item in document.part.numbering_part.element.findall(qn("w:num"))
            if item.get(qn("w:numId")) == num_id
        ).find(qn("w:abstractNumId")).get(qn("w:val"))
        abstract = next(
            item
            for item in document.part.numbering_part.element.findall(qn("w:abstractNum"))
            if item.get(qn("w:abstractNumId")) == abstract_id
        )
        level = next(
            item
            for item in abstract.findall(qn("w:lvl"))
            if item.get(qn("w:ilvl")) == "1"
        )
        level.find(qn("w:pPr")).find(qn("w:ind")).set(qn("w:left"), "720")
        structure = {
            "schema_version": "1.5",
            "numbering": {"approved": False},
            "headings": [
                {
                    "approved": True,
                    "level": 2,
                    "paragraph": 0,
                    "text_sha256": text_sha256(paragraph.text),
                }
            ],
        }
        apply_structure_map(document, structure)
        self.assertIsNotNone(paragraph._p.pPr.find(qn("w:numPr")))
        failures = audit_structure_heading_operations(document, structure)
        self.assertIn(
            "heading_direct_numbering_requires_qa",
            {failure["property"] for failure in failures},
        )


if __name__ == "__main__":
    unittest.main()
