from __future__ import annotations

import base64
import contextlib
import ctypes
import errno
import io
import json
import os
import subprocess
import sys
import tempfile
import unittest
import zipfile
import stat
from unittest.mock import MagicMock, patch
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "format-monograph"
SCRIPTS = SKILL / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import FormatMonographError, NS, field_cache_inventory  # noqa: E402
import finalize_docx  # noqa: E402
from finalize_docx import (  # noqa: E402
    field_contract_preserved,
    resolve_finalization_path_contract,
)
from structure_map import (  # noqa: E402
    candidate_structure_map,
    structure_content_fingerprint,
    text_sha256,
)
from test_v11_execution import approved_v11_profile  # noqa: E402


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


class SimulatedWindowsPublisherApi:
    """Windows capability-set simulation backed by the POSIX host rename."""

    FILE_ATTRIBUTE_DIRECTORY = 0x00000010
    FILE_ATTRIBUTE_REPARSE_POINT = 0x00000400

    def __init__(self) -> None:
        self.open_handles: set[int] = set()
        self.move_fault = None

    def open_directory(self, path: Path) -> int:
        details = path.lstat()
        if path.is_symlink() or not stat.S_ISDIR(details.st_mode):
            raise finalize_docx.FinalizationPublishError(
                f"simulated Windows authority rejected {path}"
            )
        descriptor = os.open(
            path,
            os.O_RDONLY | os.O_DIRECTORY | os.O_NOFOLLOW | os.O_CLOEXEC,
        )
        self.open_handles.add(descriptor)
        return descriptor

    def query_directory(self, handle: int) -> tuple[int, int, int]:
        details = os.fstat(handle)
        if not stat.S_ISDIR(details.st_mode):
            raise finalize_docx.FinalizationPublishError(
                "simulated Windows directory handle changed type"
            )
        return self.FILE_ATTRIBUTE_DIRECTORY, details.st_dev, details.st_ino

    def close(self, handle: int) -> None:
        os.close(handle)
        self.open_handles.discard(handle)

    @staticmethod
    def path_is_reparse(path: Path) -> bool:
        return path.is_symlink()

    @staticmethod
    def private_security_identity(path: Path) -> str:
        details = path.lstat()
        if details.st_uid != os.geteuid() or stat.S_IMODE(details.st_mode) & 0o022:
            raise finalize_docx.FinalizationPublishError(
                f"simulated Windows private ACL rejected {path}"
            )
        return f"private:{details.st_uid}:{stat.S_IMODE(details.st_mode):o}"

    def atomic_noreplace_move(self, source: Path, target: Path) -> None:
        if self.move_fault is not None and self.move_fault(source, target):
            raise self.move_fault.error
        library = ctypes.CDLL(None, use_errno=True)
        if sys.platform == "darwin":
            function_name = "renameatx_np"
            flag = finalize_docx.DARWIN_RENAME_EXCL
        elif sys.platform.startswith("linux"):
            function_name = "renameat2"
            flag = finalize_docx.LINUX_RENAME_NOREPLACE
        else:
            raise finalize_docx.FinalizationPublishError(
                "simulated Windows publisher requires a supported POSIX "
                "atomic no-replace host primitive"
            )
        function = getattr(library, function_name, None)
        if function is None:
            raise finalize_docx.FinalizationPublishError(
                f"simulated Windows publisher cannot resolve {function_name}"
            )
        function.argtypes = (
            ctypes.c_int,
            ctypes.c_char_p,
            ctypes.c_int,
            ctypes.c_char_p,
            ctypes.c_uint,
        )
        function.restype = ctypes.c_int
        if function(
            -2,
            os.fsencode(source),
            -2,
            os.fsencode(target),
            flag,
        ):
            finalize_docx._raise_atomic_move_error(
                ctypes.get_errno(), source, target
            )


def run_element(name: str, text: str | None = None) -> etree._Element:
    run = etree.Element(qn("w:r"))
    child = etree.SubElement(run, qn(name))
    if text is not None:
        child.text = text
    return run


def field_char(kind: str) -> etree._Element:
    run = etree.Element(qn("w:r"))
    child = etree.SubElement(run, qn("w:fldChar"))
    child.set(qn("w:fldCharType"), kind)
    return run


def set_paragraph_style(paragraph: etree._Element, style: str) -> None:
    p_pr = paragraph.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = etree.Element(qn("w:pPr"))
        paragraph.insert(0, p_pr)
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = etree.SubElement(p_pr, qn("w:pStyle"))
    p_style.set(qn("w:val"), style)


def expand_toc_cache(source: Path, output: Path) -> None:
    temp = output.with_suffix(".tmp")
    with zipfile.ZipFile(source) as package, zipfile.ZipFile(
        temp, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for info in package.infolist():
            data = package.read(info.filename)
            if info.filename == "word/document.xml":
                root = etree.fromstring(data)
                body = root.xpath("/w:document/w:body", namespaces=NS)[0]
                toc = body.xpath(
                    "./w:p[.//w:instrText[contains(translate(., 'toc', 'TOC'), 'TOC')]]",
                    namespaces=NS,
                )[0]
                for child in list(toc):
                    if child.tag != qn("w:pPr"):
                        toc.remove(child)
                set_paragraph_style(toc, "TOC1")
                toc.extend(
                    [
                        field_char("begin"),
                        run_element("w:instrText", ' TOC \\o "1-3" \\h \\z \\u '),
                        field_char("separate"),
                        run_element("w:t", "Generated chapter entry"),
                    ]
                )
                second = etree.Element(qn("w:p"))
                set_paragraph_style(second, "TOC2")
                second.extend(
                    [
                        run_element("w:t", "Generated section entry"),
                        field_char("end"),
                    ]
                )
                body.insert(body.index(toc) + 1, second)
                data = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
            elif info.filename == "word/settings.xml":
                root = etree.fromstring(data)
                for update in root.xpath("./w:updateFields", namespaces=NS):
                    root.remove(update)
                data = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
            target.writestr(info, data)
    os.replace(temp, output)


class V024FinalizationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.source = self.root / "synthetic-finalization.docx"
        image = self.root / "pixel.png"
        image.write_bytes(PNG_1X1)

        document = Document()
        document.add_paragraph("Synthetic book title")
        document.add_paragraph("4.1 Static contents entry", style="Heading 1")
        document.add_paragraph("4.1.1 Static contents child", style="Heading 2")
        document.add_paragraph("第4章 Synthetic chapter")
        document.add_paragraph("4.1 Synthetic section")
        polluted = document.add_paragraph("Synthetic body with inherited outline", style="Heading 2")
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "1")
        polluted._p.get_or_add_pPr().append(outline)
        figure = document.add_paragraph()
        figure.add_run().add_picture(str(image))
        document.add_paragraph("图 4.1-1 Synthetic figure")
        document.add_paragraph("表 4.1 Synthetic standalone table")
        first = document.add_table(rows=2, cols=2)
        first.cell(0, 0).text = "Header A"
        first.cell(0, 1).text = "Header B"
        first.cell(1, 0).text = "Value A"
        first.cell(1, 1).text = "Value B"
        repeated = document.add_table(rows=3, cols=2)
        repeated.cell(0, 0).merge(repeated.cell(0, 1)).text = (
            "表 4.2 Synthetic repeated title"
        )
        repeated.cell(1, 0).text = "Header C"
        repeated.cell(1, 1).text = "Header D"
        repeated.cell(2, 0).text = "Value C"
        repeated.cell(2, 1).text = "Value D"
        layout = document.add_table(rows=1, cols=1)
        layout.cell(0, 0).text = "Synthetic image-layout container"
        document.save(self.source)

        structure = candidate_structure_map(self.source)
        structure["status"] = "approved"
        structure["toc_ranges"] = [
            {
                "start_paragraph": 1,
                "end_paragraph": 2,
                "paragraph_sha256": [
                    text_sha256(document.paragraphs[1].text),
                    text_sha256(document.paragraphs[2].text),
                ],
                "levels": 3,
                "approved": True,
            }
        ]
        for entry in structure["headings"]:
            entry["approved"] = entry["paragraph"] in {3, 4}
        for entry in structure["paragraph_roles"]:
            locator = entry["locator"]
            if locator["kind"] == "body_paragraph" and locator["paragraph"] == 5:
                entry.update({"role": "body", "approved": True})
            elif entry["role"] in {
                "heading_1",
                "heading_2",
                "figure_caption",
                "table_caption",
            } and not (
                locator["kind"] == "body_paragraph"
                and locator["paragraph"] in {1, 2}
            ):
                entry["approved"] = True
        for entry in structure["captions"]:
            entry.update({"approved": True, "action": "style_only"})
        structure["tables"][0].update(
            {
                "kind": "data",
                "approved": True,
                "repeat_header_rows": [0],
                "prevent_normal_row_split": True,
            }
        )
        structure["tables"][1].update(
            {
                "kind": "data",
                "approved": True,
                "caption_row": 0,
                "header_rows": [1],
                "repeat_header_rows": [1],
                "repeat_caption_with_header": True,
                "prevent_normal_row_split": True,
            }
        )
        structure["tables"][2].update(
            {
                "kind": "layout",
                "approved": True,
                "pagination_only": True,
                "keep_rows_together": True,
            }
        )
        for group in structure["pagination_groups"]:
            group["approved"] = True
        self.structure = self.root / "structure-map.json"
        self.structure.write_text(
            json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        profile = approved_v11_profile()
        profile["rules"] = [
            rule
            for rule in profile["rules"]
            if rule["selector"]["kind"] == "paragraph_role"
        ]
        profile["rules"].append(
            {
                "id": "FMT-CAP-924",
                "category": "caption",
                "selector": {"kind": "caption_role", "value": "all"},
                "properties": {"font_size_pt": 9, "alignment": "center"},
                "source_ids": ["SRC-901"],
                "evidence_summary": "Synthetic caption formatting.",
                "confidence": "high",
                "status": "approved",
                "application": "automatic",
            }
        )
        self.profile = self.root / "profile.json"
        self.profile.write_text(
            json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        self.output_dir = self.root / "out"

    def tearDown(self) -> None:
        self.temp.cleanup()

    def run_script(self, name: str, *args: object) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(SCRIPTS / name), *(str(value) for value in args)],
            cwd=SKILL,
            capture_output=True,
            text=True,
            check=False,
        )

    def apply(self) -> Path:
        result = self.run_script(
            "apply_profile.py",
            self.source,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure,
            "--output-dir",
            self.output_dir,
            "--allow-missing-fonts",
        )
        self.assertEqual(0, result.returncode, result.stdout + result.stderr)
        return self.output_dir / "synthetic-finalization-formatted.docx"

    def test_v13_applies_toc_outline_and_pagination_patches(self) -> None:
        formatted = self.apply()
        document = Document(formatted)
        self.assertEqual("Synthetic book title", document.paragraphs[0].text)
        self.assertTrue(
            document.paragraphs[1]._p.xpath(
                ".//w:instrText[contains(translate(., 'toc', 'TOC'), 'TOC')]"
            )
        )
        self.assertEqual("第4章 Synthetic chapter", document.paragraphs[2].text)
        body = next(
            p for p in document.paragraphs if p.text == "Synthetic body with inherited outline"
        )
        self.assertEqual("Normal", body.style.name)
        self.assertIsNone(body._p.get_or_add_pPr().find(qn("w:outlineLvl")))
        figure = next(p for p in document.paragraphs if p._p.xpath(".//w:drawing"))
        self.assertIsNotNone(figure._p.get_or_add_pPr().find(qn("w:keepNext")))
        table_caption = next(
            p for p in document.paragraphs if p.text == "表 4.1 Synthetic standalone table"
        )
        self.assertIsNotNone(
            table_caption._p.get_or_add_pPr().find(qn("w:keepNext"))
        )
        for row_index in (0, 1):
            self.assertIsNotNone(
                document.tables[1].rows[row_index]
                ._tr.get_or_add_trPr()
                .find(qn("w:tblHeader"))
            )
        self.assertIsNotNone(
            document.tables[2].rows[0]._tr.get_or_add_trPr().find(qn("w:cantSplit"))
        )

    def test_expanded_toc_keeps_stable_audit_and_single_main_toc(self) -> None:
        formatted = self.apply()
        expanded = self.root / "expanded-toc.docx"
        expand_toc_cache(formatted, expanded)
        field_cache = field_cache_inventory(expanded)
        self.assertEqual("refreshed", field_cache["status"])
        self.assertEqual(1, field_cache["main_toc_fields"])
        self.assertEqual(
            structure_content_fingerprint(self.source, json.loads(self.structure.read_text(encoding="utf-8"))),
            structure_content_fingerprint(expanded, json.loads(self.structure.read_text(encoding="utf-8"))),
        )
        audited = self.run_script(
            "audit_docx.py",
            self.source,
            expanded,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure,
        )
        self.assertEqual(0, audited.returncode, audited.stdout + audited.stderr)

        damaged = self.root / "damaged.docx"
        document = Document(expanded)
        body = next(
            p for p in document.paragraphs if p.text == "Synthetic body with inherited outline"
        )
        body.add_run(" changed")
        document.save(damaged)
        failed = self.run_script(
            "audit_docx.py",
            self.source,
            damaged,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure,
        )
        self.assertEqual(1, failed.returncode)

    def test_deferred_finalization_requires_explicit_qa(self) -> None:
        formatted = self.apply()
        denied = self.run_script(
            "finalize_docx.py",
            formatted,
            "--source",
            self.source,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure,
            "--output",
            self.root / "denied.docx",
            "--field-updater",
            "deferred",
        )
        self.assertEqual(1, denied.returncode)
        self.assertIn("--approve-deferred", denied.stderr)

        finalized = self.root / "finalized.docx"
        status = self.root / "finalization.json"
        allowed = self.run_script(
            "finalize_docx.py",
            formatted,
            "--source",
            self.source,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure,
            "--output",
            finalized,
            "--status-output",
            status,
            "--field-updater",
            "deferred",
            "--approve-deferred",
        )
        self.assertEqual(0, allowed.returncode, allowed.stdout + allowed.stderr)
        payload = json.loads(status.read_text(encoding="utf-8"))
        self.assertEqual(1, payload["finalization_evidence_version"])
        self.assertEqual("deferred", payload["delivery_field_status"])
        self.assertEqual("pass", payload["content_integrity"])
        self.assertEqual("pass", payload["protected_object_integrity"])
        self.assertEqual("persisted", payload["backend_audit"]["status"])
        publication = payload["publication"]
        self.assertEqual("manual_only", publication["cleanup_policy"])
        self.assertFalse(publication["business_gate"])
        retained_staging = Path(publication["retained_staging_directory"])
        self.assertTrue(retained_staging.is_dir())
        publication_record = json.loads(
            (retained_staging / finalize_docx.PUBLICATION_RECORD_NAME).read_text(
                encoding="utf-8"
            )
        )
        self.assertEqual(
            publication["transaction_id"], publication_record["transaction_id"]
        )
        self.assertEqual("committed", publication_record["status"])
        audit_identity = payload["backend_audit"]["artifact"]
        audit_path = Path(audit_identity["path"])
        self.assertTrue(audit_path.is_file())
        self.assertEqual(audit_path.stat().st_size, audit_identity["size_bytes"])
        self.assertNotIn("attempted_backend", payload["field_backend"])

        without_status = self.run_script(
            "finalize_docx.py",
            formatted,
            "--source",
            self.source,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure,
            "--output",
            self.root / "without-status.docx",
            "--field-updater",
            "deferred",
            "--approve-deferred",
        )
        self.assertEqual(
            0, without_status.returncode, without_status.stdout + without_status.stderr
        )
        stdout_payload = json.loads(without_status.stdout)
        self.assertEqual("not_persisted", stdout_payload["backend_audit"]["status"])
        self.assertIsNone(stdout_payload["backend_audit"]["artifact"])
        self.assertNotIn("attempted_backend", stdout_payload["field_backend"])

    def test_finalization_path_contract_rejects_full_pairwise_alias_matrix(self) -> None:
        def values() -> dict:
            return {
                "input_path": self.root / "input.docx",
                "source_path": self.root / "source.docx",
                "profile_path": self.root / "profile.json",
                "structure_map_path": self.root / "structure.json",
                "output_path": self.root / "final.docx",
                "pdf_output": self.root / "target.pdf",
                "status_output": self.root / "finalization.json",
                "force": True,
            }

        output_arguments = {
            "output": "output_path",
            "pdf": "pdf_output",
            "status": "status_output",
        }
        outputs = ("output", "pdf", "status", "audit")
        for left_index, left in enumerate(outputs):
            for right in outputs[left_index + 1 :]:
                if (left, right) == ("status", "audit"):
                    contract = resolve_finalization_path_contract(**values())
                    self.assertNotEqual(
                        contract["outputs"]["status"],
                        contract["outputs"]["audit"],
                    )
                    continue
                arguments = values()
                collision = self.root / "output-collision-backend-audit.json"
                if left == "audit" or right == "audit":
                    arguments["status_output"] = self.root / "output-collision.json"
                    direct = right if left == "audit" else left
                    arguments[output_arguments[direct]] = collision
                else:
                    arguments[output_arguments[left]] = collision
                    arguments[output_arguments[right]] = collision
                with self.subTest(left=left, right=right), self.assertRaises(
                    FormatMonographError
                ):
                    resolve_finalization_path_contract(**arguments)

        input_arguments = {
            "input": "input_path",
            "source": "source_path",
            "profile": "profile_path",
            "structure_map": "structure_map_path",
        }
        for output_name in outputs:
            for input_name, input_argument in input_arguments.items():
                arguments = values()
                if output_name == "audit":
                    arguments["status_output"] = self.root / "input-collision.json"
                    arguments[input_argument] = (
                        self.root / "input-collision-backend-audit.json"
                    )
                else:
                    arguments[output_arguments[output_name]] = arguments[
                        input_argument
                    ]
                with self.subTest(
                    output=output_name, input=input_name
                ), self.assertRaises(FormatMonographError):
                    resolve_finalization_path_contract(**arguments)

    def test_finalization_path_contract_rejects_symlink_control_and_parent_matrix(self) -> None:
        real = self.root / "real-target.docx"
        real.write_bytes(b"sentinel")
        symlink = self.root / "linked-output.docx"
        symlink.symlink_to(real)
        input_symlink = self.root / "linked-input.docx"
        input_symlink.symlink_to(real)
        common = {
            "input_path": self.source,
            "source_path": self.source,
            "profile_path": self.profile,
            "structure_map_path": self.structure,
            "output_path": self.root / "final.docx",
            "pdf_output": self.root / "target.pdf",
            "status_output": self.root / "finalization.json",
            "force": True,
        }
        cases = (
            ("output_symlink", {"output_path": symlink}),
            (
                "resolved_input_alias",
                {"input_path": input_symlink, "output_path": real},
            ),
            ("nul", {"output_path": Path("bad\0.docx")}),
            ("control", {"status_output": Path("bad\nstatus.json")}),
            (
                "different_parent",
                {"pdf_output": self.root / "nested" / "target.pdf"},
            ),
        )
        for name, update in cases:
            arguments = dict(common)
            arguments.update(update)
            with self.subTest(name=name), self.assertRaises(FormatMonographError):
                resolve_finalization_path_contract(**arguments)

    def test_existing_publish_sentinels_survive_early_middle_and_write_failures(self) -> None:
        formatted = self.apply()
        real_move = finalize_docx._platform_atomic_noreplace_move
        real_atomic_write_bytes = finalize_docx.atomic_write_bytes
        for mode in (
            "early",
            "middle",
            "sidecar_write",
            "status_write",
            "publish",
        ):
            case_root = self.root / f"sentinel-{mode}"
            case_root.mkdir()
            output = case_root / "final.docx"
            pdf = case_root / "target.pdf"
            status = case_root / "finalization.json"
            sidecar = case_root / "finalization-backend-audit.json"
            sentinels = {
                output: b"old-output",
                pdf: b"old-pdf",
                status: b"old-status",
                sidecar: b"old-sidecar",
            }
            for path, payload in sentinels.items():
                path.write_bytes(payload)

            profile = self.profile
            updater = "deferred"
            approve_deferred = True
            if mode == "early":
                invalid_profile = case_root / "invalid-profile.json"
                value = json.loads(self.profile.read_text(encoding="utf-8"))
                value["approval"]["status"] = "draft"
                invalid_profile.write_text(json.dumps(value), encoding="utf-8")
                profile = invalid_profile
            elif mode == "middle":
                updater = "external"
                approve_deferred = False

            argv = [
                "finalize_docx.py",
                str(formatted),
                "--source",
                str(self.source),
                "--profile",
                str(profile),
                "--structure-map",
                str(self.structure),
                "--output",
                str(output),
                "--pdf-output",
                str(pdf),
                "--status-output",
                str(status),
                "--field-updater",
                updater,
                "--force",
            ]
            if approve_deferred:
                argv.append("--approve-deferred")

            failed_publish = False

            def move(source, target):
                nonlocal failed_publish
                if (
                    mode == "publish"
                    and target.path.resolve(strict=False) == status.resolve(strict=False)
                    and not failed_publish
                ):
                    failed_publish = True
                    raise OSError("synthetic status publish failure")
                return real_move(source, target)

            patches = [patch.object(sys, "argv", argv)]
            if mode == "sidecar_write":
                patches.append(
                    patch(
                        "finalize_docx.atomic_write_bytes",
                        side_effect=OSError("synthetic sidecar write failure"),
                    )
                )
            elif mode == "status_write":
                def fail_status_write(path, payload):
                    if Path(path).name == "finalization.json":
                        raise OSError("synthetic status write failure")
                    return real_atomic_write_bytes(path, payload)

                patches.append(
                    patch(
                        "finalize_docx.atomic_write_bytes",
                        side_effect=fail_status_write,
                    )
                )
            elif mode == "publish":
                patches.append(
                    patch(
                        "finalize_docx._platform_atomic_noreplace_move",
                        side_effect=move,
                    )
                )

            with contextlib.ExitStack() as stack:
                for active_patch in patches:
                    stack.enter_context(active_patch)
                with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(
                    io.StringIO()
                ):
                    self.assertEqual(1, finalize_docx.main(), mode)
            for path, payload in sentinels.items():
                self.assertEqual(payload, path.read_bytes(), f"{mode}: {path.name}")
            retained_staging = list(
                case_root.glob(".format-monograph-finalize-*")
            )
            if mode == "early":
                self.assertEqual([], retained_staging, mode)
            else:
                self.assertEqual(1, len(retained_staging), mode)

    def deferred_finalizer_argv(self, formatted: Path, case_root: Path) -> list[str]:
        return [
            "finalize_docx.py",
            str(formatted),
            "--source",
            str(self.source),
            "--profile",
            str(self.profile),
            "--structure-map",
            str(self.structure),
            "--output",
            str(case_root / "final.docx"),
            "--pdf-output",
            str(case_root / "target.pdf"),
            "--status-output",
            str(case_root / "finalization.json"),
            "--field-updater",
            "deferred",
            "--approve-deferred",
            "--force",
        ]

    def test_post_commit_stdout_failures_do_not_reverse_success(self) -> None:
        formatted = self.apply()

        class FailingStream(io.StringIO):
            def __init__(self, mode: str) -> None:
                super().__init__()
                self.mode = mode

            def write(self, value: str) -> int:
                if self.mode == "broken_pipe":
                    raise BrokenPipeError("synthetic closed pipe")
                if self.mode == "oserror":
                    raise OSError("synthetic stdout write failure")
                if self.mode == "encoding":
                    raise UnicodeEncodeError("ascii", value, 0, 1, "synthetic")
                return super().write(value)

            def flush(self) -> None:
                if self.mode == "flush":
                    raise OSError("synthetic stdout flush failure")
                return super().flush()

        for mode in ("broken_pipe", "oserror", "encoding", "flush", "closed"):
            with self.subTest(mode=mode):
                case_root = self.root / f"stdout-{mode}"
                case_root.mkdir()
                stream: io.StringIO = FailingStream(mode)
                if mode == "closed":
                    stream.close()
                with patch.object(
                    sys, "argv", self.deferred_finalizer_argv(formatted, case_root)
                ), patch.object(sys, "stdout", stream), contextlib.redirect_stderr(
                    io.StringIO()
                ):
                    self.assertEqual(0, finalize_docx.main())
                status = json.loads(
                    (case_root / "finalization.json").read_text(encoding="utf-8")
                )
                self.assertEqual("pass", status["status"])
                self.assertTrue((case_root / "final.docx").is_file())
                self.assertTrue(
                    (case_root / "finalization-backend-audit.json").is_file()
                )

        print_root = self.root / "stdout-print"
        print_root.mkdir()
        with patch.object(
            sys, "argv", self.deferred_finalizer_argv(formatted, print_root)
        ), patch.object(sys, "stdout", io.StringIO()), patch(
            "builtins.print", side_effect=BrokenPipeError("synthetic print")
        ):
            self.assertEqual(0, finalize_docx.main())

    def test_main_post_commit_close_warning_returns_success_and_matching_status(self) -> None:
        formatted = self.apply()
        case = self.root / "main-close-warning"
        case.mkdir()
        stdout = io.StringIO()
        stderr = io.StringIO()
        real_close = finalize_docx._close_authority
        injected = False

        def close_with_fault(authority) -> None:
            nonlocal injected
            real_close(authority)
            if not injected and authority.path.name.startswith(
                ".format-monograph-finalize-"
            ):
                injected = True
                raise OSError("synthetic staging close failure")

        with patch.object(
            sys,
            "argv",
            self.deferred_finalizer_argv(formatted, case),
        ), patch.object(
            finalize_docx, "_close_authority", close_with_fault
        ), contextlib.redirect_stdout(stdout), contextlib.redirect_stderr(stderr):
            self.assertEqual(0, finalize_docx.main())
        self.assertTrue(injected)
        disk_status = json.loads(
            (case / "finalization.json").read_text(encoding="utf-8")
        )
        self.assertEqual(disk_status, json.loads(stdout.getvalue()))
        warning = json.loads(stderr.getvalue())
        self.assertEqual(
            "committed_with_cleanup_errors",
            warning["publication_commit_state"],
        )
        self.assertEqual(1, len(warning["cleanup_errors"]))
        self.assertTrue((case / "final.docx").is_file())

    def test_commit_close_warning_console_failure_matrix_never_reverses_success(self) -> None:
        formatted = self.apply()

        class FaultyStream(io.StringIO):
            def __init__(self, mode: str) -> None:
                super().__init__()
                self.mode = mode

            def write(self, value: str) -> int:
                if self.mode == "broken_pipe":
                    raise BrokenPipeError("synthetic broken console")
                if self.mode == "oserror":
                    raise OSError("synthetic console OSError")
                return super().write(value)

        for stderr_mode in ("ok", "broken_pipe", "closed"):
            for stdout_mode in ("ok", "oserror", "closed"):
                with self.subTest(stderr=stderr_mode, stdout=stdout_mode):
                    case = self.root / f"console-{stderr_mode}-{stdout_mode}"
                    case.mkdir()
                    stdout = FaultyStream(stdout_mode)
                    stderr = FaultyStream(stderr_mode)
                    if stdout_mode == "closed":
                        stdout.close()
                    if stderr_mode == "closed":
                        stderr.close()
                    real_close = finalize_docx._close_authority
                    injected = False
                    attempted: list[str] = []

                    def close_with_fault(authority) -> None:
                        nonlocal injected
                        if authority.path.resolve() == case.resolve():
                            attempted.append("output_parent")
                        elif authority.path.name.startswith(
                            ".format-monograph-finalize-"
                        ):
                            attempted.append("staging")
                        elif authority.path.name.startswith("producer-"):
                            attempted.append("producer")
                        real_close(authority)
                        if not injected and authority.path.name.startswith(
                            ".format-monograph-finalize-"
                        ):
                            injected = True
                            raise OSError("synthetic post-commit staging close failure")

                    with patch.object(
                        sys,
                        "argv",
                        self.deferred_finalizer_argv(formatted, case),
                    ), patch.object(
                        finalize_docx, "_close_authority", close_with_fault
                    ), patch.object(sys, "stdout", stdout), patch.object(
                        sys, "stderr", stderr
                    ):
                        self.assertEqual(0, finalize_docx.main())
                    self.assertTrue(injected)
                    self.assertEqual(
                        {"producer", "staging", "output_parent"}, set(attempted)
                    )
                    disk_status = json.loads(
                        (case / "finalization.json").read_text(encoding="utf-8")
                    )
                    retained = Path(
                        disk_status["publication"]["retained_staging_directory"]
                    )
                    self.assertTrue(
                        (retained / finalize_docx.PUBLICATION_RECORD_NAME).is_file()
                    )
                    self.assertTrue((case / "final.docx").is_file())
                    if stdout_mode == "ok":
                        self.assertEqual(disk_status, json.loads(stdout.getvalue()))
                    if stderr_mode == "ok":
                        warning = json.loads(stderr.getvalue())
                        self.assertEqual(
                            "committed_with_cleanup_errors",
                            warning["publication_commit_state"],
                        )

    @unittest.skipUnless(os.name == "posix", "closed-pipe subprocess is POSIX-specific")
    def test_real_cli_closed_stdout_exits_zero_after_commit(self) -> None:
        formatted = self.apply()
        case_root = self.root / "closed-pipe-subprocess"
        case_root.mkdir()
        argv = self.deferred_finalizer_argv(formatted, case_root)
        argv[0] = str(SCRIPTS / "finalize_docx.py")
        read_fd, write_fd = os.pipe()
        process = subprocess.Popen(
            [sys.executable, *argv],
            cwd=SKILL,
            stdout=write_fd,
            stderr=subprocess.PIPE,
            close_fds=True,
        )
        os.close(write_fd)
        os.close(read_fd)
        _, stderr = process.communicate(timeout=60)
        self.assertEqual(0, process.returncode, stderr.decode(errors="replace"))
        status = json.loads(
            (case_root / "finalization.json").read_text(encoding="utf-8")
        )
        self.assertEqual("pass", status["status"])

    def test_precommit_failure_keeps_sentinels_even_with_broken_stdout(self) -> None:
        formatted = self.apply()
        case_root = self.root / "precommit-broken-output"
        case_root.mkdir()
        sentinels = {
            case_root / "final.docx": b"old-output",
            case_root / "target.pdf": b"old-pdf",
            case_root / "finalization.json": b"old-status",
            case_root / "finalization-backend-audit.json": b"old-audit",
        }
        for path, content in sentinels.items():
            path.write_bytes(content)

        class BrokenStdout(io.StringIO):
            def write(self, value: str) -> int:
                raise BrokenPipeError("synthetic precommit stdout")

        real_atomic_write_bytes = finalize_docx.atomic_write_bytes

        def fail_status_write(path, payload):
            if Path(path).name == "finalization.json":
                raise OSError("synthetic precommit status write failure")
            return real_atomic_write_bytes(path, payload)

        with patch.object(
            sys, "argv", self.deferred_finalizer_argv(formatted, case_root)
        ), patch.object(sys, "stdout", BrokenStdout()), patch(
            "finalize_docx.atomic_write_bytes",
            side_effect=fail_status_write,
        ), contextlib.redirect_stderr(io.StringIO()):
            self.assertEqual(1, finalize_docx.main())
        for path, content in sentinels.items():
            self.assertEqual(content, path.read_bytes())

    def test_producer_to_publisher_status_tamper_matrix_never_replaces_old_transaction(
        self,
    ) -> None:
        formatted = self.apply()
        mutations = (
            ("status", lambda value: value.update(status="fail")),
            ("unknown", lambda value: value.update(unknown_gate=True)),
            (
                "version",
                lambda value: value.update(finalization_evidence_version=True),
            ),
            (
                "completion",
                lambda value: value["field_completion"].update(
                    final_ready_eligible=1.0
                ),
            ),
            (
                "validation",
                lambda value: value["field_completion"].update(
                    evidence_validation={"status": "pass", "errors": []}
                ),
            ),
        )
        real_atomic_write_bytes = finalize_docx.atomic_write_bytes
        for name, mutate in mutations:
            with self.subTest(name=name):
                case_root = self.root / f"producer-publisher-{name}"
                case_root.mkdir()
                sentinels = {
                    case_root / "final.docx": b"old-output",
                    case_root / "target.pdf": b"old-pdf",
                    case_root / "finalization.json": b"old-status",
                    case_root / "finalization-backend-audit.json": b"old-audit",
                }
                for path, content in sentinels.items():
                    path.write_bytes(content)
                tampered = False

                def atomic_write(path, payload):
                    nonlocal tampered
                    real_atomic_write_bytes(path, payload)
                    if Path(path).name == "finalization.json":
                        value = json.loads(Path(path).read_text(encoding="utf-8"))
                        mutate(value)
                        Path(path).write_bytes(
                            finalize_docx.standard_json_bytes(value)
                        )
                        tampered = True

                with patch.object(
                    sys,
                    "argv",
                    self.deferred_finalizer_argv(formatted, case_root),
                ), patch(
                    "finalize_docx.atomic_write_bytes", side_effect=atomic_write
                ), contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(
                    io.StringIO()
                ):
                    self.assertEqual(1, finalize_docx.main())
                self.assertTrue(tampered)
                for path, content in sentinels.items():
                    self.assertEqual(content, path.read_bytes())
                retained = list(
                    case_root.glob(".format-monograph-finalize-*")
                )
                self.assertEqual(1, len(retained))
                self.assertTrue((retained[0] / "finalization.json").exists())
                self.assertFalse(
                    (case_root / finalize_docx.PUBLICATION_RECOVERY_DIRECTORY).exists()
                )

    def test_field_contract_rejects_removed_editable_fields(self) -> None:
        before = {
            "main_toc_fields": 1,
            "field_types": {"TOC": 1, "PAGE": 2},
        }
        self.assertTrue(field_contract_preserved(before, before))
        self.assertFalse(
            field_contract_preserved(
                before,
                {"main_toc_fields": 0, "field_types": {"PAGE": 2}},
            )
        )
        self.assertFalse(
            field_contract_preserved(
                before,
                {"main_toc_fields": 1, "field_types": {"TOC": 1, "PAGE": 1}},
            )
        )

    def test_main_retained_staging_replacement_matrix_never_deletes_unknown(self) -> None:
        formatted = self.apply()
        phases = (
            "after_staging_create",
            "after_producer_workspace_create",
            "before_producer_execution",
            "before_staged_artifact_import",
            "after_staged_artifact_import",
            "before_publisher",
            "before_staging_authority_open",
            "after_staging_authority_open",
            "before_staging_retention",
            "after_staging_retention",
        )
        for phase in phases:
            for kind in ("regular", "symlink", "empty_directory", "nonempty_directory"):
                with self.subTest(phase=phase, kind=kind):
                    case = self.root / f"main-staging-{phase}-{kind}"
                    case.mkdir()
                    marker = f"unknown-{phase}-{kind}".encode()
                    staging_path: Path | None = None
                    displaced: Path | None = None
                    replaced = False

                    def hook(event: str, name: str | None) -> None:
                        nonlocal staging_path, displaced, replaced
                        if replaced or event != phase:
                            return
                        candidates = list(
                            case.glob(".format-monograph-finalize-*")
                        )
                        self.assertEqual(1, len(candidates))
                        staging_path = candidates[0]
                        displaced = case / f"displaced-{phase}-{kind}"
                        os.rename(staging_path, displaced)
                        if kind == "regular":
                            staging_path.write_bytes(marker)
                        elif kind == "symlink":
                            staging_path.symlink_to(marker.decode())
                        else:
                            staging_path.mkdir(mode=0o700)
                            if kind == "nonempty_directory":
                                (staging_path / "unknown.marker").write_bytes(marker)
                        replaced = True

                    with patch.object(
                        sys,
                        "argv",
                        self.deferred_finalizer_argv(formatted, case),
                    ), contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(
                        io.StringIO()
                    ):
                        result = finalize_docx.main(
                            publication_event_hook=hook
                        )
                    self.assertTrue(replaced)
                    self.assertIsNotNone(staging_path)
                    self.assertIsNotNone(displaced)
                    self.assertTrue(
                        staging_path.is_symlink() or staging_path.exists()
                    )
                    self.assertTrue(displaced.exists())
                    if kind == "regular":
                        self.assertEqual(marker, staging_path.read_bytes())
                    elif kind == "symlink":
                        self.assertEqual(marker.decode(), os.readlink(staging_path))
                    elif kind == "nonempty_directory":
                        self.assertEqual(
                            marker,
                            (staging_path / "unknown.marker").read_bytes(),
                        )
                    self.assertEqual(1, result)

    def test_create_bind_window_replacement_matrix_fails_closed(self) -> None:
        formatted = self.apply()
        for authority_kind, prefix in (
            ("staging", ".format-monograph-finalize-"),
            ("producer", "producer-"),
        ):
            for replacement_kind in (
                "regular",
                "symlink",
                "empty_directory",
                "nonempty_directory",
            ):
                with self.subTest(
                    authority=authority_kind, replacement=replacement_kind
                ):
                    case = self.root / (
                        f"create-bind-{authority_kind}-{replacement_kind}"
                    )
                    case.mkdir()
                    victim = case / "external-victim.bin"
                    victim_bytes = f"victim-{authority_kind}-{replacement_kind}".encode()
                    victim.write_bytes(victim_bytes)
                    real_open = finalize_docx._open_directory_authority
                    attacked_path: Path | None = None
                    displaced: Path | None = None

                    def open_with_replacement(path: Path, **kwargs):
                        nonlocal attacked_path, displaced
                        name = kwargs.get("name")
                        if (
                            attacked_path is None
                            and isinstance(name, str)
                            and name.startswith(prefix)
                        ):
                            attacked_path = Path(path)
                            displaced = attacked_path.with_name(
                                f"displaced-{authority_kind}-{replacement_kind}"
                            )
                            os.rename(attacked_path, displaced)
                            if replacement_kind == "regular":
                                attacked_path.write_bytes(b"unknown-regular")
                            elif replacement_kind == "symlink":
                                attacked_path.symlink_to(victim)
                            else:
                                attacked_path.mkdir(mode=0o700)
                                if replacement_kind == "nonempty_directory":
                                    (attacked_path / "unknown.marker").write_bytes(
                                        b"unknown-directory"
                                    )
                        return real_open(path, **kwargs)

                    with patch.object(
                        sys,
                        "argv",
                        self.deferred_finalizer_argv(formatted, case),
                    ), patch.object(
                        finalize_docx,
                        "_open_directory_authority",
                        open_with_replacement,
                    ), contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(
                        io.StringIO()
                    ):
                        result = finalize_docx.main()
                    self.assertEqual(1, result)
                    self.assertIsNotNone(attacked_path)
                    self.assertIsNotNone(displaced)
                    self.assertTrue(displaced.is_dir())
                    self.assertEqual(victim_bytes, victim.read_bytes())
                    if replacement_kind == "regular":
                        self.assertEqual(b"unknown-regular", attacked_path.read_bytes())
                    elif replacement_kind == "symlink":
                        self.assertTrue(attacked_path.is_symlink())
                    elif replacement_kind == "nonempty_directory":
                        self.assertEqual(
                            b"unknown-directory",
                            (attacked_path / "unknown.marker").read_bytes(),
                        )
                    self.assertFalse((case / "final.docx").exists())

    def test_windows_unavailable_gate_precedes_output_parent_creation(self) -> None:
        missing_parent = self.root / "windows-unavailable" / "nested"
        with patch.object(finalize_docx.os, "name", "nt"):
            with self.assertRaisesRegex(
                finalize_docx.FinalizationPublishError,
                "before output-parent creation",
            ):
                finalize_docx._assert_publication_platform_available()
        self.assertFalse(missing_parent.exists())

    def test_main_staged_artifact_collision_matrix_never_touches_unknown(self) -> None:
        formatted = self.apply()
        artifact_names = {
            "output": "final.docx",
            "pdf": "target.pdf",
            "status": "finalization.json",
            "audit": "finalization-backend-audit.json",
        }
        for artifact, basename in artifact_names.items():
            for kind in ("regular", "symlink", "directory"):
                with self.subTest(artifact=artifact, kind=kind):
                    case = self.root / f"staged-collision-{artifact}-{kind}"
                    case.mkdir()
                    victim = case / f"victim-{artifact}.bin"
                    victim_bytes = f"victim-{artifact}-{kind}".encode()
                    victim.write_bytes(victim_bytes)
                    inserted: Path | None = None

                    def hook(event: str, name: str | None) -> None:
                        nonlocal inserted
                        if event != "after_staging_create" or inserted is not None:
                            return
                        staging = next(case.glob(".format-monograph-finalize-*"))
                        inserted = staging / basename
                        if kind == "regular":
                            inserted.write_bytes(b"unknown-regular")
                        elif kind == "symlink":
                            inserted.symlink_to(victim)
                        else:
                            inserted.mkdir()
                            (inserted / "unknown.marker").write_bytes(b"unknown-dir")

                    with patch.object(
                        sys,
                        "argv",
                        self.deferred_finalizer_argv(formatted, case),
                    ), contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(
                        io.StringIO()
                    ):
                        result = finalize_docx.main(publication_event_hook=hook)
                    self.assertEqual(1, result)
                    self.assertIsNotNone(inserted)
                    self.assertEqual(victim_bytes, victim.read_bytes())
                    if kind == "regular":
                        self.assertEqual(b"unknown-regular", inserted.read_bytes())
                    elif kind == "symlink":
                        self.assertTrue(inserted.is_symlink())
                        self.assertEqual(victim.resolve(), inserted.resolve())
                    else:
                        self.assertEqual(
                            b"unknown-dir", (inserted / "unknown.marker").read_bytes()
                        )
                    self.assertFalse((case / basename).exists())

    def test_no_status_imported_identity_mutation_matrix_never_publishes(self) -> None:
        formatted = self.apply()
        for mutation in ("bytes", "same_bytes_new_inode", "symlink"):
            with self.subTest(mutation=mutation):
                case = self.root / f"no-status-import-{mutation}"
                case.mkdir()
                argv = self.deferred_finalizer_argv(formatted, case)
                status_index = argv.index("--status-output")
                del argv[status_index : status_index + 2]
                victim = case / "external-victim.docx"
                victim.write_bytes(b"external-victim-unchanged")
                original_retained: Path | None = None

                def hook(event: str, name: str | None) -> None:
                    nonlocal original_retained
                    if event != "after_staged_artifact_import" or name != "output":
                        return
                    staging = next(case.glob(".format-monograph-finalize-*"))
                    imported = staging / "final.docx"
                    if mutation == "bytes":
                        imported.write_bytes(imported.read_bytes() + b"mutated")
                    else:
                        original = imported.read_bytes()
                        original_retained = staging / f"original-{mutation}.docx"
                        os.rename(imported, original_retained)
                        if mutation == "same_bytes_new_inode":
                            imported.write_bytes(original)
                        else:
                            imported.symlink_to(victim)

                stdout = io.StringIO()
                with patch.object(sys, "argv", argv), contextlib.redirect_stdout(
                    stdout
                ), contextlib.redirect_stderr(io.StringIO()):
                    result = finalize_docx.main(publication_event_hook=hook)
                self.assertEqual(1, result)
                self.assertEqual("", stdout.getvalue())
                self.assertFalse((case / "final.docx").exists())
                self.assertEqual(b"external-victim-unchanged", victim.read_bytes())
                if original_retained is not None:
                    self.assertTrue(original_retained.is_file())

    def test_import_close_to_first_reopen_replacement_matrix_never_publishes(self) -> None:
        formatted = self.apply()
        for mutation in (
            "same_bytes_new_inode",
            "different_bytes",
            "symlink",
            "extra_entry",
        ):
            with self.subTest(mutation=mutation):
                case = self.root / f"import-first-reopen-{mutation}"
                case.mkdir()
                argv = self.deferred_finalizer_argv(formatted, case)
                status_index = argv.index("--status-output")
                del argv[status_index : status_index + 2]
                victim = case / "external-victim.docx"
                victim.write_bytes(b"external-victim-unchanged")
                real_snapshot = finalize_docx._artifact_full_snapshot_at
                mutated = False
                retained_original: Path | None = None
                extra: Path | None = None

                def snapshot_with_replacement(location):
                    nonlocal mutated, retained_original, extra
                    if (
                        not mutated
                        and location.name == "final.docx"
                        and location.authority.path.name.startswith(
                            ".format-monograph-finalize-"
                        )
                    ):
                        mutated = True
                        imported = location.path
                        if mutation == "extra_entry":
                            extra = location.authority.path / "unknown-after-import"
                            extra.mkdir()
                            (extra / "unknown.marker").write_bytes(b"unknown-extra")
                        else:
                            original = imported.read_bytes()
                            retained_original = location.authority.path / (
                                f"original-{mutation}.docx"
                            )
                            os.rename(imported, retained_original)
                            if mutation == "same_bytes_new_inode":
                                imported.write_bytes(original)
                            elif mutation == "different_bytes":
                                imported.write_bytes(b"different-bytes")
                            else:
                                imported.symlink_to(victim)
                    return real_snapshot(location)

                stdout = io.StringIO()
                with patch.object(sys, "argv", argv), patch.object(
                    finalize_docx,
                    "_artifact_full_snapshot_at",
                    snapshot_with_replacement,
                ), contextlib.redirect_stdout(stdout), contextlib.redirect_stderr(
                    io.StringIO()
                ):
                    result = finalize_docx.main()
                self.assertTrue(mutated)
                self.assertEqual(1, result)
                self.assertEqual("", stdout.getvalue())
                self.assertFalse((case / "final.docx").exists())
                self.assertEqual(b"external-victim-unchanged", victim.read_bytes())
                if retained_original is not None:
                    self.assertTrue(retained_original.is_file())
                if extra is not None:
                    self.assertEqual(
                        b"unknown-extra", (extra / "unknown.marker").read_bytes()
                    )

    def test_main_simulated_windows_control_flow_is_not_availability_evidence(self) -> None:
        formatted = self.apply()
        for existing in (False, True):
            with self.subTest(existing=existing):
                case = self.root / f"main-windows-e2e-{int(existing)}"
                case.mkdir()
                targets = (
                    case / "final.docx",
                    case / "target.pdf",
                    case / "finalization.json",
                    case / "finalization-backend-audit.json",
                )
                if existing:
                    for index, target in enumerate(targets):
                        target.write_bytes(f"old-{index}".encode())
                events: list[tuple[str, str | None]] = []
                api = SimulatedWindowsPublisherApi()
                stdout = io.StringIO()
                with patch.object(
                    finalize_docx,
                    "_PUBLISHER_AUTHORITY_BACKEND_OVERRIDE",
                    "windows",
                ), patch.object(
                    finalize_docx, "_WINDOWS_PUBLISHER_API", api
                ), patch.object(
                    sys,
                    "argv",
                    self.deferred_finalizer_argv(formatted, case),
                ), contextlib.redirect_stdout(stdout), contextlib.redirect_stderr(
                    io.StringIO()
                ):
                    self.assertEqual(
                        0,
                        finalize_docx.main(
                            publication_event_hook=lambda event, name: events.append(
                                (event, name)
                            )
                        ),
                    )
                self.assertEqual(set(), api.open_handles)
                published = [
                    name for event, name in events if event == "after_publish"
                ]
                self.assertEqual(["output", "pdf", "audit", "status"], published)
                status = json.loads((case / "finalization.json").read_text(encoding="utf-8"))
                publication = status["publication"]
                retained = Path(publication["retained_staging_directory"])
                self.assertTrue(retained.is_dir())
                self.assertTrue(
                    (retained / finalize_docx.PUBLICATION_RECORD_NAME).is_file()
                )
                self.assertEqual("manual_only", publication["cleanup_policy"])
                recovery = case / finalize_docx.PUBLICATION_RECOVERY_DIRECTORY
                self.assertEqual(existing, recovery.is_dir())

    @unittest.skipUnless(os.name == "nt", "requires a real Windows host")
    def test_windows_host_main_publisher_fails_closed_before_production(self) -> None:
        formatted = self.apply()
        case = self.root / "windows-host-e2e"
        case.mkdir()
        targets = (
            case / "final.docx",
            case / "target.pdf",
            case / "finalization.json",
            case / "finalization-backend-audit.json",
        )
        for index, target in enumerate(targets):
            target.write_bytes(f"old-{index}".encode())
        before = {target: target.read_bytes() for target in targets}
        stdout = io.StringIO()
        stderr = io.StringIO()
        events: list[tuple[str, str | None]] = []
        move = MagicMock(
            side_effect=AssertionError("Windows publisher must not attempt a move")
        )
        with patch.object(
            sys,
            "argv",
            self.deferred_finalizer_argv(formatted, case),
        ), patch.object(
            finalize_docx, "_platform_atomic_noreplace_move", move
        ), contextlib.redirect_stdout(stdout), contextlib.redirect_stderr(stderr):
            self.assertEqual(
                1,
                finalize_docx.main(
                    publication_event_hook=lambda event, name: events.append(
                        (event, name)
                    )
                ),
            )
        self.assertIn("Windows publication is unavailable", stderr.getvalue())
        self.assertEqual([], events)
        move.assert_not_called()
        self.assertEqual(before, {target: target.read_bytes() for target in targets})
        self.assertEqual([], list(case.glob(".format-monograph-finalize-*")))

    @unittest.skipUnless(os.name == "nt", "requires a real Windows host")
    def test_windows_host_unavailable_does_not_create_output_parent(self) -> None:
        formatted = self.apply()
        missing_parent = self.root / "windows-missing-parent" / "nested"
        stderr = io.StringIO()
        with patch.object(
            sys,
            "argv",
            self.deferred_finalizer_argv(formatted, missing_parent),
        ), contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(stderr):
            self.assertEqual(1, finalize_docx.main())
        self.assertIn("before output-parent creation", stderr.getvalue())
        self.assertFalse(missing_parent.exists())


class V024AtomicPublishTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_create_bound_directory_retries_collision_without_adopting_it(self) -> None:
        parent = finalize_docx._open_directory_authority(
            self.root, require_private_owner=True
        )
        occupied_name = "bound-" + "a" * 32
        occupied = self.root / occupied_name
        occupied.mkdir()
        (occupied / "unknown.marker").write_bytes(b"unknown")

        class Token:
            def __init__(self, value: str) -> None:
                self.hex = value

        created = None
        try:
            with patch.object(
                finalize_docx.uuid,
                "uuid4",
                side_effect=(Token("a" * 32), Token("b" * 32)),
            ):
                created = finalize_docx._create_bound_directory(
                    parent, prefix="bound-", label="collision test"
                )
            self.assertEqual("bound-" + "b" * 32, created.path.name)
            self.assertEqual(b"unknown", (occupied / "unknown.marker").read_bytes())
            self.assertNotEqual(occupied.stat().st_ino, created.inode)
        finally:
            finalize_docx._close_authorities_collect(
                [("created", created), ("parent", parent)]
            )

    @contextlib.contextmanager
    def simulated_windows_backend(self):
        api = SimulatedWindowsPublisherApi()
        with patch.object(
            finalize_docx,
            "_PUBLISHER_AUTHORITY_BACKEND_OVERRIDE",
            "windows",
        ), patch.object(finalize_docx, "_WINDOWS_PUBLISHER_API", api):
            yield api
        self.assertEqual(set(), api.open_handles)

    def publish_fixture(
        self, *, existing: bool
    ) -> tuple[
        Path,
        dict[str, Path],
        dict[str, Path],
        dict[str, dict | None],
        dict[str, bytes],
    ]:
        staging = self.root / "staging"
        staging.mkdir()
        targets = {
            "output": self.root / "final.docx",
            "pdf": self.root / "target.pdf",
            "audit": self.root / "finalization-backend-audit.json",
            "status": self.root / "finalization.json",
        }
        staged = {name: staging / path.name for name, path in targets.items()}
        old = {name: f"old-{name}".encode() for name in targets}
        if existing:
            for name, path in targets.items():
                path.write_bytes(old[name])
        staged["output"].write_bytes(b"new-output")
        staged["pdf"].write_bytes(b"new-pdf")
        staged["audit"].write_bytes(b"new-audit")
        input_cache = {
            "status": "stale",
            "main_toc_fields": 1,
            "toc_entries": 1,
            "dirty_fields": 1,
            "update_on_open": True,
            "field_types": {"TOC": 1},
        }
        output_cache = {
            "status": "refreshed",
            "main_toc_fields": 1,
            "toc_entries": 1,
            "dirty_fields": 0,
            "update_on_open": False,
            "field_types": {"TOC": 1},
        }
        raw_backend = {
            "backend": "external",
            "target_id": "microsoft_word",
            "field_cache_verified": True,
            "page_count": 7,
            "selective_writeback": {
                "status": "selective_verified",
                "matched_fields": 1,
                "updated_fields": 1,
            },
            "read_only_verification": {
                "target_id": "microsoft_word",
                "operation": "verify_only",
                "read_only_verified": True,
                "repaginated": True,
                "saved": False,
                "pdf_exported": True,
                "page_count": 7,
            },
        }
        status = {
            "finalization_evidence_version": 1,
            "status": "pass",
            "delivery_field_status": "selective_verified",
            "input_field_cache": input_cache,
            "output_field_cache": output_cache,
            "field_backend": finalize_docx.canonical_backend_projection(
                raw_backend
            ),
            "field_writeback_status": "selective_verified",
            "field_completion": {
                "field_gate_completed": True,
                "final_ready_eligible": True,
                "word_verification_required": True,
                "word_verification_completed": True,
                "completion_scope": "target_word_verified",
                "evidence_validation": {"status": "pass", "errors": []},
            },
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
            "output": str(targets["output"].resolve(strict=False)),
            "workflow_state": {
                "stage": "finalized",
                "source_sha256": "1" * 64,
                "input_sha256": "2" * 64,
                "profile_sha256": "3" * 64,
                "structure_map_sha256": "4" * 64,
                "output_sha256": finalize_docx.file_sha256(staged["output"]),
            },
            "artifact_binding": {
                "version": 1,
                "finalized_docx": finalize_docx.artifact_identity(
                    staged["output"], reported_path=targets["output"]
                ),
                "word_verification_pdf": {
                    **finalize_docx.artifact_identity(
                        staged["pdf"], reported_path=targets["pdf"]
                    ),
                    "page_count": 7,
                },
            },
            "target_pdf": str(targets["pdf"].resolve(strict=False)),
            "target_layout_status": "target_pdf_ready_for_visual_qa",
            "target_software": "microsoft_word",
            "backend_audit": {
                "version": 1,
                "status": "persisted",
                "artifact": finalize_docx.artifact_identity(
                    staged["audit"], reported_path=targets["audit"]
                )
            },
        }
        payload = finalize_docx.standard_json_bytes(status)
        staged["status"].write_bytes(payload)
        self.trusted_status_identity = (
            finalize_docx.trusted_status_byte_identity(payload)
        )
        self.transaction_id = os.urandom(16).hex()
        snapshots = {
            name: finalize_docx._target_snapshot(path)
            for name, path in targets.items()
        }
        return staging, staged, targets, snapshots, old

    def publish(
        self,
        staged: dict[str, Path],
        targets: dict[str, Path],
        snapshots: dict[str, dict | None],
        staging: Path,
        *,
        event_hook=None,
    ) -> dict:
        return finalize_docx.publish_staged_artifacts(
            staged,
            targets,
            snapshots,
            staging,
            trusted_status_identity=self.trusted_status_identity,
            transaction_id=self.transaction_id,
            event_hook=event_hook,
            staged_expected={
                name: finalize_docx._target_snapshot(path, include_ctime=True)
                for name, path in staged.items()
            },
        )

    def test_publication_record_exclusive_collision_restores_targets_and_preserves_unknown(self) -> None:
        staging, staged, targets, snapshots, old = self.publish_fixture(existing=True)
        unknown = b"unknown-publication-record"

        def hook(event: str, name: str | None) -> None:
            if event == "before_staging_retention":
                (staging / finalize_docx.PUBLICATION_RECORD_NAME).write_bytes(unknown)

        with self.assertRaises(finalize_docx.FinalizationPublishError):
            self.publish(staged, targets, snapshots, staging, event_hook=hook)
        self.assertEqual(
            unknown,
            (staging / finalize_docx.PUBLICATION_RECORD_NAME).read_bytes(),
        )
        for name, target in targets.items():
            self.assertEqual(old[name], target.read_bytes())

    def test_publication_record_post_create_replacement_preserves_both_records(self) -> None:
        staging, staged, targets, snapshots, old = self.publish_fixture(existing=True)
        displaced = staging / "producer-publication-record.retained.json"
        unknown = b"unknown-replacement-record"
        producer_payload: bytes | None = None

        def hook(event: str, name: str | None) -> None:
            nonlocal producer_payload
            if event != "after_staging_retention":
                return
            record = staging / finalize_docx.PUBLICATION_RECORD_NAME
            producer_payload = record.read_bytes()
            os.rename(record, displaced)
            record.write_bytes(unknown)

        with self.assertRaises(finalize_docx.FinalizationPublishError):
            self.publish(staged, targets, snapshots, staging, event_hook=hook)
        self.assertIsNotNone(producer_payload)
        self.assertEqual(producer_payload, displaced.read_bytes())
        self.assertEqual(
            unknown,
            (staging / finalize_docx.PUBLICATION_RECORD_NAME).read_bytes(),
        )
        for name, target in targets.items():
            self.assertEqual(old[name], target.read_bytes())

    def test_publication_record_same_bytes_new_inode_is_rejected(self) -> None:
        staging, staged, targets, snapshots, old = self.publish_fixture(existing=True)
        displaced = staging / "producer-publication-record.same-bytes.json"
        producer_payload: bytes | None = None

        def hook(event: str, name: str | None) -> None:
            nonlocal producer_payload
            if event != "after_staging_retention":
                return
            record = staging / finalize_docx.PUBLICATION_RECORD_NAME
            producer_payload = record.read_bytes()
            os.rename(record, displaced)
            record.write_bytes(producer_payload)

        with self.assertRaises(finalize_docx.FinalizationPublishError):
            self.publish(staged, targets, snapshots, staging, event_hook=hook)
        self.assertIsNotNone(producer_payload)
        self.assertEqual(producer_payload, displaced.read_bytes())
        self.assertEqual(
            producer_payload,
            (staging / finalize_docx.PUBLICATION_RECORD_NAME).read_bytes(),
        )
        self.assertNotEqual(
            displaced.stat().st_ino,
            (staging / finalize_docx.PUBLICATION_RECORD_NAME).stat().st_ino,
        )
        for name, target in targets.items():
            self.assertEqual(old[name], target.read_bytes())

    def test_publication_record_extra_entry_after_commit_is_rejected_and_retained(self) -> None:
        staging, staged, targets, snapshots, old = self.publish_fixture(existing=True)
        extra = staging / "unknown-after-publication-record"

        def hook(event: str, name: str | None) -> None:
            if event == "after_staging_retention":
                extra.mkdir()
                (extra / "unknown.marker").write_bytes(b"unknown-extra")

        with self.assertRaises(finalize_docx.FinalizationPublishError):
            self.publish(staged, targets, snapshots, staging, event_hook=hook)
        self.assertEqual(b"unknown-extra", (extra / "unknown.marker").read_bytes())
        self.assertTrue(
            (staging / finalize_docx.PUBLICATION_RECORD_NAME).is_file()
        )
        for name, target in targets.items():
            self.assertEqual(old[name], target.read_bytes())

    def test_post_commit_authority_close_failures_are_aggregated_without_rollback(self) -> None:
        labels = ("transaction", "recovery", "staging", "output_parent")
        for failed_label in labels:
            with self.subTest(failed_label=failed_label):
                case = self.root / f"close-{failed_label}"
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    staging, staged, targets, snapshots, _ = self.publish_fixture(
                        existing=True
                    )
                    expected_bytes = {
                        name: path.read_bytes() for name, path in staged.items()
                    }
                    real_close = finalize_docx._close_authority
                    attempted: list[str] = []

                    def authority_label(authority) -> str:
                        if authority.path == case:
                            return "output_parent"
                        if authority.path == staging:
                            return "staging"
                        if authority.path.name == finalize_docx.PUBLICATION_RECOVERY_DIRECTORY:
                            return "recovery"
                        if authority.path.parent.name == finalize_docx.PUBLICATION_RECOVERY_DIRECTORY:
                            return "transaction"
                        return "other"

                    def close_with_fault(authority) -> None:
                        label = authority_label(authority)
                        attempted.append(label)
                        real_close(authority)
                        if label == failed_label:
                            raise OSError(f"synthetic {label} close failure")

                    with patch.object(
                        finalize_docx, "_close_authority", close_with_fault
                    ):
                        result = self.publish(staged, targets, snapshots, staging)
                    self.assertEqual(
                        "committed_with_cleanup_errors", result["commit_state"]
                    )
                    self.assertEqual(1, len(result["cleanup_errors"]))
                    self.assertIn(failed_label.replace("_", " "), result["cleanup_errors"][0])
                    self.assertEqual(
                        {"transaction", "recovery", "staging", "output_parent"},
                        set(attempted),
                    )
                    for name, target in targets.items():
                        self.assertEqual(expected_bytes[name], target.read_bytes())
                    self.assertTrue(
                        (staging / finalize_docx.PUBLICATION_RECORD_NAME).is_file()
                    )
                finally:
                    self.root = original_root

    def simple_fixture(
        self, *, existing: bool, include_pdf: bool
    ) -> tuple[
        Path,
        dict[str, Path],
        dict[str, Path],
        dict[str, dict | None],
        dict[str, bytes],
    ]:
        staging = self.root / "staging"
        staging.mkdir()
        targets = {"output": self.root / "final.docx"}
        if include_pdf:
            targets["pdf"] = self.root / "target.pdf"
        staged = {name: staging / path.name for name, path in targets.items()}
        old = {name: f"old-{name}".encode() for name in targets}
        if existing:
            for name, path in targets.items():
                path.write_bytes(old[name])
        for name, path in staged.items():
            path.write_bytes(f"new-{name}".encode())
        self.trusted_status_identity = None
        self.transaction_id = os.urandom(16).hex()
        snapshots = {
            name: finalize_docx._target_snapshot(path)
            for name, path in targets.items()
        }
        return staging, staged, targets, snapshots, old

    def no_fields_status_fixture(
        self, *, existing: bool
    ) -> tuple[
        Path,
        dict[str, Path],
        dict[str, Path],
        dict[str, dict | None],
        dict[str, bytes],
    ]:
        staging = self.root / "staging"
        staging.mkdir()
        targets = {
            "output": self.root / "final.docx",
            "audit": self.root / "finalization-backend-audit.json",
            "status": self.root / "finalization.json",
        }
        staged = {name: staging / path.name for name, path in targets.items()}
        old = {name: f"old-{name}".encode() for name in targets}
        if existing:
            for name, path in targets.items():
                path.write_bytes(old[name])
        staged["output"].write_bytes(b"new-output")
        staged["audit"].write_bytes(b"new-audit")
        absent_cache = {
            "status": "absent",
            "main_toc_fields": 0,
            "toc_entries": 0,
            "dirty_fields": 0,
            "update_on_open": False,
            "field_types": {},
        }
        status = {
            "finalization_evidence_version": 1,
            "status": "pass",
            "delivery_field_status": "absent",
            "input_field_cache": absent_cache,
            "output_field_cache": dict(absent_cache),
            "field_backend": finalize_docx.canonical_backend_projection(
                {"backend": "not_needed"}
            ),
            "backend_audit": {
                "version": 1,
                "status": "persisted",
                "artifact": finalize_docx.artifact_identity(
                    staged["audit"], reported_path=targets["audit"]
                ),
            },
            "artifact_binding": {
                "version": 1,
                "finalized_docx": finalize_docx.artifact_identity(
                    staged["output"], reported_path=targets["output"]
                ),
                "word_verification_pdf": None,
            },
            "field_writeback_status": "not_needed",
            "field_completion": {
                "field_gate_completed": True,
                "final_ready_eligible": True,
                "word_verification_required": False,
                "word_verification_completed": False,
                "completion_scope": "no_fields",
                "evidence_validation": {"status": "pass", "errors": []},
            },
            "content_integrity": "pass",
            "protected_object_integrity": "pass",
            "effective_font_integrity": "pass",
            "workflow_state": {
                "stage": "finalized",
                "source_sha256": "1" * 64,
                "input_sha256": "2" * 64,
                "profile_sha256": "3" * 64,
                "structure_map_sha256": "4" * 64,
                "output_sha256": finalize_docx.file_sha256(staged["output"]),
            },
            "target_pdf": None,
            "target_layout_status": "not_verified",
            "target_software": "microsoft_word",
            "output": str(targets["output"].resolve(strict=False)),
        }
        payload = finalize_docx.standard_json_bytes(status)
        staged["status"].write_bytes(payload)
        self.trusted_status_identity = (
            finalize_docx.trusted_status_byte_identity(payload)
        )
        self.transaction_id = os.urandom(16).hex()
        snapshots = {
            name: finalize_docx._target_snapshot(path)
            for name, path in targets.items()
        }
        return staging, staged, targets, snapshots, old

    def inject_entity(self, path: Path, kind: str, marker: bytes) -> None:
        incoming = self.root / f"incoming-{path.name}-{kind}"
        if incoming.is_symlink() or incoming.is_file():
            incoming.unlink()
        elif incoming.is_dir():
            incoming.rmdir()
        if kind == "regular":
            incoming.write_bytes(marker)
        elif kind == "symlink":
            incoming.symlink_to(marker.decode())
        elif kind == "directory":
            incoming.mkdir()
            (incoming / "marker").write_bytes(marker)
        else:
            self.fail(f"unknown injected kind: {kind}")
        try:
            os.replace(incoming, path)
        except (IsADirectoryError, NotADirectoryError, OSError):
            if path.is_symlink() or path.is_file():
                path.unlink()
            elif path.is_dir():
                path.rmdir()
            os.rename(incoming, path)

    def assert_marker_preserved(self, marker: bytes) -> None:
        found = False
        for path in self.root.rglob("*"):
            if path.is_symlink():
                found = found or os.readlink(path) == marker.decode()
            elif path.is_file():
                found = found or path.read_bytes() == marker
        self.assertTrue(found, f"concurrent marker {marker!r} was lost")

    def assert_old_preserved(self, payload: bytes) -> None:
        self.assertTrue(
            any(
                path.is_file()
                and not path.is_symlink()
                and path.read_bytes() == payload
                for path in self.root.rglob("*")
            ),
            f"old payload {payload!r} was lost",
        )

    def replace_directory_entry(
        self, path: Path, kind: str, marker: bytes
    ) -> Path:
        displaced = path.with_name(
            f"{path.name}-displaced-{os.urandom(4).hex()}"
        )
        os.rename(path, displaced)
        if kind == "regular":
            path.write_bytes(marker)
        elif kind == "symlink":
            path.symlink_to(marker.decode())
        elif kind == "directory":
            path.mkdir()
            (path / "marker").write_bytes(marker)
        else:
            self.fail(f"unknown replacement kind: {kind}")
        return displaced

    def write_open_descriptor(self, descriptor: int, marker: bytes) -> None:
        os.lseek(descriptor, 0, os.SEEK_SET)
        os.ftruncate(descriptor, 0)
        os.write(descriptor, marker)
        os.fsync(descriptor)

    def recovery_entry(self, recovery_root: Path, name: str) -> dict:
        manifest = json.loads(
            (recovery_root / "recovery-manifest.json").read_text(
                encoding="utf-8"
            )
        )
        self.assertEqual(1, manifest["version"])
        self.assertEqual(self.transaction_id, manifest["transaction_id"])
        return next(
            entry for entry in manifest["entries"] if entry["artifact"] == name
        )

    def test_open_fd_mutations_remain_reachable_in_success_recovery_matrix(self) -> None:
        for phase in (
            "after_capture",
            "after_post_commit_validation",
            "after_recovery_promotion",
        ):
            for target_name in ("output", "pdf", "audit", "status"):
                with self.subTest(phase=phase, target=target_name):
                    case = self.root / f"open-fd-success-{phase}-{target_name}"
                    case.mkdir()
                    original_root = self.root
                    self.root = case
                    descriptors: dict[str, int] = {}
                    try:
                        staging, staged, targets, snapshots, _ = self.publish_fixture(
                            existing=True
                        )
                        staged_payloads = {
                            name: path.read_bytes() for name, path in staged.items()
                        }
                        descriptors = {
                            name: os.open(path, os.O_RDWR)
                            for name, path in targets.items()
                        }
                        marker = f"open-fd-{phase}-{target_name}".encode()
                        written = False

                        def hook(event: str, name: str | None) -> None:
                            nonlocal written
                            if written or event != phase:
                                return
                            if phase in {"after_capture", "after_recovery_promotion"}:
                                if name != target_name:
                                    return
                            written = True
                            self.write_open_descriptor(
                                descriptors[target_name], marker
                            )

                        result = self.publish(
                            staged,
                            targets,
                            snapshots,
                            staging,
                            event_hook=hook,
                        )
                        self.assertTrue(written)
                        self.assertTrue(staging.is_dir())
                        self.assertTrue(
                            (staging / finalize_docx.PUBLICATION_RECORD_NAME).is_file()
                        )
                        self.assertEqual(
                            staged_payloads[target_name],
                            targets[target_name].read_bytes(),
                        )
                        recovery_root = Path(result["recovery_directory"])
                        entry = self.recovery_entry(recovery_root, target_name)
                        recovery_path = Path(entry["recovery_path"])
                        self.assertEqual(marker, recovery_path.read_bytes())
                        self.assertTrue(
                            entry["recovery_inode_may_continue_changing"]
                        )
                        self.assertNotEqual(
                            targets[target_name].stat().st_ino,
                            recovery_path.stat().st_ino,
                        )
                    finally:
                        for descriptor in descriptors.values():
                            os.close(descriptor)
                        self.root = original_root

    def test_open_fd_mutations_remain_reachable_during_rollback_matrix(self) -> None:
        for phase in (
            "after_capture",
            "after_post_commit_validation",
            "after_recovery_promotion",
        ):
            for target_name in ("output", "pdf", "audit", "status"):
                with self.subTest(phase=phase, target=target_name):
                    case = self.root / f"open-fd-rollback-{phase}-{target_name}"
                    case.mkdir()
                    original_root = self.root
                    self.root = case
                    descriptors: dict[str, int] = {}
                    try:
                        staging, staged, targets, snapshots, _ = self.publish_fixture(
                            existing=True
                        )
                        descriptors = {
                            name: os.open(path, os.O_RDWR)
                            for name, path in targets.items()
                        }
                        marker = f"rollback-fd-{phase}-{target_name}".encode()
                        injected = False

                        def hook(event: str, name: str | None) -> None:
                            nonlocal injected
                            if injected or event != phase:
                                return
                            if phase in {"after_capture", "after_recovery_promotion"}:
                                if name != target_name:
                                    return
                            injected = True
                            self.write_open_descriptor(
                                descriptors[target_name], marker
                            )
                            raise OSError("synthetic post-capture rollback")

                        with self.assertRaises(
                            finalize_docx.FinalizationPublishError
                        ) as caught:
                            self.publish(
                                staged,
                                targets,
                                snapshots,
                                staging,
                                event_hook=hook,
                            )
                        self.assertTrue(injected)
                        self.assertTrue(caught.exception.preserve_staging)
                        self.assert_marker_preserved(marker)
                    finally:
                        for descriptor in descriptors.values():
                            os.close(descriptor)
                        self.root = original_root

    def test_publish_adversarial_interleavings_preserve_unknown_entities(self) -> None:
        phases = (
            "after_start_snapshot",
            "before_capture",
            "after_capture",
            "after_publish",
        )
        for existing in (False, True):
            for target_name in ("output", "pdf", "audit", "status"):
                for phase in phases:
                    for kind in ("regular", "symlink", "directory"):
                        with self.subTest(
                            existing=existing,
                            target=target_name,
                            phase=phase,
                            kind=kind,
                        ):
                            case = self.root / (
                                f"case-{int(existing)}-{target_name}-{phase}-{kind}"
                            )
                            case.mkdir()
                            original_root = self.root
                            self.root = case
                            try:
                                staging, staged, targets, snapshots, old = (
                                    self.publish_fixture(existing=existing)
                                )
                                marker = (
                                    f"concurrent-{existing}-{target_name}-{phase}-{kind}"
                                ).encode()
                                injected = False

                                def hook(event: str, name: str | None) -> None:
                                    nonlocal injected
                                    matches = event == phase and (
                                        event == "after_start_snapshot"
                                        or name == target_name
                                    )
                                    if matches and not injected:
                                        injected = True
                                        self.inject_entity(
                                            targets[target_name], kind, marker
                                        )

                                with self.assertRaises(
                                    finalize_docx.FinalizationPublishError
                                ) as caught:
                                    self.publish(
                                        staged,
                                        targets,
                                        snapshots,
                                        staging,
                                        event_hook=hook,
                                    )
                                self.assertTrue(injected)
                                self.assertIn("phase=", str(caught.exception))
                                self.assertTrue(caught.exception.preserve_staging)
                                self.assert_marker_preserved(marker)
                                if existing and phase in {
                                    "after_capture",
                                    "after_publish",
                                }:
                                    self.assert_old_preserved(old[target_name])
                            finally:
                                self.root = original_root

    def test_absent_targets_use_atomic_no_replace_publish(self) -> None:
        staging, staged, targets, snapshots, _ = self.publish_fixture(existing=False)
        blocker = b"concurrent-exclusive-create"

        def hook(event: str, name: str | None) -> None:
            if event == "after_capture" and name == "output":
                targets["output"].write_bytes(blocker)

        with self.assertRaises(finalize_docx.FinalizationPublishError) as caught:
            self.publish(
                staged, targets, snapshots, staging, event_hook=hook
            )
        self.assertTrue(caught.exception.preserve_staging)
        self.assertEqual(blocker, targets["output"].read_bytes())

    def test_success_publishes_status_last_and_revalidates_after_commit(self) -> None:
        staging, staged, targets, snapshots, _ = self.publish_fixture(existing=False)
        events: list[tuple[str, str | None]] = []

        def hook(event: str, name: str | None) -> None:
            events.append((event, name))

        staged_payloads = {
            name: path.read_bytes() for name, path in staged.items()
        }
        result = self.publish(
            staged, targets, snapshots, staging, event_hook=hook
        )
        publishes = [name for event, name in events if event == "after_publish"]
        self.assertEqual(["output", "pdf", "audit", "status"], publishes)
        self.assertLess(
            events.index(("before_commit_validation", None)),
            events.index(("before_capture", "status")),
        )
        self.assertGreater(
            events.index(("before_post_commit_validation", None)),
            events.index(("after_publish", "status")),
        )
        self.assertFalse(result["staging_removed"])
        self.assertTrue(result["staging_retained"])
        self.assertEqual(str(staging.absolute()), result["retained_staging_directory"])
        record = json.loads(
            (staging / finalize_docx.PUBLICATION_RECORD_NAME).read_text(
                encoding="utf-8"
            )
        )
        self.assertEqual(self.transaction_id, record["transaction_id"])
        self.assertEqual("manual_only", record["cleanup_policy"])
        self.assertFalse(record["business_gate"])
        for name, target in targets.items():
            self.assertEqual(staged_payloads[name], target.read_bytes())

    def test_success_atomic_move_matrix_has_no_staged_sources_or_new_inode_in_recovery(self) -> None:
        cases = (
            (False, False, False),
            (True, False, False),
            (False, True, False),
            (True, True, False),
            (False, False, True),
            (True, False, True),
            (False, True, True),
            (True, True, True),
        )
        for existing, include_pdf, include_status in cases:
            with self.subTest(
                existing=existing,
                include_pdf=include_pdf,
                include_status=include_status,
            ):
                case = self.root / (
                    f"atomic-move-matrix-{int(existing)}-{int(include_pdf)}-"
                    f"{int(include_status)}"
                )
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    if include_status and include_pdf:
                        fixture = self.publish_fixture(existing=existing)
                    elif include_status:
                        fixture = self.no_fields_status_fixture(existing=existing)
                    else:
                        fixture = self.simple_fixture(
                            existing=existing, include_pdf=include_pdf
                        )
                    staging, staged, targets, snapshots, _ = fixture
                    staged_payloads = {
                        name: path.read_bytes() for name, path in staged.items()
                    }
                    result = self.publish(staged, targets, snapshots, staging)
                    self.assertTrue(staging.is_dir())
                    self.assertTrue(
                        (staging / finalize_docx.PUBLICATION_RECORD_NAME).is_file()
                    )
                    for name, target in targets.items():
                        self.assertEqual(staged_payloads[name], target.read_bytes())
                        self.assertFalse(staged[name].exists())
                    recovery_value = result["recovery_directory"]
                    self.assertEqual(existing, recovery_value is not None)
                    if recovery_value is not None:
                        recovery_root = Path(recovery_value)
                        target_inodes = {
                            (path.stat().st_dev, path.stat().st_ino)
                            for path in targets.values()
                        }
                        for recovery_path in recovery_root.glob("*.previous"):
                            self.assertNotIn(
                                (
                                    recovery_path.stat().st_dev,
                                    recovery_path.stat().st_ino,
                                ),
                                target_inodes,
                            )
                finally:
                    self.root = original_root

    def test_atomic_no_replace_failure_matrix_rolls_back(self) -> None:
        failures = (
            ("exists", FileExistsError(errno.EEXIST, "synthetic collision")),
            ("cross_device", OSError(errno.EXDEV, "synthetic cross-device")),
            ("permission", PermissionError(errno.EACCES, "synthetic denial")),
            (
                "unsupported",
                finalize_docx.FinalizationPublishError(
                    "synthetic unsupported no-replace primitive"
                ),
            ),
        )
        for failure_name, failure in failures:
            with self.subTest(failure=failure_name):
                case = self.root / f"atomic-move-failure-{failure_name}"
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    staging, staged, targets, snapshots, old = self.simple_fixture(
                        existing=True, include_pdf=False
                    )
                    real_move = finalize_docx._platform_atomic_noreplace_move
                    failed = False

                    def move(source, target):
                        nonlocal failed
                        if source.name == staged["output"].name and not failed:
                            failed = True
                            raise failure
                        return real_move(source, target)

                    with patch.object(
                        finalize_docx,
                        "_platform_atomic_noreplace_move",
                        side_effect=move,
                    ), self.assertRaises(finalize_docx.FinalizationPublishError):
                        self.publish(staged, targets, snapshots, staging)
                    self.assertTrue(failed)
                    for name, target in targets.items():
                        self.assertEqual(old[name], target.read_bytes())
                finally:
                    self.root = original_root

    def test_staged_source_swap_before_and_inside_atomic_primitive_is_rejected(self) -> None:
        for window in ("before_atomic_publish", "after_source_validation"):
            for kind in ("regular", "symlink", "directory"):
                with self.subTest(window=window, kind=kind):
                    case = self.root / f"source-swap-{window}-{kind}"
                    case.mkdir()
                    original_root = self.root
                    self.root = case
                    try:
                        staging, staged, targets, snapshots, old = self.simple_fixture(
                            existing=True, include_pdf=False
                        )
                        marker = f"source-swap-{window}-{kind}".encode()
                        swapped = False
                        real_move = finalize_docx._platform_atomic_noreplace_move

                        def hook(event: str, name: str | None) -> None:
                            nonlocal swapped
                            if (
                                window == "before_atomic_publish"
                                and event == window
                                and name == "output"
                                and not swapped
                            ):
                                swapped = True
                                self.inject_entity(staged["output"], kind, marker)

                        def move(source, target):
                            nonlocal swapped
                            if (
                                window == "after_source_validation"
                                and source.authority.path == staging
                                and source.name == staged["output"].name
                                and not swapped
                            ):
                                swapped = True
                                self.inject_entity(staged["output"], kind, marker)
                            return real_move(source, target)

                        with patch.object(
                            finalize_docx,
                            "_platform_atomic_noreplace_move",
                            side_effect=move,
                        ), self.assertRaises(
                            finalize_docx.FinalizationPublishError
                        ) as caught:
                            self.publish(
                                staged,
                                targets,
                                snapshots,
                                staging,
                                event_hook=hook,
                            )
                        self.assertTrue(swapped)
                        self.assertTrue(caught.exception.preserve_staging)
                        self.assert_marker_preserved(marker)
                        self.assert_old_preserved(old["output"])
                    finally:
                        self.root = original_root

    def test_atomic_no_replace_errno_and_unsupported_platform_mapping(self) -> None:
        source = self.root / "source"
        target = self.root / "target"
        mappings = (
            (errno.EEXIST, FileExistsError),
            (errno.ENOTEMPTY, FileExistsError),
            (errno.EACCES, PermissionError),
            (errno.EPERM, PermissionError),
            (errno.EXDEV, OSError),
            (errno.EINVAL, finalize_docx.FinalizationPublishError),
            (getattr(errno, "ENOSYS", errno.EINVAL), finalize_docx.FinalizationPublishError),
        )
        for error_code, expected_type in mappings:
            with self.subTest(errno=error_code):
                with self.assertRaises(expected_type):
                    finalize_docx._raise_atomic_move_error(
                        error_code, source, target
                    )
        with patch.object(finalize_docx.sys, "platform", "unsupported-test"):
            authority = finalize_docx._DirectoryAuthority(
                self.root, -1, 1, 1, os.geteuid(), 0o700
            )
            with self.assertRaises(finalize_docx.FinalizationPublishError):
                finalize_docx._platform_atomic_noreplace_move(
                    finalize_docx._ArtifactLocation(authority, "source"),
                    finalize_docx._ArtifactLocation(authority, "target"),
                )

    def test_platform_atomic_no_replace_uses_official_exclusive_flags(self) -> None:
        source_authority = finalize_docx._DirectoryAuthority(
            self.root / "source-parent", 11, 1, 1, os.geteuid(), 0o700
        )
        target_authority = finalize_docx._DirectoryAuthority(
            self.root / "target-parent", 12, 1, 1, os.geteuid(), 0o700
        )
        source = finalize_docx._ArtifactLocation(source_authority, "source.bin")
        target = finalize_docx._ArtifactLocation(target_authority, "target.bin")

        for platform_name, function_name, flag in (
            ("darwin", "renameatx_np", 0x00000004),
            ("linux", "renameat2", 0x00000001),
        ):
            with self.subTest(platform=platform_name):
                function = MagicMock(return_value=0)
                library = MagicMock()
                setattr(library, function_name, function)
                with patch.object(finalize_docx.sys, "platform", platform_name), patch.object(
                    finalize_docx.ctypes, "CDLL", return_value=library
                ):
                    finalize_docx._platform_atomic_noreplace_move(source, target)
                self.assertEqual(flag, function.call_args.args[-1])
                self.assertEqual(11, function.call_args.args[0])
                self.assertEqual(12, function.call_args.args[2])

        with patch.object(finalize_docx.sys, "platform", "win32"), patch.object(
            finalize_docx.os, "name", "nt"
        ):
            with self.assertRaisesRegex(
                finalize_docx.FinalizationPublishError,
                "authority-bound no-replace rename",
            ):
                finalize_docx._platform_atomic_noreplace_move(source, target)

    def test_windows_private_dacl_is_never_authorized_by_string_matrix(self) -> None:
        cases = (
            "NULL DACL",
            "empty DACL",
            "inherited allow",
            "numeric write rights",
            "Domain Users",
            "Interactive",
            "Authenticated Users",
            "Builtin Users",
            "Everyone",
            "OA object allow",
            "XA callback allow",
            "ZA callback-object allow",
            "deny then allow",
            "owner changed",
        )
        api = object.__new__(finalize_docx._WindowsPublisherApi)
        for case in cases:
            with self.subTest(case=case), self.assertRaisesRegex(
                finalize_docx.FinalizationPublishError,
                "AccessCheck-based effective-permission proof",
            ):
                api.private_security_identity(Path(case))

    def test_simulated_windows_full_publisher_existing_absent_and_status_last(self) -> None:
        for existing in (False, True):
            with self.subTest(existing=existing):
                case = self.root / f"windows-e2e-{int(existing)}"
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    staging, staged, targets, snapshots, _ = self.publish_fixture(
                        existing=existing
                    )
                    events: list[tuple[str, str | None]] = []
                    with self.simulated_windows_backend():
                        result = self.publish(
                            staged,
                            targets,
                            snapshots,
                            staging,
                            event_hook=lambda event, name: events.append(
                                (event, name)
                            ),
                        )
                    published = [
                        name for event, name in events if event == "after_publish"
                    ]
                    self.assertEqual(
                        ["output", "pdf", "audit", "status"], published
                    )
                    self.assertTrue(staging.is_dir())
                    self.assertFalse(result["staging_removed"])
                    self.assertTrue(result["staging_retained"])
                    self.assertTrue(Path(result["publication_record"]).is_file())
                    self.assertEqual(
                        existing, result["recovery_directory"] is not None
                    )
                    if existing:
                        recovery = Path(result["recovery_directory"])
                        self.assertTrue(
                            (recovery / "recovery-manifest.json").is_file()
                        )
                finally:
                    self.root = original_root

    def test_simulated_windows_full_publisher_rollback_winerror_matrix(self) -> None:
        failures = (
            FileExistsError(183, "destination exists"),
            PermissionError(5, "access denied"),
            OSError(errno.EXDEV, "cross-volume"),
            finalize_docx.FinalizationPublishError(
                "MoveFileExW unsupported WinError 120"
            ),
        )
        for failure in failures:
            with self.subTest(failure=type(failure).__name__):
                case = self.root / f"windows-rollback-{len(list(self.root.iterdir()))}"
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    staging, staged, targets, snapshots, old = self.publish_fixture(
                        existing=True
                    )

                    class Fault:
                        error = failure

                        def __call__(self, source: Path, target: Path) -> bool:
                            return (
                                source.parent.name == staging.name
                                and source.name == staged["output"].name
                                and target.name == targets["output"].name
                            )

                    with self.simulated_windows_backend() as api:
                        api.move_fault = Fault()
                        with self.assertRaises(
                            finalize_docx.FinalizationPublishError
                        ) as caught:
                            self.publish(staged, targets, snapshots, staging)
                    self.assertTrue(caught.exception.preserve_staging)
                    self.assertTrue(staging.is_dir())
                    for name, target in targets.items():
                        self.assertEqual(old[name], target.read_bytes())
                finally:
                    self.root = original_root

    def test_simulated_windows_staging_authority_replacement_is_preserved(self) -> None:
        for kind in ("regular", "symlink", "empty_directory", "nonempty_directory"):
            with self.subTest(kind=kind):
                case = self.root / f"windows-staging-replacement-{kind}"
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    staging, staged, targets, snapshots, old = self.publish_fixture(
                        existing=True
                    )
                    marker = f"windows-staging-{kind}".encode()
                    displaced = self.root / f"displaced-{kind}"
                    replaced = False

                    def hook(event: str, name: str | None) -> None:
                        nonlocal replaced
                        if replaced or event != "after_staging_authority_open":
                            return
                        os.rename(staging, displaced)
                        if kind == "regular":
                            staging.write_bytes(marker)
                        elif kind == "symlink":
                            staging.symlink_to(marker.decode())
                        else:
                            staging.mkdir(mode=0o700)
                            if kind == "nonempty_directory":
                                (staging / "unknown.marker").write_bytes(marker)
                        replaced = True

                    with self.simulated_windows_backend(), self.assertRaises(
                        finalize_docx.FinalizationPublishError
                    ) as caught:
                        self.publish(
                            staged,
                            targets,
                            snapshots,
                            staging,
                            event_hook=hook,
                        )
                    self.assertTrue(replaced)
                    self.assertTrue(caught.exception.preserve_staging)
                    self.assertTrue(staging.is_symlink() or staging.exists())
                    self.assertTrue(displaced.exists())
                    if kind == "regular":
                        self.assertEqual(marker, staging.read_bytes())
                    elif kind == "symlink":
                        self.assertEqual(marker.decode(), os.readlink(staging))
                    elif kind == "nonempty_directory":
                        self.assertEqual(
                            marker, (staging / "unknown.marker").read_bytes()
                        )
                    self.assert_old_preserved(old["output"])
                finally:
                    self.root = original_root

    def test_recovery_directory_authority_replacement_matrix_fails_closed(self) -> None:
        cases = (
            ("before_recovery_parent_open", "parent"),
            ("after_recovery_parent_open", "parent"),
            ("after_recovery_transaction_mkdir", "transaction"),
            ("after_recovery_transaction_create", "transaction"),
            ("before_recovery_backup_move", "transaction"),
            ("after_recovery_promotion", "transaction"),
            ("before_recovery_manifest", "transaction"),
            ("after_recovery_manifest", "transaction"),
            ("before_recovery_return", "transaction"),
        )
        for phase, scope in cases:
            for kind in ("regular", "symlink", "directory"):
                with self.subTest(phase=phase, scope=scope, kind=kind):
                    case = self.root / f"recovery-authority-{phase}-{kind}"
                    case.mkdir()
                    original_root = self.root
                    self.root = case
                    try:
                        staging, staged, targets, snapshots, old = self.simple_fixture(
                            existing=True, include_pdf=False
                        )
                        marker = f"authority-{phase}-{kind}".encode()
                        replaced = False

                        def hook(event: str, name: str | None) -> None:
                            nonlocal replaced
                            if replaced or event != phase:
                                return
                            if phase in {
                                "before_recovery_backup_move",
                                "after_recovery_promotion",
                            } and name != "output":
                                return
                            recovery = (
                                self.root
                                / finalize_docx.PUBLICATION_RECOVERY_DIRECTORY
                            )
                            entry = (
                                recovery
                                if scope == "parent"
                                else recovery / self.transaction_id
                            )
                            self.replace_directory_entry(entry, kind, marker)
                            replaced = True

                        with self.assertRaises(
                            finalize_docx.FinalizationPublishError
                        ) as caught:
                            self.publish(
                                staged,
                                targets,
                                snapshots,
                                staging,
                                event_hook=hook,
                            )
                        self.assertTrue(replaced)
                        self.assertTrue(caught.exception.preserve_staging)
                        self.assertIn(self.transaction_id, str(caught.exception))
                        self.assert_marker_preserved(marker)
                        self.assert_old_preserved(old["output"])
                    finally:
                        self.root = original_root

    def test_output_parent_authority_replacement_matrix_fails_closed(self) -> None:
        phases = (
            "after_start_snapshot",
            "after_post_commit_validation",
            "after_recovery_promotion",
            "before_recovery_return",
        )
        for phase in phases:
            for kind in ("regular", "symlink", "directory"):
                with self.subTest(phase=phase, kind=kind):
                    outer = self.root / f"output-parent-{phase}-{kind}"
                    outer.mkdir()
                    output_parent = outer / "output"
                    output_parent.mkdir(mode=0o700)
                    original_root = self.root
                    self.root = output_parent
                    try:
                        staging, staged, targets, snapshots, old = self.simple_fixture(
                            existing=True, include_pdf=False
                        )
                        marker = f"output-parent-{phase}-{kind}".encode()
                        replaced = False

                        def hook(event: str, name: str | None) -> None:
                            nonlocal replaced
                            if replaced or event != phase:
                                return
                            if phase == "after_recovery_promotion" and name != "output":
                                return
                            self.replace_directory_entry(
                                output_parent, kind, marker
                            )
                            replaced = True

                        with self.assertRaises(
                            finalize_docx.FinalizationPublishError
                        ) as caught:
                            self.publish(
                                staged,
                                targets,
                                snapshots,
                                staging,
                                event_hook=hook,
                            )
                        self.assertTrue(replaced)
                        self.assertTrue(caught.exception.preserve_staging)
                        found_marker = False
                        found_old = False
                        for path in outer.rglob("*"):
                            if path.is_symlink():
                                found_marker = found_marker or os.readlink(path) == marker.decode()
                            elif path.is_file():
                                payload = path.read_bytes()
                                found_marker = found_marker or payload == marker
                                found_old = found_old or payload == old["output"]
                        self.assertTrue(found_marker)
                        self.assertTrue(found_old)
                    finally:
                        self.root = original_root

    def test_recovery_mode_owner_existing_transaction_and_fault_matrix(self) -> None:
        recovery_name = finalize_docx.PUBLICATION_RECOVERY_DIRECTORY

        case = self.root / "recovery-world-writable"
        case.mkdir()
        original_root = self.root
        self.root = case
        try:
            recovery = self.root / recovery_name
            recovery.mkdir(mode=0o700)
            recovery.chmod(0o777)
            staging, staged, targets, snapshots, old = self.simple_fixture(
                existing=True, include_pdf=False
            )
            with self.assertRaises(finalize_docx.FinalizationPublishError):
                self.publish(staged, targets, snapshots, staging)
            self.assertEqual(0o777, recovery.stat().st_mode & 0o777)
            self.assert_old_preserved(old["output"])
        finally:
            self.root = original_root

        owner_case = self.root / "recovery-owner"
        owner_case.mkdir(mode=0o700)
        with patch.object(
            finalize_docx.os, "geteuid", return_value=os.geteuid() + 1
        ), self.assertRaises(finalize_docx.FinalizationPublishError):
            finalize_docx._open_directory_authority(
                owner_case, require_private_owner=True
            )

        for fault in ("existing_transaction", "manifest_write", "mkdir_permission"):
            with self.subTest(fault=fault):
                case = self.root / f"recovery-{fault}"
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    staging, staged, targets, snapshots, old = self.simple_fixture(
                        existing=True, include_pdf=False
                    )
                    recovery = self.root / recovery_name
                    patches = contextlib.ExitStack()
                    if fault == "existing_transaction":
                        recovery.mkdir(mode=0o700)
                        (recovery / self.transaction_id).mkdir(mode=0o700)
                    elif fault == "manifest_write":
                        patches.enter_context(
                            patch.object(
                                finalize_docx,
                                "_write_manifest_exclusive",
                                side_effect=PermissionError(
                                    errno.EACCES, "synthetic manifest denial"
                                ),
                            )
                        )
                    else:
                        real_mkdir = finalize_docx.os.mkdir

                        def mkdir(path, *args, **kwargs):
                            if path == recovery_name:
                                raise PermissionError(
                                    errno.EACCES, "synthetic recovery mkdir denial"
                                )
                            return real_mkdir(path, *args, **kwargs)

                        patches.enter_context(
                            patch.object(finalize_docx.os, "mkdir", side_effect=mkdir)
                        )
                    with patches, self.assertRaises(
                        finalize_docx.FinalizationPublishError
                    ) as caught:
                        self.publish(staged, targets, snapshots, staging)
                    self.assertTrue(caught.exception.preserve_staging)
                    self.assert_old_preserved(old["output"])
                finally:
                    self.root = original_root

    def test_recovery_artifact_and_manifest_entry_replacement_is_rejected(self) -> None:
        for phase, entry_name in (
            ("after_recovery_promotion", "output.previous"),
            ("after_recovery_manifest", "recovery-manifest.json"),
        ):
            for kind in ("regular", "symlink", "directory"):
                with self.subTest(phase=phase, kind=kind):
                    case = self.root / f"recovery-entry-{phase}-{kind}"
                    case.mkdir()
                    original_root = self.root
                    self.root = case
                    try:
                        staging, staged, targets, snapshots, old = self.simple_fixture(
                            existing=True, include_pdf=False
                        )
                        marker = f"recovery-entry-{phase}-{kind}".encode()
                        replaced = False

                        def hook(event: str, name: str | None) -> None:
                            nonlocal replaced
                            if replaced or event != phase:
                                return
                            if phase == "after_recovery_promotion" and name != "output":
                                return
                            transaction = (
                                self.root
                                / finalize_docx.PUBLICATION_RECOVERY_DIRECTORY
                                / self.transaction_id
                            )
                            entry = transaction / entry_name
                            displaced = transaction / f"{entry_name}.displaced"
                            os.rename(entry, displaced)
                            if kind == "regular":
                                entry.write_bytes(marker)
                            elif kind == "symlink":
                                entry.symlink_to(marker.decode())
                            else:
                                entry.mkdir()
                                (entry / "marker").write_bytes(marker)
                            replaced = True

                        with self.assertRaises(
                            finalize_docx.FinalizationPublishError
                        ) as caught:
                            self.publish(
                                staged,
                                targets,
                                snapshots,
                                staging,
                                event_hook=hook,
                            )
                        self.assertTrue(replaced)
                        self.assertTrue(caught.exception.preserve_staging)
                        self.assert_marker_preserved(marker)
                        self.assert_old_preserved(old["output"])
                    finally:
                        self.root = original_root

    def test_recovery_manifest_writer_identity_and_entry_set_matrix(self) -> None:
        for phase in ("after_recovery_manifest", "before_recovery_return"):
            for mutation in (
                "same_bytes_new_inode",
                "different_bytes",
                "symlink",
                "directory",
                "extra_entry",
            ):
                with self.subTest(phase=phase, mutation=mutation):
                    case = self.root / f"recovery-manifest-{phase}-{mutation}"
                    case.mkdir()
                    original_root = self.root
                    self.root = case
                    try:
                        staging, staged, targets, snapshots, old = self.simple_fixture(
                            existing=True, include_pdf=False
                        )
                        victim = case / "external-victim.bin"
                        victim.write_bytes(b"external-victim-unchanged")
                        changed = False
                        displaced: Path | None = None
                        extra: Path | None = None
                        original_payload: bytes | None = None

                        def hook(event: str, name: str | None) -> None:
                            nonlocal changed, displaced, extra, original_payload
                            if changed or event != phase:
                                return
                            transaction = (
                                self.root
                                / finalize_docx.PUBLICATION_RECOVERY_DIRECTORY
                                / self.transaction_id
                            )
                            manifest = transaction / "recovery-manifest.json"
                            if mutation == "extra_entry":
                                extra = transaction / "unknown-extra-entry"
                                extra.mkdir()
                                (extra / "unknown.marker").write_bytes(
                                    b"unknown-extra"
                                )
                            else:
                                original_payload = manifest.read_bytes()
                                displaced = transaction / (
                                    f"recovery-manifest.{mutation}.retained.json"
                                )
                                os.rename(manifest, displaced)
                                if mutation == "same_bytes_new_inode":
                                    manifest.write_bytes(original_payload)
                                elif mutation == "different_bytes":
                                    manifest.write_bytes(b"different-manifest")
                                elif mutation == "symlink":
                                    manifest.symlink_to(victim)
                                else:
                                    manifest.mkdir()
                                    (manifest / "unknown.marker").write_bytes(
                                        b"unknown-directory"
                                    )
                            changed = True

                        with self.assertRaises(
                            finalize_docx.FinalizationPublishError
                        ) as caught:
                            self.publish(
                                staged,
                                targets,
                                snapshots,
                                staging,
                                event_hook=hook,
                            )
                        self.assertTrue(changed)
                        self.assertTrue(caught.exception.preserve_staging)
                        self.assertEqual(b"external-victim-unchanged", victim.read_bytes())
                        if displaced is not None:
                            self.assertTrue(displaced.is_file())
                            self.assertEqual(original_payload, displaced.read_bytes())
                        if extra is not None:
                            self.assertEqual(
                                b"unknown-extra",
                                (extra / "unknown.marker").read_bytes(),
                            )
                        self.assert_old_preserved(old["output"])
                    finally:
                        self.root = original_root

    @unittest.skipUnless(os.name == "posix", "chmod semantics are POSIX-specific")
    def test_chmod_changed_staging_authority_fails_closed_with_evidence(self) -> None:
        staging, staged, targets, snapshots, old = self.publish_fixture(existing=True)
        changed_mode = False

        def hook(event: str, name: str | None) -> None:
            nonlocal changed_mode
            if event == "before_atomic_publish" and name == "output":
                staging.chmod(0o500)
                changed_mode = True

        try:
            with self.assertRaises(finalize_docx.FinalizationPublishError) as caught:
                self.publish(
                    staged,
                    targets,
                    snapshots,
                    staging,
                    event_hook=hook,
                )
            self.assertTrue(changed_mode)
            self.assertTrue(caught.exception.preserve_staging)
            self.assertTrue(staging.exists())
            self.assert_old_preserved(old["output"])
            self.assertEqual(old["status"], targets["status"].read_bytes())
        finally:
            staging.chmod(0o700)

    def test_partial_cleanup_and_recreated_staged_source_fail_closed(self) -> None:
        for mode in ("unknown_residue", "recreated_source"):
            with self.subTest(mode=mode):
                case = self.root / mode
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    staging, staged, targets, snapshots, old = self.publish_fixture(
                        existing=True
                    )
                    marker = f"{mode}-marker".encode()
                    injected = False

                    def hook(event: str, name: str | None) -> None:
                        nonlocal injected
                        if injected:
                            return
                        if mode == "unknown_residue" and event == "after_post_commit_validation":
                            injected = True
                            (staging / "unknown-concurrent-residue").write_bytes(marker)
                        elif (
                            mode == "recreated_source"
                            and event == "after_atomic_publish"
                            and name == "output"
                        ):
                            injected = True
                            staged["output"].write_bytes(marker)

                    with self.assertRaises(
                        finalize_docx.FinalizationPublishError
                    ) as caught:
                        self.publish(
                            staged,
                            targets,
                            snapshots,
                            staging,
                            event_hook=hook,
                        )
                    self.assertTrue(injected)
                    self.assertTrue(caught.exception.preserve_staging)
                    self.assert_marker_preserved(marker)
                    if mode == "unknown_residue":
                        for name, target in targets.items():
                            self.assertEqual(old[name], target.read_bytes())
                finally:
                    self.root = original_root

    def test_rollback_race_preserves_concurrent_target_old_backup_and_quarantine(
        self,
    ) -> None:
        staging, staged, targets, snapshots, old = self.publish_fixture(existing=True)
        first = b"first-concurrent-object"
        late = b"late-rollback-object"
        injected = set()

        def hook(event: str, name: str | None) -> None:
            key = (event, name)
            if event == "after_publish" and name == "output" and key not in injected:
                injected.add(key)
                self.inject_entity(targets["output"], "regular", first)
            if (
                event == "after_rollback_capture"
                and name == "output"
                and key not in injected
            ):
                injected.add(key)
                targets["output"].write_bytes(late)

        with self.assertRaises(finalize_docx.FinalizationPublishError) as caught:
            self.publish(
                staged, targets, snapshots, staging, event_hook=hook
            )
        self.assertTrue(caught.exception.preserve_staging)
        self.assertEqual(late, targets["output"].read_bytes())
        self.assert_marker_preserved(first)
        self.assert_old_preserved(old["output"])

    def test_status_binding_is_rechecked_before_and_after_commit(self) -> None:
        for phase in ("before_commit_validation", "after_publish"):
            with self.subTest(phase=phase):
                case = self.root / phase
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    staging, staged, targets, snapshots, old = self.publish_fixture(
                        existing=True
                    )
                    changed = False

                    def hook(event: str, name: str | None) -> None:
                        nonlocal changed
                        if changed:
                            return
                        if event == phase and (
                            phase == "before_commit_validation" or name == "status"
                        ):
                            changed = True
                            if phase == "before_commit_validation":
                                value = json.loads(
                                    staged["status"].read_text(encoding="utf-8")
                                )
                                value["artifact_binding"]["finalized_docx"][
                                    "sha256"
                                ] = "0" * 64
                                staged["status"].write_text(
                                    json.dumps(value), encoding="utf-8"
                                )
                            else:
                                self.inject_entity(
                                    targets["output"], "regular", b"post-status-race"
                                )

                    with self.assertRaises(finalize_docx.FinalizationPublishError):
                        self.publish(
                            staged,
                            targets,
                            snapshots,
                            staging,
                            event_hook=hook,
                        )
                    self.assertTrue(changed)
                    self.assert_old_preserved(old["status"])
                finally:
                    self.root = original_root

    def test_staged_status_artifact_binding_conflict_matrix_is_rejected(self) -> None:
        mutations = (
            ("output_path", lambda value: value.update(output="wrong.docx")),
            (
                "workflow_output_hash",
                lambda value: value["workflow_state"].update(
                    output_sha256="0" * 64
                ),
            ),
            (
                "docx_path",
                lambda value: value["artifact_binding"]["finalized_docx"].update(
                    path="wrong.docx"
                ),
            ),
            (
                "docx_hash",
                lambda value: value["artifact_binding"]["finalized_docx"].update(
                    sha256="0" * 64
                ),
            ),
            (
                "docx_size",
                lambda value: value["artifact_binding"]["finalized_docx"].update(
                    size_bytes=999
                ),
            ),
            ("pdf_path", lambda value: value.update(target_pdf="wrong.pdf")),
            (
                "pdf_hash",
                lambda value: value["artifact_binding"][
                    "word_verification_pdf"
                ].update(sha256="0" * 64),
            ),
            (
                "pdf_size",
                lambda value: value["artifact_binding"][
                    "word_verification_pdf"
                ].update(size_bytes=999),
            ),
            (
                "audit_hash",
                lambda value: value["backend_audit"]["artifact"].update(
                    sha256="0" * 64
                ),
            ),
            (
                "audit_size",
                lambda value: value["backend_audit"]["artifact"].update(
                    size_bytes=999
                ),
            ),
        )
        for name, mutate in mutations:
            with self.subTest(name=name):
                case = self.root / f"status-binding-{name}"
                case.mkdir()
                original_root = self.root
                self.root = case
                try:
                    staging, staged, targets, snapshots, _ = self.publish_fixture(
                        existing=False
                    )
                    value = json.loads(staged["status"].read_text(encoding="utf-8"))
                    mutate(value)
                    staged["status"].write_text(json.dumps(value), encoding="utf-8")
                    with self.assertRaises(finalize_docx.FinalizationPublishError):
                        self.publish(
                            staged, targets, snapshots, staging
                        )
                    self.assertFalse(targets["status"].exists())
                finally:
                    self.root = original_root

    def test_full_production_status_mutation_matrix_rejected_with_and_without_anchor_tamper(
        self,
    ) -> None:
        mutations = (
            ("status_fail", lambda value: value.update(status="fail")),
            ("unknown_top", lambda value: value.update(unknown_gate=True)),
            (
                "version_bool",
                lambda value: value.update(finalization_evidence_version=True),
            ),
            (
                "version_float",
                lambda value: value.update(finalization_evidence_version=1.0),
            ),
            (
                "completion_bool",
                lambda value: value["field_completion"].update(
                    field_gate_completed=False
                ),
            ),
            (
                "completion_float",
                lambda value: value["field_completion"].update(
                    final_ready_eligible=1.0
                ),
            ),
            (
                "validation_conflict",
                lambda value: value["field_completion"].update(
                    evidence_validation={"status": "incomplete", "errors": ["x"]}
                ),
            ),
            (
                "validation_missing",
                lambda value: value["field_completion"].pop(
                    "evidence_validation"
                ),
            ),
            (
                "unknown_nested",
                lambda value: value["field_completion"].update(extra=True),
            ),
        )
        for name, mutate in mutations:
            for anchor_mode in ("original", "mutated"):
                with self.subTest(name=name, anchor=anchor_mode):
                    case = self.root / f"full-status-{name}-{anchor_mode}"
                    case.mkdir()
                    original_root = self.root
                    self.root = case
                    try:
                        staging, staged, targets, snapshots, old = self.publish_fixture(
                            existing=True
                        )
                        value = json.loads(
                            staged["status"].read_text(encoding="utf-8")
                        )
                        mutate(value)
                        payload = finalize_docx.standard_json_bytes(value)
                        staged["status"].write_bytes(payload)
                        if anchor_mode == "mutated":
                            self.trusted_status_identity = (
                                finalize_docx.trusted_status_byte_identity(payload)
                            )
                        with self.assertRaises(
                            finalize_docx.FinalizationPublishError
                        ):
                            self.publish(staged, targets, snapshots, staging)
                        for artifact, target in targets.items():
                            self.assertEqual(old[artifact], target.read_bytes())
                        self.assertFalse(
                            (
                                self.root
                                / finalize_docx.PUBLICATION_RECOVERY_DIRECTORY
                            ).exists()
                        )
                    finally:
                        self.root = original_root

    def test_status_entity_mutation_after_status_publish_rolls_back(self) -> None:
        staging, staged, targets, snapshots, old = self.publish_fixture(existing=True)
        mutated = False

        def hook(event: str, name: str | None) -> None:
            nonlocal mutated
            if event == "after_publish" and name == "status" and not mutated:
                mutated = True
                value = json.loads(targets["status"].read_text(encoding="utf-8"))
                value["status"] = "fail"
                targets["status"].write_bytes(
                    finalize_docx.standard_json_bytes(value)
                )

        with self.assertRaises(finalize_docx.FinalizationPublishError) as caught:
            self.publish(
                staged, targets, snapshots, staging, event_hook=hook
            )
        self.assertTrue(mutated)
        self.assertTrue(caught.exception.preserve_staging)
        for name, payload in old.items():
            self.assert_old_preserved(payload)


if __name__ == "__main__":
    unittest.main()
