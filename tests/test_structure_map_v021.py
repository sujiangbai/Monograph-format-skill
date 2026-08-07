from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "format-monograph"
SCRIPTS = SKILL / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import content_fingerprint  # noqa: E402
from structure_map import candidate_structure_map, text_sha256  # noqa: E402
from test_v11_execution import approved_v11_profile  # noqa: E402


class StructureMapV021Tests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.source = self.root / "synthetic-book.docx"
        document = Document()
        document.add_paragraph("Synthetic book title")
        document.add_paragraph("1.1 Static contents entry 1", style="Heading 1")
        document.add_paragraph("第1章 Synthetic chapter")
        document.add_paragraph("1.1 Synthetic section")
        document.add_paragraph("1.1.1 Synthetic subsection")
        document.add_paragraph("1.1.1.1 Synthetic topic")
        document.add_paragraph("Synthetic body text remains unchanged.")
        document.add_paragraph("图 1.1.1-1 Synthetic figure caption")
        document.add_paragraph("[[PAGE]]")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Value A"
        table.cell(1, 1).text = "Value B"
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_section(WD_SECTION.NEW_PAGE)
        document.save(self.source)

        structure = candidate_structure_map(self.source)
        structure["status"] = "approved"
        structure["toc_ranges"] = [
            {
                "start_paragraph": 1,
                "end_paragraph": 1,
                "paragraph_sha256": [text_sha256("1.1 Static contents entry 1")],
                "levels": 4,
                "approved": True,
            }
        ]
        for entry in structure["headings"]:
            entry["approved"] = entry["paragraph"] in {2, 3, 4, 5}
        for entry in structure["captions"]:
            entry["approved"] = True
        structure["tables"][0].update(
            {
                "approved": True,
                "repeat_header": True,
                "prevent_normal_row_split": True,
            }
        )
        for entry in structure["trailing_empty_sections"]:
            entry["approved_delete"] = True
        self.structure_path = self.root / "structure-map.json"
        self.structure_path.write_text(
            json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        profile = approved_v11_profile()
        profile["rules"] = [
            rule
            for rule in profile["rules"]
            if rule["selector"]["kind"] != "table_role"
        ]
        self.profile = self.root / "profile.json"
        self.profile.write_text(
            json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        self.output = self.root / "out"

    def tearDown(self) -> None:
        self.temp.cleanup()

    def run_script(self, name: str, *args: object) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(SCRIPTS / name), *(str(value) for value in args)],
            cwd=SKILL,
            capture_output=True,
            text=True,
            check=False,
        )

    def test_candidate_map_contains_no_authored_text(self) -> None:
        value = self.structure_path.read_text(encoding="utf-8")
        self.assertNotIn("Synthetic chapter", value)
        self.assertNotIn("Synthetic figure caption", value)
        self.assertIn("text_sha256", value)

    def test_apply_approved_structure_map(self) -> None:
        original = self.source.read_bytes()
        result = self.run_script(
            "apply_profile.py",
            self.source,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure_path,
            "--output-dir",
            self.output,
            "--allow-missing-fonts",
        )
        self.assertEqual(0, result.returncode, result.stdout + result.stderr)
        self.assertEqual(original, self.source.read_bytes())

        formatted = self.output / "synthetic-book-formatted.docx"
        document = Document(formatted)
        self.assertEqual(1, len(document.sections))
        self.assertEqual("Synthetic chapter", document.paragraphs[2].text)
        self.assertEqual("Heading 1", document.paragraphs[2].style.name)
        self.assertEqual("Heading 2", document.paragraphs[3].style.name)
        self.assertEqual("Heading 3", document.paragraphs[4].style.name)
        self.assertEqual("Heading 4", document.paragraphs[5].style.name)
        self.assertEqual("Caption", document.paragraphs[7].style.name)

        first_row = document.tables[0].rows[0]._tr.get_or_add_trPr()
        self.assertIsNotNone(first_row.find(qn("w:tblHeader")))
        for row in document.tables[0].rows:
            self.assertIsNotNone(row._tr.get_or_add_trPr().find(qn("w:cantSplit")))

        with zipfile.ZipFile(formatted) as package:
            xml = package.read("word/document.xml")
        self.assertIn(b"TOC", xml)
        self.assertIn(b"STYLEREF 3", xml)
        self.assertIn(b"SEQ Figure", xml)
        self.assertIn(b"PAGE", xml)
        self.assertNotIn("Static contents entry".encode(), xml)

        audit = self.run_script(
            "audit_docx.py",
            self.source,
            formatted,
            "--profile",
            self.profile,
            "--structure-map",
            self.structure_path,
        )
        self.assertEqual(0, audit.returncode, audit.stdout + audit.stderr)

    def test_changed_source_blocks_map(self) -> None:
        document = Document(self.source)
        document.paragraphs[2].add_run(" changed")
        changed = self.root / "changed.docx"
        document.save(changed)
        result = self.run_script(
            "validate_structure_map.py",
            self.structure_path,
            "--source",
            changed,
        )
        self.assertEqual(1, result.returncode)
        self.assertIn("fingerprint", result.stderr.lower())

    def test_environment_reports_independent_capabilities(self) -> None:
        result = self.run_script("check_environment.py", "--json")
        self.assertEqual(0, result.returncode, result.stderr)
        capabilities = json.loads(result.stdout)["capabilities"]
        self.assertEqual(
            {"inspection", "profile_validation", "docx_editing", "rendering"},
            set(capabilities),
        )


if __name__ == "__main__":
    unittest.main()
