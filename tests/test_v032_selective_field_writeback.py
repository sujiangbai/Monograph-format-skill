from __future__ import annotations

import tempfile
import unittest
import zipfile
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_HEADER_FOOTER, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "format-monograph" / "scripts"
import sys

sys.path.insert(0, str(SCRIPTS))

from _common import NS, FormatMonographError  # noqa: E402
from field_writeback import (  # noqa: E402
    parse_fields,
    selective_field_result_writeback,
)
from finalize_docx import (  # noqa: E402
    apply_measured_layout_adjustments,
    apply_page_display_offsets,
    external_measure,
    external_verify,
    remove_measured_block_spacers,
)


def add_complex_field(paragraph, instruction: str, value: str, *, dirty: bool = True) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)
    instruction_run = paragraph.add_run()
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = f" {instruction} "
    instruction_run._r.append(instruction_node)
    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    paragraph.add_run(value)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_simple_field(paragraph, instruction: str, value: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = value
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def add_nested_formula(
    paragraph,
    nested_instruction: str,
    value: str,
) -> None:
    begin = paragraph.add_run()
    outer_begin = OxmlElement("w:fldChar")
    outer_begin.set(qn("w:fldCharType"), "begin")
    begin._r.append(outer_begin)
    instruction = paragraph.add_run()
    instruction_node = OxmlElement("w:instrText")
    instruction_node.text = " = "
    instruction._r.append(instruction_node)
    add_complex_field(paragraph, nested_instruction, "2", dirty=False)
    suffix = paragraph.add_run()
    suffix_node = OxmlElement("w:instrText")
    suffix_node.text = " - 1 "
    suffix._r.append(suffix_node)
    separator = paragraph.add_run()
    separate_node = OxmlElement("w:fldChar")
    separate_node.set(qn("w:fldCharType"), "separate")
    separator._r.append(separate_node)
    paragraph.add_run(value)
    end = paragraph.add_run()
    end_node = OxmlElement("w:fldChar")
    end_node.set(qn("w:fldCharType"), "end")
    end._r.append(end_node)


def rewrite_package(source: Path, output: Path, transform) -> None:
    with zipfile.ZipFile(source) as package, zipfile.ZipFile(
        output, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for info in package.infolist():
            target.writestr(info, transform(info.filename, package.read(info.filename)))


def field_values(path: Path, part: str = "word/document.xml") -> list[str]:
    with zipfile.ZipFile(path) as package:
        root = etree.fromstring(package.read(part))
    values = []
    elements = list(root.iter())
    for record in parse_fields(root):
        if record.form != "complex":
            continue
        start = elements.index(record.separate)
        end = elements.index(record.end)
        values.append(
            "".join(
                element.text or ""
                for element in elements[start + 1 : end]
                if element.tag == qn("w:t")
            )
        )
    return values


class V032SelectiveFieldWritebackTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_backend_serialization_noise_is_discarded(self) -> None:
        baseline = self.root / "baseline.docx"
        document = Document()
        document.add_paragraph("Authored content")
        add_complex_field(document.add_paragraph(), "PAGE", "1", dirty=False)
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        document.settings._element.append(update)
        document.save(baseline)
        refreshed = self.root / "refreshed.docx"

        def transform(name: str, data: bytes) -> bytes:
            if name != "word/document.xml":
                return data
            root = etree.fromstring(data)
            root.xpath(".//w:p[1]", namespaces=NS)[0].set(qn("w:rsidR"), "00ABCDEF")
            authored_run = root.xpath(".//w:p[1]/w:r[1]", namespaces=NS)[0]
            properties = OxmlElement("w:rPr")
            properties.append(OxmlElement("w:b"))
            authored_run.insert(0, properties)
            record = next(item for item in parse_fields(root) if item.field_type == "PAGE")
            elements = list(root.iter())
            start = elements.index(record.separate)
            end = elements.index(record.end)
            next(
                item
                for item in elements[start + 1 : end]
                if item.tag == qn("w:t")
            ).text = "2"
            return etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        rewrite_package(baseline, refreshed, transform)
        output = self.root / "output.docx"
        report = selective_field_result_writeback(baseline, refreshed, output)
        self.assertEqual("selective_verified", report["status"])
        self.assertEqual(["2"], field_values(output))
        with zipfile.ZipFile(output) as package:
            root = etree.fromstring(package.read("word/document.xml"))
            self.assertFalse(root.xpath(".//w:p[1]/w:r[1]/w:rPr/w:b", namespaces=NS))
            self.assertIsNone(root.xpath(".//w:p[1]", namespaces=NS)[0].get(qn("w:rsidR")))
            settings = etree.fromstring(package.read("word/settings.xml"))
            self.assertFalse(settings.xpath("./w:updateFields", namespaces=NS))

    def test_changed_instruction_or_authored_text_is_rejected(self) -> None:
        baseline = self.root / "baseline.docx"
        document = Document()
        document.add_paragraph("Protected authored text")
        add_complex_field(document.add_paragraph(), "PAGE", "1")
        document.save(baseline)

        for mode in ("instruction", "authored"):
            refreshed = self.root / f"{mode}.docx"

            def transform(name: str, data: bytes, selected: str = mode) -> bytes:
                if name != "word/document.xml":
                    return data
                root = etree.fromstring(data)
                if selected == "instruction":
                    root.xpath(".//w:instrText", namespaces=NS)[0].text = " NUMPAGES "
                else:
                    root.xpath(".//w:p[1]//w:t", namespaces=NS)[0].text = "Changed"
                return etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )

            rewrite_package(baseline, refreshed, transform)
            with self.assertRaises(FormatMonographError):
                selective_field_result_writeback(
                    baseline, refreshed, self.root / f"{mode}-output.docx"
                )

    def test_duplicate_scalar_fields_match_by_unique_authored_context(self) -> None:
        baseline = self.root / "duplicate.docx"
        document = Document()
        add_complex_field(document.add_paragraph("First "), "PAGE", "1", dirty=False)
        add_complex_field(document.add_paragraph("Second "), "PAGE", "1", dirty=False)
        document.save(baseline)
        refreshed = self.root / "duplicate-refreshed.docx"

        def transform(name: str, data: bytes) -> bytes:
            if name != "word/document.xml":
                return data
            root = etree.fromstring(data)
            elements = list(root.iter())
            for value, record in zip(("2", "3"), parse_fields(root)):
                start = elements.index(record.separate)
                end = elements.index(record.end)
                next(
                    item
                    for item in elements[start + 1 : end]
                    if item.tag == qn("w:t")
                ).text = value
            return etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        rewrite_package(baseline, refreshed, transform)
        output = self.root / "duplicate-output.docx"
        report = selective_field_result_writeback(baseline, refreshed, output)
        self.assertEqual(2, report["matched_fields"])
        self.assertEqual(["2", "3"], field_values(output))

    def test_ambiguous_duplicate_scalar_fields_are_rejected(self) -> None:
        baseline = self.root / "ambiguous.docx"
        document = Document()
        add_complex_field(document.add_paragraph(), "PAGE", "1")
        add_complex_field(document.add_paragraph(), "PAGE", "1")
        document.save(baseline)
        refreshed = self.root / "ambiguous-refreshed.docx"
        rewrite_package(baseline, refreshed, lambda _name, data: data)
        with self.assertRaisesRegex(FormatMonographError, "Duplicate fields"):
            selective_field_result_writeback(
                baseline,
                refreshed,
                self.root / "ambiguous-output.docx",
            )

    def test_simple_field_can_match_word_complex_field_serialization(self) -> None:
        baseline = self.root / "simple.docx"
        document = Document()
        add_simple_field(document.add_paragraph(), "PAGE", "1")
        document.save(baseline)
        refreshed = self.root / "complex.docx"

        def transform(name: str, data: bytes) -> bytes:
            if name != "word/document.xml":
                return data
            root = etree.fromstring(data)
            paragraph = root.xpath(".//w:p[w:fldSimple]", namespaces=NS)[0]
            field = paragraph.find(qn("w:fldSimple"))
            paragraph.remove(field)
            temporary = Document()
            temporary_paragraph = temporary.add_paragraph()
            add_complex_field(temporary_paragraph, "PAGE", "8", dirty=False)
            for child in list(temporary_paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph.append(child)
            return etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        rewrite_package(baseline, refreshed, transform)
        output = self.root / "simple-output.docx"
        report = selective_field_result_writeback(baseline, refreshed, output)
        self.assertEqual(1, report["matched_fields"])
        with zipfile.ZipFile(output) as package:
            root = etree.fromstring(package.read("word/document.xml"))
            self.assertEqual(
                "8",
                "".join(root.xpath(".//w:fldSimple//w:t/text()", namespaces=NS)),
            )

    def test_header_footer_results_are_patched_without_importing_other_xml(self) -> None:
        baseline = self.root / "footer.docx"
        document = Document()
        document.add_paragraph("Body")
        add_complex_field(
            document.sections[0].footer.paragraphs[0],
            "PAGE",
            "1",
            dirty=False,
        )
        document.save(baseline)
        with zipfile.ZipFile(baseline) as package:
            footer_part = next(
                name for name in package.namelist() if name.startswith("word/footer")
            )
        refreshed = self.root / "footer-refreshed.docx"

        def transform(name: str, data: bytes) -> bytes:
            if name != footer_part:
                return data
            root = etree.fromstring(data)
            record = parse_fields(root)[0]
            elements = list(root.iter())
            start = elements.index(record.separate)
            end = elements.index(record.end)
            next(
                item
                for item in elements[start + 1 : end]
                if item.tag == qn("w:t")
            ).text = "9"
            root.set("backend-noise", "discard-me")
            return etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        rewrite_package(baseline, refreshed, transform)
        output = self.root / "footer-output.docx"
        report = selective_field_result_writeback(baseline, refreshed, output)
        self.assertIn(footer_part, report["patched_parts"])
        self.assertEqual(["9"], field_values(output, footer_part))
        with zipfile.ZipFile(output) as package:
            root = etree.fromstring(package.read(footer_part))
            self.assertIsNone(root.get("backend-noise"))

    def test_unapproved_dirty_seq_is_reported_and_not_updated(self) -> None:
        baseline = self.root / "seq.docx"
        document = Document()
        add_complex_field(document.add_paragraph(), "SEQ Figure", "1")
        document.save(baseline)
        refreshed = self.root / "seq-refreshed.docx"
        rewrite_package(baseline, refreshed, lambda _name, data: data)
        output = self.root / "seq-output.docx"
        report = selective_field_result_writeback(baseline, refreshed, output)
        self.assertEqual(["SEQ"], report["unapproved_field_types"])
        self.assertEqual(1, report["unapproved_dirty_fields"])
        self.assertEqual(["1"], field_values(output))

    def test_global_update_on_open_is_removed_with_clean_unapproved_fields(self) -> None:
        baseline = self.root / "mixed-fields.docx"
        document = Document()
        add_complex_field(document.add_paragraph(), "PAGE", "1", dirty=False)
        add_complex_field(document.add_paragraph("Figure "), "SEQ Figure", "1", dirty=False)
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        document.settings._element.append(update)
        document.save(baseline)
        refreshed = self.root / "mixed-fields-refreshed.docx"
        rewrite_package(baseline, refreshed, lambda _name, data: data)
        output = self.root / "mixed-fields-output.docx"
        report = selective_field_result_writeback(baseline, refreshed, output)
        self.assertEqual(0, report["unapproved_dirty_fields"])
        self.assertEqual(["SEQ"], report["unapproved_field_types"])
        with zipfile.ZipFile(output) as package:
            settings = etree.fromstring(package.read("word/settings.xml"))
        self.assertFalse(settings.xpath("./w:updateFields", namespaces=NS))

    def test_dirty_approved_result_is_rejected(self) -> None:
        baseline = self.root / "dirty-page.docx"
        document = Document()
        add_complex_field(document.add_paragraph(), "PAGE", "1")
        document.save(baseline)
        refreshed = self.root / "dirty-page-refreshed.docx"
        rewrite_package(baseline, refreshed, lambda _name, data: data)
        with self.assertRaisesRegex(FormatMonographError, "left field PAGE dirty"):
            selective_field_result_writeback(
                baseline,
                refreshed,
                self.root / "dirty-page-output.docx",
            )

    def test_arbitrary_formula_is_rejected_even_when_formula_type_is_allowed(self) -> None:
        baseline = self.root / "formula.docx"
        document = Document()
        add_complex_field(document.add_paragraph(), "= 2 + 2", "4")
        document.save(baseline)
        refreshed = self.root / "formula-refreshed.docx"
        rewrite_package(baseline, refreshed, lambda _name, data: data)

        with self.assertRaisesRegex(
            FormatMonographError,
            "Only the exact core-generated PAGE-minus-one display formula",
        ):
            selective_field_result_writeback(
                baseline,
                refreshed,
                self.root / "formula-output.docx",
                allowed_field_types={"="},
            )

    def test_non_page_nested_formula_and_non_numeric_result_are_rejected(self) -> None:
        for nested, value in (("NUMPAGES", "1"), ("PAGE", "FORGED")):
            baseline = self.root / f"nested-{nested}-{value}.docx"
            document = Document()
            add_nested_formula(document.add_paragraph(), nested, value)
            document.save(baseline)
            refreshed = self.root / f"nested-{nested}-{value}-refreshed.docx"
            rewrite_package(baseline, refreshed, lambda _name, data: data)
            with self.assertRaises(FormatMonographError):
                selective_field_result_writeback(
                    baseline,
                    refreshed,
                    self.root / f"nested-{nested}-{value}-output.docx",
                    allowed_field_types={"=", "PAGE", "NUMPAGES"},
                )

    def test_external_verify_requires_no_save_contract_and_matching_page_count(self) -> None:
        source = self.root / "verify.docx"
        Document().save(source)
        helper = self.root / "verify_backend.py"
        helper.write_text(
            "import json, pathlib, sys\n"
            "request=json.load(sys.stdin)\n"
            "assert request['operation']=='verify_only'\n"
            "pathlib.Path(request['pdf_output_path']).write_bytes(b'%PDF-1.4\\n%%EOF')\n"
            "print(json.dumps({'status':'success','operation':'verify_only',"
            "'backend':'test_word','software':'Microsoft Word','repaginated':True,"
            "'saved':False,'read_only_verified':True,'pdf_exported':True,"
            "'page_count':7}))\n",
            encoding="utf-8",
        )
        pdf = self.root / "verify.pdf"
        response = external_verify(
            source,
            json.dumps([sys.executable, str(helper)]),
            self.root / "profile.json",
            self.root / "map.json",
            pdf,
            "Microsoft Word 2021",
            expected_page_count=7,
        )
        self.assertTrue(response["read_only_verified"])
        with self.assertRaises(FormatMonographError):
            external_verify(
                source,
                json.dumps([sys.executable, str(helper)]),
                self.root / "profile.json",
                self.root / "map.json",
                self.root / "mismatch.pdf",
                "Microsoft Word 2021",
                expected_page_count=8,
            )
        helper.write_text(
            "import json, pathlib, sys\n"
            "request=json.load(sys.stdin)\n"
            "pathlib.Path(request['pdf_output_path']).write_bytes(b'%PDF-1.4\\n%%EOF')\n"
            "print(json.dumps({'status':'success','operation':'verify_only',"
            "'backend':'test_word','software':'Microsoft Word','repaginated':True,"
            "'saved':False,'read_only_verified':True,'pdf_exported':True}))\n",
            encoding="utf-8",
        )
        with self.assertRaisesRegex(FormatMonographError, "valid page count"):
            external_verify(
                source,
                json.dumps([sys.executable, str(helper)]),
                self.root / "profile.json",
                self.root / "map.json",
                self.root / "missing-count.pdf",
                "Microsoft Word 2021",
                expected_page_count=7,
            )

    def test_layout_measurement_and_core_spacer_removal_are_separate(self) -> None:
        source = self.root / "spacers.docx"
        document = Document()
        style = document.styles.add_style(
            "Monograph Figure Table Spacer", WD_STYLE_TYPE.PARAGRAPH
        )
        document.add_paragraph("Body before")
        document.add_paragraph(style=style)
        document.add_paragraph("Body after")
        document.save(source)
        helper = self.root / "measure_backend.py"
        helper.write_text(
            "import json, sys\n"
            "request=json.load(sys.stdin)\n"
            "assert request['operation']=='measure_layout'\n"
            "print(json.dumps({'status':'success','operation':'measure_layout',"
            "'backend':'test_word','software':'Microsoft Word','repaginated':True,"
            "'saved':False,'read_only_verified':True,'structural_changes_applied':0,"
            "'page_count':3,'sections':[{'section_index':0}],"
            "'page_boundary_spacer_ordinals':[0]}))\n",
            encoding="utf-8",
        )
        measured = external_measure(
            source,
            json.dumps([sys.executable, str(helper)]),
            self.root / "profile.json",
            self.root / "map.json",
            "Microsoft Word 2021",
        )
        self.assertEqual([0], measured["page_boundary_spacer_ordinals"])
        output = self.root / "spacers-normalized.docx"
        self.assertEqual(
            2,
            apply_measured_layout_adjustments(
                source,
                output,
                [0],
                {0: "evenPage"},
            ),
        )
        self.assertEqual(
            ["Body before", "Body after"],
            [paragraph.text for paragraph in Document(output).paragraphs],
        )
        section_type = Document(output).sections[0]._sectPr.find(qn("w:type"))
        self.assertEqual("evenPage", section_type.get(qn("w:val")))

        spacer_only = self.root / "spacer-only.docx"
        self.assertEqual(1, remove_measured_block_spacers(source, spacer_only, [0]))

        offset_output = self.root / "page-offset.docx"
        self.assertEqual(2, apply_page_display_offsets(source, offset_output, {0: 1}))
        with zipfile.ZipFile(offset_output) as package:
            field_types = []
            for name in package.namelist():
                if not name.startswith("word/footer"):
                    continue
                field_types.extend(
                    record.field_type
                    for record in parse_fields(etree.fromstring(package.read(name)))
                )
        self.assertEqual(2, field_types.count("="))
        self.assertEqual(2, field_types.count("PAGE"))

    def test_page_offset_footer_is_isolated_from_adjacent_sections(self) -> None:
        source = self.root / "shared-footers.docx"
        document = Document()
        document.settings.odd_and_even_pages_header_footer = True
        document.add_paragraph("Title")
        for footer in (
            document.sections[0].footer,
            document.sections[0].even_page_footer,
        ):
            add_simple_field(footer.paragraphs[0], "PAGE", "1")
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Contents")
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Body")
        first = document.sections[0]
        for section in document.sections[1:]:
            for footer_type in (
                WD_HEADER_FOOTER.PRIMARY,
                WD_HEADER_FOOTER.EVEN_PAGE,
            ):
                if section._sectPr.get_footerReference(footer_type) is not None:
                    section._sectPr.remove_footerReference(footer_type)
                reference = first._sectPr.get_footerReference(footer_type)
                section._sectPr.add_footerReference(footer_type, reference.rId)
        document.save(source)

        output = self.root / "isolated-footers.docx"
        self.assertEqual(2, apply_page_display_offsets(source, output, {1: 1}))
        result = Document(output)
        for footer_name in ("footer", "even_page_footer"):
            field_types = []
            for section in result.sections:
                footer = getattr(section, footer_name)
                root = etree.fromstring(footer._element.xml.encode("utf-8"))
                field_types.append([record.field_type for record in parse_fields(root)])
            self.assertEqual([["PAGE"], ["=", "PAGE"], ["PAGE"]], field_types)


if __name__ == "__main__":
    unittest.main()
