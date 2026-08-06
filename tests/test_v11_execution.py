from __future__ import annotations

import base64
import json
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "format-monograph"
SCRIPTS = SKILL / "scripts"
sys.path.insert(0, str(SCRIPTS))

from _common import (  # noqa: E402
    content_fingerprint,
    equation_inventory,
    protected_object_manifest,
)
from validate_profile import validate  # noqa: E402


def approved_v11_profile() -> dict:
    source = {
        "id": "SRC-901",
        "type": "user_requirement",
        "label": "Synthetic caller requirements",
        "summary": "Synthetic V1.1 requirements for automated execution tests.",
        "public": True,
    }

    def rule(identifier: str, category: str, kind: str, value: str, properties: dict) -> dict:
        return {
            "id": identifier,
            "category": category,
            "selector": {"kind": kind, "value": value},
            "properties": properties,
            "source_ids": ["SRC-901"],
            "evidence_summary": "Synthetic approved rule.",
            "confidence": "high",
            "status": "approved",
            "application": "automatic",
        }

    return {
        "schema_version": "1.1",
        "profile_id": "synthetic-v11-execution",
        "name": "Synthetic V1.1 execution profile",
        "locale": "zh-CN",
        "scope": {"document_type": "monograph", "input_format": "DOCX"},
        "target_applications": ["Microsoft 365"],
        "runtime_policy": {
            "caller_requirements_highest": True,
            "editable_equations_required": True,
            "formula_image_policy": "block",
            "legacy_equation_policy": "qa",
            "field_rebuild_policy": "explicit_or_unambiguous",
        },
        "source_precedence": ["user_requirement", "written_requirement", "sample_book"],
        "sources": [source],
        "rules": [
            rule(
                "FMT-PAGE-901",
                "page",
                "document",
                "all",
                {
                    "page_size_policy": "preserve",
                    "mirror_margins": True,
                    "margin_inner_ratio": 0.12,
                    "margin_outer_ratio": 0.105,
                    "margin_top_ratio": 0.096,
                    "margin_bottom_ratio": 0.085,
                    "header_distance_ratio": 0.055,
                    "footer_distance_ratio": 0.05,
                    "gutter_mm": 0,
                    "different_first_page_header_footer": True,
                    "odd_and_even_pages_header_footer": True,
                },
            ),
            rule(
                "FMT-BODY-901",
                "body",
                "paragraph_role",
                "body_text",
                {
                    "font_name_east_asia": "SimSun",
                    "font_name_ascii": "Times New Roman",
                    "font_name_complex_script": "Times New Roman",
                    "font_size_pt": 10.5,
                    "alignment": "justify",
                    "line_spacing_rule": "exact",
                    "line_spacing_pt": 18,
                    "first_line_indent_chars": 2,
                    "space_before_pt": 0,
                    "space_after_pt": 0,
                },
            ),
            rule(
                "FMT-HEAD-901",
                "heading",
                "paragraph_role",
                "chapter_title",
                {
                    "font_name_east_asia": "SimHei",
                    "font_name_ascii": "Arial",
                    "font_size_pt": 18,
                    "bold": True,
                    "alignment": "center",
                    "page_break_before": True,
                },
            ),
            rule(
                "FMT-TABLE-901",
                "table",
                "table_role",
                "all",
                {
                    "alignment": "center",
                    "repeat_header_row": True,
                    "prevent_row_split": True,
                },
            ),
            rule(
                "FMT-FIELD-901",
                "field",
                "field_role",
                "derived",
                {
                    "update_on_open": True,
                    "mark_fields_dirty": True,
                    "convert_explicit_markers": True,
                    "rebuild_heading_numbering": True,
                    "heading_levels": 4,
                    "strip_manual_heading_prefixes": True,
                },
            ),
            rule(
                "FMT-EQ-901",
                "equation",
                "equation_role",
                "all",
                {
                    "require_editable_equations": True,
                    "preserve_editable_objects": True,
                    "block_formula_images": True,
                },
            ),
        ],
        "conflicts": [],
        "open_questions": [],
        "approval": {
            "status": "approved",
            "approved_by": "test-suite",
            "approved_at": "2026-08-06T00:00:00Z",
        },
    }


def add_omml(paragraph) -> None:
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.text = "x+y"
    run.append(text)
    math.append(run)
    paragraph._p.append(math)


class V11ExecutionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.input = self.root / "v11.docx"
        document = Document()
        document.add_heading("第1章 Synthetic chapter", level=1)
        document.add_paragraph("中西文 mixed body text remains unchanged.")
        document.add_paragraph("[[PAGE]]")
        equation = document.add_paragraph()
        add_omml(equation)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "表头 A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "数据"
        table.cell(1, 1).text = "Value"
        document.save(self.input)

        self.profile_path = self.root / "profile.json"
        self.profile_path.write_text(
            json.dumps(approved_v11_profile(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        self.out = self.root / "out"
        self.original_bytes = self.input.read_bytes()

    def tearDown(self) -> None:
        self.temp.cleanup()

    def run_script(self, name: str, *args: object) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(SCRIPTS / name), *(str(value) for value in args)],
            cwd=SKILL,
            capture_output=True,
            text=True,
            check=False,
        )

    def test_profile_validates_and_v10_remains_supported(self) -> None:
        errors, parsed = validate(self.profile_path)
        self.assertEqual([], errors)
        self.assertEqual("1.1", parsed["schema_version"])

    def test_full_v11_structural_execution(self) -> None:
        before_objects = protected_object_manifest(self.input)
        result = self.run_script(
            "apply_profile.py",
            self.input,
            "--profile",
            self.profile_path,
            "--output-dir",
            self.out,
        )
        self.assertEqual(0, result.returncode, result.stderr)
        self.assertEqual(self.original_bytes, self.input.read_bytes())

        formatted = self.out / "v11-formatted.docx"
        review = self.out / "v11-review.docx"
        report = self.out / "v11-format-report.md"
        self.assertTrue(formatted.exists())
        self.assertTrue(review.exists())
        self.assertTrue(report.exists())

        self.assertEqual(
            content_fingerprint(self.input, normalize_derived=True),
            content_fingerprint(formatted, normalize_derived=True),
        )
        self.assertEqual(before_objects, protected_object_manifest(formatted))
        self.assertEqual(before_objects, protected_object_manifest(review))
        self.assertEqual(1, equation_inventory(formatted)["omml"])

        document = Document(formatted)
        normal = document.styles["Normal"]
        r_fonts = normal.element.rPr.rFonts
        self.assertEqual("Times New Roman", r_fonts.get(qn("w:ascii")))
        self.assertEqual("SimSun", r_fonts.get(qn("w:eastAsia")))
        self.assertAlmostEqual(10.5, normal.font.size.pt)
        self.assertAlmostEqual(18, normal.paragraph_format.line_spacing.pt)
        self.assertEqual(WD_LINE_SPACING.EXACTLY, normal.paragraph_format.line_spacing_rule)
        indent = normal.element.pPr.find(qn("w:ind"))
        self.assertEqual("200", indent.get(qn("w:firstLineChars")))

        heading = next(p for p in document.paragraphs if p.style.name == "Heading 1")
        self.assertEqual("Synthetic chapter", heading.text)
        self.assertIsNotNone(document.styles["Heading 1"].element.pPr.find(qn("w:numPr")))

        section = document.sections[0]
        width = section.page_width.mm
        height = section.page_height.mm
        self.assertAlmostEqual(width * 0.12, section.left_margin.mm, places=1)
        self.assertAlmostEqual(width * 0.105, section.right_margin.mm, places=1)
        self.assertAlmostEqual(height * 0.096, section.top_margin.mm, places=1)
        self.assertAlmostEqual(height * 0.085, section.bottom_margin.mm, places=1)
        self.assertIsNotNone(document.settings.element.find(qn("w:mirrorMargins")))
        self.assertTrue(document.settings.odd_and_even_pages_header_footer)

        first_row = document.tables[0].rows[0]._tr.get_or_add_trPr()
        self.assertIsNotNone(first_row.find(qn("w:tblHeader")))
        for row in document.tables[0].rows:
            self.assertIsNotNone(row._tr.get_or_add_trPr().find(qn("w:cantSplit")))

        with zipfile.ZipFile(formatted) as package:
            xml = package.read("word/document.xml")
        self.assertIn(b'PAGE', xml)
        self.assertNotIn(b'[[PAGE]]', xml)

        inventory = self.root / "inventory.json"
        inspected = self.run_script(
            "inspect_docx.py", formatted, "--output", inventory
        )
        self.assertEqual(0, inspected.returncode, inspected.stderr)
        parsed_inventory = json.loads(inventory.read_text(encoding="utf-8"))
        self.assertEqual(1, parsed_inventory["equations"]["omml"])
        self.assertEqual(1, parsed_inventory["fields"]["types"]["PAGE"])

        audit = self.root / "audit.json"
        audited = self.run_script(
            "audit_docx.py",
            self.input,
            formatted,
            "--profile",
            self.profile_path,
            "--output",
            audit,
        )
        self.assertEqual(0, audited.returncode, audited.stderr)
        self.assertTrue(json.loads(audit.read_text(encoding="utf-8"))["passed"])

    def test_formula_image_candidate_blocks_application(self) -> None:
        image_docx = self.root / "formula-image.docx"
        png = self.root / "formula.png"
        png.write_bytes(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wl2nuwAAAAASUVORK5CYII="
            )
        )
        document = Document()
        document.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
        paragraph = document.add_paragraph(style="Equation")
        paragraph.add_run().add_picture(str(png))
        document.save(image_docx)
        original = image_docx.read_bytes()

        result = self.run_script(
            "apply_profile.py",
            image_docx,
            "--profile",
            self.profile_path,
            "--output-dir",
            self.root / "blocked",
        )
        self.assertEqual(1, result.returncode)
        self.assertIn("formula image candidate", result.stderr.lower())
        self.assertEqual(original, image_docx.read_bytes())
        self.assertFalse((self.root / "blocked" / "formula-image-formatted.docx").exists())


if __name__ == "__main__":
    unittest.main()
